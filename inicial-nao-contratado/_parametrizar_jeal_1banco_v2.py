# -*- coding: utf-8 -*-
"""V2 — parametrização CIRÚRGICA do inicial-jeal-1banco.docx.
Substituições por trecho pequeno cabendo em run de formatação uniforme,
preservando Segoe UI Bold do título/banco e Cambria normal do texto."""
import sys, io, shutil, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from docx.shared import Pt
from helpers_docx import substituir_in_run
from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

CAMINHO = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jeal-1banco.docx")

# ---- Restaurar do backup primeiro (garantir estado limpo) ----
backup = CAMINHO.with_suffix(CAMINHO.suffix + ".bak_pre_parametrizacao")
shutil.copy2(backup, CAMINHO)
print(f"Restaurado de backup")

# ---- Mapas de substituição POR PARÁGRAFO (cada chave fica em UM run de
# formatação uniforme) ----

SUBS_POR_PARAGRAFO = {
    # PAR 0 — cabeçalho (tudo Segoe UI Bold) — substituições preservam
    0: {
        "___":       "{{vara_protocolo}}",
        "Cidade/AL": "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
    # PAR 6 — qualificação (FULANO em Segoe UI+B; resto em Cambria)
    6: {
        "FULANO DE TAL":                                    "{{nome_autor}}",
        "brasileiro, casado, aposentado":                   "{{nacionalidade}}, {{estado_civil}}, {{profissao}}",
        "xxx.xxx.xxx-xx":                                   "{{cpf_autor}}",
        "xxxxxxxx":                                          "{{rg_autor}}",
        "SSP/AL":                                            "{{orgao_expedidor}}",
        "rua tal":                                           "{{logradouro_autor}}",
        "n° xx, bairro tal, em cidade/AL, CEP xxxxx-xxx":   "nº {{numero_autor}}, {{bairro_autor}}, em {{cidade_autor}}/{{uf_autor}}, CEP {{cep_autor}}",
    },
    # PAR 10 — banco réu
    10: {
        "BANCO PAN S/A":                                    "{{banco_reu_nome}}",
        "pessoa jurídica de direito privado":               "{{banco_reu_descricao_pj}}",
        "59.285.411/0006-28":                               "{{banco_reu_cnpj}}",
        # endereço fragmentado em runs - juntar e substituir
        "Av. Ephigenio Salles, nº 1.327, bairro Aleixo, Manaus, CEP 69.060-020": "{{banco_reu_endereco}}",
    },
    # PAR 12 — benefício
    12: {
        "aposentadoria por idade":                          "{{tipo_beneficio}}",
        "149.139.433-9":                                    "{{nb_beneficio}}",
        "agência 3706, conta corrente nº 0000211974":       "agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}",
        "ao banco bradesco sa":                             "ao {{banco_pagador}}",
    },
    # PAR 14 — constatação (já era bold no original, manter)
    14: {
        "BANCO PAN S/A":  "{{banco_reu_nome}}",
        "3880089838":     "{{contrato_numero}}",
    },
    # PAR 16 — detalhes do contrato
    16: {
        "01/06/2024":                                                          "{{contrato_competencia_inicio}}",
        "84 parcelas":                                                          "{{contrato_qtd_parcelas}} parcelas",
        "R$ 49,00 (quarenta e nove reais)":                                    "R$ {{contrato_valor_parcela}} ({{contrato_valor_parcela_extenso}})",
        "R$ 2.171,24 (dois mil, cento e setenta e um reais e vinte e quatro centavos)": "R$ {{contrato_valor_emprestado}} ({{contrato_valor_emprestado_extenso}})",
        "contrato nº 3880089838":                                               "contrato nº {{contrato_numero}}",
        "BANCO PAN S/A, ora requerido":                                         "{{banco_reu_nome}}, ora requerido",
    },
    # PAR 44 — renda (R$ 1.093,99 está num run BOLD; o extenso também)
    44: {
        "R$ 1.093,99":                                                          "R$ {{valor_renda_liquida}}",
        "um mil e noventa e três reais e noventa e nove centavos":              "{{valor_renda_liquida_extenso}}",
    },
    # PAR 214 — pedido de prioridade (envolver com placeholder condicional)
    214: {
        "A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de Processo Civil": "{{pedido_prioridade}}",
    },
    # PAR 228 — pedido (declaração)
    228: {
        "R$ 2.171,24 (dois mil, cento e setenta e um reais e vinte e quatro centavos)": "R$ {{contrato_valor_emprestado}} ({{contrato_valor_emprestado_extenso}})",
        "contrato nº 3880089838":     "contrato nº {{contrato_numero}}",
        "R$ 49,00 (quarenta e nove reais)": "R$ {{contrato_valor_parcela}} ({{contrato_valor_parcela_extenso}})",
        "31/05/2024":                 "{{contrato_data_inclusao}}",
        "01/06/2024":                 "{{contrato_competencia_inicio}}",
        "149.139.433-9":              "{{nb_beneficio}}",
    },
    # PAR 248 — fecho (manter data fixa)
    248: {
        "Cidade/AL": "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
}

doc = Document(CAMINHO)
total = 0
for par_idx, mapa in SUBS_POR_PARAGRAFO.items():
    p = doc.paragraphs[par_idx]
    if substituir_in_run(p._p, mapa):
        total += 1
        print(f"  par {par_idx}: substituições aplicadas")

# ---- Forçar Segoe UI nos placeholders do PAR 10 (banco réu) ----
# Critério: se o run contém algum dos 4 placeholders, aplicar fonte Segoe UI
PLACEHOLDERS_BANCO_REU = ["{{banco_reu_nome}}", "{{banco_reu_descricao_pj}}",
                           "{{banco_reu_cnpj}}", "{{banco_reu_endereco}}"]

def aplicar_segoe_ui_se_contem(p_xml, placeholders):
    """Aplica fonte Segoe UI aos runs que contêm qualquer placeholder da lista."""
    n_aplicados = 0
    for r in p_xml.findall('.//' + W + 'r'):
        txt = ''.join((t.text or '') for t in r.findall(W + 't'))
        if any(ph in txt for ph in placeholders):
            rpr = r.find(W + 'rPr')
            if rpr is None:
                rpr = etree.SubElement(r, W + 'rPr')
                # rPr deve ser o primeiro filho
                r.insert(0, rpr)
            # rFonts
            rfonts = rpr.find(W + 'rFonts')
            if rfonts is None:
                rfonts = etree.SubElement(rpr, W + 'rFonts')
            for atr in ['ascii', 'hAnsi', 'cs']:
                rfonts.set(W + atr, 'Segoe UI')
            n_aplicados += 1
    return n_aplicados

p10 = doc.paragraphs[10]
n = aplicar_segoe_ui_se_contem(p10._p, PLACEHOLDERS_BANCO_REU)
print(f"  par 10: Segoe UI aplicado em {n} run(s) com placeholders banco réu")

doc.save(CAMINHO)
print(f"\nSalvo. Total de parágrafos modificados: {total}")

# ---- Verificação ----
import re
d2 = Document(CAMINHO)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r"\{\{[^}]+\}\}", p.text):
        phs.add(m.group(0))
print(f"\nPlaceholders únicos: {len(phs)}")
for ph in sorted(phs):
    print(f"  {ph}")

print("\nVerificação dados pilotos remanescentes:")
for i, p in enumerate(d2.paragraphs):
    t = p.text
    for ruim in ["FULANO DE TAL", "xxx.xxx.xxx-xx", "rua tal", "bairro tal",
                  "BANCO PAN S/A", "149.139.433-9", "3880089838",
                  "Cidade/AL", "Ephigenio", "59.285.411", "1.093,99"]:
        if ruim in t:
            print(f"  ⚠️ par {i}: '{ruim}' ainda presente")
            break
print("(vazio = OK)")
