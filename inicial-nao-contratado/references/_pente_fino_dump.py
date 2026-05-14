# -*- coding: utf-8 -*-
"""Dump completo de cada template para análise."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
from docx import Document

VAULT = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates")
OUT_DIR = Path(__file__).parent / "_pente_fino"
OUT_DIR.mkdir(exist_ok=True)

for nome in ['inicial-jeal-1banco.docx', 'inicial-jeal-2bancos.docx',
              'inicial-jemg-1banco.docx', 'inicial-jfal-1banco.docx',
              'inicial-jfal-2bancos.docx']:
    d = Document(VAULT / nome)
    out = OUT_DIR / nome.replace('.docx', '.txt')
    linhas = []
    for i, p in enumerate(d.paragraphs):
        t = p.text
        marca = ''
        if not t.strip():
            marca = ' [VAZIO]'
        linhas.append(f'{i:4d}{marca}: {t}')
    out.write_text('\n'.join(linhas), encoding='utf-8')
    print(f'{nome}: {len(d.paragraphs)} pars salvos em {out.name}')
