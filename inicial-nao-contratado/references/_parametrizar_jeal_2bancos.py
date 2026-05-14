# -*- coding: utf-8 -*-
"""Parametriza inicial-jeal-2bancos.docx — estrutura idêntica ao jeal-1banco.
Nota: o template originalmente tinha apenas 1 banco no par 10. Mantive
estrutura. Se precisar de 2 bancos depois, adicionar manualmente.
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

CAMINHO = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jeal-2bancos.docx")

backup = CAMINHO.with_suffix(CAMINHO.suffix + ".bak_pre_parametrizacao")
if not backup.exists():
    shutil.copy2(CAMINHO, backup)
    print(f"Backup: {backup.name}")

SUBS = {
    0: {
        "___":          "{{vara_protocolo}}",
        "Cidade/AL":    "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
    6: {
        "FULANO DE TAL":                                    "{{nome_autor}}",
        "brasileiro, casado, aposentado":                   "{{nacionalidade}}, {{estado_civil}}, {{profissao}}",
        "xxx.xxx.xxx-xx":                                   "{{cpf_autor}}",
        "xxxxxxxx":                                          "{{rg_autor}}",
        "SSP/AL":                                            "{{orgao_expedidor}}",
        "rua tal":                                           "{{logradouro_autor}}",
        "n° xx, bairro tal, em cidade/AL, CEP xxxxx-xxx":   "nº {{numero_autor}}, {{bairro_autor}}, em {{cidade_autor}}/{{uf_autor}}, CEP {{cep_autor}}",
    },
    10: {
        "BANCO PAN S/A":                                    "{{banco_reu_nome}}",
        "pessoa jurídica de direito privado":               "{{banco_reu_descricao_pj}}",
        "59.285.411/0006-28":                               "{{banco_reu_cnpj}}",
        "Av. Ephigenio Salles, nº 1.327, bairro Aleixo, Manaus, CEP 69.060-020": "{{banco_reu_endereco}}",
    },
    12: {
        "aposentadoria por idade":                          "{{tipo_beneficio}}",
        "149.139.433-9":                                    "{{nb_beneficio}}",
        "agência 3706, conta corrente nº 0000211974":       "agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}",
        "ao banco bradesco sa":                             "ao {{banco_pagador}}",
    },
    14: {
        "à um empréstimo que não contratou junto ao BANCO PAN S/A, CONTRATO Nº 3880089838:":
            "a empréstimo(s) que não contratou junto ao {{banco_reu_nome}}, conforme detalhamento abaixo:",
    },
    222: {
        "A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de Processo Civil":
            "{{pedido_prioridade}}",
    },
    236: {
        "do empréstimo consignado no valor de R$ 2.171,24 (dois mil, cento e setenta e um reais e vinte e quatro centavos), contrato nº 3880089838 - com descontos de R$ 49,00 (quarenta e nove reais) mensais, com inclusão em 31/05/2024, início de desconto em 01/06/2024, no benefício previdenciário 149.139.433-9":
            "{{BLOCO_PEDIDO_DECLARACAO}}",
    },
    254: {
        "R$ 23.877,98 (vinte e três mil, oitocentos e setenta e sete reais e noventa e oito centavos)":
            "R$ {{valor_causa}} ({{valor_causa_extenso}})",
    },
    256: {
        "Cidade/AL":  "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
}

doc = Document(CAMINHO)
for par_idx, mapa in SUBS.items():
    if par_idx >= len(doc.paragraphs):
        print(f"  ⚠️ par {par_idx} fora do range")
        continue
    p = doc.paragraphs[par_idx]
    if substituir_in_run(p._p, mapa):
        print(f"  par {par_idx}: substituído")

# PAR 16 — vira BLOCO_CONTRATOS_FRAUDULENTOS
SUB16_FULL = (
    "No que diz respeito ao referido empréstimo, cumpre informar que: "
    "a primeira parcela descontada do benefício da parte autora foi na competência "
    "01/06/2024, de um total de 84 parcelas, no valor de R$ 49,00 (quarenta e nove reais), "
    "relativas a um empréstimo consignado no valor de R$ 2.171,24 "
    "(dois mil, cento e setenta e um reais e vinte e quatro centavos), contrato nº 3880089838, "
    "cuja operação foi realizada pelo BANCO PAN S/A, ora requerido."
)
p16 = doc.paragraphs[16]
if substituir_in_run(p16._p, {SUB16_FULL: "{{BLOCO_CONTRATOS_FRAUDULENTOS}}"}):
    print("  par 16: virou {{BLOCO_CONTRATOS_FRAUDULENTOS}}")

# Remover smallCaps do run com {{banco_reu_nome}} no par 10
p10 = doc.paragraphs[10]
for r in p10._p.findall('.//' + W + 'r'):
    txt = ''.join((t.text or '') for t in r.findall(W + 't'))
    if "{{banco_reu_nome}}" in txt:
        rpr = r.find(W + 'rPr')
        if rpr is not None:
            sc = rpr.find(W + 'smallCaps')
            if sc is not None:
                rpr.remove(sc)

doc.save(CAMINHO)
print()

import re
d2 = Document(CAMINHO)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r"\{\{[^}]+\}\}", p.text):
        phs.add(m.group(0))
print(f"Placeholders únicos: {len(phs)}")

print("\nDados pilotos remanescentes (excluindo ementas):")
sobras = []
for i, p in enumerate(d2.paragraphs):
    t = p.text
    for ruim in ["FULANO", "xxx.xxx.xxx-xx", "rua tal", "BANCO PAN S/A",
                  "59.285.411", "149.139.433-9", "3880089838",
                  "Cidade/AL", "Ephigenio", "1.093,99", "23.877,98",
                  "2.171,24", "01/06/2024", "31/05/2024"]:
        if ruim in t and not any(em in t for em in ["DIREITO CIVIL", "Apelação Cível", "RECURSO INOMINADO", "Tribunal de Justiça"]):
            sobras.append((i, ruim, t[:120]))
            break
if sobras:
    for i, r, t in sobras:
        print(f"  ⚠️ par {i}: '{r}' em: {t}")
else:
    print("  (vazio = OK)")
