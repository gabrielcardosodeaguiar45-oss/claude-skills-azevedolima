# -*- coding: utf-8 -*-
"""Parametriza inicial-jemg-1banco.docx (MG Estadual / 1 banco / N contratos):
substitui dados pilotos do EXEMPLO DA SILVA + Banco C6 + Ipatinga/MG
por placeholders canônicos. Adota Opção 4 (BLOCO_CONTRATOS / BLOCO_PEDIDO).

Backup: <arquivo>.bak_pre_parametrizacao
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

CAMINHO = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jemg-1banco.docx")

backup = CAMINHO.with_suffix(CAMINHO.suffix + ".bak_pre_parametrizacao")
if not backup.exists():
    shutil.copy2(CAMINHO, backup)
    print(f"Backup: {backup.name}")

# === Substituições por parágrafo (chaves curtas, dentro de runs uniformes) ===
SUBS = {
    # PAR 0 — cabeçalho
    0: {
        "___":          "{{vara_protocolo}}",
        "Ipatinga/MG":  "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
    # PAR 5 — qualificação EXEMPLO DA SILVA
    5: {
        "EXEMPLO DA SILVA":              "{{nome_autor}}",
        "brasileiro, casado, aposentado":        "{{nacionalidade}}, {{estado_civil}}, {{profissao}}",
        "000.000.001-11":                         "{{cpf_autor}}",
        "MG-13.921.800":                          "{{rg_autor}}",
        "PC/MG":                                   "{{orgao_expedidor}}",
        "Rua Panaceia":                           "{{logradouro_autor}}",
        "65, Cx B":                               "{{numero_autor}}",
        "Chacaras Madalena":                      "{{bairro_autor}}",
        "Município de Ipatinga/MG":               "Município de {{cidade_autor}}/{{uf_autor}}",
    },
    # PAR 9 — banco réu C6
    9: {
        "BANCO C6 CONSIGNADO S.A.":               "{{banco_reu_nome}}",
        "pessoa jurídica de direito privado":     "{{banco_reu_descricao_pj}}",
        "61.348.538/0001-86":                     "{{banco_reu_cnpj}}",
        "Avenida Nove de Julho, 3148, Jardim Paulista, São Paulo/SP, CEP 01406-000": "{{banco_reu_endereco}}",
    },
    # PAR 11 — benefício
    11: {
        "aposentadoria por invalidez previdenciária":  "{{tipo_beneficio}}",
        "626.148.020-0":                                "{{nb_beneficio}}",
        "agência 1, conta corrente nº 0010197981":      "agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}",
        "junto ao banco Mercantil do Brasil S.A":       "junto ao {{banco_pagador}}",
    },
    # PAR 13 — constatação (vira genérico com BLOCO)
    13: {
        "à empréstimos que não contratou junto ao BANCO C6 CONSIGNADO S/A, CONTRATOS Nº 901364 25359 e 901305 38347:":
            "a empréstimo(s) que não contratou junto ao {{banco_reu_nome}}, conforme detalhamento abaixo:",
    },
    # PAR 44 — renda
    44: {
        "R$ 1.548,00":                                  "R$ {{valor_renda_liquida}} ({{valor_renda_liquida_extenso}})",
    },
    # PAR 216 — intro do pedido (remover "declarar a inexistência dos seguintes empréstimos consignados")
    216: {
        "para o fim de declarar a inexistência dos seguintes empréstimos consignados":
            "para o fim de",
    },
    # PAR 235 — valor causa
    235: {
        "R$ 59.022,10 (cinquenta e nove mil, vinte e dois reais e dez centavos)":
            "R$ {{valor_causa}} ({{valor_causa_extenso}})",
    },
    # PAR 236 — fecho
    236: {
        "Ipatinga/MG":                                 "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
}

doc = Document(CAMINHO)
for par_idx, mapa in SUBS.items():
    p = doc.paragraphs[par_idx]
    if substituir_in_run(p._p, mapa):
        print(f"  par {par_idx}: substituições aplicadas")

# === PAR 15 → "{{BLOCO_INTRO}}" — vai ser absorvido (ou eliminado pelo helper) ===
# Estratégia: par 15 atual diz "No que diz respeito ao referido empréstimo,
# cumpre informar que:" — vou DEIXAR ESSE PAR fixo (já é texto neutro);
# os pares 16 e 17 viram um único marcador BLOCO_CONTRATOS_FRAUDULENTOS
# Como par 15 já diz "No que diz respeito... cumpre informar que:" e meu
# helper começa com "No que diz respeito ao referido empréstimo, cumpre
# informar que: a primeira parcela...", vou:
#   - Substituir par 15 inteiro por "{{BLOCO_CONTRATOS_FRAUDULENTOS}}" (helper já contém intro)
#   - Substituir par 16 e 17 por strings vazias

# Substituir par 15 por marcador único (helper já gera "No que diz respeito...")
p15 = doc.paragraphs[15]
SUB15 = {"No que diz respeito ao referido empréstimo, cumpre informar que:":
         "{{BLOCO_CONTRATOS_FRAUDULENTOS}}"}
if substituir_in_run(p15._p, SUB15):
    print("  par 15: virou {{BLOCO_CONTRATOS_FRAUDULENTOS}}")

# Pars 16 e 17: limpar conteúdo inteiro (deixar parágrafo em branco — Word
# absorve naturalmente; ou removemos via XML)
for pidx in [16, 17]:
    p = doc.paragraphs[pidx]
    # Remover todos os runs do parágrafo (mantém parágrafo, esvazia conteúdo)
    for r in list(p._p.findall(W + 'r')):
        p._p.remove(r)
    print(f"  par {pidx}: esvaziado")

# === PAR 217 → "Declarar a inexistência {{BLOCO_PEDIDO_DECLARACAO}};" ===
p217 = doc.paragraphs[217]
SUB217 = {"No valor de R$ 20.429,93, contrato nº 901364 25359 - com descontos de R$ 438,30 mensais, com inclusão em 13/08/2024, início de desconto em 09/2024, no benefício previdenciário 626.148.020-0":
          "Declarar a inexistência {{BLOCO_PEDIDO_DECLARACAO}}"}
if substituir_in_run(p217._p, SUB217):
    print("  par 217: virou Declarar a inexistência {{BLOCO_PEDIDO_DECLARACAO}};")

# PAR 218 → esvaziar
p218 = doc.paragraphs[218]
for r in list(p218._p.findall(W + 'r')):
    p218._p.remove(r)
print("  par 218: esvaziado")

doc.save(CAMINHO)
print()

# Verificar
import re
d2 = Document(CAMINHO)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r"\{\{[^}]+\}\}", p.text):
        phs.add(m.group(0))
print(f"Placeholders únicos: {len(phs)}")
for ph in sorted(phs):
    print(f"  {ph}")

print("\nVerificação dados pilotos remanescentes:")
sobras = []
for i, p in enumerate(d2.paragraphs):
    t = p.text
    for ruim in ["WELINGTON", "080.268.696", "MG-13.921", "Panaceia",
                  "Chacaras Madalena", "BANCO C6 CONSIGNADO S",
                  "61.348.538", "626.148.020", "Mercantil do Brasil",
                  "1.548,00", "59.022,10", "Ipatinga/MG", "901364 25359",
                  "901305 38347"]:
        if ruim in t:
            sobras.append((i, ruim, t[:120]))
            break
if sobras:
    for i, r, t in sobras:
        print(f"  ⚠️ par {i}: '{r}' em: {t}")
else:
    print("  (vazio = OK)")
