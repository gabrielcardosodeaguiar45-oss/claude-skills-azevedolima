"""Helpers de REDAÇÃO compartilhados entre os 3 pipelines (BA, AM, AL).

REGRAS GLOBAIS GRAVADAS (não pode violar):

1. CABEÇALHO/ENDEREÇAMENTO ('Ao Juízo...'): Segoe UI Bold INLINE (NÃO usar
   rStyle 2TtuloChar, que tem caps automático e estraga o texto).

2. NOMES EM DESTAQUE (autor, banco-réu, INSS): Segoe UI Bold via
   rStyle 2TtuloChar (estilo do template já inclui caps + bold).

3. RESTO DO TEXTO: Cambria.

4. TODA alteração da skill: highlight amarelo.

5. Conjugação masculina/feminina automática (nacionalidade 'brasileira'
   → 'inscrita / domiciliada'; 'brasileiro' → 'inscrito / domiciliado').

6. Estado_civil vazio: omitir limpamente (sem deixar vírgula dupla).

7. Procuração é a ÚNICA fonte autoritativa dos contratos a impugnar.
   NUNCA pegar 'todos os contratos do banco' como fallback silencioso.
"""
import re
from copy import deepcopy
from typing import Dict, List, Optional
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
XMLSPC = '{http://www.w3.org/XML/1998/namespace}space'


# ============================================================
#  PRIMITIVAS DE BAIXO NÍVEL
# ============================================================

def make_run(texto: str, *, bold: bool = False, fonte: str = 'Cambria',
              grifo: bool = True, preserve: bool = True,
              usar_rstyle_titulo: bool = False,
              tamanho_pt: Optional[int] = None):
    """Cria um <w:r> com a configuração desejada.

    Args:
        usar_rstyle_titulo: True para usar rStyle=2TtuloChar (Segoe UI Bold
            COM CAIXA ALTA aplicada automaticamente pelo estilo). USAR APENAS
            para nomes em destaque (NOME DO AUTOR, NOME DO BANCO, INSS).
            Para outros destaques (cabeçalho, títulos comuns), usar
            `bold=True, fonte='Segoe UI', usar_rstyle_titulo=False` para
            Segoe UI Bold inline SEM caps.
        tamanho_pt: tamanho da fonte em pontos (ex.: 11). Se None, herda do
            estilo do parágrafo. OOXML usa meio-pontos (sz=22 → 11pt).
    """
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    r.append(rpr)
    if usar_rstyle_titulo:
        rs = OxmlElement('w:rStyle')
        rs.set(qn('w:val'), '2TtuloChar')
        rpr.append(rs)
    else:
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), fonte)
        rfonts.set(qn('w:hAnsi'), fonte)
        rpr.append(rfonts)
        if bold:
            rpr.append(OxmlElement('w:b'))
            rpr.append(OxmlElement('w:bCs'))
    if tamanho_pt is not None:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(tamanho_pt) * 2))  # OOXML usa meio-pontos
        rpr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(tamanho_pt) * 2))
        rpr.append(szCs)
    if grifo:
        hl = OxmlElement('w:highlight')
        hl.set(qn('w:val'), 'yellow')
        rpr.append(hl)
    t = OxmlElement('w:t')
    t.text = texto
    if preserve:
        t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


def limpar_paragrafo_preservando_pPr(p_elem):
    """Remove todos os filhos do parágrafo MENOS o <w:pPr>."""
    for child in list(p_elem):
        if child.tag != W + 'pPr':
            p_elem.remove(child)


def substituir_paragrafo_completo(p_elem, novo_texto: str, grifo: bool = True,
                                     fonte: str = 'Cambria', bold: bool = False):
    """Substitui o conteúdo inteiro de um parágrafo por 1 run.
    Default: Cambria não-bold.
    Para CABEÇALHO/ENDEREÇAMENTO, usar fonte='Segoe UI', bold=True.
    """
    limpar_paragrafo_preservando_pPr(p_elem)
    p_elem.append(make_run(novo_texto, bold=bold, fonte=fonte, grifo=grifo))


# ============================================================
#  REDAÇÃO CANÔNICA (qualificação, polo passivo, intro, pedidos)
# ============================================================

def _conjugacao_genero(nacionalidade: str):
    """Retorna ('inscrita'/'inscrito', 'residente e domiciliada à'/'residente
    e domiciliado à') conforme nacionalidade.
    """
    fem = (nacionalidade or '').endswith('a')
    return ('inscrita' if fem else 'inscrito',
            'residente e domiciliada à' if fem else 'residente e domiciliado à')


def substituir_qualificacao_autor(p_elem, autora: Dict, end_escritorio: str,
                                     grifo: bool = True, uf_default: str = ''):
    """Reescreve o parágrafo de qualificação do autor preservando o padrão:
    NOME (Segoe UI Bold via rStyle 2TtuloChar) + resto Cambria.

    Compatível com BA/AM/AL. O parâmetro `end_escritorio` é a string completa
    do endereço da filial onde o procurador atua.
    """
    limpar_paragrafo_preservando_pPr(p_elem)
    nome = autora['nome']
    nac = autora.get('nacionalidade', 'brasileiro')
    ec = autora.get('estado_civil') or ''
    prof = autora.get('profissao') or 'aposentado'
    inscrita, domiciliada = _conjugacao_genero(nac)

    quali_partes = [nac]
    if ec:
        quali_partes.append(ec)
    if prof:
        quali_partes.append(prof)
    quali_pessoal = ', '.join(quali_partes)

    uf = autora.get('uf', uf_default)
    cpf = (autora.get('cpf') or '').strip()
    rg = (autora.get('rg') or '').strip()
    orgao = (autora.get('orgao_expedidor') or '').strip()

    # NUNCA escrever placeholders fictícios no RG (ex.: 'XX.XXX.XXX', '[PENDENTE]').
    # Se o RG for igual ao CPF (modelo novo de RG = CPF unificado), inválido,
    # vazio ou marcado como pendente, OMITIR a parte "Cédula de Identidade...".
    # O alerta correspondente vai para o relatório paralelo (responsabilidade
    # do chamador via dados['alertas']).
    rg_invalido = (
        not rg
        or rg.upper().startswith('[PENDENTE')
        or 'XX' in rg or 'xx' in rg
        or rg == cpf
    )
    if rg_invalido:
        bloco_cedula = ''  # OMITIR — relatório paralelo deve alertar
    else:
        bloco_cedula = f'Cédula de Identidade sob nº {rg}, órgão expedidor {orgao}, '

    resto = (
        f', {quali_pessoal}, '
        f'{inscrita} no CPF sob o nº {cpf}, '
        f'{bloco_cedula}'
        f'{domiciliada} {autora.get("logradouro", "")}, '
        f'n° {autora.get("numero", "s/nº")}, bairro {autora.get("bairro", "")}, '
        f'em {autora.get("cidade", "")}/{uf}, '
        f'CEP {autora.get("cep", "")}, não possui endereço eletrônico, por '
        f'seus advogados que assinam digitalmente a presente peça (instrumento '
        f'de procuração anexo), com escritório profissional em {end_escritorio}, '
        f'local onde recebem avisos e intimações, vem, respeitosamente, '
        f'perante Vossa Excelência, propor a presente:'
    )
    p_elem.append(make_run(nome, bold=True, fonte='Segoe UI', grifo=grifo,
                              preserve=False, usar_rstyle_titulo=True))
    p_elem.append(make_run(resto, bold=False, fonte='Cambria', grifo=grifo))


def substituir_polo_passivo(p_elem, bancos_reus: List[Dict],
                                inclui_inss: bool,
                                end_inss: str = ('Av. Sete de Setembro, 1078 - '
                                                  'Mercês, Salvador/BA'),
                                grifo: bool = True):
    """Reescreve o polo passivo:
    - "em face de " (Cambria)
    - NOME BANCO 1 (Segoe UI Bold via rStyle 2TtuloChar)
    - ", pessoa jurídica de direito privado..." (Cambria)
    - se 2 bancos: ", e " + NOME BANCO 2 + ", ..."
    - se Federal: ", e " + INSS (Segoe UI Bold) + ", Autarquia..."
    - " pelos motivos de fato e de direito a seguir expostos:" (Cambria)
    """
    limpar_paragrafo_preservando_pPr(p_elem)
    p_elem.append(make_run('em face de ', bold=False, fonte='Cambria', grifo=False))

    for i, b in enumerate(bancos_reus):
        if i > 0:
            p_elem.append(make_run(', e ', bold=False, fonte='Cambria', grifo=grifo))
        p_elem.append(make_run(b['nome'], bold=True, fonte='Segoe UI', grifo=grifo,
                                 preserve=False, usar_rstyle_titulo=True))
        resto_b = (
            f', {b["descricao_pj"]}, '
            f'inscrita no CNPJ sob o nº {b["cnpj"]}, '
            f'com endereço na {b["endereco"]}'
        )
        p_elem.append(make_run(resto_b, bold=False, fonte='Cambria', grifo=grifo))

    if inclui_inss:
        p_elem.append(make_run(', e ', bold=False, fonte='Cambria', grifo=grifo))
        p_elem.append(make_run('INSTITUTO NACIONAL DO SEGURO SOCIAL - INSS',
                                 bold=True, fonte='Segoe UI', grifo=grifo,
                                 preserve=False, usar_rstyle_titulo=True))
        p_elem.append(make_run(
            f', Autarquia Federal, com sede à {end_inss}',
            bold=False, fonte='Cambria', grifo=grifo))

    p_elem.append(make_run(' pelos motivos de fato e de direito a seguir expostos:',
                             bold=False, fonte='Cambria', grifo=False))


def substituir_intro_contratos(p_elem, nome_banco: str, numeros: List[str],
                                  grifo: bool = True):
    """Reescreve o parágrafo "Nessa oportunidade... tomou conhecimento dos
    descontos referentes a empréstimo(s) que não contratou junto ao BANCO X,
    CONTRATO Nº ...:" com BANCO e CONTRATO em Segoe UI Bold + amarelo.
    """
    limpar_paragrafo_preservando_pPr(p_elem)
    n = len(numeros)
    if n == 0:
        rotulo = ''
    elif n == 1:
        rotulo = f'CONTRATO Nº {numeros[0]}'
    elif n == 2:
        rotulo = f'CONTRATOS Nº {numeros[0]} e {numeros[1]}'
    else:
        rotulo = f'CONTRATOS Nº {", ".join(numeros[:-1])} e {numeros[-1]}'

    artigo = 'a empréstimo' if n == 1 else 'a empréstimos'

    # NOME DO BANCO e CONTRATO Nº — em Cambria Bold + CAPS (regra fixa
    # 07/05/2026, Gabriel — substituiu o Segoe UI Bold via rStyle 2TtuloChar).
    p_elem.append(make_run(
        f'Nessa oportunidade, após informações, tomou conhecimento dos '
        f'descontos referentes {artigo} que não contratou junto ao ',
        bold=False, fonte='Cambria', grifo=grifo))
    p_elem.append(make_run(nome_banco.upper(), bold=True, fonte='Cambria',
                              grifo=grifo, preserve=False, usar_rstyle_titulo=False))
    p_elem.append(make_run(', ', bold=False, fonte='Cambria', grifo=grifo))
    p_elem.append(make_run(rotulo.upper(), bold=True, fonte='Cambria',
                              grifo=grifo, preserve=False, usar_rstyle_titulo=False))
    p_elem.append(make_run(':', bold=False, fonte='Cambria', grifo=False))


def substituir_intro_contratos_multi_banco(p_elem, grupos: List[Dict],
                                             grifo: bool = True):
    """Reescreve a intro fática quando há contratos de 2+ bancos diferentes.

    `grupos`: lista de dicts no formato:
        [{'banco_nome': 'BANCO X', 'numeros': ['111', '222']},
         {'banco_nome': 'BANCO Y', 'numeros': ['333']}]

    Frase gerada:
        "Nessa oportunidade, após informações, tomou conhecimento dos
        descontos referentes a empréstimos que não contratou junto ao
        BANCO X, CONTRATOS Nº 111 E 222, e ao BANCO Y, CONTRATO Nº 333:"
    """
    limpar_paragrafo_preservando_pPr(p_elem)
    if not grupos:
        return
    if len(grupos) == 1:
        substituir_intro_contratos(p_elem, grupos[0]['banco_nome'],
                                     grupos[0]['numeros'], grifo=grifo)
        return
    p_elem.append(make_run(
        'Nessa oportunidade, após informações, tomou conhecimento dos '
        'descontos referentes a empréstimos que não contratou junto ao ',
        bold=False, fonte='Cambria', grifo=grifo))
    for i, g in enumerate(grupos):
        if i > 0:
            sep = ' e ao ' if i == len(grupos) - 1 else ', ao '
            p_elem.append(make_run(sep, bold=False, fonte='Cambria', grifo=grifo))
        p_elem.append(make_run(g['banco_nome'].upper(), bold=True, fonte='Cambria',
                                  grifo=grifo, preserve=False,
                                  usar_rstyle_titulo=False))
        p_elem.append(make_run(', ', bold=False, fonte='Cambria', grifo=grifo))
        nums = g['numeros']
        n = len(nums)
        if n == 1:
            rotulo = f'CONTRATO Nº {nums[0]}'
        elif n == 2:
            rotulo = f'CONTRATOS Nº {nums[0]} e {nums[1]}'
        else:
            rotulo = f'CONTRATOS Nº {", ".join(nums[:-1])} e {nums[-1]}'
        p_elem.append(make_run(rotulo.upper(), bold=True, fonte='Cambria',
                                  grifo=grifo, preserve=False,
                                  usar_rstyle_titulo=False))
    p_elem.append(make_run(':', bold=False, fonte='Cambria', grifo=False))


def aplicar_intro_fatica(p_elem, contratos_fmt: List[Dict],
                            fallback_banco_nome: str, grifo: bool = True):
    """Aplica a intro fática automaticamente — decide entre 1-banco ou multi-banco
    a partir dos `banco_nome`/`banco` dos contratos.

    Compartilhado por TODOS os pipelines (BA, AM, AL) para garantir comportamento
    idêntico. Quando há ≥2 bancos distintos, agrupa por banco e gera:

        "Nessa oportunidade, após informações, tomou conhecimento dos descontos
        referentes a empréstimos que não contratou junto ao BANCO X, CONTRATOS Nº
        111 E 222, e ao BANCO Y, CONTRATO Nº 333:"

    Quando há 1 banco só, cai no caminho clássico `substituir_intro_contratos`.

    Args:
        p_elem: <w:p> do parágrafo da intro fática
        contratos_fmt: lista de contratos formatados (cada um com 'numero' e
            'banco_nome' ou 'banco')
        fallback_banco_nome: nome do banco default (banco_reu['nome']),
            usado quando o contrato não tem banco_nome próprio
        grifo: aplica highlight amarelo nas substituições
    """
    grupos_dict = {}
    for c in contratos_fmt:
        if not c.get('numero'):
            continue
        bn = (c.get('banco_nome') or c.get('banco')
              or fallback_banco_nome or '').strip()
        if not bn:
            continue
        grupos_dict.setdefault(bn, []).append(c['numero'])
    grupos = [{'banco_nome': bn, 'numeros': nums}
              for bn, nums in grupos_dict.items()]
    if len(grupos) >= 2:
        substituir_intro_contratos_multi_banco(p_elem, grupos, grifo=grifo)
    elif len(grupos) == 1:
        substituir_intro_contratos(p_elem, grupos[0]['banco_nome'],
                                     grupos[0]['numeros'], grifo=grifo)
    elif fallback_banco_nome:
        # contratos sem números — usar todos do fallback
        nums = [c.get('numero', '') for c in contratos_fmt if c.get('numero')]
        substituir_intro_contratos(p_elem, fallback_banco_nome, nums, grifo=grifo)


def modalidade_extenso(tipo_origem: str) -> str:
    """Mapeia tipo_origem do contrato para o termo nos pedidos:
    'empréstimo' / 'refinanciamento' / 'empréstimo (portabilidade)'.
    Default: 'empréstimo' (mais conservador).
    """
    return {
        'original':         'empréstimo',
        'averbacao_nova':   'empréstimo',
        'avn':              'empréstimo',
        'refinanciamento':  'refinanciamento',
        'refin':            'refinanciamento',
        'portabilidade':    'empréstimo (portabilidade)',
        'migracao':         'empréstimo',
    }.get((tipo_origem or '').lower(), 'empréstimo')


def preencher_pedidos_declaratorios(doc, contratos_fmt: List[Dict],
                                       nb_beneficio: str,
                                       grifo: bool = True) -> int:
    """Reescreve o(s) parágrafo(s) "Declarar a inexistência do
    empréstimo/refinanciamento consignado..." com:
    - escolha empréstimo vs refinanciamento conforme tipo_origem
    - dados reais (valor, número, parcela, data, NB)
    - 1 pedido por contrato (duplica o parágrafo se N > 1)
    """
    feitos = 0
    pars = list(doc.paragraphs)
    idx_pedido = None
    formato_mg = False  # padrão MG: cabeçalho "...inexistência dos seguintes empréstimos:" + N sub-itens
    for ip, p in enumerate(pars):
        # Padrão BA/AM/AL: "Declarar a inexistência do [empréstimo|refinanciamento]"
        if ('Declarar a inexistência do empréstimo' in p.text
            or 'Declarar a inexistência do refinanciamento' in p.text):
            idx_pedido = ip
            break
        # Padrão MG: "...declarar a inexistência dos seguintes empréstimos consignados:"
        if ('declarar a inexistência dos seguintes empréstimos' in p.text.lower()
            or 'declarar a inexistência dos seguintes refinanciamentos' in p.text.lower()):
            idx_pedido = ip
            formato_mg = True
            break
    if idx_pedido is None:
        return 0

    # Formato MG: o parágrafo PEDIDO é só o cabeçalho. Os SUB-ITENS
    # ("No valor de R$ X, contrato nº Y...") vêm nos parágrafos seguintes
    # com pStyle de lista. Preencher sub-itens (1 por contrato; remover ou
    # adicionar conforme N).
    if formato_mg:
        return _preencher_pedidos_formato_mg(
            doc, idx_pedido, contratos_fmt, nb_beneficio, grifo)

    p_orig = pars[idx_pedido]._element
    pPr_orig = p_orig.find(W + 'pPr')
    pPr_xml = deepcopy(pPr_orig) if pPr_orig is not None else None

    try:
        from extenso import extenso_moeda as _ext
    except Exception:
        _ext = lambda v: ''

    def _montar(c):
        modalidade = modalidade_extenso(c.get('tipo_origem'))
        try:
            ve = float((c.get('valor_emprestado_str') or '').replace('.', '').replace(',', '.'))
            ve_ext = _ext(ve)
        except Exception:
            ve_ext = ''
        try:
            vp = float((c.get('valor_parcela_str') or '').replace('.', '').replace(',', '.'))
            vp_ext = _ext(vp)
        except Exception:
            vp_ext = ''
        return (
            f'Declarar a inexistência do {modalidade} consignado no valor de '
            f'R$ {c.get("valor_emprestado_str", "")} ({ve_ext}), contrato nº '
            f'{c.get("numero", "")}, com descontos de R$ {c.get("valor_parcela_str", "")} '
            f'({vp_ext}) mensais, com inclusão em {c.get("data_inclusao_str", "")}, '
            f'início de desconto em {c.get("competencia_inicio_str", "")}, no benefício '
            f'previdenciário {nb_beneficio};'
        )

    # 1º contrato → reescreve o parágrafo original
    limpar_paragrafo_preservando_pPr(p_orig)
    p_orig.append(make_run(_montar(contratos_fmt[0]),
                              bold=False, fonte='Cambria', grifo=grifo))
    feitos += 1

    # 2º+ contratos → INSERIR novos parágrafos APÓS o original
    elem_anterior = p_orig
    for c in contratos_fmt[1:]:
        novo_par = OxmlElement('w:p')
        if pPr_xml is not None:
            novo_par.append(deepcopy(pPr_xml))
        novo_par.append(make_run(_montar(c), bold=False, fonte='Cambria', grifo=grifo))
        elem_anterior.addnext(novo_par)
        elem_anterior = novo_par
        feitos += 1

    return feitos


def _preencher_pedidos_formato_mg(doc, idx_cabecalho: int,
                                     contratos_fmt: List[Dict],
                                     nb_beneficio: str,
                                     grifo: bool = True) -> int:
    """Preenche os pedidos no FORMATO MG:
        [N]   "...declarar a inexistência dos seguintes empréstimos consignados:"
        [N+1] "No valor de R$ X, contrato nº Y - com descontos de R$ Z mensais,
                com inclusão em DD/MM/AAAA, início de desconto em MM/AAAA,
                no benefício previdenciário NB;"
        [N+2] "No valor de R$ ..."
        ...

    A skill identifica os parágrafos sub-item adjacentes ao cabeçalho (que
    começam com "No valor de R$") e:
      - Reescreve o 1º com dados do 1º contrato
      - Reescreve o 2º com dados do 2º contrato (se existir)
      - Remove sub-itens excedentes ou adiciona novos conforme N de contratos
    """
    feitos = 0
    pars = list(doc.paragraphs)

    # Encontrar todos os sub-itens "No valor de R$" CONSECUTIVOS após o cabeçalho
    idx_sub_itens = []
    pPr_sub_xml = None
    for ip in range(idx_cabecalho + 1, min(idx_cabecalho + 20, len(pars))):
        t = pars[ip].text.strip()
        if t.startswith('No valor de R$') or t.startswith('no valor de R$'):
            idx_sub_itens.append(ip)
            if pPr_sub_xml is None:
                pPr_orig = pars[ip]._element.find(W + 'pPr')
                if pPr_orig is not None:
                    pPr_sub_xml = deepcopy(pPr_orig)
        elif t and not t.startswith('No valor de R$'):
            break  # parou a sequência de sub-itens

    if not idx_sub_itens:
        return 0

    try:
        from extenso import extenso_moeda as _ext
    except Exception:
        _ext = lambda v: ''

    def _montar_subitem(c):
        try:
            ve = float((c.get('valor_emprestado_str') or '').replace('.', '').replace(',', '.'))
            ve_ext = _ext(ve)
        except Exception:
            ve_ext = ''
        return (
            f'No valor de R$ {c.get("valor_emprestado_str", "")} ({ve_ext}), '
            f'contrato nº {c.get("numero", "")} — com descontos de R$ '
            f'{c.get("valor_parcela_str", "")} mensais, com inclusão em '
            f'{c.get("data_inclusao_str", "")}, início de desconto em '
            f'{c.get("competencia_inicio_str", "")}, no benefício previdenciário '
            f'{nb_beneficio};'
        )

    n_contratos = len(contratos_fmt)
    n_sub_existentes = len(idx_sub_itens)

    # 1. Reescrever sub-itens existentes (até o min entre N existentes e N contratos)
    for i in range(min(n_sub_existentes, n_contratos)):
        p_elem = pars[idx_sub_itens[i]]._element
        limpar_paragrafo_preservando_pPr(p_elem)
        p_elem.append(make_run(_montar_subitem(contratos_fmt[i]),
                                  bold=False, fonte='Cambria', grifo=grifo))
        feitos += 1

    # 2a. Se há MAIS contratos do que sub-itens: ADICIONAR sub-itens novos
    if n_contratos > n_sub_existentes:
        elem_anterior = pars[idx_sub_itens[-1]]._element
        for c in contratos_fmt[n_sub_existentes:]:
            novo_par = OxmlElement('w:p')
            if pPr_sub_xml is not None:
                novo_par.append(deepcopy(pPr_sub_xml))
            novo_par.append(make_run(_montar_subitem(c),
                                       bold=False, fonte='Cambria', grifo=grifo))
            elem_anterior.addnext(novo_par)
            elem_anterior = novo_par
            feitos += 1

    # 2b. Se há MENOS contratos do que sub-itens: REMOVER sub-itens excedentes
    if n_contratos < n_sub_existentes:
        for i in range(n_contratos, n_sub_existentes):
            p_excedente = pars[idx_sub_itens[i]]._element
            p_excedente.getparent().remove(p_excedente)
            feitos += 1

    return feitos


def preencher_bloco_fatico_formato_mg(doc, contratos_fmt: List[Dict],
                                          nome_banco: str,
                                          grifo: bool = True) -> int:
    """Preenche bloco fático no FORMATO MG/AL-2bancos:
        [N]   "No que diz respeito ao referido empréstimo, cumpre informar que:"
        [N+1] "Do contrato nº A: a primeira parcela descontada do benefício...
                competência X, total de Y parcelas, no valor de R$ Z, relativas
                a um empréstimo consignado no valor de R$ W, cuja operação foi
                realizada pelo banco BANCO_NOME, ora requerido."
        [N+2] "Do contrato nº B: ..."
        ...

    Reescreve os sub-itens "Do contrato nº..." adjacentes ao cabeçalho,
    preservando pPr (numeração de lista). Adiciona ou remove sub-itens
    conforme o número de contratos passado.
    """
    feitos = 0
    pars = list(doc.paragraphs)

    # Achar cabeçalho "No que diz respeito"
    idx_cab = None
    for ip, p in enumerate(pars):
        t = p.text.strip()
        if t.startswith('No que diz respeito'):
            idx_cab = ip
            break
    if idx_cab is None:
        return 0

    # Achar sub-itens "Do contrato" CONSECUTIVOS após cabeçalho
    idx_sub = []
    pPr_sub_xml = None
    for ip in range(idx_cab + 1, min(idx_cab + 15, len(pars))):
        t = pars[ip].text.strip()
        if t.startswith('Do contrato'):
            idx_sub.append(ip)
            if pPr_sub_xml is None:
                pPr_orig = pars[ip]._element.find(W + 'pPr')
                if pPr_orig is not None:
                    pPr_sub_xml = deepcopy(pPr_orig)
        elif t and not t.startswith('Do contrato'):
            break

    if not idx_sub:
        return 0  # template não tem formato MG (provavelmente é AL formato simples)

    try:
        from extenso import extenso_moeda as _ext
    except Exception:
        _ext = lambda v: ''

    def _montar_subitem(c):
        try:
            vp = float((c.get('valor_parcela_str') or '').replace('.', '').replace(',', '.'))
            vp_ext = _ext(vp)
        except Exception:
            vp_ext = ''
        try:
            ve = float((c.get('valor_emprestado_str') or '').replace('.', '').replace(',', '.'))
            ve_ext = _ext(ve)
        except Exception:
            ve_ext = ''
        return (
            f'Do contrato nº {c.get("numero", "")}: a primeira parcela '
            f'descontada do benefício da parte autora foi na competência '
            f'{c.get("competencia_inicio_str", "")}, de um total de '
            f'{c.get("qtd_parcelas", "")} parcelas, no valor de '
            f'R$ {c.get("valor_parcela_str", "")} ({vp_ext}), relativas a um '
            f'empréstimo consignado no valor de R$ {c.get("valor_emprestado_str", "")} '
            f'({ve_ext}), cuja operação foi realizada pelo {nome_banco}, ora requerido.'
        )

    n_contratos = len(contratos_fmt)
    n_sub = len(idx_sub)

    # 1. Reescrever sub-itens existentes
    for i in range(min(n_sub, n_contratos)):
        p_elem = pars[idx_sub[i]]._element
        limpar_paragrafo_preservando_pPr(p_elem)
        p_elem.append(make_run(_montar_subitem(contratos_fmt[i]),
                                  bold=False, fonte='Cambria', grifo=grifo))
        feitos += 1

    # 2a. Mais contratos do que sub-itens: ADICIONAR
    if n_contratos > n_sub:
        elem_anterior = pars[idx_sub[-1]]._element
        for c in contratos_fmt[n_sub:]:
            novo_par = OxmlElement('w:p')
            if pPr_sub_xml is not None:
                novo_par.append(deepcopy(pPr_sub_xml))
            novo_par.append(make_run(_montar_subitem(c),
                                       bold=False, fonte='Cambria', grifo=grifo))
            elem_anterior.addnext(novo_par)
            elem_anterior = novo_par
            feitos += 1

    # 2b. Menos contratos do que sub-itens: REMOVER
    if n_contratos < n_sub:
        for i in range(n_contratos, n_sub):
            p_excedente = pars[idx_sub[i]]._element
            p_excedente.getparent().remove(p_excedente)
            feitos += 1

    return feitos


def inserir_prioridade_idoso_se_faltando(doc, eh_idoso: bool,
                                            grifo: bool = True) -> int:
    """Se o autor é idoso E o template NÃO tem o parágrafo de prioridade no
    cabeçalho, INSERE o parágrafo logo após o cabeçalho "Ao Juízo...".
    Estilo: Cambria 11pt (não Segoe UI Bold — regra fixa, gravada 07/05/2026).
    Devolve número de parágrafos inseridos (0 ou 1).
    """
    if not eh_idoso:
        return 0
    pars = list(doc.paragraphs)
    # Verificar se já tem
    for p in pars:
        if 'Prioridade de tramitação: art. 1.048' in p.text:
            return 0  # já tem, nada a fazer

    # Localizar cabeçalho "Ao Juízo..." para inserir DEPOIS
    idx_cabecalho = None
    for ip, p in enumerate(pars):
        if 'Ao Juízo' in p.text or 'Juízo do Juizado' in p.text:
            idx_cabecalho = ip
            break
    if idx_cabecalho is None:
        return 0

    cabecalho_elem = pars[idx_cabecalho]._element

    # Criar novo parágrafo (Cambria 11pt + DIREITA + RECUO 4cm).
    # Regras fixas do escritório (07/05/2026):
    #   - alinhamento à DIREITA
    #   - Cambria 11pt (não Segoe UI Bold)
    #   - recuo esquerdo de 4cm = 2268 twips (1 cm = 567 twips)
    novo_par = OxmlElement('w:p')
    novo_pPr = OxmlElement('w:pPr')
    # Recuo esquerdo de 4 cm
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '2268')
    novo_pPr.append(ind)
    # Forçar alinhamento à direita
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'right')
    novo_pPr.append(jc)
    novo_par.append(novo_pPr)
    novo_par.append(make_run(
        'Prioridade de tramitação: art. 1.048 do Código de Processo Civil (Idoso).',
        bold=False, fonte='Cambria', tamanho_pt=11, grifo=grifo))
    cabecalho_elem.addnext(novo_par)
    return 1


def inserir_pedido_prioridade_idoso_se_faltando(doc, eh_idoso: bool,
                                                    grifo: bool = True) -> int:
    """Se o autor é idoso E o template NÃO tem o PEDIDO de prioridade idoso
    nos pedidos, INSERE o pedido como PRIMEIRO item da lista de pedidos
    (numeração romana I, II, III...), HERDANDO o pStyle do primeiro item
    da lista para preservar a numeração automática.

    Estratégia: localiza o PRIMEIRO parágrafo da lista de pedidos (com
    pStyle=5Listaalfabtica ou similar de lista numerada). Insere o novo
    pedido ANTES dele, com o mesmo pPr. O Word renumera automaticamente.

    Texto inserido:
        "A prioridade na tramitação, tendo em vista que a parte autora é
         pessoa idosa, nos termos do art. 1.048, inciso I, do Código de
         Processo Civil;"

    Devolve número de parágrafos inseridos (0 ou 1).
    """
    if not eh_idoso:
        return 0
    pars = list(doc.paragraphs)

    # Verificar se já tem o pedido
    for p in pars:
        t = p.text.strip()
        if (t.startswith('A prioridade na tramitação')
            or 'prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa' in t):
            return 0  # já tem, nada a fazer

    # Localizar "DOS PEDIDOS" para começar a busca
    idx_dos_pedidos = None
    for ip, p in enumerate(pars):
        if p.text.strip() == 'DOS PEDIDOS':
            idx_dos_pedidos = ip
            break
    if idx_dos_pedidos is None:
        return 0

    # PROCURAR o PRIMEIRO parágrafo de LISTA NUMERADA após "DOS PEDIDOS"
    # (pStyle=5Listaalfabtica é o padrão do escritório para lista romana I,II,III)
    idx_primeiro_item = None
    pPr_lista_xml = None
    for ip in range(idx_dos_pedidos + 1, min(idx_dos_pedidos + 10, len(pars))):
        pPr_p = pars[ip]._element.find(W + 'pPr')
        if pPr_p is None:
            continue
        pStyle_p = pPr_p.find(W + 'pStyle')
        if pStyle_p is None:
            continue
        style_val = pStyle_p.get(W + 'val', '')
        # Padrões de lista numerada do escritório:
        #   '5Listaalfabtica', 'PargrafodaLista', etc.
        if 'Lista' in style_val or 'lista' in style_val:
            idx_primeiro_item = ip
            pPr_lista_xml = deepcopy(pPr_p)
            break

    if idx_primeiro_item is None:
        # Fallback: usar pPr do parágrafo logo após "DOS PEDIDOS"
        if idx_dos_pedidos + 1 < len(pars):
            pPr_proximo = pars[idx_dos_pedidos + 1]._element.find(W + 'pPr')
            if pPr_proximo is not None:
                pPr_lista_xml = deepcopy(pPr_proximo)
                idx_primeiro_item = idx_dos_pedidos + 1

    if idx_primeiro_item is None:
        return 0

    # Criar novo parágrafo HERDANDO o pPr do primeiro item da lista
    # (assim entra na numeração automática como item I).
    # Texto em NEGRITO (regra fixa 07/05/2026, Gabriel).
    novo_par = OxmlElement('w:p')
    if pPr_lista_xml is not None:
        novo_par.append(pPr_lista_xml)
    novo_par.append(make_run(
        'A prioridade na tramitação, tendo em vista que a parte autora é '
        'pessoa idosa, nos termos do art. 1.048, inciso I, do Código de '
        'Processo Civil;',
        bold=True, fonte='Cambria', grifo=grifo))
    # INSERIR ANTES do primeiro item (vira o novo item I)
    primeiro_item_elem = pars[idx_primeiro_item]._element
    primeiro_item_elem.addprevious(novo_par)
    return 1


def remover_prioridade_pedidos(doc) -> int:
    """Remove o parágrafo dos pedidos que pleitea prioridade idoso quando
    o autor não é idoso.
    """
    feitos = 0
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if (t.startswith('A prioridade na tramitação')
            or 'prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa' in t):
            p._element.getparent().remove(p._element)
            feitos += 1
    return feitos


def montar_paragrafo_recebimento_beneficio(hiscon_cab: Dict,
                                                autora: Dict = None,
                                                tipo_benef_default: str = 'aposentadoria por idade') -> str:
    """Monta o texto canônico do parágrafo de síntese fática que descreve
    como o autor recebe seu benefício previdenciário.

    Duas variantes (regra fixa do escritório, gravada 07/05/2026):

    (a) CONTA BANCÁRIA — quando o autor recebe POR DEPÓSITO em conta:
        "A parte autora recebe benefício previdenciário de {tipo_benef} –
         NB {nb}, o qual é depositado em conta bancária, agência {ag},
         conta corrente nº {cc}, junto ao {banco_pagador}, sendo que
         sobrevive basicamente do que recebe da previdência."

    (b) CARTÃO MAGNÉTICO — quando o autor recebe POR CARTÃO MAGNÉTICO
        (não tem conta bancária):
        "A parte autora recebe benefício previdenciário de {tipo_benef} –
         NB {nb}, o qual é sacado por meio de cartão magnético na agência
         {ag}, junto ao {banco_pagador}, sendo que sobrevive basicamente
         do que recebe da previdência."

    Detecção da forma:
      1. Se `autora['forma_recebimento']` for explicitamente passado:
         - 'cartao_magnetico' → variante (b)
         - 'conta_bancaria'   → variante (a)
      2. Se NÃO foi passado, detecta automaticamente:
         - hiscon_cab['conta_pagador'] vazio → variante (b)
         - hiscon_cab['conta_pagador'] preenchido → variante (a)

    Args:
        hiscon_cab: dict com `tipo_beneficio`, `nb_beneficio`,
                    `agencia_pagador`, `conta_pagador`, `banco_pagador`.
        autora: dict da autora. Aceita chave opcional `forma_recebimento`
                ('cartao_magnetico' / 'conta_bancaria') para forçar a variante.
        tipo_benef_default: usado quando `tipo_beneficio` está vazio.

    Returns:
        str pronta para `substituir_paragrafo_completo`.
    """
    tipo_benef = (hiscon_cab.get('tipo_beneficio') or tipo_benef_default).lower()
    nb = hiscon_cab.get('nb_beneficio') or '___'
    ag = hiscon_cab.get('agencia_pagador') or '___'
    cc = hiscon_cab.get('conta_pagador') or ''
    bp = hiscon_cab.get('banco_pagador') or 'BANCO ___'

    forma_explicita = (autora or {}).get('forma_recebimento', '').strip().lower()
    if forma_explicita in ('cartao_magnetico', 'cartão_magnético', 'cartao', 'cartão'):
        usar_cartao = True
    elif forma_explicita in ('conta_bancaria', 'conta_bancária', 'conta', 'deposito'):
        usar_cartao = False
    else:
        # Detecção automática: sem conta_pagador no HISCON → presume cartão
        usar_cartao = not cc

    if usar_cartao:
        return (
            f'A parte autora recebe benefício previdenciário de {tipo_benef} – '
            f'NB {nb}, o qual é sacado por meio de cartão magnético na agência '
            f'{ag}, junto ao {bp}, sendo que sobrevive basicamente do que '
            f'recebe da previdência.'
        )
    return (
        f'A parte autora recebe benefício previdenciário de {tipo_benef} – '
        f'NB {nb}, o qual é depositado em conta bancária, agência {ag}, '
        f'conta corrente nº {cc}, junto ao {bp}, sendo que sobrevive '
        f'basicamente do que recebe da previdência.'
    )


def inserir_unidade_apoio_se_faltando(doc, uf: str, grifo: bool = True) -> int:
    """Padroniza o trecho do endereço do escritório nos templates HARDCODED
    (BA e AM) para o formato canônico:
        "Rua Frei Rogério, 541, Centro, Joaçaba/SC, CEP 89600-000, e
         unidade de apoio em [endereço da filial da UF do cliente]"

    Substitui run-aware via `substituir_in_run` para preservar o restante do
    parágrafo (placeholders, rPr de outros runs, etc.). Idempotente: se o
    parágrafo já tem o endereço composto correto para a UF, não faz nada.

    Casos cobertos:
      - Template AM: 'Rua Frei Rogério, 541, Centro, no município de
        Joaçaba-SC, CEP 89.600-000' (sem apoio)
      - Template BA: 'Rua Frei Rogério, 541, Centro, no município de
        Joaçaba-SC, CEP 89.600-000 e Rua Portugal, 5 - Sala 505 - Comercio,
        Salvador - BA, CEP 40015-000' (com apoio antigo Salvador)
      - Variantes com pontuação ligeiramente distinta.

    Devolve número de parágrafos atualizados.
    """
    from helpers_docx import substituir_in_run
    from escritorios import (montar_endereco_escritorio_completo,
                              FILIAL_APOIO_POR_UF, ENDERECOS_FILIAIS)
    novo_endereco = montar_endereco_escritorio_completo(uf)
    cidade_apoio = FILIAL_APOIO_POR_UF.get((uf or '').upper())

    # Padrões antigos a substituir — ordenar do mais longo p/ mais curto
    # para evitar match parcial (ex.: o do BA contém o do AM).
    PADROES_OLD = [
        ('Rua Frei Rogério, 541, Centro, no município de Joaçaba-SC, '
         'CEP 89.600-000 e Rua Portugal, 5 - Sala 505 - Comercio, '
         'Salvador - BA, CEP 40015-000'),
        ('Rua Frei Rogério, 541, Centro, no município de Joaçaba-SC, '
         'CEP 89.600-000'),
        ('Rua Frei Rogério, 541, Centro, Joaçaba/SC, CEP 89600-000'),
    ]

    feitos = 0
    for p in list(doc.paragraphs):
        if 'Frei Rogério' not in p.text:
            continue
        # Se a cidade da filial-apoio já aparece, considerar OK
        if cidade_apoio:
            cid_curta = cidade_apoio.split('/')[0]
            cep_apoio = ENDERECOS_FILIAIS[cidade_apoio]['cep']
            cep_norm = cep_apoio.replace('-', '')
            if (cid_curta in p.text
                and 'unidade de apoio em' in p.text
                and (cep_apoio in p.text or cep_norm in p.text.replace('-', ''))):
                continue
        # Procurar primeiro padrão que casa
        for padrao in PADROES_OLD:
            if padrao in p.text:
                if substituir_in_run(p._element, {padrao: novo_endereco},
                                          grifo=grifo):
                    feitos += 1
                break
    return feitos
