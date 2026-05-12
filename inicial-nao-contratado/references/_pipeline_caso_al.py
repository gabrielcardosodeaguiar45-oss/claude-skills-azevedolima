"""Pipeline para iniciais AL (JEF Federal AL OU TJAL Estadual rito comum).

Diferenças vs pipelines BA e AM:
  - 4 templates: Federal/Estadual × 1banco/2bancos
  - Foro decidido por valor da causa (≤60 SM = JEF AL; >60 SM = TJAL)
    OU pelo procurador (override por sorteio/escolha estratégica)
  - Procurador default: Tiago (OAB/AL 20906A); transição → Alexandre
  - 12 blocos fáticos pré-prontos no template (1/2 contratos × Ativo/Inativo
    × Com/Sem depósito × AVN/Refin) — a skill SELECIONA o bloco apropriado
    e REMOVE os outros 11
  - Polo passivo no Federal inclui INSS; no Estadual só o banco

ESTRATÉGIA: NÃO instrumentar templates AL com placeholders {{...}}. Em vez
disso, fazer substituições TARGETED no DOCX usando string matching dos
textos literais do caso piloto (FULANO DE TAL, BANCO BRADESCO, 0123506012709,
"Subseção de  /AL", etc.).

Uso:
    dados = montar_dados_inicial_al(
        pasta_cliente='.../EDMUNDA LIMA DOS SANTOS',
        autora=AUTORA_EDMUNDA,
        comarca='Arapiraca',
        forcar_foro=None,        # 'federal' | 'estadual' | None (auto)
        forcar_procurador='tiago' # 'tiago' | 'alexandre'
    )
    res = gerar_inicial_al(dados, output_path='.../INICIAL_EDMUNDA.docx')
"""
import os, re, glob, sys, shutil
from copy import deepcopy
from datetime import datetime
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers_docx import substituir_in_run
from extenso import extenso_moeda
from extrator_hiscon import (parse_hiscon, filtrar_contratos_por_numero,
                              formatar_contrato_para_template,
                              auditar_procuracoes_vs_hiscon)
from extrator_hiscre import parse_hiscre
from extrator_calculo import parse_calculo, eh_idoso
from extrator_procuracao import extrair_numeros_contrato_de_pasta
from bancos_canonicos import resolver_banco
from helpers_redacao import (
    make_run as _make_run,
    limpar_paragrafo_preservando_pPr as _limpar_paragrafo_preservando_pPr,
    substituir_paragrafo_completo as _substituir_paragrafo_completo,
    substituir_qualificacao_autor,
    substituir_polo_passivo as _substituir_polo_passivo,
    substituir_intro_contratos as _substituir_intro_contratos,
    modalidade_extenso as _modalidade_extenso,
    preencher_pedidos_declaratorios as _preencher_pedidos_declaratorios,
    remover_prioridade_pedidos as _remover_prioridade_pedidos,
    preencher_bloco_fatico_formato_mg as _preencher_bloco_fatico_formato_mg,
    inserir_prioridade_idoso_se_faltando as _inserir_prioridade_idoso_se_faltando,
    inserir_pedido_prioridade_idoso_se_faltando as _inserir_pedido_prioridade_idoso_se_faltando,
    montar_paragrafo_recebimento_beneficio as _montar_paragrafo_recebimento_beneficio,
)
from auditor_dano_moral import calcular_dano_moral, auditar_dano_moral
from verificador_dados_pessoais import comparar_doc_vs_hiscre, consolidar_dados_autora
from escritorios import (PROCURADORES, selecionar_procurador,
                          selecionar_template_por_uf, decidir_foro_al,
                          ENDERECOS_FILIAIS)


class ProcuracaoSemFiltroError(RuntimeError):
    """Erro CRÍTICO: pipeline tentou rodar sem filtro válido de contratos.

    A procuração é a ÚNICA fonte autoritativa do que o cliente nos autorizou
    a impugnar. Se o pipeline não conseguir extrair os números de contrato:
      - do nome do arquivo da procuração; OU
      - do CONTEÚDO da procuração (text-layer ou OCR via easyocr); OU
      - de `numeros_contrato_explicitos` passado pelo chamador,
    DEVE ABORTAR — nunca cair em "pegar todos os contratos do banco".

    Para resolver: ler o PDF da procuração manualmente e passar os números
    via `numeros_contrato_explicitos=[...]` no `montar_dados_inicial_al`.
    """
    pass

VAULT_TEMPLATES = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates'

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
XMLSPC = '{http://www.w3.org/XML/1998/namespace}space'


# ============================================================
#  HELPERS DE PASTA E ARQUIVO
# ============================================================

def encontrar_pdf(pasta: str, *padroes: str) -> Optional[str]:
    """Encontra PDF na pasta cujo nome (lower) contém qualquer um dos padrões."""
    if not os.path.isdir(pasta):
        return None
    for arq in os.listdir(pasta):
        if not arq.lower().endswith('.pdf'):
            continue
        nome_lower = arq.lower()
        if any(p.lower() in nome_lower for p in padroes):
            return os.path.join(pasta, arq)
    return None


def listar_procuracoes(pasta_cliente: str) -> List[Dict]:
    """Lista os arquivos de procuração na pasta do cliente.

    Convenção AL (ex. EDMUNDA LIMA DOS SANTOS):
        '2 - Procuração — N°1.pdf', '2 - Procuração — N°2.pdf', etc.
        SEM número de contrato no nome.

    Convenção BA (compatibilidade):
        '2 - PROCURAÇÃO BANCO 0123456789.pdf'
        COM número de contrato no nome.

    Returns:
        Lista de dicts: [{ 'arquivo': caminho, 'numero_contrato': str | None }]
    """
    out = []
    for arq in os.listdir(pasta_cliente):
        if not arq.lower().endswith('.pdf'):
            continue
        if 'procura' not in arq.lower():
            continue
        # Tentar extrair número de contrato (5+ dígitos seguidos)
        m = re.search(r'\b(\d{5,})\b', arq)
        numero = m.group(1) if m else None
        out.append({'arquivo': os.path.join(pasta_cliente, arq), 'numero_contrato': numero})
    return out


# ============================================================
#  CLASSIFICAÇÃO DO CENÁRIO (qual bloco fático usar)
# ============================================================

# SIMPLIFICADO (Gabriel, 07/05/2026 — Opção B): templates AL têm 1 bloco
# fático genérico. A skill apenas preenche os placeholders xxxxxxxx, xxx,xx,
# etc. Sem variações por cenário (ATIVO/INATIVO × COM/SEM DEPÓSITO).

# Mantemos CENARIOS_BLOCO/TODAS_ANCORAS_BLOCO vazios para compatibilidade.
CENARIOS_BLOCO = {}
TODAS_ANCORAS_BLOCO: List[str] = []


def classificar_cenario(contratos: List[Dict], assume_com_deposito: bool = False) -> Dict:
    """Versão SIMPLIFICADA (Opção B, 07/05/2026): retorna apenas metadados
    informativos e alertas. Não há mais "bloco_ancora" — o template tem
    apenas 1 bloco fático genérico que será preenchido por _preencher_bloco_fatico.
    """
    alertas = []
    n = len(contratos)
    tipos = {c.get('tipo_origem') for c in contratos}
    situacoes = {c.get('situacao', '').lower() for c in contratos}

    if 'refinanciamento' in tipos and tipos != {'refinanciamento'}:
        alertas.append(
            f'⚠ Mistura de contratos AVN + REFIN ({tipos}). O bloco fático '
            f'genérico fala em "empréstimo" — revisar se cabe complementar.'
        )
    if not (situacoes <= {'ativo'} or situacoes <= {'excluído', 'encerrado', 'inativo'}):
        alertas.append(
            f'⚠ Mistura de contratos ATIVOS + INATIVOS ({situacoes}). '
            f'Revisar narração manualmente.'
        )
    if n >= 3:
        alertas.append(
            f'🚨 {n} contratos — bloco fático genérico lista todos. CONFERIR '
            f'a redação resultante (3+ contratos é cenário raro em AL).'
        )

    return {
        'n_contratos': n,
        'tipos': sorted(t for t in tipos if t),
        'situacoes': sorted(situacoes),
        'com_deposito': assume_com_deposito,
        'alertas': alertas,
    }


# ============================================================
#  PIPELINE PRINCIPAL — montagem dos dados
# ============================================================

def montar_dados_inicial_al(pasta_cliente: str, autora: Dict, comarca: str,
                              forcar_foro: Optional[str] = None,
                              forcar_procurador: str = 'tiago',
                              assume_com_deposito: bool = False,
                              numeros_contrato_explicitos: Optional[List[str]] = None,
                              uf_override: Optional[str] = None) -> Dict:
    """Monta o dicionário completo para gerar a inicial AL.

    Args:
        pasta_cliente: pasta com HISCON, procurações, RG, etc.
        autora: dict com qualificação da autora (nome, CPF, RG, endereço, etc.)
        comarca: nome da comarca AL (ex.: 'Arapiraca', 'Maceió')
        forcar_foro: 'federal' | 'estadual' | None (auto pelo valor da causa)
        forcar_procurador: 'tiago' (default) | 'alexandre' (em transição)
        assume_com_deposito: False por padrão. SÓ usar True quando HISCRE +
            extrato bancário confirmarem expressamente o crédito em conta.
            Sem confirmação, vai bloco SEM DEPÓSITO (mais conservador — não
            podemos afirmar que ela recebeu).
        numeros_contrato_explicitos: lista de números de contrato a filtrar
            do HISCON. Quando informado, sobrescreve a detecção via nome de
            arquivo da procuração. Use quando a procuração NÃO traz o número
            no nome do arquivo (ex.: '2 - Procuração — N°1.pdf') mas sim no
            CONTEÚDO do PDF (extrair via OCR antes de chamar este pipeline).

    Returns:
        dict de dados para gerar a inicial.
    """
    alertas = []

    # 1. HISCON
    hiscon_path = (encontrar_pdf(pasta_cliente, 'histórico de empréstimo')
                   or encontrar_pdf(pasta_cliente, 'historico de emprestimo')
                   or encontrar_pdf(pasta_cliente, 'extrato_emprestimo')
                   or encontrar_pdf(pasta_cliente, 'extrato emprestimo'))
    if not hiscon_path:
        raise FileNotFoundError(f'HISCON não encontrado em {pasta_cliente}')
    hiscon = parse_hiscon(hiscon_path)

    # 2. HISCRE (opcional)
    hiscre_path = (encontrar_pdf(pasta_cliente, 'histórico de crédito')
                   or encontrar_pdf(pasta_cliente, 'historico de credito'))
    hiscre = parse_hiscre(hiscre_path) if hiscre_path else {}
    if not hiscre:
        alertas.append('⚠ HISCRE não encontrado. CPF e renda podem estar incompletos.')

    # 3. Cálculo Jurídico (opcional)
    calc_path = (encontrar_pdf(pasta_cliente, 'cálculo')
                 or encontrar_pdf(pasta_cliente, 'calculo'))
    calculo = parse_calculo(calc_path) if calc_path else {
        'valor_total_geral': None, 'dano_moral_pleiteado_pdf': None,
        'idade': None, 'data_nascimento': None,
    }

    # 4. Procurações — extrair números de contrato VIA OCR do conteúdo
    # (regra OBRIGATÓRIA gravada na SKILL.md §9-quater)
    procuracoes = listar_procuracoes(pasta_cliente)
    nums_proc_arquivo = [p['numero_contrato'] for p in procuracoes if p['numero_contrato']]

    # PRIORIDADE 1: chamador passou explicitamente
    if numeros_contrato_explicitos:
        nums_filtro = list(numeros_contrato_explicitos)
        alertas.append(
            f'ℹ️ Usando {len(nums_filtro)} números de contrato passados '
            f'explicitamente: {nums_filtro}. Procurações na pasta: '
            f'{len(procuracoes)}.'
        )
    # PRIORIDADE 2: nome do arquivo já trazia número
    elif nums_proc_arquivo:
        nums_filtro = nums_proc_arquivo
        alertas.append(
            f'ℹ️ Filtro pelos números do nome do arquivo das '
            f'procurações: {nums_filtro}.'
        )
    # PRIORIDADE 3: ler conteúdo da(s) procuração(ões) via text-layer/OCR
    else:
        info_proc = extrair_numeros_contrato_de_pasta(pasta_cliente, usar_easyocr=True)
        nums_filtro = info_proc['numeros_unicos']
        alertas.extend(info_proc['alertas'])
        if nums_filtro:
            alertas.append(
                f'ℹ️ Filtro extraído via leitura do CONTEÚDO da procuração '
                f'({len(nums_filtro)} número(s)): {nums_filtro}. CONFERIR.'
            )

    # REGRA CRÍTICA (SKILL.md §9-quater): se mesmo após todas as tentativas
    # o filtro de contratos ficou vazio, ABORTAR. A procuração é a única
    # fonte autoritativa — NUNCA pegar "todos os contratos do banco" como
    # fallback silencioso.
    if not nums_filtro:
        raise ProcuracaoSemFiltroError(
            f'🚨 IMPOSSÍVEL extrair números de contrato das procurações em '
            f'{pasta_cliente}. Nenhum número no nome do arquivo, e a leitura '
            f'do conteúdo (text-layer + OCR) também falhou. AÇÃO: abrir o(s) '
            f'PDF(s) das procurações manualmente, ler o número do contrato '
            f'outorgado, e chamar novamente passando '
            f'`numeros_contrato_explicitos=[...]` para `montar_dados_inicial_al`. '
            f'NUNCA pegamos "todos os contratos do banco" como fallback.'
        )

    # 5. Filtrar contratos do HISCON
    contratos_brutos = filtrar_contratos_por_numero(
        hiscon['contratos'], nums_filtro, fuzzy_dist=1)
    if not contratos_brutos:
        raise ProcuracaoSemFiltroError(
            f'🚨 Nenhum contrato do HISCON casou com os números das '
            f'procurações {nums_filtro}. Os números podem estar errados, '
            f'ou os contratos podem não estar no HISCON do INSS deste '
            f'cliente. CONFERIR antes de prosseguir.'
        )

    if not contratos_brutos:
        raise RuntimeError('Nenhum contrato encontrado para a inicial AL.')

    contratos_fmt = [formatar_contrato_para_template(c) for c in contratos_brutos]

    # 6. Identificar banco-réu (do PRIMEIRO contrato)
    banco_nome_hiscon = contratos_brutos[0].get('banco_nome', '')
    banco_reu = (resolver_banco(banco_nome_hiscon, 'AL')
                 or resolver_banco(banco_nome_hiscon, 'matriz'))
    if not banco_reu:
        raise RuntimeError(f'Banco réu não identificado: "{banco_nome_hiscon}"')

    # 6-bis. Detectar 2º banco se houver
    bancos_unicos_codigos = sorted({c.get('banco_codigo') for c in contratos_brutos
                                     if c.get('banco_codigo')})
    n_bancos = len(bancos_unicos_codigos)
    bancos_reus = [banco_reu]
    if n_bancos >= 2:
        # Pegar o nome do 2º banco
        for c in contratos_brutos:
            if c.get('banco_codigo') == bancos_unicos_codigos[1]:
                banco2 = (resolver_banco(c.get('banco_nome'), 'AL')
                          or resolver_banco(c.get('banco_nome'), 'matriz'))
                if banco2:
                    bancos_reus.append(banco2)
                break
    if n_bancos >= 3:
        alertas.append(
            f'🚨 {n_bancos} bancos diferentes no HISCON. Template AL atende '
            f'no máximo 2 bancos no polo passivo. Os contratos do(s) banco(s) '
            f'extras DEVEM ser AJUSTADOS MANUALMENTE.'
        )

    # 7. AUDITAR procurações vs HISCON
    audit_proc = None
    if nums_filtro:
        audit_proc = auditar_procuracoes_vs_hiscon(
            hiscon['contratos'], nums_filtro,
            contratos_brutos[0].get('banco_codigo', ''))

    # 8. Verificação cruzada doc vs HISCRE
    divergencias = comparar_doc_vs_hiscre(autora, hiscre) if hiscre else []
    autora_consolidada = consolidar_dados_autora(autora, hiscre) if hiscre else dict(autora)

    # 9. Dano moral (regra: 1 contrato = R$ 15k; 2+ = R$ 5k × N)
    dm = calcular_dano_moral(contratos_brutos)
    audit_dm = (auditar_dano_moral(contratos_brutos, calculo['dano_moral_pleiteado_pdf'])
                if calculo.get('dano_moral_pleiteado_pdf') else
                {'divergencia': False, 'alerta': None})

    # 10. Valor da causa
    if calculo.get('valor_total_geral'):
        vc = calculo['valor_total_geral']
        fonte_vc = 'PDF de cálculo'
    else:
        soma_dobros = sum(
            (c.get('valor_parcela_float', 0) or 0)
            * (c.get('qtd_parcelas', 0) or 0) * 2
            for c in contratos_fmt
        )
        vc = soma_dobros + dm['total']
        fonte_vc = (f'estimado: soma dos dobros (R$ {soma_dobros:,.2f}) + dano '
                    f'moral (R$ {dm["total"]:,.2f}) — sem PDF de cálculo')

    # 11. Decidir foro
    decisao_foro = decidir_foro_al(vc, forcar=forcar_foro)
    foro = decisao_foro['foro']

    # 12. Selecionar template
    # uf_override permite reusar este pipeline para outras UFs com templates
    # estruturalmente similares (ex.: MG_ESTADUAL usa 'inicial-jemg-1banco.docx').
    uf_para_template = uf_override or 'AL'
    template_nome = selecionar_template_por_uf(
        uf_para_template, '1contrato', foro=foro, n_bancos=n_bancos)
    template_path = os.path.join(VAULT_TEMPLATES, template_nome)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'Template não encontrado: {template_path}')

    # 13. Procurador
    procurador = selecionar_procurador('AL', override_chave=forcar_procurador)
    if not procurador:
        raise RuntimeError(f'Procurador AL não cadastrado: {forcar_procurador}')

    # 14. Idade
    eh_id = False
    if hiscre.get('data_nascimento'):
        d = hiscre['data_nascimento']
        idade = (datetime.now() - d).days // 365
        eh_id = idade >= 60
    elif autora.get('data_nascimento'):
        d = autora['data_nascimento']
        idade = (datetime.now() - d).days // 365
        eh_id = idade >= 60

    # 15. Cenário do bloco fático
    cenario = classificar_cenario(contratos_fmt, assume_com_deposito=assume_com_deposito)
    alertas.extend(cenario['alertas'])

    # 16. Renda líquida
    renda_liquida = (hiscre.get('valor_liquido')
                     or autora.get('renda_liquida')
                     or hiscon['margens'].get('base_calculo'))
    fonte_renda = ('HISCRE (líquido)' if hiscre.get('valor_liquido')
                   else 'AUTORA' if autora.get('renda_liquida')
                   else 'HISCON BASE_CÁLCULO (BRUTO — confirmar)')

    return {
        'pasta_cliente': pasta_cliente,
        'hiscon': hiscon,
        'hiscre': hiscre,
        'calculo': calculo,
        'autora_consolidada': autora_consolidada,
        'divergencias_pessoais': divergencias,
        'procuracoes': procuracoes,
        'contratos_brutos': contratos_brutos,
        'contratos_questionados': contratos_fmt,
        'banco_reu': banco_reu,
        'bancos_reus': bancos_reus,
        'n_bancos': n_bancos,
        'audit_procuracoes': audit_proc,
        'template': template_path,
        'comarca': comarca,
        'foro': foro,
        'decisao_foro': decisao_foro,
        'procurador': procurador,
        'dano_moral': dm,
        'audit_dm': audit_dm,
        'valor_causa': vc,
        'fonte_vc': fonte_vc,
        'eh_idoso': eh_id,
        'cenario': cenario,
        'renda_liquida': renda_liquida,
        'fonte_renda': fonte_renda,
        'alertas': alertas,
    }


# ============================================================
#  GERAÇÃO DO DOCX
# ============================================================

def _fmt_brl(v: float) -> str:
    """Formata float para 'X.XXX,XX' (padrão BR)."""
    if v is None:
        return ''
    return f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')


def _fmt_data_extenso(d: datetime) -> str:
    """01/05/2026 → '1 de maio de 2026'"""
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    return f'{d.day} de {meses[d.month - 1]} de {d.year}'


# (funções _montar_qualificacao e _montar_polo_passivo foram substituídas
# por _substituir_qualificacao e _substituir_polo_passivo, que constroem os
# runs diretamente preservando Segoe UI Bold para nome do autor e dos réus.)


def _preencher_bloco_fatico(doc, contratos_fmt: List[Dict], cenario: Dict,
                              hiscre: Dict = None, nome_banco: str = '') -> int:
    """Substitui os placeholders xxxxxxxx, xxx,xx, xx parcelas, xx/xxxx
    DENTRO do bloco fático selecionado pelos valores reais dos contratos.

    A estratégia: usar matching com CONTEXTO único para diferenciar a 1ª
    ocorrência de "R$ xxx,xx (valor por extenso)" (parcela) da 2ª (valor
    emprestado), já que `substituir_in_run` substitui TODAS as ocorrências.

    Args:
        doc: Document do python-docx
        contratos_fmt: lista de contratos formatados
        cenario: dict do classificar_cenario (com 'bloco_ancora', 'n_contratos')
        hiscre: dict do parse_hiscre (opcional)
        nome_banco: string com nome do banco-réu para substituir "banco xxxxx"

    Returns:
        Número de substituições feitas.
    """
    n = cenario['n_contratos']
    feitos = 0

    # Template SIMPLIFICADO (Opção B): identificar o ÚNICO bloco fático
    # genérico, que começa com "No que diz respeito ao referido empréstimo"
    # (singular) e termina antes do "Sabe-se que tal fato ocorre".
    pars = list(doc.paragraphs)
    idx_inicio = None
    idx_fim = None
    for ip, p in enumerate(pars):
        if 'No que diz respeito ao referido empréstimo' in p.text:
            idx_inicio = ip
            break
        if 'No que diz respeito aos referidos empréstimos' in p.text:
            idx_inicio = ip
            break
    if idx_inicio is None:
        return 0
    for ip in range(idx_inicio + 1, len(pars)):
        if 'Sabe-se que tal fato ocorre' in pars[ip].text:
            idx_fim = ip
            break
    if idx_fim is None:
        idx_fim = idx_inicio + 10  # fallback

    # Substituir em ordem dentro de [idx_inicio, idx_fim)
    # Construir fila de substituições conforme cenário
    substituicoes = []
    for c in contratos_fmt[:max(2, n)]:  # limita ao N do bloco
        # Para cada contrato:
        # primeiro xxxxxxxx do parágrafo  → numero
        # segundo  xxxxxxxx do parágrafo  → competência inicial
        # xx parcelas                     → qtd parcelas
        # primeiro  xxx,xx                → valor parcela
        # segundo   xxx,xx                → valor emprestado
        # xx/xxxx                         → competência fim (só inativos)
        substituicoes.append({
            'numero': c.get('numero', ''),
            'competencia_inicio': c.get('competencia_inicio_str', ''),
            'qtd_parcelas': str(c.get('qtd_parcelas', '')),
            'valor_parcela': c.get('valor_parcela_str', ''),
            'valor_emprestado': c.get('valor_emprestado_str', ''),
            'competencia_fim': c.get('competencia_fim_str', ''),
            # Nome real do banco do contrato (litisconsórcio passivo) —
            # quando há múltiplos bancos, cada bloco fático identifica
            # o banco específico daquele contrato.
            # Após formatar_contrato_para_template, o campo vira 'banco'.
            'banco_nome': c.get('banco_nome') or c.get('banco') or '',
        })

    # Função auxiliar: tenta construir extenso de moeda
    try:
        from extenso import extenso_moeda as _ext
    except Exception:
        _ext = lambda v: ''

    # Para cada contrato, vamos substituir as ocorrências em sequência.
    # Estratégia: iterar SOMENTE pelos parágrafos do bloco, e para cada
    # parágrafo substituir os placeholders na ORDEM em que aparecem.
    # Para 1 contrato: tudo num parágrafo só.
    # Para 2 contratos: cada contrato em uma linha "Do contrato sob n°...".

    if n == 1 and substituicoes:
        s = substituicoes[0]
        for ip in range(idx_inicio, idx_fim):
            p = pars[ip]
            txt = p.text
            if 'xxxxxxxx' not in txt and 'xxx,xx' not in txt:
                continue
            # Substituições gerais dentro do parágrafo (ordem importa)
            try:
                vp = float(s['valor_parcela'].replace('.', '').replace(',', '.'))
                vp_ext = _ext(vp)
            except Exception:
                vp_ext = ''
            try:
                ve = float(s['valor_emprestado'].replace('.', '').replace(',', '.'))
                ve_ext = _ext(ve)
            except Exception:
                ve_ext = ''
            # IMPORTANTE: substituir_in_run substitui TODAS as ocorrências,
            # então usamos CONTEXTO único pra diferenciar parcela vs empréstimo.
            sequencia = [
                # (1) "no valor de R$ xxx,xx (valor por extenso), relativas"
                #     → valor da PARCELA
                ('no valor de R$ xxx,xx (valor por extenso), relativas',
                 f'no valor de R$ {s["valor_parcela"]} ({vp_ext}), relativas'),
                # (2) "empréstimo consignado no valor de R$ xxx,xx (valor por extenso)"
                #     → valor EMPRESTADO
                ('empréstimo consignado no valor de R$ xxx,xx (valor por extenso)',
                 f'empréstimo consignado no valor de R$ {s["valor_emprestado"]} ({ve_ext})'),
                # (3) competência inicial
                ('competência xxxxxxxx', f'competência {s["competencia_inicio"]}'),
                # (4) qtd parcelas
                ('xx parcelas', f'{s["qtd_parcelas"]} parcelas'),
                # (5) competência fim (apenas em blocos INATIVOS)
                ('encerradas em xx/xxxx', f'encerradas em {s["competencia_fim"]}'),
                # (6) número do contrato
                ('contrato n° xxxxxxx', f'contrato n° {s["numero"]}'),
                # (7) nome do banco — preencher para evitar "pelo , ora requerido"
                ('cuja operação foi realizada pelo banco xxxxx',
                 f'cuja operação foi realizada pelo {nome_banco}'),
            ]
            for old, new in sequencia:
                if old in p.text:
                    substituir_in_run(p._element, {old: new}, grifo=True)
                    feitos += 1
            # Substitui depósito (xxxxxxxx → "[DATA DO DEPÓSITO — A CONFIRMAR no extrato]")
            if 'verificou o depósito do montante, no dia xxxxxxxx' in p.text:
                substituir_in_run(p._element,
                    {'no dia xxxxxxxx': 'no dia [A CONFIRMAR no extrato bancário]'},
                    grifo=True)
                feitos += 1

    elif n >= 2 and len(substituicoes) >= 2:
        # CAMINHO A: template traz sub-blocos "Do contrato sob n° xxxxxxxx..."
        # (formato 2bancos) — usa o iterador clássico de substituições.
        # CAMINHO B: template traz APENAS o bloco singular "No que diz respeito
        # ao referido empréstimo..." (formato 1banco com N contratos do mesmo
        # banco — caso refin/refin/refin). Aí precisamos DUPLICAR o parágrafo
        # singular N vezes (uma cópia por contrato) e preencher cada cópia.
        tem_sub_blocos = any(
            ('Do contrato sob n°' in pars[ip].text and 'xxxxxxxx' in pars[ip].text)
            for ip in range(idx_inicio, idx_fim)
        )
        if not tem_sub_blocos:
            # === CAMINHO B: 1banco × N contratos — duplicar bloco singular ===
            p_template = None
            for ip in range(idx_inicio, idx_fim):
                t = pars[ip].text
                if 'No que diz respeito ao referido empréstimo' in t and 'xxxxxxxx' in t:
                    p_template = pars[ip]
                    break
                if 'No que diz respeito aos referidos empréstimos' in t and 'xxxxxxxx' in t:
                    p_template = pars[ip]
                    break
            if p_template is None:
                return feitos
            elem_template = p_template._element
            # Cria (n-1) cópias logo após o original (mantém o original como 1ª)
            elementos = [elem_template]
            last = elem_template
            for _ in range(n - 1):
                novo = deepcopy(elem_template)
                last.addnext(novo)
                elementos.append(novo)
                last = novo
            # Para cada cópia, preencher com os dados do contrato correspondente.
            # Para múltiplos contratos, prefixar com "a) ", "b) ", "c) "...
            # (até "z)" — se for mais de 26, rola para "aa)" mas não esperamos isso).
            def _letra(idx: int) -> str:
                if idx < 26:
                    return chr(ord('a') + idx)
                # fallback double-letter ('aa', 'ab', ...)
                return chr(ord('a') + idx // 26 - 1) + chr(ord('a') + idx % 26)
            # Remover numeração automática (numPr) dos parágrafos do bloco
            # fático duplicado — senão sai "5. a) ... 6. b) ... 7. c) ...".
            # Queremos só "a) ... b) ... c) ..." sem numeração lateral.
            _W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            for elem in elementos:
                pPr = elem.find(f'{_W}pPr')
                if pPr is None:
                    continue
                numPr = pPr.find(f'{_W}numPr')
                if numPr is not None:
                    pPr.remove(numPr)
                # Também limpa pStyle de lista (que faz indentação automática)
                pStyle = pPr.find(f'{_W}pStyle')
                if pStyle is not None and 'Lista' in (pStyle.get(f'{_W}val', '') or ''):
                    pPr.remove(pStyle)
            # Detectar múltiplos bancos para decidir o nome a usar no bloco
            bancos_distintos = {s.get('banco_nome', '').strip() for s in substituicoes[:n]
                                if s.get('banco_nome')}
            multi_banco = len(bancos_distintos) >= 2
            for idx, (elem, s) in enumerate(zip(elementos, substituicoes[:n])):
                # 1) Prefixo "a) " / "b) " / ... no início do parágrafo (antes de "No que diz")
                prefixo = f'{_letra(idx)}) '
                substituir_in_run(elem,
                    {'No que diz respeito ao referido empréstimo':
                     f'{prefixo}No que diz respeito ao referido empréstimo',
                     'No que diz respeito aos referidos empréstimos':
                     f'{prefixo}No que diz respeito aos referidos empréstimos'},
                    grifo=True)
                # Quando há múltiplos bancos, substituir "banco xxxxx" pelo
                # nome REAL do banco deste contrato — ANTES da sequência padrão
                # rodar (para não cair no banco-principal).
                if multi_banco and s.get('banco_nome'):
                    nome_banco_contrato = s['banco_nome'].upper()
                    substituir_in_run(elem, {
                        'cuja operação foi realizada pelo banco xxxxx':
                        f'cuja operação foi realizada pelo {nome_banco_contrato}',
                    }, grifo=True)
                try:
                    vp = float(s['valor_parcela'].replace('.', '').replace(',', '.'))
                    vp_ext = _ext(vp)
                except Exception:
                    vp_ext = ''
                try:
                    ve = float(s['valor_emprestado'].replace('.', '').replace(',', '.'))
                    ve_ext = _ext(ve)
                except Exception:
                    ve_ext = ''
                sequencia = [
                    ('no valor de R$ xxx,xx (valor por extenso), relativas',
                     f'no valor de R$ {s["valor_parcela"]} ({vp_ext}), relativas'),
                    ('empréstimo consignado no valor de R$ xxx,xx (valor por extenso)',
                     f'empréstimo consignado no valor de R$ {s["valor_emprestado"]} ({ve_ext})'),
                    ('competência xxxxxxxx', f'competência {s["competencia_inicio"]}'),
                    ('xx parcelas', f'{s["qtd_parcelas"]} parcelas'),
                    ('encerradas em xx/xxxx', f'encerradas em {s["competencia_fim"]}'),
                    ('contrato n° xxxxxxx', f'contrato n° {s["numero"]}'),
                    ('cuja operação foi realizada pelo banco xxxxx',
                     f'cuja operação foi realizada pelo {nome_banco}'),
                ]
                # Lê o texto agregado dos w:t descendentes para checar match
                _ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                texto_elem = ''.join(t.text or '' for t in elem.iter(f'{_ns}t'))
                for old, new in sequencia:
                    if old in texto_elem:
                        substituir_in_run(elem, {old: new}, grifo=True)
                        feitos += 1
                        texto_elem = texto_elem.replace(old, new, 1)
            return feitos
        # === CAMINHO A: template tem sub-blocos "Do contrato sob n°..." ===
        contratos_iter = iter(substituicoes)
        depositos_iter = iter(substituicoes)  # para datas de depósito
        for ip in range(idx_inicio, idx_fim):
            p = pars[ip]
            txt = p.text
            if 'Do contrato sob n°' in txt and 'xxxxxxxx' in txt:
                try:
                    s = next(contratos_iter)
                except StopIteration:
                    continue
                try:
                    vp = float(s['valor_parcela'].replace('.', '').replace(',', '.'))
                    vp_ext = _ext(vp)
                except Exception:
                    vp_ext = ''
                try:
                    ve = float(s['valor_emprestado'].replace('.', '').replace(',', '.'))
                    ve_ext = _ext(ve)
                except Exception:
                    ve_ext = ''
                # Mesma proteção contra "substituir TODAS as ocorrências":
                # diferenciar parcela vs empréstimo via contexto único.
                sequencia = [
                    ('Do contrato sob n° xxxxxxxx', f'Do contrato sob n° {s["numero"]}'),
                    ('competência xxxxxxxx', f'competência {s["competencia_inicio"]}'),
                    ('xx parcelas', f'{s["qtd_parcelas"]} parcelas'),
                    ('no valor de R$ xxx,xx (valor por extenso), relativas',
                     f'no valor de R$ {s["valor_parcela"]} ({vp_ext}), relativas'),
                    ('empréstimo consignado no valor de R$ xxx,xx (valor por extenso)',
                     f'empréstimo consignado no valor de R$ {s["valor_emprestado"]} ({ve_ext})'),
                    ('última parcela descontada do benefício da parte autora em xxxxxx',
                     f'última parcela descontada do benefício da parte autora em {s["competencia_fim"]}'),
                ]
                for old, new in sequencia:
                    if old in p.text:
                        substituir_in_run(p._element, {old: new}, grifo=True)
                        feitos += 1
            elif 'No dia xxxxxxxx, o valor de R$ xxx,xx' in txt:
                # Linhas de depósito por contrato
                try:
                    s = next(depositos_iter)
                except StopIteration:
                    continue
                substituir_in_run(p._element, {
                    'No dia xxxxxxxx': 'No dia [A CONFIRMAR]',
                    'R$ xxx,xx': f'R$ {s["valor_emprestado"]}',
                }, grifo=True)
                feitos += 1

    return feitos


# ============================================================
# REDAÇÃO: usa helpers_redacao (compartilhados com BA e AM).
# Aqui só um wrapper de compat para `_substituir_qualificacao`
# (a versão comum recebe end_escritorio direto; o pipeline AL passa
# o procurador inteiro, então extraímos o endereço aqui).
# ============================================================

def _substituir_qualificacao(p_elem, autora: Dict, procurador: Dict, grifo: bool = True):
    """Wrapper local: monta endereço composto (MATRIZ SC + unidade de apoio
    na UF do cliente) e delega ao helper.
    """
    from helpers_redacao import substituir_qualificacao_autor
    from escritorios import montar_endereco_escritorio_completo
    uf = autora.get('uf', 'AL')
    end = montar_endereco_escritorio_completo(uf)
    substituir_qualificacao_autor(p_elem, autora, end, grifo=grifo, uf_default=uf)


def _identificar_blocos_faticos(doc) -> Dict[str, Tuple[int, int]]:
    """Para cada âncora de bloco fático, identifica (idx_inicio, idx_fim)
    no doc.paragraphs (incluindo todas as linhas desde a âncora até o
    parágrafo ANTES da próxima âncora ou da seção 'Sabe-se que tal fato').

    Returns:
        {ancora: (idx_start, idx_end_exclusive)}
    """
    indices = {}
    pars = list(doc.paragraphs)

    # Encontrar posição inicial de cada âncora
    for ip, p in enumerate(pars):
        for anc in TODAS_ANCORAS_BLOCO:
            if anc in p.text:
                indices[anc] = ip
                break

    # Determinar fim de cada bloco (próxima âncora OU "Sabe-se que tal fato")
    sentinela_fim_idx = None
    for ip, p in enumerate(pars):
        if 'Sabe-se que tal fato ocorre' in p.text:
            sentinela_fim_idx = ip
            break

    ordered = sorted(indices.items(), key=lambda kv: kv[1])
    out = {}
    for i, (anc, idx_start) in enumerate(ordered):
        if i + 1 < len(ordered):
            idx_end = ordered[i + 1][1]
        else:
            idx_end = sentinela_fim_idx if sentinela_fim_idx else len(pars)
        out[anc] = (idx_start, idx_end)
    return out


def _remover_paragrafos_intervalo(doc, idx_start: int, idx_end_exclusive: int):
    """Remove os parágrafos no intervalo [idx_start, idx_end_exclusive)."""
    pars = list(doc.paragraphs)
    for ip in range(idx_end_exclusive - 1, idx_start - 1, -1):
        if ip >= len(pars):
            continue
        p = pars[ip]
        p._element.getparent().remove(p._element)


def gerar_inicial_al(dados_caso: Dict, output_path: str) -> Dict:
    """Gera o DOCX da inicial AL aplicando substituições targeted no template.

    Etapas:
      1. Copiar template para output_path
      2. Cabeçalho (comarca + foro)
      3. Prioridade idoso (remover se não-idoso)
      4. Qualificação do autor (parágrafo inteiro)
      5. Polo passivo (banco(s) + INSS se Federal)
      6. Síntese fática (NB, agência, conta, banco_pagador)
      7. Lista de contratos no preâmbulo
      8. Selecionar bloco fático apropriado, REMOVER os outros 11
      9. Justiça gratuita (renda)
      10. Quantum dano moral (5k × N para 2+, 15k para 1)
      11. Pedidos finais (valor causa)
      12. Rodapé (cláusula intimação + assinatura)
    """
    template = dados_caso['template']
    autora = dados_caso['autora_consolidada']
    bancos_reus = dados_caso['bancos_reus']
    banco_reu = dados_caso['banco_reu']  # principal (1º) — para o bloco fático
    foro = dados_caso['foro']
    inclui_inss = (foro == 'federal')
    procurador = dados_caso['procurador']
    comarca = dados_caso['comarca']
    cenario = dados_caso['cenario']
    contratos_fmt = dados_caso['contratos_questionados']
    hiscon_cab = dados_caso['hiscon']['cabecalho']
    valor_causa = dados_caso['valor_causa']
    dm = dados_caso['dano_moral']
    eh_idoso = dados_caso['eh_idoso']
    renda_liquida = dados_caso['renda_liquida']

    shutil.copy(template, output_path)
    doc = Document(output_path)
    modificados = 0

    # ---- (1) Cabeçalho com comarca (preserva Segoe UI Bold do template) ----
    # Aceita qualquer UF (AL/MG/PE/SE/...) — o regex captura `/{UF}` no fim e
    # substitui pelo `{comarca}/{uf_da_autora}`.
    uf_autora = (autora.get('uf') or 'AL').upper()
    for p in doc.paragraphs[:5]:
        t = p.text
        if 'Subseção de' in t and re.search(r'/[A-Z]{2}', t):
            # Federal — pStyle 1Pargrafo + Segoe UI Bold
            novo = re.sub(r'Subseção de\s*\S*\s*/[A-Z]{2}',
                            f'Subseção de {comarca}/{uf_autora}', t)
            if novo != t:
                _substituir_paragrafo_completo(p._element, novo, grifo=True,
                                                fonte='Segoe UI', bold=True)
                modificados += 1
            break
        if 'Vara Cível da Comarca de' in t:
            # Estadual — mesmo padrão
            novo = re.sub(r'Comarca de\s*\S*\s*/[A-Z]{2}',
                            f'Comarca de {comarca}/{uf_autora}', t)
            if novo != t:
                _substituir_paragrafo_completo(p._element, novo, grifo=True,
                                                fonte='Segoe UI', bold=True)
                modificados += 1
            break

    # ---- (2) Prioridade idoso (remover se não-idoso) ----
    if not eh_idoso:
        for p in list(doc.paragraphs[:5]):
            if 'Prioridade de tramitação: art. 1.048' in p.text:
                p._element.getparent().remove(p._element)
                modificados += 1

    # ---- (3) Qualificação do autor (NOME em Segoe UI Bold + resto Cambria) ----
    # Detecta o parágrafo de qualificação de forma genérica: começa com NOME
    # EM CAPS (3+ palavras) seguido de ", brasileiro/brasileira,". Suporta
    # qualquer cliente piloto, não só FULANO/EDMUNDA/LOURDES.
    padrao_quali = re.compile(r'^[A-ZÀ-Ú]{3,}(?:\s[A-ZÀ-Ú]{2,})+\s*,\s*brasileir[oa]')
    for p in doc.paragraphs[:15]:
        if padrao_quali.match(p.text.strip()):
            _substituir_qualificacao(p._element, autora, procurador, grifo=True)
            modificados += 1
            break

    # ---- (4) Polo passivo (NOMES em Segoe UI Bold + resto Cambria) ----
    for p in doc.paragraphs[:20]:
        if 'em face de' in p.text:
            _substituir_polo_passivo(p._element, bancos_reus,
                                       inclui_inss=inclui_inss, grifo=True)
            modificados += 1
            break

    # ---- (5) Síntese fática (NB, agência, conta/cartão, banco_pagador) ----
    # Helper escolhe automaticamente entre 'depositado em conta bancária' e
    # 'sacado por meio de cartão magnético' conforme `autora['forma_recebimento']`
    # ou conta_pagador vazio no HISCON. Regra fixa do escritório (07/05/2026).
    for p in doc.paragraphs[:30]:
        t = p.text
        if 'recebe benefício previdenciário' in t and 'NB' in t:
            novo = _montar_paragrafo_recebimento_beneficio(hiscon_cab, autora)
            _substituir_paragrafo_completo(p._element, novo, grifo=True)
            modificados += 1
            break

    # ---- (6) Intro "CONTRATO(S) Nº ..." — agrupa por banco quando há ≥2 bancos
    # (litisconsórcio passivo). Helper compartilhado com BA/AM. ----
    from helpers_redacao import aplicar_intro_fatica as _aplicar_intro_fatica
    for p in doc.paragraphs[:30]:
        t = p.text
        if (('CONTRATO Nº' in t) or ('CONTRATO N°' in t)
            or ('CONTRATOS Nº' in t) or ('CONTRATOS N°' in t)):
            _aplicar_intro_fatica(p._element, contratos_fmt,
                                    banco_reu['nome'], grifo=True)
            modificados += 1
            break

    # ---- (7) PREENCHER o bloco fático ----
    # Tentar PRIMEIRO o formato MG (cabeçalho "No que diz respeito" +
    # sub-itens "Do contrato nº A: ..."); se não houver sub-itens, cair no
    # formato AL (xxxxxxxx, xxx,xx, etc.).
    feitos_bloco_mg = _preencher_bloco_fatico_formato_mg(
        doc, contratos_fmt, banco_reu['nome'], grifo=True)
    if feitos_bloco_mg > 0:
        modificados += feitos_bloco_mg
    else:
        modificados += _preencher_bloco_fatico(doc, contratos_fmt, cenario,
                                                  hiscre=dados_caso.get('hiscre'),
                                                  nome_banco=banco_reu['nome'])

    # ---- (7-bis) Pedidos declaratórios (1 por contrato; "empréstimo" vs "refinanciamento") ----
    nb = dados_caso['hiscon']['cabecalho'].get('nb_beneficio') or '___'
    modificados += _preencher_pedidos_declaratorios(doc, contratos_fmt, nb, grifo=True)

    # ---- (7-ter) Prioridade idoso ----
    # Se autor é idoso e o template NÃO tem o parágrafo (cabeçalho OU pedido),
    # INSERIR. Se autor não é idoso, REMOVER do pedido (cabeçalho tratado em (2)).
    if eh_idoso:
        modificados += _inserir_prioridade_idoso_se_faltando(doc, eh_idoso, grifo=True)
        modificados += _inserir_pedido_prioridade_idoso_se_faltando(doc, eh_idoso, grifo=True)
    else:
        modificados += _remover_prioridade_pedidos(doc)

    # ---- (8) Justiça gratuita — renda líquida ----
    for p in doc.paragraphs:
        if 'no valor líquido de R$ ()' in p.text or 'valor líquido de R$' in p.text:
            renda_str = _fmt_brl(renda_liquida) if renda_liquida else '___'
            try:
                renda_ext = extenso_moeda(float(renda_liquida)) if renda_liquida else '___'
            except Exception:
                renda_ext = '___'
            substituir_in_run(p._element, {
                'R$ ()': f'R$ {renda_str} ({renda_ext})',
                'R$ ( )': f'R$ {renda_str} ({renda_ext})',
            }, grifo=True)
            modificados += 1
            break

    # ---- (9) Dano moral — pleito (R$ 15k para 1 contrato; R$ 5k×N para 2+) ----
    dm_total = dm['total']
    dm_str = _fmt_brl(dm_total)
    try:
        dm_ext = extenso_moeda(dm_total)
    except Exception:
        dm_ext = ''
    for p in doc.paragraphs:
        if 'R$ 15.000,00 (quinze mil reais)' in p.text and dm_total != 15000.0:
            substituir_in_run(p._element, {
                'R$ 15.000,00 (quinze mil reais)': f'R$ {dm_str} ({dm_ext})',
            }, grifo=True)
            modificados += 1

    # ---- (10) Valor da causa "Dá-se a causa o valor de R$ (centavos)." ----
    vc_str = _fmt_brl(valor_causa)
    try:
        vc_ext = extenso_moeda(valor_causa)
    except Exception:
        vc_ext = ''
    for p in doc.paragraphs:
        if 'Dá-se a causa o valor' in p.text:
            novo = f'Dá-se a causa o valor de R$ {vc_str} ({vc_ext}).'
            _substituir_paragrafo_completo(p._element, novo, grifo=True)
            modificados += 1
            break

    # ---- (11) Cidade + data ao final "Cidade/AL, 31 de março de 2026" ----
    for p in doc.paragraphs[-15:]:
        if re.match(r'^\s*\S+/AL,\s+\d', p.text) or 'Cidade/AL' in p.text:
            data_str = _fmt_data_extenso(datetime.now())
            novo = f'{comarca}/AL, {data_str}'
            _substituir_paragrafo_completo(p._element, novo, grifo=True)
            modificados += 1
            break

    # ---- (12) Cláusula intimação + assinatura ----
    # Trocar 'Tiago de Azevedo Lima' e 'OAB/AL 20906A' caso o procurador
    # forçado seja outro (Alexandre).
    if procurador['chave'] != 'tiago':
        for p in doc.paragraphs:
            if 'Tiago de Azevedo Lima' in p.text:
                substituir_in_run(p._element,
                    {'Tiago de Azevedo Lima': procurador['nome']}, grifo=True)
                modificados += 1
            if 'OAB/AL 20906A' in p.text:
                oab_uf = procurador.get('oab_uf', procurador.get('oab', ''))
                substituir_in_run(p._element,
                    {'OAB/AL 20906A': oab_uf}, grifo=True)
                substituir_in_run(p._element,
                    {'OAB/AL  sob o n.º 20906A': f'{oab_uf} sob o n.º {oab_uf.split()[-1]}'},
                    grifo=True)
                modificados += 1

    # ---- Salvar ----
    doc.save(output_path)

    # Verificar residuais (xxxxxxxx, xxx,xx, etc.) no bloco selecionado —
    # esses são os campos que o procurador precisa preencher manualmente.
    doc2 = Document(output_path)
    placeholders_residuais = []
    for p in doc2.paragraphs:
        for ph in re.findall(r'xxxxxxxx|xxx,xx|xx parcelas|xx/xxxx', p.text):
            placeholders_residuais.append(ph)

    return {
        'output': output_path,
        'modificados': modificados,
        'placeholders_para_preencher_no_bloco': len(placeholders_residuais),
        'placeholders_amostra': placeholders_residuais[:10],
    }
