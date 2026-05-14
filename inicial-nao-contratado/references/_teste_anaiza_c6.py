# -*- coding: utf-8 -*-
"""Teste de geração da inicial jeal-1banco para ANAIZA × Banco C6 APOSENTADORIA.
Usa dados do _estado_cliente.json + valores demo onde faltar."""
import sys, io, json, re, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
from docx import Document
from helpers_docx import substituir_in_run
from _blocos_narrativos import (
    gerar_bloco_contratos_fraudulentos,
    gerar_bloco_pedido_declaracao,
    normalizar_banco_reu,
)

PASTA_CLIENTE = Path(r"C:\Users\gabri\OneDrive\Área de Trabalho\APP - ORGANIZAÇÃO PASTA AL\TESTE - Fazer inicial\MARIA EXEMPLO DA SILVA")
TEMPLATE = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jeal-1banco.docx")
OUT = PASTA_CLIENTE / "APOSENTADORIA" / "BANCO C6 CONSIGNADO" / "INICIAL_TESTE.docx"

estado = json.loads((PASTA_CLIENTE / '_estado_cliente.json').read_text(encoding='utf-8'))
nome_cliente = estado['cliente']['nome_completo']

# Pegar 1º contrato C6/APOSENTADORIA (descartar duplicata)
contrato_dict = next(c for c in estado['contratos']
                      if c['banco_chave'] == 'C6' and c['beneficio_pasta'] == 'APOSENTADORIA')

def _strip_brl(v):
    """'R$ R$37,10' -> 37.10  ;  'R$1.178,55' -> 1178.55"""
    if v is None:
        return None
    s = re.sub(r'R\$\s*', '', str(v)).strip()
    s = s.replace('.', '').replace(',', '.')
    try:
        return float(s)
    except ValueError:
        return None

contratos_p_blocos = [{
    'numero': contrato_dict['contrato'],
    'valor_emprestado': _strip_brl(contrato_dict.get('valor_emprestado')),
    'valor_parcela': _strip_brl(contrato_dict.get('valor_parcela')),
    'qtd_parcelas': contrato_dict.get('qtd_parcelas'),
    'competencia_inicio': contrato_dict.get('competencia_inicio'),
    'data_inclusao': contrato_dict.get('data_inclusao'),
}]

# Benefício APOSENTADORIA
b = next(b for b in estado['beneficios_inss'] if 'APOSENTADORIA' in b['pasta_label'])

# Banco réu
banco_reu_nome_raw = contrato_dict['banco_nome_completo']  # 'BANCO C6 CONSIGNADO'
banco_reu_nome = normalizar_banco_reu(banco_reu_nome_raw)

# Dict de placeholders
dados = {
    # Cabeçalho
    '{{vara_protocolo}}':    '___',
    '{{cidade_protocolo}}':  'Arapiraca',
    '{{uf_protocolo}}':      'AL',
    # Autor (dados parciais — JSON não tem todos)
    '{{nome_autor}}':         nome_cliente,
    '{{nacionalidade}}':      'brasileira',
    '{{estado_civil}}':       '[ESTADO_CIVIL]',
    '{{profissao}}':          'aposentada',
    '{{cpf_autor}}':          '[CPF]',
    '{{rg_autor}}':           '[RG]',
    '{{orgao_expedidor}}':    'SSP/AL',
    '{{logradouro_autor}}':   '[LOGRADOURO]',
    '{{numero_autor}}':       '[Nº]',
    '{{bairro_autor}}':       '[BAIRRO]',
    '{{cidade_autor}}':       '[CIDADE]',
    '{{uf_autor}}':           'AL',
    '{{cep_autor}}':          '[CEP]',
    # Banco réu (forçado MAIÚSCULA)
    '{{banco_reu_nome}}':         banco_reu_nome,
    '{{banco_reu_descricao_pj}}': 'pessoa jurídica de direito privado',
    '{{banco_reu_cnpj}}':         '31.872.495/0001-72',
    '{{banco_reu_endereco}}':     'Av. Nove de Julho, 3186, Jardim Paulista, São Paulo/SP, CEP 01406-000',
    # Benefício
    '{{tipo_beneficio}}':   b['especie_nome'].lower(),
    '{{nb_beneficio}}':     b['nb'],
    '{{banco_pagador}}':    b['banco_pagador'],
    '{{agencia_pagador}}':  b['agencia_pagadora'],
    '{{conta_pagador}}':    b['conta_pagadora'],
    # Renda
    '{{valor_renda_liquida}}':         '[VALOR_RENDA]',
    '{{valor_renda_liquida_extenso}}': '[VALOR_RENDA_EXTENSO]',
    # Prioridade (assumindo idosa)
    '{{pedido_prioridade}}': 'A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de Processo Civil',
    # BLOCOS NARRATIVOS — chamada dos helpers
    '{{BLOCO_CONTRATOS_FRAUDULENTOS}}': gerar_bloco_contratos_fraudulentos(
        contratos_p_blocos, banco_reu_nome),
    '{{BLOCO_PEDIDO_DECLARACAO}}': gerar_bloco_pedido_declaracao(
        contratos_p_blocos, b['nb']),
}

# Copiar template e aplicar
shutil.copy2(TEMPLATE, OUT)
doc = Document(OUT)
modificados = 0
for p in doc.paragraphs:
    for k, v in dados.items():
        if k in p.text:
            if substituir_in_run(p._p, {k: v}):
                modificados += 1
doc.save(OUT)
print(f"OK — gerado: {OUT}")
print(f"Substituições aplicadas: {modificados}")
print()

# Verificar residuais
import re
d2 = Document(OUT)
phs_rest = set()
for p in d2.paragraphs:
    for m in re.finditer(r'\{\{[^}]+\}\}', p.text):
        phs_rest.add(m.group(0))
if phs_rest:
    print(f"⚠️ Placeholders NÃO substituídos: {phs_rest}")
else:
    print("✓ Todos placeholders substituídos")

# Mostrar parágrafos críticos
print()
print('=== PAR 6 (qualificação) ===')
print(d2.paragraphs[6].text[:400])
print()
print('=== PAR 10 (banco réu) ===')
print(d2.paragraphs[10].text)
print()
print('=== PAR 14 (constatação) ===')
print(d2.paragraphs[14].text)
print()
print('=== PAR 16 (BLOCO CONTRATOS) ===')
print(d2.paragraphs[16].text)
print()
print('=== PAR 228 (BLOCO PEDIDO) ===')
print(d2.paragraphs[228].text)
