# -*- coding: utf-8 -*-
"""Adiciona segundo banco réu no par 10 dos templates 2bancos.
Renomeia {{banco_reu_*}} → {{banco_reu_1_*}} e insere {{banco_reu_2_*}}.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run

VAULT = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates")

# === jeal-2bancos ===
sub_jeal = {
    "{{banco_reu_nome}}, {{banco_reu_descricao_pj}}, inscrita no CNPJ sob o nº {{banco_reu_cnpj}}, com endereço na {{banco_reu_endereco}}, pelos motivos":
        "{{banco_reu_1_nome}}, {{banco_reu_1_descricao_pj}}, inscrita no CNPJ sob o nº {{banco_reu_1_cnpj}}, com endereço na {{banco_reu_1_endereco}}, e em face de {{banco_reu_2_nome}}, {{banco_reu_2_descricao_pj}}, inscrita no CNPJ sob o nº {{banco_reu_2_cnpj}}, com endereço na {{banco_reu_2_endereco}}, pelos motivos"
}
caminho_jeal = VAULT / "inicial-jeal-2bancos.docx"
d = Document(caminho_jeal)
ok = substituir_in_run(d.paragraphs[10]._p, sub_jeal)
d.save(caminho_jeal)
print(f"jeal-2bancos par 10: {'OK' if ok else 'NAO bateu'}")

# === jfal-2bancos ===
sub_jfal = {
    "{{banco_reu_nome}}, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº. {{banco_reu_cnpj}}, com endereço na {{banco_reu_endereco}}, e INSTITUTO NACIONAL":
        "{{banco_reu_1_nome}}, {{banco_reu_1_descricao_pj}}, inscrita no CNPJ sob o nº. {{banco_reu_1_cnpj}}, com endereço na {{banco_reu_1_endereco}}, e em face de {{banco_reu_2_nome}}, {{banco_reu_2_descricao_pj}}, inscrita no CNPJ sob o nº. {{banco_reu_2_cnpj}}, com endereço na {{banco_reu_2_endereco}}, e INSTITUTO NACIONAL"
}
caminho_jfal = VAULT / "inicial-jfal-2bancos.docx"
d = Document(caminho_jfal)
ok = substituir_in_run(d.paragraphs[10]._p, sub_jfal)
d.save(caminho_jfal)
print(f"jfal-2bancos par 10: {'OK' if ok else 'NAO bateu'}")

# Verificar
print()
for nome in ["inicial-jeal-2bancos.docx", "inicial-jfal-2bancos.docx"]:
    d = Document(VAULT/nome)
    print(f"=== {nome} PAR 10 ===")
    print(d.paragraphs[10].text)
    print()
