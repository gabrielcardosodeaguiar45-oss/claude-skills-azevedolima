"""Pipeline para iniciais AM (Justiça Estadual rito comum, TJAM).

Diferenças vs pipeline BA:
- Polo passivo SÓ banco (sem INSS)
- Comarca = município do autor (Maués, Manaus, Boa Vista do Ramos, etc.)
- Procurador AM (Patrick OAB/AM A2638 ou Eduardo OAB/AM A2118)
- Naming diferente dos placeholders (ver adaptador_am.py)
- Suporte a representante legal (autor menor de idade)
- Templates: inicial-jeam-base.docx (1 contrato AVN/inativo) ou inicial-jeam-refin.docx (refin)

Uso:
    dados = montar_dados_inicial_am(
        pasta_banco='.../FABIO MARINHO/AGIBANK',
        autora=AUTORA_FABIO,
        comarca='Presidente Figueiredo',
        procurador_chave='eduardo',  # ou 'patrick'
        representante_legal=None,
    )
    res = gerar_inicial_am(dados, output_path='.../INICIAL_FABIO_AGIBANK.docx')
"""
import os, re, glob, sys
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers_docx import substituir_in_run
from extenso import extenso_moeda
from extrator_hiscon import parse_hiscon, filtrar_contratos_por_numero, formatar_contrato_para_template
from extrator_hiscre import parse_hiscre
from extrator_calculo import parse_calculo, eh_idoso
from extrator_procuracao import extrair_numeros_contrato_de_pasta
from bancos_canonicos import resolver_banco
from helpers_redacao import (
    substituir_intro_contratos,
    preencher_pedidos_declaratorios,
    remover_prioridade_pedidos,
    inserir_unidade_apoio_se_faltando,
    inserir_prioridade_idoso_se_faltando,
    inserir_pedido_prioridade_idoso_se_faltando,
    montar_paragrafo_recebimento_beneficio,
    substituir_paragrafo_completo,
)


class ProcuracaoSemFiltroError(RuntimeError):
    """Erro CRÍTICO: pipeline tentou rodar sem filtro válido de contratos.
    Ver SKILL.md §9-quater. A procuração é a ÚNICA fonte autoritativa do
    que o cliente nos autorizou a impugnar — NUNCA pegar "todos os contratos
    do banco" como fallback silencioso.
    """
    pass
from auditor_dano_moral import auditar_dano_moral, calcular_dano_moral
from verificador_dados_pessoais import comparar_doc_vs_hiscre, consolidar_dados_autora
from escritorios import PROCURADORES
from adaptador_am import adaptar_dados_para_am, montar_quali_banco

VAULT_TEMPLATES = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates'


def fmt_brl(v: Optional[float]) -> str:
    if v is None: return ''
    return f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')


def listar_procuracoes(pasta_banco: str) -> List[str]:
    """Lista os arquivos de procuração e extrai números de contrato.

    AM convention: '2. Procuração – AGIBANK.pdf' (sem número no nome).
    Para descobrir o número do contrato, vamos confiar no HISCON
    (todos os contratos do banco questionado) — caso AM tipicamente é 1 contrato.
    """
    candidatos = []
    for f in os.listdir(pasta_banco):
        if f.lower().startswith(('2.', '2 -', '2-')) and 'procura' in f.lower():
            candidatos.append(f)
    return candidatos


def encontrar_pdf(pasta: str, padrao: str) -> Optional[str]:
    for f in os.listdir(pasta):
        if padrao.lower() in f.lower() and f.lower().endswith('.pdf'):
            return os.path.join(pasta, f)
    return None


def gerar_data_extenso(d: datetime) -> str:
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    return f'{d.day} de {meses[d.month - 1]} de {d.year}'


def montar_dados_inicial_am(pasta_banco: str, autora: Dict, comarca: str,
                            procurador_chave: str = 'patrick',
                            representante_legal: Dict = None,
                            tipo_template: str = 'auto',
                            numeros_contrato_explicitos: Optional[List[str]] = None,
                            permitir_contrato_virtual: bool = False,
                            contrato_virtual_overrides: Optional[Dict] = None,
                            banco_codigo_override: Optional[str] = None) -> Dict:
    """Monta o dicionário completo de dados para uma inicial AM.

    Args:
        pasta_banco: ex. '.../FABIO/AGIBANK/'
        autora: dict com qualificação do autor
        comarca: 'Maués', 'Manaus', 'Boa Vista do Ramos', etc.
        procurador_chave: 'patrick' (default) ou 'eduardo'
        representante_legal: opcional (autor menor de idade)
        tipo_template: 'auto' (decide pelo HISCON), 'base' ou 'refin'
        numeros_contrato_explicitos: lista de números de contrato a filtrar
            do HISCON. OBRIGATÓRIO quando a procuração não traz o número no
            nome do arquivo nem dá pra extrair via OCR (caso típico AM).
            REGRA CRÍTICA (SKILL.md §9-quater): a procuração é a ÚNICA fonte
            autoritativa — NUNCA pegar "todos os contratos do banco" como
            fallback silencioso.

    Returns:
        dict completo para gerar inicial AM
    """
    # 1. HISCON
    hiscon_path = (encontrar_pdf(pasta_banco, 'histórico de empréstimo') or
                   encontrar_pdf(pasta_banco, 'historico de emprestimo'))
    if not hiscon_path:
        raise FileNotFoundError(f'HISCON não encontrado em {pasta_banco}')
    hiscon = parse_hiscon(hiscon_path)

    # 2. Identificar banco — primeiro pelo nome da pasta, depois pela procuração
    nome_pasta = os.path.basename(pasta_banco).upper()
    candidato_banco = None
    BANCOS_KW = ['AGIBANK', 'C6', 'ITAU', 'ITAÚ', 'FACTA', 'PAN', 'BRADESCO',
                 'DAYCOVAL', 'BMG', 'OLE', 'SANTANDER', 'SAFRA', 'MERCANTIL',
                 'INTER', 'INBURSA', 'PARANA', 'PARATI', 'SENFF', 'SICOOB',
                 'SICRED', 'PICPAY', 'CETELEM', 'QI', 'CREFISA',
                 'MASTER', 'NUBANK', 'CAPITAL CONSIGNADO', 'BNP', 'CIFRA',
                 'DO BRASIL', 'BANRISUL']
    for kw in BANCOS_KW:
        if kw in nome_pasta:
            candidato_banco = kw
            break
    # Se não achou no nome da pasta, procurar no nome das procurações
    if not candidato_banco:
        for f in os.listdir(pasta_banco):
            if f.lower().startswith(('2.', '2 -', '2-')) and 'procura' in f.lower():
                f_up = f.upper()
                for kw in BANCOS_KW:
                    if kw in f_up:
                        candidato_banco = kw
                        break
                if candidato_banco:
                    break
    if not candidato_banco:
        raise RuntimeError(
            f'Não foi possível identificar o banco. Nem o nome da pasta '
            f'"{nome_pasta}" nem as procurações 2.* contêm uma palavra-chave '
            f'reconhecida ({BANCOS_KW}).'
        )

    # Filtrar contratos do HISCON pelo banco identificado
    BANCO_MATCH = {
        'AGIBANK': ['AGIBANK'],
        'C6': ['C6'],
        'ITAU': ['ITAU', 'ITAÚ'],
        'ITAÚ': ['ITAU', 'ITAÚ'],
        'FACTA': ['FACTA'],
        'PAN': ['PAN'],
        'BRADESCO': ['BRADESCO'],
        'DAYCOVAL': ['DAYCOVAL'],
        'BMG': ['BMG'],
        'OLE': ['OLE'],
        'SANTANDER': ['SANTANDER'],
        'SAFRA': ['SAFRA'],
        'MERCANTIL': ['MERCANTIL'],
        'DO BRASIL': ['DO BRASIL', 'BRASIL S/A', 'BRASIL S.A'],
        'BANRISUL': ['BANRISUL'],
    }
    palavras_match = BANCO_MATCH.get(candidato_banco, [candidato_banco])
    # Match TOLERANTE A ESPAÇOS: pdfplumber às vezes quebra nomes no meio
    # ('BANCO INBURS A SA' em vez de 'BANCO INBURSA SA'). Removemos espaços de
    # ambos os lados antes de comparar para garantir match robusto.
    # Caso paradigma: EULALIA / INBURSA 2026-05-13.
    def _strip_spaces(s):
        return re.sub(r'\s+', '', s or '').upper()
    # OVERRIDE por código FEBRABAN (cadeia com bancos distintos na mesma pasta):
    # quando o JSON da kit-juridico identifica contratos predecessores (QI SCD,
    # Bradesco origem, etc.) dentro da pasta do banco final (INBURSA), filtramos
    # pelo código FEBRABAN do contrato específico em vez do banco da pasta.
    # Caso paradigma: EULALIA / cadeia INBURSA 2026-05-13.
    if banco_codigo_override:
        contratos_do_banco = [
            c for c in hiscon['contratos']
            if c.get('banco_codigo') == banco_codigo_override
        ]
        if not contratos_do_banco:
            raise RuntimeError(
                f'Nenhum contrato com código FEBRABAN {banco_codigo_override!r} '
                f'encontrado no HISCON (override de banco do contrato).'
            )
    else:
        contratos_do_banco = [
            c for c in hiscon['contratos']
            if any(_strip_spaces(kw) in _strip_spaces(c.get('banco_nome'))
                   or _strip_spaces(kw) in _strip_spaces(c.get('banco_nome_raw'))
                   for kw in palavras_match)
        ]
        if not contratos_do_banco:
            raise RuntimeError(f'Nenhum contrato do banco "{candidato_banco}" encontrado no HISCON')

    # REGRA CRÍTICA (SKILL.md §9-quater): a procuração é a única fonte
    # autoritativa do que o cliente nos autorizou a impugnar.
    # Hierarquia para obter os números de contrato a filtrar:
    #   1. numeros_contrato_explicitos (prioridade máxima — explícito do chamador)
    #   2. extrair via text-layer/OCR do conteúdo da procuração
    #   3. ABORTAR (NUNCA pegar todos os contratos do banco)
    if numeros_contrato_explicitos:
        nums_filtro = list(numeros_contrato_explicitos)
    else:
        info_proc = extrair_numeros_contrato_de_pasta(pasta_banco, usar_easyocr=True)
        nums_filtro = info_proc['numeros_unicos']
        if not nums_filtro:
            raise ProcuracaoSemFiltroError(
                f'🚨 IMPOSSÍVEL extrair números de contrato das procurações em '
                f'{pasta_banco}. Nem o nome do arquivo nem o conteúdo (text-layer + '
                f'OCR easyocr) trouxeram números válidos. AÇÃO: abrir o(s) PDF(s) '
                f'da(s) procuração(ões) manualmente, ler o número do(s) contrato(s) '
                f'outorgado(s), e chamar novamente passando '
                f'`numeros_contrato_explicitos=[...]`. NUNCA pegamos "todos os '
                f'contratos do banco" como fallback (caso paradigma: FABIO/C6 '
                f'07/05/2026 — pegava 7 contratos quando procuração só autorizava 1).'
            )

    contratos_brutos = filtrar_contratos_por_numero(
        contratos_do_banco, nums_filtro, fuzzy_dist=1)
    if not contratos_brutos:
        if permitir_contrato_virtual:
            # Contrato consta na procuração mas NÃO está no HISCON.
            # Regra do escritório (2026-05-13): NÃO abortar. Gerar inicial
            # com valores ESTIMADOS (via contrato_virtual_overrides) e marcar
            # pendência para juntar HISCON do período do empréstimo.
            ov = contrato_virtual_overrides or {}
            contratos_brutos = [{
                'numero': n,
                'banco_codigo': candidato_banco or '???',
                'banco_nome': candidato_banco or '?',
                'situacao': ov.get('situacao', 'Ativo (estimado — pendente HISCON)'),
                'origem_averbacao': ov.get('origem_averbacao', 'Averbação nova (estimada)'),
                'data_inclusao': ov.get('data_inclusao', '[A CONFIRMAR — pendente HISCON]'),
                'competencia_inicio': ov.get('competencia_inicio', ''),
                'competencia_fim': ov.get('competencia_fim', ''),
                'qtd_parcelas': ov.get('qtd_parcelas', 84),
                'valor_parcela': ov.get('valor_parcela', 0.0),
                'valor_emprestado': ov.get('valor_emprestado', 0.0),
                '_virtual': True,
                '_pendencia_hiscon': True,
            } for n in nums_filtro]
        else:
            raise ProcuracaoSemFiltroError(
                f'🚨 Nenhum contrato do HISCON casou com os números das procurações '
                f'{nums_filtro} para o banco {candidato_banco}. CONFERIR. '
                f'Para gerar inicial mesmo assim com pendência HISCON, passar '
                f'`permitir_contrato_virtual=True`.'
            )

    contratos_fmt = [formatar_contrato_para_template(c) for c in contratos_brutos]

    # 3. PDF de cálculo (pode estar na pasta ou não)
    calc_path = (encontrar_pdf(pasta_banco, '10- cálculo') or
                 encontrar_pdf(pasta_banco, '9- cálculo') or
                 encontrar_pdf(pasta_banco, 'cálculo') or
                 encontrar_pdf(pasta_banco, 'calculo'))
    calculo = parse_calculo(calc_path) if calc_path else {
        'valor_total_geral': None, 'dano_moral_pleiteado_pdf': None, 'idade': None,
        'data_nascimento': None,
    }

    # 4. HISCRE (preferencial em pasta_banco; fallback para 0. Kit do cliente)
    hiscre_path = (encontrar_pdf(pasta_banco, 'histórico de crédito') or
                   encontrar_pdf(pasta_banco, 'historico de credito') or
                   encontrar_pdf(pasta_banco, 'historico-creditos'))
    if not hiscre_path:
        # Procurar em pastas adjacentes
        cliente_root = os.path.dirname(pasta_banco)
        for raiz, _, arqs in os.walk(cliente_root):
            for a in arqs:
                if 'crédit' in a.lower() and a.lower().endswith('.pdf'):
                    hiscre_path = os.path.join(raiz, a)
                    break
                if 'credito' in a.lower() and a.lower().endswith('.pdf'):
                    hiscre_path = os.path.join(raiz, a)
                    break
            if hiscre_path: break
    hiscre = parse_hiscre(hiscre_path) if hiscre_path else {}

    # 5. Verificação cruzada doc vs HISCRE
    divergencias = comparar_doc_vs_hiscre(autora, hiscre)
    autora_consolidada = consolidar_dados_autora(autora, hiscre)

    # 6. Identificar banco-réu canônico
    banco_nome_hiscon = contratos_brutos[0].get('banco_nome', '')
    banco_reu = resolver_banco(banco_nome_hiscon, 'AM')
    if not banco_reu:
        banco_reu = resolver_banco(banco_nome_hiscon, 'matriz')
    if not banco_reu:
        raise RuntimeError(f'Banco réu não identificado: "{banco_nome_hiscon}"')

    # 7. Selecionar template
    n = len(contratos_brutos)
    tem_refin = any(c.get('tipo_origem') == 'refinanciamento' for c in contratos_brutos)
    if tipo_template == 'refin' or (tipo_template == 'auto' and tem_refin and n == 1):
        template_nome = 'inicial-jeam-refin.docx'
    else:
        template_nome = 'inicial-jeam-base.docx'
    template_path = os.path.join(VAULT_TEMPLATES, template_nome)

    # Alertas
    alertas = []
    if n > 1:
        alertas.append(
            f'AM: {n} contratos do mesmo banco — template AM ainda não suporta '
            f'múltiplos contratos automaticamente. Gerada inicial com APENAS o '
            f'primeiro contrato. CONFERIR e adaptar manualmente para incluir '
            f'os demais ({n-1} contratos restantes).'
        )
    if tem_refin and tipo_template == 'auto' and n > 1:
        alertas.append('Caso contém refinanciamento(s). REVISAR template selecionado.')

    # 8. Dano moral
    dm = calcular_dano_moral(contratos_brutos[:1])  # AM = 1 contrato por enquanto
    audit_dm = auditar_dano_moral(contratos_brutos[:1], calculo.get('dano_moral_pleiteado_pdf'))

    # 9. Idade (do HISCRE se disponível, senão do PDF cálculo)
    eh_id = False
    if hiscre.get('data_nascimento'):
        d = hiscre['data_nascimento']
        if isinstance(d, str):
            try:
                d = datetime.fromisoformat(d)
            except ValueError:
                d = None
        if d:
            from datetime import datetime as dt
            hoje = dt.today()
            idade = hoje.year - d.year - ((hoje.month, hoje.day) < (d.month, d.day))
            eh_id = idade >= 60

    # 10. Procurador
    procurador = PROCURADORES.get(procurador_chave, PROCURADORES['patrick'])

    # 11. Valor da causa
    vc = calculo.get('valor_total_geral')
    if vc is None:
        # Calcular: dano moral + soma das parcelas × 2 (dobro CDC)
        soma_parc_dobro = sum(
            (c['valor_parcela_float'] or 0) * (c['qtd_parcelas'] or 0) * 2
            for c in contratos_fmt[:1]
        )
        vc = dm['total'] + soma_parc_dobro
        alertas.append(
            f'PDF de cálculo NÃO encontrado. Valor da causa estimado pela '
            f'fórmula (parcelas × N × 2 + dano moral) = R$ {fmt_brl(vc)}. '
            f'CONFERIR antes do protocolo.'
        )
    vc_extenso = extenso_moeda(vc)

    # 12. Montar placeholders BA-style (para usar no adaptador)
    cpf_final = (autora_consolidada.get('cpf') or '').strip()
    renda = hiscre.get('valor_liquido') or autora.get('renda_liquida')
    dados_ba_intermed = {
        '{{cpf_autor}}': cpf_final,
        '{{agencia_pagador}}': hiscon['cabecalho'].get('agencia_pagador', ''),
        '{{conta_pagador}}': hiscon['cabecalho'].get('conta_pagador', ''),
        '{{banco_pagador}}': hiscon['cabecalho'].get('banco_pagador', ''),
        '{{valor_renda_liquida}}': fmt_brl(renda),
    }

    # 13. Adaptar para AM
    placeholders_am = adaptar_dados_para_am(
        dados_ba=dados_ba_intermed,
        hiscre=hiscre,
        autora=autora_consolidada,
        contrato=contratos_fmt[0],
        banco_reu=banco_reu,
        comarca_am=comarca,
        procurador=procurador,
        valor_causa=vc,
        valor_causa_extenso=vc_extenso,
        representante_legal=representante_legal,
    )

    return {
        'pasta_banco': pasta_banco,
        'hiscon': hiscon,
        'hiscre': hiscre,
        'autora_consolidada': autora_consolidada,
        'divergencias_pessoais': divergencias,
        'contratos_questionados': contratos_fmt,
        'contratos_brutos': contratos_brutos,
        'calculo': calculo,
        'banco_reu': banco_reu,
        'template': template_path,
        'alertas': alertas,
        'dano_moral': dm,
        'audit_dm': audit_dm,
        'eh_idoso': eh_id,
        'placeholders_am': placeholders_am,
        'comarca': comarca,
        'procurador': procurador,
        'representante_legal': representante_legal,
        'valor_causa': vc,
    }


def gerar_inicial_am(dados_caso: Dict, output_path: str) -> Dict:
    """Aplica o template AM, substituindo placeholders e tratando casos
    especiais (representante legal, procurador override, idoso/não-idoso).

    Particularidade: o {{quali_banco}} é substituído com um SEPARADOR `¤¤¤`
    entre o nome do banco e o resto. Depois, no pós-processamento, esse
    separador é detectado e o run quebrado em 2: nome em Segoe UI Bold,
    resto em Cambria.
    """
    import shutil
    from copy import deepcopy
    from lxml import etree
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from adaptador_am import QUALI_BANCO_SEP, classificar_menor

    template = dados_caso['template']
    placeholders = dados_caso['placeholders_am']
    eh_idoso = dados_caso['eh_idoso']
    procurador = dados_caso['procurador']
    banco_reu = dados_caso['banco_reu']
    representante_legal = dados_caso.get('representante_legal')
    autora = dados_caso.get('autora_consolidada') or {}

    shutil.copy(template, output_path)
    doc = Document(output_path)

    W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    XMLSPC = '{http://www.w3.org/XML/1998/namespace}space'

    # 1. Remover prioridade idoso se autor não-idoso
    if not eh_idoso:
        for p in list(doc.paragraphs):
            if 'Prioridade de tramitação: art. 1.048' in p.text:
                p._element.getparent().remove(p._element)
            elif 'A prioridade na tramitação' in p.text:
                p._element.getparent().remove(p._element)

    # 1-bis. Reescrever parágrafo de qualificação se houver REPRESENTANTE LEGAL
    # (autor menor de idade — impúbere é REPRESENTADO; púbere é ASSISTIDO).
    # Esta etapa precisa rodar ANTES da substituição de placeholders, porque ela
    # remove os {{nome_completo}}, {{cpf}}, {{rg}} etc. do parágrafo 6 e os
    # substitui por uma estrutura customizada de runs (5 runs: parte autora,
    # nome autora bold, conector, nome mãe bold, resto).
    if representante_legal:
        cls = classificar_menor(autora.get('data_nascimento'))
        estado_menor = f'menor {cls["classe"]}'
        verbo = cls['verbo_repr'] or 'representada por sua genitora'
        prof_aut = autora.get('profissao') or 'beneficiária'
        rep = representante_legal

        # Texto antes do nome da autora: "" (nada — começa com o nome)
        # Run 1: NOME AUTORA — Segoe UI Bold (rStyle 2TtuloChar) + amarelo
        # Run 2: ", brasileira, menor impúbere, beneficiária, inscrita no CPF
        #         sob o nº 095.239.132-55, neste ato representada por sua genitora "
        #         Cambria + amarelo
        # Run 3: NOME MÃE — Segoe UI Bold (rStyle 2TtuloChar) + amarelo
        # Run 4: ", brasileira, solteira, beneficiária do INSS, inscrita no
        #         CPF sob o nº 031.113.782-25, Cédula de Identidade nº
        #         2.943.060-7, órgão expedidor SSP/AM, residente e domiciliada
        #         na ..."
        #         Cambria + amarelo

        texto_meio = (
            f', {autora.get("nacionalidade", "brasileira")}, {estado_menor}, '
            f'{prof_aut}, inscrita no CPF sob o nº {autora["cpf"]}, '
            f'neste ato {verbo} '
        )
        partes_rep = [rep.get('nacionalidade', 'brasileira')]
        if rep.get('estado_civil'):
            partes_rep.append(rep['estado_civil'])
        if rep.get('profissao'):
            partes_rep.append(rep['profissao'])
        partes_rep.append(f'inscrita no CPF sob o nº {rep["cpf"]}')
        partes_rep.append(
            f'Cédula de Identidade nº {rep["rg"]}, '
            f'órgão expedidor {rep.get("orgao_expedidor", "SSP/AM")}'
        )
        endereco_str = (
            f'residente e domiciliada na {autora.get("logradouro", "")}'
            f', nº {autora.get("numero", "s/nº")}'
            f', bairro {autora.get("bairro", "")}'
            f', {autora.get("cidade", "")}/{autora.get("uf", "AM")}'
            f', CEP {autora.get("cep", "")}'
        )
        # Endereço composto matriz Joaçaba/SC + unidade de apoio em Maués/AM
        # (regra fixa do escritório, 07/05/2026)
        from escritorios import montar_endereco_escritorio_completo
        end_escritorio = montar_endereco_escritorio_completo('AM')
        texto_fim = ', ' + ', '.join(partes_rep) + ', ' + endereco_str + (
            ', não possui endereço eletrônico, por seus advogados que assinam '
            'digitalmente a presente peça (instrumento de procuração anexo), '
            f'com escritório profissional em {end_escritorio}, local onde '
            'recebem avisos e intimações, vem, respeitosamente, perante Vossa '
            'Excelência, propor a presente'
        )

        # Localizar o parágrafo de qualificação (contém {{nome_completo}})
        for p in list(doc.paragraphs):
            if '{{nome_completo}}' not in p.text:
                continue
            # Limpar runs do parágrafo
            pPr = p._element.find(W + 'pPr')
            for child in list(p._element):
                if child.tag != W + 'pPr':
                    p._element.remove(child)

            def _make_run(texto, *, bold=False, preserve=True):
                r = OxmlElement('w:r')
                rpr = OxmlElement('w:rPr')
                r.append(rpr)
                if bold:
                    rs = OxmlElement('w:rStyle')
                    rs.set(qn('w:val'), '2TtuloChar')
                    rpr.append(rs)
                else:
                    # Cambria explícito p/ resto (idem aos outros runs do template)
                    rfonts = OxmlElement('w:rFonts')
                    rfonts.set(qn('w:ascii'), 'Cambria')
                    rfonts.set(qn('w:hAnsi'), 'Cambria')
                    rpr.append(rfonts)
                hl = OxmlElement('w:highlight')
                hl.set(qn('w:val'), 'yellow')
                rpr.append(hl)
                t = OxmlElement('w:t')
                if preserve:
                    t.set(qn('xml:space'), 'preserve')
                t.text = texto
                r.append(t)
                return r

            p._element.append(_make_run(autora['nome'], bold=True))
            p._element.append(_make_run(texto_meio, bold=False))
            p._element.append(_make_run(rep['nome'], bold=True))
            p._element.append(_make_run(texto_fim, bold=False))
            break

    # 2. Substituir TODOS os placeholders (incluindo {{quali_banco}} com separador ¤¤¤)
    placeholders_validos = {k: v for k, v in placeholders.items() if not k.startswith('_')}
    modificados = 0
    for p in doc.paragraphs:
        for k, v in placeholders_validos.items():
            if k in p.text:
                if substituir_in_run(p._element, {k: v}, grifo=True):
                    modificados += 1

    # 3. Pós-processamento: encontrar o separador ¤¤¤ e quebrar em 2 runs
    # (nome do banco em Segoe UI Bold + resto em Cambria, ambos com highlight amarelo)
    for p in list(doc.paragraphs):
        if QUALI_BANCO_SEP not in p.text:
            continue
        # Localizar o run que contém o separador
        for r in list(p._element.findall('.//' + W + 'r')):
            t_el = r.find(W + 't')
            if t_el is None or not t_el.text or QUALI_BANCO_SEP not in t_el.text:
                continue
            txt = t_el.text
            idx_sep = txt.find(QUALI_BANCO_SEP)
            parte_banco = txt[:idx_sep]  # nome do banco
            parte_resto = txt[idx_sep + len(QUALI_BANCO_SEP):]  # resto da qualificação

            # rPr original (já tem highlight amarelo aplicado pelo substituir_in_run)
            rpr_original = r.find(W + 'rPr')

            # Atualizar o run atual com APENAS o resto (Cambria + highlight)
            t_el.text = parte_resto
            if parte_resto and (parte_resto[0].isspace() or parte_resto[-1].isspace()):
                t_el.set(XMLSPC, 'preserve')

            # Inserir ANTES um novo run com o nome do banco em Segoe UI Bold + amarelo
            # IMPORTANTE: usar OxmlElement (não etree.Element) para manter o proxy
            # class do python-docx; do contrário, p.text quebra com TypeError.
            parent = r.getparent()
            idx_pos = list(parent).index(r)

            r_banco = OxmlElement('w:r')
            rpr_banco = OxmlElement('w:rPr')
            r_banco.append(rpr_banco)

            rfonts_b = OxmlElement('w:rFonts')
            rfonts_b.set(qn('w:ascii'), 'Segoe UI')
            rfonts_b.set(qn('w:hAnsi'), 'Segoe UI')
            rpr_banco.append(rfonts_b)

            rpr_banco.append(OxmlElement('w:b'))
            rpr_banco.append(OxmlElement('w:bCs'))

            hl_b = OxmlElement('w:highlight')
            hl_b.set(qn('w:val'), 'yellow')
            rpr_banco.append(hl_b)

            t_banco = OxmlElement('w:t')
            t_banco.text = parte_banco
            t_banco.set(qn('xml:space'), 'preserve')
            r_banco.append(t_banco)

            parent.insert(idx_pos, r_banco)
            break

    # 4. Trocar procurador no rodapé (se diferente do default Patrick)
    if procurador['oab'] != 'OAB/AM A2638':
        for p in doc.paragraphs:
            if 'Patrick Willian da Silva' in p.text:
                substituir_in_run(p._element, {'Patrick Willian da Silva': procurador['nome']}, grifo=True)
            if 'OAB/AM 02638' in p.text or 'OAB/AM A2638' in p.text:
                substituir_in_run(p._element,
                    {'OAB/AM 02638': procurador['oab'], 'OAB/AM A2638': procurador['oab']},
                    grifo=True)
            if 'Patrick' in p.text and procurador['nome'] != 'Patrick Willian da Silva':
                substituir_in_run(p._element, {'Patrick Willian da Silva': procurador['nome']}, grifo=True)
                substituir_in_run(p._element, {'Patrick': procurador['nome'].split()[0]}, grifo=True)

    # 5. Reescrever a INTRO FÁTICA — agrupa por banco quando há ≥2 bancos
    # (litisconsórcio passivo). Helper compartilhado com BA/AL.
    contratos_fmt_caso = dados_caso.get('contratos_questionados') or []
    nums_contratos = [c.get('numero', '') for c in contratos_fmt_caso if c.get('numero')]
    if nums_contratos:
        from helpers_redacao import aplicar_intro_fatica
        for p in list(doc.paragraphs):
            if ('tomou conhecimento dos descontos referentes' in p.text
                or 'constatou a existência de descontos referentes' in p.text):
                aplicar_intro_fatica(p._element, contratos_fmt_caso,
                                       banco_reu['nome'], grifo=True)
                modificados += 1
                break

    # 6. Reescrever PEDIDOS DECLARATÓRIOS — escolher empréstimo vs
    # refinanciamento conforme tipo_origem; 1 pedido por contrato.
    nb = (dados_caso.get('hiscon') or {}).get('cabecalho', {}).get('nb_beneficio') or '___'
    if contratos_fmt_caso:
        modificados += preencher_pedidos_declaratorios(
            doc, contratos_fmt_caso, nb, grifo=True)

    # 7. Remover prioridade dos PEDIDOS quando não-idoso (o cabeçalho já
    # foi tratado na etapa 1).
    if not eh_idoso:
        modificados += remover_prioridade_pedidos(doc)

    # 7-bis. Padronizar endereço escritório: matriz Joaçaba/SC + unidade apoio
    # da UF do cliente (atualiza hardcoded do template para versão composta).
    modificados += inserir_unidade_apoio_se_faltando(doc, 'AM', grifo=True)

    # 7-quater. Reescrever PARÁGRAFO DE RECEBIMENTO: cartão magnético vs
    # conta bancária (regra fixa 07/05/2026, Gabriel). O template AM usa
    # placeholders {{conta_agencia_conta}} + {{banco_que_recebe}} já
    # substituídos antes, mas o texto fixo "agência X junto Y" não cabe
    # quando o autor recebe por cartão magnético.
    hiscon_cab_p = (dados_caso.get('hiscon') or {}).get('cabecalho', {}) or {}
    for p in list(doc.paragraphs):
        t = p.text
        if 'recebe benefício previdenciário' in t and ('NB' in t or 'sobrevive basicamente' in t):
            novo_recebimento = montar_paragrafo_recebimento_beneficio(
                hiscon_cab_p, autora)
            substituir_paragrafo_completo(p._element, novo_recebimento, grifo=True)
            modificados += 1
            break

    # 7-ter. Garantir prioridade no CABEÇALHO + nos PEDIDOS quando idoso.
    if eh_idoso:
        modificados += inserir_prioridade_idoso_se_faltando(doc, eh_idoso, grifo=True)
        modificados += inserir_pedido_prioridade_idoso_se_faltando(doc, eh_idoso, grifo=True)

    # 8. Verificar residuais
    residuais = []
    for p in doc.paragraphs:
        for ph in re.findall(r'\{\{[^}]+\}\}', p.text):
            if ph not in residuais:
                residuais.append(ph)

    doc.save(output_path)
    return {'modificados': modificados, 'residuais': residuais, 'output': output_path}
