# -*- coding: utf-8 -*-
"""Teste de geração da inicial jfal-1banco para JOSÉ EXEMPLO DA SILVA × Itaú.
Federal AL — JFAL Subseção Arapiraca. Caso real:
- Aposentadoria por idade NB 198.882.188-3
- Recebido por CARTÃO MAGNÉTICO (não conta corrente — par 12 ajustado)
- 1 contrato fraudulento de refinanciamento ITAÚ 652247564
"""
import sys, io, shutil, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
from docx import Document
from helpers_docx import substituir_in_run
from extenso import extenso_moeda
from _blocos_narrativos import (
    gerar_bloco_contratos_fraudulentos,
    gerar_bloco_pedido_declaracao,
    normalizar_banco_reu,
)

PASTA = Path(r"C:\Users\gabri\OneDrive\Área de Trabalho\APP - NÃO CONTRATADO\JOSÉ EXEMPLO DA SILVA - ALEXANDRE-ESCRITÓRIO\APOSENTADORIA")
TEMPLATE = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jfal-1banco.docx")
OUT = PASTA / "INICIAL_CICERO_JFAL_TESTE.docx"

# ===== Dados extraídos dos documentos =====
contrato = {
    'numero': '652247564',
    'valor_emprestado': 17278.49,
    'valor_parcela': 384.90,
    'qtd_parcelas': 84,
    'competencia_inicio': '11/2024',
    'data_inclusao': '07/10/2024',
}
banco_reu_nome = normalizar_banco_reu('BANCO ITAÚ CONSIGNADO S.A.')

renda_liq = 1100.00
nb = '198.882.188-3'

dados = {
    # Cabeçalho federal — Subseção Arapiraca/AL
    '{{vara_protocolo}}':    '___',  # número da vara — pendência
    '{{cidade_protocolo}}':  'Arapiraca',
    '{{uf_protocolo}}':      'AL',

    # Autor — TODOS extraídos da procuração + RG + comprovante
    '{{nome_autor}}':         'JOSÉ EXEMPLO DA SILVA',
    '{{nacionalidade}}':      'brasileiro',
    '{{estado_civil}}':       'viúvo',
    '{{profissao}}':          'aposentado',
    '{{cpf_autor}}':          '000.000.006-16',
    '{{rg_autor}}':           '4458132-7',
    '{{orgao_expedidor}}':    'SSP/AL',
    '{{logradouro_autor}}':   'Rua Floriano Leite',
    '{{numero_autor}}':       '28',
    '{{bairro_autor}}':       'Centro',
    '{{cidade_autor}}':       'Lagoa da Canoa',
    '{{uf_autor}}':           'AL',
    '{{cep_autor}}':          '57330-000',

    # Banco réu (uppercase forçado)
    '{{banco_reu_nome}}':         banco_reu_nome,
    '{{banco_reu_descricao_pj}}': 'pessoa jurídica de direito privado',
    '{{banco_reu_cnpj}}':         '33.885.724/0001-19',
    '{{banco_reu_endereco}}':     'Praça Alfredo Egydio de Souza Aranha, nº 100, Torre Itaúsa, Parque Jabaquara, São Paulo/SP, CEP 04344-902',

    # Benefício — placeholders padrão (par 12 será REESCRITO depois)
    '{{tipo_beneficio}}':   'aposentadoria por idade',
    '{{nb_beneficio}}':     nb,
    '{{banco_pagador}}':    'BANCO ITAÚ',
    '{{agencia_pagador}}':  '550149',
    '{{conta_pagador}}':    'CARTÃO MAGNÉTICO',

    # Renda
    '{{valor_renda_liquida}}':         '1.100,00',
    '{{valor_renda_liquida_extenso}}': extenso_moeda(renda_liq),

    # Prioridade — autora 73 anos → idoso
    '{{pedido_prioridade}}': 'A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de Processo Civil',

    # Valor da causa — soma do dobro dos descontos + dano moral
    # 1 parcela × 12 (estimado, ainda não temos cálculo) = ~R$ 4.618,80; dobro=9237,60; dano=15000
    # Vou usar VALOR PROVISÓRIO baseado em estimativa simples
    '{{valor_causa}}':         '15.000,00',  # placeholder mínimo — pipeline real calcula
    '{{valor_causa_extenso}}': extenso_moeda(15000.00),

    # BLOCOS NARRATIVOS
    '{{BLOCO_CONTRATOS_FRAUDULENTOS}}': gerar_bloco_contratos_fraudulentos(
        [contrato], banco_reu_nome),
    '{{BLOCO_PEDIDO_DECLARACAO}}': gerar_bloco_pedido_declaracao(
        [contrato], nb),
}

shutil.copy2(TEMPLATE, OUT)
doc = Document(OUT)

# === Aplicar todas substituições padrão ===
modificados = 0
for p in doc.paragraphs:
    for k, v in dados.items():
        if k in p.text:
            if substituir_in_run(p._p, {k: v}):
                modificados += 1

# === REESCRITA ESPECÍFICA do PAR 12 — cartão magnético em vez de conta corrente ===
NOVO_PAR12 = (
    'A parte autora recebe benefício previdenciário de aposentadoria por idade '
    '– NB 198.882.188-3, recebido por meio de cartão magnético junto ao '
    'BANCO ITAÚ — OP 550149 ARAPIRACA/AL, sendo que sobrevive basicamente '
    'do que recebe da previdência.'
)
# Substituir o conteúdo do par 12 inteiro
p12 = doc.paragraphs[12]
TROCA_PAR12 = {p12.text: NOVO_PAR12}
substituir_in_run(p12._p, TROCA_PAR12)

doc.save(OUT)
print(f"OK gerado: {OUT.name}")
print(f"Substituições: {modificados}")
print()

# Verificar
d2 = Document(OUT)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r'\{\{[^}]+\}\}', p.text):
        phs.add(m.group(0))
print(f"Placeholders residuais: {phs if phs else 'NENHUM ✓'}")

marcadores = set()
for p in d2.paragraphs:
    for m in re.finditer(r'\[[A-Z_º]+\]', p.text):
        marcadores.add(m.group(0))
print(f"Marcadores [XXX] residuais: {marcadores if marcadores else 'NENHUM ✓'}")

print()
for i in [0, 6, 10, 12, 15, 17, 257, 270, 271, 291]:
    if i < len(d2.paragraphs):
        t = d2.paragraphs[i].text
        if t.strip():
            print(f'PAR {i}: {t[:300]}')
            print()
