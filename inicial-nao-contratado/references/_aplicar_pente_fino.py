# -*- coding: utf-8 -*-
"""Aplica correções #1 a #15 do pente fino nos 5 templates não-contratado.
Backup: <arquivo>.bak_pre_pente_fino
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

VAULT = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates")

# === Correções universais (todos os 5 templates) ===
SUBS_UNIVERSAIS = {
    # #3 PRINT DOS CONTRATOS — esvaziar parágrafo será feito separadamente
    # #4 "do demando" → "do demandado"
    "do demando em produzir":  "do demandado em produzir",
    # #5 "ào banco" → "ao banco"
    "cabendo ào banco":        "cabendo ao banco",
    # #6 "dispões" → "disposições"
    "conforme dispões legais apontadas":  "conforme disposições legais apontadas",
    # #7 "que vem sendo gerados" → "que vêm sendo gerados"
    "que vem sendo gerados":   "que vêm sendo gerados",
    # #8 "Requerida... continuam cometendo" → "continua"
    "porém, mesmo assim, continuam cometendo":   "porém, mesmo assim, continua cometendo",
    # #9 "Tribunal de Justiça do Estado de Minas Gerais, têm aplicado" → "tem aplicado"
    "Tribunal de Justiça do Estado de Minas Gerais, têm aplicado":
        "Tribunal de Justiça do Estado de Minas Gerais, tem aplicado",
    # #10 "TOTALMENTE PROCEDENTE;" → "TOTALMENTE PROCEDENTES;"
    "TOTALMENTE PROCEDENTE;":  "TOTALMENTE PROCEDENTES;",
    # #11 "hipossuficiente em relação a ré" → "à ré"
    "hipossuficiente em relação a ré":  "hipossuficiente em relação à ré",
    # #12 "Excelência, que a parte autora não contratou tal empréstimo." → "Excelência, importa registrar que..."
    "Excelência, que a parte autora não contratou tal empréstimo.":
        "Excelência, importa registrar que a parte autora não contratou tal empréstimo.",
    # #13 "Tribunal e Justiça" → "Tribunal de Justiça"
    "Tribunal e Justiça de Alagoas":  "Tribunal de Justiça de Alagoas",
}

# === Correções específicas por template ===
SUBS_ESPECIFICAS = {
    'inicial-jeal-1banco.docx': {
        # #1 valor causa hardcoded → parametrizado
        "Dá-se a causa o valor de R$ 23.877,98 (vinte e três mil, oitocentos e setenta e sete reais e noventa e oito centavos).":
            "Dá-se a causa o valor de R$ {{valor_causa}} ({{valor_causa_extenso}}).",
        # #14 "aos réus" → "ao réu" (1 banco)
        "em relação aos réus.":  "em relação ao réu.",
        # #15 "dos requeridos" → "do requerido" (1 banco)
        "a condenação dos requeridos ao pagamento de custas":
            "a condenação do requerido ao pagamento de custas",
        "Condenar os requeridos ao pagamento das custas":
            "Condenar o requerido ao pagamento das custas",
    },
    'inicial-jeal-2bancos.docx': {
        # 2 bancos = "aos réus" e "os requeridos" CORRETOS, manter
        # Só correções universais aplicadas
    },
    'inicial-jemg-1banco.docx': {
        # #14 "aos réus" → "ao réu" (1 banco)
        "em relação aos réus.":  "em relação ao réu.",
        # #15 "dos requeridos" → "do requerido"
        "a condenação dos requeridos ao pagamento de custas":
            "a condenação do requerido ao pagamento de custas",
        "Condenar os requeridos ao pagamento das custas":
            "Condenar o requerido ao pagamento das custas",
    },
    'inicial-jfal-1banco.docx': {
        # #2 "R$ (centavos)" → parametrizar
        "Dá-se a causa o valor de R$ (centavos).":
            "Dá-se a causa o valor de R$ {{valor_causa}} ({{valor_causa_extenso}}).",
        # 1 banco + INSS = 2 réus, "aos réus" CORRETO, manter
    },
    'inicial-jfal-2bancos.docx': {
        # 2 bancos + INSS = 3 réus, "aos réus" CORRETO, manter
    },
}


def esvaziar_par_print(doc):
    """#3 — Esvazia o parágrafo 'PRINT DOS CONTRATOS SUBLINHADOS'."""
    n = 0
    for p in doc.paragraphs:
        if p.text.strip() == "PRINT DOS CONTRATOS SUBLINHADOS":
            for r in list(p._p.findall(W + 'r')):
                p._p.remove(r)
            n += 1
    return n


total_geral = 0
for nome in ['inicial-jeal-1banco.docx', 'inicial-jeal-2bancos.docx',
              'inicial-jemg-1banco.docx', 'inicial-jfal-1banco.docx',
              'inicial-jfal-2bancos.docx']:
    caminho = VAULT / nome
    backup = caminho.with_suffix(caminho.suffix + ".bak_pre_pente_fino")
    if not backup.exists():
        shutil.copy2(caminho, backup)

    doc = Document(caminho)

    # Aplicar todas as universais + específicas
    todas_subs = {**SUBS_UNIVERSAIS, **SUBS_ESPECIFICAS.get(nome, {})}

    n_par = 0
    for p in doc.paragraphs:
        if substituir_in_run(p._p, todas_subs):
            n_par += 1

    # #3 esvaziar PRINT DOS CONTRATOS
    n_print = esvaziar_par_print(doc)

    doc.save(caminho)
    total_geral += n_par + n_print
    print(f"{nome}: {n_par} parágrafos modificados + {n_print} 'PRINT' esvaziado(s)")

print(f"\nTotal: {total_geral} alterações")

# === Verificação final — listar erros remanescentes ===
import re
print("\n=== Verificação final dos erros corrigidos ===")
ERROS_BUSCAR = [
    "do demando em produzir", "ào banco", "dispões legais",
    "que vem sendo gerados", "continuam cometendo",
    "Estado de Minas Gerais, têm aplicado",
    "TOTALMENTE PROCEDENTE;", "em relação a ré",
    "Excelência, que a parte autora não contratou",
    "Tribunal e Justiça",
    "23.877,98", "R$ (centavos)",
    "PRINT DOS CONTRATOS",
]
for nome in ['inicial-jeal-1banco.docx', 'inicial-jeal-2bancos.docx',
              'inicial-jemg-1banco.docx', 'inicial-jfal-1banco.docx',
              'inicial-jfal-2bancos.docx']:
    d = Document(VAULT / nome)
    sobras = []
    for i, p in enumerate(d.paragraphs):
        for ruim in ERROS_BUSCAR:
            if ruim in p.text:
                sobras.append((i, ruim))
                break
    if sobras:
        print(f"  ⚠️ {nome}:")
        for i, r in sobras:
            print(f"    par {i}: '{r}' ainda presente")
    else:
        print(f"  ✓ {nome}: limpo")
