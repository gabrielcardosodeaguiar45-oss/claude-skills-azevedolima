"""Orquestrador do pipeline de geração da inicial NÃO CONTRATADO.

Recebe:
- caminho da pasta do banco (ex.: "GEORGE/BANCO ITAÚ/2 AVERBAÇÃO NOVA INATIVO/")
- dados de qualificação do autor (manualmente codificados ou OCR-extraídos do KIT)
- jurisdição/subseção JEF

Produz:
- DOCX da inicial (template selecionado + placeholders preenchidos)
- DOCX do relatório paralelo com pendências/alertas
"""
import os, re, glob, sys
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers_docx import substituir_in_run
from extenso import extenso_moeda, extenso_cardinal
from extrator_hiscon import (parse_hiscon, filtrar_contratos_por_numero,
                              formatar_contrato_para_template,
                              auditar_procuracoes_vs_hiscon,
                              validar_contratos_obrigatorios)
from extrator_procuracao import extrair_numeros_contrato_de_pasta


class ProcuracaoSemFiltroError(RuntimeError):
    """Erro CRÍTICO: pipeline tentou rodar sem filtro válido de contratos.
    Ver SKILL.md §9-quater. A procuração é a ÚNICA fonte autoritativa do
    que o cliente nos autorizou a impugnar — NUNCA pegar "todos os contratos
    do banco" como fallback silencioso.
    """
    pass
from extrator_hiscre import parse_hiscre
from extrator_calculo import parse_calculo, eh_idoso
from _blocos_narrativos import (
    gerar_bloco_contratos_fraudulentos,
    gerar_bloco_pedido_declaracao,
    normalizar_banco_reu,
)
from bancos_canonicos import resolver_banco, INSS_FIXO
from seletor_template import selecionar_template, descrever_caso
from auditor_dano_moral import auditar_dano_moral, calcular_dano_moral
from verificador_dados_pessoais import comparar_doc_vs_hiscre, consolidar_dados_autora
from helpers_redacao import (
    substituir_intro_contratos,
    preencher_pedidos_declaratorios,
    remover_prioridade_pedidos,
    inserir_unidade_apoio_se_faltando,
    inserir_prioridade_idoso_se_faltando,
    inserir_pedido_prioridade_idoso_se_faltando,
    montar_paragrafo_recebimento_beneficio,
    substituir_paragrafo_completo,
)


def fmt_brl(v: Optional[float]) -> str:
    if v is None:
        return ''
    return f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')


def listar_procuracoes(pasta_banco: str) -> List[str]:
    """Lista os arquivos `2 - PROCURAÇÃO XXX <NÚMERO>.pdf` da pasta do banco
    e extrai os NÚMEROS dos contratos questionados (do nome do arquivo)."""
    padrao = os.path.join(pasta_banco, '2 - PROCURAÇÃO *.pdf')
    arquivos = glob.glob(padrao)
    # Suporta também variantes "2- Procuração ..." (sem espaço, minúsculas)
    if not arquivos:
        for nome in os.listdir(pasta_banco) if os.path.isdir(pasta_banco) else []:
            if 'procura' in nome.lower() and nome.lower().endswith('.pdf'):
                arquivos.append(os.path.join(pasta_banco, nome))
    numeros = []
    for f in arquivos:
        nome = os.path.basename(f)
        # Padrão clássico: "2 - PROCURAÇÃO BANCO XXX 0123456789.pdf"
        m = re.search(r'(\d{6,15}(?:-\d+)?)\.pdf$', nome)
        if m:
            numeros.append(m.group(1))
            continue
        # Padrão variante: "2- Procuração - Banco X - Contrato 12345.pdf"
        for m2 in re.finditer(r'[Cc]ontrato\s+(?:n[º°]\s*)?(\d{6,15}(?:-\d+)?)', nome):
            numeros.append(m2.group(1))
    return sorted(set(numeros))


def _numeros_de_contratos_impugnar_json(pasta_banco: str) -> tuple[List[str], str]:
    """Lê contratos_impugnar_ids do _estado_cliente.json (se houver) e
    converte para números de contrato.

    Retorna (numeros, origem). origem ∈ {'manual', 'sugestao_automatica_revisada',
    'sugestao_automatica', ''}.
    """
    # Sobe diretórios procurando _estado_cliente.json
    cur = pasta_banco
    json_path = None
    for _ in range(5):  # max 5 níveis
        candidate = os.path.join(cur, '_estado_cliente.json')
        if os.path.isfile(candidate):
            json_path = candidate
            break
        parent = os.path.dirname(cur)
        if parent == cur:
            break
        cur = parent
    if not json_path:
        return ([], '')

    try:
        import json as _json
        with open(json_path, encoding='utf-8') as f:
            estado = _json.load(f)
    except Exception:
        return ([], '')

    pasta_cliente = os.path.dirname(json_path)
    pasta_banco_norm = os.path.normpath(pasta_banco).lower()

    pastas_acao = estado.get('pastas_acao') or []
    contratos = {c.get('id_interno'): c for c in (estado.get('contratos') or []) if c.get('id_interno')}
    for pa in pastas_acao:
        path_rel = pa.get('path_relativo') or ''
        candidato_abs = os.path.normpath(os.path.join(pasta_cliente, path_rel)).lower()
        if candidato_abs != pasta_banco_norm:
            continue
        ids = pa.get('contratos_impugnar_ids') or []
        if not ids:
            return ([], pa.get('contratos_impugnar_origem', ''))
        nums = []
        for cid in ids:
            c = contratos.get(cid)
            if c and c.get('contrato'):
                nums.append(str(c['contrato']))
        return (sorted(set(nums)), pa.get('contratos_impugnar_origem', ''))
    return ([], '')


def encontrar_pdf(pasta_banco: str, padrao_nome: str) -> Optional[str]:
    """Encontra o PDF que casa com o padrão (case-insensitive)."""
    for f in os.listdir(pasta_banco):
        if padrao_nome.lower() in f.lower() and f.lower().endswith('.pdf'):
            return os.path.join(pasta_banco, f)
    return None


def gerar_data_extenso(d: datetime) -> str:
    """01/05/2026 → '1 de maio de 2026'"""
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    return f'{d.day} de {meses[d.month - 1]} de {d.year}'


def montar_dados_inicial(pasta_banco: str, autora: Dict, subsecao: str,
                         banco_jurisdicao: str = 'matriz',
                         numeros_contrato_explicitos: Optional[List[str]] = None) -> Dict:
    """Monta o dicionário completo de dados para preencher o template.

    Args:
        pasta_banco: ex. 'GEORGE/BANCO ITAÚ/2 AVERBAÇÃO NOVA INATIVO/'
        autora: dict com dados de qualificação (nome, cpf, rg, endereço, etc.)
        subsecao: 'Salvador', 'Manaus', 'Maceió', etc.
        banco_jurisdicao: 'matriz' (default) / 'AL' / 'AM' / 'BA'
        numeros_contrato_explicitos: opcional. Quando informado, sobrescreve
            a detecção pelo nome do arquivo da procuração / OCR.

    Returns:
        {
            'pasta_banco': str,
            'numeros_procuracoes': [str],
            'hiscon': dict do parse_hiscon,
            'contratos_questionados': [dict] formatados,
            'calculo': dict do parse_calculo,
            'banco_reu': dict do resolver_banco,
            'template': str (caminho),
            'alertas_seletor': [str],
            'dano_moral': dict do calcular_dano_moral,
            'audit_dm': dict do auditar_dano_moral,
            'eh_idoso': bool,
            'dados_template': dict {placeholder: valor} pronto para aplicar
        }
    """
    # 1. Hierarquia para definir números a impugnar:
    #   1. numeros_contrato_explicitos (chamador) — vence tudo
    #   2. JSON com origem manual/sugestao_automatica_revisada — vence procuração
    #      (revisão humana já decidiu, prevalece sobre arquivos físicos)
    #   3. Procurações na pasta_acao (nome do arquivo)
    #   4. JSON com origem sugestao_automatica (sem revisão) — só se sem procuração
    #   5. OCR do PDF das procurações
    #   6. Sem nada → ProcuracaoSemFiltroError (NUNCA "todos do banco")
    if numeros_contrato_explicitos:
        numeros = list(numeros_contrato_explicitos)
    else:
        numeros = []
        nums_json, origem_json = _numeros_de_contratos_impugnar_json(pasta_banco)
        if nums_json and origem_json in ('manual', 'sugestao_automatica_revisada'):
            numeros = nums_json
        else:
            numeros = listar_procuracoes(pasta_banco)
            if not numeros and nums_json:  # JSON automático como fallback
                numeros = nums_json

    if not numeros:
        info_proc = extrair_numeros_contrato_de_pasta(pasta_banco, usar_easyocr=True)
        numeros = info_proc['numeros_unicos']
        if not numeros:
            raise ProcuracaoSemFiltroError(
                f'🚨 IMPOSSÍVEL extrair números de contrato em {pasta_banco}. '
                f'Tentado: numeros_contrato_explicitos do chamador, '
                f'_estado_cliente.json (contratos_impugnar_ids), nome do arquivo '
                f'da procuração e OCR easyocr. AÇÃO: abrir os PDFs das procurações '
                f'manualmente, ler o número do contrato outorgado, e renomear o '
                f'arquivo para incluir o número (ex.: "2 - PROCURAÇÃO BANCO '
                f'0123456789.pdf"). NUNCA pegamos "todos os contratos do banco" '
                f'como fallback.'
            )

    # 2. Localizar HISCON
    hiscon_path = encontrar_pdf(pasta_banco, 'HISTÓRICO DE EMPRÉSTIMO')
    if not hiscon_path:
        raise FileNotFoundError(f'HISCON não encontrado em {pasta_banco}')

    # 3. Parsear HISCON
    hiscon = parse_hiscon(hiscon_path)

    # 4. Filtrar contratos questionados (FUZZY MATCH ativo — tolera 1 typo no
    # número de contrato no nome do arquivo da procuração).
    contratos_brutos = filtrar_contratos_por_numero(
        hiscon['contratos'], numeros, fuzzy_dist=1
    )
    if not contratos_brutos:
        raise RuntimeError(f'Nenhum dos contratos {numeros} encontrado no HISCON')

    # 4-bis. AUDITAR procurações vs HISCON. Esta auditoria detecta:
    # - typos no nome do arquivo da procuração (já corrigidos por fuzzy)
    # - contratos do mesmo banco no HISCON que NÃO foram referidos por
    #   nenhuma procuração (suspeitos por proximidade temporal/prefixo)
    # - procurações cujos números NÃO existem no HISCON (alerta crítico)
    # Os alertas vão para o relatório paralelo + console.
    audit_proc = None
    if contratos_brutos:
        banco_codigo_principal = contratos_brutos[0].get('banco_codigo', '')
        if banco_codigo_principal:
            audit_proc = auditar_procuracoes_vs_hiscon(
                hiscon['contratos'], numeros, banco_codigo_principal
            )

    # 5. Formatar contratos para o template
    contratos_fmt = [formatar_contrato_para_template(c) for c in contratos_brutos]

    # Patch C (2026-05-16) — Validador pré-geração: aborta se algum contrato
    # estiver com valor zero, qtd inválida, competência vazia ou data com
    # placeholder. Caso paradigma VILSON/BANRISUL.
    validar_contratos_obrigatorios(contratos_fmt)

    # 6. Localizar PDF de cálculo (opcional — sem, calcula valor da causa
    # pela fórmula soma_dobros + dano_moral)
    calc_path = (encontrar_pdf(pasta_banco, '10- CÁLCULO') or
                 encontrar_pdf(pasta_banco, '9- CÁLCULO') or
                 encontrar_pdf(pasta_banco, 'CÁLCULO') or
                 encontrar_pdf(pasta_banco, 'cálculo') or
                 encontrar_pdf(pasta_banco, 'calculo'))
    if calc_path:
        calculo = parse_calculo(calc_path)
    else:
        calculo = {
            'valor_total_geral': None,
            'dano_moral_pleiteado_pdf': None,
            'idade': None,
            'data_nascimento': None,
        }

    # 6.5. Localizar HISCRE — fonte mais confiável para CPF + renda LÍQUIDA
    hiscre_path = encontrar_pdf(pasta_banco, 'HISTÓRICO DE CRÉDITO')
    hiscre = {}
    if hiscre_path:
        try:
            hiscre = parse_hiscre(hiscre_path)
        except Exception as e:
            print(f'⚠ Erro ao parsear HISCRE: {e}')
            hiscre = {}

    # 7. Identificar banco-réu (pelo banco_nome do PRIMEIRO contrato — assumindo
    # que todos são do mesmo banco)
    banco_nome_hiscon = contratos_brutos[0].get('banco_nome', '')
    banco_reu = resolver_banco(banco_nome_hiscon, banco_jurisdicao)
    if not banco_reu:
        # Tentar com o código + matriz
        codigo = contratos_brutos[0].get('banco_codigo', '')
        banco_reu = resolver_banco(banco_nome_hiscon, 'matriz')
    if not banco_reu:
        raise RuntimeError(f'Banco réu não identificado: "{banco_nome_hiscon}"')

    # 8. Selecionar template
    template, alertas = selecionar_template(contratos_brutos)

    # 9. Calcular dano moral conforme regra + auditar contra PDF
    dm = calcular_dano_moral(contratos_brutos)
    audit_dm = auditar_dano_moral(contratos_brutos, calculo['dano_moral_pleiteado_pdf'])

    # 10. Idade
    eh_id = eh_idoso(calculo)

    # 11. Montar dicionário de placeholders
    # Patch 2026-05-16 (caso paradigma VILSON): pipeline BA caía em R$ 0,00
    # quando não havia PDF de cálculo. Agora usa `calcular_valor_causa_nc`
    # com a MESMA fórmula que o XLSX.
    vc = calculo.get('valor_total_geral')
    if not vc:
        try:
            import sys as _sys
            _common_dir = os.path.normpath(os.path.join(
                os.path.dirname(os.path.abspath(__file__)), '..', '..', '_common'))
            if _common_dir not in _sys.path:
                _sys.path.insert(0, _common_dir)
            from calculadora_indebito import calcular_valor_causa_nc
            calc_vc = calcular_valor_causa_nc(contratos_fmt)
            vc = calc_vc['valor_causa']
        except Exception:
            vc = 0.0  # fallback de segurança (template avisa via residual)
    inss_endereco = INSS_FIXO['enderecos_subsecao'].get(subsecao, INSS_FIXO['enderecos_subsecao']['Salvador'])

    # Lista compactada de contratos para a frase intro do MULT
    nums_fmt = [c['numero'] for c in contratos_fmt if c['numero']]
    if len(nums_fmt) == 1:
        contratos_lista_breve = f'CONTRATO Nº {nums_fmt[0]}'
    elif len(nums_fmt) == 2:
        contratos_lista_breve = f'CONTRATOS Nº {nums_fmt[0]} e {nums_fmt[1]}'
    else:
        contratos_lista_breve = f'CONTRATOS Nº {", ".join(nums_fmt[:-1])} e {nums_fmt[-1]}'

    # === HIERARQUIA DE FONTES + VERIFICAÇÃO CRUZADA (SKILL.md §9-bis) ===
    # 1. Comparar dados de AUTORA (doc físico) vs HISCRE
    divergencias_pessoais = comparar_doc_vs_hiscre(autora, hiscre)
    # 2. Consolidar: doc primário, HISCRE subsidiário
    autora_consolidada = consolidar_dados_autora(autora, hiscre)

    # CPF: usa o de AUTORA se preenchido (doc físico = primário); senão HISCRE
    cpf_final = (autora_consolidada.get('cpf') or '').strip()
    # Nome
    nome_final = autora_consolidada.get('nome') or hiscre.get('nome_autor') or ''
    # Renda LÍQUIDA: prioridade HISCRE (oficial INSS) > autora.renda_liquida > BASE DE CÁLCULO (BRUTO — alerta)
    renda_liquida = (hiscre.get('valor_liquido') or
                     autora.get('renda_liquida') or
                     hiscon['margens'].get('base_calculo'))
    fonte_renda = ('HISCRE (líquido)' if hiscre.get('valor_liquido')
                   else 'AUTORA' if autora.get('renda_liquida')
                   else 'HISCON BASE_CALCULO (BRUTO — confirmar)')
    # Tipo de benefício: HISCRE tem a descrição oficial; fallback para HISCON
    tipo_beneficio = (hiscre.get('especie_descricao') or
                      hiscon['cabecalho']['tipo_beneficio'] or '')
    # NB: HISCRE > HISCON
    nb_beneficio = hiscre.get('nb_beneficio') or hiscon['cabecalho']['nb_beneficio'] or ''

    dados_template = {
        '{{cidade_subsecao}}': subsecao,
        '{{uf_subsecao}}': 'BA',  # TODO: parametrizar conforme subsecao
        # Autor (com hierarquia doc > HISCRE)
        '{{nome_autor}}': nome_final,
        '{{nacionalidade}}': autora.get('nacionalidade', 'brasileiro'),
        '{{estado_civil}}': autora.get('estado_civil', ''),
        '{{profissao}}': autora.get('profissao', ''),
        '{{cpf_autor}}': cpf_final,
        '{{rg_autor}}': autora.get('rg', ''),
        '{{orgao_expedidor_autor}}': autora.get('orgao_expedidor', 'SSP/BA'),
        '{{logradouro_autor}}': autora.get('logradouro', ''),
        '{{numero_autor}}': autora.get('numero', ''),
        '{{bairro_autor}}': autora.get('bairro', ''),
        '{{cidade_autor}}': autora.get('cidade', ''),
        '{{uf_autor}}': autora.get('uf', 'BA'),
        '{{cep_autor}}': autora.get('cep', ''),
        # Banco réu (nome forçado em CAIXA ALTA)
        '{{banco_reu_nome}}': normalizar_banco_reu(banco_reu['nome']),
        '{{banco_reu_descricao_pj}}': banco_reu['descricao_pj'],
        '{{banco_reu_cnpj}}': banco_reu['cnpj'],
        '{{banco_reu_endereco}}': banco_reu['endereco'],
        '{{inss_endereco_subsecao}}': inss_endereco,
        # Blocos narrativos (Opção 4) — usados pelos templates AL/MG
        '{{BLOCO_CONTRATOS_FRAUDULENTOS}}': gerar_bloco_contratos_fraudulentos(
            contratos_fmt, normalizar_banco_reu(banco_reu['nome'])),
        '{{BLOCO_PEDIDO_DECLARACAO}}': gerar_bloco_pedido_declaracao(
            contratos_fmt, nb_beneficio),
        # Benefício
        '{{tipo_beneficio}}': tipo_beneficio.lower() if tipo_beneficio else '',
        '{{nb_beneficio}}': nb_beneficio,
        '{{banco_pagador}}': hiscon['cabecalho']['banco_pagador'] or '',
        '{{agencia_pagador}}': hiscon['cabecalho']['agencia_pagador'] or '',
        '{{conta_pagador}}': hiscon['cabecalho']['conta_pagador'] or '',
        # Renda LÍQUIDA (do HISCRE, NÃO do BASE DE CÁLCULO bruto do HISCON)
        '{{valor_renda_liquida}}': fmt_brl(renda_liquida),
        # Dano moral
        '{{dano_moral_total}}': fmt_brl(dm['total']),
        '{{dano_moral_total_extenso}}': extenso_moeda(dm['total']),
        '{{dano_moral_unitario}}': fmt_brl(dm['unitario']) if dm['unitario'] else '',
        '{{dano_moral_unitario_extenso}}': extenso_moeda(dm['unitario']) if dm['unitario'] else '',
        # Valor causa
        '{{valor_causa}}': fmt_brl(vc),
        '{{valor_causa_extenso}}': extenso_moeda(vc),
        # Cidade + data
        '{{cidade_protocolo}}': subsecao,
        '{{uf_protocolo}}': 'BA',
        '{{data_protocolo}}': gerar_data_extenso(datetime.today()),
        # Lista compacta (só MULT)
        '{{contratos_lista_breve}}': contratos_lista_breve,
        # Para template de 1 contrato (BASE/REFIN), placeholders por contrato:
        # (preenchidos com o ÚNICO contrato — para MULT, esses serão substituídos
        # ao DUPLICAR o bloco repetível)
    }

    # Para BASE/REFIN: placeholders por contrato preenchidos com o ÚNICO contrato
    if len(contratos_fmt) == 1:
        c = contratos_fmt[0]
        dados_template.update({
            '{{contrato_numero}}': c['numero'] or '',
            '{{contrato_banco}}': c['banco'] or '',
            '{{contrato_qtd_parcelas}}': str(c['qtd_parcelas']) if c['qtd_parcelas'] else '',
            '{{contrato_valor_parcela}}': c['valor_parcela_str'] or '',
            '{{contrato_valor_parcela_extenso}}': extenso_moeda(c['valor_parcela_float']) if c['valor_parcela_float'] else '',
            '{{contrato_valor_emprestado}}': c['valor_emprestado_str'] or '',
            '{{contrato_valor_emprestado_extenso}}': extenso_moeda(c['valor_emprestado_float']) if c['valor_emprestado_float'] else '',
            '{{contrato_data_inclusao}}': c['data_inclusao_str'] or '',
            '{{contrato_competencia_inicio}}': c['competencia_inicio_str'] or '',
            '{{contrato_competencia_fim}}': c['competencia_fim_str'] or '',
        })

    return {
        'pasta_banco': pasta_banco,
        'numeros_procuracoes': numeros,
        'hiscon': hiscon,
        'hiscre': hiscre,
        'autora_consolidada': autora_consolidada,
        'divergencias_pessoais': divergencias_pessoais,
        'fonte_renda': fonte_renda,
        'contratos_questionados': contratos_fmt,
        'contratos_brutos': contratos_brutos,
        'calculo': calculo,
        'banco_reu': banco_reu,
        'template': template,
        'alertas_seletor': alertas,
        'audit_procuracoes': audit_proc,  # auditoria procurações vs HISCON
        'dano_moral': dm,
        'audit_dm': audit_dm,
        'eh_idoso': eh_id,
        'dados_template': dados_template,
        'subsecao': subsecao,
    }


def gerar_inicial(dados_caso: Dict, output_path: str) -> Dict:
    """Aplica o template selecionado, preenchendo placeholders e duplicando
    blocos repetíveis quando necessário (template MULT).

    Returns:
        {'modificados': int, 'residuais': [str], 'output_path': str,
         'paragrafos_removidos': [str descrição]}
    """
    import shutil
    template = dados_caso['template']
    contratos = dados_caso['contratos_questionados']
    placeholders = dados_caso['dados_template']
    eh_idoso = dados_caso['eh_idoso']

    # Copiar template para destino
    shutil.copy(template, output_path)
    doc = Document(output_path)

    paragrafos_removidos = []

    eh_mult = 'multiplos' in template.lower()

    # === Remover marcadores de "BLOCO REPETÍVEL" do MULT ===
    # (parágrafos de nota visual que aparecem no template para conferência)
    if eh_mult:
        n_marcadores = _remover_marcadores_visuais(doc)
        if n_marcadores > 0:
            paragrafos_removidos.append(f'{n_marcadores} marcador(es) "BLOCO REPETÍVEL" removido(s)')

    # === Duplicar blocos repetíveis (MULT) ===
    if eh_mult and len(contratos) > 1:
        _duplicar_blocos_mult(doc, contratos)

    # === FIX 4: Tratar marcador {{SE_IDOSO}} ===
    # Se idoso → remover só o marcador (mantém parágrafo)
    # Se NÃO idoso → remover parágrafo INTEIRO
    n_remov = _processar_marcador_se_idoso(doc, eh_idoso)
    if n_remov > 0:
        if eh_idoso:
            paragrafos_removidos.append(f'{n_remov} marcador(es) {{{{SE_IDOSO}}}} consumido(s) (autor idoso — parágrafos mantidos)')
        else:
            paragrafos_removidos.append(f'{n_remov} parágrafo(s) com {{{{SE_IDOSO}}}} removido(s) (autor não-idoso)')

    # === Aplicar substituições (placeholders simples) ===
    modificados = 0
    for p in doc.paragraphs:
        for k, v in placeholders.items():
            if k in p.text:
                if substituir_in_run(p._element, {k: v}):
                    modificados += 1

    # === Reescrever PARÁGRAFO DE RECEBIMENTO: cartão magnético vs conta
    # bancária (regra fixa 07/05/2026, Gabriel). O template BA traz placeholders
    # {{banco_pagador}}, {{agencia_pagador}}, {{conta_pagador}} já substituídos
    # acima, MAS o texto fixo "depositado em conta bancária" não cabe quando
    # o autor recebe por cartão magnético. Helper monta a variante correta. ===
    autora_consol = dados_caso.get('autora_consolidada', {}) or {}
    hiscon_cab_p = (dados_caso.get('hiscon') or {}).get('cabecalho', {}) or {}
    for p in list(doc.paragraphs):
        t = p.text
        if 'recebe benefício previdenciário' in t and 'NB' in t:
            novo_recebimento = montar_paragrafo_recebimento_beneficio(
                hiscon_cab_p, autora_consol)
            substituir_paragrafo_completo(p._element, novo_recebimento, grifo=True)
            modificados += 1
            break

    # === Reescrever INTRO FÁTICA — agrupa por banco quando há ≥2 bancos
    # (litisconsórcio passivo). Helper compartilhado com AM/AL. ===
    banco_reu = dados_caso.get('banco_reu') or {}
    nome_banco = banco_reu.get('nome', '')
    nums_contratos = [c.get('numero', '') for c in contratos if c.get('numero')]
    if nome_banco and nums_contratos:
        from helpers_redacao import aplicar_intro_fatica
        for p in list(doc.paragraphs):
            if ('tomou conhecimento dos descontos referentes' in p.text
                or 'constatou a existência de descontos referentes' in p.text):
                aplicar_intro_fatica(p._element, contratos, nome_banco, grifo=True)
                modificados += 1
                break

    # === Reescrever PEDIDOS DECLARATÓRIOS — escolher empréstimo vs
    # refinanciamento conforme tipo_origem; 1 pedido por contrato ===
    nb = (dados_caso.get('hiscon') or {}).get('cabecalho', {}).get('nb_beneficio') or '___'
    if contratos:
        modificados += preencher_pedidos_declaratorios(
            doc, contratos, nb, grifo=True)

    # === Remover prioridade dos PEDIDOS quando não-idoso ===
    if not eh_idoso:
        modificados += remover_prioridade_pedidos(doc)

    # === Padronizar endereço escritório: matriz Joaçaba/SC + unidade apoio
    # da UF (atualiza trecho hardcoded do template para versão composta) ===
    uf_destino = (dados_caso.get('autora_consolidada', {}) or {}).get('uf') \
                 or dados_caso.get('subsecao_uf') or 'BA'
    modificados += inserir_unidade_apoio_se_faltando(doc, uf_destino, grifo=True)

    # === Garantir prioridade no CABEÇALHO + nos PEDIDOS quando idoso ===
    if eh_idoso:
        modificados += inserir_prioridade_idoso_se_faltando(doc, eh_idoso, grifo=True)
        modificados += inserir_pedido_prioridade_idoso_se_faltando(doc, eh_idoso, grifo=True)

    # === Verificar residuais ===
    residuais = []
    for p in doc.paragraphs:
        for ph in re.findall(r'\{\{[^}]+\}\}', p.text):
            if ph not in residuais:
                residuais.append(ph)

    doc.save(output_path)

    # Patch D (2026-05-16) — Validação pós-DOCX
    from helpers_docx import validar_docx_gerado
    validar_docx_gerado(output_path, abortar=True)

    return {
        'modificados': modificados,
        'residuais': residuais,
        'output_path': output_path,
        'paragrafos_removidos': paragrafos_removidos,
    }


def _remover_marcadores_visuais(doc) -> int:
    """Remove parágrafos de nota visual '⤵ BLOCO REPETÍVEL ...' inseridos no
    template MULT para comunicar a estrutura repetível ao leitor humano."""
    removidos = 0
    # Iterar de trás para frente para não bagunçar índices
    for p in list(doc.paragraphs):
        if 'BLOCO REPETÍVEL' in p.text and '⤵' in p.text:
            p._element.getparent().remove(p._element)
            removidos += 1
    return removidos


def _processar_marcador_se_idoso(doc, eh_idoso: bool) -> int:
    """Processa o marcador {{SE_IDOSO}} no início de parágrafos condicionais.

    - Se eh_idoso=True: remove APENAS o marcador (parágrafo mantido)
    - Se eh_idoso=False: REMOVE o parágrafo inteiro
    """
    n = 0
    W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    for p in list(doc.paragraphs):
        if '{{SE_IDOSO}}' not in p.text:
            continue
        if not eh_idoso:
            # Remover parágrafo inteiro
            p._element.getparent().remove(p._element)
            n += 1
        else:
            # Remover apenas o marcador do primeiro <w:t> que o contém
            for t in p._element.findall('.//' + W_NS + 't'):
                if t.text and '{{SE_IDOSO}}' in t.text:
                    t.text = t.text.replace('{{SE_IDOSO}}', '', 1)
                    n += 1
                    break
    return n


def _duplicar_blocos_mult(doc, contratos: List[Dict]):
    """Para o template MULT: localiza os parágrafos com {{contrato_numero}} e
    duplica N vezes (uma por contrato), substituindo placeholders por contrato
    em cada cópia.

    Atualmente o template MULT tem 2 parágrafos repetíveis:
    - p18 = "Do contrato nº {{contrato_numero}}: a primeira parcela..." (síntese)
    - p230 = "Declarar a inexistência do empréstimo consignado..." (pedido)

    Para cada um, duplica N-1 cópias logo após o original, e substitui os
    placeholders {{contrato_*}} por valores específicos. As substituições são
    GRIFADAS em amarelo (regra obrigatória da skill — ver SKILL.md §10).
    """
    from copy import deepcopy

    # Localizar parágrafos repetíveis pelo placeholder {{contrato_numero}}
    indices_repetiveis = []
    for i, p in enumerate(doc.paragraphs):
        if '{{contrato_numero}}' in p.text:
            indices_repetiveis.append(i)

    # Duplicar cada parágrafo repetível N-1 vezes
    n = len(contratos)
    for idx in sorted(indices_repetiveis, reverse=True):
        p_original = doc.paragraphs[idx]
        elem_original = p_original._element
        for _ in range(n - 1):
            copia = deepcopy(elem_original)
            elem_original.addnext(copia)

    # Re-listar índices (agora cada repetível tem N parágrafos consecutivos)
    indices_grupos = []
    i = 0
    paragrafos = doc.paragraphs
    while i < len(paragrafos):
        if '{{contrato_numero}}' in paragrafos[i].text:
            grupo = [i]
            j = i + 1
            while j < len(paragrafos) and '{{contrato_numero}}' in paragrafos[j].text:
                grupo.append(j)
                j += 1
            indices_grupos.append(grupo)
            i = j
        else:
            i += 1

    # Para cada grupo, substituir 1-a-1 com os contratos (grifo=True implícito)
    for grupo in indices_grupos:
        for k, idx_par in enumerate(grupo):
            if k >= len(contratos):
                break
            c = contratos[k]
            par = doc.paragraphs[idx_par]
            mapa = {
                '{{contrato_numero}}': c['numero'] or '',
                '{{contrato_banco}}': c['banco'] or '',
                '{{contrato_qtd_parcelas}}': str(c['qtd_parcelas']) if c['qtd_parcelas'] else '',
                '{{contrato_valor_parcela}}': c['valor_parcela_str'] or '',
                '{{contrato_valor_parcela_extenso}}': extenso_moeda(c['valor_parcela_float']) if c['valor_parcela_float'] else '',
                '{{contrato_valor_emprestado}}': c['valor_emprestado_str'] or '',
                '{{contrato_valor_emprestado_extenso}}': extenso_moeda(c['valor_emprestado_float']) if c['valor_emprestado_float'] else '',
                '{{contrato_data_inclusao}}': c['data_inclusao_str'] or '',
                '{{contrato_competencia_inicio}}': c['competencia_inicio_str'] or '',
                '{{contrato_competencia_fim}}': c['competencia_fim_str'] or '',
            }
            substituir_in_run(par._element, mapa, grifo=True)


def gerar_relatorio_paralelo(dados_caso: Dict, output_path: str):
    """Gera DOCX com pendências, alertas e checklist."""
    doc = Document()
    doc.add_heading(f'RELATÓRIO DE PENDÊNCIAS — {os.path.basename(dados_caso["pasta_banco"])}', level=1)

    p = doc.add_paragraph()
    p.add_run('Cliente: ').bold = True
    p.add_run(dados_caso['hiscon']['cabecalho']['nome_autor'] or '?')
    p = doc.add_paragraph()
    p.add_run('Banco-réu: ').bold = True
    p.add_run(dados_caso['banco_reu']['nome'])
    p = doc.add_paragraph()
    p.add_run('Subseção: ').bold = True
    p.add_run(f'{dados_caso["subsecao"]}/BA')
    p = doc.add_paragraph()
    p.add_run('Resumo do caso: ').bold = True
    p.add_run(descrever_caso(dados_caso['contratos_brutos']))

    # Resumo
    doc.add_heading('1. RESUMO', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    tbl.rows[0].cells[0].text = 'Campo'
    tbl.rows[0].cells[1].text = 'Valor'
    pdf = dados_caso['calculo']
    rows = [
        ('Procurações encontradas', f'{len(dados_caso["numeros_procuracoes"])} ({", ".join(dados_caso["numeros_procuracoes"])})'),
        ('Contratos no HISCON', f'{len(dados_caso["contratos_questionados"])}'),
        ('Idoso?', 'Sim' if dados_caso['eh_idoso'] else 'Não'),
        ('Idade (PDF cálculo)', str(pdf.get('idade'))),
        ('Total Geral (VC)', f'R$ {fmt_brl(pdf.get("valor_total_geral"))}'),
        ('Dano moral pleiteado (PDF)', f'R$ {fmt_brl(pdf.get("dano_moral_pleiteado_pdf"))}'),
        ('Dano moral pela regra do escritório', f'R$ {fmt_brl(dados_caso["dano_moral"]["total"])}'),
        ('Template selecionado', os.path.basename(dados_caso['template'])),
    ]
    for k, v in rows:
        row = tbl.add_row().cells
        row[0].text = k
        row[1].text = v

    # === SEÇÃO PRIORITÁRIA: DIVERGÊNCIAS DOC vs HISCRE ===
    divergencias = dados_caso.get('divergencias_pessoais', [])
    if divergencias:
        doc.add_heading('🚨 DIVERGÊNCIAS DOC vs HISCRE — REVISAR ANTES DE PROTOCOLAR', level=2)
        doc.add_paragraph(
            'A skill comparou os dados pessoais do KIT (documento físico) com o '
            'HISCRE (oficial do INSS). Divergências detectadas — pode ser '
            'documento de OUTRA pessoa na pasta, OCR errado, homônimo ou mudança '
            'de nome. CONFIRMAR antes de protocolar.', style='Intense Quote'
        )
        tbl_div = doc.add_table(rows=1, cols=4)
        tbl_div.style = 'Light Grid Accent 1'
        for i, h in enumerate(['Severidade', 'Campo', 'Doc físico (KIT)', 'HISCRE (INSS)']):
            tbl_div.rows[0].cells[i].text = h
        for d in divergencias:
            row = tbl_div.add_row().cells
            row[0].text = d['severidade']
            row[1].text = d['campo']
            row[2].text = str(d.get('doc') or '(não preenchido)')
            row[3].text = str(d.get('hiscre') or '(não preenchido)')
        for d in divergencias:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f'[{d["severidade"]}] {d["campo"]}: ').bold = True
            p.add_run(d['msg'])

    # Alertas demais
    todos_alertas = list(dados_caso['alertas_seletor'])
    if dados_caso['audit_dm']['alerta']:
        todos_alertas.insert(0, dados_caso['audit_dm']['alerta'])
    # Alertas da auditoria de procurações vs HISCON (typos, contratos faltantes)
    audit_p = dados_caso.get('audit_procuracoes')
    if audit_p:
        for a in audit_p.get('alertas', []):
            todos_alertas.append(a)
    # Alertas de qualidade do HISCRE (Patch 2026-05-16, caso VILSON):
    # competência muito antiga, valor líquido faltando, etc.
    for a in (dados_caso.get('hiscre') or {}).get('alertas_qualidade') or []:
        todos_alertas.append(a)
    # Auditoria procurações órfãs (Patch 2026-05-16, caso VILSON 2ª Banrisul)
    pasta_banco_caso = dados_caso.get('pasta_banco') or ''
    if pasta_banco_caso:
        try:
            import sys as _sys
            skill_dir = r'C:\Users\gabri\.claude\skills\kit-juridico\scripts'
            if skill_dir not in _sys.path:
                _sys.path.insert(0, skill_dir)
            from auditor_procuracoes_orfas import auditar_cliente  # type: ignore
            pasta_cliente_aud = os.path.dirname(os.path.dirname(
                os.path.abspath(pasta_banco_caso)))
            if os.path.isdir(os.path.join(pasta_cliente_aud, '0. Kit')):
                rel_aud = auditar_cliente(pasta_cliente_aud)
                if 'erro' not in rel_aud:
                    for orfa in rel_aud.get('orfas') or []:
                        todos_alertas.append(
                            f'🚨 PROCURAÇÃO ÓRFÃ: pág {orfa["pagina"]} · '
                            f'banco {orfa["banco"]} · contrato '
                            f'`{orfa["contrato"]}`. AÇÃO: criar pasta-banco e '
                            f'rodar inicial separada.'
                        )
        except Exception:
            pass
    if 'BRUTO' in dados_caso.get('fonte_renda', ''):
        todos_alertas.append(
            f'Renda usada veio do BASE DE CÁLCULO do HISCON (valor BRUTO). '
            f'O HISCRE não estava disponível. CONFERIR — o ideal é usar o '
            f'valor LÍQUIDO depositado em conta.'
        )
    if not dados_caso['eh_idoso']:
        todos_alertas.append('Cliente NÃO é idoso (>=60). A skill removeu automaticamente '
                             'os parágrafos de prioridade idoso (cabeçalho + pedido). Verificar '
                             'se cabe pedido de prioridade por outra hipótese (deficiência, '
                             'doença grave, etc.).')

    doc.add_heading('2. PENDÊNCIAS / ALERTAS', level=2)
    if not todos_alertas:
        doc.add_paragraph('Nenhuma pendência crítica detectada. Conferir checklist abaixo.')
    else:
        for a in todos_alertas:
            doc.add_paragraph(a, style='List Bullet')

    # Checklist
    doc.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
    items = [
        'Conferir nome / CPF / RG / endereço da parte autora.',
        'Conferir CNPJ + endereço do banco-réu.',
        'Conferir lista de contratos questionados (procurações).',
        'Conferir valor da causa contra PDF de cálculo.',
        'Conferir dano moral pleiteado (regra vs PDF).',
        'Verificar idade e necessidade de remover bloco de prioridade idoso (se não for >= 60).',
        'Anexar HISCON, HISCRE, RG, CPF, comp. residência, declaração hipossuficiência, cálculo, procurações.',
        'Verificar competência: subseção JEF correta?',
    ]
    if dados_caso['alertas_seletor']:
        items.insert(0, 'REVISAR ALERTAS DA SEÇÃO 2 (refinanciamentos / contratos ativos / outros).')
    for it in items:
        doc.add_paragraph(it, style='List Number')

    p = doc.add_paragraph()
    p.add_run('Conclusão: ').bold = True
    if todos_alertas:
        p.add_run('APTA COM RESSALVAS — ').italic = True
        p.add_run('revisar alertas antes do protocolo.').bold = True
    else:
        p.add_run('APTA — ').italic = True
        p.add_run('protocolar após conferência do checklist.').bold = True

    doc.save(output_path)
