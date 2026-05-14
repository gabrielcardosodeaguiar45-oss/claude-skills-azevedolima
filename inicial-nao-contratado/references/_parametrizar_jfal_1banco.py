# -*- coding: utf-8 -*-
"""Parametriza inicial-jfal-1banco.docx (AL Federal / 1 banco + INSS).
- Mantém INSS como texto fixo no par 10 (já presente no template).
- Parametriza endereço subseção INSS via {{inss_endereco_subsecao}}.
- Adota Opção 4 (BLOCO_CONTRATOS / BLOCO_PEDIDO).
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

CAMINHO = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jfal-1banco.docx")

backup = CAMINHO.with_suffix(CAMINHO.suffix + ".bak_pre_parametrizacao")
if not backup.exists():
    shutil.copy2(CAMINHO, backup)
    print(f"Backup: {backup.name}")

SUBS = {
    # PAR 0 — cabeçalho federal (Subseção em vez de Comarca)
    0: {
        "Subseção de  /AL":  "Subseção de {{cidade_protocolo}}/{{uf_protocolo}}",
    },
    # PAR 6 — qualificação FULANO
    6: {
        "FULANO DE TAL":                          "{{nome_autor}}",
        "brasileiro, casado, aposentado":         "{{nacionalidade}}, {{estado_civil}}, {{profissao}}",
        "xxx.xxx.xxx-xx":                          "{{cpf_autor}}",
        "xxxxxxxx":                                 "{{rg_autor}}",
        "SSP/SC":                                   "{{orgao_expedidor}}",
        "rua tal":                                  "{{logradouro_autor}}",
        "n° xx, bairro tal, em cidade/BA, CEP xxxxx-xxx":
            "nº {{numero_autor}}, {{bairro_autor}}, em {{cidade_autor}}/{{uf_autor}}, CEP {{cep_autor}}",
    },
    # PAR 10 — banco réu BRADESCO + INSS (mantém texto fixo do INSS,
    # parametriza endereço da subseção)
    10: {
        "BANCO BRADESCO S/A":                     "{{banco_reu_nome}}",
        "60.746.948/0320-73":                      "{{banco_reu_cnpj}}",
        "Avenida Sete de Setembro, nº 895, Centro, Salvador/BA, CEP 69.005-140":
            "{{banco_reu_endereco}}",
        "Av. Sete de Setembro, 1078 - Mercês, Salvador/BA":
            "{{inss_endereco_subsecao}}",
    },
    # PAR 12 — benefício
    12: {
        "aposentadoria por idade":                "{{tipo_beneficio}}",
        "149.139.433-9":                           "{{nb_beneficio}}",
        "agência 3706, conta corrente nº 0000211974":
            "agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}",
        "junto ao BANCO BRADESCO SA":             "junto ao {{banco_pagador}}",
    },
    # PAR 15 — constatação
    15: {
        "Nessa oportunidade, após informações, tomou conhecimento dos descontos referentes à um empréstimo que não contratou junto ao BANCO BRADESCO S/A, CONTRATO Nº 0123506012709:":
            "Nessa oportunidade, após informações, tomou conhecimento dos descontos referentes a empréstimo(s) que não contratou junto ao {{banco_reu_nome}}, conforme detalhamento abaixo:",
    },
    # PAR 257 — prioridade
    257: {
        "A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de Processo Civil":
            "{{pedido_prioridade}}",
    },
    # PAR 271 — pedido (vira BLOCO)
    271: {
        "do empréstimo/refinanciamento consignado no valor de R$ 1.394,78 (um mil, trezentos e noventa e quatro reais e setenta e oito centavos), contrato nº 0123506012709 - com descontos de R$ 31,21 (trinta e um reais e vinte e um centavos) mensais, com inclusão em 24/07/2024, início de desconto em 01/08/2024, no benefício previdenciário 149.139.433-9":
            "{{BLOCO_PEDIDO_DECLARACAO}}",
    },
    # PAR 291 — fecho
    291: {
        "Cidade/AL":  "{{cidade_protocolo}}/{{uf_protocolo}}",
    },
}

doc = Document(CAMINHO)
for par_idx, mapa in SUBS.items():
    p = doc.paragraphs[par_idx]
    if substituir_in_run(p._p, mapa):
        print(f"  par {par_idx}: substituído")

# === PAR 17 — substituir conteúdo inteiro pelo marcador BLOCO ===
p17 = doc.paragraphs[17]
SUB17 = {
    "No que diz respeito ao referido empréstimo, cumpre informar que a primeira parcela descontada do benefício da parte autora foi na competência xxxxxxxx, de um total de xx parcelas, no valor de R$ xxx,xx (valor por extenso), relativas a um empréstimo consignado no valor de R$ xxx,xx (valor por extenso), contrato n° xxxxxxx, cuja operação foi realizada pelo banco xxxxx, ora requerido.":
        "{{BLOCO_CONTRATOS_FRAUDULENTOS}}"
}
if substituir_in_run(p17._p, SUB17):
    print("  par 17: virou {{BLOCO_CONTRATOS_FRAUDULENTOS}}")

# === Tratar par 10: aplicar Segoe UI no run que contém {{banco_reu_nome}} ===
# (smallCaps removido como no jeal-1banco)
p10 = doc.paragraphs[10]
for r in p10._p.findall('.//' + W + 'r'):
    txt = ''.join((t.text or '') for t in r.findall(W + 't'))
    if "{{banco_reu_nome}}" in txt:
        rpr = r.find(W + 'rPr')
        if rpr is not None:
            sc = rpr.find(W + 'smallCaps')
            if sc is not None:
                rpr.remove(sc)

doc.save(CAMINHO)
print()

# Verificar
import re
d2 = Document(CAMINHO)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r"\{\{[^}]+\}\}", p.text):
        phs.add(m.group(0))
print(f"Placeholders únicos: {len(phs)}")
for ph in sorted(phs):
    print(f"  {ph}")

print("\nDados pilotos remanescentes:")
sobras = []
for i, p in enumerate(d2.paragraphs):
    t = p.text
    for ruim in ["FULANO", "xxx.xxx.xxx-xx", "rua tal", "bairro tal",
                  "BANCO BRADESCO", "60.746.948", "149.139.433",
                  "0123506012709", "Cidade/AL", "Sete de Setembro, nº 895",
                  "Mercês, Salvador", "1.394,78", "31,21", "24/07/2024",
                  "Avenida Sete de Setembro"]:
        if ruim in t:
            sobras.append((i, ruim, t[:120]))
            break
if sobras:
    for i, r, t in sobras:
        print(f"  ⚠️ par {i}: '{r}' em: {t}")
else:
    print("  (vazio = OK)")
