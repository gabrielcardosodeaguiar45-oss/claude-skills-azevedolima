# -*- coding: utf-8 -*-
"""Parametriza inicial-jfal-2bancos.docx — Federal AL com INSS, dados pilotos
da EXEMPLA DA SILVA + Banco C6 + 2 contratos hardcoded.

Nota: par 6 original NÃO tem RG/órgão expedidor (só CPF). Mantido assim.
Se precisar adicionar RG, fazer manualmente após esta parametrização.
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

CAMINHO = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jfal-2bancos.docx")

backup = CAMINHO.with_suffix(CAMINHO.suffix + ".bak_pre_parametrizacao")
if not backup.exists():
    shutil.copy2(CAMINHO, backup)
    print(f"Backup: {backup.name}")

SUBS = {
    # PAR 0 — cabeçalho federal
    0: {"Subseção de Arapiraca/AL":  "Subseção de {{cidade_protocolo}}/{{uf_protocolo}}"},
    # PAR 6 — qualificação LOURDES (sem RG no template original)
    6: {
        "EXEMPLA DA SILVA":               "{{nome_autor}}",
        "brasileira, casada, aposentada":         "{{nacionalidade}}, {{estado_civil}}, {{profissao}}",
        "000.000.002-12":                          "{{cpf_autor}}",
        "Vila Santa Izabel":                       "{{logradouro_autor}}",
        "nº 2, Zona Rural":                        "nº {{numero_autor}}, {{bairro_autor}}",
        "em Lagoa da Canoa/AL":                   "em {{cidade_autor}}/{{uf_autor}}",
        "57.330-000":                              "{{cep_autor}}",
    },
    # PAR 10 — banco C6 réu + INSS
    10: {
        "BANCO C6 CONSIGNADO S.A.":               "{{banco_reu_nome}}",
        "61.348.538/0001-86":                      "{{banco_reu_cnpj}}",
        "Avenida Nove de Julho, 3148, Jardim Paulista, São Paulo/SP, CEP 01406-000": "{{banco_reu_endereco}}",
        "Av. Sete de Setembro, 1078 - Mercês, Salvador/BA": "{{inss_endereco_subsecao}}",
    },
    # PAR 12 — benefício
    12: {
        "aposentadoria por idade":                "{{tipo_beneficio}}",
        "168.722.319-7":                           "{{nb_beneficio}}",
        "agência 3169, conta corrente nº 000098681P": "agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}",
        "junto ao BANCO BRADESCO SA":             "junto ao {{banco_pagador}}",
    },
    # PAR 15 — constatação
    15: {
        "Nessa oportunidade, após informações, tomou conhecimento dos descontos referentes à empréstimos que não contratou junto ao BANCO C6 CINSIGNADO S/A, CONTRATOS Nº 903234 5962 e 902673 2183:":
            "Nessa oportunidade, após informações, tomou conhecimento dos descontos referentes a empréstimo(s) que não contratou junto ao {{banco_reu_nome}}, conforme detalhamento abaixo:",
    },
    # PAR 46 — renda
    46: {
        "R$ 996,40":  "R$ {{valor_renda_liquida}} ({{valor_renda_liquida_extenso}})",
    },
    # PAR 262 — prioridade
    262: {
        "A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de\xa0Processo\xa0Civil":
            "{{pedido_prioridade}}",
    },
    # PAR 295 — valor causa
    295: {
        "R$ 29.193,73 (vinte e nove mil, cento e noventa e três reais e setenta e três centavos)":
            "R$ {{valor_causa}} ({{valor_causa_extenso}})",
    },
}

doc = Document(CAMINHO)
for par_idx, mapa in SUBS.items():
    if par_idx >= len(doc.paragraphs):
        print(f"  ⚠️ par {par_idx} fora do range")
        continue
    p = doc.paragraphs[par_idx]
    if substituir_in_run(p._p, mapa):
        print(f"  par {par_idx}: substituído")
    else:
        print(f"  par {par_idx}: NÃO bateu — txt: {p.text[:120]!r}")

# === PAR 17 — vira intro do bloco (substitui inteiro pelo marcador) ===
p17 = doc.paragraphs[17]
SUB17 = {"No que diz respeito aos referidos empréstimos, cumpre informar que:":
         "{{BLOCO_CONTRATOS_FRAUDULENTOS}}"}
if substituir_in_run(p17._p, SUB17):
    print("  par 17: virou {{BLOCO_CONTRATOS_FRAUDULENTOS}}")

# Pars 18 e 19: esvaziar (já absorvidos pelo BLOCO no par 17)
for pidx in [18, 19]:
    p = doc.paragraphs[pidx]
    for r in list(p._p.findall(W + 'r')):
        p._p.remove(r)
    print(f"  par {pidx}: esvaziado")

# PAR 276 — vira pedido único com BLOCO
p276 = doc.paragraphs[276]
SUB276 = {"do empréstimo consignado no valor de R$ 1.327,78, contrato nº 903234 5962- com descontos de R$ 32,20 mensais, com inclusão em 12/02/2024, início de desconto em 03/2024, no benefício previdenciário 168.722.319-7":
          "{{BLOCO_PEDIDO_DECLARACAO}}"}
if substituir_in_run(p276._p, SUB276):
    print("  par 276: virou {{BLOCO_PEDIDO_DECLARACAO}}")

# PAR 277 — esvaziar (segundo pedido absorvido pelo BLOCO)
p277 = doc.paragraphs[277]
for r in list(p277._p.findall(W + 'r')):
    p277._p.remove(r)
print("  par 277: esvaziado")

# Remover smallCaps no run do {{banco_reu_nome}} no par 10
p10 = doc.paragraphs[10]
for r in p10._p.findall('.//' + W + 'r'):
    txt = ''.join((t.text or '') for t in r.findall(W + 't'))
    if "{{banco_reu_nome}}" in txt:
        rpr = r.find(W + 'rPr')
        if rpr is not None:
            sc = rpr.find(W + 'smallCaps')
            if sc is not None:
                rpr.remove(sc)

doc.save(CAMINHO)
print()

import re
d2 = Document(CAMINHO)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r"\{\{[^}]+\}\}", p.text):
        phs.add(m.group(0))
print(f"Placeholders únicos: {len(phs)}")
for ph in sorted(phs):
    print(f"  {ph}")

print("\nDados pilotos remanescentes (excluindo ementas):")
sobras = []
for i, p in enumerate(d2.paragraphs):
    t = p.text
    for ruim in ["LOURDES", "926.906.964", "Vila Santa", "Lagoa da Canoa",
                  "BANCO C6 CONSIGNADO S", "61.348.538", "168.722.319",
                  "Mercês, Salvador", "996,40", "29.193,73",
                  "903234 5962", "902673 2183", "Subseção de Arapiraca"]:
        if ruim in t and not any(em in t for em in ["DIREITO CIVIL", "Apelação", "RECURSO"]):
            sobras.append((i, ruim, t[:100]))
            break
if sobras:
    for i, r, t in sobras:
        print(f"  ⚠️ par {i}: '{r}' em: {t}")
else:
    print("  (vazio = OK)")
