"""Regressão para o bug 2026-05-10:
- Template uniformizado em {{NOME_COMPLETO}} (UPPERCASE)
- Dict com chave 'nome_completo' (lowercase)
- ANTES: o motor abortava o parágrafo no primeiro placeholder não encontrado,
  deixando bloco de qualificação inteiro cru.
- AGORA: lookup é case-insensitive, placeholder desconhecido é mascarado mas
  outros do mesmo parágrafo são resolvidos, e aplicar_template levanta
  PlaceholdersResiduaisError se sobrar qualquer {{...}} no fim.

Para rodar:
    cd skills/inicial-bradesco
    python -m pytest tests/test_placeholders_residuais.py -v
ou simplesmente:
    python tests/test_placeholders_residuais.py
"""
import os, sys, tempfile, zipfile
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'references'))

from docx import Document

import helpers_docx
from helpers_docx import (
    aplicar_template,
    PlaceholdersResiduaisError,
    processar_paragrafo,
)
from lxml import etree

W = helpers_docx.W


def _docx_minimo(placeholders, dst):
    """Cria um docx minimal com um único parágrafo contendo os placeholders."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run(' '.join(placeholders))
    doc.save(dst)


def test_lookup_case_insensitive(tmpdir=None):
    """{{NOME_COMPLETO}} deve casar com a chave 'nome_completo' do dict."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl.docx')
    out = os.path.join(tmpdir, 'out.docx')
    _docx_minimo(['{{NOME_COMPLETO}}', '{{cpf}}'], template)

    res = aplicar_template(template, {'nome_completo': 'JOÃO DA SILVA', 'cpf': '111.222.333-44'}, out)
    assert res['residuais'] == [], f'esperado zero residuais, veio {res["residuais"]}'
    txt = '\n'.join(p.text for p in Document(out).paragraphs)
    assert 'JOÃO DA SILVA' in txt
    assert '111.222.333-44' in txt
    assert '{{' not in txt
    print('OK case-insensitive')


def test_skip_unknown_continua(tmpdir=None):
    """{{INEXISTENTE}} no MEIO do parágrafo NÃO impede que {{cpf}} seja resolvido."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl.docx')
    out = os.path.join(tmpdir, 'out_falhou.docx')

    doc = Document()
    p = doc.add_paragraph()
    p.add_run('{{NOME_COMPLETO}} {{INEXISTENTE}} CPF {{cpf}}')
    doc.save(template)

    try:
        aplicar_template(template, {'nome_completo': 'X', 'cpf': '999.000.111-22'}, out)
        assert False, 'devia ter levantado PlaceholdersResiduaisError'
    except PlaceholdersResiduaisError as e:
        # Deve ter renomeado para _FALHOU_PLACEHOLDERS
        assert os.path.exists(e.dst_path), f'falha esperada em {e.dst_path}'
        assert '_FALHOU_PLACEHOLDERS' in e.dst_path
        # Deve listar APENAS o placeholder desconhecido como residual
        assert e.residuais == ['INEXISTENTE'], (
            f'esperado [INEXISTENTE], veio {e.residuais}'
        )
        # E o resto deve ter sido substituído (NÃO ABORTOU o parágrafo!)
        txt = '\n'.join(par.text for par in Document(e.dst_path).paragraphs)
        assert 'X' in txt, 'NOME_COMPLETO devia ter sido substituído por "X"'
        assert '999.000.111-22' in txt, 'cpf devia ter sido substituído'
        assert '{{INEXISTENTE}}' in txt, 'placeholder desconhecido devia ter sido restaurado'
        print('OK skip-unknown-continua')


def test_strict_raise(tmpdir=None):
    """Por padrão (strict=True) DEVE levantar exceção em residuais."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl.docx')
    out = os.path.join(tmpdir, 'out_strict.docx')
    _docx_minimo(['{{NUNCA_EXISTIRA}}'], template)

    try:
        aplicar_template(template, {'nome_completo': 'X'}, out)
        assert False, 'deveria ter levantado'
    except PlaceholdersResiduaisError as e:
        assert 'NUNCA_EXISTIRA' in e.residuais
        print('OK strict-raise')


def test_strict_false_devolve(tmpdir=None):
    """Com strict=False, devolve o dict mas não levanta — modo legado."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl.docx')
    out = os.path.join(tmpdir, 'out_lazy.docx')
    _docx_minimo(['{{NUNCA_EXISTIRA}}'], template)

    res = aplicar_template(template, {'x': 'y'}, out, strict=False)
    assert 'NUNCA_EXISTIRA' in res['residuais']
    assert os.path.exists(out)
    print('OK strict=False')


def test_qualificacao_completa(tmpdir=None):
    """Cenário real do bug: bloco de qualificação inteiro deve ser preenchido."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl_qual.docx')
    out = os.path.join(tmpdir, 'out_qual.docx')

    doc = Document()
    p = doc.add_paragraph()
    p.add_run(
        '{{NOME_COMPLETO}}, {{nacionalidade}}, {{profissao}}, '
        'inscrito(a) no CPF sob o nº {{cpf}}, '
        'Cédula de Identidade nº {{rg}}{{orgao_expedidor_prefixo}}, '
        'residente e domiciliado(a) à {{logradouro}}, nº {{numero}}, '
        '{{bairro}}, Município de {{cidade_de_residencia}}, CEP {{cep}}'
    )
    doc.save(template)

    dados = {
        'nome_completo': 'MARIA DA SILVA',
        'nacionalidade': 'brasileira',
        'profissao': 'aposentada',
        'cpf': '111.222.333-44',
        'rg': '12345-6',
        'orgao_expedidor_prefixo': ' SSP/AM',
        'logradouro': 'Rua A',
        'numero': '100',
        'bairro': 'Centro',
        'cidade_de_residencia': 'Maués',
        'cep': '69.190-000',
    }
    res = aplicar_template(template, dados, out)
    assert res['residuais'] == [], f'sobrou: {res["residuais"]}'
    txt = '\n'.join(p.text for p in Document(out).paragraphs)
    for valor in dados.values():
        assert valor in txt, f'{valor!r} não está no docx final'
    assert '{{' not in txt
    print('OK qualificacao-completa')


def test_orfao_cedula_removido(tmpdir=None):
    """Quando RG é OPCIONAL omitido, "Cédula de Identidade nº " órfão deve sumir."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl_orfao.docx')
    out = os.path.join(tmpdir, 'out_orfao.docx')

    doc = Document()
    p = doc.add_paragraph()
    p.add_run(
        '{{NOME_COMPLETO}}, {{nacionalidade}}, '
        'inscrito(a) no CPF sob o nº {{cpf}}, '
        'Cédula de Identidade nº {{rg}}{{orgao_expedidor_prefixo}}, '
        'residente e domiciliado(a) à {{logradouro}}.'
    )
    doc.save(template)

    dados = {
        'nome_completo': 'JOÃO',
        'nacionalidade': 'brasileiro',
        'cpf': '111.222.333-44',
        'rg': '',  # OPCIONAL omitido
        'orgao_expedidor_prefixo': '',
        'logradouro': 'Rua A',
    }
    res = aplicar_template(template, dados, out)
    assert res['residuais'] == [], f'sobrou: {res["residuais"]}'
    txt = '\n'.join(p.text for p in Document(out).paragraphs)
    assert 'Cédula de Identidade nº residente' not in txt, (
        f'órfão NÃO foi removido: ...{txt}...'
    )
    assert 'Cédula de Identidade nº ,' not in txt
    # texto correto: "JOÃO, brasileiro, inscrito(a) no CPF sob o nº 111.222.333-44, residente e domiciliado(a) à Rua A."
    assert 'inscrito(a) no CPF sob o nº 111.222.333-44, residente e domiciliado(a) à Rua A' in txt
    print('OK orfão-cédula-removido')


def test_dedup_rubrica(tmpdir=None):
    """Runs adjacentes com mesmo texto separados por '; ' devem dedupar."""
    tmpdir = tmpdir or tempfile.mkdtemp()
    template = os.path.join(tmpdir, 'tpl_dedup.docx')
    out = os.path.join(tmpdir, 'out_dedup.docx')

    doc = Document()
    p = doc.add_paragraph()
    p.add_run('Rubricas: “')
    p.add_run('{{rubrica_a}}')
    p.add_run('; ')
    p.add_run('{{rubrica_a}}')
    p.add_run('; ')
    p.add_run('{{rubrica_a}}')
    p.add_run('; ')
    p.add_run('{{rubrica_b}}')
    p.add_run('”')
    doc.save(template)

    res = aplicar_template(template, {'rubrica_a': 'ENC LIM CRÉDITO', 'rubrica_b': 'MORA CRED PESS'}, out)
    assert res['residuais'] == [], f'sobrou: {res["residuais"]}'
    txt = '\n'.join(p.text for p in Document(out).paragraphs)
    assert txt.count('ENC LIM CRÉDITO') == 1, (
        f'rubrica duplicada não foi dedupada: count={txt.count("ENC LIM CRÉDITO")}, txt={txt!r}'
    )
    assert 'ENC LIM CRÉDITO; MORA CRED PESS' in txt or 'ENC LIM CRÉDITO;MORA CRED PESS' in txt
    print('OK dedup-rubrica')


if __name__ == '__main__':
    tmp = tempfile.mkdtemp()
    test_lookup_case_insensitive(tmp)
    test_skip_unknown_continua(tmp)
    test_strict_raise(tmp)
    test_strict_false_devolve(tmp)
    test_qualificacao_completa(tmp)
    test_orfao_cedula_removido(tmp)
    test_dedup_rubrica(tmp)
    print('\nTODOS OS TESTES PASSARAM.')
