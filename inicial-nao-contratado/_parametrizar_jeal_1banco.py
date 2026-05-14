# -*- coding: utf-8 -*-
"""Parametriza inicial-jeal-1banco.docx (AL Estadual / 1 banco):
substitui dados pilotos hardcoded por placeholders canônicos,
preservando formatação (fonte/layout) via substituir_in_run.

Backup: <arquivo>.bak_pre_parametrizacao
"""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\inicial-bradesco\references")
from docx import Document
from helpers_docx import substituir_in_run

CAMINHO = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jeal-1banco.docx")

# Backup
backup = CAMINHO.with_suffix(CAMINHO.suffix + ".bak_pre_parametrizacao")
if not backup.exists():
    shutil.copy2(CAMINHO, backup)
    print(f"Backup criado: {backup.name}")

# Mapa de substituições — blocos LONGOS para ancoragem segura.
# Ordem: do mais específico (longo) ao mais genérico.
MAPA = {
    # ----- PAR 0 (cabeçalho) -----
    "Ao Juízo da ___ Vara Cível da Comarca de Cidade/AL":
        "Ao Juízo da {{vara_protocolo}} Vara Cível da Comarca de {{cidade_protocolo}}/{{uf_protocolo}}",

    # ----- PAR 6 (qualificação do autor) — bloco longo com ancoragem -----
    "FULANO DE TAL, brasileiro, casado, aposentado, inscrito no CPF sob o nº xxx.xxx.xxx-xx, Cédula de Identidade sob nº xxxxxxxx, órgão expedidor SSP/AL, residente e domiciliado à rua tal, n° xx, bairro tal, em cidade/AL, CEP xxxxx-xxx":
        "{{nome_autor}}, {{nacionalidade}}, {{estado_civil}}, {{profissao}}, inscrito no CPF sob o nº {{cpf_autor}}, Cédula de Identidade sob nº {{rg_autor}}, órgão expedidor {{orgao_expedidor}}, residente e domiciliado à {{logradouro_autor}}, nº {{numero_autor}}, {{bairro_autor}}, em {{cidade_autor}}/{{uf_autor}}, CEP {{cep_autor}}",

    # ----- PAR 10 (banco réu — qualificação) -----
    "em face de BANCO PAN S/A, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 59.285.411/0006-28, com endereço na Av. Ephigenio Salles, nº 1.327, bairro Aleixo, Manaus, CEP 69.060-020":
        "em face de {{banco_reu_nome}}, {{banco_reu_descricao_pj}}, inscrita no CNPJ sob o nº {{banco_reu_cnpj}}, com endereço na {{banco_reu_endereco}}",

    # ----- PAR 12 (benefício INSS) -----
    "A parte autora recebe benefício previdenciário de aposentadoria por idade – NB 149.139.433-9, o qual é depositado em conta bancária, agência 3706, conta corrente nº 0000211974, junto ao banco bradesco sa":
        "A parte autora recebe benefício previdenciário de {{tipo_beneficio}} – NB {{nb_beneficio}}, o qual é depositado em conta bancária, agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}, junto ao {{banco_pagador}}",

    # ----- PAR 14 (constatação de fraude) -----
    "empréstimo que não contratou junto ao BANCO PAN S/A, CONTRATO Nº 3880089838":
        "empréstimo que não contratou junto ao {{banco_reu_nome}}, CONTRATO Nº {{contrato_numero}}",

    # ----- PAR 16 (detalhe do contrato fraudulento) -----
    "a primeira parcela descontada do benefício da parte autora foi na competência 01/06/2024, de um total de 84 parcelas, no valor de R$ 49,00 (quarenta e nove reais), relativas a um empréstimo consignado no valor de R$ 2.171,24 (dois mil, cento e setenta e um reais e vinte e quatro centavos), contrato nº 3880089838, cuja operação foi realizada pelo BANCO PAN S/A, ora requerido.":
        "a primeira parcela descontada do benefício da parte autora foi na competência {{contrato_competencia_inicio}}, de um total de {{contrato_qtd_parcelas}} parcelas, no valor de R$ {{contrato_valor_parcela}} ({{contrato_valor_parcela_extenso}}), relativas a um empréstimo consignado no valor de R$ {{contrato_valor_emprestado}} ({{contrato_valor_emprestado_extenso}}), contrato nº {{contrato_numero}}, cuja operação foi realizada pelo {{banco_reu_nome}}, ora requerido.",

    # ----- PAR 44 (renda líquida) -----
    "oriunda de benefício previdenciário no valor líquido de R$ 1.093,99 (um mil e noventa e três reais e noventa e nove centavos)":
        "oriunda de benefício previdenciário no valor líquido de R$ {{valor_renda_liquida}} ({{valor_renda_liquida_extenso}})",

    # ----- PAR 228 (pedido — declaração de inexistência) -----
    "Declarar a inexistência do empréstimo consignado no valor de R$ 2.171,24 (dois mil, cento e setenta e um reais e vinte e quatro centavos), contrato nº 3880089838 - com descontos de R$ 49,00 (quarenta e nove reais) mensais, com inclusão em 31/05/2024, início de desconto em 01/06/2024, no benefício previdenciário 149.139.433-9":
        "Declarar a inexistência do empréstimo consignado no valor de R$ {{contrato_valor_emprestado}} ({{contrato_valor_emprestado_extenso}}), contrato nº {{contrato_numero}} - com descontos de R$ {{contrato_valor_parcela}} ({{contrato_valor_parcela_extenso}}) mensais, com inclusão em {{contrato_data_inclusao}}, início de desconto em {{contrato_competencia_inicio}}, no benefício previdenciário {{nb_beneficio}}",

    # ----- PAR 248 (fecho — só cidade/UF; data fica fixa para campo Word) -----
    "Cidade/AL, 9 de maio de 2026.":
        "{{cidade_protocolo}}/{{uf_protocolo}}, 9 de maio de 2026.",
}

doc = Document(CAMINHO)
modificados = []
for i, p in enumerate(doc.paragraphs):
    if substituir_in_run(p._p, MAPA):
        modificados.append(i)

doc.save(CAMINHO)
print(f"\nParágrafos modificados: {len(modificados)}")
for i in modificados:
    print(f"  par {i}")

# Verificação: contar placeholders agora
import re
d2 = Document(CAMINHO)
phs = set()
for p in d2.paragraphs:
    for m in re.finditer(r"\{\{[^}]+\}\}", p.text):
        phs.add(m.group(0))
print(f"\nPlaceholders únicos agora: {len(phs)}")
for ph in sorted(phs):
    print(f"  {ph}")

# Conferir se sobrou algum dado piloto
print("\nVerificação: dados pilotos remanescentes (deveria estar VAZIO):")
for i, p in enumerate(d2.paragraphs):
    t = p.text
    for ruim in ["FULANO DE TAL", "xxx.xxx.xxx-xx", "rua tal", "bairro tal",
                  "BANCO PAN S/A", "149.139.433-9", "3880089838",
                  "Cidade/AL", "Av. Ephigenio", "59.285.411"]:
        if ruim in t:
            print(f"  ⚠️ par {i}: '{ruim}' ainda presente")
            break
