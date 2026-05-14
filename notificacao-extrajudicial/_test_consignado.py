"""
Teste 1 — Tese consignado-nao-contratado.

Gera 2 versões usando os modelos REAIS enviados pelo escritório:
  - COM ESCRITÓRIO (Tiago/Eduardo) → base Tiago AL contra Bradesco
  - SEM ESCRITÓRIO (Patrick/Gabriel/Alexandre) → base Patrick AM contra Bradesco

Substituições GRANULARES preservam fonte Segoe UI no nome + Cambria no resto.
Inclui:
  - RG/órgão expedidor
  - Placeholders de gênero
  - Dano moral fixo R$ 15.000 por contrato
  - Limpeza automática de parágrafos vazios excessivos
"""
import sys, os, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, 'scripts')

from scripts.docx_replace import substituir_em_docx

BASE = os.path.dirname(__file__)
ASSETS = os.path.join(BASE, 'assets')

# === COM ESCRITÓRIO — base Tiago AL Bradesco ===
COM_BASE = os.path.join(ASSETS, '__BASE_com-escritorio__modelo-real-enviado.docx')
COM_TEMPLATE = os.path.join(ASSETS, 'template_consignado-nao-contratado__com-escritorio.docx')

mapa_com = {
    # IMPORTANTE: ordenar das strings MAIS LONGAS para as MAIS CURTAS, pra evitar
    # que substituições curtas corrompam os runs antes das longas casarem.

    # ENDEREÇO ESCRITÓRIO (string longa — substitui PRIMEIRO antes que outras mexam nos runs)
    'com escritório na Rua Nossa Senhora da Salete, nº 597, Sala 04, Térreo, Itapoã, CEP 57.314-175, Arapiraca/AL':
        'com escritório na {{ESCRITORIO_ENDERECO_COMPOSTO}}',

    # Contexto fático — corrige "com a" → "com o"
    'celebrado com a BANCO BRADESCO S/A, sob os Contratos nº 0123435237547 e nº 0123435237014':
        'celebrado com o {{NOME_BANCO_CONTRATO}}, sob o(s) Contrato(s) nº {{CONTRATO_NUMEROS}}',

    # Cancelamento
    'O Cancelamento definitivo dos Contratos nº 0123435237547 e nº 0123435237014':
        'O Cancelamento definitivo do(s) Contrato(s) nº {{CONTRATO_NUMEROS}}',

    # Qualificação cliente — string longa cirúrgica
    'aposentada, inscrita no CPF sob o nº 037.079.194-04':
        '{{CLIENTE_PROFISSAO}}, {{INSCRITO_A}} no CPF sob o nº {{CLIENTE_CPF}}, Cédula de Identidade nº {{CLIENTE_RG}}, órgão expedidor {{CLIENTE_RG_ORGAO}}',
    'residente e domiciliado na Rua Povoado Pau Ferro, nº 04':
        'residente e {{DOMICILIADO_A}} na {{CLIENTE_LOGRADOURO}}',
    'Laranjeiras, em Arapiraca/AL': '{{CLIENTE_BAIRRO}}, em {{CLIENTE_MUNICIPIO}}/{{CLIENTE_UF}}',
    'CEP 57.330-000': 'CEP {{CLIENTE_CEP}}',

    # Cabeçalho data
    'Arapiraca/AL, 29 de janeiro de 2026': '{{CIDADE_ASSINATURA}}/{{UF_ASSINATURA}}, {{DATA_EXTENSO}}',
    'Arapiraca/AL 29 de janeiro de 2026': '{{CIDADE_ASSINATURA}}/{{UF_ASSINATURA}}, {{DATA_EXTENSO}}',

    # Endereçamento ao banco — formato 4 linhas
    'À Ouvidoria do BANCO BRADESCO S/A': 'À Ouvidoria do {{BANCO_NOME_QUALIFICADO}}',
    'CNPJ: 60.746.948/0001-12': 'CNPJ: {{BANCO_CNPJ}}',
    'Cidade de Deus': '{{BANCO_LOGRADOURO}}',
    'Bairro Vila Yara': 'Bairro {{BANCO_BAIRRO}}',
    'CEP: 06.029-900 – Osasco/SP': 'CEP: {{BANCO_CEP}} – {{BANCO_MUNICIPIO}}/{{BANCO_UF}}',

    # Strings curtas no fim
    'MARIA DULCELINA DE OLIVEIRA': '{{CLIENTE_NOME}}',
    'brasileira,': '{{CLIENTE_NACIONALIDADE_GENERO}},',
    'viúva,': '{{CLIENTE_ESTADO_CIVIL}},',
    'por intermédio de seu advogado': 'por intermédio de {{SEU_SUA_ADVOGADO_A}}',
    'Tiago de Azevedo Lima': '{{ADVOGADO_NOME}}',

    # OAB — DUAS substituições. O modelo Bradesco tem "OAB/AL  36672" no meio
    # do parágrafo (erro do escrivão original — esse é o nº SC) E "AL 20906A"
    # no rodapé. Ambas viram {{ADVOGADO_OAB_UF}}.
    'OAB/AL  36672': '{{ADVOGADO_OAB_UF}}',
    'AL 20906A': '{{ADVOGADO_OAB_UF}}',

    # Danos morais — fixo R$ 15k por contrato
    'no valor de R$ 15.000,00 (quinze mil reais)':
        'no valor de R$ 15.000,00 (quinze mil reais) por cada contrato',
}

# === SEM ESCRITÓRIO — base Patrick AM Bradesco ===
SEM_BASE = os.path.join(ASSETS, '__BASE_sem-escritorio__modelo-original.docx')
SEM_TEMPLATE = os.path.join(ASSETS, 'template_consignado-nao-contratado__sem-escritorio.docx')

mapa_sem = {
    # ENDEREÇO RESIDENCIAL CLIENTE (string longa — substitui PRIMEIRO)
    'residente e domiciliado à Comunidade Fortaleza, Rio Paracuni, Polo 9, Município de Maués, CEP 69.190-000, estado do Amazonas':
        'residente e {{DOMICILIADO_A}} na {{CLIENTE_LOGRADOURO}}, {{CLIENTE_BAIRRO}}, em {{CLIENTE_MUNICIPIO}}/{{CLIENTE_UF}}, CEP {{CLIENTE_CEP}}',

    # ENDEREÇO ESCRITÓRIO (substitui pelo composto matriz + apoio UF)
    'com escritório na Travessa Michiles, s/n, Centro, na cidade de Maués/AM, CEP 69.195-000':
        'com escritório na {{ESCRITORIO_ENDERECO_COMPOSTO}}',

    # Qualificação cliente — string longa cirúrgica
    'aposentado, inscrito no CPF sob o nº 473.178.092-68, Cédula de Identidade nº 0651903-2, órgão expedidor SSP/AM':
        '{{CLIENTE_PROFISSAO}}, {{INSCRITO_A}} no CPF sob o nº {{CLIENTE_CPF}}, Cédula de Identidade nº {{CLIENTE_RG}}, órgão expedidor {{CLIENTE_RG_ORGAO}}',

    # Cabeçalho data
    'Maués/AM, 9 de fevereiro de 2026': '{{CIDADE_ASSINATURA}}/{{UF_ASSINATURA}}, {{DATA_EXTENSO}}',

    # Endereçamento ao banco
    'À Ouvidoria do BANCO BRADESCO S.A': 'À Ouvidoria do {{BANCO_NOME_QUALIFICADO}}',
    'CNPJ: 60.746.948/0320-73': 'CNPJ: {{BANCO_CNPJ}}',
    'Avenida Sete de Setembro, nº 895,': '{{BANCO_LOGRADOURO}},',
    'Centro, Manaus/AM, CEP 69.005-140':
        '{{BANCO_BAIRRO}}, {{BANCO_MUNICIPIO}}/{{BANCO_UF}}, CEP {{BANCO_CEP}}',

    # Qualificação cliente — strings curtas no fim
    'ADAUTO ANTONIO PINHEIRO DOS SANTOS': '{{CLIENTE_NOME}}',
    'brasileiro,': '{{CLIENTE_NACIONALIDADE_GENERO}},',
    'solteiro,': '{{CLIENTE_ESTADO_CIVIL}},',

    # Advogado
    'PATRICK WILLIAN DA SILVA': '{{ADVOGADO_NOME}}',  # template tem em CAIXA ALTA
    'Patrick Willian da Silva': '{{ADVOGADO_NOME}}',  # rodapé pode estar em title case
    'OAB/AM A2638': '{{ADVOGADO_OAB_UF}}',

    # Banco no contexto fático — corrige "com a" → "com o"
    'celebrado com a BANCO BRADESCO S.A':
        'celebrado com o {{NOME_BANCO_CONTRATO}}',

    # Detalhes do contrato (parágrafo OPCIONAL — só presente em modelos como esse)
    'no valor de R$ 20.770,93, com 96 parcelas de R$ 447,01, com inclusão em 07/04/2025':
        'no valor de R$ {{CONTRATO_VALOR_EMPRESTIMO}}, com {{CONTRATO_QTD_PARCELAS}} parcelas de R$ {{CONTRATO_VALOR_PARCELA}}, com inclusão em {{CONTRATO_DATA_INCLUSAO}}',

    # Cancelamento (singular nesse modelo)
    'O Cancelamento definitivo do Contrato nº 0123528031058':
        'O Cancelamento definitivo do(s) Contrato(s) nº {{CONTRATO_NUMEROS}}',

    # Danos morais
    'no valor de R$ 15.000,00 (quinze mil reais)':
        'no valor de R$ 15.000,00 (quinze mil reais) por cada contrato',
}


# === Helpers ===

def map_genero_cliente(genero: str) -> dict:
    if genero.upper() == 'F':
        return {
            '{{CLIENTE_NACIONALIDADE_GENERO}}': 'brasileira',
            '{{INSCRITO_A}}': 'inscrita',
            '{{DOMICILIADO_A}}': 'domiciliada',
        }
    return {
        '{{CLIENTE_NACIONALIDADE_GENERO}}': 'brasileiro',
        '{{INSCRITO_A}}': 'inscrito',
        '{{DOMICILIADO_A}}': 'domiciliado',
    }


def map_genero_advogado(genero_adv: str) -> dict:
    if genero_adv.upper() == 'F':
        return {'{{SEU_SUA_ADVOGADO_A}}': 'sua advogada'}
    return {'{{SEU_SUA_ADVOGADO_A}}': 'seu advogado'}


# === ETAPA 1: criar templates placeholderizados ===
print('=== ETAPA 1A: Template COM escritório ===')
rel = substituir_em_docx(COM_BASE, mapa_com, COM_TEMPLATE)
print(f'  Substituições: {rel["total_substituicoes"]}')
print(f'  Parágrafos vazios removidos: {rel["paragrafos_vazios_removidos"]}')

# Mudar wrap da text box "Assunto" pra topAndBottom (não desvia texto lateralmente)
import zipfile
def _patch_textbox_wrap(docx_path):
    with zipfile.ZipFile(docx_path, 'r') as z:
        contents = {n: z.read(n) for n in z.namelist()}
    xml = contents['word/document.xml'].decode('utf-8')
    xml = xml.replace('<wp:wrapSquare wrapText="bothSides"/>', '<wp:wrapTopAndBottom/>')
    contents['word/document.xml'] = xml.encode('utf-8')
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for name, data in contents.items():
            z.writestr(name, data)
_patch_textbox_wrap(COM_TEMPLATE)
print(f'  Wrap textbox: bothSides → topAndBottom')

print('\n=== ETAPA 1B: Template SEM escritório ===')
rel = substituir_em_docx(SEM_BASE, mapa_sem, SEM_TEMPLATE)
print(f'  Substituições: {rel["total_substituicoes"]}')
print(f'  Parágrafos vazios removidos: {rel["paragrafos_vazios_removidos"]}')

_patch_textbox_wrap(SEM_TEMPLATE)
print(f'  Wrap textbox: bothSides → topAndBottom')

# Padronização: adicionar parágrafo "Outrossim email" no SEM (faltava no original)
from scripts.docx_replace import inserir_paragrafo_antes
from docx import Document
d = Document(SEM_TEMPLATE)
texto_outrossim = ('Outrossim, informa-se que eventual resposta poderá ser '
                   'encaminhada diretamente ao e-mail institucional do escritório, '
                   'qual seja: notificacoesescritorio@gmail.com.')
inseriu = inserir_paragrafo_antes(d, 'Em caso de não atendimento desta solicitação',
                                  texto_outrossim, herdar_estilo_de='anterior')
print(f'  Padronização "Outrossim email" no SEM: {"OK" if inseriu else "ÂNCORA NÃO ACHADA"}')
d.save(SEM_TEMPLATE)


# === ETAPA 2: gerar testes (1 com cada template) ===

cliente_F = {
    '{{CIDADE_ASSINATURA}}': 'Arapiraca',
    '{{UF_ASSINATURA}}': 'AL',
    '{{DATA_EXTENSO}}': '15 de maio de 2026',
    '{{BANCO_NOME_QUALIFICADO}}': 'BANCO BRADESCO S/A, sociedade anônima',
    '{{BANCO_CNPJ}}': '60.746.948/0001-12',
    '{{BANCO_LOGRADOURO}}': 'Cidade de Deus',
    '{{BANCO_BAIRRO}}': 'Vila Yara',
    '{{BANCO_CEP}}': '06.029-900',
    '{{BANCO_MUNICIPIO}}': 'Osasco',
    '{{BANCO_UF}}': 'SP',
    '{{CLIENTE_NOME}}': 'MARIA APARECIDA',  # nome curto pra testar layout
    '{{CLIENTE_ESTADO_CIVIL}}': 'viúva',
    '{{CLIENTE_PROFISSAO}}': 'aposentada',
    '{{CLIENTE_CPF}}': '000.111.222-33',
    '{{CLIENTE_RG}}': '1.234.567',
    '{{CLIENTE_RG_ORGAO}}': 'SSP/AL',
    '{{CLIENTE_LOGRADOURO}}': 'Rua Exemplo das Flores, nº 100',
    '{{CLIENTE_BAIRRO}}': 'Centro',
    '{{CLIENTE_MUNICIPIO}}': 'Maceió',
    '{{CLIENTE_UF}}': 'AL',
    '{{CLIENTE_CEP}}': '57000-000',
    '{{ESCRITORIO_ENDERECO_COMPOSTO}}': 'Rua Frei Rogério, 541, Centro, Joaçaba/SC, CEP 89600-000, e unidade de apoio em Rua Nossa Senhora da Salete, 597, Sala 04, Itapuã, Arapiraca/AL, CEP 57314-175',
    '{{NOME_BANCO_CONTRATO}}': 'BANCO BRADESCO S/A',
    '{{CONTRATO_NUMEROS}}': '0123435237547 e 0123435237014',
    **map_genero_cliente('F'),
}

# Cliente para teste COM escritório (Tiago AL — mantém o caso original)
TESTE_COM = os.path.join(BASE, '_TESTE_consignado__com-escritorio.docx')
cliente_com = {
    **cliente_F,
    '{{ADVOGADO_NOME}}': 'Tiago de Azevedo Lima',
    '{{ADVOGADO_OAB_UF}}': 'OAB/AL 20906A',
    **map_genero_advogado('M'),
}
print('\n=== ETAPA 2A: Gerando teste COM escritório (Tiago AL) ===')
rel = substituir_em_docx(COM_TEMPLATE, cliente_com, TESTE_COM)
print(f'  Substituições: {rel["total_substituicoes"]}')

# Cliente para teste SEM escritório (Patrick AM)
TESTE_SEM = os.path.join(BASE, '_TESTE_consignado__sem-escritorio.docx')
cliente_sem = {
    **cliente_F,
    '{{CIDADE_ASSINATURA}}': 'Maués',
    '{{UF_ASSINATURA}}': 'AM',
    '{{ADVOGADO_NOME}}': 'Patrick Willian da Silva',
    '{{ADVOGADO_OAB_UF}}': 'OAB/AM A2638',
    # Detalhes do contrato (só esse template tem)
    '{{CONTRATO_VALOR_EMPRESTIMO}}': '20.770,93',
    '{{CONTRATO_QTD_PARCELAS}}': '96',
    '{{CONTRATO_VALOR_PARCELA}}': '447,01',
    '{{CONTRATO_DATA_INCLUSAO}}': '07/04/2025',
    '{{CONTRATO_DATA_PRIMEIRO_DESCONTO}}': '07/05/2025',
    **map_genero_advogado('M'),
}
print('\n=== ETAPA 2B: Gerando teste SEM escritório (Patrick AM) ===')
rel = substituir_em_docx(SEM_TEMPLATE, cliente_sem, TESTE_SEM)
print(f'  Substituições: {rel["total_substituicoes"]}')

# Verificar placeholders restantes
import re
from docx import Document
print('\n=== Placeholders pendentes ===')
for label, path in [('COM', TESTE_COM), ('SEM', TESTE_SEM)]:
    d = Document(path)
    restantes = []
    for p in d.paragraphs:
        m = re.findall(r'\{\{[A-Z_]+\}\}', p.text)
        if m:
            restantes.extend(m)
    if restantes:
        print(f'  {label}: {set(restantes)}')
    else:
        print(f'  {label}: tudo preenchido!')

print(f'\nTestes salvos em:')
print(f'  {TESTE_COM}')
print(f'  {TESTE_SEM}')
