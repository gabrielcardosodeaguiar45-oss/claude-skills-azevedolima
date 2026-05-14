# -*- coding: utf-8 -*-
"""Dump dos 10 templates de notificação para análise."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
from docx import Document

ASSETS = Path(r"C:\Users\gabri\.claude\skills\notificacao-extrajudicial\assets")
OUT_DIR = Path(__file__).parent / "_pente_fino"
OUT_DIR.mkdir(exist_ok=True)

TEMPLATES = [
    "template_consignado-nao-contratado__com-escritorio.docx",
    "template_consignado-nao-contratado__sem-escritorio.docx",
    "template_rmc__com-escritorio.docx",
    "template_rmc__sem-escritorio.docx",
    "template_rcc__com-escritorio.docx",
    "template_rcc__sem-escritorio.docx",
    "template_bradesco-tarifas__sem-escritorio.docx",
    "template_bradesco-encargos__sem-escritorio.docx",
    "template_bradesco-capitalizacao__sem-escritorio.docx",
    "template_bradesco-pe__sem-escritorio.docx",
]

for nome in TEMPLATES:
    caminho = ASSETS / nome
    if not caminho.exists():
        print(f'AUSENTE: {nome}')
        continue
    d = Document(caminho)
    out = OUT_DIR / nome.replace('.docx', '.txt')
    linhas = []
    for i, p in enumerate(d.paragraphs):
        t = p.text
        marca = ' [VAZIO]' if not t.strip() else ''
        linhas.append(f'{i:4d}{marca}: {t}')
    # Tabelas
    for ti, table in enumerate(d.tables):
        linhas.append(f'\n--- TABELA {ti} ({len(table.rows)} linhas × {len(table.columns)} colunas) ---')
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    t = p.text.strip()
                    if t:
                        linhas.append(f'  T{ti}R{ri}C{ci}P{pi}: {t}')
    out.write_text('\n'.join(linhas), encoding='utf-8')
    print(f'{nome}: {len(d.paragraphs)} pars, {len(d.tables)} tabelas')
