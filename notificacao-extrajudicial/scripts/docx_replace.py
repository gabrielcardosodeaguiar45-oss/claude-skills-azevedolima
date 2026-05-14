"""
Find-and-replace em DOCX preservando formatação.

Lida com o caso comum em que python-docx quebra texto em múltiplos `runs`
mesmo dentro de um único parágrafo (isso acontece quando há mudança de
fonte, negrito, sublinhado etc., ou às vezes por nada — só herança do
editor que criou o arquivo).

Estratégia:
1. Para cada parágrafo, computar o texto inteiro (concat dos runs)
2. Se o texto contém o `old_str`, fazer substituição
3. Re-distribuir o resultado pelos runs originais, preservando o estilo
   do PRIMEIRO run que continha pedaço do `old_str`. Demais runs cobertos
   pela busca ficam vazios (text="").

Uso CLI:
    python docx_replace.py <input.docx> <mapeamento.json> <output.docx>

Onde mapeamento.json é tipo:
    {"valor antigo": "{{NOVO_PLACEHOLDER}}", ...}

Uso programático:
    from docx_replace import substituir_em_docx
    substituir_em_docx("in.docx", {"x": "y"}, "out.docx")
"""
import sys
import os
import json
import re
from pathlib import Path

try:
    from docx import Document
except ImportError as e:
    raise ImportError(
        f"Dependência ausente: {e}. "
        f"Instale via: pip install -r requirements.txt"
    ) from e


def substituir_em_paragrafo(paragrafo, mapeamento: dict) -> int:
    """
    Aplica substituições em um parágrafo, preservando formatação dos runs.
    Retorna número de substituições feitas.
    """
    n = 0
    if not paragrafo.runs:
        return 0
    for old_str, new_str in mapeamento.items():
        if not old_str:
            continue
        # Loop while: substitui TODAS as ocorrências de old_str neste parágrafo,
        # não só a primeira. Safeguard: se new_str contiver old_str (caso raro),
        # limita a 100 iterações para evitar loop infinito.
        max_iter = 100
        while old_str in paragrafo.text and max_iter > 0:
            max_iter -= 1
            # Texto inteiro do parágrafo
            texto_completo = "".join(r.text for r in paragrafo.runs)
            if old_str not in texto_completo:
                break
            # Achar onde começa o old_str em termos de runs
            pos_global = 0
            run_inicio = None
            offset_inicio = None
            for idx, r in enumerate(paragrafo.runs):
                r_len = len(r.text)
                if pos_global + r_len > texto_completo.index(old_str):
                    run_inicio = idx
                    offset_inicio = texto_completo.index(old_str) - pos_global
                    break
                pos_global += r_len
            if run_inicio is None:
                break
            # Achar até onde se estende
            old_len = len(old_str)
            pos_global = 0
            run_fim = None
            offset_fim = None
            for idx, r in enumerate(paragrafo.runs):
                r_len = len(r.text)
                if pos_global + r_len >= texto_completo.index(old_str) + old_len:
                    run_fim = idx
                    offset_fim = (texto_completo.index(old_str) + old_len) - pos_global
                    break
                pos_global += r_len
            if run_fim is None:
                run_fim = len(paragrafo.runs) - 1
                offset_fim = len(paragrafo.runs[-1].text)

            # Caso A: tudo em 1 run só
            if run_inicio == run_fim:
                r = paragrafo.runs[run_inicio]
                r.text = r.text[:offset_inicio] + new_str + r.text[offset_fim:]
            else:
                # Caso B: vários runs
                primeiro = paragrafo.runs[run_inicio]
                ultimo = paragrafo.runs[run_fim]
                primeiro.text = primeiro.text[:offset_inicio] + new_str
                ultimo.text = ultimo.text[offset_fim:]
                for idx in range(run_inicio + 1, run_fim):
                    paragrafo.runs[idx].text = ""
            n += 1
            # Safeguard contra loop infinito quando new_str contém old_str
            if old_str in new_str:
                break
    return n


class PlaceholdersResiduaisError(RuntimeError):
    """Levantada por substituir_em_docx quando o DOCX gerado ainda contém
    placeholders {{...}} não substituídos. Em modo strict (padrão) isso
    impede que uma notificação incompleta seja entregue ao usuário ou
    enviada por e-mail.

    Causa raiz típica: o template foi atualizado com novos placeholders
    mas montar_mapa_placeholders() em _run_notificacoes.py não foi
    atualizada para fornecê-los (vide commit 1c10a59 — bug recorrente).
    """

    def __init__(self, residuais, dst_path):
        self.residuais = list(residuais)
        self.dst_path = dst_path
        super().__init__(
            f'NOTIFICACAO INCOMPLETA — placeholders {{{{...}}}} restaram em '
            f'{Path(dst_path).name}: {self.residuais}. '
            f'Causa típica: o mapeamento não cobre o token, ou o template '
            f'usou um nome novo que ainda não foi mapeado em '
            f'_run_notificacoes.montar_mapa_placeholders. Adicione o(s) '
            f'token(s) ao mapa OU remova-os do template.'
        )


def substituir_em_docx(input_path: str, mapeamento: dict, output_path: str,
                       max_paragrafos_vazios: int = 1, strict: bool = True) -> dict:
    """
    Aplica todas as substituições do mapeamento em todos os parágrafos do DOCX.
    Inclui parágrafos dentro de tabelas e cabeçalhos/rodapés.

    Após substituir, limita parágrafos vazios consecutivos no corpo a no máximo
    `max_paragrafos_vazios` (1 por padrão), eliminando espaçamentos exagerados
    de modelos.

    Args:
        input_path: caminho do template DOCX
        mapeamento: dict {string_a_buscar: string_substituta}
        output_path: caminho de saída
        max_paragrafos_vazios: limite de parágrafos vazios consecutivos
        strict: se True (padrão), LEVANTA :class:`PlaceholdersResiduaisError`
            quando sobrar qualquer ``{{xxx}}`` no DOCX final E renomeia o
            arquivo para ``*_FALHOU_PLACEHOLDERS.docx``. Esse comportamento
            evita que notificações com placeholders crus sejam enviadas
            (bug recorrente — vide commit 1c10a59 e E15 do inicial-bradesco).

    Retorna dict com estatísticas, incluindo a chave ``residuais`` com a
    lista (vazia se OK) de placeholders não substituídos.
    """
    doc = Document(input_path)
    total = 0

    # Parágrafos do corpo
    for p in doc.paragraphs:
        total += substituir_em_paragrafo(p, mapeamento)

    # Tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += substituir_em_paragrafo(p, mapeamento)

    # Cabeçalhos e rodapés
    for section in doc.sections:
        for p in section.header.paragraphs:
            total += substituir_em_paragrafo(p, mapeamento)
        for p in section.footer.paragraphs:
            total += substituir_em_paragrafo(p, mapeamento)

    # Limitar parágrafos vazios consecutivos
    paragrafos_removidos = limitar_paragrafos_vazios(doc, max_paragrafos_vazios)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    # Verifica residuais — varre TODO o documento (corpo + tabelas +
    # headers/footers) procurando {{xxx}} crus.
    import zipfile
    with zipfile.ZipFile(output_path, 'r') as z:
        partes = [n for n in z.namelist()
                  if n.endswith('.xml') and ('document' in n or 'header' in n or 'footer' in n)]
        xml_total = ''
        for n in partes:
            try:
                xml_total += z.read(n).decode('utf-8')
            except Exception:
                pass
    residuais = sorted(set(re.findall(r'\{\{([^{}]+)\}\}', xml_total)))

    if residuais and strict:
        # Renomeia para deixar visualmente óbvio que NÃO pode ser usada.
        # Mantém o arquivo (não deleta) para inspeção.
        from os import rename as _os_rename, remove as _os_remove
        falha_path = str(Path(output_path).with_name(
            Path(output_path).stem + '_FALHOU_PLACEHOLDERS' + Path(output_path).suffix
        ))
        if Path(falha_path).exists():
            _os_remove(falha_path)
        _os_rename(output_path, falha_path)
        raise PlaceholdersResiduaisError(residuais, falha_path)

    return {
        "total_substituicoes": total,
        "paragrafos_vazios_removidos": paragrafos_removidos,
        "residuais": residuais,
        "output": output_path,
    }


def padronizar_fontes(doc, fonte_corpo: str = 'Cambria',
                      fonte_destaque: str = 'Segoe UI') -> dict:
    """
    Padroniza fontes do documento seguindo o padrão das notificações do
    escritório:
        - Corpo do texto (parágrafos sem estilo de título): Cambria
        - Títulos de seção (estilo '2. Título', etc.): Segoe UI (já do estilo)
        - Nome do cliente (parágrafo da qualificação): Segoe UI
        - Cabeçalho "À Ouvidoria do BANCO": Segoe UI bold
        - "Notificação Extrajudicial" (título principal): Segoe UI

    Usa a heurística: aplica Cambria em runs sem fonte explícita,
    EXCETO no nome do cliente (até a primeira vírgula do parágrafo
    de qualificação) e em parágrafos de título.

    Retorna dict com estatísticas.
    """
    estatisticas = {'cambria_aplicada': 0, 'segoe_aplicada': 0, 'mantidas': 0}
    estilos_titulo = ['2. Título', '3. Subtítulo', '3.1 Subtítulo', 'Heading']

    for p in doc.paragraphs:
        if not p.runs:
            continue
        texto = p.text
        style_name = p.style.name if p.style else 'Normal'

        # Parágrafos de título: NÃO mexer (mantém Segoe UI do estilo)
        if any(t in style_name for t in estilos_titulo):
            estatisticas['mantidas'] += len(p.runs)
            continue

        # Parágrafo do título "Notificação Extrajudicial" e subtítulo da tese
        if texto.strip() in ('Notificação Extrajudicial',) or \
           any(t in texto for t in ('Empréstimo Não Contratado', 'Cartão de Crédito RMC',
                                     'Cartão de Crédito RCC', 'Descontos indevidos')):
            for r in p.runs:
                if not r.font.name:
                    r.font.name = fonte_destaque
                    estatisticas['segoe_aplicada'] += 1
            continue

        # Parágrafo "À Ouvidoria do BANCO X" — Segoe UI bold
        if texto.startswith('À Ouvidoria') or texto.startswith('Ao'):
            for r in p.runs:
                if not r.font.name:
                    r.font.name = fonte_destaque
                    estatisticas['segoe_aplicada'] += 1
            continue

        # Parágrafo da qualificação (tem 'inscrita/inscrito no CPF' OU os
        # placeholders {{INSCRITO_A}} / {{CLIENTE_CPF}}). Case-insensitive
        # para funcionar antes E depois da substituição de placeholders.
        if 'inscrit' in texto.lower() and 'cpf' in texto.lower():
            # Heurística: setar Segoe UI no(s) primeiro(s) run(s) até a
            # primeira vírgula (que marca o fim do nome). Depois, Cambria.
            primeira_virgula = texto.find(',')
            pos_atual = 0
            nome_terminou = False
            for r in p.runs:
                run_text = r.text
                run_inicio = pos_atual
                run_fim = pos_atual + len(run_text)
                if not nome_terminou and run_inicio < primeira_virgula:
                    # Run faz parte do nome (ou contém a vírgula final do nome)
                    r.font.name = fonte_destaque
                    estatisticas['segoe_aplicada'] += 1
                    if run_fim >= primeira_virgula:
                        nome_terminou = True
                else:
                    # Run depois do nome → Cambria
                    if not r.font.name:
                        r.font.name = fonte_corpo
                        estatisticas['cambria_aplicada'] += 1
                pos_atual = run_fim
            continue

        # Parágrafo padrão de corpo: Cambria nos runs sem fonte
        for r in p.runs:
            if not r.font.name:
                r.font.name = fonte_corpo
                estatisticas['cambria_aplicada'] += 1
            else:
                estatisticas['mantidas'] += 1

    return estatisticas


def normalizar_tema_corpo(docx_path: str) -> dict:
    """
    Corrige documentos cujo corpo foi salvo apontando para majorFont
    (geralmente Calibri) quando deveria apontar para minorFont (Cambria
    em temas Office padrão).

    Substitui no word/document.xml:
        asciiTheme="majorHAnsi"  →  asciiTheme="minorHAnsi"
        hAnsiTheme="majorHAnsi"  →  hAnsiTheme="minorHAnsi"
        cstheme="majorHAnsi"     →  cstheme="minorHAnsi"

    Runs com w:ascii/w:hAnsi/w:cs explícitos (ex: Segoe UI nos títulos)
    NÃO são afetados — só os que herdam do tema.

    Usa-se quando o BASE foi salvo pelo Word com majorHAnsi em vez de
    minorHAnsi. Sintoma: documento renderiza tudo em Calibri apesar de
    minorFont=Cambria no theme1.xml.

    Sobrescreve o arquivo. Retorna dict com estatísticas.
    """
    import zipfile
    with zipfile.ZipFile(docx_path, 'r') as z:
        contents = {n: z.read(n) for n in z.namelist()}
    if 'word/document.xml' not in contents:
        return {'major_substituidos': 0}
    xml = contents['word/document.xml'].decode('utf-8')
    n_ascii = xml.count('asciiTheme="majorHAnsi"')
    n_hansi = xml.count('hAnsiTheme="majorHAnsi"')
    n_cs = xml.count('cstheme="majorHAnsi"')
    xml = xml.replace('asciiTheme="majorHAnsi"', 'asciiTheme="minorHAnsi"')
    xml = xml.replace('hAnsiTheme="majorHAnsi"', 'hAnsiTheme="minorHAnsi"')
    xml = xml.replace('cstheme="majorHAnsi"', 'cstheme="minorHAnsi"')
    contents['word/document.xml'] = xml.encode('utf-8')
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for name, data in contents.items():
            z.writestr(name, data)
    return {'ascii': n_ascii, 'hAnsi': n_hansi, 'cs': n_cs,
            'total': n_ascii + n_hansi + n_cs}


def aplicar_timbrado_neutro(template_com_path: str, timbrado_sem_base: str,
                             output_path: str) -> str:
    """
    Cria uma versão SEM-escritório de um template, copiando o CORPO do
    template COM e os HEADERS/FOOTERS/IMAGENS do timbrado neutro base.

    Estratégia:
      - Abre o template COM (corpo da tese correto)
      - Abre o timbrado SEM (Patrick AM original — header/footer neutros)
      - No output: corpo do COM + headers/footers/imagens do SEM

    Limitações:
      - Os IDs de relacionamento (rId) entre header.xml e image.png precisam
        ser consistentes. O template SEM-base usa 'media/image1.png',
        'media/image2.png', 'media/image3.png'. Se o template COM tiver
        outras imagens, podem haver conflitos.
      - Para a maioria dos templates do escritório isso funciona porque
        a única imagem do COM é 'image1.png' (logo Azevedo Lima), que será
        substituída pelo conjunto neutro do SEM.
    """
    import zipfile
    # Ler conteúdo de ambos os zips
    with zipfile.ZipFile(template_com_path, 'r') as z_com:
        com_files = {n: z_com.read(n) for n in z_com.namelist()}
    with zipfile.ZipFile(timbrado_sem_base, 'r') as z_sem:
        sem_files = {n: z_sem.read(n) for n in z_sem.namelist()}

    # Arquivos a tomar do SEM (timbrado): headers, footers, _rels, imagens
    arquivos_timbrado = [
        n for n in sem_files
        if n.startswith('word/header') or n.startswith('word/footer')
        or n.startswith('word/_rels/header') or n.startswith('word/_rels/footer')
        or n.startswith('word/media/image')
    ]

    # Construir output: corpo do COM + timbrado do SEM
    saida = dict(com_files)
    for n in arquivos_timbrado:
        saida[n] = sem_files[n]
    # word/_rels/document.xml.rels precisa apontar pros headers/footers que vão pro doc
    # Mantém o do COM (que já referencia headers1/2/3 e footers1/2)

    # Salvar
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as z_out:
        for name, data in saida.items():
            z_out.writestr(name, data)
    return output_path


def inserir_paragrafo_antes(doc, ancora_str: str, novo_texto: str,
                            herdar_estilo_de: str = "anterior") -> bool:
    """
    Insere um novo parágrafo ANTES do parágrafo que contém `ancora_str`.

    O novo parágrafo tem o texto `novo_texto`. Para manter formatação
    consistente, herda o estilo do parágrafo `anterior` ao âncora (default)
    ou do `proximo` (parágrafo âncora em si).

    Retorna True se inseriu, False se não achou a âncora.
    """
    paragrafos = list(doc.paragraphs)
    for i, p in enumerate(paragrafos):
        if ancora_str in p.text:
            # Identificar parágrafo modelo de estilo
            if herdar_estilo_de == "anterior" and i > 0:
                modelo = paragrafos[i-1]
            else:
                modelo = p
            # Inserir novo parágrafo ANTES do âncora usando python-docx
            novo_p = p.insert_paragraph_before(novo_texto, style=modelo.style)
            # Copiar alinhamento
            novo_p.paragraph_format.alignment = modelo.paragraph_format.alignment
            return True
    return False


def limitar_paragrafos_vazios(doc, max_consecutivos: int = 1) -> int:
    """
    Remove parágrafos vazios CONSECUTIVOS que excedam max_consecutivos.

    PRESERVA parágrafos que contenham elementos críticos de seção:
        - <w:sectPr> (define header/footer da seção)
        - <w:bookmarkStart/End>
        - drawings, fields, etc.

    Sem essa proteção, removeria os <w:sectPr> ao limpar parágrafos
    aparentemente vazios e quebraria o layout (timbrados, header/footer).

    Retorna número de parágrafos removidos.
    """
    from docx.oxml.ns import qn

    # Tags que NUNCA devem ser removidas mesmo se o parágrafo está sem texto
    TAGS_CRITICAS = {
        qn('w:sectPr'),
        qn('w:bookmarkStart'), qn('w:bookmarkEnd'),
        qn('w:drawing'), qn('w:pict'),
        qn('w:fldChar'), qn('w:instrText'),
        qn('w:hyperlink'),
    }
    # mc:AlternateContent não está no nsmap padrão do python-docx
    # — adicionar manualmente
    TAGS_CRITICAS.add('{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent')

    paragrafos = list(doc.paragraphs)
    removidos = 0
    consec_vazios = 0
    a_remover = []
    for p in paragrafos:
        # Verifica se o parágrafo contém algum elemento crítico
        elem = p._element
        tem_critico = False
        for tag in TAGS_CRITICAS:
            if elem.find(f'.//{tag}') is not None:
                tem_critico = True
                break
        if tem_critico:
            consec_vazios = 0  # reset, parágrafo não conta como "vazio"
            continue

        if not p.text.strip():
            consec_vazios += 1
            if consec_vazios > max_consecutivos:
                a_remover.append(p)
        else:
            consec_vazios = 0

    for p in a_remover:
        elem = p._element
        elem.getparent().remove(elem)
        removidos += 1
    return removidos


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    inp, mapfile, outp = sys.argv[1:4]
    with open(mapfile, encoding="utf-8") as f:
        mapeamento = json.load(f)
    rel = substituir_em_docx(inp, mapeamento, outp)
    print(f"OK: {rel['total_substituicoes']} substituições -> {outp}")


if __name__ == "__main__":
    main()
