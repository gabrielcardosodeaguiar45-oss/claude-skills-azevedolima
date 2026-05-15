"""Auditoria automática de iniciais RMC/RCC e notificações extrajudiciais.

Verifica em cada DOCX:
  1. Placeholders restantes ({{...}})
  2. Parágrafos proibidos (lista canônica)
  3. Padrões de bug conhecidos (R$ R$, vírgulas duplas, CPF/RG vazio)
  4. Campos críticos ausentes

Saída: relatório DOCX consolidado por cliente.
"""
import io, sys, os, re
from datetime import datetime
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor

ROOTS = [
    r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - NÃO CONTRATADO\RMC - RCC - NÃO CONTRATADO',
    r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - NÃO CONTRATADO\RMC - RCC - NÃO CONTRATADO - BACKUP\RMC - RCC - NÃO CONTRATADO',
]

# Parágrafos/frases que NÃO podem aparecer (lista canônica de banidos)
PARAGRAFOS_PROIBIDOS = [
    "Conforme informações constantes do próprio extrato de benefício",  # extrato literal
    "Tais descontos foram identificados",  # Bradesco antigo
]

# Padrões de bug típicos
BUGS_REGEX = [
    (r'R\$\s*R\$', 'duplo "R$ R$"'),
    (r'CPF sob o nº ,', 'CPF vazio (CPF sob o nº ,)'),
    (r'CPF sob o nº\s*\.', 'CPF vazio (CPF sob o nº.)'),
    (r'Cédula de Identidade nº ,', 'RG vazio (Cédula nº ,)'),
    (r'\{\{[^}]+\}\}', 'placeholder não substituído'),
    (r'brasileir[oa]?,\s*,', 'estado civil vazio (brasileir(o/a), ,)'),
    (r'(\bnº \d+) e \1\b', 'número duplicado'),
]


def coletar_paragrafos(doc):
    """Retorna lista de textos: parágrafos + cells da tabela."""
    textos = []
    for p in doc.paragraphs:
        if p.text.strip():
            textos.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        textos.append(p.text)
    return textos


def auditar_docx(path):
    """Retorna dict com achados do arquivo."""
    achados = {
        'placeholders': set(),
        'paragrafos_proibidos': [],
        'bugs': [],
        'arquivo': path,
    }
    try:
        d = Document(path)
    except Exception as e:
        achados['erro_leitura'] = str(e)
        return achados
    textos = coletar_paragrafos(d)
    full_text = '\n'.join(textos)
    # 1. Placeholders
    for m in re.findall(r'\{\{[^}]+\}\}', full_text):
        achados['placeholders'].add(m)
    # 2. Parágrafos proibidos
    for frase in PARAGRAFOS_PROIBIDOS:
        if frase in full_text:
            # Pegar trecho de até 200 chars onde aparece
            idx = full_text.find(frase)
            achados['paragrafos_proibidos'].append({
                'frase': frase,
                'contexto': full_text[max(0,idx-20):idx+200],
            })
    # 3. Bugs por regex
    for padrao, descr in BUGS_REGEX:
        for m in re.finditer(padrao, full_text):
            ctx_start = max(0, m.start() - 30)
            ctx_end = min(len(full_text), m.end() + 30)
            achados['bugs'].append({
                'tipo': descr,
                'match': m.group(0),
                'contexto': full_text[ctx_start:ctx_end].replace('\n', ' '),
            })
    achados['placeholders'] = sorted(achados['placeholders'])
    return achados


def classificar(filename):
    """Determina se é inicial RMC/RCC, notificação, etc."""
    fn = filename.upper()
    if 'NOTIF' in fn:
        return 'notificacao'
    if 'INICIAL' in fn:
        if '_RMC.DOCX' in fn or '_RMC_' in fn:
            return 'inicial-rmc'
        if '_RCC.DOCX' in fn or '_RCC_' in fn:
            return 'inicial-rcc'
        return 'inicial-outro'
    return None


def varrer():
    docs = []
    for root in ROOTS:
        if not os.path.exists(root): continue
        for cur, _, files in os.walk(root):
            for f in files:
                if not f.lower().endswith('.docx'): continue
                if f.startswith('~$'): continue  # lock Word
                if '.bak' in f.lower(): continue
                tipo = classificar(f)
                if not tipo: continue
                docs.append((tipo, os.path.join(cur, f)))
    return docs


def gerar_relatorio(achados_por_arquivo, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Cambria'
    style.font.size = Pt(10)

    h = doc.add_paragraph()
    r = h.add_run('AUDITORIA — RMC/RCC + NOTIFICAÇÕES')
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph(f'Total de arquivos analisados: {len(achados_por_arquivo)}')

    # Resumo
    total_placeholders = sum(len(a['placeholders']) for a in achados_por_arquivo if 'placeholders' in a)
    total_proibidos = sum(len(a['paragrafos_proibidos']) for a in achados_por_arquivo if 'paragrafos_proibidos' in a)
    total_bugs = sum(len(a['bugs']) for a in achados_por_arquivo if 'bugs' in a)
    total_arquivos_com_problema = sum(1 for a in achados_por_arquivo
        if a.get('placeholders') or a.get('paragrafos_proibidos') or a.get('bugs'))

    p = doc.add_paragraph()
    p.add_run('RESUMO').bold = True
    doc.add_paragraph(f'• Arquivos COM problemas: {total_arquivos_com_problema}')
    doc.add_paragraph(f'• Placeholders {{...}} não substituídos: {total_placeholders} ocorrências')
    doc.add_paragraph(f'• Parágrafos proibidos detectados: {total_proibidos} ocorrências')
    doc.add_paragraph(f'• Bugs/padrões problemáticos: {total_bugs} ocorrências')
    doc.add_paragraph()

    # Agrupar por cliente
    por_cliente = {}
    for a in achados_por_arquivo:
        path = a['arquivo']
        cliente = None
        partes = path.split(os.sep)
        for p_ in partes:
            if ' - ' in p_ and 'APP' not in p_ and 'RMC' not in p_[:5] and not p_.startswith('1.'):
                cliente = p_; break
        if not cliente:
            cliente = '(?)'
        por_cliente.setdefault(cliente, []).append(a)

    for cliente in sorted(por_cliente):
        achs = por_cliente[cliente]
        problemas_aqui = [a for a in achs if a.get('placeholders') or a.get('paragrafos_proibidos') or a.get('bugs')]
        if not problemas_aqui:
            continue
        p = doc.add_paragraph()
        r = p.add_run(cliente)
        r.bold = True; r.font.size = Pt(12)
        for a in problemas_aqui:
            rel = a['arquivo']
            # mostrar só os 3 últimos níveis
            partes = rel.split(os.sep)
            rel_curto = os.sep.join(partes[-4:])
            pp = doc.add_paragraph()
            pp.add_run(f'  📄 {rel_curto}').italic = True
            if a.get('placeholders'):
                doc.add_paragraph(f'     ⚠ Placeholders restantes ({len(a["placeholders"])}): ' + ', '.join(a['placeholders']))
            if a.get('paragrafos_proibidos'):
                for pr in a['paragrafos_proibidos']:
                    doc.add_paragraph(f'     ⚠ Parágrafo proibido: "{pr["frase"]}"')
                    doc.add_paragraph(f'       Contexto: ...{pr["contexto"][:200]}...')
            if a.get('bugs'):
                for b in a['bugs']:
                    doc.add_paragraph(f'     ⚠ {b["tipo"]} — match: "{b["match"]}"')
                    doc.add_paragraph(f'       Contexto: ...{b["contexto"]}...')
        doc.add_paragraph()
    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    docs = varrer()
    print(f'Encontrados {len(docs)} DOCX para auditar')
    achados = []
    for tipo, path in docs:
        a = auditar_docx(path)
        a['tipo'] = tipo
        achados.append(a)
    out = os.path.join(os.path.dirname(__file__), f'AUDITORIA_{datetime.now().strftime("%Y%m%d_%H%M")}.docx')
    gerar_relatorio(achados, out)
    print(f'Relatório gerado: {out}')

    # Sumário no console
    com_problema = [a for a in achados if a.get('placeholders') or a.get('paragrafos_proibidos') or a.get('bugs')]
    print(f'\n=== {len(com_problema)} arquivos COM problema (de {len(achados)} totais) ===')
    for a in com_problema[:30]:  # primeiros 30
        partes = a['arquivo'].split(os.sep)
        curto = os.sep.join(partes[-3:])
        flags = []
        if a.get('placeholders'): flags.append(f'plh:{len(a["placeholders"])}')
        if a.get('paragrafos_proibidos'): flags.append(f'proib:{len(a["paragrafos_proibidos"])}')
        if a.get('bugs'): flags.append(f'bugs:{len(a["bugs"])}')
        print(f'  [{",".join(flags)}] {curto}')
    if len(com_problema) > 30:
        print(f'  ... e mais {len(com_problema)-30}')
