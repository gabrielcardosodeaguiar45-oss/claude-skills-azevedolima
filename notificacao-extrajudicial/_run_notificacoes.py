"""
Pipeline de geração de notificações extrajudiciais em batch.

Para cada cliente em <PASTA_BATCH>:
  1. Carrega _estado_cliente.json
  2. Localiza procuração no KIT (PDF mais provável)
  3. Extrai qualificação via OCR
  4. Para cada pasta_acao:
     - Identifica tese (heurística do nome da pasta)
     - Filtra contratos relevantes
     - Resolve banco no cadastro
     - Monta dict de placeholders
     - Gera notificação .docx em <pasta_acao>/notificacao/

Pasta batch é hardcoded para evitar problemas de encoding em argumentos.
"""
import os
import sys
import io
import json
import re
import locale
from datetime import date

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'scripts'))

from docx_replace import substituir_em_docx
from bancos import obter_endereco, resolver_chave
from extrair_qualificacao import extrair_qualificacao

PASTA_BATCH = os.environ.get(
    'NOTIF_PASTA_BATCH',
    r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - ORGANIZAÇÃO PASTA AL\TESTE - Fazer inicial',
)
SKILL_DIR = os.path.dirname(__file__)
ASSETS = os.path.join(SKILL_DIR, 'assets')
OABS_DIR = os.path.join(ASSETS, 'oabs')

# Mapeamento procurador → PDF da OAB em assets/oabs/
OAB_PDF_POR_PROCURADOR = {
    'tiago':     'OAB TIAGO.pdf',
    'patrick':   'OAB PATRICK.pdf',
    'gabriel':   'OAB GABRIEL.pdf',
    'alexandre': 'OAB ALEXANDRE.pdf',
    'eduardo':   'OAB EDUARDO.pdf',
}

# Templates por tese e versão (COM/SEM)
TEMPLATES = {
    ('consignado-nao-contratado', 'COM'): os.path.join(ASSETS, 'template_consignado-nao-contratado__com-escritorio.docx'),
    ('consignado-nao-contratado', 'SEM'): os.path.join(ASSETS, 'template_consignado-nao-contratado__sem-escritorio.docx'),
    ('rmc', 'COM'): os.path.join(ASSETS, 'template_rmc__com-escritorio.docx'),
    ('rmc', 'SEM'): os.path.join(ASSETS, 'template_rmc__sem-escritorio.docx'),
    ('rcc', 'COM'): os.path.join(ASSETS, 'template_rcc__com-escritorio.docx'),
    ('rcc', 'SEM'): os.path.join(ASSETS, 'template_rcc__sem-escritorio.docx'),
}

# Cadastros centrais em skills/_common/ — fonte única de verdade compartilhada
# com inicial-nao-contratado, inicial-bradesco e demais skills do escritório.
sys.path.insert(0, os.path.join(SKILL_DIR, '..', '_common'))
try:
    from procuradores import ADVOGADO_POR_UF
    from escritorios_cadastro import montar_endereco_escritorio_completo
except ImportError as e:
    raise ImportError(
        f"Cadastro central não encontrado em skills/_common/: {e}. "
        f"Verifique se a pasta _common/ existe ao lado de notificacao-extrajudicial/ "
        f"e contém procuradores.py e escritorios_cadastro.py."
    ) from e


def encontrar_procuracao(pasta_kit: str, pasta_cliente: str | None = None) -> str | None:
    """Procura procuração — tenta múltiplas fontes em ordem de preferência:
    1. PDF "Procuração*" no KIT
    2. PDF original referenciado em <pasta_cliente>/_proc_crops/manifesto.json
    3. PNG já recortado em <pasta_cliente>/_proc_crops/crop_pag_01.png
    """
    candidatos = []

    # 1. PDF na pasta KIT
    if pasta_kit and os.path.isdir(pasta_kit):
        for nome in os.listdir(pasta_kit):
            if not nome.lower().endswith('.pdf'):
                continue
            nome_lower = nome.lower()
            if 'procura' in nome_lower:
                score = 100  # alta prioridade
                if 'procurações' in nome_lower or 'procuracoes' in nome_lower:
                    score += 10
                if re.search(r'n[º°]\d+', nome_lower):
                    score += 5
                candidatos.append((score, os.path.join(pasta_kit, nome)))

    # 2. PDF original via _proc_crops/manifesto.json
    if pasta_cliente:
        manifesto_path = os.path.join(pasta_cliente, '_proc_crops', 'manifesto.json')
        if os.path.exists(manifesto_path):
            try:
                with open(manifesto_path, encoding='utf-8') as f:
                    manif = json.load(f)
                pdf_orig = manif.get('pdf_origem', '')
                if pdf_orig and os.path.exists(pdf_orig):
                    candidatos.append((90, pdf_orig))
            except Exception:
                pass

        # 3. PNG já recortado (último recurso, mas funciona)
        crop = os.path.join(pasta_cliente, '_proc_crops', 'crop_pag_01.png')
        if os.path.exists(crop):
            candidatos.append((50, crop))

    if not candidatos:
        return None
    candidatos.sort(reverse=True)
    return candidatos[0][1]


CHAVES_BANCO = ['BMG', 'BRADESCO', 'ITAU', 'C6', 'PAN', 'CAIXA',
                'SANTANDER', 'DAYCOVAL', 'MERCANTIL', 'FACTA',
                'AGIBANK', 'OLE', 'BGN', 'CETELEM', 'MASTER',
                'SAFRA', 'CREFISA', 'PARANA', 'DIGIO', 'BANRISUL',
                'CAPITAL', 'CAPITALCONSIG', 'INBURSA', 'INTER',
                'BB', 'BANCODOBRASIL', 'DOBRASIL', 'BRB', 'PINE']


def _resolver_pasta_acao(pasta_cliente: str, path_relativo: str) -> str | None:
    """Resolve path_relativo para pasta real no disco com fuzzy matching.

    A skill kit-juridico às vezes gera nomes com espaços ruins
    ('BANCO BRADESCO FINANC IAMENT OS') que não batem com o nome real.
    Esta função normaliza ignorando whitespace e compara.
    """
    direto = os.path.join(pasta_cliente, path_relativo)
    if os.path.isdir(direto):
        return direto

    def _norm(s):
        return re.sub(r'\s+', '', s.upper())

    partes = path_relativo.replace('/', '\\').split('\\')
    base_dir = pasta_cliente
    for parte in partes:
        if not os.path.isdir(base_dir):
            return None
        candidatos_disco = os.listdir(base_dir)
        # 1) Match exato
        match_exato = next((n for n in candidatos_disco if n == parte), None)
        if match_exato:
            base_dir = os.path.join(base_dir, match_exato)
            continue
        # 2) Match normalizado (ignora whitespace)
        target = _norm(parte)
        match_fuzzy = next((n for n in candidatos_disco if _norm(n) == target), None)
        if match_fuzzy:
            base_dir = os.path.join(base_dir, match_fuzzy)
            continue
        return None
    return base_dir if os.path.isdir(base_dir) else None


def extrair_numeros_contrato_da_pasta(pasta_acao_abs: str) -> list:
    """Extrai números de contrato a partir dos nomes dos arquivos de procuração
    presentes na pasta_acao. Padrão típico:
        '2- Procuração - Banco Itaú Consignado - Contrato 562732399.pdf'

    Estratégia 1 (primária): regex no nome do arquivo.
    Estratégia 2 (fallback): ESTUDO DE CADEIA.docx, se houver.
    Retorna lista única de strings.
    """
    if not os.path.isdir(pasta_acao_abs):
        return []
    nums = set()
    # 1. Nomes de arquivo
    for nome in os.listdir(pasta_acao_abs):
        if 'procura' not in nome.lower():
            continue
        if not nome.lower().endswith('.pdf'):
            continue
        # Captura número + sufixo opcional "-N" (alguns bancos usam ex: 326994938-8)
        for m in re.finditer(r'[Cc]ontrato\s+(?:n[º°]\s*)?(\d{6,}(?:-\d+)?)', nome):
            nums.add(m.group(1))
    if nums:
        return list(nums)
    # 2. Fallback: ESTUDO DE CADEIA.docx
    for nome in os.listdir(pasta_acao_abs):
        if 'estudo' in nome.lower() and 'cadeia' in nome.lower() and nome.lower().endswith('.docx'):
            try:
                from docx import Document
                d = Document(os.path.join(pasta_acao_abs, nome))
                texto = '\n'.join(p.text for p in d.paragraphs)
                # Padrão: "Contrato: 12345" ou "Contrato nº 12345"
                for m in re.finditer(r'[Cc]ontrato[:\s]+(?:n[º°]\s*)?(\d{6,})', texto):
                    nums.add(m.group(1))
            except Exception:
                pass
            break
    return list(nums)


def agrupar_contratos_por_banco_tese(contratos: list, path_relativo: str,
                                      pasta_acao_abs: str | None = None,
                                      contratos_impugnar_ids: list | None = None,
                                      contratos_impugnar_origem: str | None = None) -> dict:
    """Agrupa contratos por (banco_chave, tese) aplicáveis à pasta_acao.

    Filtros aplicados:
      1. banco_chave do path (suporta litisconsórcio: "BANCO X + BANCO Y")
      2. beneficio_pasta == APOSENTADORIA/PENSAO conforme path
      3. tipo do contrato:
           - se path contém "RMC-RCC" → só tipo RMC ou RCC; tese = rmc/rcc
           - senão → só tipo CONSIGNADO; tese = consignado-nao-contratado
      4. **autoridade da procuração**: se pasta_acao_abs for fornecida, filtra
         contratos pelos números extraídos das procurações (fonte autoritativa
         do escritório). Se nenhum número encontrado, retorna {} (skip).
      5. de-duplica por número de contrato (preferindo Ativo sobre Excluído)

    Retorna {(banco_chave, tese_slug): [contratos]}.
    """
    partes = path_relativo.replace('/', '\\').split('\\')
    if not partes:
        return {}
    # Formatos suportados:
    #   1 nível: "BANCO X" (cliente sem benefício explícito)
    #   2 níveis: "BENEFÍCIO/BANCO X" OU "TESE/BANCO X"
    #   3 níveis: "BENEFÍCIO/TESE/BANCO X" (paradigma Guilherme 2026-05-14)
    BENEFICIOS_KW = {'APOSENTADORIA', 'PENSÃO', 'PENSAO'}
    TESES_KW = {'NÃO CONTRATADO', 'NAO CONTRATADO', 'RMC', 'RCC'}
    # Banco é sempre o último segmento (começa com "BANCO ")
    pasta_banco = partes[-1].strip()
    # Benefício: primeiro segmento se for um dos benefícios conhecidos
    beneficio = ''
    if len(partes) >= 2:
        p0 = partes[0].strip().upper()
        if p0 in BENEFICIOS_KW:
            beneficio = p0
    pasta_banco_norm = re.sub(r'\s+', '', pasta_banco.upper())
    # Tirar acentos para match com CHAVES_BANCO (que é ASCII)
    import unicodedata as _ud
    pasta_banco_norm = _ud.normalize('NFD', pasta_banco_norm).encode('ascii', 'ignore').decode('ascii')
    beneficio = _ud.normalize('NFD', beneficio).encode('ascii', 'ignore').decode('ascii')
    # Detectar tese: legacy (nome banco contém RMC-RCC) OU novo (segmento path = RMC/RCC).
    segmentos_upper = [_ud.normalize('NFD', p.strip().upper()).encode('ascii', 'ignore').decode('ascii')
                       for p in partes]
    is_rmc = 'RMC' in segmentos_upper or 'RMC-RCC' in pasta_banco_norm
    is_rcc = 'RCC' in segmentos_upper or 'RMC-RCC' in pasta_banco_norm
    is_rmc_rcc = is_rmc or is_rcc

    chaves_pasta = [ch for ch in CHAVES_BANCO if ch in pasta_banco_norm]
    # Aliases: se a pasta tem "BANCODOBRASIL" ou "DOBRASIL", também aceita banco_chave='BB'
    if 'BANCODOBRASIL' in chaves_pasta or 'DOBRASIL' in chaves_pasta:
        if 'BB' not in chaves_pasta:
            chaves_pasta.append('BB')
    if not chaves_pasta:
        return {}

    # Autoridade dos contratos a impugnar — hierarquia:
    #   1. Procurações na pasta_acao (autoridade física, assinada pelo cliente)
    #   2. contratos_impugnar_ids do JSON, MAS apenas se origem revisada/manual
    #      (sugestao_automatica sem revisão NÃO tem precedência)
    #   3. contratos_impugnar_ids do JSON (sugestao_automatica) — usado só
    #      quando não há procurações na pasta_acao
    #
    # REGRA CRÍTICA: sem nenhuma das 3 fontes, NÃO se gera notificação
    # (advogado não tem mandato para incluir contratos não outorgados).
    nums_procuracao = []
    if pasta_acao_abs:
        nums_procuracao = extrair_numeros_contrato_da_pasta(pasta_acao_abs)

    ids_impugnar = None
    if contratos_impugnar_ids:
        if contratos_impugnar_origem in ('sugestao_automatica_revisada', 'manual'):
            # Sempre vence — revisão humana já aconteceu
            ids_impugnar = list(contratos_impugnar_ids)
            nums_procuracao = []  # ignora procurações; JSON revisado é fonte
        elif not nums_procuracao:
            # Heurística automática só vale se não houver procurações
            ids_impugnar = list(contratos_impugnar_ids)

    # Sem autoridade explícita → skip (sem mandato, sem ação)
    if not ids_impugnar and not nums_procuracao:
        return {}

    grupos = {}
    for c in contratos:
        bc = c.get('banco_chave')
        # Se há autoridade explícita (ids_impugnar ou procuração), pula filtro
        # de banco_chave — alguns JSONs têm banco_chave=None mas o número da
        # procuração é suficiente para identificar.
        autoridade = bool(ids_impugnar) or bool(nums_procuracao)
        if not autoridade:
            if not bc or bc.upper() not in chaves_pasta:
                continue
        else:
            # Banco da pasta vai ser inferido de outra forma se bc é None
            if bc and bc.upper() not in chaves_pasta:
                continue
        # Filtro benefício
        bp = (c.get('beneficio_pasta') or '').upper()
        if beneficio in ('APOSENTADORIA', 'PENSAO') and bp != beneficio:
            continue
        # Filtro tipo + decisão de tese
        tipo = (c.get('tipo') or '').upper()
        if is_rmc and not is_rcc:
            # Pasta dedicada RMC — só aceita tipo RMC
            if tipo != 'RMC':
                continue
            tese_slug = 'rmc'
        elif is_rcc and not is_rmc:
            # Pasta dedicada RCC — só aceita tipo RCC
            if tipo != 'RCC':
                continue
            tese_slug = 'rcc'
        elif is_rmc_rcc:
            # Pasta legacy RMC-RCC mista
            if tipo not in ('RMC', 'RCC'):
                continue
            tese_slug = tipo.lower()
        else:
            if tipo != 'CONSIGNADO':
                continue
            tese_slug = 'consignado-nao-contratado'
        # Autoridade dos contratos a impugnar
        if ids_impugnar:
            if c.get('id_interno') not in ids_impugnar:
                continue
        elif nums_procuracao:
            num_contrato = str(c.get('contrato', ''))
            # Match exato OU pelo prefixo numérico (tolera sufixo "-N" só de um lado)
            raiz = num_contrato.split('-')[0]
            nums_raiz = {n.split('-')[0] for n in nums_procuracao}
            if num_contrato not in nums_procuracao and raiz not in nums_raiz:
                continue
        # Banco para agrupamento: usa banco_chave do contrato. Se None, tenta
        # inferir do banco_nome_completo do contrato (importante para
        # litisconsórcio passivo "BANCO X + BANCO Y" onde o kit-juridico
        # deixou banco_chave=None mas o nome completo está preenchido).
        # Fallback final: primeiro banco da pasta.
        if bc:
            banco_grupo = bc.upper()
        else:
            bn = (c.get('banco_nome_completo') or '').upper()
            bn_norm = re.sub(r'\s+', '', bn)
            banco_grupo = None
            for chave in chaves_pasta:
                if chave in bn_norm:
                    banco_grupo = chave
                    break
            if not banco_grupo:
                banco_grupo = chaves_pasta[0]
        key = (banco_grupo, tese_slug)
        grupos.setdefault(key, []).append(c)

    # De-duplicar por número de contrato — manter o Ativo se houver
    grupos_dedup = {}
    for key, lista in grupos.items():
        seen = {}
        for c in lista:
            num = c.get('contrato')
            if not num:
                continue
            atual = seen.get(num)
            if atual is None:
                seen[num] = c
            elif c.get('situacao') == 'Ativo' and atual.get('situacao') != 'Ativo':
                seen[num] = c
        # Sempre passa — só chegou aqui quem tem autoridade explícita
        # (procuração ou contratos_impugnar_ids no JSON). A decisão de
        # impugnar já foi feita; situação Ativo/Excluído não restringe.
        if seen:
            grupos_dedup[key] = list(seen.values())
    return grupos_dedup


# UF por extenso COM preposição (templates dizem "estado {{CLIENTE_UF_EXTENSO}}",
# então o valor já entrega "do Amazonas", "de Minas Gerais", etc.)
UF_EXTENSO_COM_PREPOSICAO = {
    'AC': 'do Acre', 'AL': 'de Alagoas', 'AM': 'do Amazonas', 'AP': 'do Amapá',
    'BA': 'da Bahia', 'CE': 'do Ceará', 'DF': 'do Distrito Federal',
    'ES': 'do Espírito Santo', 'GO': 'de Goiás', 'MA': 'do Maranhão',
    'MG': 'de Minas Gerais', 'MS': 'de Mato Grosso do Sul', 'MT': 'de Mato Grosso',
    'PA': 'do Pará', 'PB': 'da Paraíba', 'PE': 'de Pernambuco', 'PI': 'do Piauí',
    'PR': 'do Paraná', 'RJ': 'do Rio de Janeiro', 'RN': 'do Rio Grande do Norte',
    'RO': 'de Rondônia', 'RR': 'de Roraima', 'RS': 'do Rio Grande do Sul',
    'SC': 'de Santa Catarina', 'SE': 'de Sergipe', 'SP': 'de São Paulo',
    'TO': 'do Tocantins',
}


def extrair_numero_logradouro(logradouro: str) -> tuple[str, str]:
    """Separa o número do endereço do logradouro.

    Procura padrões "n° NNN", "nº NNN", ", NNN," ou ", NNN" (no fim).
    Retorna (logradouro_sem_numero, numero). Se não achar número, devolve
    (logradouro_original, '') — o caller decide se grava 'S/N' ou deixa vazio.
    """
    if not logradouro:
        return '', ''
    # Padrão 1: "n° NNN" ou "nº NNN" (com a abreviatura explícita)
    m = re.search(r',?\s*n[º°ºo]\s*(\d+(?:[-/]?\w+)?)', logradouro, re.IGNORECASE)
    if m:
        numero = m.group(1)
        novo = (logradouro[:m.start()] + logradouro[m.end():]).strip().rstrip(',').strip()
        return novo, numero
    # Padrão 2: número solto entre vírgulas — ", 123,"
    m = re.search(r',\s*(\d{1,5}(?:[-/]?\w+)?)\s*(?:,|$)', logradouro)
    if m:
        numero = m.group(1)
        novo = (logradouro[:m.start()] + ',' + logradouro[m.end():]).strip().rstrip(',').strip()
        # Limpa vírgulas duplas que possam ter sobrado
        novo = re.sub(r',\s*,', ',', novo)
        return novo, numero
    # Padrão 3: S/N explícito
    if re.search(r'\bs/?n[º°ºo]?\b', logradouro, re.IGNORECASE):
        novo = re.sub(r',?\s*\bs/?n[º°ºo]?\b', '', logradouro, flags=re.IGNORECASE).strip().rstrip(',').strip()
        return novo, 'S/N'
    return logradouro, ''


def montar_mapa_placeholders(qual: dict, banco_info: dict, advogado: dict,
                              contratos: list, hoje: str, uf_acao: str) -> dict:
    """Monta dict completo de placeholders para passar ao substituir_em_docx.

    O endereço completo do escritório é resolvido a partir da `uf_acao`
    via `_common/escritorios_cadastro.py` (matriz + unidade de apoio na UF).
    """
    # Gênero do cliente
    g = qual.get('genero', 'M')
    if g == 'F':
        nacionalidade = 'brasileira'
        inscrito_a = 'inscrita'
        domiciliado_a = 'domiciliada'
        o_a_notificante = 'a'
        do_da_notificante = 'da'
    else:
        nacionalidade = 'brasileiro'
        inscrito_a = 'inscrito'
        domiciliado_a = 'domiciliado'
        o_a_notificante = 'o'
        do_da_notificante = 'do'

    # Contratos — concatena com "e"
    nums = [c.get('contrato', '?') for c in contratos]
    contrato_numeros = ' e '.join(nums) if nums else '?'

    # Competência inicial dos descontos (template_consignado-nao-contratado__sem
    # tem {{CONTRATO_COMPETENCIA_INICIO}} no parágrafo "início dos descontos em ...").
    # Usa a competência mais antiga entre os contratos do grupo. Formato esperado: MM/AAAA.
    def _comp_key(comp: str) -> str:
        # 'MM/AAAA' -> 'AAAAMM' para ordenar cronologicamente
        if not comp or '/' not in comp:
            return ''
        mm, aaaa = comp.split('/', 1)
        return f'{aaaa}{mm.zfill(2)}'

    competencias = [c.get('competencia_inicio') for c in contratos if c.get('competencia_inicio')]
    contrato_competencia_inicio = min(competencias, key=_comp_key) if competencias else ''

    # Dados do primeiro contrato (mais antigo) — usados em templates que pedem
    # CONTRATO_DATA_INCLUSAO, CONTRATO_QTD_PARCELAS, CONTRATO_VALOR_EMPRESTIMO, CONTRATO_VALOR_PARCELA.
    # Quando há múltiplos contratos, usa o mais antigo por competência.
    contrato_referencia = None
    for c in contratos:
        if contrato_referencia is None:
            contrato_referencia = c
        elif c.get('competencia_inicio') and contrato_referencia.get('competencia_inicio'):
            if _comp_key(c['competencia_inicio']) < _comp_key(contrato_referencia['competencia_inicio']):
                contrato_referencia = c
    if contrato_referencia is None:
        contrato_referencia = {}
    contrato_data_inclusao = contrato_referencia.get('data_inclusao') or ''
    contrato_qtd_parcelas = str(contrato_referencia.get('qtd_parcelas') or '')
    contrato_valor_emprestimo = contrato_referencia.get('valor_emprestado') or ''
    contrato_valor_parcela = contrato_referencia.get('valor_parcela') or ''

    # Endereço — separa número do logradouro (templates Bradesco usam
    # {{CLIENTE_LOGRADOURO}}, n° {{CLIENTE_NUMERO}}, bairro ...).
    logradouro_full = qual.get('logradouro', '')
    logradouro_sem_numero, cliente_numero = extrair_numero_logradouro(logradouro_full)
    if not cliente_numero:
        cliente_numero = 'S/N'

    # UF por extenso com preposição (template_bradesco-* usa
    # "estado {{CLIENTE_UF_EXTENSO}}").
    cliente_uf_extenso = UF_EXTENSO_COM_PREPOSICAO.get(qual.get('uf', '').upper(), '')

    # Nome do banco no contexto fático (versão sem qualificação completa)
    nome_banco_contrato = banco_info['nome_qualificado'].split(',')[0]

    # Blocos condicionais — removem trechos do template quando o campo
    # correspondente está vazio (evita "Cédula de Identidade nº ,").
    # Aplicados ANTES dos placeholders individuais (ordem do dict).
    condicionais = {}
    if not qual.get('rg'):
        # Remove bloco "Cédula de Identidade nº __, órgão expedidor __," INCLUINDO
        # a vírgula final antes de "residente", mantendo só a vírgula que vem antes
        # de "Cédula" (delimitador do CPF anterior).
        condicionais[', Cédula de Identidade nº {{CLIENTE_RG}}, órgão expedidor {{CLIENTE_RG_ORGAO}},'] = ','
        # Variantes (caso template não tenha órgão expedidor)
        condicionais[', Cédula de Identidade nº {{CLIENTE_RG}},'] = ','
    elif not qual.get('rg_orgao'):
        condicionais[', órgão expedidor {{CLIENTE_RG_ORGAO}},'] = ','

    # Endereço incompleto (logradouro vazio)
    if not qual.get('logradouro'):
        condicionais[', residente e {{DOMICILIADO_A}} na {{CLIENTE_LOGRADOURO}}, {{CLIENTE_BAIRRO}}, em {{CLIENTE_MUNICIPIO}}/{{CLIENTE_UF}}, CEP {{CLIENTE_CEP}},'] = ','

    individuais = {
        '{{CIDADE_ASSINATURA}}': advogado['cidade'],
        '{{UF_ASSINATURA}}': advogado['uf'],
        '{{DATA_EXTENSO}}': hoje,

        '{{BANCO_NOME_QUALIFICADO}}': banco_info['nome_qualificado'],
        '{{BANCO_CNPJ}}': banco_info['cnpj'],
        '{{BANCO_LOGRADOURO}}': banco_info['logradouro'],
        '{{BANCO_BAIRRO}}': banco_info['bairro'],
        '{{BANCO_MUNICIPIO}}': banco_info['municipio'],
        '{{BANCO_UF}}': banco_info['uf'],
        '{{BANCO_CEP}}': banco_info['cep'],
        # Banco no meio do texto (curto, sem qualificação completa)
        '{{BANCO_NOME}}': nome_banco_contrato,
        # Compatibilidade temporária com templates antigos (caso algum não tenha sido atualizado)
        '{{NOME_BANCO_CONTRATO}}': nome_banco_contrato,

        '{{CLIENTE_NOME}}': qual.get('nome', ''),
        '{{CLIENTE_NACIONALIDADE_GENERO}}': nacionalidade,
        '{{CLIENTE_ESTADO_CIVIL}}': qual.get('estado_civil', ''),
        '{{CLIENTE_PROFISSAO}}': qual.get('profissao', ''),
        '{{CLIENTE_CPF}}': qual.get('cpf', ''),
        '{{CLIENTE_RG}}': qual.get('rg', ''),
        '{{CLIENTE_RG_ORGAO}}': qual.get('rg_orgao', ''),
        '{{CLIENTE_LOGRADOURO}}': logradouro_sem_numero or qual.get('logradouro', ''),
        '{{CLIENTE_NUMERO}}': cliente_numero,
        '{{CLIENTE_BAIRRO}}': qual.get('bairro', ''),
        '{{CLIENTE_MUNICIPIO}}': qual.get('municipio', ''),
        '{{CLIENTE_UF}}': qual.get('uf', ''),
        '{{CLIENTE_UF_EXTENSO}}': cliente_uf_extenso,
        '{{CLIENTE_CEP}}': qual.get('cep', ''),
        '{{INSCRITO_A}}': inscrito_a,
        '{{DOMICILIADO_A}}': domiciliado_a,
        # Notificante — gênero flexível (criado no pente fino)
        '{{O_A_NOTIFICANTE}}': o_a_notificante,
        '{{DO_DA_NOTIFICANTE}}': do_da_notificante,

        '{{ADVOGADO_NOME}}': advogado['nome'],
        # Compatibilidade temporária — templates novos só usam {{ADVOGADO_NOME}}
        '{{ADVOGADO_NOME_MAIUSCULO}}': advogado['nome_maiusculo'],
        '{{ADVOGADO_OAB_UF}}': advogado['oab_uf'],
        '{{SEU_SUA_ADVOGADO_A}}': 'sua advogada' if advogado['genero'] == 'F' else 'seu advogado',
        '{{ESCRITORIO_ENDERECO_COMPOSTO}}': montar_endereco_escritorio_completo(uf_acao),

        '{{CONTRATO_NUMEROS}}': contrato_numeros,
        '{{CONTRATO_COMPETENCIA_INICIO}}': contrato_competencia_inicio,
        '{{CONTRATO_DATA_INCLUSAO}}': contrato_data_inclusao,
        '{{CONTRATO_QTD_PARCELAS}}': contrato_qtd_parcelas,
        '{{CONTRATO_VALOR_EMPRESTIMO}}': contrato_valor_emprestimo,
        '{{CONTRATO_VALOR_PARCELA}}': contrato_valor_parcela,
    }

    # Condicionais ANTES de individuais — substitui blocos completos quando campo vazio
    return {**condicionais, **individuais}


def data_extenso(d: date | None = None) -> str:
    if d is None:
        d = date.today()
    meses = ['janeiro', 'fevereiro', 'março', 'abril', 'maio', 'junho',
             'julho', 'agosto', 'setembro', 'outubro', 'novembro', 'dezembro']
    return f'{d.day} de {meses[d.month - 1]} de {d.year}'


def montar_dossie_notificacao(pasta_acao_abs: str, output_dir_banco: str,
                                 advogado_chave: str, contratos_pasta: list) -> dict:
    """Copia os documentos de suporte para a pasta da notificação.

    O dossiê contém:
      - OAB do procurador (assinatura da notificação)
      - Procurações dos contratos específicos do banco
      - Documentos pessoais (RG/CPF, hipossuficiência, comprovante)
      - HISCON (histórico de empréstimo) e HISCRE (histórico de créditos/pagamento)

    Args:
        pasta_acao_abs: pasta_acao da kit-juridico (ex.: '.../APOSENTADORIA/BANCO X')
        output_dir_banco: pasta de destino (ex.: '.../notificacao/BRADESCO')
        advogado_chave: chave do procurador ('tiago' / 'patrick' / etc.)
        contratos_pasta: lista de contratos do banco (cada um com 'contrato')

    Returns:
        dict {copiados: int, faltantes: [str]}
    """
    import shutil
    copiados = 0
    faltantes = []

    # 1. OAB do procurador
    oab_nome = OAB_PDF_POR_PROCURADOR.get(advogado_chave)
    if oab_nome:
        oab_src = os.path.join(OABS_DIR, oab_nome)
        if os.path.exists(oab_src):
            shutil.copy2(oab_src, os.path.join(output_dir_banco, oab_nome))
            copiados += 1
        else:
            faltantes.append(f'OAB do procurador {advogado_chave}')
    else:
        faltantes.append(f'mapeamento de OAB para procurador {advogado_chave!r}')

    # 2. Procurações dos contratos do banco
    if os.path.isdir(pasta_acao_abs):
        numeros_banco = {str(c.get('contrato', '')) for c in contratos_pasta if c.get('contrato')}
        numeros_raiz = {n.split('-')[0] for n in numeros_banco}
        for nome in os.listdir(pasta_acao_abs):
            if not nome.lower().endswith('.pdf'):
                continue
            if 'procura' not in nome.lower():
                continue
            # Procura número do contrato no nome
            m = re.search(r'[Cc]ontrato\s+(?:n[º°]\s*)?(\d{4,}(?:-\d+)?)', nome)
            if not m:
                continue
            num = m.group(1)
            raiz = num.split('-')[0]
            if num in numeros_banco or raiz in numeros_raiz:
                # Encurta nome do dossiê para evitar Windows path-260 com client folders longos
                dest_name = nome
                dest_full = os.path.join(output_dir_banco, dest_name)
                if len(dest_full) > 240:
                    dest_name = f'2. Procuração {num}.pdf'
                    dest_full = os.path.join(output_dir_banco, dest_name)
                shutil.copy2(
                    os.path.join(pasta_acao_abs, nome),
                    dest_full,
                )
                copiados += 1

    # 3. Documentos comuns na pasta_acao — APENAS:
    #    - RG / CPF do cliente (item 3.) e do rogado/testemunhas (item 3.1, 3.2, 3.3)
    #    - HISCON (Histórico de empréstimo grifado) — comprova os descontos
    #    - HISCRE (Histórico de pagamento/crédito) — extrato dos descontos
    # REMOVIDOS do dossiê extrajudicial:
    # - Hipossuficiência (4.) — só para inicial
    # - Comprovante de residência (5.) — só para inicial
    # - Declaração de domicílio (5.1) — só para inicial
    # - RG do declarante terceiro (5.2) — vinculado ao comprovante/declaração que
    #   não vão na notif, então também não envia (decisão 13/05/2026, Gabriel)
    if os.path.isdir(pasta_acao_abs):
        for nome in os.listdir(pasta_acao_abs):
            if not nome.lower().endswith('.pdf'):
                continue
            nome_upper = nome.upper()
            # Pula procurações (já tratadas acima)
            if 'PROCURA' in nome_upper:
                continue
            # Pula hipossuficiência, comprovante e declaração de domicílio + RG
            # do declarante terceiro (item 5.x — vinculados ao comprovante)
            if any(k in nome_upper for k in ['HIPOSS', 'COMPROV', 'RESID', 'DOMICI', 'DECLARANTE TERCEIR']):
                continue
            # Pula explicitamente prefixos "5." e "5.1"/"5.2" (item 5 inteiro fica para inicial)
            if re.match(r'^5(\.\d+)?\s*[-.\s]', nome):
                continue
            # Inclui se for RG/CPF ou HISCON/HISCRE
            if any(k in nome_upper for k in [
                'RG', 'CPF',
                'HIST', 'EMPRESTIMO', 'EMPRÉSTIMO', 'PAGAMENTO', 'CRÉDITO', 'CREDITO',
            ]):
                dest_full = os.path.join(output_dir_banco, nome)
                # Defesa Windows path-260
                if len(dest_full) > 240:
                    # Encurta — extrai número canônico (3, 3.1, 3.2, 6, 7) + sufixo curto
                    import re as _re
                    m_num = _re.match(r'^(\d+(?:\.\d+)?)\.\s*', nome)
                    prefix = m_num.group(1) if m_num else 'X'
                    if 'HIST' in nome_upper and ('CREDIT' in nome_upper or 'PAGAMENT' in nome_upper):
                        dest_full = os.path.join(output_dir_banco, f'7. HISCRE.pdf')
                    elif 'HIST' in nome_upper:
                        dest_full = os.path.join(output_dir_banco, f'6. HISCON.pdf')
                    elif 'ROGAD' in nome_upper:
                        dest_full = os.path.join(output_dir_banco, f'{prefix} RG rogado.pdf')
                    elif 'TESTEM' in nome_upper:
                        dest_full = os.path.join(output_dir_banco, f'{prefix} RG testemunha.pdf')
                    else:
                        dest_full = os.path.join(output_dir_banco, f'{prefix} RG.pdf')
                shutil.copy2(
                    os.path.join(pasta_acao_abs, nome),
                    dest_full,
                )
                copiados += 1

    return {'copiados': copiados, 'faltantes': faltantes}


def processar_cliente(pasta_cliente: str, log: list):
    """Processa um cliente: extrai qualificação + gera notificação por pasta_acao."""
    nome_cliente = os.path.basename(pasta_cliente)
    print(f'\n{"="*70}\n{nome_cliente}\n{"="*70}')
    log.append({'cliente': nome_cliente, 'status': 'processando', 'notificacoes': []})

    estado_path = os.path.join(pasta_cliente, '_estado_cliente.json')
    if not os.path.exists(estado_path):
        print(f'  [SKIP] sem _estado_cliente.json')
        log[-1]['status'] = 'skip-sem-json'
        return

    with open(estado_path, encoding='utf-8') as f:
        estado = json.load(f)

    pastas_acao = estado.get('pastas_acao', [])
    contratos = estado.get('contratos', [])
    if not pastas_acao:
        print(f'  [SKIP] sem pastas_acao')
        log[-1]['status'] = 'skip-sem-pastas'
        return

    # 1. Localizar pasta KIT (pode ter nome variável: KIT, "0. Kit", Kit)
    pasta_kit = None
    for nome_pasta in os.listdir(pasta_cliente):
        full_path = os.path.join(pasta_cliente, nome_pasta)
        if os.path.isdir(full_path) and 'kit' in nome_pasta.lower():
            pasta_kit = full_path
            break

    procuracao = encontrar_procuracao(pasta_kit, pasta_cliente=pasta_cliente)
    if not procuracao:
        print(f'  [SKIP] procuração não localizada')
        log[-1]['status'] = 'skip-sem-procuracao'
        return

    print(f'  Procuração: {os.path.basename(procuracao)}')
    qual = extrair_qualificacao(procuracao, max_pages=3)
    if not qual:
        print(f'  [SKIP] qualificação vazia')
        log[-1]['status'] = 'skip-qualif-vazia'
        return
    qual['nome'] = estado.get('cliente', {}).get('nome_completo', '')
    print(f'  Qualificação: nome={qual.get("nome")!r} cpf={qual.get("cpf")!r} '
          f'rg={qual.get("rg")!r} estado_civil={qual.get("estado_civil")!r}')

    log[-1]['qualificacao'] = {k: v for k, v in qual.items() if k != '_texto_extraido'}

    # 2. Determinar UF da ação (a partir do escritório no JSON)
    uf_acao = estado.get('advogado_responsavel', {}).get('uf_atuacao', 'AL')
    advogado = ADVOGADO_POR_UF.get(uf_acao, ADVOGADO_POR_UF['AL'])
    versao = advogado['versao']

    hoje_extenso = data_extenso()

    # 3-pre. Agrupar entradas de pastas_acao com mesmo path_relativo (caso
    # de banco com múltiplos contratos: o JSON tem 1 entrada por contrato,
    # mas a notificação deve listar TODOS os contratos juntos em 1 só docx).
    # Regra acrescentada 2026-05-14 (paradigma FERNANDO FACTA com 2 contratos).
    from collections import defaultdict as _dd
    _agg = _dd(lambda: {'tese': '', 'contratos_impugnar_ids': [],
                         'contratos_impugnar_origem': '', 'path_relativo': ''})
    for _pa in pastas_acao:
        _p = _pa.get('path_relativo', '')
        _a = _agg[_p]
        _a['path_relativo'] = _p
        _a['tese'] = _pa.get('tese', '') or _a['tese']
        _a['contratos_impugnar_origem'] = (
            _pa.get('contratos_impugnar_origem', '') or _a['contratos_impugnar_origem']
        )
        for _cid in (_pa.get('contratos_impugnar_ids') or []):
            if _cid not in _a['contratos_impugnar_ids']:
                _a['contratos_impugnar_ids'].append(_cid)
    pastas_acao = list(_agg.values())

    # 3. Para cada pasta_acao, gerar notificação
    for pa in pastas_acao:
        path_rel = pa.get('path_relativo', '')
        pasta_acao_abs = _resolver_pasta_acao(pasta_cliente, path_rel)
        if not pasta_acao_abs:
            print(f'  [SKIP] pasta inexistente: {path_rel}')
            continue

        grupos = agrupar_contratos_por_banco_tese(
            contratos, path_rel,
            pasta_acao_abs=pasta_acao_abs,
            contratos_impugnar_ids=pa.get('contratos_impugnar_ids'),
            contratos_impugnar_origem=pa.get('contratos_impugnar_origem'),
        )
        if not grupos:
            print(f'  [SKIP] {path_rel}: sem contratos correspondentes')
            continue

        # Regra escritório (2026-05-14): se a pasta de ação tem só 1 banco,
        # não criar subpasta de banco dentro de notificacao/ — gerar direto.
        # Se tem múltiplos bancos (caso de cadeia), manter subpasta por banco.
        bancos_distintos_na_pasta = len(set(b for (b, _t) in grupos.keys()))

        # Para cada (banco, tese) gera 1 notificação
        for (banco_chave, tese), contratos_pasta in grupos.items():
            banco_info = obter_endereco(banco_chave, uf_acao=uf_acao)
            if not banco_info:
                print(f'  [SKIP] {path_rel}/{banco_chave}: banco não cadastrado')
                continue

            template = TEMPLATES.get((tese, versao))
            # Fallback COM → SEM quando o template COM-escritório não existe
            # (caso típico: Bradesco só tem versão SEM)
            if (not template or not os.path.exists(template)) and versao == 'COM':
                template_sem = TEMPLATES.get((tese, 'SEM'))
                if template_sem and os.path.exists(template_sem):
                    print(f'  [INFO] {path_rel}/{banco_chave}/{tese}: '
                          f'template {tese}/COM não existe → fallback para SEM')
                    template = template_sem
            if not template or not os.path.exists(template):
                print(f'  [SKIP] {path_rel}/{banco_chave}/{tese}: template {tese}/{versao} não existe')
                continue

            mapa = montar_mapa_placeholders(qual, banco_info, advogado, contratos_pasta, hoje_extenso, uf_acao)

            # Subpasta de banco SÓ quando há múltiplos bancos na pasta de ação
            # (regra escritório 2026-05-14: evitar profundidade desnecessária).
            if bancos_distintos_na_pasta > 1:
                output_dir = os.path.join(pasta_acao_abs, 'notificacao', banco_chave)
            else:
                output_dir = os.path.join(pasta_acao_abs, 'notificacao')
            os.makedirs(output_dir, exist_ok=True)
            # Nome curto para evitar Windows path-260 em pastas profundas.
            # Versão longa antiga: 'Notificação Extrajudicial - {banco} - {tese}.docx'
            nome_output = f'Notif - {banco_chave} - {tese.upper()}.docx'
            output_path = os.path.join(output_dir, nome_output)
            # Defesa extra para Windows path-260 (na verdade limite efetivo ~230
            # pois validador depois adiciona "_FALHOU_PLACEHOLDERS.docx" como
            # sufixo se houver placeholders restantes).
            if len(output_path) > 220:
                tese_short = 'RMC' if 'rmc' in tese else ('RCC' if 'rcc' in tese else 'NC')
                nome_output = f'Notif {banco_chave[:6]} {tese_short}.docx'
                output_path = os.path.join(output_dir, nome_output)

            try:
                rel = substituir_em_docx(template, mapa, output_path)
                # Monta o dossiê (OAB + procurações + docs pessoais + HISCON/HISCRE)
                dossie = montar_dossie_notificacao(
                    pasta_acao_abs, output_dir,
                    advogado.get('chave', ''), contratos_pasta,
                )
                print(f'  [OK] {path_rel} → {banco_chave}/{nome_output} '
                      f'({rel["total_substituicoes"]} subs, {dossie["copiados"]} docs anexos)')
                if dossie['faltantes']:
                    print(f'       ⚠ faltantes: {dossie["faltantes"]}')
                log[-1]['notificacoes'].append({
                    'pasta': path_rel,
                    'banco': banco_chave,
                    'tese': tese,
                    'output': output_path,
                    'substituicoes': rel['total_substituicoes'],
                    'dossie_copiados': dossie['copiados'],
                    'dossie_faltantes': dossie['faltantes'],
                })
            except PermissionError:
                # Arquivo já aberto no Word — gera com sufixo
                output_path_alt = output_path.replace('.docx', '_v2.docx')
                try:
                    rel = substituir_em_docx(template, mapa, output_path_alt)
                    dossie = montar_dossie_notificacao(
                        pasta_acao_abs, output_dir,
                        advogado.get('chave', ''), contratos_pasta,
                    )
                    print(f'  [OK*] {path_rel} → {banco_chave}/{os.path.basename(output_path_alt)} '
                          f'(orig estava aberto, {dossie["copiados"]} docs anexos)')
                    log[-1]['notificacoes'].append({
                        'pasta': path_rel, 'banco': banco_chave, 'tese': tese,
                        'output': output_path_alt,
                        'substituicoes': rel['total_substituicoes'],
                        'dossie_copiados': dossie['copiados'],
                        'dossie_faltantes': dossie['faltantes'],
                    })
                except Exception as e:
                    print(f'  [ERR] {path_rel}/{banco_chave}: {e}')
            except Exception as e:
                import traceback
                print(f'  [ERR] {path_rel}/{banco_chave}: {e}')
                traceback.print_exc()

    log[-1]['status'] = 'ok' if log[-1]['notificacoes'] else 'sem-notificacoes'


def gerar_relatorio_pendencias(log: list, pasta_batch: str):
    """Gera _pendencias_notificacoes.docx (1 página, simples) listando
    o que falta preencher manualmente em cada cliente.
    """
    from docx import Document
    from docx.shared import Pt

    CAMPOS = [
        ('cpf', 'CPF'),
        ('rg', 'RG'),
        ('rg_orgao', 'Órgão RG'),
        ('logradouro', 'Logradouro'),
        ('bairro', 'Bairro'),
        ('municipio', 'Município'),
        ('uf', 'UF'),
        ('cep', 'CEP'),
        ('estado_civil', 'Estado civil'),
        ('profissao', 'Profissão'),
    ]

    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    titulo = d.add_paragraph()
    r = titulo.add_run('Pendências — Notificações Extrajudiciais')
    r.bold = True
    r.font.size = Pt(13)

    sub = d.add_paragraph()
    sub.add_run(f'{date.today().strftime("%d/%m/%Y")} — campos não extraídos automaticamente das procurações. '
                'Verificar nos documentos físicos (RG, comprovante de residência, procuração) e preencher nas notificações.')

    tabela = d.add_table(rows=1, cols=3)
    tabela.style = 'Light Grid Accent 1'
    hdr = tabela.rows[0].cells
    hdr[0].text = 'Cliente'
    hdr[1].text = 'Notif geradas'
    hdr[2].text = 'Campos faltantes'

    for entry in log:
        qual = entry.get('qualificacao', {})
        n_notif = len(entry.get('notificacoes', []))
        faltam = [label for k, label in CAMPOS if not qual.get(k)]
        if not faltam and n_notif:
            continue  # cliente OK, omite
        row = tabela.add_row().cells
        row[0].text = entry['cliente']
        row[1].text = str(n_notif) if n_notif else '0 (ver log)'
        row[2].text = ', '.join(faltam) if faltam else '—'

    out = os.path.join(pasta_batch, '_pendencias_notificacoes.docx')
    d.save(out)
    print(f'\nRelatório de pendências: {out}')


def main():
    if not os.path.isdir(PASTA_BATCH):
        print(f'PASTA_BATCH não existe: {PASTA_BATCH}')
        return

    clientes = sorted(
        os.path.join(PASTA_BATCH, n) for n in os.listdir(PASTA_BATCH)
        if os.path.isdir(os.path.join(PASTA_BATCH, n))
    )
    # Suporte a filtro de cliente único via env var: NOTIF_FILTRO_CLIENTE=ANAIZA
    filtro = os.environ.get('NOTIF_FILTRO_CLIENTE')
    if filtro:
        clientes = [c for c in clientes if filtro.lower() in os.path.basename(c).lower()]
    print(f'{len(clientes)} clientes encontrados em {PASTA_BATCH}')

    log = []
    for pasta in clientes:
        processar_cliente(pasta, log)

    # Salvar log
    log_path = os.path.join(PASTA_BATCH, '_log_notificacoes.json')
    with open(log_path, 'w', encoding='utf-8') as f:
        json.dump(log, f, indent=2, ensure_ascii=False)

    # Relatório de pendências (campos vazios em cada cliente)
    gerar_relatorio_pendencias(log, PASTA_BATCH)

    # Sumário
    print('\n' + '='*70)
    print('SUMÁRIO')
    print('='*70)
    for entry in log:
        n = len(entry.get('notificacoes', []))
        print(f'  {entry["cliente"]:50s} {entry["status"]:25s} {n} notif')
    total_notif = sum(len(e.get('notificacoes', [])) for e in log)
    print(f'\nTotal: {total_notif} notificações geradas em {len(log)} clientes')
    print(f'Log: {log_path}')


if __name__ == '__main__':
    main()
