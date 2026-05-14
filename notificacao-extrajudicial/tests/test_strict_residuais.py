"""Teste de regressão para a trava strict adicionada em substituir_em_docx."""
import os, sys, tempfile
sys.path.insert(0, r"C:\Users\gabri\.claude\skills\notificacao-extrajudicial\scripts")
from docx import Document
from docx_replace import substituir_em_docx, PlaceholdersResiduaisError

tmp = tempfile.mkdtemp()
template = os.path.join(tmp, 'tpl.docx')
out = os.path.join(tmp, 'out.docx')

doc = Document()
p = doc.add_paragraph()
p.add_run('Olá {{NOME}}, seu CPF é {{CPF}} e {{TOKEN_QUE_NAO_EXISTE}}.')
doc.save(template)

# 1) Sem strict, devolve residuais
res = substituir_em_docx(template, {'{{NOME}}': 'João', '{{CPF}}': '111'}, out, strict=False)
assert res['residuais'] == ['TOKEN_QUE_NAO_EXISTE'], res
print('OK strict=False devolve residuais')

# 2) Com strict (default), levanta E renomeia
out2 = os.path.join(tmp, 'out2.docx')
try:
    substituir_em_docx(template, {'{{NOME}}': 'João'}, out2)
    raise AssertionError('devia ter levantado')
except PlaceholdersResiduaisError as e:
    assert 'CPF' in e.residuais, e.residuais
    assert 'TOKEN_QUE_NAO_EXISTE' in e.residuais
    assert e.dst_path.endswith('_FALHOU_PLACEHOLDERS.docx'), e.dst_path
    assert os.path.exists(e.dst_path)
    assert not os.path.exists(out2), 'arquivo limpo deveria ter sido renomeado'
    print('OK strict=True levanta + renomeia')

# 3) Tudo preenchido, devolve normalmente
out3 = os.path.join(tmp, 'out3.docx')
res = substituir_em_docx(template, {
    '{{NOME}}': 'João', '{{CPF}}': '111', '{{TOKEN_QUE_NAO_EXISTE}}': 'X'
}, out3)
assert res['residuais'] == [], res
assert os.path.exists(out3)
print('OK preenchido salva sem levantar')

print('\nTODOS OS TESTES PASSARAM.')
