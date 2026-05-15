"""Teste: garante que nenhum template ativo contém frases banidas.

Como o usuário pede para remover um parágrafo de um template, é fácil
esquecer outras versões (com-escritório/sem-escritório, RMC/RCC/NC,
Bradesco-tarifas/encargos/PE/capitalização). Este teste varre TODOS os
`template_*.docx` em `assets/` e quebra se encontrar qualquer frase da
lista PARAGRAFOS_PROIBIDOS.

Roda standalone:
    python -m pytest tests/test_templates_sem_paragrafos_proibidos.py -v
    OU
    python tests/test_templates_sem_paragrafos_proibidos.py

Origem (paradigma): LADIMIR DA SILVA 2026-05-15 — o parágrafo "Conforme
informações constantes do próprio extrato de benefício..." havia sido
removido dos templates Bradesco e RMC/RCC, mas o NC sem-escritório
(Patrick AM) ficou intocado e reapareceu em 21 notificações do lote.
"""
import os
import sys
try:
    import pytest
    _HAS_PYTEST = True
except ImportError:
    _HAS_PYTEST = False
    # Stub para permitir importação sem pytest (modo CLI)
    class _PytestStub:
        class mark:
            @staticmethod
            def parametrize(*a, **kw):
                def deco(f): return f
                return deco
    pytest = _PytestStub()
from docx import Document

# Frases que NÃO podem aparecer em nenhum template ativo
PARAGRAFOS_PROIBIDOS = [
    # Removido em 2026-05-13 (Bradesco) e 2026-05-15 (NC sem-escritório)
    # — caso paradigma: LADIMIR DA SILVA + 20 outros clientes do lote AM
    "Conforme informações constantes do próprio extrato de benefício",
    # Removido em 2026-05-13 (Bradesco) — não pode voltar
    "Tais descontos foram identificados",
]

ASSETS_DIR = os.path.abspath(
    os.path.join(os.path.dirname(__file__), '..', 'assets')
)


def listar_templates_ativos():
    """Retorna lista de templates atualmente em uso (ignora .bak)."""
    if not os.path.isdir(ASSETS_DIR):
        return []
    out = []
    for f in sorted(os.listdir(ASSETS_DIR)):
        if not f.startswith('template_'):
            continue
        if not f.endswith('.docx'):
            continue
        if '.bak' in f.lower():
            continue
        out.append(os.path.join(ASSETS_DIR, f))
    return out


def extrair_texto(docx_path):
    """Junta todos os parágrafos + cells da tabela num único string."""
    d = Document(docx_path)
    partes = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        partes.append(p.text)
    return '\n'.join(partes)


@pytest.mark.parametrize('template_path', listar_templates_ativos())
@pytest.mark.parametrize('frase', PARAGRAFOS_PROIBIDOS)
def test_template_nao_contem_frase_proibida(template_path, frase):
    """Cada par (template, frase) é um caso de teste isolado."""
    texto = extrair_texto(template_path)
    nome = os.path.basename(template_path)
    assert frase not in texto, (
        f'\n\nTemplate "{nome}" ainda contém a frase proibida:\n'
        f'  "{frase}"\n\n'
        f'Remova a linha do XML do template (use python-docx) e atualize '
        f'a lista PARAGRAFOS_PROIBIDOS aqui se for caso de regra nova.'
    )


def main():
    """Modo CLI (sem pytest) — varre tudo e imprime relatório."""
    templates = listar_templates_ativos()
    if not templates:
        print(f'[WARN] Nenhum template encontrado em {ASSETS_DIR}')
        return 1
    achados = []
    for t in templates:
        texto = extrair_texto(t)
        for frase in PARAGRAFOS_PROIBIDOS:
            if frase in texto:
                achados.append((os.path.basename(t), frase))
    if achados:
        print(f'[WARN] {len(achados)} ocorrência(s) de frase proibida:')
        for nome, frase in achados:
            print(f'  - {nome}: "{frase[:60]}..."')
        return 1
    print(f'[OK] {len(templates)} templates limpos (sem frases proibidas).')
    return 0


if __name__ == '__main__':
    sys.exit(main())
