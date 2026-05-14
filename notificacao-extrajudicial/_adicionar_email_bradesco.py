# -*- coding: utf-8 -*-
"""Adiciona o parágrafo do email institucional nos 4 templates Bradesco
(que não tinham). Insere antes do 'Em caso de não atendimento'."""
import sys, io, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
from docx import Document
from lxml import etree
W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

ASSETS = Path(r"C:\Users\gabri\.claude\skills\notificacao-extrajudicial\assets")

EMAIL_TXT = (
    "Outrossim, informa-se que eventual resposta poderá ser encaminhada "
    "diretamente ao e-mail institucional do escritório, qual seja: "
    "notificacoesescritorio@gmail.com."
)

ALVOS = [
    'template_bradesco-tarifas__sem-escritorio.docx',
    'template_bradesco-encargos__sem-escritorio.docx',
    'template_bradesco-capitalizacao__sem-escritorio.docx',
    'template_bradesco-pe__sem-escritorio.docx',
]

for nome in ALVOS:
    caminho = ASSETS / nome
    doc = Document(caminho)
    # Localizar o parágrafo "Em caso de não atendimento"
    alvo_idx = None
    for i, p in enumerate(doc.paragraphs):
        if 'Em caso de não atendimento' in p.text:
            alvo_idx = i
            break
    if alvo_idx is None:
        print(f"  ⚠️ {nome}: parágrafo 'Em caso de não atendimento' NÃO localizado")
        continue
    # Verificar se já tem o email
    ja_tem = any('notificacoesescritorio' in p.text for p in doc.paragraphs)
    if ja_tem:
        print(f"  {nome}: já contém email")
        continue
    # Pegar parágrafo de referência (alvo) para clonar formatação
    p_alvo = doc.paragraphs[alvo_idx]
    # Criar novo parágrafo CLONANDO o p_alvo (preserva pPr/estilo)
    novo_xml = copy.deepcopy(p_alvo._p)
    # Remover runs do clone, deixando apenas pPr
    for r in list(novo_xml.findall(W + 'r')):
        novo_xml.remove(r)
    # Adicionar 1 run com o texto do email
    novo_run = etree.SubElement(novo_xml, W + 'r')
    novo_t = etree.SubElement(novo_run, W + 't')
    novo_t.text = EMAIL_TXT
    # Inserir o novo parágrafo IMEDIATAMENTE ANTES do p_alvo
    p_alvo._p.addprevious(novo_xml)
    doc.save(caminho)
    print(f"  ✓ {nome}: email adicionado antes do par {alvo_idx}")

# Verificação
print("\n=== Verificação ===")
for nome in ALVOS:
    d = Document(ASSETS/nome)
    encontrou = False
    for i, p in enumerate(d.paragraphs):
        if 'notificacoesescritorio' in p.text:
            print(f"  {nome} par {i}: {p.text[:120]}...")
            encontrou = True
            break
    if not encontrou:
        print(f"  ⚠️ {nome}: ainda sem email")
