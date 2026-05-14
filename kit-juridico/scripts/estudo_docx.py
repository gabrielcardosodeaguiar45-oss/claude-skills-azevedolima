"""
Gera ESTUDO DE CADEIA - <Banco>.docx por pasta de ação.

Recebe os componentes (cadeias + isolados) que pertencem àquela pasta
de ação e produz um documento Word profissional contendo:
  - Cabeçalho (cliente, benefício, banco/s)
  - Resumo (qtd procurações, contratos, cadeias)
  - Cadeias detectadas (uma seção por cadeia, com diagrama, tabela, observações)
  - Contratos isolados
  - Observações finais

Uso:
    python estudo_docx.py <output.docx> <componentes.json> <metadata.json>

metadata.json contém:
{
    "cliente": "ANAIZA MARIA DA CONCEIÇÃO",
    "beneficio": "PENSÃO POR MORTE PREVIDENCIÁRIA",
    "nb": "041.645.683-9",
    "banco_pasta": "BANCO ITAU CONSIGNADO",
    "procuracoes": [
        {"pagina": 1, "contrato": "631248310"},
        ...
    ]
}
"""
import sys
import os
import json
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    raise ImportError(
        f"Dependência ausente: {e}. "
        f"Instale via: pip install -r requirements.txt"
    ) from e


# Mapa cor RGB → hex pra Word
def rgb_to_hex(rgb_tuple) -> str:
    r, g, b = rgb_tuple
    return f"{int(r*255):02X}{int(g*255):02X}{int(b*255):02X}"


def add_cell_color(cell, color_hex: str):
    """Define cor de fundo de uma célula."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def gerar_estudo(output_path: str, componentes: list[dict], metadata: dict):
    doc = Document()

    # Estilo padrão Cambria 11
    style = doc.styles["Normal"]
    style.font.name = "Cambria"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === Cabeçalho ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ESTUDO DE CADEIA DE CONTRATOS")
    run.bold = True
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(metadata.get("banco_pasta", ""))
    run.bold = True
    run.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Cliente: {metadata.get('cliente', '')}").italic = True

    if metadata.get("nb"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"Benefício: NB {metadata.get('nb')} — "
                  f"{metadata.get('beneficio', '')}").italic = True

    doc.add_paragraph()

    # === Resumo ===
    procs = metadata.get("procuracoes", [])
    cadeias = [c for c in componentes if c.get("tipo") == "CADEIA"]
    isolados = [c for c in componentes if c.get("tipo") == "ISOLADO"]

    h = doc.add_paragraph()
    h.add_run("RESUMO").bold = True

    doc.add_paragraph(f"• Procurações outorgadas neste banco/benefício: {len(procs)}")
    contratos_total = sum(len(c.get("contratos", [])) for c in componentes)
    doc.add_paragraph(f"• Contratos identificados no extrato: {contratos_total}")
    doc.add_paragraph(f"• Cadeias detectadas: {len(cadeias)}")
    doc.add_paragraph(f"• Contratos isolados (sem cadeia): {sum(len(c.get('contratos',[])) for c in isolados)}")

    doc.add_paragraph()

    # === Cadeias ===
    if cadeias:
        h = doc.add_paragraph()
        h.add_run("CADEIAS DE CONTRATOS").bold = True

        for cad in cadeias:
            cor_rgb = cad.get("cor_grifo", (1.0, 0.95, 0.4))
            cor_hex = rgb_to_hex(cor_rgb)
            cor_nome = cad.get("cor_nome", "Amarelo")
            subtipo = cad.get("subtipo", "CADEIA")

            p = doc.add_paragraph()
            run = p.add_run(f"Cadeia {cad['id']} — {_humanizar_subtipo(subtipo)}")
            run.bold = True
            run.font.size = Pt(12)

            doc.add_paragraph(f"Cor no extrato: {cor_nome}")
            doc.add_paragraph(f"Valor de parcela de referência: {cad.get('valor_parcela_referencia', '?')}")
            doc.add_paragraph(f"Data do refinanciamento: {cad.get('data_referencia', '?')}")

            if subtipo == "PORTABILIDADE_INTER_BANCO":
                bancos_str = " → ".join(cad.get("bancos", []))
                p = doc.add_paragraph()
                p.add_run("Atenção: ").bold = True
                p.add_run(
                    f"Esta cadeia atravessa bancos diferentes ({bancos_str}). "
                    "Os bancos envolvidos devem ser tratados em uma única ação por se "
                    "tratarem de operações sucessivas e relacionadas. A litigância "
                    "fragmentada por banco pode caracterizar abuso processual."
                )
            elif subtipo == "SUBSTITUICAO_BANCO":
                p = doc.add_paragraph()
                p.add_run("Observação: ").bold = True
                p.add_run(
                    "Trata-se de movimentação suspeita: o banco excluiu administrativamente "
                    "o contrato e reabriu novo contrato com mesmo valor de reserva no dia "
                    "seguinte (ou próximo). Padrão típico de renovação artificial sem nova "
                    "manifestação do segurado."
                )
            elif subtipo == "CONSOLIDACAO":
                p = doc.add_paragraph()
                p.add_run("Observação: ").bold = True
                p.add_run(
                    "Consolidação de múltiplos contratos em um único refinanciamento "
                    "(N→1). Verificar se o valor liberado ao consumidor justificou a "
                    "quitação da dívida anterior ou se houve apenas alongamento de prazo."
                )
            elif subtipo == "FRACIONAMENTO":
                p = doc.add_paragraph()
                p.add_run("Observação: ").bold = True
                p.add_run(
                    "Fracionamento de um contrato anterior em múltiplos novos contratos "
                    "(1→N). Pode indicar tentativa de pulverização de saldo para mascarar "
                    "anatocismo."
                )

            # Diagrama
            seq = cad.get("contratos", [])
            diag = doc.add_paragraph()
            diag.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for i, c in enumerate(seq):
                situacao = c.get("situacao", "?")
                run = diag.add_run(f"{c['contrato']} [{situacao}]")
                run.bold = True
                if i < len(seq) - 1:
                    diag.add_run("   →   ")

            # Tabela detalhada
            tabela = doc.add_table(rows=1, cols=6)
            tabela.style = "Light Grid Accent 1"
            hdr = tabela.rows[0].cells
            hdr[0].text = "Contrato"
            hdr[1].text = "Origem"
            hdr[2].text = "Inclusão"
            hdr[3].text = "Exclusão"
            hdr[4].text = "Parcela"
            hdr[5].text = "Empréstimo"
            for cell in hdr:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

            for c in seq:
                row = tabela.add_row().cells
                row[0].text = c.get("contrato", "")
                row[1].text = c.get("origem") or "-"
                row[2].text = c.get("data_inclusao") or "-"
                row[3].text = c.get("data_exclusao") or "-"
                row[4].text = c.get("valor_parcela") or "-"
                row[5].text = c.get("valor_emprestado") or "-"
                add_cell_color(row[0], cor_hex)

            doc.add_paragraph()

    # === Contratos isolados ===
    if isolados:
        h = doc.add_paragraph()
        h.add_run("CONTRATOS SEM CADEIA DETECTADA").bold = True

        for comp in isolados:
            for c in comp.get("contratos", []):
                p = doc.add_paragraph()
                run = p.add_run(f"Contrato {c.get('contrato', '?')}")
                run.bold = True

                t = doc.add_table(rows=0, cols=2)
                t.style = "Light List Accent 1"
                for label, key in [
                    ("Banco", "banco"),
                    ("Situação", "situacao"),
                    ("Origem", "origem"),
                    ("Data inclusão", "data_inclusao"),
                    ("Data exclusão", "data_exclusao"),
                    ("Motivo exclusão", "motivo_exclusao"),
                    ("Valor parcela", "valor_parcela"),
                    ("Valor empréstimo", "valor_emprestado"),
                    ("Quantidade parcelas", "qtd_parcelas"),
                    ("Competência início", "competencia_inicio"),
                    ("Competência fim", "competencia_fim"),
                    ("Tipo", "tipo"),
                ]:
                    v = c.get(key)
                    if v not in (None, "", "-"):
                        row = t.add_row().cells
                        row[0].text = label
                        for run in row[0].paragraphs[0].runs:
                            run.bold = True
                        row[1].text = str(v)
                doc.add_paragraph()

    # === Observações finais ===
    h = doc.add_paragraph()
    h.add_run("OBSERVAÇÕES").bold = True

    if cadeias:
        port = sum(1 for c in cadeias if c["subtipo"] == "PORTABILIDADE_INTER_BANCO")
        consol = sum(1 for c in cadeias if c["subtipo"] == "CONSOLIDACAO")
        frac = sum(1 for c in cadeias if c["subtipo"] == "FRACIONAMENTO")
        subst = sum(1 for c in cadeias if c["subtipo"] == "SUBSTITUICAO_BANCO")
        rec = sum(1 for c in cadeias if c["subtipo"] == "CADEIA_RECURSIVA")

        partes = [f"{len(cadeias)} cadeia(s) de refinanciamento"]
        if port:
            partes.append(f"{port} portabilidade entre bancos")
        if consol:
            partes.append(f"{consol} consolidação(ões) N→1")
        if frac:
            partes.append(f"{frac} fracionamento(s) 1→N")
        if subst:
            partes.append(f"{subst} substituição(ões) administrativa(s) de cartão")
        if rec:
            partes.append(f"{rec} cadeia(s) recursiva(s)")

        doc.add_paragraph(
            "Foram detectadas " + ", ".join(partes) + " envolvendo os contratos "
            f"do {metadata.get('banco_pasta', '')} no benefício "
            f"NB {metadata.get('nb', '')}."
        )

        doc.add_paragraph(
            "A identificação das cadeias permite tratar todos os contratos relacionados "
            "em uma única peça processual, evitando fragmentação de ações e reduzindo o "
            "risco de caracterização de litigância abusiva."
        )

    if isolados:
        doc.add_paragraph(
            f"Os contratos sem cadeia detectada constam como averbações novas isoladas no "
            f"extrato. Sua impugnação deve ser feita junto com os contratos em cadeia, "
            f"mantendo a unidade da ação por banco."
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(date.today().strftime("Data do estudo: %d/%m/%Y"))
    run.italic = True

    doc.save(output_path)
    return output_path


def _humanizar_subtipo(s: str) -> str:
    return {
        "REFIN_DIRETO": "Refinanciamento direto (1→1)",
        "CONSOLIDACAO": "Consolidação (N→1)",
        "FRACIONAMENTO": "Fracionamento (1→N)",
        "PORTABILIDADE_INTER_BANCO": "Portabilidade entre bancos",
        "SUBSTITUICAO_BANCO": "Substituição administrativa de cartão",
        "CADEIA_RECURSIVA": "Cadeia recursiva (refins sucessivos)",
        "ISOLADO": "Contrato isolado",
    }.get(s, s)


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    out = sys.argv[1]
    with open(sys.argv[2], encoding="utf-8") as f:
        comps = json.load(f)
    with open(sys.argv[3], encoding="utf-8") as f:
        meta = json.load(f)
    gerar_estudo(out, comps, meta)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
