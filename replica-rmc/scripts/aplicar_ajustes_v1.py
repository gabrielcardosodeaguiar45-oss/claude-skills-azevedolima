#!/usr/bin/env python3
"""aplicar_ajustes_v1.py — aplica os 4 ajustes cirúrgicos na réplica da Edilce.

Ajustes:
  1. Parágrafos 35 e 39 → Heading 2
  2. Copiar header com timbre do modelo BMG/AL para todas as seções
  3. Margens laterais 3,0 cm → 2,5 cm em todas as seções
  4. Inserir Heading 2 "DA INEXISTÊNCIA DE VENDA CASADA — ITEM 5.5 DA CONTESTAÇÃO" antes do parágrafo 113

Uso:
  python aplicar_ajustes_v1.py --replica <replica.docx> --modelo <modelo_com_timbre.docx> --saida <saida.docx>
"""
from __future__ import annotations

import argparse
import json
import shutil
import sys
from copy import deepcopy
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print(json.dumps({"erro": "python-docx nao instalado. pip install python-docx"}))
    sys.exit(2)

FONT = "Cambria"
SIZE_CORPO = Pt(12)


def force_font_run(run_elem, size_pt: Pt = SIZE_CORPO) -> None:
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        rPr = run_elem.makeelement(qn("w:rPr"), {})
        run_elem.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), FONT)
    val = str(int(size_pt.pt * 2))
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        rPr.append(rPr.makeelement(qn("w:sz"), {qn("w:val"): val}))
    else:
        sz.set(qn("w:val"), val)
    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        rPr.append(rPr.makeelement(qn("w:szCs"), {qn("w:val"): val}))
    else:
        szCs.set(qn("w:val"), val)


def remover_bold(run_elem) -> bool:
    rPr = run_elem.find(qn("w:rPr"))
    if rPr is None:
        return False
    alterado = False
    for tag in ("w:b", "w:bCs"):
        el = rPr.find(qn(tag))
        if el is not None:
            rPr.remove(el)
            alterado = True
    return alterado


def aplicar_heading2(p, doc: Document) -> None:
    """Aplica estilo Heading 2 a um parágrafo, sem bold manual, Cambria 12pt, CENTER."""
    h2 = doc.styles["Heading 2"]
    p.style = h2
    # Remover numPr se existir (evita auto-numbering duplicado)
    pPr = p._element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)
    # Aplicar alinhamento e fonte
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        remover_bold(r._element)
        force_font_run(r._element, SIZE_CORPO)


def copiar_header_do_modelo(doc_destino: Document, doc_modelo: Document) -> int:
    """Copia o header da primeira seção do modelo para todas as seções do destino.
    Retorna a quantidade de imagens copiadas."""
    # Pegar o header da seção 0 do modelo
    sec_modelo = doc_modelo.sections[0]
    header_modelo_elem = sec_modelo.header._element

    imgs_no_modelo = len(header_modelo_elem.findall(".//" + qn("w:drawing")))
    if imgs_no_modelo == 0:
        return 0  # modelo não tem imagens no header — não há nada a copiar

    for sec in doc_destino.sections:
        # Garantir que header não está vinculado à seção anterior
        sec.header.is_linked_to_previous = False
        header_elem = sec.header._element
        # Remover todos os parágrafos existentes do header (exceto o hdr tag em si)
        for child in list(header_elem):
            if child.tag in (qn("w:p"), qn("w:tbl")):
                header_elem.remove(child)
        # Copiar parágrafos do header do modelo
        for child in header_modelo_elem:
            if child.tag in (qn("w:p"), qn("w:tbl")):
                header_elem.append(deepcopy(child))

    return imgs_no_modelo


def ajustar_margens(doc: Document, esquerda_cm: float, direita_cm: float) -> list:
    """Ajusta margens laterais de todas as seções. Retorna lista de alterações."""
    alteracoes = []
    esq_emu = int(esquerda_cm / 2.54 * 914400)
    dir_emu = int(direita_cm / 2.54 * 914400)
    for i, sec in enumerate(doc.sections):
        esq_ant = round(sec.left_margin / 914400 * 2.54, 2)
        dir_ant = round(sec.right_margin / 914400 * 2.54, 2)
        sec.left_margin = esq_emu
        sec.right_margin = dir_emu
        alteracoes.append({
            "secao": i,
            "antes": f"L={esq_ant}cm R={dir_ant}cm",
            "depois": f"L={esquerda_cm}cm R={direita_cm}cm"
        })
    return alteracoes


def criar_paragrafo_heading2_antes(doc: Document, idx_ref: int, texto: str) -> None:
    """Insere um novo parágrafo Heading 2 imediatamente antes do parágrafo de índice idx_ref."""
    p_ref = doc.paragraphs[idx_ref]
    h2 = doc.styles["Heading 2"]

    # Criar novo parágrafo usando XML diretamente (para poder inserir antes)
    novo_p = doc.add_paragraph()
    novo_p.style = h2
    novo_p.text = texto
    novo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Remover numPr do novo parágrafo se existir
    pPr = novo_p._element.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)

    for r in novo_p.runs:
        remover_bold(r._element)
        force_font_run(r._element, SIZE_CORPO)

    # Mover o elemento XML para antes do parágrafo de referência
    p_ref._element.addprevious(novo_p._element)


def verificar_resultado(doc: Document) -> dict:
    result = {
        "header_imgs_por_secao": [],
        "footer_imgs_por_secao": [],
        "margens_por_secao": [],
        "headings2_encontrados": [],
    }
    for i, sec in enumerate(doc.sections):
        hi = len(sec.header._element.findall(".//" + qn("w:drawing")))
        fi = len(sec.footer._element.findall(".//" + qn("w:drawing")))
        esq = round(sec.left_margin / 914400 * 2.54, 2)
        dir_ = round(sec.right_margin / 914400 * 2.54, 2)
        result["header_imgs_por_secao"].append(hi)
        result["footer_imgs_por_secao"].append(fi)
        result["margens_por_secao"].append(f"L={esq}cm R={dir_}cm")

    for i, p in enumerate(doc.paragraphs):
        if p.style.name == "Heading 2":
            result["headings2_encontrados"].append({"idx": i, "texto": p.text.strip()[:80]})

    result["timbre_header_ok"] = all(x >= 1 for x in result["header_imgs_por_secao"])
    result["margens_ok"] = all("2.5" in m for m in result["margens_por_secao"])
    return result


def main() -> int:
    ap = argparse.ArgumentParser(description="Aplica ajustes cirúrgicos v1 na réplica Edilce")
    ap.add_argument("--replica", type=Path, required=True, help="Réplica .docx a corrigir")
    ap.add_argument("--modelo", type=Path, required=True, help=".docx modelo com timbre no header")
    ap.add_argument("--saida", type=Path, required=True, help=".docx final corrigido")
    args = ap.parse_args()

    if not args.replica.exists():
        print(json.dumps({"erro": f"réplica não encontrada: {args.replica}"}))
        return 2
    if not args.modelo.exists():
        print(json.dumps({"erro": f"modelo não encontrado: {args.modelo}"}))
        return 2

    # Copiar réplica para saída (trabalhar na cópia)
    shutil.copy2(args.replica, args.saida)
    doc = Document(args.saida)
    doc_modelo = Document(args.modelo)

    metrics = {}

    # AJUSTE 1: Parágrafos 35 e 39 → Heading 2
    # Busca dinâmica pelos títulos (índices podem variar)
    ajuste1_alvos = [
        "DA INÉPCIA POR SUPOSTA FALTA DE DELIMITAÇÃO DA CONTROVÉRSIA",
        "DA CARÊNCIA DE AÇÃO E SUPOSTA AUSÊNCIA DE PRETENSÃO RESISTIDA",
    ]
    ajuste1_aplicados = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        for alvo in ajuste1_alvos:
            # Comparar normalizando (remover acentos para comparação robusta)
            import unicodedata
            def norm(s):
                return unicodedata.normalize("NFD", s).encode("ascii", "ignore").decode("ascii").upper()
            if norm(alvo) in norm(t) and p.style.name != "Heading 2":
                aplicar_heading2(p, doc)
                ajuste1_aplicados.append({"idx": i, "texto": t[:80]})
                break

    metrics["ajuste1_heading2_aplicados"] = ajuste1_aplicados

    # AJUSTE 2: Copiar header do modelo para todas as seções
    imgs_copiadas = copiar_header_do_modelo(doc, doc_modelo)
    metrics["ajuste2_header_imgs_copiadas"] = imgs_copiadas

    # AJUSTE 3: Margens laterais 3,0 cm → 2,5 cm
    alteracoes_margens = ajustar_margens(doc, esquerda_cm=2.5, direita_cm=2.5)
    metrics["ajuste3_margens"] = alteracoes_margens

    # AJUSTE 4: Inserir Heading 2 antes do parágrafo da venda casada
    # Busca dinâmica: parágrafo após o qual a tese de venda casada começa
    # O revisor indica parágrafo 113 — mas buscamos dinamicamente pelo conteúdo
    ajuste4_feito = False
    texto_heading_venda_casada = "DA INEXISTÊNCIA DE VENDA CASADA — ITEM 5.5 DA CONTESTAÇÃO"

    # Primeiro verificar se o heading já existe
    heading_venda_ja_existe = any(
        "VENDA CASADA" in p.text.upper() and p.style.name == "Heading 2"
        for p in doc.paragraphs
    )

    if heading_venda_ja_existe:
        metrics["ajuste4_venda_casada"] = "heading já existe — nenhuma ação necessária"
    else:
        # Procurar o parágrafo que fala de venda casada (sem ser heading)
        idx_venda_casada = None
        for i, p in enumerate(doc.paragraphs):
            t_upper = p.text.upper()
            # Indicadores de início do bloco de venda casada
            if ("VENDA CASADA" in t_upper or "ITEM 5.5" in t_upper) and p.style.name != "Heading 2":
                # Se é um parágrafo de texto (não outro heading), é o alvo
                if len(p.text.strip()) > 30:  # parágrafo real, não título curto
                    idx_venda_casada = i
                    break

        if idx_venda_casada is not None:
            criar_paragrafo_heading2_antes(doc, idx_venda_casada, texto_heading_venda_casada)
            ajuste4_feito = True
            metrics["ajuste4_venda_casada"] = {
                "inserido_antes_do_paragrafo": idx_venda_casada,
                "texto_heading": texto_heading_venda_casada
            }
        else:
            metrics["ajuste4_venda_casada"] = "ATENÇÃO: parágrafo de venda casada não localizado automaticamente — verificar manualmente"

    # Salvar
    doc.save(args.saida)

    # Verificar resultado
    doc_verif = Document(args.saida)
    verif = verificar_resultado(doc_verif)

    result = {
        "saida": str(args.saida),
        "metrics": metrics,
        "verificacao": verif,
        "status": "OK" if (verif["timbre_header_ok"] and verif["margens_ok"] and len(ajuste1_aplicados) >= 2) else "VERIFICAR",
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
