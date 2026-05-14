#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Pós-processamento cirúrgico da réplica ECI DOS SANTOS SACRAMENTO."""
from __future__ import annotations
import sys, re, json
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

DOCX = r'C:\Users\gabri\OneDrive\Área de Trabalho\Correção\5. Réplica de RMC 1\Réplica - 0001462-16.2026 - ECI DOS SANTOS SACRAMENTO.docx'

doc = Document(DOCX)
FONT = 'Cambria'
SIZE_CORPO = Pt(12)


def force_font_run(run_elem, size_pt=SIZE_CORPO):
    rPr = run_elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = run_elem.makeelement(qn('w:rPr'), {})
        run_elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), FONT)
    val = str(int(size_pt.pt * 2))
    for tag in ('w:sz', 'w:szCs'):
        el = rPr.find(qn(tag))
        if el is None:
            rPr.append(rPr.makeelement(qn(tag), {qn('w:val'): val}))
        else:
            el.set(qn('w:val'), val)


def get_h2_style(doc):
    for s in doc.styles:
        if s.name == 'Heading 2':
            return s
    return doc.styles['Normal']


def get_h1_style(doc):
    for s in doc.styles:
        if s.name == 'Heading 1':
            return s
    return doc.styles['Normal']


# ============================================================
# 1) Corrigir estrutura H1/H2 no documento
#    Identificar quais seções são "subtítulos dentro do mérito"
#    e reclassificar como H2
# ============================================================

# Títulos que são BLOCOS PRINCIPAIS (devem ser H1 — seções grandes)
H1_TITULOS = {
    'SÍNTESE DA CONTESTAÇÃO',
    'DA TEMPESTIVIDADE',
    'DAS PRELIMINARES — IMPUGNAÇÃO ÀS PRELIMINARES DO BANCO RÉU',
    'DAS PRELIMINARES DA PARTE AUTORA',
    'DOS FUNDAMENTOS JURÍDICOS DOS PEDIDOS',
    'DOS HONORÁRIOS ADVOCATÍCIOS',
    'DO JULGAMENTO ANTECIPADO DO MÉRITO',
    'DOS PEDIDOS FINAIS',
}

# Tudo que está dentro dos fundamentos jurídicos deve ser H2
# Detectar pelo intervalo: de 'DOS FUNDAMENTOS JURÍDICOS DOS PEDIDOS' até 'DOS HONORÁRIOS'

paragrafos = doc.paragraphs
idx_fundamentos = None
idx_honorarios = None
idx_tempestividade_placeholder = None

rx_roman = re.compile(r'^[IVXLCDM]+\s*[—\-–]\s*')

for i, p in enumerate(paragrafos):
    t = p.text.strip()
    t_sem_roman = rx_roman.sub('', t).strip()

    if 'FUNDAMENTOS JUR' in t_sem_roman and 'PEDIDOS' in t_sem_roman:
        idx_fundamentos = i
    if 'HONORÁRIOS' in t_sem_roman and 'ADVOCATÍCIOS' in t_sem_roman:
        idx_honorarios = i
    if 'CONFERIR DATA DE INTIMA' in t:
        idx_tempestividade_placeholder = i

print(f'idx_fundamentos={idx_fundamentos}, idx_honorarios={idx_honorarios}')
print(f'idx_tempestividade_placeholder={idx_tempestividade_placeholder}')

h1_style = get_h1_style(doc)
h2_style = get_h2_style(doc)

fixes = 0

for i, p in enumerate(paragrafos):
    if p.style.name != 'Heading 1':
        continue

    t = p.text.strip()
    t_sem_roman = rx_roman.sub('', t).strip()

    # O placeholder de tempestividade não deve ser Heading — deve ser parágrafo vermelho
    if 'CONFERIR DATA DE INTIMA' in t:
        # Reclassificar como Normal com formatação de marcador vermelho
        p.style = doc.styles['Normal']
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        novo_run = p.add_run('[CONFERIR DATA DE INTIMAÇÃO DA CONTESTAÇÃO DO AGIBANK — 02/03/2026]')
        novo_run.bold = True
        novo_run.italic = True
        novo_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = None
        force_font_run(novo_run._element, SIZE_CORPO)
        fixes += 1
        print(f'  Convertido placeholder tempestividade [{i}]')
        continue

    # Subtítulos dentro do mérito (entre DOS FUNDAMENTOS e DOS HONORÁRIOS) → H2
    if idx_fundamentos is not None and idx_honorarios is not None:
        if idx_fundamentos < i < idx_honorarios:
            # Verificar se NÃO é um título principal (títulos muito longos são subtítulos)
            # Todos os subtítulos do mérito aqui são H2
            p.style = h2_style
            # Remover numeração romana do H2
            t_limpo = rx_roman.sub('', t).strip()
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            p.add_run(t_limpo)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                force_font_run(r._element, SIZE_CORPO)
            fixes += 1
            continue

    # Re-numerar os H1 que ficaram fora do bloco do mérito (já estão corretos como H1)

print(f'Total fixes fase 1: {fixes}')

# ============================================================
# 2) Re-numerar todos os H1 em ordem romana correta
# ============================================================
cont_h1 = 0
for p in doc.paragraphs:
    if p.style.name != 'Heading 1':
        continue
    t = p.text.strip()
    # Remover numeração antiga
    t_sem_roman = rx_roman.sub('', t).strip()
    cont_h1 += 1

    def int_to_roman(n):
        tab = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),(100,'C'),(90,'XC'),
               (50,'L'),(40,'XL'),(10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
        out = ''
        for v, s in tab:
            while n >= v:
                out += s; n -= v
        return out

    novo_t = f'{int_to_roman(cont_h1)} — {t_sem_roman}'
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(novo_t)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        force_font_run(r._element, SIZE_CORPO)

print(f'H1 renumerados: {cont_h1}')

# ============================================================
# 3) Remover parágrafo duplicado OAB/RS no fecho
# ============================================================
paragrafos = doc.paragraphs
oab_rs_count = 0
for i, p in enumerate(paragrafos):
    if p.text.strip() == 'OAB/RS 139330A':
        oab_rs_count += 1
        if oab_rs_count > 1:
            # Remover o duplicado
            p._element.getparent().remove(p._element)
            print(f'Removido OAB/RS duplicado no parágrafo [{i}]')
            break

# ============================================================
# 4) Garantir que o fecho está centralizado
# ============================================================
paragrafos = doc.paragraphs
for p in paragrafos:
    t = p.text.strip()
    if t in ('Aracaju/SE, 23 de abril de 2026.', 'Aracaju/SE, 23 de abril de 2026'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = None
    if t.startswith('Tiago de Azevedo Lima'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = None
        for r in p.runs:
            r.bold = True
            force_font_run(r._element, SIZE_CORPO)
    if t.startswith('OAB/'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = None
        pf.left_indent = None

# ============================================================
# 5) Garantir linha em branco antes de cada H1
# ============================================================
paragrafos = doc.paragraphs
to_insert = []
for i, p in enumerate(paragrafos):
    if p.style.name == 'Heading 1' and i > 0:
        prev = paragrafos[i - 1]
        if prev.text.strip():
            to_insert.append(i)

for idx in reversed(to_insert):
    h = doc.paragraphs[idx]
    novo = deepcopy(doc.paragraphs[0]._element)
    for c in list(novo):
        novo.remove(c)
    h._element.addprevious(novo)

print(f'Linhas em branco adicionadas antes de H1: {len(to_insert)}')

doc.save(DOCX)
print(f'SALVO: {DOCX}')

# Verificar resultado final
doc2 = Document(DOCX)
h1s = [p for p in doc2.paragraphs if p.style.name == 'Heading 1']
h2s = [p for p in doc2.paragraphs if p.style.name == 'Heading 2']
print(f'H1: {len(h1s)}, H2: {len(h2s)}, Total parágrafos: {len(doc2.paragraphs)}')
print()
print('=== ESTRUTURA H1 ===')
for p in h1s:
    print(f'  {p.text[:80]}')
print()
print('=== ESTRUTURA H2 ===')
for p in h2s[:5]:
    print(f'  {p.text[:80]}')
if len(h2s) > 5:
    print(f'  ... (+{len(h2s)-5} mais)')
