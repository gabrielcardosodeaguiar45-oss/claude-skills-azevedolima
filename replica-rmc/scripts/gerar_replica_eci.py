#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Gerador da réplica ECI DOS SANTOS SACRAMENTO x BANCO AGIBANK S.A."""
from __future__ import annotations
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\Correção\5. Réplica de RMC 1'
SAIDA_TMP = PASTA + r'\replica_tmp.docx'

doc = Document()
sec = doc.sections[0]
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(3.25)
sec.bottom_margin = Cm(2.75)


def h1(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto.upper())
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def h2(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto.upper())
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def body(texto, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = bold
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def cite(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.left_indent = Emu(1439545)
    pf.line_spacing = 1.0
    return p


def red_marker(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def yellow_marker(texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


# ===========================================================================
# ENDEREÇAMENTO
# ===========================================================================
p0 = doc.add_paragraph()
r0 = p0.add_run(
    'EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) FEDERAL DA 5ª VARA FEDERAL '
    'DA SUBSEÇÃO JUDICIÁRIA DE ARACAJU/SE'
)
r0.bold = True
p0.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p_proc = doc.add_paragraph()
r_proc = p_proc.add_run('Processo nº 0001462-16.2026.4.05.8500')
r_proc.bold = True
p_proc.alignment = WD_ALIGN_PARAGRAPH.RIGHT

body(
    'ECI DOS SANTOS SACRAMENTO, já qualificada nos autos do processo em epígrafe, '
    'beneficiária do INSS (benefício B42 nº 162.964.218-2), CPF nº 336.198.245-68, '
    'RG nº 845.634 SSP/SE, residente e domiciliada na Rua do Meio, s/n, Área Rural, '
    'São Cristóvão/SE, representada pelo advogado que esta subscreve, vem, '
    'respeitosamente, à presença de Vossa Excelência, apresentar '
    'RÉPLICA À CONTESTAÇÃO DO BANCO AGIBANK S.A. '
    '(CNPJ 10.664.513/0001-50), nos autos da presente ação, '
    'pelas razões de fato e de direito a seguir expostas.'
)

# ===========================================================================
# I — SÍNTESE DA CONTESTAÇÃO
# ===========================================================================
h1('SÍNTESE DA CONTESTAÇÃO')

body(
    'O BANCO AGIBANK S.A., representado pelo escritório Renno Machado Advogados '
    '(MG), apresentou contestação em 02/03/2026, suscitando as seguintes teses:'
)
body(
    'No campo das preliminares, o banco réu levantou: a) ausência de comprovante '
    'de residência atualizado, sob o argumento de que o documento datado de '
    '04/07/2025 seria "muito antigo" para atestar o domicílio da autora em '
    'janeiro de 2026; b) suposta fragilidade da procuração, outorgada em '
    '09/07/2025, com insinuação de que o patrono teria ajuizado múltiplos '
    'processos com procurações genéricas; c) incompetência absoluta do Juizado '
    'Especial Federal, sob o pretexto de que a causa exigiria perícia técnica '
    'complexa incompatível com o rito do JEF.'
)
body(
    'No mérito, o banco réu sustentou: a) regularidade absoluta do contrato de '
    'cartão de crédito consignado RMC nº 1517444869, averbado em 27/08/2024, '
    'alegando que foi firmado eletronicamente por biometria facial; b) que a '
    'autora teria utilizado o cartão para compras via Pix; c) validade da '
    'assinatura eletrônica por biometria facial nos termos da IN INSS nº '
    '138/2022; d) presunção de autenticidade da assinatura biométrica não '
    'impugnada, invocando o art. 411, III, do CPC; e) impossibilidade de '
    'inversão do ônus da prova; f) inexistência de danos morais in re ipsa; '
    'g) impossibilidade de converter o RMC em empréstimo consignado tradicional; '
    'h) necessidade de compensação do valor liberado; i) base de cálculo '
    'restrita para eventual repetição em dobro; j) "engano justificável" para '
    'afastar a restituição em dobro.'
)
body(
    'A presente réplica rebate, ponto a ponto, cada uma das teses levantadas '
    'pelo banco réu, demonstrando que nenhuma delas merece acolhimento.'
)

# ===========================================================================
# II — DA TEMPESTIVIDADE
# ===========================================================================
h1('DA TEMPESTIVIDADE')

body(
    'A contestação do BANCO AGIBANK S.A. foi protocolada em 02/03/2026. '
    'O prazo legal para a apresentação de réplica, nos termos dos arts. 350 e '
    '351 do Código de Processo Civil, é de 15 (quinze) dias úteis, contados '
    'da intimação da contestação.'
)
red_marker(
    '[CONFERIR DATA DE INTIMAÇÃO DA CONTESTAÇÃO DO AGIBANK — 02/03/2026]'
)
body(
    'A presente réplica é apresentada dentro do prazo legal, razão pela qual '
    'é tempestiva e deve ser recebida e processada.'
)

# ===========================================================================
# III — DAS PRELIMINARES — IMPUGNAÇÃO ÀS PRELIMINARES DO BANCO RÉU
# ===========================================================================
h1('DAS PRELIMINARES — IMPUGNAÇÃO ÀS PRELIMINARES DO BANCO RÉU')

# 3.1
h2('DA AUSÊNCIA DE COMPROVANTE DE RESIDÊNCIA — PRELIMINAR IMPROCEDENTE')

body(
    'O BANCO AGIBANK S.A. sustenta, à fl. 133 da contestação, que "a parte '
    'autora deixou de apresentar comprovante de residência atualizado em seu '
    'nome, o que se faz indispensável para atestar, de forma idônea, sua '
    'residência e domicílio, na forma dos artigos 319, II do Código de '
    'Processo Civil e 1º da Lei 6.629/79", afirmando ainda que "o comprovante '
    'apresentado é datado de 04/07/2025, sendo muito antigo".'
)
body(
    'A preliminar não merece acolhimento por quatro razões objetivas:'
)
body(
    'a) A legislação processual brasileira não impõe prazo de validade para '
    'comprovantes de residência. Os arts. 319, II, do CPC e 1º da Lei nº '
    '6.629/79, invocados pelo banco réu, exigem apenas a indicação de '
    'endereço, o que foi cumprido pela autora. Não existe qualquer norma que '
    'limite o comprovante a 30, 60 ou 90 dias antes do ajuizamento — a '
    'limitação temporal é construção do banco réu, desprovida de amparo legal.'
)
body(
    'b) O comprovante de residência está em nome da própria autora, ECI DOS '
    'SANTOS SACRAMENTO, referente ao endereço situado no município de São '
    'Cristóvão/SE. O documento é idôneo para comprovar o domicílio da '
    'requerente.'
)
body(
    'c) As aparentes divergências de logradouro (Rua do Meio s/n vs. Estrada '
    'de Rita Cacete 1, Zona Rural) referem-se à mesma localidade rural, no '
    'mesmo município de São Cristóvão/SE, com o mesmo CEP 49100-000. Em '
    'localidades rurais, é comum que o mesmo imóvel receba denominações '
    'distintas conforme a fonte cadastral utilizada. Trata-se de variação de '
    'denominação, não de endereços distintos.'
)
body(
    'd) Ainda que houvesse algum defeito formal, o art. 73 do CPC determina '
    'que o juízo deve intimar a parte para sanar irregularidades de '
    'representação, sendo vedada a extinção sumária. A tentativa de converter '
    'questão eminentemente sanável em "pressuposto processual" é manobra '
    'dilatória incompatível com a boa-fé processual.'
)
body(
    'Por todo o exposto, requer-se a rejeição da preliminar de ausência de '
    'comprovante de residência.'
)

# 3.2
h2('DA PROCURAÇÃO — PRELIMINAR MANIFESTAMENTE PROTELATÓRIA')

body(
    'O banco réu, à fl. 134 da contestação, insinua que "não são raros os '
    'casos em que, munido de uma procuração genérica, o advogado ajuíza '
    'diversos processos em nome da parte que a outorgou, existindo casos em '
    'que estes sequer são conhecidos pela parte", sustentando que a procuração '
    'outorgada em 09/07/2025 seria desatualizada por ter sido firmada mais de '
    '90 dias antes do ajuizamento (27/01/2026).'
)
body(
    'A alegação é totalmente improcedente por razões igualmente objetivas:'
)
body(
    'a) O ordenamento jurídico brasileiro não prevê prazo de validade para '
    'procurações. Os arts. 653 e seguintes do Código Civil e os arts. 104 e '
    '106 do CPC não impõem qualquer limitação temporal à eficácia da outorga '
    'de poderes. O prazo de "90 dias" invocado pelo banco réu é construção de '
    'ficção jurídica sem qualquer fundamento legal.'
)
body(
    'b) A procuração foi outorgada de forma expressa e clara pela autora ao '
    'patrono Tiago de Azevedo Lima, com identificação precisa da outorgante '
    'e poderes específicos para a condução desta demanda judicial.'
)
body(
    'c) A alegação genérica de "múltiplos processos" e suspeita de advocacia '
    'predatória é absolutamente abstrata — o banco réu não apontou um único '
    'processo, um único autor, uma única prova concreta que sustente essa '
    'grave acusação contra o patrono da autora. Trata-se de tentativa de '
    'transformar a defesa legítima de consumidora hipossuficiente em '
    '"fraude processual" sem lastro algum nos autos.'
)
body(
    'd) O incidente próprio para arguir falsidade documental é regulado pelo '
    'art. 430 do CPC. Se o banco réu tivesse qualquer dúvida fundada sobre '
    'a autenticidade da procuração, deveria ter instaurado o incidente '
    'correspondente — não usar a contestação como veículo de insinuações '
    'sem prova.'
)
body(
    'Por todo o exposto, requer-se a rejeição integral da preliminar.'
)

# 3.3
h2('DA INCOMPETÊNCIA DO JEF — PRELIMINAR QUE REFORÇA O PEDIDO DE PERÍCIA DA AUTORA')

body(
    'O banco réu, à fl. 143 da contestação, alega "incompetência absoluta do '
    'Juizado Especial Cível para processar esta causa" sob o argumento de que '
    '"caso a parte autora negue ter usado sua senha pessoal e biometria para '
    'contratar o cartão de crédito, a lide só poderá ser dirimida com a '
    'realização de perícia técnica em arquivos digitais, o que é incompatível '
    'com o rito dos Juizados Especiais".'
)
body(
    'O argumento não merece prosperar, pelo contrário, fortalece a posição '
    'da autora:'
)
body(
    'a) Este Juízo Federal tem competência plena para processar a causa. O '
    'valor da causa é de R$ 26.231,59, dentro do limite de 60 salários '
    'mínimos estabelecido pelo art. 3º da Lei nº 10.259/2001, que rege os '
    'Juizados Especiais Federais.'
)
body(
    'b) O Juizado Especial Federal admite produção de prova pericial quando '
    'necessária. Diferentemente do JEC estadual, o JEF federal tem regramento '
    'próprio (Lei 10.259/2001) e admite perícias técnicas por analogia ao CPC '
    '(arts. 464 e seguintes). A mera necessidade de perícia não afasta a '
    'competência do JEF federal.'
)
body(
    'c) A questão fática central é objetivamente delimitada: o banco réu '
    'juntou laudo digital com graves inconsistências — OS e Browser '
    'registrados como "undefined@undefined", ausência de hash SHA-256, '
    'ausência de geolocalização, canal "App do Consultor" — e não comprovou '
    'que foi a própria autora quem realizou a biometria. O ônus de comprovar '
    'a autenticidade da assinatura é inteiramente do banco (Tema 1061 STJ).'
)
body(
    'd) O banco está usando sua própria incapacidade de comprovar a assinatura '
    'para tentar deslocar a competência do juízo. É comportamento contraditório, '
    'vedado pela boa-fé processual (art. 5º do CPC): quem não consegue provar '
    'o que alega não pode usar esse fracasso probatório como argumento '
    'processual.'
)
body(
    'e) A presente réplica reitera o pedido de designação de perícia '
    'grafotécnica e, subsidiariamente, perícia digital nos arquivos '
    'eletrônicos do banco réu — plenamente cabíveis neste Juízo.'
)
body(
    'Por todo o exposto, requer-se a rejeição da preliminar de incompetência.'
)

# ===========================================================================
# IV — DAS PRELIMINARES DA PARTE AUTORA
# ===========================================================================
h1('DAS PRELIMINARES DA PARTE AUTORA')

h2('DA AUSÊNCIA DO TERMO DE CONSENTIMENTO ESCLARECIDO — NULIDADE FORMAL')

body(
    'O contrato de cartão RMC nº 1517444869 foi averbado em 27/08/2024, muito '
    'após dezembro de 2018. Nos termos do art. 21-A da Instrução Normativa '
    'INSS/PRES nº 28/2008, introduzido pela IN nº 83/2018, é obrigatória a '
    'apresentação do Termo de Consentimento Esclarecido (TCE) para a '
    'contratação de cartão de crédito consignado, de modo que a autora seja '
    'devidamente informada sobre a natureza, custos e efeitos do produto '
    'ofertado.'
)
body(
    'Ocorre que o banco réu não juntou aos autos o TCE. Os documentos '
    'acostados à contestação incluem o contrato formal e o laudo de biometria, '
    'mas não o Termo de Consentimento Esclarecido — peça obrigatória para '
    'demonstrar o cumprimento do dever de informação.'
)
body(
    'A ausência do TCE configura nulidade formal autônoma, independente dos '
    'demais vícios demonstrados nesta réplica. A exigência decorre da '
    'ACP 0106890-28.2015.4.01.3700, que firmou obrigação de caráter geral '
    'para todas as instituições financeiras que operam crédito consignado '
    'pelo INSS.'
)
body(
    'Requer-se que este Juízo reconheça a nulidade formal da contratação por '
    'ausência do TCE, com as consequências legais pertinentes.'
)

# ===========================================================================
# V — DOS FUNDAMENTOS JURÍDICOS DOS PEDIDOS
# ===========================================================================
h1('DOS FUNDAMENTOS JURÍDICOS DOS PEDIDOS')

# 5.1
h2('DA INEXISTÊNCIA DO NEGÓCIO JURÍDICO — IMPUGNAÇÃO DA ASSINATURA E TEMA 1061 STJ')

body(
    'A tese principal da parte autora é a inexistência do negócio jurídico '
    'celebrado com o BANCO AGIBANK S.A. por ausência de assinatura válida '
    'da consumidora.'
)
body(
    'O banco réu juntou aos autos o contrato de cartão RMC nº 1517444869. '
    'A parte autora, neste ato, impugna expressamente a assinatura aposta '
    'nesse contrato, requerendo seja designado perito grafotécnico para '
    'confrontar: a) a assinatura do contrato Agibank nº 1517444869; '
    'b) a assinatura constante do RG da autora (nº 845.634 SSP/SE); '
    'c) a assinatura da procuração outorgada em 09/07/2025.'
)
yellow_marker(
    '[MARCADOR PARA O ADVOGADO — INSERIR AQUI: (a) texto descritivo da '
    'comparação gráfica entre a assinatura do contrato Agibank × assinatura '
    'do RG × assinatura da procuração; (b) três imagens lado a lado das '
    'assinaturas para análise comparativa. O redator não inventa essa '
    'comparação — o advogado deve preencher antes do protocolo.]'
)
body(
    'Sobre o ônus probatório, o Superior Tribunal de Justiça, ao julgar o '
    'Tema 1061, firmou a seguinte tese:'
)
cite(
    '"Impugnada pelo consumidor a autenticidade de assinatura em contrato '
    'bancário, cabe à instituição financeira o ônus de comprová-la." '
    '(Tema 1061 STJ)'
)
body(
    'A aplicação do Tema 1061 ao caso é direta: a autora impugna a assinatura '
    'nesta réplica, e o ônus de demonstrar a autenticidade passa integralmente '
    'ao banco réu. Os documentos juntados pelo banco não satisfazem esse ônus '
    '— pelo contrário, o laudo digital apresentado é repleto de '
    'inconsistências que confirmam a irregularidade da contratação, como se '
    'demonstrará a seguir.'
)
body(
    'Diante do exposto, requer-se a designação de perícia grafotécnica para '
    'confronto das assinaturas, nos termos do Tema 1061 STJ, bem como a '
    'declaração de inexistência do negócio jurídico caso a perícia confirme '
    'a inautenticidade da assinatura.'
)

# 5.2
h2('DA IMPRESTABILIDADE DO LAUDO DIGITAL — "UNDEFINED@UNDEFINED" COMO SINAL DE CORRESPONDENTE BANCÁRIO')

body(
    'O banco réu juntou laudo digital relativo à contratação de 27/08/2024. '
    'Esse laudo apresenta inconsistências graves que o tornam imprestável '
    'como prova de contratação válida pela própria autora.'
)
body(
    'A primeira e mais reveladora inconsistência é o registro de OS e Browser '
    'como "undefined@undefined". Em qualquer dispositivo pessoal legítimo '
    '— smartphone, tablet ou computador —, o sistema operacional (OS) e o '
    'navegador (Browser) são identificados automaticamente pelo servidor no '
    'momento da conexão. Os valores "undefined@undefined" indicam ausência '
    'completa de user-agent na requisição, o que ocorre quando se usa '
    'ferramenta automatizada ou plataforma de correspondente bancário que '
    'deliberadamente omite as informações do dispositivo do operador.'
)
body(
    'O laudo digital registra expressamente que o canal de contratação foi '
    '"App do Consultor / Navegador/Link". O "App do Consultor" é a plataforma '
    'utilizada pelos correspondentes bancários credenciados ao Agibank — não '
    'é o aplicativo pessoal da consumidora. Portanto, o laudo não prova que '
    'a autora contratou: prova que um correspondente bancário operou o sistema '
    'usando as credenciais da autora.'
)
body(
    'A segunda inconsistência é a ausência de geolocalização. O laudo não '
    'registra latitude, longitude ou precisão em metros. A geolocalização é '
    'elemento essencial para vincular a contratação ao dispositivo e à '
    'localização da consumidora — sua ausência impede qualquer conclusão '
    'sobre onde e por quem o contrato foi efetivamente firmado.'
)
body(
    'A terceira inconsistência é a ausência de hash SHA-256 de integridade do '
    'documento eletrônico. O Superior Tribunal de Justiça, no REsp 2.159.442/PR '
    '(Rel. Min. Nancy Andrighi, 3ª Turma, j. 24/09/2024), decidiu que a função '
    'hash SHA-256 garante a cadeia de custódia e a integridade do documento '
    'eletrônico — sem o hash, não há como verificar se o documento foi '
    'alterado após a assinatura. O laudo do Agibank não contém qualquer '
    'hash SHA-256, o que compromete irremediavelmente a sua integridade.'
)
body(
    'A quarta inconsistência é a ausência de selfie liveness com certificação '
    'iBeta 2 / IEEE Std 2790/2020. O banco alega que houve "biometria facial", '
    'mas não comprovou que a selfie passou por processo de liveness detection '
    'certificado. Captura de imagem sem liveness detection certificado é '
    'vulnerável a spoofing (uso de foto ou deepfake para enganar o sistema).'
)
body(
    'A quinta inconsistência é a divergência entre a conta de liberação '
    'indicada no laudo e a conta onde o crédito efetivamente caiu. O laudo '
    'digital indica como conta de liberação o Banco do Brasil, Ag. 2611, '
    'Conta 0000367192 — mas o extrato da conta Agibank da própria autora '
    '(ag. 0001, conta 127906610) mostra o crédito de R$ 1.096,94 em '
    '27/08/2024. São contas distintas. Essa divergência interna revela que '
    'os dados registrados no laudo não correspondem ao fluxo real da '
    'operação, sugerindo triangulação intermediada pelo correspondente '
    'bancário.'
)
body(
    'Diante de todas essas inconsistências, o laudo digital é imprestável '
    'como prova de contratação válida pela autora. Requer-se a declaração '
    'de sua ineficácia probatória e, subsidiariamente, a designação de '
    'perícia digital nos arquivos eletrônicos do banco réu.'
)

# 5.3
h2('DA VIOLAÇÃO DA RESOLUÇÃO CNJ Nº 159/2024 — VIDEOCHAMADA OBRIGATÓRIA NÃO COMPROVADA')

body(
    'O contrato de cartão RMC nº 1517444869 foi averbado em 27/08/2024 '
    '— data posterior a 01/03/2024, quando entrou em vigor a Resolução '
    'CNJ nº 159/2024.'
)
body(
    'A Resolução CNJ nº 159/2024 estabeleceu como requisito obrigatório '
    'a realização de videochamada para a averbação de contratos de cartão '
    'de crédito consignado (RMC/RCC) realizados a partir de 01/03/2024. '
    'A videochamada tem por finalidade assegurar que o próprio titular do '
    'benefício está ciente e consente com a contratação.'
)
body(
    'Ocorre que o banco réu, em sua contestação, sequer mencionou a '
    'Resolução CNJ nº 159/2024. Não juntou qualquer comprovante de '
    'realização de videochamada. O silêncio do banco réu sobre esse '
    'requisito, somado à ausência de documentação, é confissão tácita '
    'do descumprimento da norma.'
)
body(
    'Cumpre destacar que a IN INSS/PRES nº 138/2022, art. 5º, II e III, '
    'já havia tornado obrigatório o reconhecimento biométrico certificado '
    'a partir de 01/12/2022. A averbação de 27/08/2024 deveria atender '
    'cumulativamente: biometria certificada (IN 138/2022) e videochamada '
    '(Res. CNJ 159/2024). O banco réu não comprovou nenhum dos dois de '
    'forma suficiente.'
)
body(
    'A ausência de videochamada é vício insanável que contamina toda a '
    'contratação. Requer-se o reconhecimento da nulidade da averbação '
    'por descumprimento da Resolução CNJ nº 159/2024.'
)

# 5.4 — Margem livre
h2('DA MARGEM LIVRE PARA EMPRÉSTIMO CONSIGNADO TRADICIONAL')

body(
    'A autora ECI DOS SANTOS SACRAMENTO, beneficiária do INSS (benefício '
    'B42, nº 162.964.218-2), possui renda mensal de R$ 1.518,00 (um mil, '
    'quinhentos e dezoito reais). Nos termos da legislação de crédito '
    'consignado, a margem disponível para empréstimo consignado tradicional '
    'corresponde a 30% da renda, ou seja, R$ 455,40 (quatrocentos e '
    'cinquenta e cinco reais e quarenta centavos).'
)
body(
    'Na data de averbação do contrato (27/08/2024), havia margem disponível '
    'para empréstimo consignado tradicional — modalidade mais vantajosa '
    'para o consumidor, por ter prazo fixo, parcelas que amortizam '
    'completamente a dívida e custo financeiro significativamente menor '
    'que o cartão RMC.'
)
body(
    'O banco réu, entretanto, não ofereceu à autora a modalidade de '
    'empréstimo consignado tradicional. Em vez disso, averbou o cartão '
    'RMC — produto estruturalmente mais oneroso, no qual as parcelas '
    'descontadas em folha representam apenas o pagamento mínimo de uma '
    'dívida rotativa que não se encerra.'
)
body(
    'Essa conduta viola o dever de informação e transparência (CDC, '
    'arts. 6º, III, e 52) e configura prática abusiva por omissão de '
    'alternativa mais adequada e economicamente mais favorável '
    'ao consumidor hipossuficiente.'
)
red_marker(
    '[INSERIR AQUI: print do extrato INSS / HISCON competência 08/2024 '
    '— mostrando MARGEM DISPONÍVEL e margem consignável de empréstimo '
    'consignado tradicional]'
)

# 5.5 — Onerosidade excessiva
h2('DA ONEROSIDADE EXCESSIVA DO CARTÃO RMC — NÚMEROS CONCRETOS DO CASO')

body(
    'Ponto central e quantitativo desta réplica: os números concretos '
    'demonstram que o cartão RMC é estruturalmente oneroso e que a autora '
    'já pagou mais do que recebeu, sem que a dívida se extinga.'
)

# Tabela de onerosidade (6 colunas)
tbl = doc.add_table(rows=2, cols=6)
tbl.style = 'Table Grid'
cabecalhos = [
    'Valor emprestado',
    'Início dos descontos',
    'Parcela média',
    'Nº parcelas pagas',
    'Valor pago até o momento',
    'Valor refin. maquiador'
]
dados_tbl = [
    'R$ 1.096,94',
    '09/2024',
    'R$ 74,90',
    '20 parcelas (set/2024–abr/2026)',
    '~R$ 1.498,00',
    'R$ 1.062,40 (23/01/2025)'
]
for i, cab in enumerate(cabecalhos):
    cell = tbl.rows[0].cells[i]
    cell.text = cab
    for par in cell.paragraphs:
        for run in par.runs:
            run.bold = True
for i, dado in enumerate(dados_tbl):
    tbl.rows[1].cells[i].text = dado

body(
    'Os números são objetivos: o BANCO AGIBANK S.A. liberou R$ 1.096,94 '
    '(um mil, noventa e seis reais e noventa e quatro centavos) em '
    '27/08/2024. Até abril de 2026 — apenas 20 meses depois — a autora '
    'já pagou aproximadamente R$ 1.498,00 (um mil, quatrocentos e noventa '
    'e oito reais) em parcelas de R$ 74,90 descontadas de seu benefício '
    'do INSS. Isso representa um excedente de R$ 401,06 — aproximadamente '
    '36,56% acima do valor originalmente recebido — e a dívida continua '
    'não quitada.'
)
body(
    'Em um empréstimo consignado tradicional, após 20 parcelas equivalentes, '
    'a dívida estaria integralmente quitada. No cartão RMC, após 20 parcelas, '
    'a autora não apenas não quitou a dívida como ainda sofreu o '
    'refinanciamento maquiador de R$ 1.062,40 em 23/01/2025 — que '
    'artificialmente resetou o saldo devedor, perpetuando os descontos '
    'indefinidamente. Essa é a característica estrutural mais perversa do '
    'cartão RMC: a dívida nunca termina.'
)

# 5.6 — Maquiagem contábil
h2('DA MAQUIAGEM CONTÁBIL PELO REFINANCIAMENTO — "CRÉDITO PESSOAL RENOVADO"')

body(
    'Em 23/01/2025, o extrato da conta Agibank da autora registrou o '
    'lançamento "Crédito Empréstimo - Crédito Pessoal Renovado" no valor '
    'de R$ 1.062,40. Esse lançamento é o que o mercado de crédito chama '
    'de refinanciamento maquiador.'
)
body(
    'O mecanismo funciona da seguinte forma: o banco identifica que o saldo '
    'devedor do cartão (principal + juros acumulados) atingiu determinado '
    'patamar; registra um "novo crédito" do valor do saldo devedor como se '
    'fosse liberação de recursos, mas não há dinheiro novo saindo para o '
    'consumidor. O que ocorre é a rolagem da dívida: o saldo antigo é '
    '"zerado" contabilmente e um novo saldo de R$ 1.062,40 é gerado, '
    'reiniciando o ciclo de descontos.'
)
body(
    'O resultado prático para a autora: mesmo tendo pago parcelas mensais '
    'de R$ 74,90 por vários meses, o saldo devedor foi artificialmente '
    'reinflado a R$ 1.062,40 em janeiro de 2025. Os descontos continuam — '
    'agora para pagar essa "nova dívida" que nada mais é que a dívida '
    'antiga reciclada contabilmente.'
)
body(
    'Essa prática viola os arts. 6º, III, e 51 do CDC (cláusulas abusivas '
    'que colocam o consumidor em desvantagem exagerada). Requer-se a '
    'declaração de nulidade do refinanciamento de 23/01/2025 e a exclusão '
    'de R$ 1.062,40 da base de cálculo de eventual "compensação" '
    'pretendida pelo banco réu.'
)

# 5.7 — Venda casada
h2('DA VENDA CASADA — SEGURO "AGI PROTEGE" EMBUTIDO SEM SOLICITAÇÃO')

body(
    'A primeira fatura do cartão RMC registra cobrança do produto '
    '"Agi Protege" no valor de R$ 176,39 (cento e setenta e seis reais '
    'e trinta e nove centavos), seguida de cobranças mensais de R$ 17,99 '
    '(dezessete reais e noventa e nove centavos) a título de '
    '"Débito de Seguro" nas faturas subsequentes.'
)
body(
    'A autora não solicitou qualquer seguro. O produto "Agi Protege" foi '
    'embutido na contratação do cartão RMC sem o consentimento expresso '
    'da consumidora, configurando típica venda casada vedada pelo '
    'art. 39, I e III, do Código de Defesa do Consumidor.'
)
body(
    'O Superior Tribunal de Justiça, por meio da Súmula 532, consagrou que:'
)
cite(
    '"Constitui prática abusiva o envio de cartão de crédito sem prévia '
    'e expressa solicitação do consumidor, configurando-se ato ilícito '
    'indenizável e sujeito à aplicação de multa administrativa." '
    '(Súmula 532 STJ)'
)
body(
    'O caso presente é ainda mais grave: além do cartão não solicitado, '
    'o banco embutiu um seguro igualmente não solicitado, cobrado '
    'imediatamente na primeira fatura e continuado mensalmente. '
    'Trata-se de dupla prática abusiva: cartão RMC não solicitado + '
    'seguro "Agi Protege" não solicitado.'
)
body(
    'Requer-se: a) a declaração de nulidade da cláusula de seguro '
    '"Agi Protege"; b) a cessação imediata das cobranças do seguro; '
    'c) a restituição em dobro de todos os valores cobrados a esse título '
    '(R$ 176,39 na 1ª fatura + R$ 17,99/mês nas faturas subsequentes), '
    'em regime integral de dobro, por se tratarem de cobranças integralmente '
    'posteriores ao marco de 30/03/2021 (EAREsp 676.608/RS — STJ).'
)

# 5.8 — Invalidade digital
h2('DA INVALIDADE DA ASSINATURA ELETRÔNICA — HASH, GEOLOCALIZAÇÃO E LIVENESS')

body(
    'A Lei nº 14.063/2020 estabelece três modalidades de assinatura '
    'eletrônica: a) simples; b) avançada; c) qualificada (ICP-Brasil). '
    'Para que uma assinatura eletrônica seja classificada como "avançada" '
    '— e portanto idônea para a contratação de serviços financeiros de '
    'alto impacto patrimonial —, é necessário que utilize dados criados '
    'com alto nível de confiança e vinculados univocamente ao signatário, '
    'incluindo rastreabilidade e integridade verificável do documento.'
)
body(
    'O laudo digital do Agibank não satisfaz os requisitos de assinatura '
    'avançada porque: a) não há hash SHA-256 — sem hash, não há '
    'integridade verificável do documento (REsp 2.159.442/PR, Min. Nancy '
    'Andrighi, STJ, j. 24/09/2024); b) não há geolocalização — sem '
    'latitude/longitude registrados, é impossível vincular a assinatura '
    'ao dispositivo e localização da autora; c) OS e Browser '
    '"undefined@undefined" — ausência de user-agent impede a identificação '
    'do dispositivo real utilizado na operação.'
)
body(
    'Sem esses elementos, a assinatura eletrônica eventualmente capturada '
    'pelo correspondente bancário não atinge sequer o padrão de assinatura '
    '"avançada" da Lei nº 14.063/2020, sendo imprestável para gerar '
    'obrigações da magnitude dos descontos compulsórios em benefício '
    'previdenciário.'
)
body(
    'Reforça esse entendimento a jurisprudência analógica do TJAL em '
    'casos envolvendo contratos eletrônicos similares:'
)
cite(
    'TJAL, Apelação 0745371-43.2022.8.02.0001 (4ª CC, Des. Márcio Roberto '
    'Tenório, j. 15/05/2024): "Nulidade por ausência de selfie, protocolo, '
    'geolocalização e ID de usuário em contrato eletrônico."'
)
cite(
    'TJAL, Apelação 0740736-48.2024.8.02.0001 (4ª CC, Des. Márcio Tenório, '
    'j. 05/02/2025): "Contrato eletrônico exige geolocalização, ID de '
    'sessão e IP/terminal identificável."'
)

# 5.9 — Prática abusiva
h2('DA PRÁTICA ABUSIVA E NULIDADE DO CONTRATO DE CARTÃO RMC')

body(
    'Nos termos do art. 39 do CDC, é vedado ao fornecedor, dentre outras '
    'práticas abusivas: "condicionar o fornecimento de produto ou serviço '
    'ao fornecimento de outro produto ou serviço" (inciso I); '
    '"prevalecer-se da fraqueza ou ignorância do consumidor" (inciso IV); '
    'e "exigir do consumidor vantagem manifestamente excessiva" (inciso V).'
)
body(
    'O banco réu violou esses três incisos de forma simultânea: '
    'a) vinculou o crédito à contratação de cartão RMC em vez de oferecer '
    'o empréstimo consignado tradicional (inciso I); '
    'b) prevaleceu-se da hipossuficiência da autora — aposentada '
    '(benefício B42), 59 anos, residente em área rural, renda de '
    'R$ 1.518,00 — para impor produto desvantajoso (inciso IV); '
    'c) impôs condição estruturalmente mais onerosa que resulta em '
    'pagamentos perpétuos sem amortização real da dívida (inciso V).'
)
body(
    'As cláusulas contratuais que estabelecem o modelo de cobrança do '
    'cartão RMC são nulas de pleno direito por violação ao art. 51, IV, '
    'do CDC, que considera nula toda cláusula "que estabeleça obrigações '
    'consideradas iníquas, abusivas, que coloquem o consumidor em '
    'desvantagem exagerada, ou sejam incompatíveis com a boa-fé ou '
    'a equidade".'
)

# 5.10 — Vício de consentimento (subsidiária)
h2('DA PRESENÇA DE VÍCIO DE CONSENTIMENTO NA MODALIDADE — TESE SUBSIDIÁRIA')

body(
    'Ainda que Vossa Excelência, por alguma razão, não acolha a tese '
    'principal de inexistência do negócio jurídico — o que se admite '
    'apenas para argumentar —, a autora faz jus à declaração de nulidade '
    'do contrato de cartão RMC por vício de consentimento quanto à '
    'modalidade contratada.'
)
body(
    'A autora ECI DOS SANTOS SACRAMENTO buscou, em agosto de 2024, '
    'contratar um empréstimo consignado tradicional — modalidade que '
    'conhece e que já utilizou ao longo de sua vida previdenciária. '
    'O empréstimo consignado tradicional tem características claras: '
    'valor emprestado fixo, prazo determinado, parcelas fixas que '
    'amortizam o saldo devedor até a quitação completa.'
)
body(
    'O banco réu, entretanto, averbou um cartão de crédito consignado '
    '(RMC) — produto funcionalmente distinto: não tem prazo fixo, não '
    'amortiza completamente a dívida pelo desconto mensal, e o que a '
    'autora recebeu foi um crédito de R$ 1.096,94 em sua conta Agibank '
    'estruturado como saque de cartão, não como empréstimo parcelado '
    'com data de quitação.'
)
body(
    'Esse erro substancial sobre a natureza do negócio configura vício '
    'de consentimento previsto nos arts. 138 e 139, I, do Código Civil. '
    'A autora foi induzida a erro sobre a espécie do produto contratado '
    '— foi levada a acreditar que celebrava empréstimo consignado '
    'tradicional e se viu presa em cartão rotativo de desconto '
    'compulsório perpétuo.'
)
body(
    'Nesta tese subsidiária, a parte autora reconhece expressamente o '
    'recebimento do valor de R$ 1.096,94 creditado em sua conta Agibank '
    '(ag. 0001, conta 127906610) em 27/08/2024, e requer, desde já, a '
    'compensação desse valor em eventual condenação à restituição dos '
    'valores indevidamente descontados. O que se contesta não é o '
    'recebimento do crédito, mas a modalidade do produto imposto.'
)
body(
    'Registra-se, por fim, que a conta de liberação indicada no laudo '
    'digital (Banco do Brasil, Ag. 2611, Conta 0000367192) é diversa '
    'da conta Agibank onde o crédito efetivamente foi depositado '
    '(ag. 0001, conta 127906610). Essa divergência interna — laudo '
    'aponta para o BB, mas o dinheiro caiu na conta Agibank da própria '
    'autora — sugere triangulação intermediada pelo correspondente '
    'bancário que operou o "App do Consultor". O Pix imediato para '
    'a própria autora no mesmo dia reforça a hipótese de que o '
    'correspondente intermediou o recebimento e o repasse do valor.'
)

# 5.11 — Hipervulnerabilidade dupla
h2('DA HIPERVULNERABILIDADE DUPLA — RCC BANCO PAN ATIVO')

body(
    'Importa noticiar, para a correta compreensão da situação '
    'patrimonial da autora, que ela também é titular do cartão de '
    'crédito consignado RCC do Banco PAN (contrato nº 764211206-9, '
    'averbado em 19/09/2022), ainda ativo e gerando descontos mensais '
    'em seu benefício INSS.'
)
body(
    'Esse contrato gêmeo (RCC PAN) não é objeto de impugnação na '
    'presente ação — será tratado em ação autônoma. Todavia, a '
    'existência simultânea de dois contratos de cartão consignado '
    '(RMC Agibank + RCC PAN) consumindo a margem da autora é fator '
    'agravante de hipervulnerabilidade que este Juízo deve considerar '
    'ao apreciar os pedidos desta ação.'
)
body(
    'A autora, com renda de R$ 1.518,00 mensais, tem sua margem '
    'comprometida por pelo menos dois contratos de cartão consignado '
    'simultâneos de bancos diferentes, além do seguro "Agi Protege". '
    'Esse cenário de compressão múltipla do benefício previdenciário '
    'revela o estado de hipervulnerabilidade extrema da consumidora '
    'e agrava o dano moral sofrido.'
)

# 5.12 — 2ª via massiva
h2('DAS FATURAS DO CARTÃO — 2ª VIA MASSIVA EMITIDA PARA FINS PROCESSUAIS')

body(
    'O BANCO AGIBANK S.A. acostou à contestação 17 (dezessete) faturas '
    'do cartão RMC nº 1517444869 (terminação 7839). Ocorre que todas '
    'as 17 (dezessete) faturas foram emitidas e postadas na mesma data: '
    '10/02/2026 — aproximadamente 20 dias antes do protocolo da própria '
    'contestação (02/03/2026).'
)
body(
    'Essa circunstância evidencia que se trata de 2ª via gerada em lote '
    'pelo banco réu, sem qualquer prova de que as faturas mensais foram '
    'regularmente enviadas à autora no momento de cada competência. '
    'O banco extraiu do sistema 17 meses de faturas de uma só vez, '
    'imediatamente antes de protocolar sua defesa.'
)
body(
    'A emissão concentrada das faturas em data única demonstra que: '
    'a) as faturas originais não foram enviadas à autora mensalmente, '
    'descumprindo o art. 52, V, do CDC, que exige informação prévia e '
    'ostensiva sobre encargos; b) a autora nunca recebeu as cobranças '
    'mensais regularmente, o que agrava sua condição de hipossuficiência '
    'informacional; c) as faturas juntadas são cópias de segunda via '
    'produzidas exclusivamente para fins processuais, sem valor '
    'probatório autônomo para comprovar o envio regular ao consumidor.'
)
body(
    'Requer-se que este Juízo desconsidere as faturas como prova '
    'suficiente da regularidade da contratação e dos descontos, '
    'reconhecendo sua natureza de segunda via processual. Nos termos '
    'do art. 6º, III, do CDC, o banco deveria ter demonstrado que '
    'prestou à autora informação adequada e ostensiva sobre cada '
    'cobrança mensal no momento de sua realização — o que não ocorreu.'
)

# 5.13 — Restituição
h2('DA RESTITUIÇÃO DOS VALORES INDEVIDAMENTE DESCONTADOS — DOBRO INTEGRAL')

body(
    'Todos os descontos realizados pelo BANCO AGIBANK S.A. no benefício '
    'previdenciário da autora, a título do cartão RMC nº 1517444869, '
    'são posteriores a 27/08/2024 — muito depois do marco de 30/03/2021 '
    'estabelecido pelo Superior Tribunal de Justiça no julgamento do '
    'EAREsp 676.608/RS.'
)
body(
    'O EAREsp 676.608/RS fixou a tese de que a restituição em dobro do '
    'art. 42, parágrafo único, do CDC é devida para cobranças indevidas '
    'realizadas após 30/03/2021, independentemente de comprovação de '
    'má-fé subjetiva do fornecedor — basta a ilicitude objetiva da '
    'cobrança. Como todos os descontos impugnados são de agosto de 2024 '
    'em diante, o regime de restituição é integralmente em dobro.'
)
body(
    'O banco réu alega "engano justificável" para afastar a dobra. '
    'O argumento não prospera. O banco embutiu seguro não solicitado, '
    'realizou refinanciamento maquiador, operou por correspondente '
    'não identificado e deixou a autora sem as faturas mensais regulares. '
    'Esses fatos, em conjunto, afastam qualquer "engano justificável" '
    'e confirmam a natureza sistemática da cobrança indevida.'
)
body(
    'No que toca à compensação: a parte autora EXPRESSAMENTE CONCORDA '
    'com a compensação do valor de R$ 1.096,94 recebido em 27/08/2024 '
    'em conta Agibank, e requer desde já que esse valor seja deduzido '
    'da base de cálculo da restituição em dobro. O que se contesta é '
    'a fórmula proposta pelo banco réu — que pretende compensar apenas '
    'o "excesso" dos descontos acima do valor recebido —, quando a base '
    'de cálculo correta para a dobra é o total das parcelas cobradas '
    'indevidamente, deduzido o TED recebido, e não apenas o saldo '
    'remanescente após a compensação simples.'
)
body(
    'Requer-se a restituição em dobro de: a) todas as parcelas '
    'descontadas do benefício a título do RMC nº 1517444869, '
    'a partir de 27/08/2024; b) todos os valores cobrados a título '
    'do seguro "Agi Protege" (R$ 176,39 na 1ª fatura + R$ 17,99/mês '
    'nas faturas subsequentes); c) deduzido o valor de R$ 1.096,94 '
    'recebido via TED em 27/08/2024 — valor exato a apurar em '
    'cumprimento de sentença, nos termos do Tema 929 STJ.'
)

# 5.14 — Danos morais
h2('DOS DANOS MORAIS IN RE IPSA — R$ 15.000,00')

body(
    'O banco réu sustenta que "a ocorrência de descontos em folha '
    'derivados de contratação anulada não constitui hipótese de dano '
    'in re ipsa". O argumento é equivocado e contraria a jurisprudência '
    'consolidada.'
)
body(
    'O dano moral in re ipsa em cobranças indevidas sobre benefício '
    'previdenciário de consumidora hipossuficiente é reconhecido pela '
    'jurisprudência nacional, que dispensa a prova individualizada do '
    'sofrimento quando a ilicitude da cobrança é manifesta e afeta '
    'diretamente a subsistência do consumidor.'
)
body(
    'No caso concreto, o dano se manifesta com especial intensidade: '
    'a) a autora é beneficiária do INSS com renda de R$ 1.518,00, '
    'que representa sua única fonte de sustento; '
    'b) os descontos compulsórios do RMC (R$ 74,90/mês) somados ao '
    'seguro "Agi Protege" (R$ 17,99/mês) e ao RCC PAN simultaneamente '
    'ativo comprometem percentual significativo desse benefício; '
    'c) a autora reside em área rural (São Cristóvão/SE) em condição '
    'de hipossuficiência comprovada; '
    'd) a contratação foi realizada por correspondente bancário usando '
    'o "App do Consultor" sem que a autora fosse adequadamente informada '
    'sobre a natureza do produto.'
)
body(
    'O cerceamento da capacidade econômica de consumidora hipossuficiente '
    'que não anuiu ao produto imposto é suficiente para caracterizar dano '
    'moral in re ipsa — basta a ilicitude objetiva da cobrança. '
    'Requer-se a condenação do BANCO AGIBANK S.A. ao pagamento de '
    'R$ 15.000,00 (quinze mil reais) a título de danos morais, '
    'valor pedido na inicial e compatível com a extensão do dano e '
    'com a capacidade econômica da instituição financeira ré.'
)

# 5.15 — Danos temporais
h2('DOS DANOS TEMPORAIS — DESVIO PRODUTIVO — R$ 5.000,00')

body(
    'A teoria do desvio produtivo do consumidor, desenvolvida pelo '
    'professor Marcos Dessaune e reconhecida pelo Superior Tribunal de '
    'Justiça (REsp 1.737.412 e julgados subsequentes), impõe a '
    'indenização pelo tempo produtivo perdido pelo consumidor na '
    'tentativa de resolver problema gerado pelo fornecedor.'
)
body(
    'A autora ECI DOS SANTOS SACRAMENTO foi compelida a despender tempo '
    'e energia para identificar os descontos indevidos, buscar orientação '
    'jurídica, outorgar procuração, deslocar-se de área rural até o '
    'escritório do advogado e acompanhar a tramitação da ação — tudo '
    'isso em razão exclusiva da conduta ilícita do banco réu. Esse '
    'desvio de tempo e atenção de consumidora com 59 anos, residente '
    'em zona rural, tem impacto concreto e mensurável em sua rotina.'
)
body(
    'Requer-se a condenação do BANCO AGIBANK S.A. ao pagamento de '
    'R$ 5.000,00 (cinco mil reais) a título de danos temporais, '
    'conforme pedido formulado na petição inicial.'
)

# 5.16 — Impugnação presunção autenticidade
h2('DA IMPUGNAÇÃO À TESE DE PRESUNÇÃO DE AUTENTICIDADE — ART. 411, III, DO CPC')

body(
    'O banco réu invoca o art. 411, III, do CPC para sustentar que a '
    'assinatura biométrica não impugnada deve ser presumida autêntica. '
    'A invocação é equivocada por três razões:'
)
body(
    'a) A parte autora IMPUGNA expressamente a assinatura biométrica — '
    'na petição inicial (item 55.c, com pedido de IP, geolocalização '
    'e selfie/vídeo) e na presente réplica. O pressuposto do '
    'art. 411, III (ausência de impugnação) simplesmente não se '
    'verifica no caso concreto.'
)
body(
    'b) O Tema 1061 STJ distribui o ônus de forma diversa: cabe ao '
    'banco provar a autenticidade da assinatura, não à autora provar '
    'sua ausência. O art. 411, III, do CPC não revoga o Tema 1061 '
    '— a distribuição especial de ônus firmada pelo STJ em matéria '
    'consumerista bancária prevalece sobre a regra geral processual.'
)
body(
    'c) Biometria "Ativa BPO" realizada pelo canal "App do Consultor" '
    'não é biometria do titular: é o correspondente bancário que '
    'opera o sistema e captura a imagem usando os dados e o rosto do '
    'cliente, sem que este compreenda estar assinando um contrato de '
    'cartão de crédito consignado. A biometria coletada por '
    'correspondente não vincula o consumidor da mesma forma que a '
    'biometria realizada pelo próprio titular em seu dispositivo pessoal.'
)

# 5.17 — Inversão ônus
h2('DA INVERSÃO DO ÔNUS DA PROVA — OBRIGATÓRIA NO CASO CONCRETO')

body(
    'O banco réu sustenta que a inversão do ônus da prova "deve ser '
    'indeferida quando as alegações autorais carecem de verossimilhança". '
    'O argumento ignora os fatos do caso concreto.'
)
body(
    'Nos termos da Súmula 297 do STJ, o CDC aplica-se integralmente '
    'às instituições financeiras. Nos termos do art. 6º, VIII, do CDC, '
    'a inversão do ônus da prova é devida quando o consumidor for '
    'hipossuficiente OU quando sua alegação for verossímil — condições '
    'alternativas, não cumulativas.'
)
body(
    'A autora preenche ambas as condições: a) hipossuficiência — renda '
    'de R$ 1.518,00, residente em zona rural, beneficiária do INSS, '
    '59 anos, sem expertise em serviços financeiros; b) verossimilhança '
    '— laudo com "undefined@undefined", sem hash, sem geolocalização, '
    'canal "App do Consultor", divergência entre conta no laudo e '
    'conta do crédito efetivo, seguro não solicitado, refinanciamento '
    'maquiador. Cada um desses elementos, individualmente, já tornaria '
    'verossímil a alegação da autora.'
)
body(
    'Adicionalmente, o art. 373, § 1º, do CPC consagra a teoria da '
    'distribuição dinâmica do ônus da prova: quando uma das partes '
    'tiver maior aptidão para produzir a prova, o ônus recai sobre ela. '
    'O banco réu detém todos os arquivos digitais, logs de acesso, '
    'registros de sessão, dados biométricos e demais evidências da '
    'contratação — enquanto a autora não tem acesso a nenhum desses '
    'elementos. O próprio Tema 1061 STJ já determina que o ônus da '
    'prova da assinatura cabe ao banco — a inversão é norma aplicável '
    'ao caso, não exceção.'
)

# 5.18 — Impossibilidade de conversão
h2('DA IMPOSSIBILIDADE DE CONVERSÃO — IMPUGNAÇÃO À TESE DO BANCO')

body(
    'O banco réu alega, à fl. 146 da contestação, que seria materialmente '
    'impossível a conversão do RMC em empréstimo consignado tradicional '
    '"haja vista que a parte autora não possui margem consignável '
    'disponível para suportar novos descontos".'
)
body(
    'O argumento é impertinente por duas razões: a) o pedido principal '
    'da autora é a declaração de inexistência do negócio jurídico — não '
    'a conversão. A conversão é pedido subsidiário da inicial e não '
    'altera o mérito principal desta réplica; b) se o banco pretende '
    'alegar ausência de margem, cabe a ele demonstrar com extrato '
    'HISCON atualizado qual a real situação da margem da autora. '
    'A mera afirmação na contestação, sem documento probatório, '
    'não tem valor algum.'
)

# ===========================================================================
# VI — DOS HONORÁRIOS ADVOCATÍCIOS
# ===========================================================================
h1('DOS HONORÁRIOS ADVOCATÍCIOS')

body(
    'Nos termos do art. 55 da Lei nº 9.099/95, aplicável subsidiariamente '
    'ao rito do JEF federal, são devidos honorários advocatícios em '
    'caso de condenação, na hipótese de recurso interposto pela parte '
    'vencida. Requer-se a condenação do BANCO AGIBANK S.A. ao pagamento '
    'de honorários advocatícios sucumbenciais, a ser arbitrado por este '
    'Juízo nos termos do art. 85 do CPC.'
)

# ===========================================================================
# VII — DO JULGAMENTO ANTECIPADO
# ===========================================================================
h1('DO JULGAMENTO ANTECIPADO DO MÉRITO')

body(
    'Caso Vossa Excelência entenda que a matéria é exclusivamente de '
    'direito ou que já há prova suficiente nos autos para a formação '
    'do convencimento judicial, pugna-se pelo julgamento antecipado do '
    'mérito nos termos do art. 355, I, do CPC, com a procedência '
    'integral dos pedidos formulados na inicial e reiterados nesta réplica.'
)
body(
    'Subsidiariamente, caso Vossa Excelência entenda necessária a '
    'produção de prova, requer-se a designação de:'
)
body(
    'a) perícia grafotécnica para confronto das assinaturas (contrato '
    'Agibank × RG × procuração), nos termos do Tema 1061 STJ;'
)
body(
    'b) subsidiariamente à grafotécnica, perícia digital nos arquivos '
    'eletrônicos do banco réu (logs de acesso, user-agent, geolocalização, '
    'hash, gravação de videochamada);'
)
body(
    'c) em caso de designação de audiência instrutória, oitiva da '
    'autora e depoimento pessoal do preposto do banco réu.'
)

# ===========================================================================
# VIII — DOS PEDIDOS FINAIS
# ===========================================================================
h1('DOS PEDIDOS FINAIS')

body(
    'Diante de todo o exposto, requer a autora ECI DOS SANTOS SACRAMENTO '
    'a Vossa Excelência:'
)
body(
    'a) A REJEIÇÃO de todas as preliminares suscitadas pelo BANCO AGIBANK '
    'S.A. (ausência de comprovante de residência, procuração desatualizada '
    'e incompetência do JEF), pelos fundamentos expostos nesta réplica;'
)
body(
    'b) O RECONHECIMENTO da nulidade formal da contratação por ausência '
    'do Termo de Consentimento Esclarecido (TCE), nos termos do '
    'art. 21-A da IN INSS/PRES nº 28/2008;'
)
body(
    'c) A DESIGNAÇÃO de perícia grafotécnica para confronto das '
    'assinaturas (contrato Agibank × RG × procuração), nos termos do '
    'Tema 1061 STJ, e, subsidiariamente, a designação de perícia digital '
    'nos arquivos eletrônicos do banco réu;'
)
body(
    'd) A DECLARAÇÃO de inexistência do negócio jurídico consubstanciado '
    'no contrato de cartão RMC nº 1517444869, em caso de confirmação da '
    'inautenticidade da assinatura pela perícia grafotécnica, com a '
    'consequente cessação dos descontos no benefício previdenciário da '
    'autora e liberação integral da margem consignável;'
)
body(
    'e) Subsidiariamente, caso não acolhida a tese de inexistência, '
    'a DECLARAÇÃO de nulidade do contrato de cartão RMC nº 1517444869 '
    'por vício de consentimento quanto à modalidade contratada '
    '(arts. 138 e 139, I, do Código Civil), com ANUÊNCIA expressa '
    'à compensação do TED de R$ 1.096,94 recebido em 27/08/2024;'
)
body(
    'f) A DECLARAÇÃO de nulidade da cobertura de seguro "Agi Protege" '
    'embutida na contratação sem solicitação da autora (venda casada — '
    'CDC art. 39, I e III, e Súmula 532 STJ), com cessação imediata '
    'das cobranças a esse título;'
)
body(
    'g) A DECLARAÇÃO de nulidade do refinanciamento maquiador de '
    'R$ 1.062,40 lançado em 23/01/2025 como "Crédito Pessoal Renovado";'
)
body(
    'h) A CONDENAÇÃO do BANCO AGIBANK S.A. à RESTITUIÇÃO EM DOBRO de: '
    'i) todas as parcelas descontadas do benefício previdenciário a '
    'título do RMC nº 1517444869, a partir de 27/08/2024; '
    'ii) todos os valores cobrados a título do seguro "Agi Protege" '
    '(R$ 176,39 na 1ª fatura + R$ 17,99/mês nas faturas subsequentes); '
    'iii) DEDUZIDO o valor de R$ 1.096,94 recebido via TED em 27/08/2024 '
    '— valores a apurar em cumprimento de sentença, nos termos do '
    'Tema 929 STJ;'
)
body(
    'i) A CONDENAÇÃO do BANCO AGIBANK S.A. ao pagamento de '
    'R$ 15.000,00 (quinze mil reais) a título de danos morais in re ipsa;'
)
body(
    'j) A CONDENAÇÃO do BANCO AGIBANK S.A. ao pagamento de '
    'R$ 5.000,00 (cinco mil reais) a título de danos temporais '
    'pelo desvio produtivo da autora;'
)
body(
    'k) A INVERSÃO do ônus da prova, determinando ao banco réu a '
    'juntada de: IP com identificação do titular, geolocalização, '
    'selfie liveness certificada iBeta 2/IEEE Std 2790/2020, hash '
    'SHA-256 do contrato, gravação da videochamada exigida pela '
    'Resolução CNJ nº 159/2024, e dados completos de sessão;'
)
body(
    'l) A CONDENAÇÃO do BANCO AGIBANK S.A. ao pagamento de honorários '
    'advocatícios sucumbenciais, nos termos do art. 55 da Lei nº 9.099/95 '
    'e art. 85 do CPC;'
)
body(
    'm) O PREQUESTIONAMENTO expresso dos seguintes dispositivos: '
    'arts. 138, 139, I, e 421 do Código Civil; '
    'arts. 6º, III e VIII, 39, I, III, IV e V, 42, parágrafo único, '
    '51, IV, e 52, V, do CDC; '
    'arts. 104, 106, 373, § 1º, 411, III, 430 e 464 do CPC; '
    'art. 5º, II e III, da IN INSS/PRES nº 138/2022; '
    'Resolução CNJ nº 159/2024; '
    'Tema 1061 STJ; Tema 183 TNU; Tema 929 STJ; '
    'EAREsp 676.608/RS; REsp 2.159.442/PR; '
    'Súmulas 297 e 532 do STJ;'
)
body(
    'n) Em caso de designação de audiência instrutória, o depoimento '
    'pessoal do preposto do BANCO AGIBANK S.A. e a oitiva da '
    'autora como informante.'
)
body(
    'Termos em que pede e espera deferimento.'
)

# FECHO
doc.add_paragraph('')
p_fecho = doc.add_paragraph()
r_fecho = p_fecho.add_run('Aracaju/SE, 23 de abril de 2026.')
p_fecho.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')

p_linha = doc.add_paragraph()
r_linha = p_linha.add_run('_' * 50)
p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_nome = doc.add_paragraph()
r_nome = p_nome.add_run('Tiago de Azevedo Lima')
r_nome.bold = True
p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_oab_se = doc.add_paragraph()
r_oab_se = p_oab_se.add_run('OAB/SE 1850A')
r_oab_se.bold = True
p_oab_se.alignment = WD_ALIGN_PARAGRAPH.CENTER

for oab_txt in ['OAB/SC 36672', 'OAB/AL 20906A', 'OAB/BA 80006',
                'OAB/MG 228433', 'OAB/RS 139330A']:
    p_oab = doc.add_paragraph()
    p_oab.add_run(oab_txt)
    p_oab.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(SAIDA_TMP)
print(f'SALVO: {SAIDA_TMP}')
print(f'Paragrafos: {len(doc.paragraphs)}')
print(f'Tabelas: {len(doc.tables)}')
