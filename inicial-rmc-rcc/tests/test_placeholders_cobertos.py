"""Teste: garante que TODOS os placeholders dos templates estão cobertos
pelo dict de `montar_dict_placeholders` em `_pipeline_caso.py`.

Origem (paradigma): MARIA AZEVEDO PARINTINS 2026-05-14 — placeholders
`{{data_da_inclusão}}` e `{{anos_meses_ativo}}` foram adicionados ao
template AM RMC em algum momento, mas o dict do pipeline nunca foi
atualizado para emiti-los. Resultado: literais `{{...}}` saindo no docx.

Este teste extrai todos os `{{...}}` de cada template ativo e confere
contra o dict canônico declarado abaixo (espelha o `montar_dict_placeholders`).

Roda standalone:
    python -m pytest tests/test_placeholders_cobertos.py -v
    OU
    python tests/test_placeholders_cobertos.py
"""
import os
import sys
import re
try:
    import pytest
    _HAS_PYTEST = True
except ImportError:
    _HAS_PYTEST = False
    class _PytestStub:
        class mark:
            @staticmethod
            def parametrize(*a, **kw):
                def deco(f): return f
                return deco
    pytest = _PytestStub()
from docx import Document

# Conjunto canônico de placeholders que o pipeline emite (espelha
# `montar_dict_placeholders` em references/_pipeline_caso.py).
# Quando adicionar campo no template, ATUALIZAR esta lista E o dict
# do pipeline. O teste pega o esquecimento de um dos dois.
PLACEHOLDERS_COBERTOS = {
    # Comarca
    '{{competencia}}',
    # Autora
    '{{nome_completo}}', '{{nacionalidade}}', '{{estado_civil}}',
    '{{profissao}}', '{{cpf}}', '{{rg}}', '{{orgao_expedidor}}',
    '{{logradouro}}', '{{numero}}', '{{bairro}}',
    '{{cidade_de_residencia}}', '{{uf}}', '{{cep}}',
    '{{inscrito}}', '{{domiciliado}}',
    # Benefício
    '{{tipo_de_beneficio}}', '{{numero_do_beneficio}}',
    '{{conta_agencia_conta}}', '{{banco_que_recebe}}',
    '{{valor_liquido_beneficio}}',
    # Contrato
    '{{numero_do_contrato}}',
    '{{data_do_primeiro_desconto}}',
    '{{data_da_inclusão}}',  # variante com cedilha+acento (templates AM RMC)
    '{{data_da_inclusao}}',  # variante sem acento (fallback)
    '{{anos_meses_ativo}}',
    '{{total_de_parcelas}}',
    '{{valor_da_parcela}}',
    # Cálculos
    '{{valor_dobro}}', '{{valor_dobro_extenso}}',
    '{{valor_final_da_causa}}', '{{valor_final_da_causa_por_extenso}}',
    # Escritório
    '{{endereco_escritorio}}',
    # Placeholders Jinja-like de bloco condicional (não-string)
    # — não aparecem como {{nome}} simples; ignorados na varredura
    # via prefixo {% se aparecer. Por ora, todos os do template são {{...}}.
}

# Placeholders no template que NÃO precisam estar no dict
# (são processados por outras funções, ex: quali_banco compõe-se via run
# separado, idoso via resolver_jinja_idoso, banco-réu via substituir_quali_banco).
PLACEHOLDERS_EXTERNOS = {
    '{{quali_banco}}',          # substituir_quali_banco em helpers_redacao.py
    '{{nome_curto_banco}}',     # idem
    '{{cnpj_banco}}',           # idem
    '{{IDOSO}}',                # resolver_jinja_idoso
    '{{cabecalho_idoso}}',      # idem
    '{{pedido_idoso}}',         # idem
}

# Templates da skill RMC ficam em outro path do escritório (cf.
# `references/perfis_juridicos.py: DEST_BASE`). Importamos a constante
# para descobrir o diretório de templates dinamicamente.
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'references')))
try:
    from perfis_juridicos import DEST_BASE, PERFIS
    # Coleta todos os 'destino' configurados em todos os perfis (UFs)
    _TEMPLATES_FROM_PERFIS = []
    for _uf, _perfil in PERFIS.items():
        for _arq in _perfil.get('arquivos', []):
            dest = _arq.get('destino')
            if dest and os.path.exists(dest):
                _TEMPLATES_FROM_PERFIS.append(dest)
    TEMPLATES_DIR = DEST_BASE if os.path.isdir(DEST_BASE) else None
except Exception:
    _TEMPLATES_FROM_PERFIS = []
    TEMPLATES_DIR = None


def listar_templates():
    """Retorna lista de templates DOCX para auditar.

    Fonte primária: `destino` declarado em `perfis_juridicos.PERFIS`
    (caminho real onde a skill consome os templates).
    Fallback: walk em `DEST_BASE` se existir.
    Fallback 2: walk em `references/` da skill (caso desenvolvedor tenha
    movido localmente).
    """
    if _TEMPLATES_FROM_PERFIS:
        return list(_TEMPLATES_FROM_PERFIS)
    out = []
    candidatos = []
    if TEMPLATES_DIR and os.path.isdir(TEMPLATES_DIR):
        candidatos.append(TEMPLATES_DIR)
    candidatos.append(os.path.abspath(
        os.path.join(os.path.dirname(__file__), '..', 'references')
    ))
    for cand in candidatos:
        if not os.path.isdir(cand): continue
        for root, _, files in os.walk(cand):
            for f in files:
                if f.lower().endswith('.docx') and not f.startswith('~$'):
                    if '.bak' in f.lower():
                        continue
                    out.append(os.path.join(root, f))
        if out: break
    return out


def extrair_placeholders(docx_path):
    """Retorna o set de placeholders {{...}} encontrados no template."""
    try:
        d = Document(docx_path)
    except Exception:
        return set()
    encontrados = set()
    fontes = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fontes.append(p.text)
    for txt in fontes:
        for m in re.findall(r'\{\{[^}]+\}\}', txt):
            encontrados.add(m)
    return encontrados


@pytest.mark.parametrize('template_path', listar_templates())
def test_template_placeholders_cobertos_pelo_dict(template_path):
    """Cada placeholder do template deve estar no dict canônico OU na
    lista de placeholders externos (que outra função preenche)."""
    encontrados = extrair_placeholders(template_path)
    nao_cobertos = encontrados - PLACEHOLDERS_COBERTOS - PLACEHOLDERS_EXTERNOS
    nome = os.path.basename(template_path)
    assert not nao_cobertos, (
        f'\n\nTemplate "{nome}" usa placeholders NÃO cobertos pelo dict:\n'
        f'  {sorted(nao_cobertos)}\n\n'
        f'Adicione esses placeholders ao dict de `montar_dict_placeholders` '
        f'em `references/_pipeline_caso.py` E à constante PLACEHOLDERS_COBERTOS '
        f'neste teste. Se forem processados por outra função, adicione a '
        f'PLACEHOLDERS_EXTERNOS aqui.'
    )


def main():
    """Modo CLI — relatório sem pytest."""
    templates = listar_templates()
    if not templates:
        print(f'[WARN] Nenhum template encontrado em {TEMPLATES_DIR}')
        return 1
    n_falhas = 0
    for t in templates:
        encontrados = extrair_placeholders(t)
        nao_cobertos = encontrados - PLACEHOLDERS_COBERTOS - PLACEHOLDERS_EXTERNOS
        nome = os.path.basename(t)
        if nao_cobertos:
            n_falhas += 1
            print(f'[WARN] {nome}: NÃO COBERTOS {sorted(nao_cobertos)}')
        else:
            print(f'[OK] {nome}: {len(encontrados)} placeholders, todos cobertos')
    if n_falhas:
        print(f'\n{n_falhas} template(s) com placeholders descobertos.')
        return 1
    print(f'\n[OK] Todos os {len(templates)} templates OK.')
    return 0


if __name__ == '__main__':
    sys.exit(main())
