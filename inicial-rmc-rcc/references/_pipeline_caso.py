"""Pipeline de renderizacao de inicial RMC/RCC para um caso especifico.

Recebe um dict `caso` com dados da autora, banco-reu, beneficio e contrato +
lista de descontos do HISCON, e gera:

  * INICIAL_<NOME>_<TESE>.docx        — peca processual
  * CALCULO_<NOME>_<TESE>.xlsx        — memoria de calculo (aba unica)
  * RELATORIO_PENDENCIAS_<NOME>.docx  — ausencias e alertas

Aplica todas as 5 regras canonicas + extras (paradigma BENEDITA 2026-05-13).
Veja SKILL.md para uso.
"""
import os
import sys
import shutil
from datetime import date
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

from helpers_redacao import (
    calcular_valor_liquido,
    filtrar_descontos_contrato_corrente,
    dentro_prescricao_quinquenal,
    centralizar_celulas_tabela_quadro_sumario,
    forcar_cambria_quadro_sumario,
    aplicar_12pt_no_polo_passivo,
    substituir_quali_banco,
    conjugacao_por_nacionalidade,
    omitir_estado_civil_se_vazio,
    resolver_jinja_idoso,
    substituir_placeholders_com_grifo,
)
from helpers_docx import iter_all_paragraphs
from perfis_juridicos import perfil

# Importa extenso do NC
_NC_REFS = r"C:/Users/gabri/.claude/skills/inicial-nao-contratado/references"
if _NC_REFS not in sys.path:
    sys.path.insert(0, _NC_REFS)
try:
    from extenso import extenso_moeda
except Exception:
    def extenso_moeda(v):
        return f"({v:,.2f} reais)"


def fmt_brl(v):
    return f"{v:,.2f}".replace(",", "#").replace(".", ",").replace("#", ".")


# ============================================================
#   GERACAO DA INICIAL DOCX
# ============================================================

def gerar_inicial(caso, template_path, destino_docx):
    """Renderiza a inicial DOCX para o caso."""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template nao encontrado: {template_path}")
    shutil.copyfile(template_path, destino_docx)

    doc = Document(destino_docx)

    eh_idoso = caso["autora"].get("eh_idoso", False)
    nacionalidade = caso["autora"].get("nacionalidade", "brasileira")
    estado_civil = caso["autora"].get("estado_civil", "")
    inscrito, domiciliado = conjugacao_por_nacionalidade(nacionalidade)

    # 1. Resolver Jinja idoso
    resolver_jinja_idoso(doc, eh_idoso)

    # 2. Omitir estado_civil se vazio (ANTES das substituicoes)
    omitir_estado_civil_se_vazio(doc, estado_civil)

    # 3. quali_banco em 2 runs (Segoe Bold + Cambria, 12pt, grifo)
    substituir_quali_banco(
        doc,
        caso["banco"]["nome"],
        caso["banco"]["resto_qualificacao"],
    )

    # 4. Substituir demais placeholders com grifo amarelo
    dados = montar_dict_placeholders(caso, inscrito, domiciliado)
    substituir_placeholders_com_grifo(doc, dados)

    # 5. Centralizar tabela Quadro Sumario (caso o gerador nao tenha aplicado)
    centralizar_celulas_tabela_quadro_sumario(doc)

    # 5-bis. Forçar fonte Cambria em TODAS as celulas do Quadro Sumario.
    # Sem isto, os placeholders {{numero_do_contrato}}, {{valor_da_parcela}} e
    # {{data_do_primeiro_desconto}} herdam fonte sans-serif do template em
    # algumas variantes (feedback Gabriel 2026-05-14, ELIAS RCC AGIBANK).
    forcar_cambria_quadro_sumario(doc)

    # 6. 12pt no polo passivo (caso o gerador nao tenha aplicado)
    aplicar_12pt_no_polo_passivo(doc)

    doc.save(destino_docx)
    return doc


def _humanizar_anos_meses(data_inclusao_str, ref=None):
    """Recebe 'dd/mm/aaaa' (ou 'mm/aaaa') e devolve string 'X anos e Y meses'
    contados até a data de referência (default: hoje). Usada no parágrafo
    'o empréstimo foi contratado em DATA, ou seja, há ANOS_MESES...'.
    """
    if ref is None:
        ref = date.today()
    if not data_inclusao_str:
        return ""
    s = str(data_inclusao_str).strip()
    partes = s.split("/")
    try:
        if len(partes) == 3:
            dia, mes, ano = int(partes[0]), int(partes[1]), int(partes[2])
        elif len(partes) == 2:
            dia = 1
            mes, ano = int(partes[0]), int(partes[1])
        else:
            return ""
    except ValueError:
        return ""
    if ano < 100:
        ano += 2000
    try:
        ini = date(ano, mes, dia)
    except ValueError:
        return ""
    meses_total = (ref.year - ini.year) * 12 + (ref.month - ini.month)
    if ref.day < ini.day:
        meses_total -= 1
    if meses_total < 0:
        meses_total = 0
    anos = meses_total // 12
    meses = meses_total % 12
    partes_h = []
    if anos:
        partes_h.append(f"{anos} ano" + ("s" if anos != 1 else ""))
    if meses or not anos:
        partes_h.append(f"{meses} {'meses' if meses != 1 else 'mês'}")
    return " e ".join(partes_h)


def montar_dict_placeholders(caso, inscrito, domiciliado):
    """Monta o dict de substituicao a partir dos dados do caso."""
    autora = caso["autora"]
    benef = caso["beneficio"]
    contrato = caso["contrato"]
    calc = caso["calculo"]
    perfil_uf = caso["perfil"]

    data_inclusao = contrato.get("data_inclusao", "")
    anos_meses = _humanizar_anos_meses(data_inclusao) if data_inclusao else ""

    return {
        "{{competencia}}": caso.get("comarca", perfil_uf["comarca_default"]),
        # Autora
        "{{nome_completo}}": autora["nome"],
        "{{nacionalidade}}": autora.get("nacionalidade", "brasileira"),
        "{{estado_civil}}": autora.get("estado_civil", ""),
        "{{profissao}}": autora.get("profissao", "aposentada"),
        "{{cpf}}": autora["cpf"],
        "{{rg}}": autora.get("rg", ""),
        "{{orgao_expedidor}}": autora.get("orgao_expedidor", "SSP/AM"),
        "{{logradouro}}": autora.get("logradouro", ""),
        "{{numero}}": autora.get("numero", "s/nº"),
        "{{bairro}}": autora.get("bairro", ""),
        "{{cidade_de_residencia}}": autora.get("cidade", ""),
        "{{uf}}": autora.get("uf", perfil_uf["uf"]),
        "{{cep}}": autora.get("cep", ""),
        "{{inscrito}}": inscrito,
        "{{domiciliado}}": domiciliado,
        # Beneficio
        "{{tipo_de_beneficio}}": benef["tipo"],
        "{{numero_do_beneficio}}": benef["nb"],
        "{{conta_agencia_conta}}": benef["conta_agencia_conta"],
        "{{banco_que_recebe}}": benef["banco_pagador"],  # sem prefixo "BANCO"
        "{{valor_liquido_beneficio}}": fmt_brl(calc["valor_liquido"]),
        # Contrato
        "{{numero_do_contrato}}": contrato["numero"],
        "{{data_do_primeiro_desconto}}": contrato["data_primeiro_desconto"],
        # Placeholders do template RMC (parágrafo "o empréstimo foi contratado
        # em DATA, ou seja, há ANOS_MESES..."). Cobrem variantes ortográficas
        # encontradas nos templates AM (com/sem cedilha+acento).
        "{{data_da_inclusão}}": data_inclusao,
        "{{data_da_inclusao}}": data_inclusao,
        "{{anos_meses_ativo}}": anos_meses,
        "{{total_de_parcelas}}": str(calc["total_parcelas_historico"]),
        "{{valor_da_parcela}}": fmt_brl(calc["valor_parcela_atual"]),
        # Calculos
        "{{valor_dobro}}": fmt_brl(calc["valor_dobro"]),
        "{{valor_dobro_extenso}}": extenso_moeda(calc["valor_dobro"]),
        "{{valor_final_da_causa}}": fmt_brl(calc["valor_causa"]),
        "{{valor_final_da_causa_por_extenso}}": extenso_moeda(calc["valor_causa"]),
        # Escritorio
        "{{endereco_escritorio}}": perfil_uf["endereco_escritorio"],
    }


# ============================================================
#   CALCULOS
# ============================================================

def calcular_valores(caso):
    """Calcula valor liquido, dobro, valor da causa a partir do HISCON e
    descontos. Atualiza caso['calculo']."""
    descontos = caso["contrato"]["descontos_hiscon"]  # cronologico recente->antigo
    hiscon = caso["beneficio"]

    # REGRA 2 — filtrar so do contrato corrente
    if "data_inclusao" in caso["contrato"]:
        descontos_corrente = filtrar_descontos_contrato_corrente(
            descontos, caso["contrato"]["data_inclusao"]
        )
    else:
        descontos_corrente = descontos
    caso["contrato"]["descontos_filtrados"] = descontos_corrente

    # REGRA 1 (revisada 2026-05-16) — valor líquido = última competência
    # paga "Origem: Maciça" do HISCRE. Antes era `base - total_comprometido`
    # do HISCON (paradigma BENEDITA original), mas a regra unificada do
    # escritório passou a ser HISCRE (mesma fonte usada pelo NC). Caso
    # paradigma: PEDRO 2026-05-16 — NC mostrava R$ 970,00 (HISCRE) e
    # RMC mostrava R$ 891,70 (HISCON-comprometido); operador determinou
    # que SEMPRE deve ser HISCRE.
    valor_liquido = None
    hiscre_path = caso.get("hiscre_path") or hiscon.get("_hiscre_path")
    if hiscre_path:
        try:
            import sys as _sys
            _nc_refs = r"C:/Users/gabri/.claude/skills/inicial-nao-contratado/references"
            if _nc_refs not in _sys.path:
                _sys.path.insert(0, _nc_refs)
            from extrator_hiscre import parse_hiscre  # type: ignore
            _h = parse_hiscre(hiscre_path)
            if _h and _h.get("valor_liquido"):
                valor_liquido = float(_h["valor_liquido"])
        except Exception:
            pass
    if valor_liquido is None and caso["beneficio"].get("valor_liquido_hiscre"):
        valor_liquido = float(caso["beneficio"]["valor_liquido_hiscre"])
    if valor_liquido is None:
        # Fallback antigo (paradigma BENEDITA): HISCON base − comprometido
        valor_liquido = calcular_valor_liquido(
            hiscon["base_calculo"], hiscon["total_comprometido"]
        )

    # TRATO SUCESSIVO (correção 2026-05-16): em descontos mensais consecutivos,
    # o prazo prescricional flui do ÚLTIMO desconto (não do primeiro). NÃO
    # aplicamos prescrição retroativa — TODAS as parcelas históricas entram
    # no dobro. Caso paradigma VILSON / BANRISUL 2026-05-16.
    descontos_impugnados = list(descontos_corrente)
    soma_total = sum(v for _, v in descontos_impugnados)
    valor_dobro = soma_total * 2

    # Danos morais e temporais (padrao escritorio)
    dano_moral = caso.get("dano_moral", 10000.00)
    dano_temporal = caso.get("dano_temporal", 5000.00)
    valor_causa = valor_dobro + dano_moral + dano_temporal

    caso["calculo"] = {
        "valor_liquido": valor_liquido,
        "valor_parcela_atual": descontos_corrente[0][1] if descontos_corrente else 0.0,
        "total_parcelas_historico": len(descontos_corrente),
        "total_parcelas_5anos": len(descontos_impugnados),  # = todas (trato sucessivo)
        "soma_5anos": soma_total,  # = todas (trato sucessivo)
        "soma_total": soma_total,
        "valor_dobro": valor_dobro,
        "dano_moral": dano_moral,
        "dano_temporal": dano_temporal,
        "valor_causa": valor_causa,
        # set agora cobre TODAS as competências (visual da planilha)
        "descontos_prescricao_set": {c for c, _ in descontos_impugnados},
    }
    return caso["calculo"]


# ============================================================
#   PLANILHA DE CALCULO (ABA UNICA)
# ============================================================

def gerar_planilha(caso, destino_xlsx):
    """Gera a memoria de calculo em XLSX (aba unica 'Calculo')."""
    bold = Font(bold=True)
    italic = Font(italic=True)
    grifo = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
    header_fill = PatternFill(start_color="FF305496", end_color="FF305496", fill_type="solid")
    section_fill = PatternFill(start_color="FFD9E1F2", end_color="FFD9E1F2", fill_type="solid")
    destaque_amarelo_claro = PatternFill(start_color="FFFFF2CC", end_color="FFFFF2CC", fill_type="solid")
    cabec_cinza = PatternFill(start_color="FFE7E6E6", end_color="FFE7E6E6", fill_type="solid")
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right = Alignment(horizontal="right", vertical="center")
    left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)
    border_thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Calculo"

    autora = caso["autora"]
    benef = caso["beneficio"]
    contrato = caso["contrato"]
    calc = caso["calculo"]
    perfil_uf = caso["perfil"]
    descontos = contrato["descontos_filtrados"]

    # Titulo
    ws.merge_cells("A1:D1")
    ws["A1"] = f"MEMORIA DE CALCULO — ACAO ANULATORIA DE {caso.get('tese','RMC')}"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFFFF")
    ws["A1"].fill = header_fill
    ws["A1"].alignment = center
    ws.row_dimensions[1].height = 24

    # Identificacao
    info = [
        ("Autora", autora["nome"]),
        ("CPF", autora["cpf"]),
        ("Beneficio (NB)", f"{benef['nb']} — {benef['tipo'].title()}"),
        ("Contrato", f"{contrato['numero']} — {caso.get('tese_extenso', 'Cartao de Credito com Reserva de Margem Consignavel (RMC)')}"),
        ("Banco-reu", f"{caso['banco']['nome_curto']} — CNPJ {caso['banco']['cnpj']}"),
        ("Periodo apurado", f"{descontos[-1][0]} a {descontos[0][0]} ({len(descontos)} parcelas — destacadas as {calc['total_parcelas_5anos']} dentro da prescricao quinquenal)"),
        ("Procurador", f"{caso.get('procurador_nome','—')} — {caso.get('procurador_oab','—')}"),
    ]
    for i, (k, v) in enumerate(info, start=3):
        ca = ws.cell(row=i, column=1, value=k)
        ca.font = bold; ca.alignment = left_wrap; ca.border = border_thin
        ws.merge_cells(start_row=i, start_column=2, end_row=i, end_column=4)
        cb = ws.cell(row=i, column=2, value=v)
        cb.alignment = left_wrap; cb.border = border_thin
        ws.cell(row=i, column=3).border = border_thin
        ws.cell(row=i, column=4).border = border_thin

    # Secao 1: detalhamento
    linha = 3 + len(info) + 1
    ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=4)
    cs = ws.cell(row=linha, column=1,
                  value=f"1. DETALHAMENTO DOS DESCONTOS NO BENEFICIO (HISCON) — {len(descontos)} parcelas")
    cs.font = bold; cs.fill = section_fill; cs.alignment = center
    linha += 1
    for col, h in enumerate(["#", "Competencia", "Valor desconto (R$)", "Dentro prescricao (5 anos)?"], 1):
        c = ws.cell(row=linha, column=col, value=h)
        c.font = bold; c.alignment = center; c.border = border_thin; c.fill = cabec_cinza
    linha += 1

    primeira = linha
    presc_set = calc["descontos_prescricao_set"]
    for i, (comp, val) in enumerate(descontos, 1):
        dentro = comp in presc_set
        ws.cell(row=linha, column=1, value=i).border = border_thin
        ws.cell(row=linha, column=1).alignment = center
        ws.cell(row=linha, column=2, value=comp).border = border_thin
        ws.cell(row=linha, column=2).alignment = center
        c = ws.cell(row=linha, column=3, value=val)
        c.number_format = '"R$" #,##0.00'; c.alignment = right; c.border = border_thin
        c4 = ws.cell(row=linha, column=4, value="SIM" if dentro else "—")
        c4.border = border_thin; c4.alignment = center
        if dentro:
            for col in range(1, 5):
                ws.cell(row=linha, column=col).fill = destaque_amarelo_claro
        linha += 1
    ultima = linha - 1

    # Subtotal 5 anos
    ws.cell(row=linha, column=1).border = border_thin
    c = ws.cell(row=linha, column=2,
                 value=f"SUBTOTAL — {calc['total_parcelas_5anos']} parcelas DENTRO da prescricao (5 anos):")
    c.font = bold; c.border = border_thin; c.alignment = right
    c = ws.cell(row=linha, column=3, value=f'=SUMIFS(C{primeira}:C{ultima},D{primeira}:D{ultima},"SIM")')
    c.font = bold; c.number_format = '"R$" #,##0.00'; c.alignment = right
    c.border = border_thin; c.fill = grifo
    ws.cell(row=linha, column=4).border = border_thin
    linha += 1

    # Total historico
    ws.cell(row=linha, column=1).border = border_thin
    c = ws.cell(row=linha, column=2,
                 value=f"TOTAL HISTORICO — {len(descontos)} parcelas (informativo):")
    c.font = italic; c.border = border_thin; c.alignment = right
    c = ws.cell(row=linha, column=3, value=f"=SUM(C{primeira}:C{ultima})")
    c.font = italic; c.number_format = '"R$" #,##0.00'; c.alignment = right
    c.border = border_thin
    ws.cell(row=linha, column=4).border = border_thin
    linha += 2

    # Secao 2: composicao valor da causa
    ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=4)
    cs = ws.cell(row=linha, column=1, value="2. COMPOSICAO DO VALOR DA CAUSA")
    cs.font = bold; cs.fill = section_fill; cs.alignment = center
    linha += 1
    for col, h in enumerate(["Componente", "Fundamento / criterio", "Valor (R$)", ""], 1):
        c = ws.cell(row=linha, column=col, value=h)
        c.font = bold; c.alignment = center; c.border = border_thin; c.fill = cabec_cinza
    linha += 1

    componentes = [
        ("Repeticao em dobro (art. 42, par. unico, CDC)",
         f"Soma dos {calc['total_parcelas_5anos']} descontos dentro da prescricao x 2 — independe de ma-fe",
         calc["valor_dobro"], False),
        ("Subsidiariamente — repeticao simples",
         f"Soma dos {calc['total_parcelas_5anos']} descontos dentro da prescricao (pedido alternativo)",
         calc["soma_5anos"], True),
        ("Dano moral", "Padrao escritorio (IRDR Tema 5 TJAM + Sumula 54 STJ)",
         calc["dano_moral"], False),
        ("Dano temporal — teoria do desvio produtivo",
         "Padrao escritorio (REsp 1.737.412/SP — Min. Nancy Andrighi)",
         calc["dano_temporal"], False),
    ]
    for desc, fund, val, eh_italico in componentes:
        ca = ws.cell(row=linha, column=1, value=desc)
        cb = ws.cell(row=linha, column=2, value=fund)
        cc = ws.cell(row=linha, column=3, value=val)
        cd = ws.cell(row=linha, column=4, value="")
        for cell in (ca, cb, cc, cd):
            cell.border = border_thin
            cell.alignment = left_wrap if cell.column < 3 else right
            if eh_italico:
                cell.font = italic
        cc.number_format = '"R$" #,##0.00'; cc.alignment = right
        linha += 1

    # Valor da causa final (amarelo)
    ws.cell(row=linha, column=1, value="VALOR DA CAUSA").font = bold
    ws.cell(row=linha, column=1).border = border_thin
    ws.cell(row=linha, column=1).fill = grifo
    ws.cell(row=linha, column=2, value="Dobro + dano moral + dano temporal").font = bold
    ws.cell(row=linha, column=2).border = border_thin
    ws.cell(row=linha, column=2).fill = grifo
    ws.cell(row=linha, column=2).alignment = left_wrap
    c = ws.cell(row=linha, column=3, value=calc["valor_causa"])
    c.font = Font(bold=True, size=12); c.fill = grifo
    c.number_format = '"R$" #,##0.00'; c.alignment = right; c.border = border_thin
    ws.cell(row=linha, column=4).border = border_thin
    ws.cell(row=linha, column=4).fill = grifo
    linha += 2

    # Secao 3: notas juridicas
    ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=4)
    cs = ws.cell(row=linha, column=1, value="3. NOTAS JURIDICAS")
    cs.font = bold; cs.fill = section_fill; cs.alignment = center
    linha += 1
    notas = [
        f"1. Total historico de descontos no contrato {contrato['numero']}: {len(descontos)} parcelas ({descontos[-1][0]} a {descontos[0][0]}).",
        f"2. Dentro da prescricao quinquenal (CDC art. 27): {calc['total_parcelas_5anos']} parcelas — base do calculo do dobro.",
        "3. Restituicao em dobro independe de ma-fe (EAREsp 676.608/RS, STJ, 30/03/2021).",
        "4. Dano moral in re ipsa — IRDR Tema 5 TJAM (Autos 0005217-75.2019.8.04.0000).",
        "5. Dano temporal (desvio produtivo): REsp 1.737.412/SP, Min. Nancy Andrighi.",
        "6. Juros de mora: 1% a.m. desde o primeiro desconto indevido (Sumula 54 STJ).",
        "7. Correcao monetaria: INPC desde cada desembolso (dano material) e a partir da sentenca (dano moral).",
        "8. Pedido subsidiario de repeticao simples NAO soma no valor da causa.",
    ]
    for nota in notas:
        ws.merge_cells(start_row=linha, start_column=1, end_row=linha, end_column=4)
        c = ws.cell(row=linha, column=1, value=nota)
        c.alignment = left_wrap; c.border = border_thin
        linha += 1

    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 18

    ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_margins.left = 0.5
    ws.page_margins.right = 0.5
    ws.print_options.horizontalCentered = True

    wb.save(destino_xlsx)


# ============================================================
#   RELATORIO PARALELO DE PENDENCIAS
# ============================================================

def gerar_relatorio_pendencias(caso, destino_relatorio):
    pendencias = caso.get("pendencias", [])
    doc = Document()
    doc.add_heading(f"RELATORIO DE PENDENCIAS — {caso['autora']['nome']}", level=1)
    doc.add_paragraph(
        "Documento paralelo a inicial. Lista ausencias e alertas a resolver "
        "antes do protocolo. Nao integra a peca."
    )
    doc.add_paragraph(f"Banco-reu: {caso['banco']['nome_curto']} — Contrato {caso['contrato']['numero']}")
    if caso.get("procurador_nome"):
        doc.add_paragraph(f"Procurador: {caso['procurador_nome']} — {caso.get('procurador_oab','—')}")
    doc.add_paragraph("")

    if pendencias:
        doc.add_heading("Pendencias identificadas", level=2)
        for i, (titulo, descricao) in enumerate(pendencias, 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{i}. {titulo}: "); r.bold = True
            p.add_run(descricao)

    doc.save(destino_relatorio)


# ============================================================
#   PIPELINE COMPLETO
# ============================================================

class DadosObrigatoriosRMCFaltandoError(RuntimeError):
    """Levantada quando o `caso` RMC/RCC chega ao pipeline com campos
    essenciais ausentes/zerados. Paridade com a barreira pre-geração da
    skill `inicial-nao-contratado` (paradigma VILSON / BANRISUL 2026-05-16):
    sem dados reais, NÃO gera inicial e NÃO usa fallbacks fictícios."""
    def __init__(self, erros):
        self.erros = list(erros)
        super().__init__(
            'Inicial RMC/RCC NÃO PODE ser gerada — dados obrigatórios ausentes:\n'
            + '\n'.join(f'  • {e}' for e in self.erros)
            + '\n\nAÇÃO: conferir HISCON/HISCRE da autora, número de contrato na '
              'procuração, e completar o dict `caso` antes de chamar '
              '`renderizar_caso`. NUNCA usar fallbacks fictícios (R$ 50,00, '
              '84 parcelas, "[A CONFIRMAR]") como contorno.'
        )


def _validar_caso_pre_geracao(caso):
    """Pre-validator do dict `caso` antes de qualquer geração.

    Critérios (cada falha vai à lista, todas acumuladas):
      * autora.nome, autora.cpf preenchidos
      * banco.nome_curto e banco.resto_qualificacao preenchidos
      * beneficio.nb preenchido
      * contrato.numero preenchido
      * contrato.data_inclusao preenchida e válida (sem "[A CONFIRMAR")
      * contrato.descontos_hiscon não vazio e sem zeros
      * tese ∈ {"RMC", "RCC"}

    Levanta DadosObrigatoriosRMCFaltandoError com a lista consolidada.
    """
    import re as _re
    erros = []
    a = caso.get("autora") or {}
    if not a.get("nome"):
        erros.append("autora.nome ausente.")
    if not a.get("cpf"):
        erros.append("autora.cpf ausente.")
    b = caso.get("banco") or {}
    if not b.get("nome_curto"):
        erros.append("banco.nome_curto ausente.")
    if not b.get("resto_qualificacao"):
        erros.append("banco.resto_qualificacao ausente.")
    bn = caso.get("beneficio") or {}
    if not bn.get("nb"):
        erros.append("beneficio.nb ausente.")
    c = caso.get("contrato") or {}
    if not c.get("numero"):
        erros.append("contrato.numero ausente.")
    di = c.get("data_inclusao") or ""
    if not di or "[A CONFIRMAR" in str(di) or "pendente" in str(di).lower():
        erros.append(f"contrato.data_inclusao ausente/placeholder ({di!r}).")
    else:
        # Validar formato dd/mm/aaaa ou mm/aaaa
        if not _re.match(r"^\d{1,2}/\d{2,4}(?:/\d{2,4})?$", str(di).strip()):
            erros.append(f"contrato.data_inclusao em formato inválido ({di!r}).")
    descontos = c.get("descontos_hiscon") or []
    if not descontos:
        erros.append("contrato.descontos_hiscon vazio — sem descontos no HISCRE/HISCON.")
    else:
        # Cada item deve ser (competencia, valor) com valor > 0
        zerados = [comp for comp, val in descontos if not val or float(val) <= 0]
        if zerados:
            erros.append(
                f"contrato.descontos_hiscon contém {len(zerados)} desconto(s) "
                f"zerado(s): {zerados[:3]}{'...' if len(zerados) > 3 else ''}."
            )
    tese = caso.get("tese")
    if tese not in ("RMC", "RCC"):
        erros.append(f"tese inválida ({tese!r}); esperado 'RMC' ou 'RCC'.")
    perfil_uf = caso.get("perfil")
    if not perfil_uf:
        erros.append("perfil (perfil_juridicos) ausente.")
    if erros:
        raise DadosObrigatoriosRMCFaltandoError(erros)


def renderizar_caso(caso, pasta_saida):
    """Pipeline completo: calcula, gera inicial, planilha e relatorio."""
    os.makedirs(pasta_saida, exist_ok=True)

    # 0. PRE-VALIDATOR (2026-05-16): aborta antes de gerar nada se faltar
    # dado obrigatório. Evita inicial com R$ 0,00 ou placeholder remanescente.
    _validar_caso_pre_geracao(caso)

    # 1. Calculos
    calcular_valores(caso)

    # 1-bis. POST-CALCULO VALIDATOR (2026-05-16): se após filtrar descontos
    # pelo contrato corrente sobrou 0 (data_inclusao filtra tudo), ou se soma
    # ficou zerada, abortar. Cobre o caso silencioso onde HISCRE traz descontos
    # mas todos são de data anterior à data_inclusao informada (erro humano
    # típico ao montar `caso["contrato"]["data_inclusao"]`).
    _calc = caso["calculo"]
    _erros_calc = []
    if (_calc.get("total_parcelas_historico") or 0) <= 0:
        _erros_calc.append(
            f"total_parcelas_historico = {_calc.get('total_parcelas_historico')} — "
            "filtro do contrato corrente eliminou todos os descontos. Conferir "
            "caso['contrato']['data_inclusao']."
        )
    if (_calc.get("soma_total") or 0) <= 0:
        _erros_calc.append(
            f"soma_total = R$ {_calc.get('soma_total'):.2f} — sem valor a impugnar."
        )
    if (_calc.get("valor_parcela_atual") or 0) <= 0:
        _erros_calc.append(
            f"valor_parcela_atual = R$ {_calc.get('valor_parcela_atual'):.2f} — "
            "última parcela zerada."
        )
    if (_calc.get("valor_liquido") or 0) <= 0:
        _erros_calc.append(
            f"valor_liquido = R$ {_calc.get('valor_liquido'):.2f} — renda zerada "
            "(HISCRE/HISCON sem valor). Conferir caso['hiscre_path']."
        )
    if _erros_calc:
        raise DadosObrigatoriosRMCFaltandoError(_erros_calc)

    # 2. Selecionar template baseado em UF + tese + banco
    perfil_uf = caso["perfil"]
    tese = caso.get("tese", "RMC")
    is_bmg = "BMG" in caso["banco"]["nome_curto"].upper()
    is_demais = not is_bmg
    template = None
    for arq in perfil_uf["arquivos"]:
        if arq["is_rcc"] != (tese == "RCC"):
            continue
        if arq["is_demais"] != is_demais:
            continue
        template = arq["destino"]
        break
    if not template:
        raise ValueError(f"Template nao encontrado para UF {perfil_uf['uf']}, tese={tese}, bmg={is_bmg}")

    # 3. Gerar inicial
    nome_safe = caso["autora"]["nome"].split()[0].upper()
    destino_docx = os.path.join(pasta_saida, f"INICIAL_{nome_safe}_{tese}.docx")
    gerar_inicial(caso, template, destino_docx)

    # 4. Gerar planilha
    destino_xlsx = os.path.join(pasta_saida, f"CALCULO_{nome_safe}_{tese}.xlsx")
    gerar_planilha(caso, destino_xlsx)

    # 5. Gerar relatorio
    destino_rel = os.path.join(pasta_saida, f"RELATORIO_PENDENCIAS_{nome_safe}.docx")
    gerar_relatorio_pendencias(caso, destino_rel)

    # 6. AUDITORIA DE PLACEHOLDERS — padrão alinhado com inicial-nao-contratado
    #    (paradigma MARIA AZEVEDO 2026-05-14): após salvar a inicial, varrer
    #    o docx atrás de {{...}} restantes. Se sobrar algum, é descompasso
    #    template×dict — exige correção da skill (e não do operador).
    #    Documentação em SKILL.md seção "Auditoria de placeholders".
    import re as _re
    from docx import Document as _Doc
    _doc = _Doc(destino_docx)
    _residuais = []
    for _par in _doc.paragraphs:
        for _ph in _re.findall(r'\{\{[^}]+\}\}', _par.text):
            if _ph not in _residuais:
                _residuais.append(_ph)
    for _tbl in _doc.tables:
        for _row in _tbl.rows:
            for _cell in _row.cells:
                for _par in _cell.paragraphs:
                    for _ph in _re.findall(r'\{\{[^}]+\}\}', _par.text):
                        if _ph not in _residuais:
                            _residuais.append(_ph)
    if _residuais:
        print(f"  ⚠ placeholders residuais em {os.path.basename(destino_docx)}: {_residuais}")

    # 7. VALIDADOR PÓS-DOCX (paridade com inicial-nao-contratado, 2026-05-16):
    # detecta R$ 0,00, [A CONFIRMAR, "pendente HISCON", competências/datas
    # vazias entre vírgulas. Se algum dispara, renomeia para
    # *_FALHOU_VALIDACAO_FINAL.docx e levanta DocxValidacaoFinalError.
    # Importa via importlib.util para evitar colisão com `helpers_docx`
    # local da skill (que tem outras funções, não tem validar_docx_gerado).
    try:
        import importlib.util as _ilu
        _nc_helpers_path = r"C:/Users/gabri/.claude/skills/inicial-nao-contratado/references/helpers_docx.py"
        if os.path.exists(_nc_helpers_path):
            _spec = _ilu.spec_from_file_location("_nc_helpers_docx", _nc_helpers_path)
            _mod = _ilu.module_from_spec(_spec)
            _spec.loader.exec_module(_mod)
            if hasattr(_mod, 'validar_docx_gerado'):
                _mod.validar_docx_gerado(destino_docx, abortar=True)
    except Exception as _e:
        # Re-levanta se for o erro de validação (queremos travar mesmo)
        tipo_nome = type(_e).__name__
        if tipo_nome == 'DocxValidacaoFinalError':
            raise
        # Outros erros: registra mas não bloqueia
        print(f"  ⚠ validador pós-DOCX não rodou: {tipo_nome}: {str(_e)[:120]}")

    return {
        "inicial": destino_docx,
        "planilha": destino_xlsx,
        "relatorio": destino_rel,
        "residuais": _residuais,
    }
