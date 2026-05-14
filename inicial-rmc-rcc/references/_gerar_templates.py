"""Gerador de templates padronizados para a skill `inicial-rmc-rcc`.

A partir dos templates originais do escritorio (em `APP - RMC-RCC/Tese R*/`),
gera versoes padronizadas em `APP - RMC-RCC/Templates Padronizados/<UF>/`
aplicando TODAS as 5 regras canonicas + extras do paradigma BENEDITA.

Uso:
    python _gerar_templates.py [UF]

Sem argumento: gera para TODAS as UFs ativas (com templates em disco).
Com UF: gera so para aquela UF (ex: python _gerar_templates.py AM).

Configuracao por UF: perfis_juridicos.py.
"""
import os
import shutil
import sys
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from helpers_redacao import (
    centralizar_celulas_tabela_quadro_sumario,
    aplicar_12pt_no_polo_passivo,
    limpar_highlights_pre_existentes,
)
from perfis_juridicos import PERFIS, perfil, listar_ufs_ativas


# ============================================================
#   MAPEAMENTOS CANONICOS (universais — todos os Estados)
# ============================================================

# Caixa baixa + nomes canonicos NC (alinhado com inicial-nao-contratado)
REPLACEMENTS_CANONICOS = [
    ("{{NOME_COMPLETO}}", "{{nome_completo}}"),
    ("{{QUALI_BANCO}}", "{{quali_banco}}"),
    ("{{numero_contrato}}", "{{numero_do_contrato}}"),
    ("{{valor_parcela}}", "{{valor_da_parcela}}"),
    ("{{parcela_mensal}}", "{{valor_da_parcela}}"),
    ("{{quantidade_parcelas}}", "{{total_de_parcelas}}"),
    ("{{inicio_desconto}}", "{{data_do_primeiro_desconto}}"),
    ("{{data_inicio_desconto}}", "{{data_do_primeiro_desconto}}"),
    ("{{data_averbacao}}", "{{data_da_inclusão}}"),
]

REPLACEMENTS_CONTA_CONSOLIDADA = [
    ("agência {{agencia}}, conta corrente nº {{numero_conta_corrente}}",
     "{{conta_agencia_conta}}"),
]

# Conjugacao genero (placeholders que o renderer resolve depois)
REPLACEMENTS_CONJUGACAO_GENERO = [
    ("inscrito (a)", "{{inscrito}}"),
    ("residente e domiciliado (a)", "residente e {{domiciliado}}"),
    ("domiciliado (a)", "{{domiciliado}}"),
]

# Pedido idoso — texto canonico NC
TEXTO_PEDIDO_IDOSO_NC = (
    "A prioridade na tramitação, tendo em vista que a parte autora é "
    "pessoa idosa, nos termos do art. 1.048, inciso I, do Código de "
    "Processo Civil;"
)
TEXTOS_PEDIDO_IDOSO_ANTIGOS = [
    "Seja concedida prioridade de tramitação, tendo em vista a parte Autora ser pessoa idosa, nos termos do artigo 1.048 do Novo Código de Processo Civil;",
]
TEXTO_CABECALHO_IDOSO = "Prioridade de tramitação: art. 1.048 do Código de Processo Civil (Idoso)."

# Frases que disparam wrap idoso (atualmente ausente em RCC)
WRAP_IDOSO_TARGETS = [
    "Prioridade de tramitação: art. 1.048 do Código de Processo Civil (Idoso).",
    "Seja concedida prioridade de tramitação, tendo em vista a parte Autora ser pessoa idosa, nos termos do artigo 1.048 do Novo Código de Processo Civil;",
]


# ============================================================
#   FUNCOES AUXILIARES
# ============================================================

def replace_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        return 0
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return 1
    full = paragraph.text.replace(old, new)
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full
    return 1


def make_jinja_paragraph(text, ref_paragraph):
    from copy import deepcopy
    p = OxmlElement("w:p")
    ref_pPr = ref_paragraph._p.find(qn("w:pPr"))
    if ref_pPr is not None:
        p.append(deepcopy(ref_pPr))
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    p.append(r)
    return p


def wrap_with_idoso(paragraph):
    p_open = make_jinja_paragraph("{% if idoso %}", paragraph)
    p_close = make_jinja_paragraph("{% endif %}", paragraph)
    paragraph._p.addprevious(p_open)
    paragraph._p.addnext(p_close)


def aplicar_formato_cabecalho_idoso(paragraph):
    """Cabecalho idoso: Cambria 11pt + alinhamento direita + recuo 4cm (2268 twips)."""
    p_elem = paragraph._p
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    for tag in ("w:ind", "w:jc"):
        existente = pPr.find(qn(tag))
        if existente is not None:
            pPr.remove(existente)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "2268")
    pPr.append(ind)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)
    for run in paragraph.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        for tag in ("w:rStyle", "w:rFonts"):
            existente = rpr.find(qn(tag))
            if existente is not None:
                rpr.remove(existente)
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Cambria")
        rfonts.set(qn("w:hAnsi"), "Cambria")
        rpr.insert(0, rfonts)
        for tag in ("w:sz", "w:szCs"):
            existente = rpr.find(qn(tag))
            if existente is not None:
                rpr.remove(existente)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")  # 11pt
        rpr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), "22")
        rpr.append(szCs)


def aplicar_formato_pedido_idoso(paragraph):
    """Pedido idoso: Cambria Bold."""
    for run in paragraph.runs:
        rpr = run._element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._element.insert(0, rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        rfonts.set(qn("w:ascii"), "Cambria")
        rfonts.set(qn("w:hAnsi"), "Cambria")
        if rpr.find(qn("w:b")) is None:
            rpr.append(OxmlElement("w:b"))
        if rpr.find(qn("w:bCs")) is None:
            rpr.append(OxmlElement("w:bCs"))


# ============================================================
#   PIPELINE DE GERACAO POR ARQUIVO
# ============================================================

def processa(spec, perfil_uf):
    print(f"\n=== {spec['label']} ({perfil_uf['uf']}) ===")
    print(f"   origem : {spec['origem']}")
    print(f"   destino: {spec['destino']}")

    if not os.path.exists(spec["origem"]):
        print(f"   !! ARQUIVO ORIGEM NAO EXISTE — pulando")
        return

    os.makedirs(os.path.dirname(spec["destino"]), exist_ok=True)
    shutil.copyfile(spec["origem"], spec["destino"])

    doc = Document(spec["destino"])

    # 0. Limpar highlights legacy do template
    n_clean = limpar_highlights_pre_existentes(doc)
    print(f"   [0] highlights legacy removidos: {n_clean}")

    # 1. Mapeamentos canonicos (caixa baixa + nomes alinhados com NC)
    count_canon = 0
    for paragraph in doc.paragraphs:
        for old, new in REPLACEMENTS_CANONICOS:
            count_canon += replace_in_paragraph(paragraph, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in REPLACEMENTS_CANONICOS:
                        count_canon += replace_in_paragraph(paragraph, old, new)
    print(f"   [1] substituicoes canonicas: {count_canon}")

    # 1b. Consolidacao agencia+conta -> {{conta_agencia_conta}}
    count_conta = 0
    for paragraph in doc.paragraphs:
        for old, new in REPLACEMENTS_CONTA_CONSOLIDADA:
            count_conta += replace_in_paragraph(paragraph, old, new)
    print(f"   [1b] consolidacao conta: {count_conta}")

    # 2. Bug UF errada (template AM-Demais tem /BA — outros UF podem ter outras inconsistencias)
    if perfil_uf["uf"] == "AM" and spec["is_demais"]:
        count_uf = 0
        for paragraph in doc.paragraphs:
            count_uf += replace_in_paragraph(
                paragraph,
                "Comarca de {{competencia}}/BA",
                "Comarca de {{competencia}}/AM"
            )
        print(f"   [2] correcao UF (BA->AM): {count_uf}")

    # 3. RCC: adiciona {% if idoso %}/{% endif %} onde ausente
    if spec["is_rcc"]:
        count_idoso = 0
        for paragraph in list(doc.paragraphs):
            for target in WRAP_IDOSO_TARGETS:
                if target in paragraph.text:
                    prev = paragraph._p.getprevious()
                    prev_text = ""
                    if prev is not None:
                        ts = prev.findall(".//" + qn("w:t"))
                        prev_text = "".join(t.text or "" for t in ts)
                    if "{% if idoso %}" in prev_text:
                        continue
                    wrap_with_idoso(paragraph)
                    count_idoso += 1
        print(f"   [3] wraps idoso adicionados: {count_idoso}")

    # 4. Texto canonico NC do PEDIDO idoso
    count_pedido = 0
    for paragraph in doc.paragraphs:
        for antigo in TEXTOS_PEDIDO_IDOSO_ANTIGOS:
            if antigo in paragraph.text:
                replace_in_paragraph(paragraph, antigo, TEXTO_PEDIDO_IDOSO_NC)
                count_pedido += 1
    print(f"   [4] pedido idoso (texto canonico NC): {count_pedido}")

    # 5. Formato CABECALHO idoso (Cambria 11pt + direita + recuo 4cm)
    count_cab = 0
    for paragraph in doc.paragraphs:
        if TEXTO_CABECALHO_IDOSO in paragraph.text:
            aplicar_formato_cabecalho_idoso(paragraph)
            count_cab += 1
    print(f"   [5] formato cabecalho idoso: {count_cab}")

    # 6. Formato PEDIDO idoso (Cambria Bold)
    count_ped = 0
    for paragraph in doc.paragraphs:
        if TEXTO_PEDIDO_IDOSO_NC in paragraph.text:
            aplicar_formato_pedido_idoso(paragraph)
            count_ped += 1
    print(f"   [6] formato pedido idoso: {count_ped}")

    # 7. Endereco escritorio na qualificacao
    trechos = [
        ("por seus advogados, que assinam digitalmente a presente peça (instrumento de procuração anexo), vem, respeitosamente,",
         "por seus advogados, que assinam digitalmente a presente peça (instrumento de procuração anexo), com escritório profissional em {{endereco_escritorio}}, local onde recebem avisos e intimações, vem, respeitosamente,"),
        ("por seus advogados que assinam digitalmente a presente peça (instrumento de procuração anexo), vem, respeitosamente,",
         "por seus advogados que assinam digitalmente a presente peça (instrumento de procuração anexo), com escritório profissional em {{endereco_escritorio}}, local onde recebem avisos e intimações, vem, respeitosamente,"),
    ]
    count_end = 0
    for paragraph in doc.paragraphs:
        for old, new in trechos:
            if old in paragraph.text:
                replace_in_paragraph(paragraph, old, new)
                count_end += 1
                break
    print(f"   [7] trecho escritorio inserido: {count_end}")

    # 8. Conjugacao genero — inscrito(a)/domiciliado(a) -> placeholders
    count_gen = 0
    for paragraph in doc.paragraphs:
        for old, new in REPLACEMENTS_CONJUGACAO_GENERO:
            count_gen += replace_in_paragraph(paragraph, old, new)
    print(f"   [8] conjugacao genero (placeholders): {count_gen}")

    # 9. REGRA 4 — 12pt no polo passivo
    n_12 = aplicar_12pt_no_polo_passivo(doc)
    print(f"   [9] 12pt no polo passivo: {n_12}")

    # 10. REGRA 3 — Centralizar celulas tabela Quadro Sumario
    n_cent = centralizar_celulas_tabela_quadro_sumario(doc)
    print(f"   [10] tabela Quadro Sumario centralizada: {n_cent}")

    doc.save(spec["destino"])


# ============================================================
#   MAIN
# ============================================================

def gerar_uf(uf):
    p = perfil(uf)
    if not p:
        print(f"UF nao configurada: {uf}")
        return
    print(f"\n{'='*60}\nGERANDO templates para UF: {uf}\n{'='*60}")
    for spec in p["arquivos"]:
        processa(spec, p)


def gerar_todas():
    ufs = listar_ufs_ativas()
    print(f"UFs com templates em disco: {ufs}")
    for uf in ufs:
        gerar_uf(uf)


if __name__ == "__main__":
    if len(sys.argv) > 1:
        gerar_uf(sys.argv[1])
    else:
        gerar_todas()
