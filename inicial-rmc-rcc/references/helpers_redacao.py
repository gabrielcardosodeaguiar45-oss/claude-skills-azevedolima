"""Helpers de redacao da skill `inicial-rmc-rcc`.

Consolida as 5 REGRAS CANONICAS do paradigma BENEDITA (2026-05-13, refinado
ao longo da sessao):

  FONTE AUTORITATIVA: HISCRE (Historico de Creditos do INSS).
  NAO USAR HISCON para valores ou contagem — usar APENAS HISCRE.
  O HISCRE mostra mes a mes o que foi efetivamente debitado do beneficio
  e o valor liquido REAL recebido pela autora em cada competencia.

  RUBRICAS OFICIAIS DO HISCRE (atualizado 2026-05-14, Gabriel):
    217 EMPRESTIMO SOBRE A RMC      → RMC (rotativa)
    268 CONSIGNACAO - CARTAO         → RCC (cartão consignado)
    216 CONSIGNACAO EMPRESTIMO BANCARIO  → consignado tradicional (fora do escopo desta skill)
  ATENÇÃO: rubrica 218 NÃO existe no HISCRE INSS — era erro do verificador antigo.

  1. Valor liquido do beneficio = campo 'Valor Liquido' do HISCRE da
     competencia mais recente (NUNCA calcular base HISCON - comprometido)
  2. Contagem de descontos = TODAS as ocorrencias da rubrica 217 (RMC) ou
     268 (RCC) no HISCRE, mes a mes, com o VALOR EXATO de cada uma
  3. Tabela 'Quadro Sumario' (1a tabela do doc) = conteudo de TODAS as
     celulas centralizado (tabela em si com recuo a direita, correta)
  4. Polo passivo = letra tamanho 12pt
  5. quali_banco = 2 runs (nome em Segoe UI Bold rStyle 2TtuloChar + resto
     em Cambria, ambos 12pt + grifo amarelo)

Mais:
  - Conjugacao automatica de genero (brasileira -> inscrita/domiciliada;
    brasileiro -> inscrito/domiciliado)
  - Omissao limpa de estado_civil ausente (sem aviso, sem virgula dupla)
  - Cabecalho idoso em Cambria 11pt + alinhamento direita + recuo 4cm
  - Pedido idoso em Cambria Bold com texto canonico NC
  - Endereco escritorio (matriz Joacaba/SC + unidade de apoio na UF) na
    qualificacao
  - GRIFO AMARELO em TUDO que a skill substituir
  - Limpeza de grifos amarelos legacy do template (Patrick antigos)
"""
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from helpers_docx import (
    iter_all_paragraphs,
    substituir_in_run,
    make_run_cambria_grifado,
    make_run_segoe_bold_grifado,
)


# ============================================================
#   REGRA 1 — Valor liquido do beneficio
# ============================================================

def calcular_valor_liquido(base_calculo_hiscon, total_comprometido_hiscon):
    """Valor liquido = base de calculo - total comprometido (HISCON pag 2).
    Returns: float
    """
    return float(base_calculo_hiscon) - float(total_comprometido_hiscon)


# ============================================================
#   REGRA 2 — Contagem de parcelas do contrato CORRENTE
# ============================================================

def extrair_descontos_hiscre(pdf_hiscre_path, rubrica="217"):
    """Extrai do HISCRE (Historico de Creditos do INSS) TODOS os descontos da
    rubrica especificada, com competencia, valor liquido recebido e valor
    da rubrica.

    Esta eh a FONTE AUTORITATIVA para a contagem e o valor dos descontos
    de RMC/RCC. Use SEMPRE em vez de presumir/inferir do HISCON.

    Args:
        pdf_hiscre_path: caminho do PDF do HISCRE
        rubrica: '217' para EMPRESTIMO SOBRE A RMC (default) ou
                 '268' para CONSIGNACAO - CARTAO (RCC)

    Returns:
        list[dict] cronologico (mais recente primeiro), cada item com:
            'competencia': 'MM/YYYY'
            'valor_liquido': float (valor liquido recebido na competencia)
            'valor_rubrica': float (valor descontado a titulo de RMC/RCC)
    """
    import pdfplumber
    import re
    # pdfplumber preserva estrutura linha-a-linha do HISCRE (cada bloco de
    # competência em UMA linha). pymupdf quebra cada token em linha própria
    # e o regex de competência não casa.
    with pdfplumber.open(pdf_hiscre_path) as pdf:
        texto = "\n".join((p.extract_text() or "") for p in pdf.pages)

    descricao = {
        "217": "EMPRESTIMO SOBRE A RMC",
        "268": "CONSIGNACAO - CARTAO",
    }.get(rubrica, f"RUBRICA {rubrica}")

    # Bloco de competência no HISCRE: linha tipo
    #   "MM/AAAA DD/MM/AAAA R$ valor CCF - CONTA-CORRENTE Pago ..."
    # ou
    #   "MM/AAAA DD/MM/AAAA R$ valor CMG - CARTAO MAGNETICO Pago ..."
    # A quebra de linha do pdfplumber pode separar "CCF -" do "CONTA-CORRENTE".
    # Para não depender do meio de pagamento, usamos só competência + valor
    # líquido inicial. O bloco vai até a próxima competência (ou EOF).
    padrao_bloco = re.compile(
        r"(\d{2}/\d{4})\s+\d{2}/\d{2}/\d{4}\s+"
        r"R\$\s*([\d.]+,\d{2})",
        re.MULTILINE,
    )
    padrao_rubrica = re.compile(
        rf"{rubrica}\s+{re.escape(descricao)}\s+R\$\s*([\d.]+,\d{{2}})"
    )

    matches = list(padrao_bloco.finditer(texto))
    out = []
    for i, m in enumerate(matches):
        comp = m.group(1)
        val_liquido = float(m.group(2).replace(".", "").replace(",", "."))
        inicio = m.end()
        fim = matches[i + 1].start() if i + 1 < len(matches) else len(texto)
        bloco = texto[inicio:fim]
        m_rub = padrao_rubrica.search(bloco)
        if m_rub:
            val_rub = float(m_rub.group(1).replace(".", "").replace(",", "."))
            out.append({
                "competencia": comp,
                "valor_liquido": val_liquido,
                "valor_rubrica": val_rub,
            })

    def chave(x):
        mes, ano = x["competencia"].split("/")
        return (int(ano), int(mes))
    out.sort(key=chave, reverse=True)
    return out


def filtrar_descontos_contrato_corrente(descontos, data_inclusao_str):
    """Filtra descontos pela DATA DE INCLUSAO do contrato corrente.

    USO RESTRITO: aplicar APENAS quando todos os identificadores no HISCON
    sao do mesmo formato (ou todos antigos baseados em NB, sem possibilidade
    de migracao de contrato). Em casos onde o HISCON mistura formatos novo
    (`<contrato><banco>MMYYYY`) e antigo (`<NB>MMYYYY`), preferir
    `filtrar_descontos_identificacao_explicita` — mais defensivo.

    Args:
        descontos: list[(competencia_str, valor)] cronologico (mais recente primeiro)
        data_inclusao_str: 'DD/MM/AAAA' — data de inclusao do contrato corrente

    Returns: lista filtrada (apenas competencia >= mes_inicio do contrato)
    """
    dia, mes_inc, ano_inc = data_inclusao_str.split("/")
    mes_inc = int(mes_inc)
    ano_inc = int(ano_inc)

    def _competencia_ok(comp_str):
        mes, ano = comp_str.split("/")
        mes = int(mes); ano = int(ano)
        if ano > ano_inc:
            return True
        if ano == ano_inc and mes >= mes_inc:
            return True
        return False

    return [(c, v) for c, v in descontos if _competencia_ok(c)]


def filtrar_descontos_identificacao_explicita(descontos_raw, prefixo_identificador):
    """Filtra descontos do HISCON pelo PREFIXO do identificador.

    Aplica a regra 'pega so o que esta inequivocamente no extrato' (paradigma
    BENEDITA, refinado 2026-05-13):

    O HISCON do INSS usa 2 formatos de identificador para descontos de RMC:
      * NOVO (a partir de ~11/2022): `<contrato><banco_3digitos>MMYYYY`
        ex: `12257818318032026` -> contrato 12257818, banco 318 (BMG), 03/2026
      * ANTIGO (ate ~10/2022): `<NB>MMYYYY` baseado no Numero do Beneficio
        ex: `60555992050032020` -> NB 605559920-5, 03/2020

    O formato ANTIGO eh AMBIGUO porque o mesmo NB pode ter hospedado mais de
    um contrato RMC (ex: contrato anterior excluido + contrato atual). Sem
    identificacao explicita do contrato, esses descontos NAO devem ser
    atribuidos automaticamente ao contrato corrente.

    Use esta funcao para extrair APENAS os descontos com identificador NOVO
    (que comecam com o numero do contrato corrente). Isso eh defensivo: o
    banco nao pode contestar a contagem dizendo que pertencia ao contrato
    anterior.

    Args:
        descontos_raw: list[(identificador, competencia, valor)] como vem do
                       parser do HISCON (com 3 campos)
        prefixo_identificador: str — comeco do identificador inequivoco do
                                contrato corrente. Ex: "12257818318" (numero
                                do contrato + 318 do BMG)

    Returns: list[(competencia, valor)] — somente os descontos inequivocos.
    """
    return [
        (comp, val)
        for ident, comp, val in descontos_raw
        if ident.startswith(prefixo_identificador)
    ]


def dentro_prescricao_quinquenal(competencia_str, marco_mes_ano=None):
    """True se a competencia esta dentro dos 5 anos (CDC art. 27)."""
    from datetime import date
    if marco_mes_ano is None:
        hoje = date.today()
        marco_mes = hoje.month
        marco_ano = hoje.year - 5
    else:
        marco_mes, marco_ano = marco_mes_ano
    mes, ano = competencia_str.split("/")
    mes = int(mes); ano = int(ano)
    if ano > marco_ano:
        return True
    if ano == marco_ano and mes >= marco_mes:
        return True
    return False


# ============================================================
#   REGRA 3 — Centralizar celulas da tabela Quadro Sumario
# ============================================================

def centralizar_celulas_tabela_quadro_sumario(doc):
    """Centraliza o conteudo de TODAS as celulas da primeira tabela do doc
    (Quadro Sumario, posicionada entre prioridade idoso e qualificacao).
    Retorna numero de celulas alteradas."""
    if not doc.tables:
        return 0
    tabela = doc.tables[0]
    total = 0
    for row in tabela.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                pPr = paragraph._p.find(qn("w:pPr"))
                if pPr is None:
                    pPr = OxmlElement("w:pPr")
                    paragraph._p.insert(0, pPr)
                jc = pPr.find(qn("w:jc"))
                if jc is None:
                    jc = OxmlElement("w:jc")
                    pPr.append(jc)
                if jc.get(qn("w:val")) != "center":
                    jc.set(qn("w:val"), "center")
                    total += 1
    return total


# ============================================================
#   REGRA 4 — Tamanho 12pt no polo passivo
# ============================================================

def aplicar_12pt_no_polo_passivo(doc):
    """Aplica w:sz=24 + w:szCs=24 (12pt) a TODOS os runs do paragrafo do polo
    passivo (que comeca com 'em face de' e contem '{{quali_banco}}')."""
    total = 0
    for paragraph in doc.paragraphs:
        if "em face de" in paragraph.text and "{{quali_banco}}" in paragraph.text:
            for run in paragraph.runs:
                rpr = run._element.find(qn("w:rPr"))
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    run._element.insert(0, rpr)
                for tag in ("w:sz", "w:szCs"):
                    existente = rpr.find(qn(tag))
                    if existente is not None:
                        rpr.remove(existente)
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "24")
                rpr.append(sz)
                szCs = OxmlElement("w:szCs")
                szCs.set(qn("w:val"), "24")
                rpr.append(szCs)
                total += 1
            break
    return total


# ============================================================
#   REGRA 5 — quali_banco em 2 runs (Segoe Bold + Cambria, 12pt + grifo)
# ============================================================

def substituir_quali_banco(doc, nome_banco, resto_banco):
    """Substitui {{quali_banco}} criando 2 runs separados:
       run 1: nome banco (Segoe UI Bold via rStyle 2TtuloChar, 12pt, amarelo)
       run 2: resto (Cambria 12pt, amarelo)
    """
    placeholder = "{{quali_banco}}"
    total = 0
    for paragraph in iter_all_paragraphs(doc):
        if placeholder not in paragraph.text:
            continue
        target_run_idx = None
        for i, run in enumerate(paragraph.runs):
            if placeholder in run.text:
                target_run_idx = i
                break
        if target_run_idx is None:
            full = paragraph.text
            for run in paragraph.runs[1:]:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = full
                target_run_idx = 0
            else:
                continue
        target_run = paragraph.runs[target_run_idx]
        texto = target_run.text
        idx = texto.find(placeholder)
        antes = texto[:idx]
        depois = texto[idx + len(placeholder):]
        target_run.text = antes

        r_nome = make_run_segoe_bold_grifado(nome_banco, preserve=False)
        r_resto = make_run_cambria_grifado(resto_banco + depois, preserve=True)
        target_run._element.addnext(r_resto)
        target_run._element.addnext(r_nome)
        total += 1
    return total


# ============================================================
#   CONJUGACAO DE GENERO
# ============================================================

def conjugacao_por_nacionalidade(nacionalidade):
    """Retorna ('inscrita','domiciliada') ou ('inscrito','domiciliado')."""
    fem = (nacionalidade or "").endswith("a")
    if fem:
        return "inscrita", "domiciliada"
    return "inscrito", "domiciliado"


# ============================================================
#   OMISSAO LIMPA DE ESTADO_CIVIL VAZIO
# ============================================================

def omitir_estado_civil_se_vazio(doc, estado_civil):
    """Se estado_civil estiver vazio, remove ', {{estado_civil}}' do template
    (incluindo a virgula anterior) sem grifo."""
    if estado_civil:
        return 0
    total = 0
    for p in iter_all_paragraphs(doc):
        if ", {{estado_civil}}" in p.text:
            substituir_in_run(p._element, {", {{estado_civil}}": ""}, grifo=False)
            total += 1
    return total


# ============================================================
#   IDOSO — wrap Jinja + formato cabecalho + texto canonico pedido
# ============================================================

def resolver_jinja_idoso(doc, eh_idoso):
    """Resolve {% if idoso %}...{% endif %}:
       - eh_idoso=True  -> remove os marcadores, mantem o conteudo
       - eh_idoso=False -> remove o bloco INTEIRO (marcadores + conteudo)
    """
    elementos_remover = []
    dentro_bloco = False
    for p in doc.paragraphs:
        txt = p.text
        if "{% if idoso %}" in txt:
            elementos_remover.append(p._element)
            dentro_bloco = True
            continue
        if "{% endif %}" in txt:
            elementos_remover.append(p._element)
            dentro_bloco = False
            continue
        if dentro_bloco and not eh_idoso:
            elementos_remover.append(p._element)
    for elem in elementos_remover:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)


# ============================================================
#   LIMPEZA DE GRIFOS LEGACY (template base)
# ============================================================

def limpar_highlights_pre_existentes(doc):
    """Remove highlights antigos do template base. Apos isso, qualquer grifo
    amarelo no documento veio EXCLUSIVAMENTE da skill (substituicoes)."""
    total = 0
    for p in doc.paragraphs:
        for run in p.runs:
            rpr = run._element.find(qn("w:rPr"))
            if rpr is None:
                continue
            hl = rpr.find(qn("w:highlight"))
            if hl is not None:
                rpr.remove(hl)
                total += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        rpr = run._element.find(qn("w:rPr"))
                        if rpr is None:
                            continue
                        hl = rpr.find(qn("w:highlight"))
                        if hl is not None:
                            rpr.remove(hl)
                            total += 1
    return total


# ============================================================
#   SUBSTITUICAO COM GRIFO (helper de alto nivel)
# ============================================================

def substituir_placeholders_com_grifo(doc, dados):
    """Aplica todas as substituicoes do dict `dados` em TODOS os paragrafos
    do doc (corpo + tabelas), com grifo amarelo automatico nos chars
    substituidos.
    """
    total = 0
    for paragraph in iter_all_paragraphs(doc):
        if any(ph in paragraph.text for ph in dados):
            mapa = {ph: val for ph, val in dados.items() if ph in paragraph.text}
            if substituir_in_run(paragraph._element, mapa, grifo=True):
                total += len(mapa)
    return total
