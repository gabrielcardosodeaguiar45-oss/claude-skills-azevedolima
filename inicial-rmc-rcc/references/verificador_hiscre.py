"""Verificador de HISCRE — decide se o documento esta completo o suficiente
para gerar a inicial RMC/RCC ou se precisa de pendencia (rebaixar HISCRE).

REGRA (paradigma BENEDITA + Claudio, 2026-05-13):
Sem HISCRE completo, NAO gerar inicial. Gerar RELATORIO_PENDENCIA_HISCRE.docx
informando que precisa baixar HISCRE completo para fazer o calculo correto.

Criterios de HISCRE completo:
  1. Header 'Compet. Final' >= 5 anos antes da Compet. Inicial (cobre o
     periodo prescricional de 5 anos do CDC)
  2. Os pagamentos efetivos (CCF) cobrem o periodo declarado no header
     (sem gap > 6 meses)
  3. Para o contrato a impugnar, a rubrica esperada (217 RMC / 268 RCC)
     tem ocorrencias suficientes para a contagem

  RUBRICAS OFICIAIS DO HISCRE:
  - 217 EMPRESTIMO SOBRE A RMC          → RMC (rotativa)
  - 268 CONSIGNACAO - CARTAO            → RCC (cartão consignado)
  (rubrica 218 NÃO existe no HISCRE INSS — era erro histórico do verificador)

Se qualquer falhar -> INCOMPLETO -> gerar pendencia.
"""
import re
import pymupdf
from datetime import date


def _competencia_para_tupla(comp_str):
    """'MM/AAAA' -> (ano, mes)"""
    mes, ano = comp_str.split("/")
    return (int(ano), int(mes))


def _meses_entre(comp_a, comp_b):
    """Diferenca em meses entre duas competencias (a - b)."""
    ya, ma = _competencia_para_tupla(comp_a)
    yb, mb = _competencia_para_tupla(comp_b)
    return (ya - yb) * 12 + (ma - mb)


def verificar_hiscre(pdf_hiscre_path, rubrica_esperada="217"):
    """Analisa um HISCRE e retorna diagnostico.

    Args:
        pdf_hiscre_path: caminho do PDF
        rubrica_esperada: '217' (RMC) ou '268' (RCC)

    Returns:
        dict com:
            'completo': bool
            'motivos_incompletude': list[str]
            'compet_inicial_header': 'MM/AAAA' ou None (mais recente)
            'compet_final_header': 'MM/AAAA' ou None (mais antiga)
            'compet_pagamentos_mais_antiga': ...
            'compet_pagamentos_mais_recente': ...
            'pagamentos_count': int
            'rubrica_count': int
            'rubrica_soma': float
    """
    doc = pymupdf.open(pdf_hiscre_path)
    texto = "\n".join(doc.load_page(i).get_text() for i in range(doc.page_count))

    # Header — Compet. Inicial (MAIS RECENTE) e Compet. Final (MAIS ANTIGA)
    # No HISCRE INSS, "Inicial" eh a mais recente; "Final" eh a mais antiga
    m_ini = re.search(r"Compet\.\s*Inicial[:\s]*\n?(\d{2}/\d{4})", texto)
    m_fim = re.search(r"Compet\.\s*Final[:\s]*\n?(\d{2}/\d{4})", texto)
    # Algumas variacoes (com quebra de linha entre Inicial: e o valor)
    if m_ini is None:
        m_ini = re.search(r"(\d{2}/\d{4})\s*Compet\.\s*Inicial", texto)
    if m_fim is None:
        m_fim = re.search(r"(\d{2}/\d{4})\s*Compet\.\s*Final", texto)
    compet_inicial_header = m_ini.group(1) if m_ini else None
    compet_final_header = m_fim.group(1) if m_fim else None

    # Competencias com pagamentos (CCF)
    matches_pag = re.findall(r"(\d{2}/\d{4})\s+R\$\s*[\d.]+,\d{2}\s+CCF", texto)
    comps_pag = sorted(set(matches_pag), key=_competencia_para_tupla)

    # Ocorrencias da rubrica esperada
    nome_rubrica = {"217": "EMPRESTIMO SOBRE A RMC", "268": "CONSIGNACAO - CARTAO"}.get(
        rubrica_esperada, ""
    )
    rubrica_matches = re.findall(
        rf"{rubrica_esperada}\s+{re.escape(nome_rubrica)}\s+R\$\s*([\d.]+,\d{{2}})",
        texto,
    )
    valores = [float(v.replace(".", "").replace(",", ".")) for v in rubrica_matches]

    # REGRA REFINADA (2026-05-13 noite, usuario):
    #   - rubrica >= 1 ocorrencia -> gera inicial COM O QUE HOUVER (gera_inicial=True)
    #   - rubrica == 0 ocorrencias -> pendencia (gera_inicial=False)
    # Avisos sobre cobertura parcial vao no relatorio paralelo, mas NAO impedem
    # a inicial se ha pelo menos 1 desconto.
    motivos_bloqueio = []   # impedem geracao da inicial
    avisos = []              # so registram, nao impedem
    gera_inicial = True

    if not rubrica_matches:
        gera_inicial = False
        motivos_bloqueio.append(
            f"Rubrica {rubrica_esperada} ({nome_rubrica}) NAO aparece no HISCRE — "
            f"sem prova de qualquer desconto consignado dessa modalidade. "
            f"Sem base para o calculo da repeticao em dobro."
        )

    # Aviso 1: header eh curto
    if compet_inicial_header and compet_final_header:
        meses_header = _meses_entre(compet_inicial_header, compet_final_header)
        if meses_header < 60:
            avisos.append(
                f"Periodo declarado no header eh inferior a 5 anos: "
                f"{compet_final_header} a {compet_inicial_header} "
                f"({meses_header} meses). Pode nao cobrir toda a prescricao CDC."
            )
    else:
        avisos.append("Header sem Compet. Inicial/Final detectaveis.")

    # Aviso 2: gap entre header declarado e pagamento mais antigo
    if comps_pag and compet_final_header:
        pag_mais_antigo = comps_pag[0]
        gap_inicio = _meses_entre(pag_mais_antigo, compet_final_header)
        if gap_inicio > 6:
            avisos.append(
                f"Gap entre Compet. Final declarada ({compet_final_header}) e "
                f"primeiro pagamento efetivo ({pag_mais_antigo}): {gap_inicio} "
                f"meses sem dados. HISCRE foi baixado parcialmente — pode haver "
                f"descontos anteriores nao computados."
            )

    return {
        "completo": gera_inicial,          # mantido por compat com runner antigo
        "gera_inicial": gera_inicial,
        "motivos_bloqueio": motivos_bloqueio,
        "motivos_incompletude": motivos_bloqueio + avisos,  # compat
        "avisos": avisos,
        "compet_inicial_header": compet_inicial_header,
        "compet_final_header": compet_final_header,
        "compet_pagamentos_mais_antiga": comps_pag[0] if comps_pag else None,
        "compet_pagamentos_mais_recente": comps_pag[-1] if comps_pag else None,
        "pagamentos_count": len(comps_pag),
        "rubrica_count": len(rubrica_matches),
        "rubrica_soma": sum(valores),
    }


def gerar_relatorio_pendencia_hiscre(
    destino_docx,
    cliente_nome,
    cpf,
    nb_beneficio,
    pasta_acao_label,
    contratos_impugnar,
    diagnostico,
    procurador_nome=None,
    procurador_oab=None,
):
    """Gera o RELATORIO_PENDENCIA_HISCRE.docx para um caso onde nao foi
    possivel fazer a inicial por HISCRE incompleto.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(
        f"PENDÊNCIA — HISCRE INCOMPLETO ({cliente_nome})", level=1
    )
    doc.add_paragraph(
        "Documento de PENDÊNCIA. NÃO foi possível gerar a petição inicial "
        "deste caso porque o Histórico de Créditos (HISCRE) baixado do MEU "
        "INSS está incompleto e impede o cálculo correto do valor da causa."
    )
    doc.add_paragraph(f"Cliente: {cliente_nome}")
    doc.add_paragraph(f"CPF: {cpf or '—'}")
    doc.add_paragraph(f"NB do benefício: {nb_beneficio or '—'}")
    doc.add_paragraph(f"Pasta de ação: {pasta_acao_label}")
    if procurador_nome:
        doc.add_paragraph(f"Procurador atribuído: {procurador_nome} — {procurador_oab or '—'}")

    doc.add_heading("Contratos a impugnar", level=2)
    if contratos_impugnar:
        for c in contratos_impugnar:
            doc.add_paragraph(
                f"• Contrato {c.get('contrato','?')} — {c.get('banco_nome_completo','?')} — "
                f"Tipo {c.get('tipo','?')} — Situação: {c.get('situacao','?')}",
                style="List Bullet",
            )
    else:
        doc.add_paragraph("(nenhum contrato listado na pasta de ação)")

    doc.add_heading("Diagnóstico do HISCRE", level=2)
    info = doc.add_paragraph()
    info.add_run("Competência declarada no header: ").bold = True
    info.add_run(
        f"{diagnostico['compet_final_header'] or '—'} a "
        f"{diagnostico['compet_inicial_header'] or '—'}"
    )
    info = doc.add_paragraph()
    info.add_run("Pagamentos efetivos encontrados: ").bold = True
    info.add_run(
        f"{diagnostico['pagamentos_count']} competências "
        f"({diagnostico['compet_pagamentos_mais_antiga'] or '—'} a "
        f"{diagnostico['compet_pagamentos_mais_recente'] or '—'})"
    )
    info = doc.add_paragraph()
    info.add_run("Rubrica esperada (217 RMC / 268 RCC): ").bold = True
    info.add_run(
        f"{diagnostico['rubrica_count']} ocorrências — total R$ {diagnostico['rubrica_soma']:,.2f}".replace(",", "#").replace(".", ",").replace("#", ".")
    )

    doc.add_heading("Motivos da pendência", level=2)
    if diagnostico["motivos_incompletude"]:
        for motivo in diagnostico["motivos_incompletude"]:
            doc.add_paragraph(motivo, style="List Bullet")
    else:
        doc.add_paragraph("(motivo não identificado automaticamente — revisar manualmente)")

    doc.add_heading("Ação necessária", level=2)
    doc.add_paragraph(
        "Baixar do MEU INSS (https://meu.inss.gov.br) o HISCRE COMPLETO da "
        "autora, cobrindo no mínimo 5 anos retroativos a contar da data "
        "atual (prescrição quinquenal do CDC, art. 27). Idealmente, baixar "
        "desde a data de averbação do contrato mais antigo a impugnar."
    )
    doc.add_paragraph(
        "Após obter o HISCRE completo, substituir o arquivo "
        "'7. Histórico de créditos.pdf' na pasta de ação e re-executar a "
        "geração da inicial."
    )

    doc.save(destino_docx)


if __name__ == "__main__":
    # Self-test rapido com o HISCRE do Claudio
    HISCRE = r"C:/Users/gabri/OneDrive/Área de Trabalho/APP - NÃO CONTRATADO/RMC - RCC - NÃO CONTRATADO/CLAUDIO DAS NEVES AMORIM - Elizio - RMC, RCC/0. Kit/historico-creditos (5).pdf"
    diag = verificar_hiscre(HISCRE, rubrica_esperada="217")
    print("=== DIAGNOSTICO CLAUDIO (rubrica 217 RMC) ===")
    for k, v in diag.items():
        print(f"  {k}: {v}")
    diag_rcc = verificar_hiscre(HISCRE, rubrica_esperada="268")
    print("\n=== DIAGNOSTICO CLAUDIO (rubrica 268 RCC) ===")
    for k, v in diag_rcc.items():
        print(f"  {k}: {v}")
