"""Helpers DOCX run-aware para a skill `inicial-rmc-rcc`.

Reaproveitado integralmente da skill `inicial-nao-contratado/references/helpers_docx.py`
(funcao `substituir_in_run`, linha 230). Aplica GRIFO AMARELO automatico nos
caracteres SUBSTITUIDOS, preservando rPr (formatacao) de origem dos chars que
NAO foram substituidos.

REGRA #4 do paradigma Benedita: TODO campo modificado pela skill recebe
highlight=yellow para facilitar revisao visual.
"""
import copy
from lxml import etree
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
XMLSPC = "{http://www.w3.org/XML/1998/namespace}space"


def substituir_in_run(p_elem, mapa, grifo=True):
    """Substitui chave->valor em todos os runs de um <w:p>, aplicando grifo
    amarelo APENAS nos caracteres substituidos.

    Algoritmo (do NC):
    1. Lista cada char com (char, rpr_origem, inserido=False)
    2. Para cada par do mapa, substitui no plain text, marcando os novos chars
       como `inserido=True` e preservando o rpr do char ORIGINAL na posicao.
    3. Reagrupa por (rpr, inserido) — chars consecutivos com mesma combinacao
       viram 1 run.
    4. Em runs com `inserido=True`, adiciona <w:highlight val="yellow"/>.

    Args:
        p_elem: elemento <w:p>
        mapa: dict {string_a_buscar: string_substituta}
        grifo: True para aplicar amarelo nos substituidos

    Returns:
        True se algum substituicao foi feita.
    """
    plain_chars = []
    for r in p_elem.findall(".//" + W + "r"):
        rpr = r.find(W + "rPr")
        for t in r.findall(W + "t"):
            for ch in (t.text or ""):
                plain_chars.append([ch, rpr, False])
    plain = "".join(c[0] for c in plain_chars)
    if not any(k in plain for k in mapa):
        return False

    for k, v in mapa.items():
        while True:
            indices_originais = [idx for idx, c in enumerate(plain_chars) if not c[2]]
            atual = "".join(plain_chars[idx][0] for idx in indices_originais)
            i_atual = atual.find(k)
            if i_atual < 0:
                break
            i_real_ini = indices_originais[i_atual]
            i_real_fim = indices_originais[i_atual + len(k) - 1] + 1
            rpr_origem = plain_chars[i_real_ini][1] if i_real_ini < len(plain_chars) else None
            del plain_chars[i_real_ini:i_real_fim]
            novos = [[ch, rpr_origem, True] for ch in v]
            plain_chars[i_real_ini:i_real_ini] = novos

    for child in list(p_elem):
        if child.tag != W + "pPr":
            p_elem.remove(child)

    if not plain_chars:
        return True

    grupos = []
    grupo_chars = [plain_chars[0][0]]
    grupo_rpr = plain_chars[0][1]
    grupo_inserido = plain_chars[0][2]
    for ch, rpr, inserido in plain_chars[1:]:
        if rpr is grupo_rpr and inserido == grupo_inserido:
            grupo_chars.append(ch)
        else:
            grupos.append(("".join(grupo_chars), grupo_rpr, grupo_inserido))
            grupo_chars = [ch]
            grupo_rpr = rpr
            grupo_inserido = inserido
    grupos.append(("".join(grupo_chars), grupo_rpr, grupo_inserido))

    for txt, rpr, inserido in grupos:
        if not txt:
            continue
        r = etree.SubElement(p_elem, W + "r")
        if rpr is not None:
            rpr_novo = copy.deepcopy(rpr)
            r.append(rpr_novo)
        else:
            rpr_novo = None
        if grifo and inserido:
            if rpr_novo is None:
                rpr_novo = etree.SubElement(r, W + "rPr")
                r.remove(rpr_novo)
                r.insert(0, rpr_novo)
            existing_highlight = rpr_novo.find(W + "highlight")
            if existing_highlight is None:
                hl = etree.SubElement(rpr_novo, W + "highlight")
                hl.set(W + "val", "yellow")
            else:
                existing_highlight.set(W + "val", "yellow")
        t = etree.SubElement(r, W + "t")
        t.text = txt
        t.set(XMLSPC, "preserve")
    return True


# ============================================================
#   PRIMITIVAS DE RUN — para quali_banco (2 runs em 12pt)
# ============================================================

def _set_sz(rpr, tamanho_pt):
    """Aplica tamanho de fonte em pt ao rPr. OOXML usa meio-pontos (12pt = 24)."""
    valor = str(int(tamanho_pt * 2))
    for tag in ("w:sz", "w:szCs"):
        existente = rpr.find(qn(tag))
        if existente is not None:
            rpr.remove(existente)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), valor)
    rpr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), valor)
    rpr.append(szCs)


def make_run_cambria_grifado(texto, preserve=True, tamanho_pt=12):
    """Run Cambria com grifo amarelo e tamanho fixo (default 12pt)."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Cambria")
    rfonts.set(qn("w:hAnsi"), "Cambria")
    rpr.append(rfonts)
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    rpr.append(hl)
    if tamanho_pt is not None:
        _set_sz(rpr, tamanho_pt)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = texto
    if preserve:
        t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def make_run_segoe_bold_grifado(texto, preserve=False, tamanho_pt=12):
    """Run Segoe UI Bold (rStyle 2TtuloChar) com grifo amarelo e 12pt."""
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rs = OxmlElement("w:rStyle")
    rs.set(qn("w:val"), "2TtuloChar")
    rpr.append(rs)
    hl = OxmlElement("w:highlight")
    hl.set(qn("w:val"), "yellow")
    rpr.append(hl)
    if tamanho_pt is not None:
        _set_sz(rpr, tamanho_pt)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = texto
    if preserve:
        t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def iter_all_paragraphs(doc):
    """Itera por TODOS os paragrafos do doc — corpo + tabelas."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
