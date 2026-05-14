# -*- coding: utf-8 -*-
"""
Geração dos DOCX de saída: Relatório de Conferência + Edições Sugeridas.

Padrão visual do escritório (Cambria 12pt, Segoe UI Semibold dourado para
títulos, Sitka Text para citações). Tenta usar template; se não disponível,
cai para fallback formatando manualmente via paragraph_format.
"""
from __future__ import annotations

import os
import re
from datetime import date, datetime
from typing import Dict, List, Optional, Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor


# Caminhos candidatos a template (cai no primeiro que existir)
CANDIDATOS_TEMPLATE = [
    r"C:\Users\gabri\OneDrive\Área de Trabalho\Petição desistência - Contrato Digital - 5001065-32.2025.4.04.7206.docx",
    r"C:\Users\gabri\OneDrive\Área de Trabalho\Petição desistência - Contrato Digital - 5004053-26.2025.4.04.7206.docx",
    r"C:\Users\gabri\OneDrive\Área de Trabalho\Petição desistência - Contrato Digital - 5004298-37.2025.4.04.7206.docx",
]

ESTILOS_NECESSARIOS = [
    "1. Parágrafo",
    "2. Título",
    "3. Subtítulo",
    "3.1 Subtítulo intermediário",
    "4. Citação",
]


def _abrir_template():
    """Tenta abrir um template do escritório, limpando o body. Cai para Document() em branco."""
    for tpl in CANDIDATOS_TEMPLATE:
        if os.path.exists(tpl):
            try:
                doc = Document(tpl)
                styles = [s.name for s in doc.styles]
                if all(e in styles for e in ESTILOS_NECESSARIOS):
                    body = doc.element.body
                    for child in list(body):
                        tag = child.tag
                        if tag.endswith("}p") or tag.endswith("}tbl"):
                            body.remove(child)
                    return doc, True
            except Exception:
                continue
    return Document(), False


def _set_cell_shading(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _aplicar_fallback_paragrafo(p, tamanho_pt=12):
    pf = p.paragraph_format
    pf.first_line_indent = Cm(1)
    pf.space_before = Pt(6)
    pf.line_spacing = 1.2
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Cambria"
        run.font.size = Pt(tamanho_pt)


def _add_paragraph_corrido(doc, texto: str, tem_template: bool, italico=False, negrito=False):
    if tem_template:
        try:
            p = doc.add_paragraph(style="1. Parágrafo")
        except KeyError:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Cambria"
    run.font.size = Pt(12)
    if italico:
        run.italic = True
    if negrito:
        run.bold = True
    if not tem_template:
        _aplicar_fallback_paragrafo(p)
    return p


def _add_titulo(doc, texto: str, tem_template: bool):
    if tem_template:
        try:
            p = doc.add_paragraph(style="2. Título")
        except KeyError:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Segoe UI"
    run.font.size = Pt(14)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def _add_subtitulo(doc, texto: str, tem_template: bool, intermediario=False):
    nome_estilo = "3.1 Subtítulo intermediário" if intermediario else "3. Subtítulo"
    if tem_template:
        try:
            p = doc.add_paragraph(style=nome_estilo)
        except KeyError:
            p = doc.add_paragraph()
    else:
        p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = "Segoe UI Semibold"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0xB3, 0x82, 0x4C)
    return p


def _add_citacao(doc, texto: str, tem_template: bool):
    if tem_template:
        try:
            p = doc.add_paragraph(style="4. Citação")
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(3)
    else:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3)
    run = p.add_run(texto)
    run.font.name = "Sitka Text"
    run.font.size = Pt(11)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _add_tabela(doc, cabecalho: List[str], linhas: List[List[str]]):
    tbl = doc.add_table(rows=1 + len(linhas), cols=len(cabecalho))
    tbl.style = "Table Grid"
    # Cabeçalho
    hdr = tbl.rows[0]
    for i, h in enumerate(cabecalho):
        cell = hdr.cells[i]
        _set_cell_shading(cell, "D9E1F2")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Cambria"
        run.font.size = Pt(10)
        run.bold = True
    # Linhas
    for r_idx, linha in enumerate(linhas, start=1):
        for c_idx, val in enumerate(linha):
            if c_idx >= len(cabecalho):
                break
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Cambria"
            run.font.size = Pt(10)
    return tbl


def _semaforo(status: str) -> str:
    s = (status or "").upper()
    if s in ("OK", "VERDE", "PRESENTE", "CONFERE"):
        return "🟢"
    if s in ("ALERTA", "AMARELO", "RESSALVA", "PARCIAL"):
        return "🟡"
    if s in ("INCONSISTENTE", "VERMELHO", "AUSENTE", "DIVERGENTE"):
        return "🔴"
    return "⚪"


# ----------------------------------------------------------------------- #
# Relatório
# ----------------------------------------------------------------------- #
def gerar_relatorio(
    pasta_saida: str,
    nome_cliente: str,
    contexto: Dict[str, Any],
) -> str:
    """Gera o DOCX do Relatório de Conferência.

    contexto deve ter:
      cabecalho: dict com (titulo, cliente, cpf, tipo_acao_pretty, comarca,
                            subscritor, oab, pasta, data_conferencia)
      semaforo: lista de dicts {eixo, status, observacao}
      alertas: lista de strings (alertas destacados)
      tipo_acao: dict (resultado de tipo_acao.detectar_tipo)
      cruzamentos: lista de dicts {dado, fonte_a, fonte_b, status, observacao}
      notificacao: dict (resultado de notificacao_check)
      oab: dict (resultado de oab_check)
      adaptacao: lista de alertas (modelo + placeholders)
      docs: lista de dicts {nome, caminho, observacao}
      ausencias: lista de strings
      sintese: dict com (resultado, total_edicoes, criticas, medias, baixas, resumo)

    Retorna o caminho do arquivo gerado.
    """
    os.makedirs(pasta_saida, exist_ok=True)
    nome_arq = "Relatorio_Conferencia_Inicial_" + _slug(nome_cliente) + ".docx"
    caminho = os.path.join(pasta_saida, nome_arq)

    doc, tem_template = _abrir_template()

    cab = contexto.get("cabecalho", {})
    _add_titulo(doc, "RELATÓRIO DE CONFERÊNCIA - INICIAL CONTRA BRADESCO", tem_template)
    _add_paragraph_corrido(doc, f"Cliente: {cab.get('cliente', '')}", tem_template)
    _add_paragraph_corrido(doc, f"CPF: {cab.get('cpf', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Tipo de ação: {cab.get('tipo_acao_pretty', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Comarca: {cab.get('comarca', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Subscritor: {cab.get('subscritor', '')} (OAB/{cab.get('oab_uf', '')} {cab.get('oab_numero', '')})", tem_template)
    _add_paragraph_corrido(doc, f"Pasta analisada: {cab.get('pasta', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Data da conferência: {cab.get('data_conferencia', '')}", tem_template)

    # Seção 1: Tabela Semáforo
    _add_subtitulo(doc, "1. Tabela Semáforo", tem_template)
    cabec_sem = ["#", "Eixo", "Status", "Observação"]
    linhas_sem = []
    for i, item in enumerate(contexto.get("semaforo", []), start=1):
        sem = _semaforo(item.get("status", ""))
        linhas_sem.append([str(i), item.get("eixo", ""), sem, item.get("observacao", "")])
    if linhas_sem:
        _add_tabela(doc, cabec_sem, linhas_sem)
    else:
        _add_paragraph_corrido(doc, "(sem itens)", tem_template, italico=True)

    # Seção 2: Alertas Destacados
    _add_subtitulo(doc, "2. Alertas Destacados", tem_template)
    alertas = contexto.get("alertas", [])
    if alertas:
        for a in alertas:
            _add_citacao(doc, a, tem_template)
    else:
        _add_paragraph_corrido(doc, "Nenhum alerta crítico identificado.", tem_template, italico=True)

    # Seção 3: Identificação do Tipo de Ação
    _add_subtitulo(doc, "3. Identificação do Tipo de Ação", tem_template)
    ti = contexto.get("tipo_acao", {})
    _add_paragraph_corrido(doc, f"Tipos detectados: {', '.join(ti.get('tipos_detectados', [])) or '[nenhum]'}", tem_template)
    _add_paragraph_corrido(doc, f"Subpasta de origem: {ti.get('fonte_subpasta', '') or '[direto na raiz]'}", tem_template)
    _add_paragraph_corrido(doc, f"Rubrica detectada na inicial: {ti.get('rubrica_no_texto', '') or '[não identificada]'}", tem_template)
    _add_paragraph_corrido(doc, f"IRDR/Tese citado na inicial: {ti.get('irdr_no_texto', '') or '[nenhum]'}", tem_template)
    _add_paragraph_corrido(doc, f"IRDR/Tese esperado: {', '.join(ti.get('irdr_esperado', []) or [])}", tem_template)
    _add_paragraph_corrido(doc, f"Consistência: {ti.get('consistencia', '')}", tem_template)

    # Seção 4: Cruzamento de Dados
    _add_subtitulo(doc, "4. Cruzamento de Dados Fáticos", tem_template)
    cruzamentos = contexto.get("cruzamentos", [])
    if cruzamentos:
        cabec_cz = ["Dado", "Inicial", "Documento de origem", "Status", "Observação"]
        linhas_cz = []
        for cz in cruzamentos:
            linhas_cz.append([
                cz.get("dado", ""),
                str(cz.get("inicial", "") or ""),
                str(cz.get("fonte", "") or ""),
                _semaforo(cz.get("status", "")),
                cz.get("observacao", "") or "",
            ])
        _add_tabela(doc, cabec_cz, linhas_cz)
    else:
        _add_paragraph_corrido(doc, "(sem cruzamentos registrados)", tem_template, italico=True)

    # Seção 5: Notificação Extrajudicial
    _add_subtitulo(doc, "5. Notificação Extrajudicial", tem_template)
    notif = contexto.get("notificacao", {})
    if notif:
        cabec_nt = ["Item", "Esperado", "Encontrado", "Status"]
        linhas_nt = [
            ["Notificação juntada", "sim", "sim" if notif.get("tem_notificacao") else "não", "🟢" if notif.get("tem_notificacao") else "🔴"],
            ["AR juntado", "sim", "sim" if notif.get("tem_ar") else "não", "🟢" if notif.get("tem_ar") else "🔴"],
            ["Tipo correto", str(notif.get("tipos_esperados", "")), str(notif.get("tipos_detectados", "") or "[indeterminado]"), _semaforo(notif.get("status", ""))],
            ["Data de envio", "≥ 15 dias antes da inicial", str(notif.get("data_envio", "") or "[não detectado]"), "🟡"],
            ["Destinatário Bradesco", "sim", "sim" if notif.get("destinatario_bradesco") else "não", "🟢" if notif.get("destinatario_bradesco") else "🟡"],
        ]
        _add_tabela(doc, cabec_nt, linhas_nt)
    else:
        _add_paragraph_corrido(doc, "(sem informações de notificação)", tem_template, italico=True)

    # Seção 6: OAB / Template
    _add_subtitulo(doc, "6. OAB / Template", tem_template)
    oab = contexto.get("oab", {})
    if oab:
        _add_paragraph_corrido(doc, f"Subscritor: {oab.get('advogado', '') or '[não identificado]'}", tem_template)
        _add_paragraph_corrido(doc, f"OAB: {oab.get('oab_uf', '')}/{oab.get('oab_numero', '')}", tem_template)
        _add_paragraph_corrido(doc, f"Template aplicável: {oab.get('template', '')}", tem_template)
        _add_paragraph_corrido(doc, f"Mensagem: {oab.get('mensagem', '')}", tem_template)
    else:
        _add_paragraph_corrido(doc, "(OAB não verificada)", tem_template, italico=True)

    # Seção 7: Adaptação do modelo
    _add_subtitulo(doc, "7. Adaptação do Modelo", tem_template)
    adapt = contexto.get("adaptacao", [])
    if adapt:
        cabec_ad = ["§", "Severidade", "Tipo", "Trecho", "Observação"]
        linhas_ad = []
        for a in adapt:
            linhas_ad.append([
                str(a.get("paragrafo", "")),
                a.get("severidade", ""),
                a.get("tipo") or a.get("padrao", ""),
                (a.get("trecho", "") or "")[:120],
                a.get("mensagem", "") or "",
            ])
        _add_tabela(doc, cabec_ad, linhas_ad)
    else:
        _add_paragraph_corrido(doc, "Nenhum alerta de adaptação de modelo.", tem_template, italico=True)

    # Seção 8: Documentos
    _add_subtitulo(doc, "8. Documentos Analisados e Ausências", tem_template)
    _add_subtitulo(doc, "8.1 Documentos utilizados", tem_template, intermediario=True)
    docs = contexto.get("docs", [])
    if docs:
        cabec_dc = ["#", "Documento", "Caminho", "Observação"]
        linhas_dc = []
        for i, d in enumerate(docs, start=1):
            linhas_dc.append([str(i), d.get("nome", ""), d.get("caminho", ""), d.get("observacao", "") or ""])
        _add_tabela(doc, cabec_dc, linhas_dc)

    _add_subtitulo(doc, "8.2 Documentos ausentes", tem_template, intermediario=True)
    aus = contexto.get("ausencias", [])
    if aus:
        for a in aus:
            _add_paragraph_corrido(doc, f"• {a}", tem_template)
    else:
        _add_paragraph_corrido(doc, "Nenhuma ausência relevante.", tem_template, italico=True)

    # Seção 9: Síntese
    _add_subtitulo(doc, "9. Síntese", tem_template)
    sint = contexto.get("sintese", {})
    _add_paragraph_corrido(doc, f"Resultado: {sint.get('resultado', '')}", tem_template, negrito=True)
    _add_paragraph_corrido(doc, f"Total de edições sugeridas: {sint.get('total_edicoes', 0)}", tem_template)
    _add_paragraph_corrido(doc, f"Edições críticas (🔴): {sint.get('criticas', 0)}", tem_template)
    _add_paragraph_corrido(doc, f"Edições médias (⚠️): {sint.get('medias', 0)}", tem_template)
    _add_paragraph_corrido(doc, f"Edições baixas (🟡): {sint.get('baixas', 0)}", tem_template)
    if sint.get("resumo"):
        _add_paragraph_corrido(doc, sint["resumo"], tem_template, italico=True)

    doc.save(caminho)
    return caminho


# ----------------------------------------------------------------------- #
# Edições Sugeridas
# ----------------------------------------------------------------------- #
def gerar_edicoes(
    pasta_saida: str,
    nome_cliente: str,
    cabecalho: Dict[str, Any],
    edicoes: List[Dict[str, Any]],
) -> str:
    """Gera o DOCX de Edições Sugeridas.

    Cada edição: dict com campos:
      tipo_acao (SUBSTITUIR/INSERIR ANTES/...), gravidade, eixo,
      ancoragem (texto literal), trecho_original, texto_substituto,
      destino, justificativa.
    """
    os.makedirs(pasta_saida, exist_ok=True)
    nome_arq = "Edicoes_Sugeridas_Inicial_" + _slug(nome_cliente) + ".docx"
    caminho = os.path.join(pasta_saida, nome_arq)

    doc, tem_template = _abrir_template()

    _add_titulo(doc, "EDIÇÕES SUGERIDAS À INICIAL", tem_template)
    _add_paragraph_corrido(doc, f"Cliente: {cabecalho.get('cliente', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Tipo de ação: {cabecalho.get('tipo_acao_pretty', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Subscritor: {cabecalho.get('subscritor', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Data: {cabecalho.get('data_conferencia', '')}", tem_template)
    _add_paragraph_corrido(doc, f"Total de edições: {len(edicoes)}", tem_template)

    _add_paragraph_corrido(
        doc,
        "Cada edição é ancorada pelo trecho final do parágrafo anterior, para que o "
        "responsável pela aplicação localize com facilidade. As edições estão "
        "organizadas por ordem de aparição na inicial.",
        tem_template, italico=True,
    )

    if not edicoes:
        _add_paragraph_corrido(
            doc,
            "INICIAL APTA AO PROTOCOLO. Nenhuma edição sugerida pela conferência automatizada. "
            "Recomenda-se revisão manual final antes do protocolo.",
            tem_template, negrito=True,
        )
        doc.save(caminho)
        return caminho

    # Tabela-resumo
    _add_subtitulo(doc, "Tabela-resumo das Edições", tem_template)
    cabec_tr = ["#", "Tipo", "Gravidade", "Eixo afetado", "Aplicada?"]
    linhas_tr = []
    for i, e in enumerate(edicoes, start=1):
        linhas_tr.append([
            str(i),
            e.get("tipo_acao", ""),
            e.get("gravidade", ""),
            e.get("eixo", ""),
            "☐",
        ])
    _add_tabela(doc, cabec_tr, linhas_tr)

    # Bloco de cada edição
    for i, e in enumerate(edicoes, start=1):
        _add_subtitulo(
            doc,
            f"Edição #{i} - {e.get('tipo_acao', '')} - {e.get('gravidade', '')}",
            tem_template, intermediario=True,
        )
        _add_paragraph_corrido(doc, f"Eixo afetado: {e.get('eixo', '')}", tem_template)
        if e.get("ancoragem"):
            _add_paragraph_corrido(doc, "Ancoragem (parágrafo anterior termina com):", tem_template, negrito=True)
            _add_citacao(doc, e["ancoragem"], tem_template)
        _add_paragraph_corrido(doc, f"Ação: {e.get('tipo_acao', '')}", tem_template, negrito=True)
        if e.get("trecho_original"):
            _add_paragraph_corrido(doc, "Trecho original na peça:", tem_template, negrito=True)
            _add_citacao(doc, e["trecho_original"], tem_template)
        if e.get("texto_substituto"):
            _add_paragraph_corrido(doc, "Texto substituto / novo texto:", tem_template, negrito=True)
            _add_citacao(doc, e["texto_substituto"], tem_template)
        if e.get("destino"):
            _add_paragraph_corrido(doc, f"Destino: {e['destino']}", tem_template)
        if e.get("justificativa"):
            _add_paragraph_corrido(doc, f"Justificativa: {e['justificativa']}", tem_template, italico=True)

    doc.save(caminho)
    return caminho


def _slug(s: str) -> str:
    s = re.sub(r"[^A-Za-zÀ-ÿ0-9_\-\s]+", "", s or "")
    s = re.sub(r"\s+", "_", s).strip("_")
    return s[:80] or "cliente"


# ----------------------------------------------------------------------- #
# Função principal de conveniência
# ----------------------------------------------------------------------- #
def gerar_relatorio_e_edicoes(
    pasta_saida: str, nome_cliente: str,
    contexto_relatorio: Dict[str, Any],
    edicoes: List[Dict[str, Any]],
    cabecalho_edicoes: Optional[Dict[str, Any]] = None,
) -> Dict[str, str]:
    """Gera ambos os DOCX e devolve dict com os dois caminhos."""
    rel = gerar_relatorio(pasta_saida, nome_cliente, contexto_relatorio)
    cab = cabecalho_edicoes or contexto_relatorio.get("cabecalho", {})
    cab.setdefault("data_conferencia", contexto_relatorio.get("cabecalho", {}).get("data_conferencia", ""))
    edi = gerar_edicoes(pasta_saida, nome_cliente, cab, edicoes)
    return {"relatorio": rel, "edicoes": edi}


if __name__ == "__main__":
    # Smoke test
    ctx = {
        "cabecalho": {
            "cliente": "JOAO TESTE", "cpf": "000.000.000-00",
            "tipo_acao_pretty": "Mora Cred Pess",
            "comarca": "Maués/AM",
            "subscritor": "Eduardo Rebonatto", "oab_uf": "AM", "oab_numero": "A2118",
            "pasta": ".", "data_conferencia": "25/04/2026",
        },
        "semaforo": [
            {"eixo": "1. Tipo de ação", "status": "OK", "observacao": "MORA CRED PESS detectado"},
            {"eixo": "2. Identidade", "status": "ALERTA", "observacao": "Idade não confirmada"},
        ],
        "alertas": ["Datas invertidas no período da inicial."],
        "tipo_acao": {"tipos_detectados": ["MORA_CRED_PESS"], "consistencia": "OK", "irdr_esperado": ["0004464-79.2023.8.04.0000"]},
        "cruzamentos": [
            {"dado": "Conta", "inicial": "12345-6", "fonte": "Extrato", "status": "OK", "observacao": ""},
        ],
        "notificacao": {"tem_notificacao": True, "tem_ar": True, "tipos_detectados": ["AM - Encargos"], "status": "OK"},
        "oab": {"advogado": "Eduardo Rebonatto", "oab_uf": "AM", "oab_numero": "A2118", "template": 1, "mensagem": "OK"},
        "adaptacao": [],
        "docs": [{"nome": "Inicial.docx", "caminho": "./inicial.docx", "observacao": "OK"}],
        "ausencias": [],
        "sintese": {"resultado": "✅ APTA AO PROTOCOLO", "total_edicoes": 1, "criticas": 1, "medias": 0, "baixas": 0, "resumo": "Smoke test OK."},
    }
    edicoes = [{
        "tipo_acao": "SUBSTITUIR", "gravidade": "🔴", "eixo": "Período",
        "ancoragem": "Ancoragem teste",
        "trecho_original": "07/01/2026 a 07/11/2025",
        "texto_substituto": "07/11/2025 a 07/01/2026",
        "justificativa": "Datas invertidas.",
    }]
    saida = gerar_relatorio_e_edicoes(".", "JoaoTeste", ctx, edicoes)
    print(saida)
