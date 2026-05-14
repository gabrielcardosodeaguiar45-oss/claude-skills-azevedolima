"""
Gera DOCX detalhado da análise de cadeias HISCON.
Padrão visual: Cambria 11 no corpo, Segoe UI Semibold em títulos, cor #B3824C.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from collections import defaultdict, Counter
import io

ACENTO = RGBColor(0xB3, 0x82, 0x4C)
PRETO = RGBColor(0, 0, 0)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA = RGBColor(0x80, 0x80, 0x80)
VERMELHO = RGBColor(0xB0, 0x2A, 0x2A)
AMARELO_BG = 'FFF4DC'
VERMELHO_BG = 'FBE4E4'
CAMBRIA = 'Cambria'
SEGOE = 'Segoe UI'
SEGOE_SB = 'Segoe UI Semibold'

def _set_cell_shading(cell, hex_color):
    tc = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hex_color)
    tc.append(shd)

def _set_cell_border(cell):
    tc = cell._tc.get_or_add_tcPr()
    tcb = OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{e}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'CCCCCC')
        tcb.append(el)
    tc.append(tcb)

def _run_font(run, font=CAMBRIA, size=11, bold=False, color=None):
    run.font.name = font; run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    for a in ('ascii','hAnsi','cs'): rFonts.set(qn(f'w:{a}'), font)

def _fmt_brl(v):
    if not v: return 'R$ 0,00'
    return f"R$ {v:,.2f}".replace(',','X').replace('.',',').replace('X','.')

def _fmt_date(s):
    if not s: return '—'
    if isinstance(s, str):
        try: d = datetime.fromisoformat(s)
        except: return '—'
    else: d = s
    return d.strftime('%d/%m/%Y')

def _parse_dt(s):
    if not s: return None
    if isinstance(s, str):
        try: return datetime.fromisoformat(s)
        except: return None
    return s

def gerar_docx(resultado: dict, destino: str = None) -> bytes:
    """Gera DOCX a partir do resultado do analisar_hiscon. Retorna bytes (e opcionalmente salva)."""
    doc = Document()
    for s in doc.sections:
        s.top_margin=Cm(2.5); s.bottom_margin=Cm(2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)
    style = doc.styles['Normal']; style.font.name = CAMBRIA; style.font.size = Pt(11)

    def titulo(t):
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6)
        r = p.add_run(t); _run_font(r, SEGOE, 16, bold=True, color=PRETO); r.font.all_caps=True

    def subt(t):
        p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(20)
        r = p.add_run(t); _run_font(r, SEGOE_SB, 11, color=ACENTO)

    def heading(t, size=12, sb=18):
        p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(8)
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
        b = OxmlElement('w:bottom')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:color'),'B3824C'); b.set(qn('w:space'),'1')
        pBdr.append(b); pPr.append(pBdr)
        r = p.add_run(t); _run_font(r, SEGOE_SB, size, bold=True, color=ACENTO)

    def subh(t, size=11):
        p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(4)
        r = p.add_run(t); _run_font(r, SEGOE_SB, size, bold=True, color=PRETO)

    def texto(t, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, color=None):
        p = doc.add_paragraph(); p.alignment=align
        p.paragraph_format.first_line_indent = Cm(1) if indent else Cm(0)
        p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.2
        r = p.add_run(t); _run_font(r, CAMBRIA, 11, color=color)

    def box_aviso(titulo_b, texto_b, cor_bg='FFF4DC', cor_borda='D97706', cor_txt=RGBColor(0x92,0x40,0x0E)):
        t = doc.add_table(rows=1, cols=1); c = t.rows[0].cells[0]
        _set_cell_shading(c, cor_bg)
        tc = c._tc.get_or_add_tcPr(); tcb = OxmlElement('w:tcBorders')
        for e in ('top','left','bottom','right'):
            el = OxmlElement(f'w:{e}')
            el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'12'); el.set(qn('w:color'),cor_borda)
            tcb.append(el)
        tc.append(tcb)
        p = c.paragraphs[0]; p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
        r = p.add_run(f'⚠ {titulo_b}'); _run_font(r, SEGOE, 10, bold=True, color=cor_txt)
        p2 = c.add_paragraph(); p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(4)
        r2 = p2.add_run(texto_b); _run_font(r2, CAMBRIA, 10, color=cor_txt)

    def info_tabela(pares):
        t = doc.add_table(rows=len(pares), cols=2); t.autofit=False
        for i, (k, v) in enumerate(pares):
            c1, c2 = t.rows[i].cells
            c1.width = Cm(6); c2.width = Cm(10)
            p1 = c1.paragraphs[0]; r1 = p1.add_run(k); _run_font(r1, CAMBRIA, 10, bold=True)
            p2 = c2.paragraphs[0]; r2 = p2.add_run(str(v)); _run_font(r2, CAMBRIA, 10)
            for c in (c1, c2):
                _set_cell_border(c)
                c.paragraphs[0].paragraph_format.space_before=Pt(2)
                c.paragraphs[0].paragraph_format.space_after=Pt(2)

    def tabela_contratos(conts):
        headers = ['#','Banco','Sit.','Origem','Inclusão','Exclusão','Parc.','Parcela','Emprestado','Pago','Juros m.','CET m.']
        t = doc.add_table(rows=1+len(conts), cols=len(headers)); t.autofit=False
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            _set_cell_shading(cell, '0B1120'); _set_cell_border(cell)
            p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h); _run_font(r, SEGOE, 8, bold=True, color=BRANCO)
        for ri, c in enumerate(conts, 1):
            marca = {'original':'ORIG','refinanciamento':'REFIN','portabilidade':'PORT','migracao':'MIGR'}.get(c['tipo_origem'],'?')
            ativo_ou_susp = c['situacao'] in ('Ativo','Suspenso')
            vals = [
                (c['numero'][:18], False),
                (f"{c['banco_codigo']} {c['banco_nome'][:22]}", False),
                (c['situacao'][:9], False),
                (marca, False),
                (_fmt_date(c['data_inclusao']), False),
                (_fmt_date(c['data_exclusao']), False),
                (str(c['qtd_parcelas']) if c['qtd_parcelas'] else '—', c['qtd_parcelas']==0),
                (_fmt_brl(c['valor_parcela']) if c['valor_parcela'] else '—', c['valor_parcela']==0),
                (_fmt_brl(c['valor_emprestado']) if c['valor_emprestado'] else '⚠ —', c['valor_emprestado']==0),
                (_fmt_brl(c['valor_pago']) if c['valor_pago'] else '—', c['valor_pago']==0 and c['data_exclusao']),
                (f"{c['juros_mensal']:.2f}%".replace('.',',') if c['juros_mensal'] else '⚠ —', c['juros_mensal']==0 and ativo_ou_susp),
                (f"{c['cet_mensal']:.2f}%".replace('.',',') if c['cet_mensal'] else '⚠ —', c['cet_mensal']==0 and ativo_ou_susp),
            ]
            for i, (v, flagged) in enumerate(vals):
                cell = t.rows[ri].cells[i]; _set_cell_border(cell)
                if flagged: _set_cell_shading(cell, AMARELO_BG)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 6 else WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(v); _run_font(r, CAMBRIA, 8)
                p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)

    def arvore(comp, nodes_d, edges_d, rev_edges_d, ligacoes_l):
        rs = sorted([n for n in comp if not any(p in comp for p in rev_edges_d.get(n,[]))],
                    key=lambda n: _parse_dt(nodes_d[n].get('data_inclusao')) or datetime.max)
        linhas = []
        def rec(node, prefixo, is_last, loc):
            if node in loc: return
            loc.add(node)
            c = nodes_d[node]
            ent = _parse_dt(c.get('data_inclusao'))
            sai = _parse_dt(c.get('data_exclusao'))
            marca = {'original':'ORIG','refinanciamento':'REFIN','portabilidade':'PORTAB','migracao':'MIGR'}.get(c['tipo_origem'],'?')
            periodo = f"{_fmt_date(c['data_inclusao'])} → {_fmt_date(c['data_exclusao']) if c['data_exclusao'] else 'ATIVO'}"
            info = f"[{marca}] {c['banco_nome'][:30]} · contr {c['numero']} · {c['qtd_parcelas']}x {_fmt_brl(c['valor_parcela'])} · emprest {_fmt_brl(c['valor_emprestado'])} · {periodo}"
            confianca = 'alta'
            for pred_id in rev_edges_d.get(node, []):
                if pred_id in comp:
                    for l in ligacoes_l:
                        if l['suc'] == node and l['pred'] == pred_id:
                            confianca = l['confianca']; break
            sinal = f' ⚠({confianca})' if confianca=='baixa' else ''
            conector = '└─ ' if is_last else '├─ '
            linhas.append((prefixo + conector + info + sinal, confianca))
            filhos = [f for f in edges_d.get(node,[]) if f in comp and f not in loc]
            for i, f in enumerate(filhos):
                rec(f, prefixo + ('   ' if is_last else '│  '), i==len(filhos)-1, loc)
        for i, r in enumerate(rs):
            rec(r, '', i==len(rs)-1, set())

        p = doc.add_paragraph()
        p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(8)
        for i, (ln, conf) in enumerate(linhas):
            if i > 0:
                r = p.add_run('\n'); _run_font(r, 'Consolas', 9)
            cor = VERMELHO if conf=='baixa' else (CINZA if conf=='média' else PRETO)
            r = p.add_run(ln); _run_font(r, 'Consolas', 9, color=cor)

    def red_flags(flags):
        if not flags:
            texto('Nenhum indício específico detectado nesta cadeia.', indent=False)
            return
        for f in flags:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent=Cm(0.5); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(2)
            r1 = p.add_run('● '); _run_font(r1, CAMBRIA, 11, color=ACENTO, bold=True)
            r2 = p.add_run(f['tipo'] + ': '); _run_font(r2, SEGOE_SB, 11, bold=True, color=PRETO)
            r3 = p.add_run(f['desc']); _run_font(r3, CAMBRIA, 11)

    # ============ CONTEÚDO ============
    benef = resultado['beneficiario']
    contratos = resultado['contratos']
    cadeias = resultado['cadeias']
    edges_d = defaultdict(list, resultado.get('edges', {}))
    rev_edges_d = defaultdict(list, resultado.get('rev_edges', {}))
    nodes_d = {f"{c['numero']}@{c['banco_codigo']}": c for c in contratos}
    ligacoes = resultado['ligacoes']
    avisos = resultado['avisos']
    problemas = resultado['problemas']
    red_flags_d = resultado['red_flags']
    estat = resultado['estatisticas']

    titulo('Análise de Cadeias de Empréstimo Consignado')
    subt('Histórico do Empréstimo Consignado (HISCON) — INSS')

    if avisos or problemas:
        heading('⚠ Avisos de leitura', size=11, sb=6)
        texto('Alertas gerados automaticamente. Itens em AMARELO na tabela indicam campos não disponibilizados pelo INSS (não são erros do relatório).', indent=False)
        by_tipo = defaultdict(list)
        for a in avisos: by_tipo[a['tipo']].append(a)
        nomes = {
            'campo_faltante': 'Campos ausentes no HISCON',
            'origem_desconhecida': 'Origem de averbação não classificada',
            'banco_normalizado': 'Nomes de banco reconstruídos',
            'linha_fragmentada': 'Linhas fragmentadas descartadas',
            'orfao_sem_sucessor': 'Contratos excluídos sem sucessor identificado',
        }
        for tipo, lista in by_tipo.items():
            subh(nomes.get(tipo, tipo), size=10)
            if tipo == 'orfao_sem_sucessor':
                box_aviso(
                    f'{len(lista)} contrato(s) excluído(s) sem sucessor localizado',
                    'Pode indicar: (a) portabilidade para crédito fora do consignado INSS; (b) troca de titularidade; (c) limitação do pareamento. Verificar manualmente antes de elaborar peça.',
                    cor_bg=VERMELHO_BG, cor_borda='B02A2A', cor_txt=VERMELHO
                )
                for a in lista: texto(f'• {a["texto"]}', indent=False)
            elif tipo == 'campo_faltante':
                box_aviso('Campos ausentes no próprio HISCON (não é erro do relatório)',
                          lista[0]['texto'] + ' Esses campos aparecem como "⚠ —" em amarelo nas tabelas.')
            else:
                for a in lista[:10]:
                    pref = f'[{a["contrato"]}] ' if a.get('contrato') else ''
                    texto(f'• {pref}{a["texto"]}', indent=False)
                if len(lista) > 10: texto(f'• (+ {len(lista)-10} omitidos)', indent=False)
        for p in problemas:
            box_aviso(f'Problema: {p["tipo"]}', p['texto'], cor_bg=VERMELHO_BG, cor_borda='B02A2A', cor_txt=VERMELHO)

    heading('1. Identificação do beneficiário')
    linhas = [
        ('Nome', benef.get('nome','—')),
        ('Benefício', benef.get('beneficio','—')),
        ('Número do benefício', benef.get('numero_beneficio','—')),
        ('Banco pagador', benef.get('banco_pagador','—')),
        ('Data do relatório', datetime.now().strftime('%d/%m/%Y')),
    ]
    if 'base_calculo' in benef:
        linhas += [
            ('Base de cálculo', _fmt_brl(benef['base_calculo'])),
            ('Máximo de comprometimento', _fmt_brl(benef['max_comprometimento'])),
            ('Total comprometido', _fmt_brl(benef['total_comprometido'])),
        ]
    info_tabela(linhas)

    heading('2. Resumo executivo')
    info_tabela([
        ('Total de contratos', str(estat['total_contratos'])),
        ('Cadeias identificadas', f"{estat['total_cadeias']} ({estat['cadeias_multi']} multi + {estat['cadeias_isoladas']} isoladas)"),
        ('Contratos ativos hoje', str(estat['contratos_ativos'])),
        ('Contratos excluídos/encerrados', str(estat['contratos_encerrados'])),
        ('Total liberado (somatório bruto)', _fmt_brl(estat['total_liberado'])),
        ('Total pago em refin./portab.', _fmt_brl(estat['total_pago'])),
        ('Saldo devedor projetado', _fmt_brl(estat['saldo_projetado'])),
        ('Ligações (confiança)', f"{estat['ligacoes_total']} total · {estat['ligacoes_alta']} alta · {estat['ligacoes_media']} média · {estat['ligacoes_baixa']} baixa"),
    ])

    if estat['red_flags_total']:
        todas = [f for fs in red_flags_d.values() for f in fs]
        heading('3. Indícios consolidados')
        by_t = Counter(f['tipo'] for f in todas)
        texto(f'Foram detectados {len(todas)} indícios potencialmente relevantes.', indent=False)
        info_tabela([(k, str(v)) for k, v in by_t.most_common()])

    heading('4. Cadeias identificadas')
    texto(f'Das {estat["total_contratos"]} linhas extraídas, foram formadas {estat["total_cadeias"]} cadeias. Ligações com ⚠(baixa) foram pareadas apenas por proximidade de data.', indent=False)

    cadeias_mult = [c for c in cadeias if len(c) > 1]
    cadeias_unit = [c for c in cadeias if len(c) == 1]

    for idx, comp in enumerate(cadeias_mult, 1):
        conts = sorted([nodes_d[n] for n in comp], key=lambda c: _parse_dt(c.get('data_inclusao')) or datetime.max)
        ativos = sum(1 for c in conts if 'Ativ' in c['situacao'])
        bancos = sorted(set(c['banco_nome'][:30] for c in conts))
        heading(f'Cadeia {idx}', size=12, sb=20)
        info_tabela([
            ('Contratos', str(len(conts))),
            ('Ativos hoje', str(ativos)),
            ('Início', _fmt_date(conts[0]['data_inclusao'])),
            ('Bancos envolvidos', '; '.join(bancos[:5]) + (' (...)' if len(bancos)>5 else '')),
            ('Total liberado na cadeia', _fmt_brl(sum(c['valor_liberado'] for c in conts))),
            ('Total pago em refin./portab.', _fmt_brl(sum(c['valor_pago'] for c in conts))),
        ])
        subh('Árvore da cadeia')
        arvore(comp, nodes_d, edges_d, rev_edges_d, ligacoes)
        subh('Detalhamento dos contratos')
        tabela_contratos(conts)
        subh('Indícios desta cadeia')
        red_flags(red_flags_d.get(cadeias.index(comp), []))

    if cadeias_unit:
        heading('5. Contratos isolados')
        texto(f'Os {len(cadeias_unit)} contratos abaixo não foram pareados. Geralmente contratos antigos quitados ou originais cuja ligação não pôde ser determinada.', indent=False)
        conts_iso = sorted([nodes_d[list(c)[0]] for c in cadeias_unit],
                            key=lambda c: _parse_dt(c.get('data_inclusao')) or datetime.max)
        tabela_contratos(conts_iso)

    heading('6. Nota metodológica')
    texto('Cadeias montadas por: (a) ligação explícita via "Migrado do contrato" — confiança ALTA; (b) parcela idêntica em janela ±60 dias — ALTA; (c) consolidação N→1 ou fracionamento 1→N por soma — MÉDIA; (d) proximidade temporal com mesmo banco — BAIXA. Indícios apontam padrões, não provas de ilegalidade — confirmar com cliente.', indent=False)
    texto(f'Relatório gerado automaticamente em {datetime.now().strftime("%d/%m/%Y às %H:%M")} — De Azevedo Lima & Rebonatto Advocacia.', indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = buf.read()
    if destino:
        with open(destino, 'wb') as f: f.write(data)
    return data
