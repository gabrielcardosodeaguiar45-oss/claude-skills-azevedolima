"""Pipeline reutilizável para gerar inicial PG ELETRON + relatório paralelo.

Recebe um dict completo com os dados do caso e gera:
- INICIAL_PgEletron_<TERCEIRO>_<NOME>_v1.docx
- _RELATORIO_PENDENCIAS_<TERCEIRO>_v1.docx (ou só _v1 se 1 só tese)

Uso:
    from _pipeline_caso import gerar_inicial_pg_eletron

    res = gerar_inicial_pg_eletron(
        pasta_destino=...,
        nome_arquivo_base='INICIAL_PgEletron_ASPECIR_MARIA_DALVA',
        terceiro_slug='ASPECIR',  # para o relatório paralelo
        dados=DICT_COMPLETO,
        renda_alerta=True,  # se True, sinaliza no relatório
        cobranca_anual=True/False,
        pendencias_extras=[(titulo, texto), ...],
    )
"""
import os, sys, shutil
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from helpers_docx import aplicar_template
from auditor import auditar_docx
from extenso import extenso_moeda, fmt_moeda

from docx import Document

TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-pg-eletron.docx'


def gerar_inicial_pg_eletron(
    pasta_destino: str,
    nome_arquivo_base: str,
    terceiro_slug: str,
    dados: dict,
    valores_legitimos_extra: set = None,
    datas_legitimas_extra: set = None,
    cobranca_anual: bool = False,
    renda_alerta: bool = False,
    estado_civil_omitido: bool = False,
    pendencias_extras: list = None,
):
    """Gera inicial + relatório paralelo. Retorna (caminho_docx, caminho_relatorio, alertas_auditoria)."""
    docx_path = os.path.join(pasta_destino, nome_arquivo_base + '_v1.docx')
    relatorio_path = os.path.join(pasta_destino, f'_RELATORIO_PENDENCIAS_{terceiro_slug}_v1.docx')

    # Aplicar template
    res = aplicar_template(TEMPLATE, dados, docx_path)

    # Auditar
    valores_legit = set(valores_legitimos_extra or set())
    valores_legit |= {
        dados['total_descontos'].replace('R$ ', ''),
        dados['dobro_descontos'].replace('R$ ', ''),
        dados['dano_moral_total'].replace('R$ ', ''),
        dados['valor_causa'].replace('R$ ', ''),
        dados['valor_remuneração'].replace('R$ ', ''),
        '5.000,00',  # jurisprudência citada com frequência
    }
    datas_legit = set(datas_legitimas_extra or set())
    cnpjs_legit = {'60.746.948/0001-12', dados.get('cnpj_terceiro', '')}

    dados_caso = {
        'cpf': dados['cpf'],
        'valores_legitimos': valores_legit,
        'datas_legitimas': datas_legit,
        'nomes_legitimos': {dados['nome_completo'], dados['nome_terceiro'], dados.get('rubrica_curta_caps', '')},
        'cnpjs_legitimos': cnpjs_legit,
        'valor_remuneração': dados['valor_remuneração'].replace('R$ ', ''),
    }
    alertas = auditar_docx(docx_path, dados_caso)

    # Relatório paralelo
    _gerar_relatorio(
        relatorio_path=relatorio_path,
        nome_arquivo_base=nome_arquivo_base,
        terceiro_slug=terceiro_slug,
        dados=dados,
        alertas=alertas,
        cobranca_anual=cobranca_anual,
        renda_alerta=renda_alerta,
        estado_civil_omitido=estado_civil_omitido,
        pendencias_extras=pendencias_extras or [],
    )

    return docx_path, relatorio_path, alertas


def _gerar_relatorio(relatorio_path, nome_arquivo_base, terceiro_slug, dados, alertas,
                     cobranca_anual, renda_alerta, estado_civil_omitido, pendencias_extras):
    doc = Document()
    doc.add_heading(f'RELATÓRIO DE PENDÊNCIAS — {nome_arquivo_base}', level=1)

    cab = [
        ('Cliente', dados['nome_completo']),
        ('Terceiro réu', dados['nome_terceiro']),
        ('Tese', f'PG ELETRON — {dados["rubrica_curta_caps"]}'),
        ('Arquivo gerado', nome_arquivo_base + '_v1.docx'),
    ]
    for k, v in cab:
        p = doc.add_paragraph()
        r = p.add_run(k + ': ')
        r.bold = True
        p.add_run(v)

    doc.add_heading('1. RESUMO DOS DADOS APLICADOS', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    hdr[0].text = 'Campo'
    hdr[1].text = 'Valor'
    linhas_resumo = [
        ('Comarca', f'{dados.get("competência", "")}/{dados.get("uf", "")}'),
        ('Prioridade', dados.get('prioridade_cabecalho', '(nenhuma)')),
        ('Nome', dados['nome_completo']),
        ('CPF', dados['cpf']),
        ('RG', dados.get('rg', '').strip() + ' ' + dados.get('orgao_expedidor_prefixo', '')),
        ('Endereço', f'{dados.get("logradouro", "")}, {dados.get("numero", "")}, {dados.get("bairro", "")}, {dados.get("cidade_de_residencia", "")}/{dados.get("uf", "")}, CEP {dados.get("cep", "")}'),
        ('Conta / Agência', f'{dados.get("conta", "")} / {dados.get("agencia", "")}'),
        ('Renda real (extrato)', dados['valor_remuneração']),
        ('Réu 1', f'BANCO BRADESCO S.A. — CNPJ 60.746.948/0001-12'),
        ('Réu 2 (terceiro)', f'{dados["nome_terceiro"]} — CNPJ {dados.get("cnpj_terceiro", "")}'),
        ('Endereço terceiro', f'{dados.get("logradouro_terceiro", "")}, {dados.get("numero_terceiro", "")}, {dados.get("bairro_terceiro", "")}, {dados.get("cidade_terceiro", "")}/{dados.get("uf_terceiro", "")}, CEP {dados.get("cep_terceiro", "")}'),
        ('Rubrica', dados['rubrica_curta_caps']),
        ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
        ('Total simples', dados['total_descontos']),
        ('Total em dobro', dados['dobro_descontos']),
        ('Dano moral', dados['dano_moral_total']),
        ('Valor da causa', dados['valor_causa']),
    ]
    for k, v in linhas_resumo:
        row = tbl.add_row().cells
        row[0].text = k
        row[1].text = v

    doc.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
    pendencias = []
    if estado_civil_omitido:
        pendencias.append(('Estado civil', 'Não informado em nenhum documento — placeholder OMITIDO limpamente. Confirmar com cliente.'))
    if renda_alerta:
        pendencias.append(('Renda real', 'Foi adotado o último crédito INSS líquido do extrato. Esse valor pode estar reduzido por consignações pelo INSS — a renda BRUTA do benefício pode ser superior. Conferir HISCON se houver.'))
    if cobranca_anual:
        pendencias.append(('Cobrança ANUAL', 'Padrão de 1 lançamento por ano (apólice anual de seguro/previdência). Confirmar com cliente que NUNCA recebeu apólice/proposta/contrato. Ponto que o banco mais explora em contestação.'))
    pendencias.extend(pendencias_extras)

    if pendencias:
        for titulo, txt in pendencias:
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(titulo + ': ')
            r.bold = True
            p.add_run(txt)
    else:
        doc.add_paragraph('Nenhuma pendência crítica identificada.')

    doc.add_heading('3. AUDITORIA AUTOMÁTICA', level=2)
    audit_lines = [
        ('Severidade', alertas.get('severidade', '?')),
        ('Placeholders residuais', ', '.join(alertas.get('placeholders_residuais', [])) or 'NENHUM (0)'),
        ('Valores R$ alheios', ', '.join(alertas.get('valores_suspeitos', [])) or 'NENHUM (0)'),
        ('Datas alheias', ', '.join(alertas.get('datas_suspeitas', [])) or 'NENHUM (0)'),
        ('CPFs alheios', ', '.join(alertas.get('cpfs_suspeitos', [])) or 'NENHUM (0)'),
        ('CNPJs alheios', ', '.join(alertas.get('cnpjs_suspeitos', [])) or 'NENHUM (0)'),
    ]
    for k, v in audit_lines:
        p = doc.add_paragraph()
        r = p.add_run(k + ': ')
        r.bold = True
        p.add_run(v)

    p = doc.add_paragraph()
    r = p.add_run('Observação: ')
    r.bold = True
    p.add_run('Valores como R$ 3.000,00 (BINCLUBSERVIÇOS) e R$ 500,00 (multa coercitiva) e datas de jurisprudência (30/10/2023, 28/05/2024 etc.) são estruturais do template e legítimas.')

    doc.add_heading('4. CHECKLIST PRÉ-PROTOCOLO', level=2)
    itens = [
        f'Conferir nome ({dados["nome_completo"]}) e CPF ({dados["cpf"]}).',
        f'Conferir conta/agência ({dados.get("conta", "?")} / {dados.get("agencia", "?")}).',
        f'Conferir rubrica e valores ({dados["total_descontos"]} simples / {dados["dobro_descontos"]} dobro).',
        f'Conferir dano moral ({dados["dano_moral_total"]}) e valor da causa ({dados["valor_causa"]}).',
        f'Confirmar com cliente: nunca contratou {dados["nome_terceiro"]}.',
        'Anexar: 2-Procuração específica; 3-RG; 4-Hipossuficiência; 5-Comprovante residência; 6-Extrato; 7-Tabela; 8-Notificação; 8.1-Comprovante de envio (se houver).',
        'Conferir grifo amarelo + formatação rubrica (CAPS+bold+itálico+sublinhado).',
        'Conferir nome da autora e do terceiro em Segoe UI Bold (rStyle 2TtuloChar).',
    ]
    for it in itens:
        doc.add_paragraph(it, style='List Number')

    p = doc.add_paragraph()
    r = p.add_run('Conclusão: ')
    r.bold = True
    if alertas.get('placeholders_residuais') or alertas.get('cpfs_suspeitos'):
        p.add_run('NÃO PROTOCOLAR — auditoria crítica.')
    else:
        p.add_run('APTA com ressalvas (ver pendências). Após confirmar com cliente, ')
        r2 = p.add_run('PROTOCOLAR.')
        r2.bold = True

    doc.save(relatorio_path)


def montar_dados_padrao(
    autora: dict,
    conta: dict,
    renda: dict,
    tese: dict,
    terceiro: dict,
    eh_idoso: bool = False,
    competência: str = 'Maués',
    uf: str = 'AM',
):
    """Helper para montar o dicionário completo da skill.

    autora: {nome, nacionalidade, estado_civil, profissao, cpf, rg, orgao_expedidor_prefixo,
             logradouro, numero, bairro, cidade, cep}
    conta: {agencia, numero}
    renda: {valor_float}
    tese: {rubrica, lancamentos: [(data, valor), ...]}
    terceiro: {nome, cnpj, logradouro, numero, bairro, cidade, uf, cep}
    """
    total = sum(v for _, v in tese['lancamentos'])
    dobro = total * 2
    dano_moral = 15000.0
    valor_causa = dobro + dano_moral

    datas = [d for d, _ in tese['lancamentos']]
    inicio = datas[0]
    fim = datas[-1]

    rg_field = autora.get('rg', '')
    if rg_field and not rg_field.endswith(' '):
        rg_field = rg_field + ' '

    # Inferência de gênero/condição (usada nos placeholders de pessoa autora)
    _genero = inferir_genero_pelo_prenome(autora['nome'])
    _profissao = autora.get('profissao', 'aposentado(a)')

    dados = {
        'competência': competência,
        'uf': uf,
        'prioridade_cabecalho': (
            'Prioridade de tramitação: art. 1.048 do Código de Processo Civil (Idoso)'
            if eh_idoso else ''
        ),

        'nome_completo': autora['nome'],
        'nacionalidade': autora.get('nacionalidade', 'brasileira'),
        'estado_civil': autora.get('estado_civil', ''),
        'profissao': _profissao,
        # Placeholders inferidos automaticamente
        'condicao_socioeconomica': inferir_condicao_socioeconomica(_profissao, _genero),
        'parte_autora': inferir_parte_autora(_genero),
        'parte_autora_artigo': inferir_parte_autora_artigo(_genero),
        'cpf': autora['cpf'],
        'rg': rg_field,
        'orgao_expedidor_prefixo': autora.get('orgao_expedidor_prefixo', ''),
        'logradouro': autora['logradouro'],
        'numero': autora.get('numero', 's/nº'),
        'bairro': autora.get('bairro', ''),
        'cidade_de_residencia': autora['cidade'],
        'cep': autora['cep'],

        'agencia': conta['agencia'],
        'conta': conta['numero'],

        'valor_remuneração': fmt_moeda_rs(renda['valor_float']),
        'valor_remuneração_extenso': extenso_moeda(renda['valor_float']),

        'rubrica_curta_caps': tese['rubrica'],
        'inicio_desconto': inicio,
        'fim_desconto': fim,
        'total_descontos': fmt_moeda_rs(total),
        'total_descontos_extenso': extenso_moeda(total),
        'dobro_descontos': fmt_moeda_rs(dobro),
        'dobro_descontos_extenso': extenso_moeda(dobro),
        'dano_moral_total': fmt_moeda_rs(dano_moral),
        'dano_moral_total_extenso': extenso_moeda(dano_moral),
        'valor_causa': fmt_moeda_rs(valor_causa),
        'valor_causa_extenso': extenso_moeda(valor_causa),

        'nome_terceiro': terceiro['nome'],
        'cnpj_terceiro': terceiro['cnpj'],
        'logradouro_terceiro': terceiro['logradouro'],
        'numero_terceiro': terceiro['numero'],
        'bairro_terceiro': terceiro['bairro'],
        'cidade_terceiro': terceiro['cidade'],
        'uf_terceiro': terceiro['uf'],
        'cep_terceiro': terceiro['cep'],

        'pedido_prioridade': (
            'A prioridade na tramitação, tendo em vista que a parte autora é pessoa idosa, nos termos do art. 1.048, inciso I, do Código de Processo Civil'
            if eh_idoso else ''
        ),
    }
    return dados, dict(total=total, dobro=dobro, dano_moral=dano_moral, valor_causa=valor_causa, datas=datas)


def fmt_moeda_rs(v):
    return 'R$ ' + fmt_moeda(v)


# ============================================================
# INFERÊNCIA DE GÊNERO E PLACEHOLDERS DE PESSOA AUTORA
# ============================================================

# Prenomes femininos comuns no acervo do escritório (lista vai crescendo)
PRENOMES_FEMININOS = {
    "MARIA", "ANA", "JOANA", "EDNA", "CLAUDIA", "CLÁUDIA", "JULIA", "JÚLIA",
    "ANTONIA", "ANTÔNIA", "MARINA", "LUCIA", "LÚCIA", "LAURA", "MARTA",
    "ESTER", "SARA", "RAIMUNDA", "PATRICIA", "PATRÍCIA", "CRISTINA",
    "EDINALVA", "EDMUNDA", "ANAIZA", "GEDALVA", "GILVANETE", "JOSEFA",
    "LINDINALVA", "MARCIA", "MÁRCIA", "CICERA", "CÍCERA", "MARINETE",
    "CECILIA", "CECÍLIA", "ROSANGELA", "ROSÂNGELA", "FATIMA", "FÁTIMA",
    "REGINA", "VANESSA", "JESSICA", "JÉSSICA", "BIANCA", "ADRIANA",
    "MIRIAM", "MIRIAN", "DOROTEIA", "DOROTÉIA", "ELINALDA", "TEREZINHA",
    "RAIZEL", "EDINA", "MIDIA", "MÍDIA", "NILCIENE", "ELINALDA",
    "MARLENE", "DULCE", "DULCINEIA", "DULCINÉIA", "CONCEICAO", "CONCEIÇÃO",
    "ALICE", "ALDA", "IRENE", "TEREZA", "TERESA", "FRANCISCA", "BENEDITA",
    "LEONICE", "ROSA", "MICAELA", "RAFAELA", "SIMONE", "SIMONE",
}

PRENOMES_MASCULINOS = {
    "JOAO", "JOÃO", "JOSE", "JOSÉ", "ANTONIO", "ANTÔNIO", "FRANCISCO",
    "CARLOS", "PAULO", "PEDRO", "MANUEL", "LUIZ", "LUÍS", "LUIS",
    "ROBERTO", "MARCOS", "EDVALDO", "CICERO", "CÍCERO", "IDALVO",
    "OTAVIANO", "DOMICIO", "DOMÍCIO", "CLAUDIO", "CLÁUDIO", "DENIVAL",
    "ELINALDO", "VITOR", "VÍTOR", "VICTOR", "RAIMUNDO", "SEBASTIAO",
    "SEBASTIÃO", "FERNANDO", "RICARDO", "FELIPE", "RAFAEL", "GUSTAVO",
    "DANIEL", "MARCELO", "ANDERSON", "EDUARDO", "ALEXANDRE", "GABRIEL",
    "ADAO", "ADÃO", "BENEDITO", "GERALDO", "MIGUEL", "DAVI", "ABEL",
    "EDIMILSON", "EDIVALDO", "EDMILSON", "RAIMILSON", "EUDES",
}


def inferir_genero_pelo_prenome(nome_completo: str) -> str:
    """Devolve 'F' ou 'M' a partir do primeiro nome.
    Estratégia: 1) lista de prenomes conhecidos; 2) terminação heurística;
    3) fallback 'F' (a base do escritório é majoritariamente feminina)."""
    if not nome_completo:
        return 'F'
    primeiro = nome_completo.strip().split()[0].upper()
    # Remover acentos para casar com a lista normalizada
    if primeiro in PRENOMES_FEMININOS:
        return 'F'
    if primeiro in PRENOMES_MASCULINOS:
        return 'M'
    # Heurística por terminação (não é infalível)
    if primeiro.endswith('A') and primeiro not in {'JOAO', 'JOÃO', 'ADAO', 'ADÃO', 'COSTA', 'LIMA'}:
        return 'F'
    if primeiro.endswith('O'):
        return 'M'
    # Fallback: feminino (acervo predominantemente feminino — idosas/INSS)
    return 'F'


# Flexão de termos comuns por gênero
FLEXAO_PROFISSAO = {
    # forma neutra/masculina → (masc, fem)
    'aposentado':           ('aposentado', 'aposentada'),
    'aposentada':           ('aposentado', 'aposentada'),
    'aposentado(a)':        ('aposentado', 'aposentada'),
    'pensionista':          ('pensionista', 'pensionista'),
    'agricultor':           ('agricultor', 'agricultora'),
    'agricultora':          ('agricultor', 'agricultora'),
    'do lar':               ('do lar', 'do lar'),
    'trabalhador rural':    ('trabalhador rural', 'trabalhadora rural'),
    'trabalhadora rural':   ('trabalhador rural', 'trabalhadora rural'),
    'pescador':             ('pescador', 'pescadora'),
    'pescadora':            ('pescador', 'pescadora'),
    'autônomo':             ('autônomo', 'autônoma'),
    'autônoma':             ('autônomo', 'autônoma'),
    'comerciante':          ('comerciante', 'comerciante'),
    'professor':            ('professor', 'professora'),
    'professora':           ('professor', 'professora'),
    'servidor público':     ('servidor público', 'servidora pública'),
    'servidora pública':    ('servidor público', 'servidora pública'),
}


def inferir_condicao_socioeconomica(profissao: str, genero: str) -> str:
    """Recebe profissao bruta + gênero (M/F) e devolve a forma flexionada
    pra usar como {{condicao_socioeconomica}}."""
    if not profissao:
        return 'aposentada' if genero == 'F' else 'aposentado'
    chave = profissao.strip().lower()
    par = FLEXAO_PROFISSAO.get(chave)
    if par:
        return par[1] if genero == 'F' else par[0]
    # Profissão fora da tabela: devolve como está (cliente preencheu manualmente)
    return profissao


def inferir_parte_autora(genero: str) -> str:
    return 'Autora' if genero == 'F' else 'Autor'


def inferir_parte_autora_artigo(genero: str) -> str:
    return 'da' if genero == 'F' else 'do'


# Limite de tamanho do tool Read do Claude para PDFs (32MB hard cap; usar 30MB
# como gatilho para fallback, deixando margem). PDFs maiores precisam ser
# renderizados página a página em vez de lidos como blob inteiro.
PDF_READ_MAX_BYTES = 30 * 1024 * 1024


def pdf_tem_texto(pdf_path, paginas_amostra=3):
    """Verifica se o PDF tem text-layer extraível (heurística rápida)."""
    import fitz
    doc = fitz.open(pdf_path)
    for i in range(min(paginas_amostra, len(doc))):
        if (doc[i].get_text() or '').strip():
            return True
    return False


def ler_pdf_seguro(pdf_path, out_dir_render=None, max_dim=1800):
    """Lê texto de um PDF respeitando o limite de 32MB do tool `Read`.

    Retorna dict ``{"modo": "text"|"render", "texto": str, "pngs": [paths]}``.

    - Se ``os.path.getsize(pdf_path) < PDF_READ_MAX_BYTES``: caminho normal,
      retorna texto extraído de todas as páginas (vazio se for escaneado).
    - Se ≥ 30 MB: cai no fallback de renderização página a página em PNGs
      ≤ ``max_dim`` (default 1800px) na pasta ``out_dir_render`` (default
      ``<dir do pdf>/_tmp_pages``), e devolve a lista de PNGs gerados para
      leitura visual posterior. Também tenta extrair texto via fitz mesmo
      sendo grande, porque get_text não carrega o PDF inteiro em memória.
    """
    import fitz
    size = os.path.getsize(pdf_path)
    doc = fitz.open(pdf_path)
    texto = '\n'.join((p.get_text() or '') for p in doc).strip()
    if size < PDF_READ_MAX_BYTES:
        return {'modo': 'text', 'texto': texto, 'pngs': []}
    out_dir = out_dir_render or os.path.join(os.path.dirname(pdf_path), '_tmp_pages')
    paginas = list(range(1, len(doc) + 1))
    pngs = render_paginas_pdf(pdf_path, paginas, out_dir, max_dim=max_dim)
    return {'modo': 'render', 'texto': texto, 'pngs': pngs}


# Triagem combinada de leitura: extrai textos via fitz e devolve plano para o agente.
# Resolve o erro #12 da SKILL.md (paralelismo de Read estourando o limite de tokens da request).
TAMANHO_LIMITE_LEITURA_PARALELA_MB = 8


def ler_docs_cliente_seguro(paths, max_dim=1800):
    """Lê uma lista de PDFs do cliente em SÉRIE, devolvendo um dict por path.

    Cada entrada segue o formato de :func:`ler_pdf_seguro`:
    ``{"modo": "text"|"render", "texto": str, "pngs": [paths], "tamanho_mb": float}``.

    Por que existe: vários `Read` paralelos do Claude num mesmo turno (notificação
    + RG + extrato + tabela do mesmo cliente) podem somar acima do limite de
    tokens de entrada da request, mesmo quando cada PDF tem só algumas centenas
    de KB. Isso piora em batches com 4+ clientes seguidos: o contexto principal
    acumula PDFs já lidos + DOCX gerados + relatórios.

    Este helper roda dentro de um `python` do escritório (fora do contexto do
    modelo) e:

    - Para PDFs com text-layer (notificação, tabela, declarações): devolve o
      TEXTO extraído via fitz. O agente NÃO precisa de `Read` — pode usar o
      texto direto.
    - Para PDFs escaneados (RG, extrato Bradesco): renderiza páginas em PNGs
      ≤ `max_dim` (default 1800px) e devolve a lista de PNGs. O agente lê os
      PNGs UM POR TURNO via `Read` visual.
    - Para PDFs > 30 MB: aciona o fallback do `ler_pdf_seguro` automaticamente.

    Também devolve, na chave ``__plano__``, um dicionário ``{"soma_mb": X,
    "estrategia": "paralelo"|"serial", "avisos": [...]}`` para o agente saber
    se pode paralelizar Reads visuais ou se deve serializar.
    """
    soma_bytes = 0
    resultados = {}
    avisos = []
    for p in paths:
        if not os.path.exists(p):
            resultados[p] = {'erro': 'arquivo nao encontrado'}
            avisos.append(f'arquivo ausente: {os.path.basename(p)}')
            continue
        try:
            tam = os.path.getsize(p)
            soma_bytes += tam
            r = ler_pdf_seguro(p, max_dim=max_dim)
            r['tamanho_mb'] = round(tam / (1024 * 1024), 2)
            resultados[p] = r
        except Exception as e:
            resultados[p] = {'erro': str(e)}
            avisos.append(f'erro lendo {os.path.basename(p)}: {e}')

    soma_mb = soma_bytes / (1024 * 1024)
    estrategia = 'paralelo' if soma_mb < TAMANHO_LIMITE_LEITURA_PARALELA_MB else 'serial'
    if soma_mb >= TAMANHO_LIMITE_LEITURA_PARALELA_MB:
        avisos.append(
            f'soma dos PDFs = {soma_mb:.1f} MB ≥ {TAMANHO_LIMITE_LEITURA_PARALELA_MB} MB '
            f'→ ler PNGs/Reads UM POR TURNO para evitar erro "Prompt is too long"'
        )
    resultados['__plano__'] = {
        'soma_mb': round(soma_mb, 2),
        'estrategia': estrategia,
        'avisos': avisos,
    }
    return resultados


# Helper rápido para renderizar páginas selecionadas de PDF escaneado
def render_paginas_pdf(pdf_path, paginas, out_dir, dpi=180, max_dim=1800):
    """Renderiza páginas em PNG (1-indexed). Retorna lista de PNGs gerados.

    O Claude bloqueia imagens > 2000px em qualquer dimensão (limite many-image
    requests). Para garantir que isso não estoure, calcula um zoom efetivo a
    partir do tamanho real da página e clampeia em ``max_dim`` (default 1800px,
    margem de 200px). O parâmetro ``dpi`` ainda é aceito para retrocompat, mas
    é apenas teto: se em ``dpi`` a maior dimensão passar de ``max_dim``, o zoom
    é reduzido automaticamente.
    """
    import fitz
    os.makedirs(out_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    saidas = []
    for p in paginas:
        idx = p - 1
        if idx < 0 or idx >= len(doc):
            continue
        page = doc[idx]
        rect = page.rect  # em pontos (1pt = 1/72 polegada)
        zoom_dpi = dpi / 72.0
        zoom_clamp = max_dim / max(rect.width, rect.height)
        zoom = min(zoom_dpi, zoom_clamp)
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        out = os.path.join(out_dir, f'p{p:03d}.png')
        pix.save(out)
        saidas.append(out)
    return saidas
