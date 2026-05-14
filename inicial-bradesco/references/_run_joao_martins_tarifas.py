"""Inicial TARIFAS — EXEMPLO MARTINS DA SILVA.

Comarca Manacapuru/AM (Ag 3707 / Conta 8698-3, endereço Manacapuru).
Pessoa IDOSA (afirmação expressa na notificação extrajudicial).
Aposentado pelo INSS (R$ 1.621,00 em 30/01/2026 — último crédito).
Procuração assinada A ROGO (rogada Marilene Pereira da Silva +
testemunhas Rosiane e Deilson).

Tabela: 76 lançamentos TARIFA BANCARIA - CESTA FACIL ECONOMICA entre
31/01/2017 e 15/09/2022. Total R$ 2.331,00 / dobro R$ 4.662,00.
VC R$ 19.662,00 — cabe folgadamente no JEC.

PRESCRIÇÃO: muitos lançamentos antes de 30/03/2021 (marco EAREsp
1.280.825 STJ). Procurador decide se aplica corte ou pleiteia tudo
testando art. 205 CC (10 anos).

PENDÊNCIA: cliente também tem MORA e TÍTULO DE CAPITALIZAÇÃO em
pastas separadas. Notificação previa dano moral R$ 5.000 (rateio para
3 teses combinadas). Como aqui é inicial isolada, pleitear R$ 15.000.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\EXEMPLO MARTINS DA SILVA - Ney Pedroza - TARIFAS\TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_JOAO_MARTINS_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_JOAO_MARTINS_v1.docx')

LANCAMENTOS = [
    ('31/01/2017', 19.90), ('15/02/2017', 19.90), ('31/03/2017', 19.90),
    ('13/04/2017', 22.20), ('15/05/2017', 22.20), ('14/06/2017', 22.20),
    ('14/07/2017', 22.20), ('31/08/2017', 22.20), ('15/09/2017', 22.20),
    ('13/10/2017', 22.20), ('14/11/2017', 22.20), ('15/12/2017', 24.00),
    ('15/01/2018', 24.00), ('15/02/2018', 24.00), ('29/03/2018', 24.00),
    ('13/04/2018', 24.00), ('15/05/2018', 24.00), ('15/06/2018', 28.80),
    ('13/07/2018', 28.80), ('15/08/2018', 28.80), ('14/09/2018', 28.80),
    ('31/10/2018', 28.80), ('14/11/2018', 28.80), ('14/12/2018', 28.80),
    ('15/01/2019', 22.77), ('31/01/2019', 10.23), ('15/02/2019', 33.00),
    ('15/03/2019', 33.00), ('23/04/2019', 28.13), ('30/04/2019', 4.87),
    ('15/05/2019', 7.82), ('31/05/2019', 25.18), ('28/06/2019', 33.00),
    ('15/07/2019', 13.63), ('31/07/2019', 20.37), ('15/08/2019', 34.00),
    ('13/09/2019', 34.00), ('15/10/2019', 34.00), ('14/11/2019', 34.00),
    ('13/12/2019', 34.00), ('31/01/2020', 36.70), ('28/02/2020', 31.81),
    ('23/04/2020', 36.70), ('23/04/2020', 36.70), ('23/04/2020', 4.89),
    ('15/05/2020', 36.70), ('15/06/2020', 36.70), ('15/07/2020', 36.70),
    ('14/08/2020', 36.70), ('15/09/2020', 36.70), ('15/10/2020', 24.99),
    ('30/10/2020', 14.41), ('13/11/2020', 0.33), ('25/11/2020', 39.07),
    ('15/12/2020', 39.40), ('29/01/2021', 39.40), ('26/02/2021', 39.40),
    ('15/03/2021', 39.40), ('30/04/2021', 41.90), ('31/05/2021', 41.90),
    ('15/06/2021', 41.90), ('15/07/2021', 41.90), ('13/08/2021', 41.90),
    ('15/09/2021', 41.90), ('15/10/2021', 41.90), ('12/11/2021', 41.90),
    ('15/12/2021', 41.90), ('14/01/2022', 41.90), ('15/02/2022', 44.50),
    ('15/03/2022', 44.50), ('14/04/2022', 44.50), ('16/05/2022', 44.50),
    ('15/06/2022', 44.50), ('15/07/2022', 44.50), ('15/08/2022', 46.70),
    ('15/09/2022', 46.70),
]

autora = {
    'nome': 'EXEMPLO MARTINS DA SILVA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'aposentado',
    'cpf': '000.000.013-23',
    'rg': '1000011-1',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua do Campinho',
    'numero': '1333',
    'bairro': 'Novo Manaca',
    'cidade': 'Manacapuru',
    'cep': '69.409-899',
}
conta = {'agencia': '3707', 'numero': '8698-3'}
renda = {'valor_float': 1621.00}

tese = {'rubrica': 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Manacapuru', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== EXEMPLO MARTINS — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_JOAO_MARTINS', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_JOAO_MARTINS_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC (afirmação expressa na notificação)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 30/01/2026)'),
    ('Período tabela', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO — corte de 30/03/2021 (EAREsp 1.280.825 STJ)',
     'A tabela inicia em 31/01/2017 e termina em 15/09/2022. Considerando o marco '
     'prescricional do art. 27 do CDC (5 anos) firmado pelo STJ no EAREsp 1.280.825 '
     '(30/03/2021), aproximadamente 53 lançamentos podem estar prescritos (todos '
     'antes de 30/03/2021). Lançamentos pós-30/03/2021: 23 de R$ 41,90 a R$ 46,70 = '
     '~R$ 999,80 simples (R$ 1.999,60 dobro). DECISÃO DO PROCURADOR: pleitear (a) '
     'TUDO testando art. 205 CC (10 anos — minoritária TJAM); (b) corte 5 anos '
     'reduzindo VC para ~R$ 16.999,60; (c) corte 5 anos com pedido subsidiário pelo '
     'restante. Inicial gerada com TODOS 76 lançamentos — REVISAR antes do protocolo.'),
    ('IDOSO — prioridade aplicada',
     'Notificação afirma expressamente "pessoa idosa". Prioridade do art. 1.048, I, '
     'CPC inserida no cabeçalho. Conferir RG (≥ 60 anos).'),
    ('PROCURAÇÃO ASSINADA A ROGO',
     'Pasta TARIFA tem RG da rogada (Marilene Pereira da Silva) + 2 testemunhas '
     '(Rosiane e Deilson). Cliente provavelmente analfabeto ou com dificuldade de '
     'assinar. Conferir validade da procuração. Mencionar a rogo na qualificação '
     'detalhada se for o caso.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente.'),
    ('NOTIFICAÇÃO previa dano moral R$ 5.000 — divergência intencional',
     'A notificação extrajudicial pleiteou R$ 5.000 (provavelmente projetando 3 teses '
     'combinadas: TARIFAS + MORA + TÍTULO CAPITALIZAÇÃO = R$ 15.000 total). Como esta '
     'inicial é ISOLADA por tese (TARIFAS), pleiteia R$ 15.000 individuais conforme '
     'regra § 9 da skill. CONFIRMAR com procurador se quer manter R$ 15.000 ou ajustar '
     'para R$ 5.000 (e equilibrar com as outras 2 iniciais).'),
    ('CLIENTE TEM 3 TESES SEPARADAS',
     'Cliente EXEMPLO MARTINS DA SILVA aparece nas pastas: TARIFAS (esta), MORA, e '
     'TÍTULO DE CAPITALIZAÇÃO. Comarca Manacapuru — adota INICIAL COMBINADA por '
     'padrão. Procurador organizou em pastas separadas, então gerou 3 iniciais '
     'individuais. Avaliar se deve consolidar em 1 só inicial-combinada antes do '
     'protocolo.'),
    ('TETO JEC — coberto',
     'VC R$ 19.662,00 ≈ 12,9 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
    ('LANÇAMENTOS DUPLICADOS em 23/04/2020',
     'Em 23/04/2020 a tabela registra 3 lançamentos (R$ 36,70 + 36,70 + 4,89). Pode '
     'ser cobrança em lote ou erro. Conferir no extrato original.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'DECIDIR estratégia de prescrição (corte 30/03/2021 ou pleito decenal).',
    'AVALIAR consolidação em inicial-combinada (TARIFAS + MORA + TÍTULO CAP.).',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos confirmado).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Manacapuru/AM (Justiça Estadual Delegada).',
    'Conferir validade da procuração assinada a rogo.',
    'Confirmar com cliente: nunca contratou cesta fácil econômica.',
    'Anexar 2-Procuração + 3-RG + 3.1-RG da rogada + 3.2/3.3-RGs testemunhas + 4-Hipossuficiência + 5-Comprovante + 5.1-Declaração domicílio + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir prescrição e estratégia (1 inicial vs 3 individuais). Após decisão, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
