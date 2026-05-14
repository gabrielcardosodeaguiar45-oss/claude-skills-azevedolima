"""Inicial MORA — EXEMPLO MANUEL CORDOVIL (mesma pasta da TARIFAS).

Tabela MORA agrupa MORA CRED PESSOAL + SERVIÇO CARTÃO PROTEGIDO numa
única tese (conforme procurador organizou). 18 lançamentos entre
05/02/2021 e 02/07/2025. Total R$ 2.516,22 / dobro R$ 5.032,44.
VC R$ 20.032,44 — cabe folgadamente no JEC.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\EXEMPLO MANUEL CORDOVIL - Wilson - TARIFA\TARIFA E MORA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Mora_MANUEL_LAZARO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MANUEL_LAZARO_MORA_v1.docx')

LANCAMENTOS = [
    ('05/02/2021', 0.56), ('02/03/2021', 9.43), ('05/03/2021', 8.48),
    ('05/04/2021', 2.67), ('04/07/2022', 21.09), ('04/10/2022', 297.86),
    ('03/01/2023', 296.47), ('04/04/2023', 297.88), ('03/05/2023', 296.48),
    ('02/08/2023', 108.41), ('04/09/2023', 118.97), ('03/01/2025', 152.04),
    ('04/02/2025', 145.65), ('07/03/2025', 153.34), ('02/04/2025', 151.88),
    ('05/05/2025', 153.31), ('03/06/2025', 151.58), ('02/07/2025', 150.12),
]

autora = {
    'nome': 'EXEMPLO MANUEL CORDOVIL',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'aposentado',
    'cpf': '000.000.018-28',
    'rg': '1000016-6',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Pimentel Tavares',
    'numero': '341',
    'bairro': 'CM Terra P do Limão',
    'cidade': 'Barreirinha',
    'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '2782-0'}
renda = {'valor_float': 846.22}

tese = {'rubrica': 'MORA CRÉDITO PESSOAL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Barreirinha', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['rubrica_curta'] = 'Mora Cred Pess'
dados['rubrica_curta_caps'] = 'MORA CRED PESS'
dados['rubrica_completa'] = 'Crédito Mora Pessoal'
dados['rubrica_completa_caps'] = 'MORA CREDITO PESSOAL'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== EXEMPLO MANUEL — MORA ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Mora_MANUEL_LAZARO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA — MORA CRED PESS + SERVIÇO CARTÃO PROTEGIDO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Mora_MANUEL_LAZARO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 02/07/2025)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (MORA CRED PESS + SERVIÇO CARTÃO PROTEGIDO agrupados)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('TABELA AGRUPA 2 RUBRICAS — MORA + SERVIÇO CARTÃO',
     'A tabela "7 - TABELA MORA CRED.pdf" agrupa SERVIÇO CARTÃO PROTEGIDO (4 lanç. de '
     'fev-abr/2021 = R$ 21,14) e MORA CRED PESSOAL (14 lanç. = R$ 2.495,08). Procurador '
     'agrupou por compartilharem a procuração MORA CRED PESS. A inicial usa template '
     'inicial-mora.docx tratando como tese única MORA. Se preferir, separar SERVIÇO '
     'CARTÃO em inicial-tarifas (tarifa de seguro de cartão).'),
    ('PRESCRIÇÃO — corte 30/03/2021',
     'Os 4 lançamentos SERVIÇO CARTÃO PROTEGIDO de fev-abr/2021 estão antes de '
     '30/03/2021 (alguns) — possível prescrição. MORA CRED PESS começa em 04/07/2022, '
     'todos válidos. Verificar antes do protocolo.'),
    ('IDOSO — prioridade aplicada',
     'Notificação afirma "pessoa idosa".'),
    ('CONSOLIDAÇÃO COM TARIFAS — recomendada',
     'Cliente também tem CARTAO ANUIDADE (TARIFAS) e PG ELETRON. AVALIAR inicial-'
     'combinada cobrindo as 3 famílias.'),
    ('VALORES MORA ELEVADOS',
     'Lançamentos de MORA CRED PESSOAL em 2022/2023 chegam a R$ 296-297. Indica '
     'cheque especial elevado. Conferir extrato — pode haver tese de cumulação com '
     'ENCARGOS LIMITE DE CRÉDITO se houver. Tabela atual NÃO menciona ENC LIM CRED.'),
    ('TETO JEC — coberto',
     'VC R$ 20.032,44 ≈ 13,2 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'DECIDIR consolidação (CARTAO ANUIDADE + SERVIÇO CARTÃO + MORA + PG ELETRON).',
    'DECIDIR sobre prescrição dos 4 lanç. de SERVIÇO CARTÃO em 2021.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Barreirinha/AM.',
    'Confirmar com cliente: nunca contratou seguro de cartão nem aceitou cheque especial.',
    'Anexar 2-Procuração MORA CRED PESS + 3-RG + 5-Comprovante + 6-Extrato + 7-Tabela MORA + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir consolidação. Após resolução, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
