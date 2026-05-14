"""
Auditor pós-geração: detecta valores que NÃO foram substituídos pelos placeholders
e podem ser hardcoded escapado, vazamento de outro caso, ou erro de template.

REGRA CRÍTICA: rodar SEMPRE depois de aplicar_template(). Se aparecer alerta,
NÃO entregar para o cliente sem revisão manual.

Categorias de alerta:
- VALOR_HARDCODED: valores R$ que não correspondem a nenhum {{placeholder}} substituído
- DATA_SUSPEITA: data DD/MM/AAAA não presente em parsear_tabela_descontos
- CPF_INVALIDO: CPF não corresponde ao do autor declarado
- CNPJ_DESCONHECIDO: CNPJ não está no whitelist (Bradesco + terceiros conhecidos)
- NOME_VAZADO: nome em CAIXA ALTA (8+ letras) que não é o autor / nome_terceiro
- PLACEHOLDER_RESIDUAL: {{xxx}} que sobrou no documento
- RUBRICA_SEM_FORMATO: rubrica em CAPS sem grifo amarelo (não foi reconhecida como rubrica)
"""
import re
from docx import Document


# ============================================================
# WHITELISTS — entidades conhecidas que NÃO precisam alertar
# ============================================================
CNPJS_CONHECIDOS = {
    '60.746.948/0001-12',   # Bradesco S.A. (matriz)
    '60.746.948',           # Bradesco prefixo
    '92.892.256/0001-79',   # MBM Previdência Complementar
    '58.119.199/0001-51',   # Odontoprev
    '33.067.626/0001-83',   # Aspecir Previdência (variável conforme caso)
    '17.184.037/0001-10',   # OAB/AM (escritório)
}

# Valores monetários "estruturais" que costumam aparecer no pedido (e estão OK)
VALORES_ESTRUTURAIS = {
    '15.000,00',    # dano moral isolado
    '10.000,00',    # dano moral em alguns templates
    '5.000,00',     # dano moral por tese combinada
    '1.412,00',     # custas estimativa eventualmente
}

# Padrões que indicam o texto foi corretamente substituído
PADROES_OK_NOME = [
    r'TEREZINHA BRANDÃO DA ROCHA',  # casos paradigma
    r'JOSÉ SEBASTIÃO',
    r'MARIA JOANA',
    r'ELINALDO',
    r'BANCO BRADESCO',
    r'BRADESCO',
]


# ============================================================
# REGEX DE DETECÇÃO
# ============================================================
RE_VALOR_RS = re.compile(r'R\$\s*([\d.]+,\d{2})')
RE_DATA_DDMMAAAA = re.compile(r'\b(\d{2}/\d{2}/\d{4})\b')
RE_CPF = re.compile(r'\b(\d{3}\.\d{3}\.\d{3}-\d{2})\b')
RE_CNPJ = re.compile(r'\b(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2})\b')
RE_PLACEHOLDER = re.compile(r'\{\{([^{}]+)\}\}')
RE_NOME_CAPS = re.compile(r'\b([A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}(?:\s+[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,}){2,})\b')


# ============================================================
# AUDITORIA PRINCIPAL
# ============================================================
def auditar_docx(docx_path, dados_caso):
    """Audita o DOCX gerado contra os dados esperados.

    Args:
        docx_path: caminho do .docx gerado
        dados_caso: dict com os dados que DEVIAM estar no documento.
                    Chaves esperadas: cpf, valores_legitimos (set de R$),
                    datas_legitimas (set de DD/MM/AAAA), nomes_legitimos (set str),
                    cnpjs_legitimos (set), valor_remuneração (str)

    Returns:
        dict com alertas categorizados:
            {
                'placeholders_residuais': [str],
                'valores_suspeitos': [str],
                'datas_suspeitas': [str],
                'cpfs_suspeitos': [str],
                'cnpjs_suspeitos': [str],
                'nomes_vazados': [str],
                'severidade': 'OK' | 'ATENCAO' | 'CRITICO',
                'total_alertas': int,
            }
    """
    d = Document(docx_path)
    texto = '\n'.join(p.text for p in d.paragraphs)

    # também varre tabelas
    for tab in d.tables:
        for row in tab.rows:
            for cell in row.cells:
                texto += '\n' + cell.text

    alertas = {
        'placeholders_residuais': [],
        'valores_suspeitos': [],
        'datas_suspeitas': [],
        'cpfs_suspeitos': [],
        'cnpjs_suspeitos': [],
        'nomes_vazados': [],
        'rubricas_sem_formato': [],
    }

    # 1. Placeholders residuais (mais grave de todos)
    placeholders = sorted(set(RE_PLACEHOLDER.findall(texto)))
    if placeholders:
        alertas['placeholders_residuais'] = placeholders

    # 2. Valores R$ suspeitos
    valores_legit = set(dados_caso.get('valores_legitimos', set()))
    valores_legit |= VALORES_ESTRUTURAIS
    # adiciona o valor_remuneração tbm
    if dados_caso.get('valor_remuneração'):
        valores_legit.add(dados_caso['valor_remuneração'])
    valores_no_doc = set(RE_VALOR_RS.findall(texto))
    suspeitos = valores_no_doc - valores_legit
    if suspeitos:
        alertas['valores_suspeitos'] = sorted(suspeitos)

    # 3. Datas suspeitas
    datas_legit = set(dados_caso.get('datas_legitimas', set()))
    datas_no_doc = set(RE_DATA_DDMMAAAA.findall(texto))
    datas_suspeitas = datas_no_doc - datas_legit
    if datas_suspeitas:
        alertas['datas_suspeitas'] = sorted(datas_suspeitas)

    # 4. CPF
    cpf_legit = dados_caso.get('cpf', '')
    cpfs_no_doc = set(RE_CPF.findall(texto))
    if cpf_legit in cpfs_no_doc:
        cpfs_no_doc.discard(cpf_legit)
    if cpfs_no_doc:
        alertas['cpfs_suspeitos'] = sorted(cpfs_no_doc)

    # 5. CNPJ
    cnpjs_legit = set(dados_caso.get('cnpjs_legitimos', set()))
    cnpjs_legit |= CNPJS_CONHECIDOS
    cnpjs_no_doc = set(RE_CNPJ.findall(texto))
    suspeitos_cnpj = cnpjs_no_doc - cnpjs_legit
    if suspeitos_cnpj:
        alertas['cnpjs_suspeitos'] = sorted(suspeitos_cnpj)

    # 6. Nomes em CAPS vazados (de outro caso)
    nomes_legit = set(dados_caso.get('nomes_legitimos', set()))
    for padrao_ok in PADROES_OK_NOME:
        nomes_legit.add(padrao_ok)
    nomes_no_doc = set(RE_NOME_CAPS.findall(texto))
    nomes_vazados = []
    for n in nomes_no_doc:
        # ignora se contém qualquer um dos legítimos
        if any(legit in n or n in legit for legit in nomes_legit):
            continue
        # ignora cabeçalhos/títulos comuns ("EXMO SR DR JUIZ", "VARA CIVEL", etc.)
        if re.match(r'^(EXMO|EXCEL|MM|VARA|JUSTI|TRIBUNAL|JUIZ[OA]?|FORO|COMARCA|CO?MARCA|CAUSA|VALOR|CDC|DOS|DAS|DA|DO|RG|CPF|CNPJ|UF|CEP|MORA|TARIFA|CRED|ENC|LIM|APLIC|PG|ELETRON|TITULO|CAPITALIZA|ANUIDADE|CESTA|EXPRESSO|BANCARIA|CONSIGNAD|CONTRATO|CO?BRAN[ÇC]A|FIN|LIMITE|SALDO|PESS|FÁCIL|FACIL|INVEST|RENTAB|RESGATE|SSP|AM|RS|SP|RJ|MG|BA|CARTAO|CART[ÃA]O|CR[ÉE]DITO|DEBITO|D[ÉE]BITO|BANCO|BRADESCO|S[/.]?A)\b', n):
            continue
        nomes_vazados.append(n)
    if nomes_vazados:
        alertas['nomes_vazados'] = sorted(set(nomes_vazados))

    # ====== Severidade ======
    total = sum(len(v) for v in alertas.values())
    if alertas['placeholders_residuais'] or alertas['cpfs_suspeitos']:
        severidade = 'CRITICO'
    elif total > 0:
        severidade = 'ATENCAO'
    else:
        severidade = 'OK'

    alertas['severidade'] = severidade
    alertas['total_alertas'] = total
    return alertas


# ============================================================
# AUDITORIA TABELA (NotebookLM) vs EXTRATO (parsing direto)
# ============================================================
def auditar_tabela_vs_extrato(tabela_path: str, extrato_path: str,
                              rubricas_alvo: list[str],
                              tolerancia_centavos: int = 1) -> dict:
    """Cruza a tabela do NotebookLM com o extrato bancário parseado direto.

    Para cada rubrica em `rubricas_alvo`, conta lançamentos + soma valor +
    coleta datas em ambas as fontes e detecta divergências.

    Args:
        tabela_path: PDF da tabela (gerado pelo NotebookLM)
        extrato_path: PDF do extrato bancário Bradesco
        rubricas_alvo: lista de rubricas a auditar (ex.:
            ['MORA CRED PESS', 'ENC LIM CRED'])
        tolerancia_centavos: divergência aceitável na soma (default 1 centavo)

    Returns:
        {
            'severidade': 'OK' / 'ATENCAO' / 'CRITICO',
            'rubricas': [
                {
                    'rubrica': str,
                    'tabela': {'qtd': int, 'soma': float, 'datas': [str]},
                    'extrato': {'qtd': int, 'soma': float, 'datas': [str]},
                    'divergencias': [str],   # lista de problemas detectados
                    'fonte_recomendada': 'extrato' / 'tabela' / 'paridade',
                },
                ...
            ],
            'recomendacao_global': str,
        }

    Regra de fonte_recomendada:
      - paridade: contagens iguais E somas iguais (dentro da tolerância)
      - extrato: contagem do extrato > tabela (tabela está incompleta)
      - tabela: contagem do extrato < tabela (tabela traz lançamento que o
        extrato não detectou — possível erro de filtro do parser direto)
    """
    import os
    if not os.path.exists(tabela_path) or not os.path.exists(extrato_path):
        return {'severidade': 'CRITICO', 'rubricas': [],
                'recomendacao_global': 'Tabela ou extrato ausente — não foi possível auditar.'}

    # Imports locais para evitar dependência circular
    import sys
    sys.path.insert(0, os.path.dirname(__file__))
    from extrator_documentos import parsear_tabela_descontos, parsear_lancamentos_extrato

    relatorio_rubricas = []
    severidade = 'OK'

    for rubrica in rubricas_alvo:
        # Tabela: filtra por rubrica
        descontos_tab = parsear_tabela_descontos(tabela_path, filtro_rubrica=rubrica)
        # Extrato: parseia direto
        descontos_ext = parsear_lancamentos_extrato(extrato_path, filtro_rubrica=rubrica)

        qtd_tab = len(descontos_tab)
        qtd_ext = len(descontos_ext)
        soma_tab = round(sum(d['valor'] for d in descontos_tab), 2)
        soma_ext = round(sum(d['valor'] for d in descontos_ext), 2)
        datas_tab = sorted({d['data'] for d in descontos_tab})
        datas_ext = sorted({d['data'] for d in descontos_ext})

        divergencias = []
        # Quantidade
        if qtd_tab != qtd_ext:
            divergencias.append(
                f'Quantidade diverge: tabela={qtd_tab} vs extrato={qtd_ext} '
                f'(diferença de {abs(qtd_tab - qtd_ext)} lançamentos)'
            )
        # Soma
        diff_centavos = round(abs(soma_tab - soma_ext) * 100)
        if diff_centavos > tolerancia_centavos:
            divergencias.append(
                f'Soma diverge: tabela=R$ {soma_tab:.2f} vs extrato=R$ {soma_ext:.2f} '
                f'(diferença R$ {abs(soma_tab - soma_ext):.2f})'
            )
        # Datas exclusivas
        so_extrato = set(datas_ext) - set(datas_tab)
        so_tabela = set(datas_tab) - set(datas_ext)
        if so_extrato:
            divergencias.append(
                f'{len(so_extrato)} data(s) só no extrato (faltando na tabela): '
                f'{sorted(so_extrato)[:5]}{"..." if len(so_extrato) > 5 else ""}'
            )
        if so_tabela:
            divergencias.append(
                f'{len(so_tabela)} data(s) só na tabela (não detectadas no extrato): '
                f'{sorted(so_tabela)[:5]}{"..." if len(so_tabela) > 5 else ""}'
            )

        # Decidir fonte recomendada
        if not divergencias:
            fonte = 'paridade'
        elif qtd_ext > qtd_tab:
            fonte = 'extrato'
            severidade = 'ATENCAO' if severidade == 'OK' else severidade
        elif qtd_tab > qtd_ext:
            fonte = 'tabela'
            severidade = 'ATENCAO' if severidade == 'OK' else severidade
        else:
            # Quantidades iguais mas valores divergem — investigar manual
            fonte = 'paridade'
            severidade = 'ATENCAO' if severidade == 'OK' else severidade

        if len(divergencias) >= 3:
            severidade = 'CRITICO'

        relatorio_rubricas.append({
            'rubrica': rubrica,
            'tabela': {'qtd': qtd_tab, 'soma': soma_tab, 'datas': datas_tab},
            'extrato': {'qtd': qtd_ext, 'soma': soma_ext, 'datas': datas_ext},
            'divergencias': divergencias,
            'fonte_recomendada': fonte,
        })

    if severidade == 'OK':
        recomendacao = 'Tabela bate com extrato em todas as rubricas. Ok seguir.'
    elif severidade == 'ATENCAO':
        recomendacao = (
            'Há divergências. Revisar manualmente e/ou regenerar tabela. '
            'A skill usará dados do extrato direto para gerar a inicial.'
        )
    else:
        recomendacao = (
            'Divergências graves. NÃO usar a tabela do NotebookLM como está. '
            'Solicitar nova tabela ou aceitar que skill use extrato direto + alerta no relatório.'
        )

    return {
        'severidade': severidade,
        'rubricas': relatorio_rubricas,
        'recomendacao_global': recomendacao,
    }


def gerar_relatorio_auditoria_tabela(rel: dict) -> str:
    """Formata o resultado de auditar_tabela_vs_extrato para texto humano."""
    linhas = []
    linhas.append('=' * 70)
    linhas.append('AUDITORIA — TABELA (NotebookLM) vs EXTRATO BANCÁRIO')
    linhas.append('=' * 70)
    linhas.append(f'Severidade: {rel["severidade"]}')
    linhas.append(f'Recomendação: {rel["recomendacao_global"]}')
    linhas.append('')
    for r in rel['rubricas']:
        linhas.append(f'• Rubrica: {r["rubrica"]}')
        linhas.append(f'   Tabela:  {r["tabela"]["qtd"]} lançamentos, R$ {r["tabela"]["soma"]:.2f}')
        linhas.append(f'   Extrato: {r["extrato"]["qtd"]} lançamentos, R$ {r["extrato"]["soma"]:.2f}')
        linhas.append(f'   Fonte recomendada: {r["fonte_recomendada"]}')
        if r['divergencias']:
            linhas.append('   Divergências:')
            for d in r['divergencias']:
                linhas.append(f'     - {d}')
        else:
            linhas.append('   ✓ Sem divergências')
        linhas.append('')
    return '\n'.join(linhas)


# ============================================================
# RELATÓRIO HUMANO
# ============================================================
def gerar_relatorio(alertas, docx_path):
    """Gera relatório textual da auditoria. Para colar no chat ou em .txt."""
    linhas = []
    linhas.append('=' * 70)
    linhas.append(f'AUDITORIA PÓS-GERAÇÃO — {docx_path}')
    linhas.append('=' * 70)
    linhas.append(f'SEVERIDADE: {alertas["severidade"]}')
    linhas.append(f'Total de alertas: {alertas["total_alertas"]}')
    linhas.append('')

    if alertas['placeholders_residuais']:
        linhas.append('🔴 CRÍTICO — Placeholders {{xxx}} NÃO substituídos:')
        for p in alertas['placeholders_residuais']:
            linhas.append(f'   • {{{{{p}}}}}')
        linhas.append('')

    if alertas['cpfs_suspeitos']:
        linhas.append('🔴 CRÍTICO — CPF não corresponde ao autor:')
        for c in alertas['cpfs_suspeitos']:
            linhas.append(f'   • {c}')
        linhas.append('')

    if alertas['cnpjs_suspeitos']:
        linhas.append('🟡 CNPJ desconhecido (verificar se é o terceiro do caso):')
        for c in alertas['cnpjs_suspeitos']:
            linhas.append(f'   • {c}')
        linhas.append('')

    if alertas['valores_suspeitos']:
        linhas.append('🟡 Valores R$ não esperados (podem ser hardcoded esquecido):')
        for v in alertas['valores_suspeitos']:
            linhas.append(f'   • R$ {v}')
        linhas.append('')

    if alertas['datas_suspeitas']:
        linhas.append('🟡 Datas DD/MM/AAAA não esperadas:')
        for d in alertas['datas_suspeitas']:
            linhas.append(f'   • {d}')
        linhas.append('')

    if alertas['nomes_vazados']:
        linhas.append('🟡 Nomes em CAIXA ALTA não esperados (vazamento de outro caso?):')
        for n in alertas['nomes_vazados']:
            linhas.append(f'   • {n}')
        linhas.append('')

    if alertas['severidade'] == 'OK':
        linhas.append('✅ Nenhum alerta. Documento liberado para revisão final.')

    linhas.append('=' * 70)
    return '\n'.join(linhas)


# ============================================================
# HELPER: monta dados_caso a partir das estruturas usadas pela skill
# ============================================================
def montar_dados_caso(autor_dict, descontos_list, valor_remuneração, terceiros_cnpj=None):
    """Conveniência: gera o dict dados_caso para auditar_docx().

    Args:
        autor_dict: dict com cpf, nome_completo
        descontos_list: list de tuplas (data, valor) ou dicts {data, valor}
        valor_remuneração: str já formatada (ex: '1.212,17')
        terceiros_cnpj: list de CNPJs do(s) terceiro(s) (PG ELETRON)

    Returns:
        dict pronto para auditar_docx
    """
    valores = set()
    datas = set()
    for d in descontos_list:
        if isinstance(d, dict):
            data, valor = d.get('data'), d.get('valor')
        else:
            data, valor = d
        if data:
            datas.add(data)
        if valor is not None:
            valores.add(f'{valor:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

    # adiciona total + dobro
    if descontos_list:
        if isinstance(descontos_list[0], dict):
            total = sum(d['valor'] for d in descontos_list)
        else:
            total = sum(v for _, v in descontos_list)
        for v in [total, total * 2, total * 2 + 15000, total * 2 + 5000,
                  total * 2 + 10000]:
            valores.add(f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

    nomes = set()
    if autor_dict.get('nome_completo'):
        nomes.add(autor_dict['nome_completo'])

    cnpjs = set(terceiros_cnpj or [])

    return {
        'cpf': autor_dict.get('cpf', ''),
        'valores_legitimos': valores,
        'datas_legitimas': datas,
        'nomes_legitimos': nomes,
        'cnpjs_legitimos': cnpjs,
        'valor_remuneração': valor_remuneração,
    }
