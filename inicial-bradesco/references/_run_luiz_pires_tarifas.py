"""Inicial TARIFAS — CLIENTE EXEMPLO.

Comarca Presidente Figueiredo/AM (Ag 3732 / Conta 20304-1, endereço
Presidente Figueiredo). Pessoa IDOSA, casado, aposentado pelo INSS
(R$ 988,00 último crédito 06/12/2024 — não há crédito 2025/2026).
RG do Acre (SSP/AC). Procuração assinada A ROGO (rogada + 2
testemunhas Evaristo e Nuberlândia).

Tabela: 49 lançamentos TARIFA BANCARIA - CESTA BENEFIC entre 15/01/2020
e 13/12/2024. Total R$ 979,80 / dobro R$ 1.959,60. VC R$ 16.959,60.

PRESCRIÇÃO: aproximadamente 12 lançamentos antes de 30/03/2021
(EAREsp 1.280.825 STJ). Procurador decide se aplica corte.

PENDÊNCIAS: cliente também tem MORA, TÍTULO CAPITALIZAÇÃO e PG ELETRON
em pastas separadas (4 teses no total).
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO - Ruth - TARIFA\TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_LUIZ_PIRES_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_LUIZ_PIRES_v1.docx')

LANCAMENTOS = [
    ('15/01/2020', 16.95), ('12/02/2021', 17.25), ('15/03/2021', 17.25),
    ('15/04/2021', 19.15), ('14/05/2021', 19.15), ('15/06/2021', 19.15),
    ('15/07/2021', 19.15), ('13/08/2021', 19.15), ('15/09/2021', 19.15),
    ('15/10/2021', 19.60), ('12/11/2021', 19.60), ('15/12/2021', 19.60),
    ('14/01/2022', 19.60), ('15/02/2022', 19.60), ('15/03/2022', 19.60),
    ('14/04/2022', 20.10), ('13/05/2022', 20.10), ('15/06/2022', 20.10),
    ('15/07/2022', 20.10), ('15/08/2022', 20.10), ('15/09/2022', 20.10),
    ('14/10/2022', 20.45), ('14/11/2022', 20.45), ('15/12/2022', 20.45),
    ('13/01/2023', 20.45), ('15/02/2023', 20.45), ('15/03/2023', 20.45),
    ('14/04/2023', 21.10), ('15/05/2023', 21.10), ('15/06/2023', 21.10),
    ('14/07/2023', 21.10), ('15/08/2023', 21.10), ('15/09/2023', 21.10),
    ('13/10/2023', 21.40), ('14/11/2023', 21.40), ('15/12/2023', 21.40),
    ('15/01/2024', 21.40), ('15/02/2024', 21.40), ('15/03/2024', 21.40),
    ('15/04/2024', 21.85), ('15/05/2024', 21.85), ('14/06/2024', 21.85),
    ('15/07/2024', 21.85), ('15/08/2024', 7.03), ('09/09/2024', 14.82),
    ('13/09/2024', 21.85), ('15/10/2024', 22.15), ('14/11/2024', 22.15),
    ('13/12/2024', 22.15),
]

autora = {
    'nome': 'CLIENTE EXEMPLO',
    'nacionalidade': 'brasileiro',
    'estado_civil': 'casado',
    'profissao': 'aposentado',
    'cpf': '000.000.017-27',
    'rg': '1000015-5',
    'orgao_expedidor_prefixo': 'SSP/AC',
    'logradouro': 'Av. Joaquim Cardoso',
    'numero': '646',
    'bairro': 'Aida Mendonça',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '20304-1'}
renda = {'valor_float': 988.00}

tese = {'rubrica': 'TARIFA BANCÁRIA - CESTA BENEFIC', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA BENEFIC'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_LUIZ_PIRES', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA - CESTA BENEFIC'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_LUIZ_PIRES_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 06/12/2024 — sem créditos em 2025/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO — corte 30/03/2021 (EAREsp 1.280.825 STJ)',
     'Aproximadamente 12 lançamentos antes de 30/03/2021 (R$ 16,95 + R$ 17,25 + ...). '
     'Pós-30/03/2021: 37 lançamentos válidos (~R$ 928 simples / R$ 1.856 dobro). '
     'Decisão do procurador: pleitear TUDO testando art. 205 CC ou aplicar corte de 5 '
     'anos. Inicial gerada com TODOS — REVISAR.'),
    ('IDOSO — prioridade aplicada',
     'Notificação afirma "pessoa idosa". Confirmar RG (≥ 60 anos).'),
    ('PROCURAÇÃO ASSINADA A ROGO',
     'Pasta tem RG da rogada + 2 testemunhas (Evaristo e Nuberlândia). Cliente provavelmente '
     'analfabeto ou com dificuldade de assinar. Conferir validade da procuração antes do '
     'protocolo.'),
    ('RG DO ACRE (SSP/AC) — autora reside em AM',
     'RG emitido em SSP/AC mas autora reside em Presidente Figueiredo/AM. Sem '
     'inconsistência (cidadão pode mudar de estado). Comprovante de residência atual '
     'em AM já está na pasta.'),
    ('Renda < 1 SM — sem créditos em 2025/2026',
     'INSS R$ 988,00 (06/12/2024) é último crédito. Não há créditos posteriores no '
     'extrato. CONFIRMAR com cliente: benefício foi cessado/cassado/suspenso? Pode '
     'afetar a tese de hipossuficiência mas REFORÇA o impacto do dano moral.'),
    ('CLIENTE TEM 4 TESES SEPARADAS',
     'Cliente CLIENTE EXEMPLO também aparece nas pastas: MORA, TÍTULO CAPITALIZAÇÃO, '
     'PG ELETRON. AVALIAR consolidação em 1 só inicial-combinada (Presidente Figueiredo '
     'adota combinação por padrão, ver § 4 da SKILL).'),
    ('NOTIFICAÇÃO previa dano moral R$ 5.000',
     'Notificação pleiteou R$ 5.000 (provavelmente projetando 4 teses combinadas: '
     'R$ 5k × 4 = R$ 20k, mas R$ 5k para "rateio justo"). Como esta inicial é ISOLADA, '
     'pleiteia R$ 15.000 conforme regra § 9. CONFIRMAR com procurador.'),
    ('TETO JEC — coberto',
     'VC R$ 16.959,60 ≈ 11,2 SM. Cabe folgadamente no JEC.'),
    ('NOTIFICAÇÃO MENCIONA MAUÉS — verificar competência',
     'A notificação foi escrita com "Maués/AM" no cabeçalho mas o ENDEREÇO da autora '
     'é Presidente Figueiredo/AM. A inicial usa Presidente Figueiredo (foro do '
     'consumidor — art. 101, I, CDC). Verificar antes do protocolo.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AVALIAR estratégia de prescrição.',
    'AVALIAR consolidação com MORA + TÍTULO + PG ELETRON em 1 só inicial-combinada.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade da procuração assinada a rogo.',
    'Verificar status do INSS (cessado? cassado? suspenso?).',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Confirmar com cliente: nunca contratou cesta benefic.',
    'Anexar 2-Procuração + 3-RG + 3.1-RG rogada + 3.2/3.3-RGs testemunhas + 5-Comprovante + 5.1-Declaração domicílio + 5.2-RG proprietária + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir prescrição e consolidação. Após resolução, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
