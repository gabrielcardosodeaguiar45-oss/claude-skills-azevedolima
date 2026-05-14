"""Inicial APLIC.INVEST FACIL — CLIENTE EXEMPLO DE ALMEIDA (NOVA tese, DIFERENTE
do PG ELETRON BRADESCO VIDA E PREVIDÊNCIA já processado em 06/05/2026).

Comarca Maués/AM (foro do domicílio — CM Pingo de Ouro, Maués Açu, Rio
Parauary, zona rural). IDOSA 65 anos (nascida 20/07/1960 — confirmado
do RG na sessão anterior). Estado civil omitido.

Tese: 23 lançamentos APLIC.INVEST FACIL entre 08/03/2022 e 29/05/2024
(total tabela R$ 20.200,54). Tabela na subpasta KIT/7-TABELA.xlsx aba
"APLIC INVEST FACIL". Extrato APLIC.INVEST FACIL/6-EXTRATO.pdf é
ESCANEADO (29 páginas, sem text-layer) — auditoria APLIC vs RESGATE
NÃO foi possível via parser. Estratégia (b) PADRÃO assumida com base
no padrão histórico do produto APLIC.INVEST FACIL no banco Bradesco.
Pendência crítica: confirmar via OCR ou conferência manual.

CONTA Bradesco: ATENÇÃO — extrato escaneado, número da conta não
extraído automaticamente. Adotar conta da inicial PG ELETRON anterior
(processada em 06/05/2026) ou aguardar leitura manual. Como
PROVISÓRIO, uso conta a partir de informação do checkpoint anterior.
ALERTA no relatório.

Estratégia (b) PADRÃO. VC R$ 15.000. Cabe folgadamente no JEC.

OUTRAS TESES no KIT/7-TABELA.xlsx:
- PAGTO ELETRON COBRANCA BRADESCO VIDA E PREV (já processada — 3 lanç)
- PACOTE DE SERVIÇOS (a processar — 67 lançamentos)
- MORA CREDITO PESSOAL (a processar — 14 lançamentos)
"""
import io, sys, os, copy, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\CLIENTE EXEMPLO DE ALMEIDA - Maurivã bradesco\APLICAÇÃO INVEST FÁCIL'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_APLIC_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_lancs.json')) as f:
    LANCAMENTOS = sorted(json.load(f), key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO DE ALMEIDA', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.022-32', 'rg': '1000020-0',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'CM Pingo de Ouro – Maués Açu', 'numero': 's/nº',
    'bairro': 'N. Sra. Aparecida, Rio Parauary',
    'cidade': 'Maués', 'cep': '69.190-000',
}
conta = {'agencia': '[A CONFIRMAR]', 'numero': '[A CONFIRMAR]'}  # extrato escaneado
renda = {'valor_float': 942.22}  # INSS confirmado da sessão anterior

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Maués', uf='AM')
dados['valor_causa'] = fmt_moeda_rs(15000.00)
dados['valor_causa_extenso'] = extenso_moeda(15000.00)
dados['remuneração'] = 'aposentadoria pelo INSS'

print(f'=== CLIENTE EXEMPLO DE ALMEIDA — APLIC.INVEST (b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | VC: R$ 15.000,00'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
for p in list(doc.paragraphs):
    for m in MARKERS:
        if m in (p.text or ''): p._element.getparent().remove(p._element); break

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'reiterado durante todo o período em que a renda alimentar da parte autora ficou '
        'indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = ('No caso concreto, a tabela anexa registra 23 (vinte e três) ocorrências de aplicação '
      'automática entre 08/03/2022 e 29/05/2024 (mais de 2 anos). Em todos os meses, '
      'parcela substancial da aposentadoria do INSS da parte autora — pessoa idosa de '
      '65 anos — foi automaticamente subtraída pelo banco réu sob a rubrica APLIC.INVEST '
      'FACIL, restando indisponível ao consumidor por período variável até o resgate '
      'manual. Embora os valores tenham sido restituídos via RESGATE INVEST FACIL ao '
      'longo do período, o cerne do dano moral não reside na perda patrimonial líquida — '
      'predominantemente compensada —, mas na privação reiterada da autodeterminação do '
      'consumidor sobre sua própria renda alimentar, mês após mês, durante 2 (dois) anos. '
      'Cada retenção mensal configura, autonomamente, prática abusiva vedada pelo art. 39, '
      'inciso VI, do Código de Defesa do Consumidor, sendo a recorrência sistêmica o fato '
      'gerador do abalo extrapatrimonial — agravado pela hipervulnerabilidade da consumidora '
      'idosa.')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'APLIC.INVEST FACIL — estratégia (b) PADRÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'),
             ('Arquivo', 'INICIAL_AplicInvest_CLIENTE EXEMPLO_v1.docx'),
             ('Observação', 'Cliente já tem inicial PG ELETRON BRADESCO VIDA E PREV processada em 06/05/2026; ESTA é OUTRA tese (APLIC.INVEST FACIL).')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '20/07/1960 — 65 anos (IDOSA)'),
    ('Estado civil', '(omitido)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', '[A CONFIRMAR — extrato escaneado]'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS — confirmado em sessão anterior)'),
    ('Período tabela', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC', '23 (mais de 2 anos)'),
    ('Total bruto APLIC (NÃO pleiteado)', dados['total_descontos']),
    ('Saldo líquido', '[NÃO AUDITADO — extrato escaneado, sem text-layer]'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('CONTA / AGÊNCIA — extrato escaneado, dados NÃO extraídos automaticamente',
     'O extrato 6-EXTRATO.pdf na pasta APLICAÇÃO INVEST FÁCIL é PDF escaneado de 29 páginas '
     'sem text-layer. Não foi possível extrair número de conta/agência via parser. '
     'PENDÊNCIA CRÍTICA: substituir [A CONFIRMAR] na inicial pelos números reais antes do '
     'protocolo. Sugestão: usar a mesma conta da inicial PG ELETRON BRADESCO VIDA E PREV '
     'já processada em 06/05/2026 (verificar no _checkpoint-sessao do vault).'),
    ('AUDITORIA APLIC vs RESGATE — NÃO REALIZADA',
     'Por causa do extrato escaneado, NÃO foi possível auditar APLIC vs RESGATE. A '
     'estratégia (b) padrão foi assumida com base no padrão histórico do produto '
     'APLIC.INVEST FACIL no Bradesco (todas as APLICs são resgatadas em 1-3 dias). '
     'CONFERIR via OCR do extrato OU manualmente antes do protocolo. Se houver perda '
     'líquida significativa, considerar pleito subsidiário em dobro.'),
    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Bloco "Repetição do indébito" REMOVIDO. Parágrafo doutrinário REESCRITO. '
     'Parágrafo do caso concreto INSERIDO (23 retenções, 2 anos, hipervulnerabilidade '
     'idosa). VC = R$ 15.000.'),
    ('OUTRAS 3 TESES no KIT — disponíveis para batch combinado',
     'A planilha KIT/7-TABELA.xlsx tem 4 abas: APLIC INVEST FACIL (23, processada aqui), '
     'PAGTO ELETRON COBRANCA BRADESCO VIDA E PREV (3, JÁ PROCESSADA em 06/05/2026), '
     'PACOTE DE SERVIÇOS (67), MORA CREDITO PESSOAL (14). Sobram 2 teses para batch '
     'separado: PACOTE DE SERVIÇOS (TARIFAS) e MORA CREDITO PESSOAL.'),
    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 ≈ 9,87 SM. Cabe folgadamente no JEC.'),
    ('NOTIFICAÇÃO EXTRAJUDICIAL APLIC — formato .docx',
     'Pasta tem "Notificação Extrajudicial - CLIENTE EXEMPLO DE ALMEIDA - APLICAÇÃO INVEST FÁCIL.docx". '
     'Avaliar envio antes do protocolo.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'PREENCHER conta/agência [A CONFIRMAR] (extrato escaneado).',
           'Conferir comarca: Maués/AM.', 'Conferir prioridade idosa aplicada.',
           'Conferir VC = R$ 15.000,00.', 'OCR ou auditoria manual do extrato escaneado.',
           'Decidir sobre as outras 2 teses (PACOTE + MORA) em batch separado.',
           'Avaliar envio de notificação extrajudicial.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR antes de PREENCHER a conta/agência e auditar o extrato escaneado. '
          'Após esses 2 itens críticos, '); r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
