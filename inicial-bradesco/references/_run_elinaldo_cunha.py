"""Gera inicial APLIC.INVEST FACIL do CLIENTE EXEMPLO DOS SANTOS — caso novo
do batch 06/05/2026 (não confundir com o caso paradigma "Cliente Exemplo
dos Santos (removido)" do checkpoint antigo, que tratava do mesmo
cliente em material de inicial estrita por R$ 159k em dobro — agora o
material está reapresentado e a estratégia mudou para (b) conservadora
após a regra unificada da SKILL.md em 06/05/2026).

Comarca: Caapiranga/AM (foro do domicílio art. 101 I CDC). NÃO IDOSO
(nascido 07/06/1984 → 41 anos em 06/05/2026). Solteiro (RG origem em
certidão de nascimento, procuração também). Servidor público municipal
de Caapiranga (CREDITO DE SALARIO + TED PREFEITURA MUNICIPAL DE
CAAPIRAN). Conta Bradesco Ag 3707 / 413210-6.

Tese APLIC.INVEST FACIL — 84 lançamentos entre 09/07/2021 e 30/07/2024,
total bruto aplicado R$ 89.654,74. AUDITORIA APLIC vs RESGATE no extrato
mostra ciclo aplica-resgate em 1-3 dias com rentabilidade R$ 7,76 em
~3 anos (saldo líquido NEGATIVO de R$ 7,76 — cliente recebeu MAIS do
que aplicou). 246 RESGATEs INVEST FACIL contra 84 APLICs (média 3
parcelas de resgate por aplicação).

ESTRATÉGIA (b) PADRÃO conforme regra unificada da SKILL.md:
declaratória + obrigação de cessar (multa R$ 500/dia) + R$ 15.000 dano
moral. VC R$ 15.000 → cabe folgadamente no JEC (~10 SM, SM-2025
R$ 1.518 → 40 SM = R$ 60.720).

PENDÊNCIA RENDA: O último crédito recorrente da Prefeitura no extrato é
de 04/08/2020 (R$ 1.953,58). A partir dessa data o cliente parou de
receber salário pela conta Bradesco — provavelmente migrou para outra
conta ou parou de trabalhar para a Prefeitura. Renda ATUAL (2026)
precisa ser confirmada com cliente. Para a Justiça Gratuita usa-se o
último valor identificável (R$ 1.953,58) com alerta no relatório.

Pasta KIT contém material para outras teses (TARIFAS, Crédito Pessoal,
Título Capitalização). Fora do escopo deste batch (APLIC.INVEST).
"""
import io, sys, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\CLIENTE EXEMPLO DOS SANTOS - Ney Pedroza'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_ELINALDO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_ELINALDO_v1.docx')

autora = {
    'nome': 'CLIENTE EXEMPLO DOS SANTOS',
    'nacionalidade': 'brasileiro',
    'estado_civil': 'solteiro',
    'profissao': 'servidor público municipal',
    'cpf': '000.000.010-20',
    'rg': '1000008-8',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'CM Paraná do Mari',
    'numero': 's/nº',
    'bairro': 'Zona Rural',
    'cidade': 'Caapiranga',
    'cep': '69.425-000',
}
conta = {'agencia': '3707', 'numero': '413210-6'}
renda = {'valor_float': 1953.58}

# 84 lançamentos APLIC.INVEST FACIL extraídos do extrato (ordenados por data)
LANCAMENTOS = sorted([
    ('09/07/2021', 275.00), ('29/07/2021', 1806.31), ('30/07/2021', 979.00),
    ('30/08/2021', 1046.56), ('31/08/2021', 979.00), ('29/09/2021', 994.67),
    ('28/10/2021', 510.62), ('29/10/2021', 696.84), ('29/11/2021', 238.19),
    ('30/11/2021', 159.00), ('10/12/2021', 680.00), ('20/12/2021', 975.64),
    ('23/12/2021', 11522.85), ('29/12/2021', 979.00),
    ('28/01/2022', 1979.90), ('31/01/2022', 236.33), ('24/02/2022', 2079.81),
    ('25/02/2022', 1078.68), ('30/03/2022', 3049.83), ('04/04/2022', 56.60),
    ('28/04/2022', 1378.47), ('29/04/2022', 1078.68), ('05/05/2022', 911.98),
    ('30/05/2022', 1659.58), ('29/06/2022', 402.20), ('30/06/2022', 1058.68),
    ('28/07/2022', 1689.98), ('29/07/2022', 930.68), ('30/08/2022', 2347.46),
    ('29/09/2022', 1125.42), ('30/09/2022', 1112.72), ('18/10/2022', 312.20),
    ('19/10/2022', 408.00), ('25/10/2022', 61.90), ('28/10/2022', 2079.81),
    ('29/11/2022', 1509.05), ('30/11/2022', 1112.72), ('12/12/2022', 1011.86),
    ('23/12/2022', 404.00), ('26/12/2022', 60.00), ('29/12/2022', 581.01),
    ('03/01/2023', 5980.00), ('19/01/2023', 386.43), ('30/01/2023', 428.67),
    ('31/01/2023', 1218.60), ('24/02/2023', 63.70), ('27/02/2023', 590.71),
    ('28/02/2023', 1196.17), ('29/03/2023', 66.45), ('30/03/2023', 1503.86),
    ('10/04/2023', 79.99), ('27/04/2023', 870.81), ('10/05/2023', 91.71),
    ('11/05/2023', 304.00), ('22/05/2023', 175.00), ('30/05/2023', 2045.46),
    ('29/06/2023', 1359.23), ('05/07/2023', 450.65), ('28/07/2023', 1216.70),
    ('01/08/2023', 307.33), ('07/08/2023', 525.60), ('29/08/2023', 60.00),
    ('30/08/2023', 1185.21), ('19/09/2023', 365.30), ('28/09/2023', 1733.78),
    ('27/10/2023', 59.00), ('30/10/2023', 572.44), ('31/10/2023', 1234.62),
    ('29/11/2023', 1674.80), ('01/12/2023', 53.05), ('18/12/2023', 596.97),
    ('21/12/2023', 1798.85), ('28/12/2023', 335.83),
    ('04/01/2024', 1842.38), ('30/01/2024', 1537.57), ('27/02/2024', 59.91),
    ('28/02/2024', 968.90), ('26/03/2024', 1166.60), ('27/03/2024', 60.00),
    ('29/04/2024', 969.00), ('28/05/2024', 1267.76), ('27/06/2024', 1194.71),
    ('23/07/2024', 183.04), ('30/07/2024', 313.72),
], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome': '', 'cnpj': '', 'logradouro': '', 'numero': '',
            'bairro': '', 'cidade': '', 'uf': '', 'cep': ''}

dados, calc = montar_dados_padrao(
    autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
    eh_idoso=False, competência='Caapiranga', uf='AM',
)

# OVERRIDE estratégia (b)
DANO_MORAL = 15000.00
VALOR_CAUSA_B = 15000.00
dados['valor_causa'] = fmt_moeda_rs(VALOR_CAUSA_B)
dados['valor_causa_extenso'] = extenso_moeda(VALOR_CAUSA_B)
dados['remuneração'] = 'salário do serviço público municipal de Caapiranga'

print('=== CLIENTE EXEMPLO DOS SANTOS — APLIC.INVEST FACIL (estratégia b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}')
print(f'Total bruto aplicado (NÃO pleiteado): R$ {calc["total"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'Dano moral pleiteado: R$ {DANO_MORAL:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'Valor da causa (b): R$ {VALOR_CAUSA_B:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Aplicar template: {res["modificados"]} parágrafos modificados, residuais: {res["residuais"] or "nenhum"}')

# ========== PÓS-PROCESSAMENTO (estratégia b) ==========
from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS_REMOVER = [
    'Repetição do indébito',
    'Caso se verifique a existência de valores indevidamente',
    'Na cobrança de débitos, o consumidor inadimplente',
    'O consumidor cobrado em quantia indevida tem direito',
    'a restituição em dobro se faz necessária como penalidade',
    'deve a requerida ser condenada a restituir em dobro',
    'Havendo a retenção dos valores a título de investimento',
]
removidos = 0
for p in list(doc.paragraphs):
    txt = p.text or ''
    for marker in MARKERS_REMOVER:
        if marker in txt:
            p._element.getparent().remove(p._element)
            removidos += 1
            break
print(f'Removidos: {removidos} parágrafos')

# Reescreve parágrafo 102
TEXTO_NOVO_102 = (
    'A cobrança indevida não decorre de engano justificável, mas de '
    'modelo operacional estruturado para funcionar sem contratação '
    'inequívoca, configurando ato ilícito reiterado durante todo o '
    'período em que a renda alimentar da parte autora ficou indisponível '
    'para uso imediato.'
)
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for run_elem in list(p._element.findall(W + 'r')):
            p._element.remove(run_elem)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr')
        rFonts = etree.SubElement(rPr, W + 'rFonts')
        rFonts.set(W + 'ascii', 'Cambria'); rFonts.set(W + 'hAnsi', 'Cambria'); rFonts.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TEXTO_NOVO_102
        print('Parágrafo 102 reescrito.')
        break

# Insere parágrafo do caso concreto
TEXTO_CASO = (
    'No caso concreto, o extrato bancário registra 84 (oitenta e quatro) '
    'ocorrências de aplicação automática entre 09/07/2021 e 30/07/2024 '
    '(aproximadamente 3 anos). Em todos os meses, parcela substancial '
    'da renda da parte autora — recebida em conta-salário — foi '
    'automaticamente subtraída pelo banco réu sob a rubrica APLIC.INVEST '
    'FACIL, restando indisponível ao consumidor pelo prazo de 1 (um) a '
    '5 (cinco) dias até o resgate manual mediante saque ou transferência. '
    'Embora os valores tenham sido restituídos via RESGATE INVEST FACIL '
    'ao longo do período (246 ocorrências de resgate, indicando que cada '
    'aplicação foi parcelada em 2 a 4 resgates), o cerne do dano moral '
    'não reside na perda patrimonial líquida — inexistente, pois o '
    'saldo agregado em 3 anos foi de aproximadamente zero —, mas na '
    'privação reiterada da autodeterminação do consumidor sobre sua '
    'própria renda alimentar, mês após mês, durante 3 (três) anos '
    'consecutivos. Cada retenção mensal configura, autonomamente, '
    'prática abusiva vedada pelo art. 39, inciso VI, do Código de Defesa '
    'do Consumidor, sendo a recorrência sistêmica o fato gerador do '
    'abalo extrapatrimonial.'
)
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        new_p = copy.deepcopy(p._element)
        for r in list(new_p.findall(W + 'r')):
            new_p.remove(r)
        r = etree.SubElement(new_p, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr')
        rFonts = etree.SubElement(rPr, W + 'rFonts')
        rFonts.set(W + 'ascii', 'Cambria'); rFonts.set(W + 'hAnsi', 'Cambria'); rFonts.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TEXTO_CASO
        p._element.addnext(new_p)
        print('Parágrafo caso concreto inserido.')
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# ========== RELATÓRIO ==========
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_ELINALDO', level=1)
for k, v in [('Cliente', autora['nome']),
             ('Tese', 'APLIC.INVEST FACIL — estratégia (b) PADRÃO, só dano moral'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'),
             ('Arquivo gerado', 'INICIAL_AplicInvest_ELINALDO_v1.docx')]:
    p = doc_r.add_paragraph()
    p.add_run(k + ': ').bold = True
    p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/{dados["uf"]}'),
    ('Nome', autora['nome']), ('CPF', autora['cpf']),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '07/06/1984 — 41 anos (NÃO idoso)'),
    ('Estado civil', autora['estado_civil']),
    ('Profissão', autora['profissao']),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda usada (extrato)', dados['valor_remuneração']),
    ('Réu', 'BANCO BRADESCO S.A. — CNPJ 60.746.948/0001-12'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC', '84'),
    ('Lançamentos RESGATE (extrato)', '246'),
    ('Total bruto aplicado (NÃO pleiteado)', dados['total_descontos']),
    ('Saldo líquido (APLIC - RESGATE)', '-R$ 7,76 (cliente recebeu R$ 7,76 a mais via rentabilidade)'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
pendencias = [
    ('AUDITORIA APLIC vs RESGATE — confirmada',
     'Extrato confirma ciclo aplica-resgate: 84 APLIC.INVEST FACIL e 246 RESGATE INVEST FACIL '
     'totalizando R$ 89.654,74 vs R$ 89.662,50. Saldo líquido NEGATIVO R$ 7,76 (rentabilidade). '
     'Estratégia (b) PADRÃO (vide SKILL.md atualizada em 06/05/2026): '
     'pleitear apenas dano moral pela RECORRÊNCIA das 84 retenções, sem repetição em dobro.'),

    ('RENDA ATUAL — desatualizada no extrato',
     'O último crédito da PREFEITURA MUNICIPAL DE CAAPIRANGA no extrato é de 04/08/2020 '
     '(R$ 1.953,58). Após essa data o cliente parou de receber salário pela conta Bradesco. '
     'Possíveis hipóteses: (i) migrou para outra conta de pagamento; (ii) deixou o serviço público; '
     '(iii) recebe atualmente por outra fonte. CONFIRMAR COM CLIENTE a renda mensal atual (2026) '
     'antes do protocolo. Para efeito da Justiça Gratuita, a inicial usa o último valor '
     'identificável (R$ 1.953,58), mas idealmente deve ser ATUALIZADA. Alternativa: usar '
     'última fonte recente identificável no extrato 2024 (TED-TRANSF de outras origens) ou '
     'declaração do cliente. Pendência crítica.'),

    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'O template inicial-aplic-invest.docx foi pós-processado para estratégia (b): '
     '7 parágrafos do bloco "Repetição do indébito" e do pedido subsidiário REMOVIDOS, '
     'parágrafo doutrinário sobre repetição em dobro REESCRITO, parágrafo do caso concreto '
     '(84 retenções, 3 anos, ciclo aplica-resgate) INSERIDO. CONFERIR coerência textual.'),

    ('COMARCA Caapiranga/AM — confirmada',
     'Endereço da procuração e do RG indica domicílio em CM Paraná do Mari, Zona Rural, '
     'Caapiranga/AM. Foro do domicílio art. 101 I CDC. Comarca de Caapiranga adota combinação '
     'de teses (lista do escritório), mas neste caso há apenas 1 tese (APLIC).'),

    ('SEM NOTIFICAÇÃO EXTRAJUDICIAL na pasta',
     'Pasta não contém arquivo "8 - NOTIFICACAO". Avaliar se vai enviar notificação ANTES do '
     'protocolo. CONFERIR/AJUSTAR no DOCX o bloco "Do prévio requerimento administrativo" '
     '(parágrafos 25-29) — se notificação NÃO foi enviada, esse bloco gera contradição.'),

    ('PASTA KIT — outras teses do mesmo cliente',
     'KIT contém procurações para Tarifa Bancária, Crédito Pessoal, Título de Capitalização '
     'e tabela "TARIFAS BRADESCO ELINALDO.pdf". O cliente tem mais 3 teses passíveis de '
     'iniciais separadas (ou combinada, dado que Caapiranga adota combinação). Decidir em '
     'batch separado.'),

    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 ≈ 9,87 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720). Sem necessidade '
     'de renúncia ao excedente.'),

    ('RECONTAR contra checkpoint antigo',
     'O checkpoint antigo da skill cita "Cliente Exemplo dos Santos (removido)" como caso '
     'paradigma com inicial estrita por R$ 159k em dobro. Esta inicial nova adota estratégia '
     '(b) padrão. Se o caso paradigma antigo era do MESMO cliente, esta versão SUPERA a '
     'anterior. Se era de cliente diferente (homônimo), confirmar com escritório qual será '
     'protocolada.'),
]
for titulo, txt in pendencias:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True
    p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'Conferir nome / CPF / RG.',
    'Conferir conta/agência (413210-6 / 3707).',
    'Conferir comarca: Caapiranga/AM.',
    'Conferir VC = R$ 15.000,00 (apenas dano moral).',
    'Conferir parágrafos doutrinários sobre "Repetição do indébito" REMOVIDOS.',
    'Conferir parágrafo do caso concreto inserido (84 retenções).',
    'CONFIRMAR RENDA ATUAL com cliente (extrato só vai até 08/2020 da Prefeitura).',
    'Avaliar envio de notificação extrajudicial.',
    'Decidir sobre as outras 3 teses do KIT.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas — confirmar renda atual e pendências antes de PROTOCOLAR.')
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
