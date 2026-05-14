"""Inicial MORA + ENCARGOS — CLIENTE EXEMPLO PEREIRA DA SILVA GUERRA.

Mesma autora da TARIFAS (Presidente Figueiredo/AM, IDOSO, INSS R$ 1.091,86,
Ag 3732 Conta 1525-3). Documento: CNH (não RG). Caso recente.

Tabela: 17 lançamentos (11 MORA + 6 ENCARGO) entre 31/07/2025 e
30/01/2026. Total R$ 1.976,48 / dobro R$ 3.952,96. VC R$ 18.952,96.

Caso recente — sem questão prescricional.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal, extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\2. MORA\CLIENTE EXEMPLO PEREIRA DA SILVA GUERRA - Ruth - TARIFAS\MORA CRED E ENCARGOS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora-encargo.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_MoraEncargo_LUIZ_FABIO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_LUIZ_FABIO_MORA_v1.docx')

ENCARGO = [
    ('31/10/2025', 0.86), ('28/11/2025', 28.06), ('30/12/2025', 24.74),
    ('06/01/2026', 0.78), ('06/01/2026', 25.39), ('30/01/2026', 3.27),
]
MORA = [
    ('31/07/2025', 142.86), ('03/09/2025', 145.86), ('30/09/2025', 139.48),
    ('30/09/2025', 252.38), ('31/10/2025', 4.56), ('31/10/2025', 144.10),
    ('31/10/2025', 253.65), ('28/11/2025', 249.57), ('30/12/2025', 142.15),
    ('30/12/2025', 166.36), ('30/12/2025', 252.41),
]
LANCAMENTOS = sorted(ENCARGO + MORA, key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'LUIZ FÁBIO PEREIRA DA SILVA GUERRA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': '',
    'cpf': '000.000.016-26',
    'rg': '1000014-4',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Manoel (RM) do Urubuí',
    'numero': '23 A - KM 01',
    'bairro': 'Centro',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '1525-3'}
renda = {'valor_float': 1091.86}

tese = {'rubrica': 'MORA CRED PESS / ENC LIM CRÉDITO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'benefício do INSS'
dados['rubrica_curta'] = 'Mora Cred Pess / Enc Lim Crédito'
dados['rubrica_curta_caps'] = 'MORA CRED PESS / ENC LIM CRÉDITO'
dados['rubrica_completa'] = 'Crédito Mora Pessoal / Encargos Limite de Crédito'
dados['rubrica_completa_caps'] = 'MORA CREDITO PESSOAL / ENCARGOS LIMITE DE CRÉDITO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — MORA + ENCARGOS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
def fmt_moeda(v):
    s = f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')
    return f'R$ {s}'
te=sum(v for _,v in ENCARGO); tm=sum(v for _,v in MORA)

with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {
    '{{remuneração}}': dados['remuneração'],
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
    '{{cidade_filial}}': 'Maués',
    '{{uf_filial}}': 'AM',
    '{{uf_extenso}}': 'Amazonas',
    '{{numero_desconto_encargo}}': str(len(ENCARGO)),
    '{{numero_desconto_mora}}': str(len(MORA)),
    '{{total_encargo}}': fmt_moeda(te),
    '{{total_mora}}': fmt_moeda(tm),
    '{{total_encargo_extenso}}': extenso_moeda(te),
    '{{total_mora_extenso}}': extenso_moeda(tm),
    '{{desconto_extenso_encargo}}': extenso_cardinal(len(ENCARGO)),
    '{{desconto_extenso_mora}}': extenso_cardinal(len(MORA)),
}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_MoraEncargo_LUIZ_FABIO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA + ENCARGOS LIMITE'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_MoraEncargo_LUIZ_FABIO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'NÃO IDOSO (44 anos — CNH 23/10/1981)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CNH)', '18138858 SSP/AM'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 30/01/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} ({len(ENCARGO)} encargo + {len(MORA)} mora)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('Caso RECENTE — sem prescrição',
     'Tabela cobre apenas 31/07/2025 a 30/01/2026. Nenhuma questão prescricional.'),
    ('NÃO IDOSO (44 anos) — prioridade REMOVIDA',
     'CNH mostra nascimento 23/10/1981. Notificação errou ao afirmar "pessoa idosa". '
     'Prioridade do art. 1.048 CPC FOI REMOVIDA.'),
    ('RG CORRIGIDO — 18138858 SSP/AM (CNH)',
     'Notificação trouxe RG 16138856; CNH mostra 18138858 SSP/AM. Inicial corrigida.'),
    ('CLIENTE TEM 2 TESES SEPARADAS',
     'Cliente também aparece em TARIFAS (já gerada). AVALIAR consolidação em '
     'inicial-combinada (Presidente Figueiredo adota combinação por padrão).'),
    ('VALORES MORA ELEVADOS',
     'Lançamentos MORA chegam a R$ 252-253. Indica cheque especial alto. Conferir.'),
    ('SEM 8.1 - COMPROVANTE NOTIFICAÇÃO',
     'Pasta não tem 8.1-COMPROVANTE NOTIFICAÇÃO. Verificar envio.'),
    ('TETO JEC — coberto',
     'VC R$ 18.952,96 ≈ 12,5 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AVALIAR consolidação com TARIFAS em 1 só inicial-combinada.',
    'Conferir nome / CPF / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Verificar comprovante de envio da notificação.',
    'Anexar CNH como documento de identificação.',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Confirmar com cliente: nunca contratou cheque especial.',
    'Anexar 2-Procurações (ENC LIM + MORA CRED PESS) + 3-CNH + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após decidir consolidação, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
