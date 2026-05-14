"""Inicial TARIFAS — CLIENTE EXEMPLO CAVALCANTE SANTANA.

Comarca Presidente Figueiredo/AM (Ag 3732 / Conta 510965-5). Pessoa
IDOSA (afirmação expressa na notificação). Aposentada pelo INSS
(R$ 1.621,00 em 02/02/2026 — último crédito).

Tabela: 20 lançamentos (TARIFA CESTA B EXPRESSO + PACOTE SERVIÇOS)
entre 15/09/2022 e 16/09/2025. Total R$ 310,03 / dobro R$ 620,06.
VC R$ 15.620,06 — cabe folgadamente no JEC.

PENDÊNCIAS: (1) Notificação repete CPF como número de RG — número real
do RG precisa ser extraído manualmente do 3-RG.pdf. (2) Estado civil
omitido. (3) Cliente também tem TÍTULO CAPITALIZAÇÃO em pasta separada.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO CAVALCANTE SANTANA - Ruth - TARIFAS\TARIFAS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

LANCAMENTOS = [
    ('15/09/2022', 15.16), ('14/04/2023', 9.48), ('08/05/2023', 40.42),
    ('15/05/2023', 3.58), ('07/06/2023', 46.32), ('15/06/2023', 8.68),
    ('07/07/2023', 41.22), ('07/08/2023', 15.45), ('15/08/2023', 3.98),
    ('08/09/2023', 11.47), ('15/12/2023', 8.67), ('08/01/2024', 6.78),
    ('15/01/2024', 1.22), ('27/05/2025', 15.95), ('27/05/2025', 15.95),
    ('13/06/2025', 8.10), ('16/09/2025', 16.35), ('16/09/2025', 16.35),
    ('16/09/2025', 16.65), ('16/09/2025', 8.25),
]

autora = {
    'nome': 'CLIENTE EXEMPLO CAVALCANTE SANTANA',
    'nacionalidade': 'brasileira',
    'estado_civil': '',
    'profissao': 'aposentada',
    'cpf': '000.000.015-25',
    'rg': '1000013-3',
    'orgao_expedidor_prefixo': 'CIN (CPF como Registro Geral)',
    'logradouro': 'Rua Manoel (RM) Jardim Floresta',
    'numero': 's/nº',
    'bairro': 'Centro',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '510965-5'}
renda = {'valor_float': 1621.00}

tese = {
    'rubrica': 'TARIFA BANCÁRIA - CESTA B EXPRESSO / PACOTE DE SERVIÇOS PADRONIZADO PRIORITÁRIOS',
    'lancamentos': LANCAMENTOS,
}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA B EXPRESSO / PACOTE DE SERVIÇOS PADRONIZADO PRIORITÁRIOS'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFAS (2 rubricas)'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CIN nova)', f'{autora["rg"]} — Registro Geral = CPF (Carteira Identidade Nacional)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 02/02/2026 = 1 SM)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (2 rubricas)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('RG (CIN nova) — Registro Geral = CPF',
     'O RG da autora é a nova Carteira de Identidade Nacional (CIN), em que o Registro '
     'Geral é o próprio número do CPF (000.000.015-25). A notificação NÃO errou — apenas '
     'replicou esse padrão. Inicial corrigida com RG = CPF e órgão expedidor "CIN".'),
    ('IDOSA CONFIRMADA — 60 anos (13/03/1966)',
     'RG mostra nascimento em 13/03/1966 → 60 anos completos em 07/05/2026. Prioridade '
     'art. 1.048, I, CPC mantida. Naturalidade Manaus/AM.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente.'),
    ('Renda 1 SM — hipossuficiência reforçada',
     'INSS R$ 1.621,00 = exatamente 1 salário mínimo (2026). Hipossuficiência clara. '
     'Pode reforçar o impacto do dano moral.'),
    ('CLIENTE TEM 2 TESES SEPARADAS',
     'Cliente CLIENTE EXEMPLO também aparece na pasta TÍTULO DE CAPITALIZAÇÃO. Comarca '
     'Presidente Figueiredo adota INICIAL COMBINADA por padrão (ver § 4 da SKILL). '
     'Procurador organizou em pastas separadas — gerando 2 iniciais individuais. '
     'AVALIAR consolidação em 1 só inicial-combinada antes do protocolo.'),
    ('VALOR BAIXO — combinação aconselhada',
     'Dobro R$ 620,06 é abaixo do limite R$ 400 (skill recomenda combinar). '
     'Consolidação com TÍTULO ajudaria a fortalecer o pleito. CONFERIR.'),
    ('NOTIFICAÇÃO previa dano moral R$ 5.000',
     'A notificação pleiteou R$ 5.000 (provavelmente projetando 2 teses combinadas). '
     'Como esta inicial é ISOLADA, pleiteia R$ 15.000 conforme regra § 9 da skill. '
     'CONFIRMAR com procurador.'),
    ('TETO JEC — coberto',
     'VC R$ 15.620,06 ≈ 10,3 SM. Cabe folgadamente no JEC.'),
    ('LANÇAMENTOS DE 16/09/2025 e 27/05/2025 — duplicidades',
     'Em 16/09/2025 a tabela registra 4 lançamentos (R$ 16,35×2 + R$ 16,65 + R$ 8,25). '
     'Em 27/05/2025 registra 2× R$ 15,95. Conferir extrato original — pode ser '
     'cobrança em lote ou erro de digitação.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'Conferir nome / CPF (000.000.015-25 = RG via CIN nova) / nascimento 13/03/1966 (60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'AVALIAR consolidação com TÍTULO CAPITALIZAÇÃO em 1 só inicial-combinada.',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Conferir VC = R$ 15.620,06 e dano moral R$ 15.000,00.',
    'Confirmar com cliente: nunca contratou pacote de serviços ou cesta de tarifas.',
    'Anexar 2 procurações específicas + RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — RG (CIN nova) confirmado, 60 anos confirmados. Após decidir consolidação, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
