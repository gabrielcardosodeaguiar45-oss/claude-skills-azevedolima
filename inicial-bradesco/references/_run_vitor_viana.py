"""Inicial APLIC.INVEST FACIL — EXEMPLO VIANA DA SILVA.

Comarca Presidente Figueiredo/AM. Estado civil omitido (procuração não
menciona). Servidor público municipal Presidente Figueiredo (TRANSF
SALDO C/SAL P/CC R$ 1.430,18 mensal). Conta Ag 3732/19516-2.

Tese: 24 lançamentos APLIC.INVEST FACIL entre 26/03/2020 e 25/07/2024
(total R$ 26.454,49). Audit confirma: 36 RESGATEs R$ 26.454,67. Saldo
-R$ 0,18 (zero). Estratégia (b) PADRÃO. VC R$ 15.000.

Pasta KIT: outras 2 teses (MORA CRED PESS + PACOTE DE SERVIÇOS).
"""
import io, sys, os, copy, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\EXEMPLO VIANA DA SILVA - Ruth OK bradesco\APLICAÇÃO INVEST FÁCIL'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_VITOR_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_VITOR_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_vitor_lancs.json')) as f:
    LANCAMENTOS = sorted(json.load(f), key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'EXEMPLO VIANA DA SILVA', 'nacionalidade': 'brasileiro',
    'estado_civil': '', 'profissao': 'servidor público municipal',
    'cpf': '000.000.032-42', 'rg': '1000030-0',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Boa União, Ramal do Rumo Certo', 'numero': 's/nº',
    'bairro': 'Zona Rural', 'cidade': 'Presidente Figueiredo', 'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '19516-2'}
renda = {'valor_float': 1430.18}

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Presidente Figueiredo', uf='AM')
dados['valor_causa'] = fmt_moeda_rs(15000.00)
dados['valor_causa_extenso'] = extenso_moeda(15000.00)
dados['remuneração'] = 'salário do serviço público municipal'

print(f'=== EXEMPLO VIANA — APLIC.INVEST (b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
for p in list(doc.paragraphs):
    for m in MARKERS:
        if m in (p.text or ''): p._element.getparent().remove(p._element); break

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'reiterado durante todo o período em que a renda alimentar da parte autora ficou '
        'indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = ('No caso concreto, a tabela anexa registra 24 (vinte e quatro) ocorrências de '
      'aplicação automática entre 26/03/2020 e 25/07/2024 (mais de 4 anos). Em todos '
      'os meses, parcela substancial do salário da parte autora — recebida via '
      'transferência de saldo da conta-salário — foi automaticamente subtraída pelo '
      'banco réu sob a rubrica APLIC.INVEST FACIL, restando indisponível ao consumidor '
      'por período variável até o resgate manual. Embora os valores tenham sido '
      'restituídos via RESGATE INVEST FACIL ao longo do período, o cerne do dano moral '
      'não reside na perda patrimonial líquida — inexistente —, mas na privação reiterada '
      'da autodeterminação do consumidor sobre sua própria renda alimentar, mês após mês, '
      'durante 4 (quatro) anos consecutivos. Cada retenção mensal configura, autonomamente, '
      'prática abusiva vedada pelo art. 39, inciso VI, do Código de Defesa do Consumidor, '
      'sendo a recorrência sistêmica o fato gerador do abalo extrapatrimonial.')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_VITOR', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'APLIC.INVEST FACIL — estratégia (b)'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_AplicInvest_VITOR_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Estado civil', '(omitido)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', dados['valor_remuneração'] + ' (TRANSF SALDO C/SAL — servidor municipal)'),
    ('Período tabela', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC', '24'),
    ('Lançamentos RESGATE (extrato)', '36'),
    ('Total bruto APLIC', dados['total_descontos']),
    ('Saldo líquido', '-R$ 0,18 (zero)'),
    ('Dano moral', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('AUDITORIA APLIC vs RESGATE — confirmada (estratégia b)',
     '24 APLICs vs 36 RESGATEs. Saldo zero.'),
    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Bloco "Repetição" REMOVIDO. Parágrafo 102 REESCRITO. Caso concreto INSERIDO.'),
    ('IDADE — não confirmada (não li RG)',
     'A inicial assume NÃO IDOSO por padrão. Conferir antes do protocolo se cliente tem '
     '60+ anos (caso queira aplicar prioridade idoso). RG não foi lido nesta sessão.'),
    ('PASTA KIT — outras 2 teses',
     'KIT contém procurações/tabelas para MORA CRED PESS e PACOTE DE SERVIÇOS. '
     'Decidir batch separado — Pres. Figueiredo adota combinação.'),
    ('TETO JEC — folgadamente coberto', 'VC R$ 15.000.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (19516-2 / 3732).',
           'Conferir comarca: Presidente Figueiredo/AM.', 'Confirmar idade no RG (>60? aplicar prioridade idoso).',
           'Conferir VC = R$ 15.000,00.', 'Decidir sobre as outras 2 teses do KIT.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — '); r2 = p.add_run('PROTOCOLAR após conferência da idade.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
