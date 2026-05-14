"""Inicial MORA + ENCARGO LIMITE DE CRED — CLIENTE EXEMPLO BRAZÃO NOGUEIRA FILHO.

Comarca Caapiranga/AM. Servidor público municipal Caapiranga (TRANSF
SALDO C/SAL P/CC R$ 2.068,30 mensal + R$ 500 bolsa adicional). Conta
Bradesco Ag 3707/413898-8. CIN moderna (CPF=RG=000.000.025-35).

Tese: MORA CRED PESSOAL (29 lançamentos = R$ 4.343,82) + ENCARGOS
LIMITE DE CRED (9 lançamentos = R$ 1.646,86). Total R$ 5.990,68 / dobro
R$ 11.981,36. Dano moral R$ 15.000 (Mora+Encargo é 1 só tese pelo IRDR).
VC R$ 26.981,36 — cabe no JEC.

Template: inicial-mora-encargo.docx.

PENDÊNCIA: idade não confirmada (RG não lido). Assumir NÃO IDOSO.
"""
import io, sys, os, copy, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\CLIENTE EXEMPLO BRAZÃO NOGUEIRA FILHO - Ney Pedroza'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora-encargo.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_MoraEncargo_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_lancs.json')) as f:
    data = json.load(f)
mora_lancs = sorted(data['mora'], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))
encargo_lancs = sorted(data['encargo'], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

# Combinar todas as datas para fim/inicio
todos = mora_lancs + encargo_lancs
todos = sorted(todos, key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO BRAZÃO NOGUEIRA FILHO', 'nacionalidade': 'brasileiro',
    'estado_civil': '', 'profissao': 'servidor público municipal',
    'cpf': '000.000.025-35', 'rg': '1000023-3',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Paraná do Mari', 'numero': 's/nº',
    'bairro': 'Zona Rural', 'cidade': 'Caapiranga', 'cep': '69.425-000',
}
conta = {'agencia': '3707', 'numero': '413898-8'}
renda = {'valor_float': 2068.30}

tese = {'rubrica': 'MORA CRED PESS / ENC LIM CRED', 'lancamentos': todos}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Caapiranga', uf='AM')
# Override VC: mora+encargo é 1 só tese, dobro pleiteado
DANO_MORAL = 15000.00
DOBRO = calc['dobro']
VC = DOBRO + DANO_MORAL
dados['valor_causa'] = fmt_moeda_rs(VC)
dados['valor_causa_extenso'] = extenso_moeda(VC)
dados['remuneração'] = 'salário do serviço público municipal'

# Override placeholders específicos do template mora-encargo
# Vou injetar separações de mora e encargo
total_mora = sum(v for _, v in mora_lancs)
total_enc = sum(v for _, v in encargo_lancs)
dados['rubrica_curta'] = 'Mora Cred Pess + Enc. Lim. Crédito'
dados['rubrica_curta_caps'] = 'MORA CRED PESS / ENCARGOS LIMITE DE CRÉDITO'
dados['rubrica_completa'] = 'Mora Crédito Pessoal + Encargos Limite de Crédito'
dados['rubrica_completa_caps'] = 'MORA CRÉDITO PESSOAL E ENCARGOS LIMITE DE CRÉDITO'

print(f'=== CLIENTE EXEMPLO — MORA+ENCARGO ===')
print(f'MORA: {len(mora_lancs)} = R$ {total_mora:.2f}')
print(f'ENCARGO: {len(encargo_lancs)} = R$ {total_enc:.2f}')
print(f'Total: R$ {calc["total"]:.2f} | Dobro: R$ {DOBRO:.2f} | VC: R$ {VC:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Sem pós-processamento — esse template é específico para Mora+Encargo, mantém repetição em dobro

# Pós-fix raw para placeholders com encoding
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
substituicoes_raw = {
    '{{remuneração}}': dados.get('remuneração', 'salário'),
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
    '{{total_mora}}': fmt_moeda_rs(total_mora),
    '{{total_mora_extenso}}': extenso_moeda(total_mora),
    '{{total_encargo}}': fmt_moeda_rs(total_enc),
    '{{total_encargo_extenso}}': extenso_moeda(total_enc),
    '{{numero_desconto_mora}}': str(len(mora_lancs)),
    '{{numero_desconto_encargo}}': str(len(encargo_lancs)),
    '{{desconto_extenso_mora}}': 'vinte e nove' if len(mora_lancs) == 29 else str(len(mora_lancs)),
    '{{desconto_extenso_encargo}}': 'nove' if len(encargo_lancs) == 9 else str(len(encargo_lancs)),
    '{{cidade_filial}}': 'Caapiranga',
    '{{uf_filial}}': 'AM',
    '{{uf_extenso}}': 'Amazonas',
}
fix = 0
for k, v in substituicoes_raw.items():
    if k in xml:
        cnt = xml.count(k)
        xml = xml.replace(k, v)
        fix += cnt
        print(f'Pós-fix XML: {k} ({cnt}x)')
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes:
        z.writestr(n, buf[n])
print(f'Pós-fix concluído: {fix} substituições')
print(f'OK -> {DOCX_OUT}')

# Relatório
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_MoraEncargo_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CRED PESSOAL + ENCARGOS LIMITE DE CRÉDITO (1 só tese — IRDR)'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_MoraEncargo_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CIN)', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Estado civil', '(omitido)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', dados['valor_remuneração'] + ' (TRANSF SALDO C/SAL — servidor municipal)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('MORA CRED PESSOAL', f'29 lançamentos = R$ {total_mora:.2f}'.replace('.', ',')),
    ('ENCARGOS LIMITE CRED', f'9 lançamentos = R$ {total_enc:.2f}'.replace('.', ',')),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro (pleiteado)', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('IDADE — não confirmada (RG não lido nesta sessão)',
     'Inicial assume NÃO IDOSO por padrão. Conferir RG se 60+ anos antes do protocolo.'),
    ('RENDA dupla — confirmar segundo crédito',
     'Cliente recebe 2 transferências mensais: R$ 2.068,30 (TRANSF SALDO C/SAL) + R$ 500,00 '
     '(TRANSF SALDO C/SAL menor). Possíveis hipóteses: bolsa adicional, complemento, '
     'pensão, etc. Inicial usa apenas R$ 2.068,30 como renda principal. Confirmar.'),
    ('PROCURAÇÃO MORA + ENCARGO',
     'Há 2 procurações específicas (MORA CRED PESSOAL + ENCARGOS LIMITE DE CRED) — '
     'ambas devem ser anexadas. KIT contém procuração APLIC INVEST FÁCIL adicional, '
     'fora do escopo deste batch.'),
    ('TETO JEC — coberto',
     'VC R$ 26.981,36 ≈ 17,77 SM. Cabe no JEC (40 SM = R$ 60.720).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (413898-8 / 3707).',
           'Conferir comarca: Caapiranga/AM.', 'Confirmar idade no RG.',
           f'Conferir VC = R$ {VC:.2f}'.replace('.', ','),
           'Anexar AMBAS as procurações (MORA + ENCARGO).',
           'Avaliar tese APLIC INVEST FÁCIL do KIT em batch separado.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — '); r2 = p.add_run('PROTOCOLAR após confirmar idade.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
