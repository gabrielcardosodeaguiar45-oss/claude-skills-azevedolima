"""Inicial APLIC.INVEST FACIL — EXEMPLA RAIMUNDA CRUZ.

Comarca Caapiranga/AM. NÃO idosa (55 anos, nascida 07/10/1970).
Servidora pública municipal de Caapiranga (TRANSF SALDO C/SAL R$ 1.351,02
mensal). Conta Bradesco Ag 3707/413004-9. RG = CIN moderna (mesmo nº do
CPF). Estado civil omitido.

Tese: 165 lançamentos APLIC.INVEST FACIL entre 24/02/2017 e 13/08/2024
(total tabela R$ 100.817,72). Audit: extrato detecta 165 APLIC vs 253
RESGATE, totais R$ 100.817,72 vs R$ 98.122,72. SALDO POSITIVO R$ 2.695
— diferente dos demais casos. Cliente teve perda LÍQUIDA pequena de
R$ 2.695. Não muda a estratégia (b) — ainda é "ciclo aplica-resgate"
e dano moral pela recorrência. Alerta no relatório.

Estratégia (b) PADRÃO. VC R$ 15.000.
"""
import io, sys, os, copy, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\EXEMPLA RAIMUNDA CRUZ - Ney Pedroza'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_RAIMUNDA_LIRA_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_RAIMUNDA_LIRA_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_raimunda_lira_lancs.json')) as f:
    LANCAMENTOS = sorted(json.load(f), key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'EXEMPLA RAIMUNDA CRUZ', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'servidora pública municipal',
    'cpf': '000.000.027-37', 'rg': '1000025-5',  # CIN moderno (CPF=RG)
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Dominguinhos', 'numero': '224',
    'bairro': 'Zona Rural', 'cidade': 'Caapiranga', 'cep': '69.425-000',
}
conta = {'agencia': '3707', 'numero': '413004-9'}
renda = {'valor_float': 1351.02}

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Caapiranga', uf='AM')
dados['valor_causa'] = fmt_moeda_rs(15000.00)
dados['valor_causa_extenso'] = extenso_moeda(15000.00)
dados['remuneração'] = 'salário do serviço público municipal'

print(f'=== EXEMPLA RAIMUNDA CRUZ — APLIC.INVEST (b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | VC: R$ 15.000,00'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
for p in list(doc.paragraphs):
    for m in MARKERS:
        if m in (p.text or ''): p._element.getparent().remove(p._element); break

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'reiterado durante todo o período em que a renda alimentar da parte autora ficou '
        'indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = ('No caso concreto, a tabela anexa registra 165 (cento e sessenta e cinco) ocorrências '
      'de aplicação automática entre 24/02/2017 e 13/08/2024 (mais de 7 anos). Em todos os '
      'meses, parcela substancial do salário da parte autora — recebida via transferência '
      'de saldo da conta-salário — foi automaticamente subtraída pelo banco réu sob a '
      'rubrica APLIC.INVEST FACIL, restando indisponível ao consumidor por período '
      'variável até o resgate manual. Embora a maior parte dos valores tenha sido '
      'restituída via RESGATE INVEST FACIL ao longo do período, o cerne do dano moral não '
      'reside na perda patrimonial — predominantemente compensada pelos resgates —, mas '
      'na privação reiterada da autodeterminação do consumidor sobre sua própria renda '
      'alimentar, mês após mês, durante 7 (sete) anos consecutivos. Cada retenção mensal '
      'configura, autonomamente, prática abusiva vedada pelo art. 39, inciso VI, do '
      'Código de Defesa do Consumidor, sendo a recorrência sistêmica o fato gerador do '
      'abalo extrapatrimonial.')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_RAIMUNDA_LIRA', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'APLIC.INVEST FACIL — estratégia (b)'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_AplicInvest_RAIMUNDA_LIRA_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CIN)', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '07/10/1970 — 55 anos (NÃO idosa)'),
    ('Estado civil', '(omitido — confirmar)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (TRANSF SALDO C/SAL — servidora municipal)'),
    ('Período tabela', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC', '165 (mais de 7 anos)'),
    ('Lançamentos RESGATE', '253'),
    ('Total bruto APLIC', dados['total_descontos']),
    ('Saldo líquido (APLIC - RESGATE)', '+R$ 2.695,00 (cliente teve PERDA líquida pequena)'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('AUDITORIA APLIC vs RESGATE — saldo POSITIVO R$ 2.695,00',
     'Diferentemente dos demais casos APLIC do batch (que tinham saldo zero), aqui o '
     'cliente teve perda líquida pequena de R$ 2.695,00 (165 APLICs R$ 100.817,72 vs '
     '253 RESGATEs R$ 98.122,72). Pelo teor do extrato, parece que algumas aplicações '
     'antigas não foram totalmente resgatadas. Mesmo assim, a estratégia (b) padrão é '
     'mais segura — o pleito de R$ 15.000 dano moral cobre folgadamente esse valor + '
     'fundamentação na recorrência mensal por 7 anos. Se o procurador quiser pleitear '
     'os R$ 2.695 separadamente em dobro (R$ 5.390), poderia adicionar como pedido '
     'subsidiário — mas aumenta complexidade.'),
    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Bloco "Repetição do indébito" REMOVIDO. Parágrafo doutrinário REESCRITO. '
     'Parágrafo do caso concreto INSERIDO (165 retenções, 7 anos). VC = R$ 15.000.'),
    ('PASTA KIT — outras 2 teses do mesmo cliente',
     'A planilha 7-TABELA.xlsx tem 3 abas: APLIC INVEST FACIL (R$ 100.817,72), '
     'TARIFA BANCARIA CESTA (a confirmar valor), ENCARGOS LIMITE DE CRED (a confirmar). '
     'Decidir em batch separado se vai gerar combinada (Caapiranga adota) ou separadas.'),
    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 ≈ 9,87 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (413004-9 / 3707).',
           'Conferir comarca: Caapiranga/AM.', 'Conferir VC = R$ 15.000,00.',
           'Considerar se vai pleitear R$ 2.695 perda líquida em pedido subsidiário.',
           'Decidir sobre as outras 2 teses do KIT (TARIFA + ENCARGOS).']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após pendências, '); r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
