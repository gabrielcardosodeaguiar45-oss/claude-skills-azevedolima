"""Inicial APLIC.INVEST FACIL da EXEMPLA MARIA CORREIA DA SILVA.

Comarca: Caapiranga/AM. NÃO idosa (40 anos, nascida 17/02/1986).
Servidora pública municipal de Caapiranga (TRANSF SALDO C/SAL P/CC
mensal R$ 1.416,02). Estado civil OMITIDO (RG origem cert. nascimento;
procuração não menciona — presumir solteira mas registrar).

Tese: 22 lançamentos APLIC.INVEST FACIL entre 31/01/2023 e 31/07/2024
(total tabela R$ 10.932,72). Audit do extrato detecta 80 APLICs
totalizando R$ 49.558,42 e 130 RESGATEs R$ 49.559,12 (saldo -R$ 0,70).
Diferença: tabela usa apenas alguns lançamentos selecionados,
extrato traz todos. CONFERIR antes do protocolo qual lista usar.

Estratégia (b) PADRÃO. VC R$ 15.000.

Pasta tem material para outras 4 teses (CARTAO ANUIDADE, ENCARGOS,
MORA, PACOTE SERVIÇOS, SERVIÇO CARTAO PROTEGIDO) no KIT — fora do
escopo do batch APLIC. Há subpastas NOTIFICAÇÃO ENCARGOS e
NOTIFICAÇÃO TARIFAS com material completo.
"""
import io, sys, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\EXEMPLA MARIA CORREIA DA SILVA - Ney Pedroza'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_MARIA_ADRIANA_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MARIA_ADRIANA_v1.docx')

autora = {
    'nome': 'EXEMPLA MARIA CORREIA DA SILVA',
    'nacionalidade': 'brasileira',
    'estado_civil': '',
    'profissao': 'servidora pública municipal',
    'cpf': '000.000.020-30',
    'rg': '1000018-8',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Novo Horizonte',
    'numero': 's/nº',
    'bairro': 'Novo Horizonte',
    'cidade': 'Caapiranga',
    'cep': '69.425-000',
}
conta = {'agencia': '3707', 'numero': '412844-3'}
renda = {'valor_float': 1416.02}

LANCAMENTOS = [
    ('31/01/2023', 410.68), ('10/02/2023', 987.45), ('27/02/2023', 207.47),
    ('31/03/2023', 1006.54), ('31/10/2023', 802.30), ('24/11/2023', 659.21),
    ('30/11/2023', 1089.84), ('29/01/2024', 141.81), ('30/01/2024', 693.22),
    ('16/02/2024', 76.80), ('21/02/2024', 99.30), ('04/03/2024', 164.92),
    ('28/03/2024', 773.84), ('02/05/2024', 98.50), ('03/05/2024', 1285.00),
    ('07/05/2024', 284.50), ('28/05/2024', 219.98), ('29/05/2024', 718.72),
    ('06/06/2024', 70.04), ('28/06/2024', 571.36), ('26/07/2024', 232.26),
    ('31/07/2024', 338.98),
]

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome': '', 'cnpj': '', 'logradouro': '', 'numero': '',
            'bairro': '', 'cidade': '', 'uf': '', 'cep': ''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Caapiranga', uf='AM')
dados['valor_causa'] = fmt_moeda_rs(15000.00)
dados['valor_causa_extenso'] = extenso_moeda(15000.00)
dados['remuneração'] = 'salário do serviço público municipal'

print(f'=== EXEMPLA MARIA — APLIC.INVEST FACIL (b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | VC: R$ 15.000,00'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
for p in list(doc.paragraphs):
    for m in MARKERS:
        if m in (p.text or ''): p._element.getparent().remove(p._element); break

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'reiterado durante todo o período em que a renda alimentar da parte autora ficou '
        'indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = (f'No caso concreto, a tabela anexa registra 22 (vinte e duas) ocorrências de aplicação '
      f'automática entre 31/01/2023 e 31/07/2024 (cerca de 18 meses), totalizando '
      f'R$ 10.932,72 brutos. Em cada um desses meses, parcela substancial do salário '
      f'da parte autora — recebida via transferência de saldo da conta-salário — foi '
      f'automaticamente subtraída pelo banco réu sob a rubrica APLIC.INVEST FACIL, '
      f'restando indisponível ao consumidor por período variável até o resgate manual. '
      f'Embora os valores tenham sido restituídos via RESGATE INVEST FACIL ao longo do '
      f'período, o cerne do dano moral não reside na perda patrimonial líquida — '
      f'inexistente —, mas na privação reiterada da autodeterminação do consumidor sobre '
      f'sua própria renda alimentar, mês após mês. Cada retenção mensal configura, '
      f'autonomamente, prática abusiva vedada pelo art. 39, inciso VI, do Código de '
      f'Defesa do Consumidor, sendo a recorrência sistêmica o fato gerador do abalo '
      f'extrapatrimonial.')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_MARIA_ADRIANA', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'APLIC.INVEST FACIL — estratégia (b) PADRÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_AplicInvest_MARIA_ADRIANA_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '17/02/1986 — 40 anos (NÃO idosa)'),
    ('Estado civil', '(omitido)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (TRANSF SALDO C/SAL — servidora pública municipal)'),
    ('Período tabela', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC (tabela)', '22'),
    ('Lançamentos APLIC (extrato)', '80 (extrato traz mais que a tabela)'),
    ('Lançamentos RESGATE (extrato)', '130'),
    ('Total bruto APLIC (NÃO pleiteado)', dados['total_descontos'] + ' (tabela)'),
    ('Saldo líquido (extrato)', '-R$ 0,70 (cliente recebeu de volta tudo)'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('AUDITORIA APLIC vs RESGATE — confirmada (estratégia b)',
     'Extrato confirma ciclo: 80 APLICs vs 130 RESGATEs totalizando R$ 49.558,42 vs '
     'R$ 49.559,12. Saldo NEGATIVO R$ 0,70.'),
    ('TABELA vs EXTRATO — diferença GRANDE (22 vs 80 lançamentos)',
     'A tabela 6 traz 22 lançamentos selecionados (R$ 10.932,72) entre 31/01/2023 e '
     '31/07/2024. O extrato detecta 80 APLICs totalizando R$ 49.558,42 — quase 5x mais. '
     'Possíveis hipóteses: (i) tabela usa apenas APLICs com algum critério de seleção '
     '(maiores valores, ou apenas alguns meses); (ii) extrato detecta APLICs do período '
     'inteiro 2020-2025. CONFERIR com escritório qual usar. A inicial gerada usa A '
     'TABELA (22 lançamentos, R$ 10.932,72). Se quiser maior cobertura temporal, '
     'parsear extrato completo e regenerar.'),
    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Template inicial-aplic-invest.docx pós-processado: 7 parágrafos do bloco '
     '"Repetição do indébito" REMOVIDOS, parágrafo doutrinário REESCRITO, parágrafo '
     'do caso concreto INSERIDO (22 retenções, 18 meses).'),
    ('TABELA "6 - TABELA.pdf" da pasta principal — ERRADA',
     'A tabela "6 - TABELA.pdf" da pasta PRINCIPAL é de ENCARGOS DE LIMITE DE CRÉDITO '
     '(não APLIC.INVEST). A tabela CORRETA de APLIC.INVEST está na subpasta '
     '"NOTIFICAÇÃO APLICAÇÃO INVEST FÁCIL/6 - TABELA.pdf". Esta inicial usa a '
     'tabela CORRETA. ATENÇÃO ao montar o KIT pré-protocolo.'),
    ('PASTA KIT — outras 4 teses do mesmo cliente',
     'KIT contém procurações + tabelas + notificações para: CARTAO CRED ANUIDADE, '
     'ENCARGOS LIMITE CRED, MORA CREDITO PESSOAL, PACOTE SERVIÇOS, SERVIÇO CARTAO '
     'PROTEGIDO. Subpastas NOTIFICAÇÃO ENCARGOS e NOTIFICAÇÃO TARIFAS já têm material '
     'completo. Decidir em batch separado se vai gerar combinada (Caapiranga adota) '
     'ou separadas.'),
    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 ≈ 9,87 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
    ('NOTIFICAÇÃO EXTRAJUDICIAL APLIC',
     'Pasta "NOTIFICAÇÃO APLICAÇÃO INVEST FÁCIL" tem o KIT completo (procuração, RG, '
     'extrato, OAB, tabela, xlsx) mas NÃO tem o documento de Notificação propriamente '
     'dita (.docx ou .pdf), tampouco comprovante de envio. CONFERIR e enviar antes do '
     'protocolo se for o caso.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (412844-3 / 3707).',
           'Conferir comarca: Caapiranga/AM.', 'Conferir VC = R$ 15.000,00.',
           'Conferir parágrafos doutrinários "Repetição" REMOVIDOS.',
           'Resolver TABELA vs EXTRATO (22 vs 80 lançamentos).',
           'Decidir sobre as outras 4 teses do KIT.',
           'Avaliar envio de notificação extrajudicial.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas — após pendências, '); r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
