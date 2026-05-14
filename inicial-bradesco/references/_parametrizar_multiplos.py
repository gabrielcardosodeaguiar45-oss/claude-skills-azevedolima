"""Parametriza inicial-jfba-multiplos-avn-inativo.docx (template para N contratos
da mesma rubrica AVERBAÇÃO NOVA INATIVO).

Estratégia:
- Cabeçalho/Qualificação/Polo passivo/Direito → mesmas substituições do base
- p15: frase introdutória plural (mantém placeholder `{{banco_reu_nome}}` e
  acrescenta `{{contratos_lista_breve}}` que a skill preenche com lista tipo
  "Nº X, Y e Z")
- p17: FIXO ("No que diz respeito aos referidos empréstimos, cumpre informar
  que:") — não recebe placeholder
- p18: BLOCO REPETÍVEL — parametriza com {{contrato_*}} e a skill DUPLICA o
  parágrafo N vezes antes de substituir cada cópia
- p19: DELETAR (é cópia do p18 do template original — não precisa)
- p193 + p233: dano moral mantém o texto "R$ 5.000,00 (cinco mil reais) para
  cada contrato" PARAMETRIZADO em {{dano_moral_unitario}} para flexibilidade
  futura (o cálculo final do total fica a cargo da skill)
- p231: BLOCO REPETÍVEL DECLARAR — parametriza e a skill duplica
- p232: DELETAR
"""
import io, sys, os, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from helpers_docx import substituir_in_run

def nfd(s):
    return unicodedata.normalize('NFD', s)

TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jfba-multiplos-avn-inativo.docx'

# Substituições (mesma estratégia do base + diferenças do MULTIPLOS)
SUBSTITUICOES = {
    0:  {'Salvador/BA': '{{cidade_subsecao}}/{{uf_subsecao}}'},
    7:  {
        'ROQUE CUSTODIO DA SILVA': '{{nome_autor}}',
        'brasileiro, solteiro, aposentado': '{{nacionalidade}}, {{estado_civil}}, {{profissao}}',
        '000.000.001-11': '{{cpf_autor}}',
        '01.095.230-67': '{{rg_autor}}',
        'Rua da Buganvilla, s/n, bairro Jardim Imbassai':
            '{{logradouro_autor}}, {{numero_autor}}, bairro {{bairro_autor}}',
        'Mata de São João/BA': '{{cidade_autor}}/{{uf_autor}}',
        'CEP 48289-000': 'CEP {{cep_autor}}',
    },
    11: {
        'BANCO DAYCOVAL S/A, inscrito no CNPJ/MF sob o nº 62.232.889/0001 90, na Avenida Paulista 1793, Bela Vista, São Paulo /SC, Brasil, CEP: 01311-200':
            '{{banco_reu_nome}}, {{banco_reu_descricao_pj}}, inscrito no CNPJ/MF sob o nº {{banco_reu_cnpj}}, com sede na {{banco_reu_endereco}}',
        'Av. Sete de Setembro, 1078 - Mercês, Salvador/BA':
            '{{inss_endereco_subsecao}}',
    },
    13: {
        'aposentadoria por invalidez previdenciária': '{{tipo_beneficio}}',
        '024.044.285-7': '{{nb_beneficio}}',
        'agência 3706, conta corrente nº 0000211974':
            'agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}',
        'Banco Cooperativo Sicoob SA': '{{banco_pagador}}',
    },
    # p15 — frase intro plural (mantém banco_reu_nome + acrescenta lista de contratos)
    15: {
        'BANCOS DAYCOVAL S/A, CONTRATO Nº 237344 469':
            '{{banco_reu_nome}}, {{contratos_lista_breve}}',
    },
    # p18 — BLOCO REPETÍVEL contrato (parametriza UM, skill duplica N)
    18: {
        'Do contrato nº 000000': 'Do contrato nº {{contrato_numero}}',
        'na competência 05/2017': 'na competência {{contrato_competencia_inicio}}',
        'de um total de 72 parcelas': 'de um total de {{contrato_qtd_parcelas}} parcelas',
        'no valor de R$ 58,98': 'no valor de R$ {{contrato_valor_parcela}}',
        'no valor de R$ 4.246,56': 'no valor de R$ {{contrato_valor_emprestado}}',
        'tendo as parcelas sido encerradas em 02/2023': 'tendo as parcelas sido encerradas em {{contrato_competencia_fim}}',
        'contrato n° 237344 469': 'contrato n° {{contrato_numero}}',
        'banco Daycoval S.A': 'banco {{contrato_banco}}',
    },
    46: {'R$ 974,69': 'R$ {{valor_renda_liquida}}'},
    # p193 — quantum (R$ 5.000 por contrato — placeholder permite variar)
    193: {
        'R$ 5.000,00 (cinco mil reais) para cada contrato':
            'R$ {{dano_moral_unitario}} ({{dano_moral_unitario_extenso}}) para cada contrato, totalizando R$ {{dano_moral_total}} ({{dano_moral_total_extenso}})',
    },
    # p231 — BLOCO REPETÍVEL DECLARAR (parametriza UM, skill duplica N)
    231: {
        'no valor de R$ 4.246,56': 'no valor de R$ {{contrato_valor_emprestado}}',
        'contrato nº- 237344 469': 'contrato nº {{contrato_numero}}',
        'descontos de R$ 58,98 mensais': 'descontos de R$ {{contrato_valor_parcela}} mensais',
        'com inclusão em 03/05/2017': 'com inclusão em {{contrato_data_inclusao}}',
        'início de desconto em 05/2017 a 02/2023': 'início de desconto em {{contrato_competencia_inicio}} a {{contrato_competencia_fim}}',
        'benefício previdenciário 024.044.285-7': 'benefício previdenciário {{nb_beneficio}}',
    },
    # p233 — pedido CONDENAR
    233: {
        'R$ 5.000,00 (cinco mil reais) para cada contrato':
            'R$ {{dano_moral_unitario}} ({{dano_moral_unitario_extenso}}) para cada contrato, totalizando R$ {{dano_moral_total}} ({{dano_moral_total_extenso}})',
    },
    249: {
        'R$ 40.212,65 (quarenta mil, duzentos e doze reais e sessenta e cinco centavos)':
            'R$ {{valor_causa}} ({{valor_causa_extenso}})',
    },
    251: {'Salvador/BA, 7 de maio de 2026': '{{cidade_protocolo}}/{{uf_protocolo}}, {{data_protocolo}}'},
}

# Parágrafos a REMOVER (cópias dos blocos repetíveis no template original)
REMOVER = [19, 232]  # p19 = cópia de p18; p232 = cópia de p231

print(f'Carregando: {TEMPLATE}')
doc = Document(TEMPLATE)
print(f'Total parágrafos: {len(doc.paragraphs)}')

# Substituições
falhas = []
sucessos = 0
for idx, mapa in SUBSTITUICOES.items():
    if idx >= len(doc.paragraphs):
        falhas.append(f'p{idx} fora do range')
        continue
    p = doc.paragraphs[idx]
    texto_antes = p.text
    for antigo, novo in mapa.items():
        chave = antigo if antigo in texto_antes else nfd(antigo)
        if chave not in texto_antes:
            falhas.append(f'p{idx}: NÃO ENCONTRADO "{antigo[:60]}..."')
            continue
        ok = substituir_in_run(p._element, {chave: novo})
        if ok:
            sucessos += 1
        else:
            falhas.append(f'p{idx}: substituir_in_run falhou para "{antigo[:60]}"')

# Remover parágrafos duplicados (cópias do bloco repetível)
# IMPORTANTE: remover de trás para frente para não bagunçar índices
for idx in sorted(REMOVER, reverse=True):
    if idx >= len(doc.paragraphs):
        falhas.append(f'p{idx} (REMOVER) fora do range')
        continue
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)
    print(f'Removido: p{idx}')

print(f'\n--- RESUMO ---')
print(f'Substituições bem-sucedidas: {sucessos}')
print(f'Parágrafos removidos: {len(REMOVER)}')
print(f'Falhas: {len(falhas)}')
for f in falhas:
    print(f'  ✗ {f}')

doc.save(TEMPLATE)
print(f'\nSalvo em: {TEMPLATE}')

# Verificar
import re
doc2 = Document(TEMPLATE)
texto_full = '\n'.join(p.text for p in doc2.paragraphs)
placeholders = sorted(set(re.findall(r'\{\{[^}]+\}\}', texto_full)))
print(f'\n--- PLACEHOLDERS DETECTADOS ({len(placeholders)} únicos) ---')
for ph in placeholders:
    cnt = texto_full.count(ph)
    print(f'  {ph} (×{cnt})')
