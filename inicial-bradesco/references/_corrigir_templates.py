"""Corrige 5 problemas identificados pelo usuário nos 3 templates da pasta
IniciaisNaoContratado/_templates/.

1. p0 (juízo) tem left_indent ~5cm — REMOVER
2. p11 (polo passivo): run do banco está com TUDO em Segoe UI (nome + CNPJ
   + endereço). Quebrar em: nome em Segoe UI, descrição PJ + CNPJ + endereço
   em Cambria.
3. BASE p15: substituir "BANCO DAYCOVAL S/A, CONTRATO Nº 237344 469" por
   "{{banco_reu_nome}}, CONTRATO Nº {{contrato_numero}}".
4. Prioridade idoso (p2 cabeçalho + pedido na lista): será tratada na lógica
   da skill (remoção condicional). Não exige edição no template.
5. MULT — adicionar marcador visual antes dos blocos repetíveis para que a
   revisão visual do template comunique o caráter "repetível".
"""
import io, sys, os, copy
from copy import deepcopy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt
from lxml import etree

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

BASE_DIR = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates'
T = {
    'base':  os.path.join(BASE_DIR, 'inicial-jfba-base.docx'),
    'mult':  os.path.join(BASE_DIR, 'inicial-jfba-multiplos-avn-inativo.docx'),
    'refin': os.path.join(BASE_DIR, 'inicial-jfba-refin-ativo.docx'),
}


# === FIX 1: remover left_indent do p0 ===
def fix1_remover_indent_p0(doc):
    p0 = doc.paragraphs[0]
    # Remover indentação esquerda do XML
    pPr = p0._element.find(W + 'pPr')
    if pPr is not None:
        ind = pPr.find(W + 'ind')
        if ind is not None:
            pPr.remove(ind)
    # Garantir alinhamento centralizado (já estava, mas garantir)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


# === FIX 2: separar runs do polo passivo (p11) ===
def fix2_runs_polo_passivo(doc):
    """Reconstrói os 5 runs do p11:
    - r0 (Cambria): "em face de "
    - r1 (Segoe UI Bold): "{{banco_reu_nome}}"
    - r2 (Cambria): ", {{banco_reu_descricao_pj}}, inscrito no CNPJ/MF sob o nº {{banco_reu_cnpj}}, com sede na {{banco_reu_endereco}}, e "
    - r3 (Segoe UI Bold): "INSTITUTO NACIONAL DO SEGURO SOCIAL - INSS"
    - r4 (Cambria): ", Autarquia Federal, com sede à {{inss_endereco_subsecao}} pelos motivos de fato e de direito a seguir expostos:"

    O texto exato pode variar por template (REFIN tem "inscrita" e "com endereço na"
    em vez de "inscrito" e "na" do BASE/MULT). Vou pegar o texto atual e dividir.
    """
    p = doc.paragraphs[11]
    pe = p._element

    # Texto consolidado do parágrafo
    texto_full = p.text

    # Identificar onde os 2 nomes começam/terminam para fazer o split
    pos_banco_ini = texto_full.find('{{banco_reu_nome}}')
    pos_banco_fim = pos_banco_ini + len('{{banco_reu_nome}}')
    pos_inss_ini = texto_full.find('INSTITUTO NACIONAL DO SEGURO SOCIAL')
    pos_inss_fim = pos_inss_ini + len('INSTITUTO NACIONAL DO SEGURO SOCIAL - INSS')

    if pos_banco_ini < 0 or pos_inss_ini < 0:
        print('  ⚠ FIX2: padrões não encontrados no p11')
        return False

    # 5 segmentos
    seg1 = texto_full[:pos_banco_ini]               # "em face de "
    seg2 = texto_full[pos_banco_ini:pos_banco_fim]  # "{{banco_reu_nome}}"
    seg3 = texto_full[pos_banco_fim:pos_inss_ini]   # ", desc, CNPJ, endereço, e "
    seg4 = texto_full[pos_inss_ini:pos_inss_fim]    # "INSTITUTO NACIONAL ..."
    seg5 = texto_full[pos_inss_fim:]                # ", Autarquia Federal, ..."

    # Capturar pPr (preservar formatação do parágrafo)
    pPr = pe.find(W + 'pPr')

    # Remover TODOS os runs existentes (mantém pPr)
    for child in list(pe):
        if child.tag != W + 'pPr':
            pe.remove(child)

    # Adicionar 5 runs novos
    def _add_run(texto, segoe_bold=False):
        r = etree.SubElement(pe, W + 'r')
        rpr = etree.SubElement(r, W + 'rPr')
        if segoe_bold:
            # Fonte Segoe UI + bold
            rfonts = etree.SubElement(rpr, W + 'rFonts')
            rfonts.set(W + 'ascii', 'Segoe UI')
            rfonts.set(W + 'hAnsi', 'Segoe UI')
            etree.SubElement(rpr, W + 'b')
            etree.SubElement(rpr, W + 'bCs')
        else:
            # Cambria (fonte default do corpo)
            rfonts = etree.SubElement(rpr, W + 'rFonts')
            rfonts.set(W + 'ascii', 'Cambria')
            rfonts.set(W + 'hAnsi', 'Cambria')
            rfonts.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't')
        t.text = texto
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    _add_run(seg1, segoe_bold=False)        # "em face de "
    _add_run(seg2, segoe_bold=True)         # banco
    _add_run(seg3, segoe_bold=False)        # ", descrição, CNPJ..."
    _add_run(seg4, segoe_bold=True)         # INSS
    _add_run(seg5, segoe_bold=False)        # ", Autarquia Federal..."

    return True


# === FIX 3: parametrizar p15 do BASE ===
def fix3_p15_base(doc):
    """No template BASE, o p15 ainda tem 'BANCO DAYCOVAL S/A, CONTRATO Nº 237344 469'
    sem parametrizar. Substituir pelos placeholders."""
    from helpers_docx import substituir_in_run
    p = doc.paragraphs[15]
    if 'BANCO DAYCOVAL S/A' in p.text:
        ok = substituir_in_run(p._element, {
            'BANCO DAYCOVAL S/A, CONTRATO Nº 237344 469':
                '{{banco_reu_nome}}, CONTRATO Nº {{contrato_numero}}'
        })
        return ok
    return False


# === FIX 5: adicionar marcador visual de bloco repetível no MULT ===
def fix5_marcador_bloco_repetivel_mult(doc):
    """Adiciona um parágrafo NOTA visual antes dos blocos repetíveis do MULT.
    A skill detecta esse marcador e o REMOVE antes de duplicar o bloco.
    """
    # Localizar p18 (Do contrato nº...) e adicionar nota antes
    # Localizar pedido DECLARAR e adicionar nota antes

    # Achar p17 (introdutor) e logo após inserir o marcador antes do p18
    indices_inserir = []
    for i, par in enumerate(doc.paragraphs):
        if 'Do contrato nº {{contrato_numero}}' in par.text:
            indices_inserir.append(('SINTESE_FATICA', i))
        elif 'Declarar a inexistência' in par.text and '{{contrato_numero}}' in par.text:
            indices_inserir.append(('DECLARAR', i))

    # Inserir nota ANTES de cada bloco repetível
    # Inserir de trás para frente para não bagunçar índices
    NOTA_TXT = '⤵ BLOCO REPETÍVEL — a skill duplica este parágrafo para cada contrato. Para preview manual, basta editar este parágrafo. ⤵'
    for label, idx in sorted(indices_inserir, key=lambda x: -x[1]):
        par_alvo = doc.paragraphs[idx]
        # Criar parágrafo de nota
        nota = par_alvo._element.makeelement(W + 'p', {})
        # Estilo: usar mesma indentação do par_alvo
        pPr_alvo = par_alvo._element.find(W + 'pPr')
        if pPr_alvo is not None:
            nota.append(deepcopy(pPr_alvo))
        # Run com texto e formatação destacada (itálico cinza pequeno)
        r = etree.SubElement(nota, W + 'r')
        rpr = etree.SubElement(r, W + 'rPr')
        rfonts = etree.SubElement(rpr, W + 'rFonts')
        rfonts.set(W + 'ascii', 'Cambria')
        rfonts.set(W + 'hAnsi', 'Cambria')
        etree.SubElement(rpr, W + 'i')
        sz = etree.SubElement(rpr, W + 'sz')
        sz.set(W + 'val', '20')  # 10pt (val=20 = half-points)
        color = etree.SubElement(rpr, W + 'color')
        color.set(W + 'val', '7F7F7F')
        t = etree.SubElement(r, W + 't')
        t.text = NOTA_TXT
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        # Inserir ANTES do par_alvo
        par_alvo._element.addprevious(nota)

    return len(indices_inserir)


# === EXECUÇÃO ===
def processar(label, path, fixes_aplicar):
    print(f'\n████ {label} — {os.path.basename(path)} ████')
    doc = Document(path)
    resultados = {}
    for nome_fix, fn in fixes_aplicar.items():
        ok = fn(doc)
        resultados[nome_fix] = ok
        print(f'  {nome_fix}: {ok}')
    doc.save(path)
    print(f'  SALVO em {path}')
    return resultados


if __name__ == '__main__':
    # Para os 3 templates: fix 1 (p0 centralizar) e fix 2 (p11 runs)
    fixes_3templates = {
        'fix1_remover_indent_p0': fix1_remover_indent_p0,
        'fix2_runs_polo_passivo': fix2_runs_polo_passivo,
    }

    # Apenas para BASE: fix 3 (p15)
    fixes_base = dict(fixes_3templates)
    fixes_base['fix3_p15_base'] = fix3_p15_base

    # Apenas para MULT: fix 5 (marcador visual)
    fixes_mult = dict(fixes_3templates)
    fixes_mult['fix5_marcador_bloco_repetivel_mult'] = fix5_marcador_bloco_repetivel_mult

    processar('BASE',  T['base'],  fixes_base)
    processar('MULT',  T['mult'],  fixes_mult)
    processar('REFIN', T['refin'], fixes_3templates)

    # Verificar
    print('\n=== VERIFICAÇÃO PÓS-CORREÇÃO ===')
    for label, path in T.items():
        doc = Document(path)
        p0 = doc.paragraphs[0]
        print(f'\n{label}:')
        print(f'  p0 left_indent: {p0.paragraph_format.left_indent}')
        print(f'  p0 alignment: {p0.alignment}')
        # contar runs em p11
        print(f'  p11 runs: {len(doc.paragraphs[11].runs)}')
        # Mostrar texto compactado de cada run de p11
        for i, r in enumerate(doc.paragraphs[11].runs):
            rpr = r._element.find(W + 'rPr')
            font = '?'
            bold = ''
            if rpr is not None:
                rfonts = rpr.find(W + 'rFonts')
                if rfonts is not None:
                    font = rfonts.attrib.get(W + 'ascii', '?')
                if rpr.find(W + 'b') is not None:
                    bold = ' BOLD'
            print(f'    r{i} [{font}{bold}]: {r.text[:80]!r}')
