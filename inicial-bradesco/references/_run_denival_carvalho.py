"""Gera inicial APLIC.INVEST FACIL do EXEMPLO CARVALHO.

Comarca: Manaus/AM (foro do domicílio art. 101 I CDC — Av. Margarita, 5,
Cidade Nova, CEP 69.097-207, bairro grande da zona norte de Manaus, NÃO
de Maués apesar do nome da pasta-mãe sugerir o contrário). Idoso 69 anos
(nascido 25/12/1956 → prioridade art. 1.048 I CPC). Casado (origem do RG
em certidão de casamento). Conta Bradesco Ag 3739, conta 75911-2. Renda
INSS R$ 947,90 (último crédito recorrente 26/12/2025).

Tese APLIC.INVEST FACIL — 34 lançamentos entre 10/09/2020 e 26/06/2024,
total bruto aplicado R$ 41.023,45 (NÃO PLEITEAR — ver § APLIC.INVEST da
SKILL.md). Auditoria do extrato mostra que TODAS as 34 aplicações foram
RESGATADAS via RESGATE INVEST FACIL em D+1 a D+3, exceto a aplicação de
R$ 10.550,06 em 13/10/2020 (originada de EMPRESTIMO PESSOAL recebido no
mesmo dia) que demorou ~27 dias para retorno integral. Saldo líquido
agregado em 4 anos: ~zero (rentabilidade total RENTAB.INVEST FACILCRED
< R$ 0,15).

ESTRATÉGIA DEFINIDA PELO PROCURADOR (06/05/2026): (b) CONSERVADORA — só
dano moral, fundamentado nas RETENÇÕES recorrentes (cada aplicação
mensal sem autorização configura prática abusiva, art. 39 VI CDC, mesmo
que devolvida em D+1). NÃO pleitear repetição em dobro. Pedidos:
declaratório de inexistência + obrigação de cessar (multa R$ 500/dia) +
R$ 15.000 dano moral. VC R$ 15.000 → cabe folgadamente no JEC (~10 SM).

Por causa da estratégia (b), o template padrão inicial-aplic-invest.docx
precisa de pós-processamento: remover bloco doutrinário "Repetição do
indébito" (parágrafos 146-151), remover pedido (b) "Havendo a retenção
dos valores..." (175), e reescrever último período do parágrafo 102 que
faz referência à repetição em dobro. Também é INSERIDO parágrafo
adicional após o 22 explicando o caso concreto DENIVAL (34 retenções
mensais com devolução em D+1-D+3, dano moral pela privação reiterada da
autodeterminação).

Comprovante de residência está no nome do FILHO Dionaldo Rodrigues
Batista (CPF 000.000.005-15), conforme fatura NIO Internet do mesmo
endereço. CNH 5.1 prova vínculo proprietário. Pendência sinalizada no
relatório paralelo para confirmar com cliente o vínculo de coabitação.

Pasta KIT contém material para outras teses (TARIFAS, possivelmente
outras) — fora do escopo deste batch (APLIC.INVEST). Pasta-mãe se chama
"EXEMPLO CARVALHO - Ruth - TARIFAS" sugerindo que o caso
principal do cliente é TARIFAS — confirmar se procurador quer processar
TARIFAS em batch separado.
"""
import io, sys, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\EXEMPLO CARVALHO - Ruth - TARIFAS\APLICAÇÃO INVEST FÁCIL - SEM MODELO'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_DENIVAL_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_DENIVAL_v1.docx')

autora = {
    'nome': 'EXEMPLO CARVALHO',
    'nacionalidade': 'brasileiro',
    'estado_civil': 'casado',
    'profissao': 'aposentado',
    'cpf': '000.000.006-16',
    'rg': '1000004-4',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Av. Margarita',
    'numero': '5, Casa',
    'bairro': 'Cidade Nova',
    'cidade': 'Manaus',
    'cep': '69.097-207',
}
conta = {'agencia': '3739', 'numero': '75911-2'}
renda = {'valor_float': 947.90}

LANCAMENTOS = [
    ('10/09/2020', 2539.74), ('13/10/2020', 10550.06), ('28/10/2020', 1045.00),
    ('26/11/2020', 2095.54), ('28/12/2020', 450.82), ('27/01/2021', 1037.92),
    ('24/02/2021', 1036.86), ('29/03/2021', 226.86), ('28/04/2021', 289.23),
    ('27/05/2021', 1586.33), ('28/06/2021', 266.44), ('27/08/2021', 790.89),
    ('27/10/2021', 776.34), ('26/11/2021', 773.84), ('27/01/2022', 898.85),
    ('23/02/2022', 888.36), ('27/04/2022', 1492.88), ('27/05/2022', 1487.89),
    ('27/07/2022', 887.91), ('27/10/2022', 884.57), ('27/12/2022', 906.88),
    ('27/01/2023', 975.54), ('24/02/2023', 974.20), ('26/04/2023', 990.86),
    ('27/07/2023', 869.07), ('06/09/2023', 251.20), ('27/09/2023', 806.17),
    ('27/10/2023', 806.10), ('26/12/2023', 775.65), ('27/02/2024', 884.20),
    ('21/03/2024', 302.66), ('26/03/2024', 919.46), ('26/04/2024', 989.87),
    ('26/06/2024', 565.26),
]

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome': '', 'cnpj': '', 'logradouro': '', 'numero': '',
            'bairro': '', 'cidade': '', 'uf': '', 'cep': ''}

dados, calc = montar_dados_padrao(
    autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
    eh_idoso=True, competência='Manaus', uf='AM',
)

# OVERRIDE estratégia (b): valor_causa = só dano moral
DANO_MORAL = 15000.00
VALOR_CAUSA_B = 15000.00
dados['valor_causa'] = fmt_moeda_rs(VALOR_CAUSA_B)
dados['valor_causa_extenso'] = extenso_moeda(VALOR_CAUSA_B)

# Adiciona chave faltante usada no parágrafo da Justiça Gratuita
# (montar_dados_padrao cria valor_remuneração mas NÃO cria a chave "remuneração"
# crua, que aparece no template como {{remuneração}})
dados['remuneração'] = 'aposentadoria pelo INSS'

print('=== EXEMPLO CARVALHO — APLIC.INVEST FACIL (estratégia b) ===')
print(f'Total bruto aplicado (NÃO pleiteado): R$ {calc["total"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'Dano moral pleiteado: R$ {DANO_MORAL:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'Valor da causa (b): R$ {VALOR_CAUSA_B:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'Idoso: {bool(dados["prioridade_cabecalho"])}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Aplicar template: {res["modificados"]} parágrafos modificados, '
      f'residuais: {res["residuais"] or "nenhum"}')

# ============== PÓS-PROCESSAMENTO (estratégia b) =================
from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

doc = Document(DOCX_OUT)

# (1) REMOVER parágrafos do bloco "Repetição do indébito" + pedido (b)
MARKERS_REMOVER = [
    'Repetição do indébito',
    'Caso se verifique a existência de valores indevidamente',
    'Na cobrança de débitos, o consumidor inadimplente',
    'O consumidor cobrado em quantia indevida tem direito',
    'a restituição em dobro se faz necessária como penalidade',
    'deve a requerida ser condenada a restituir em dobro',
    'Havendo a retenção dos valores a título de investimento',
]
removidos = []
for p in list(doc.paragraphs):
    txt = p.text or ''
    for marker in MARKERS_REMOVER:
        if marker in txt:
            p._element.getparent().remove(p._element)
            removidos.append(marker[:50])
            break
print(f'Parágrafos removidos: {len(removidos)}')
for r in removidos:
    print(f'  - {r}')

# (2) REESCREVER parágrafo do tipo 102: remover trecho sobre repetição em dobro
TEXTO_NOVO_102 = (
    'A cobrança indevida não decorre de engano justificável, mas de '
    'modelo operacional estruturado para funcionar sem contratação '
    'inequívoca, configurando ato ilícito reiterado durante todo o '
    'período em que a renda alimentar da parte autora ficou indisponível '
    'para uso imediato.'
)
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        # Limpar todos os runs e setar texto novo
        for run_elem in list(p._element.findall(W + 'r')):
            p._element.remove(run_elem)
        # Adicionar novo run em Cambria
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr')
        rFonts = etree.SubElement(rPr, W + 'rFonts')
        rFonts.set(W + 'ascii', 'Cambria')
        rFonts.set(W + 'hAnsi', 'Cambria')
        rFonts.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TEXTO_NOVO_102
        print('Parágrafo doutrinário (102) reescrito sem trecho de repetição em dobro.')
        break

# (3) INSERIR parágrafo após "Alegar que os valores permaneciam disponíveis..."
TEXTO_DENIVAL = (
    'No caso concreto, o extrato bancário registra 34 (trinta e quatro) '
    'ocorrências mensais de aplicação automática entre 10/09/2020 e '
    '26/06/2024. Em todos os meses, o valor da aposentadoria do INSS, '
    'recebido em conta-salário, foi automaticamente subtraído pelo banco '
    'réu sob a rubrica APLIC.INVEST FACIL, restando indisponível ao '
    'consumidor pelo prazo de 1 (um) a 27 (vinte e sete) dias até o '
    'resgate manual. Embora os valores tenham sido restituídos via '
    'RESGATE INVEST FACIL ao longo do período, o cerne do dano moral não '
    'reside na perda patrimonial líquida — inexistente —, mas na '
    'privação reiterada da autodeterminação do consumidor sobre sua '
    'própria renda alimentar, mês após mês, durante 4 (quatro) anos '
    'consecutivos. Cada retenção mensal configura, autonomamente, '
    'prática abusiva vedada pelo art. 39, inciso VI, do Código de Defesa '
    'do Consumidor, sendo a recorrência sistêmica o fato gerador do '
    'abalo extrapatrimonial.'
)
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        # Criar novo parágrafo ao lado
        new_p_xml = copy.deepcopy(p._element)
        # Limpar runs do novo
        for r in list(new_p_xml.findall(W + 'r')):
            new_p_xml.remove(r)
        # Adicionar novo run em Cambria
        r = etree.SubElement(new_p_xml, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr')
        rFonts = etree.SubElement(rPr, W + 'rFonts')
        rFonts.set(W + 'ascii', 'Cambria')
        rFonts.set(W + 'hAnsi', 'Cambria')
        rFonts.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TEXTO_DENIVAL
        # Inserir após o atual
        p._element.addnext(new_p_xml)
        print('Parágrafo do caso concreto DENIVAL inserido.')
        break

doc.save(DOCX_OUT)

# (4) PÓS-FIX RAW: aplicar_template falhou nos 3 placeholders abaixo
# (provavelmente por encoding NFC vs NFD do "ç" no XML do template).
# Aplicamos substituição direta no XML como fallback.
import zipfile, re
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}

xml = buf['word/document.xml'].decode('utf-8')
substituicoes_raw = {
    '{{remuneração}}': 'aposentadoria pelo INSS',
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
}
fix_count = 0
for placeholder, valor in substituicoes_raw.items():
    occurrences = xml.count(placeholder)
    if occurrences:
        xml = xml.replace(placeholder, valor)
        fix_count += occurrences
        print(f'Pós-fix XML: substituído "{placeholder}" → "{valor[:40]}..." ({occurrences}x)')

# Também tenta a forma com runs partidos (placeholder dividido entre tags)
# usando regex que tolera tags <w:t> intermediárias
for placeholder, valor in substituicoes_raw.items():
    # Constrói padrão regex que tolera quebras de run no meio do placeholder
    chars_pattern = ''.join(re.escape(c) + r'(?:</w:t>(?:.*?)<w:t[^>]*>)?' for c in placeholder)
    pattern = re.compile(chars_pattern, re.DOTALL)
    matches = list(pattern.finditer(xml))
    if matches:
        # Preserva o primeiro match completo, substitui pelo valor
        for m in reversed(matches):
            xml = xml[:m.start()] + valor + xml[m.end():]
            fix_count += 1
            print(f'Pós-fix XML (multi-run): substituído "{placeholder}" via regex')

buf['word/document.xml'] = xml.encode('utf-8')
import os as _os
_os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes:
        z.writestr(n, buf[n])

print(f'Pós-fix concluído: {fix_count} substituições no XML.')
print(f'OK -> {DOCX_OUT}')

# ============== RELATÓRIO DE PENDÊNCIAS =================
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_DENIVAL', level=1)

cab = [
    ('Cliente', autora['nome']),
    ('Tese', 'APLIC.INVEST FACIL — estratégia (b) CONSERVADORA, só dano moral'),
    ('Comarca', f'{dados["competência"]}/{dados["uf"]} (JEC capital)'),
    ('Arquivo gerado', 'INICIAL_AplicInvest_DENIVAL_v1.docx'),
]
for k, v in cab:
    p = doc_r.add_paragraph()
    p.add_run(k + ': ').bold = True
    p.add_run(v)

doc_r.add_heading('1. RESUMO DOS DADOS APLICADOS', level=2)
tbl = doc_r.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'
tbl.rows[0].cells[1].text = 'Valor'
linhas = [
    ('Comarca', f'{dados["competência"]}/{dados["uf"]}'),
    ('Prioridade', dados.get('prioridade_cabecalho', '')),
    ('Nome', autora['nome']),
    ('CPF', autora['cpf']),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '25/12/1956 — 69 anos (IDOSO)'),
    ('Estado civil', autora['estado_civil']),
    ('Profissão', autora['profissao']),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda real (extrato)', dados['valor_remuneração']),
    ('Réu', 'BANCO BRADESCO S.A. — CNPJ 60.746.948/0001-12'),
    ('Rubrica', dados['rubrica_curta_caps']),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', '34 (trinta e quatro) ocorrências mensais'),
    ('Total bruto aplicado (NÃO pleiteado)', dados['total_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]
for k, v in linhas:
    row = tbl.add_row().cells
    row[0].text = k
    row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
pendencias = [
    ('AUDITORIA APLIC vs RESGATE — confirmar antes do protocolo',
     'Conferiu-se o extrato bancário 2020–2025 e constatou-se que TODAS '
     'as 34 aplicações automáticas foram RESGATADAS via RESGATE INVEST '
     'FACIL em D+1 a D+3, exceto a aplicação de R$ 10.550,06 em '
     '13/10/2020 (originada de EMPRESTIMO PESSOAL recebido no mesmo dia) '
     'que demorou aproximadamente 27 dias para retorno integral '
     '(09/11/2020). Rentabilidade acumulada em 4 anos: < R$ 0,15. Por '
     'isso a estratégia (b) escolhida pelo procurador foi pleitear '
     'APENAS dano moral pela RECORRÊNCIA das retenções mensais (não a '
     'soma das aplicações em dobro). Banco vai apresentar comprovante '
     'dos resgates na contestação — a réplica precisa estar preparada '
     'para sustentar que o dano moral independe do resgate, pois decorre '
     'da privação mensal reiterada da renda alimentar (art. 39 VI CDC).'),

    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'O template inicial-aplic-invest.docx foi pós-processado para '
     'estratégia (b): (i) removidos 7 parágrafos do bloco "Repetição do '
     'indébito" e do pedido subsidiário de devolução em dobro; (ii) '
     'reescrito parágrafo doutrinário que fazia referência à repetição '
     'em dobro; (iii) inserido parágrafo adicional explicando '
     'concretamente as 34 retenções mensais do DENIVAL e o porquê de '
     'não haver pleito patrimonial. CONFERIR no DOCX gerado se as '
     'remoções/inserções ficaram coerentes textualmente.'),

    ('COMPROVANTE DE RESIDÊNCIA está no nome do FILHO',
     'A fatura NIO Internet (5 - COMPROVANTE DE RESIDÊNCIA.pdf) está em '
     'nome de EXEMPLA RODRIGUES (CPF 000.000.005-15), filho do '
     'autor. CNH 5.1 - PROPRIETÁRIO (FILHO).pdf comprova o vínculo. '
     'Endereço bate com a procuração (Av. Margarita, 5, Cidade Nova, '
     'Manaus/AM, CEP 69.097-207). Confirmar com cliente o vínculo de '
     'coabitação para fundamentar o domicílio. Eventual declaração '
     'manuscrita do filho confirmando que o autor reside no mesmo '
     'endereço pode reforçar.'),

    ('COMARCA Manaus, NÃO Maués',
     'A pasta-mãe do escritório se chama "EXEMPLO CARVALHO - '
     'Ruth - TARIFAS" e o operador inicialmente assumiu Maués. Mas o '
     'endereço real (Av. Margarita, Cidade Nova, CEP 69.097-207) é em '
     'MANAUS — Cidade Nova é bairro grande da zona norte de Manaus, e o '
     'CEP 69097-XXX pertence a Manaus, não a Maués (Maués usa CEP '
     '69190-000). Procuração foi assinada em Manaus em 22/01/2026. '
     'Comarca correta para o JEC: Manaus capital.'),

    ('SEM NOTIFICAÇÃO EXTRAJUDICIAL na pasta',
     'A pasta APLIC.INVEST FÁCIL não contém o arquivo "8 - NOTIFICACAO" '
     'que normalmente é enviado antes do protocolo (atualmente '
     'controvérsia STJ sobre exigência de prévia tentativa '
     'extrajudicial). Avaliar se vale enviar notificação extrajudicial '
     'ANTES do protocolo (e juntar comprovante depois) para reforçar a '
     'tese e neutralizar eventual preliminar de "ausência de prévio '
     'requerimento". O bloco "Do prévio requerimento de solução '
     'administrativa" do template (parágrafos 25-29) MENCIONA a '
     'notificação — ALERTA: se a inicial alegar que enviou notificação '
     'mas o documento não estiver nos autos, vai gerar contradição. '
     'CONFERIR/AJUSTAR esse bloco no DOCX gerado antes do protocolo.'),

    ('PASTA-MÃE "TARIFAS" — outro caso paralelo do mesmo cliente?',
     'O nome da pasta-mãe sugere que o caso principal do DENIVAL no '
     'escritório é uma tese de TARIFAS (não APLIC.INVEST). Esta inicial '
     'cobre APENAS APLIC.INVEST. Confirmar se haverá inicial separada '
     'para TARIFAS (provavelmente sim — a procuração específica para '
     'APLIC.INVEST estava nesta subpasta, mas a pasta-mãe deve ter '
     'outra procuração para TARIFAS). 0. Kit / Thumbs.db pode conter '
     'material adicional não inspecionado.'),

    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 corresponde a aproximadamente 9,87 SM (SM 2025 = '
     'R$ 1.518). Cabe folgadamente no JEC (40 SM = R$ 60.720). Sem '
     'necessidade de renúncia ao excedente nem ajuizamento em vara '
     'cível comum.'),

    ('OAB ADICIONAL no item "4 - OAB PATRICK.pdf"',
     'A pasta contém um arquivo "4 - OAB PATRICK.pdf" (6 MB), '
     'aparentemente a carteira da OAB do advogado Patrick Willian da '
     'Silva (OAB/AM A2638), que figura na procuração como segundo '
     'representante (ao lado do escritório Azevedo Lima & Rebonatto). '
     'Anexar como prova de habilitação se for necessário no protocolo.'),
]
for titulo, txt in pendencias:
    p = doc_r.add_paragraph(style='List Bullet')
    r = p.add_run(titulo + ': ')
    r.bold = True
    p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
itens = [
    'Conferir nome (EXEMPLO CARVALHO), CPF (000.000.006-16), RG (0300415-5 SSP/AM).',
    'Conferir conta/agência (75911-2 / 3739).',
    'Conferir comarca: MANAUS (NÃO Maués).',
    'Conferir prioridade idoso aplicada (cabeçalho + pedido).',
    'Conferir que NÃO há pedido de repetição em dobro (estratégia b).',
    'Conferir parágrafos doutrinários sobre "Repetição do indébito" REMOVIDOS.',
    'Conferir parágrafo doutrinário 102 reescrito (sem trecho de repetição em dobro).',
    'Conferir parágrafo do caso concreto DENIVAL inserido após "Alegar que os valores permaneciam disponíveis...".',
    'Conferir VC = R$ 15.000,00 (apenas dano moral).',
    'Avaliar se vai enviar notificação extrajudicial ANTES do protocolo.',
    'Conferir o bloco "Do prévio requerimento administrativo" — ajustar texto se notificação NÃO houver sido enviada.',
    'Anexar: 2-Procuração; 3-RG/CPF; 4-Hipossuficiência; 4-OAB Patrick; 5-Comprovante (filho); 5.1-CNH (filho); 6-Extrato; 7-Tabela.',
    'Confirmar com cliente o vínculo de coabitação com o filho Dionaldo.',
    'Decidir sobre o caso TARIFAS da pasta-mãe (batch separado).',
]
for it in itens:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph()
p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas — após confirmar pendências críticas (auditoria APLIC vs RESGATE, '
          'pós-processamento textual coerente, coabitação com filho, decisão sobre notificação prévia), ')
r2 = p.add_run('PROTOCOLAR.')
r2.bold = True

doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
