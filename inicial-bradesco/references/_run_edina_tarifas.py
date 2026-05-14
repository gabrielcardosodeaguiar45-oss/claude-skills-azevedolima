"""Inicial TARIFAS — CLIENTE EXEMPLO RODRIGUES PINTO.

Comarca Maués/AM. Pessoa IDOSA (afirmação expressa na notificação extrajudicial).
Conta Bradesco Ag 3706 / 21299-7. INSS R$ 947,57 (último crédito 27/02/2026 — líquido).
Tabela: 55 lançamentos entre 13/05/2022 e 27/02/2026 (3 rubricas).
Total simples R$ 1.700,50 / dobro R$ 3.401,00. VC R$ 18.401,00 — cabe no JEC.

Rubricas:
  TARIFA BANCARIA - CESTA B EXPRESSO
  CARTAO CREDITO ANUIDADE
  PACOTE DE SERVICOS PADRONIZADO PRIORITARIOS
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO RODRIGUES PINTO - Maurivã - TARIFAS\TARIFAS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

LANCAMENTOS = [
    ('13/05/2022', 44.50), ('15/06/2022', 1.20), ('30/06/2022', 43.30),
    ('18/07/2022', 44.50), ('10/08/2022', 19.25), ('17/08/2022', 44.50),
    ('15/09/2022', 44.50), ('18/10/2022', 49.90), ('18/11/2022', 49.90),
    ('15/12/2022', 49.90), ('13/01/2023', 49.90), ('15/02/2023', 49.90),
    ('15/03/2023', 49.90), ('14/04/2023', 49.90), ('15/05/2023', 49.90),
    ('15/06/2023', 49.90), ('14/07/2023', 49.90), ('15/08/2023', 51.60),
    ('15/09/2023', 51.60), ('13/10/2023', 51.60), ('14/11/2023', 51.60),
    ('15/12/2023', 51.60), ('15/01/2024', 51.60), ('15/02/2024', 56.75),
    ('15/03/2024', 56.75), ('15/04/2024', 56.75), ('15/05/2024', 56.75),
    ('14/06/2024', 56.75), ('15/07/2024', 56.75), ('15/08/2024', 15.95),
    ('13/09/2024', 15.95), ('15/10/2024', 15.95), ('14/11/2024', 15.95),
    ('13/12/2024', 15.95), ('30/05/2025', 16.35), ('30/05/2025', 15.95),
    ('30/05/2025', 16.35), ('30/05/2025', 15.95), ('30/05/2025', 16.35),
    ('13/06/2025', 9.26), ('30/06/2025', 7.09), ('15/07/2025', 6.11),
    ('31/07/2025', 10.24), ('15/08/2025', 2.96), ('29/08/2025', 13.39),
    ('15/09/2025', 0.81), ('19/09/2025', 15.84), ('15/10/2025', 16.65),
    ('14/11/2025', 0.22), ('28/11/2025', 16.43), ('15/12/2025', 3.53),
    ('30/12/2025', 13.12), ('30/01/2026', 16.65), ('13/02/2026', 0.48),
    ('27/02/2026', 16.17),
]

autora = {
    'nome': 'CLIENTE EXEMPLO RODRIGUES PINTO',
    'nacionalidade': 'brasileira',
    'estado_civil': '',
    'profissao': 'aposentada',
    'cpf': '000.000.009-19',
    'rg': '1000007-7',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Ponta Alegre Apocuitaua',
    'numero': '2601',
    'bairro': 'Rio Apocuitaua',
    'cidade': 'Maués',
    'cep': '69.190-000',
}
conta = {'agencia': '3706', 'numero': '21299-7'}
renda = {'valor_float': 947.57}

tese = {
    'rubrica': 'TARIFA BANCÁRIA - CESTA B EXPRESSO / CARTÃO CRÉDITO ANUIDADE / PACOTE DE SERVIÇOS PADRONIZADO PRIORITÁRIOS',
    'lancamentos': LANCAMENTOS,
}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Maués', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA B EXPRESSO / CARTÃO CRÉDITO ANUIDADE / PACOTE DE SERVIÇOS PADRONIZADO PRIORITÁRIOS'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw para placeholders no XML que escapam ao run-aware
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

# Relatório paralelo
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFAS (3 rubricas)'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA — art. 1.048, I, CPC (afirmação expressa na notificação)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 27/02/2026 — líquido)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (3 rubricas combinadas)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('IDOSA — prioridade aplicada',
     'Notificação extrajudicial afirma expressamente que a parte autora é "pessoa idosa". '
     'Prioridade do art. 1.048, I, do CPC inserida no cabeçalho. Conferir no RG a data '
     'de nascimento (≥ 60 anos) antes do protocolo.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente. Confirmar com '
     'cliente antes do protocolo.'),
    ('Renda INSS LÍQUIDA — pode ser bruta superior',
     f'O extrato traz INSS R$ 947,57 em 27/02/2026 (último crédito). Valor reduzido '
     'sugere descontos consignados pelo próprio INSS antes do crédito. Renda BRUTA '
     'do benefício pode ser maior. Conferir HISCON se houver e ajustar se necessário '
     '(reforça hipossuficiência e impacto do dano moral).'),
    ('3 RUBRICAS combinadas em 1 só inicial',
     'Tabela cobre TARIFA BANCARIA CESTA B EXPRESSO + CARTAO CREDITO ANUIDADE + '
     'PACOTE DE SERVICOS PADRONIZADO PRIORITARIOS no MESMO período. Tratadas como '
     '1 SÓ TESE (TARIFAS) por afinidade temática (cobrança não autorizada de tarifas '
     'bancárias) — IRDR 0005053 TJAM. VC dobro = R$ 3.401,00 + dano moral R$ 15.000 '
     '= R$ 18.401,00. Cabe folgadamente no JEC.'),
    ('5 lançamentos em 30/05/2025 — datas duplicadas',
     'Em 30/05/2025 a tabela registra 5 lançamentos do PACOTE DE SERVIÇOS '
     '(R$ 16,35 + 15,95 + 16,35 + 15,95 + 16,35 = R$ 80,95 num único dia). Padrão '
     'atípico — o banco fez um lote de débitos retroativos. Confirmar no extrato '
     'original que esses 5 débitos efetivamente ocorreram no mesmo dia.'),
    ('Procurações específicas por rubrica',
     'A pasta tem 3 procurações específicas (CARTAO CREDITO ANUIDADE / PACOTE DE '
     'SERVICOS / TARIFA BANCARIA CESTA B EXPRESSO). Anexar TODAS as 3 antes do '
     'protocolo (cobertura para cada rubrica).'),
    ('TETO JEC — coberto',
     f'VC R$ 18.401,00 ≈ 12,1 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos confirmado).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Maués/AM.',
    'Conferir VC = R$ 18.401,00 e dano moral R$ 15.000,00.',
    'Confirmar com cliente que NUNCA contratou pacote de serviços, cartão de crédito ou cesta de tarifas.',
    'Conferir grifo amarelo + formatação rubrica (CAPS+bold+itálico+sublinhado).',
    'Conferir nome da autora em Segoe UI Bold (rStyle 2TtuloChar).',
    'Anexar 3 procurações específicas + RG + Hipossuficiência + Comprovante residência + Extrato + Tabela + Notificação + Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas (idade no RG, renda HISCON). Após confirmar com cliente, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
