"""Inicial TÍTULO DE CAPITALIZAÇÃO — CLIENTE EXEMPLO CAVALCANTE SANTANA."""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\3. Título de capitalização\CLIENTE EXEMPLO CAVALCANTE SANTANA - Ruth - TARIFAS\TÍTULO DE CAPITALIZAÇÃO'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Titulo_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_TITULO_v1.docx')

LANCAMENTOS = [('16/09/2025', 100.00), ('03/11/2025', 20.00), ('03/12/2025', 20.00)]

autora = {
    'nome': 'CLIENTE EXEMPLO CAVALCANTE SANTANA', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.015-25', 'rg': '1000013-3',
    'orgao_expedidor_prefixo': 'CIN (CPF como Registro Geral)',
    'logradouro': 'Rua Manoel (RM) Jardim Floresta', 'numero': 's/nº',
    'bairro': 'Centro', 'cidade': 'Presidente Figueiredo', 'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '510965-5'}
renda = {'valor_float': 1621.00}

tese = {'rubrica': 'TÍTULO DE CAPITALIZAÇÃO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TÍTULO DE CAPITALIZAÇÃO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TÍTULO ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Titulo_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TÍTULO DE CAPITALIZAÇÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Titulo_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CIN nova)', f'{autora["rg"]} — Registro Geral = CPF'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]}'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)}'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('TEMPLATE GENÉRICO ADAPTADO', 'Sem template próprio para TÍTULO. Usado inicial-tarifas. Ajustar fundamentação para TC SUSEP nº 415, art. 39 VI CDC, jurisprudência sobre venda casada.'),
    ('RG (CIN nova) confirmado', 'Carteira de Identidade Nacional: Registro Geral = CPF (000.000.015-25). Nascimento 13/03/1966 (60 anos — IDOSA).'),
    ('Caso recente — sem prescrição', 'Lançamentos de set-dez/2025.'),
    ('LANÇAMENTO INICIAL R$ 100', 'Em 16/09/2025 a tabela registra R$ 100 (vs R$ 20 dos demais). Pode ser parcela inicial. Confirmar.'),
    ('CLIENTE TEM 2 TESES', 'CLIENTE EXEMPLO também aparece em TARIFAS (já gerada). AVALIAR consolidação.'),
    ('TETO JEC — coberto', 'VC R$ 15.280,00 ≈ 10 SM.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'AJUSTAR fundamentação para TÍTULO.',
    'AVALIAR consolidação com TARIFAS.',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Confirmar com cliente: nunca contratou título de capitalização.',
    'Anexar 2-Procuração + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem ajustes. ')
r2 = p.add_run('Atenção.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
