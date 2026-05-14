"""Inicial COMBINADA TARIFAS — CLIENTE EXEMPLO CAVALCANTE SANTANA.

2 procurações:
  1. PACOTE DE SERVIÇOS - PADRONIZADO PRIORITÁRIOS
  2. TARIFA BANCÁRIA - CESTA B EXPRESSO

Comarca Pres. Figueiredo/AM, IDOSA 60 anos, Ag 3732 / Conta 510965-5,
INSS R$ 1.621.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _combinada_helper import gerar_combinada
from docx import Document

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO CAVALCANTE SANTANA - Ruth - TARIFAS\TARIFAS'

CESTA = [
    ('15/09/2022', 15.16), ('14/04/2023', 9.48), ('08/05/2023', 40.42),
    ('15/05/2023', 3.58), ('07/06/2023', 46.32), ('15/06/2023', 8.68),
    ('07/07/2023', 41.22),
]
PACOTE = [
    ('07/08/2023', 15.45), ('15/08/2023', 3.98), ('08/09/2023', 11.47),
    ('15/12/2023', 8.67), ('08/01/2024', 6.78), ('15/01/2024', 1.22),
    ('27/05/2025', 15.95), ('27/05/2025', 15.95), ('13/06/2025', 8.10),
    ('16/09/2025', 16.35), ('16/09/2025', 16.35), ('16/09/2025', 16.65),
    ('16/09/2025', 8.25),
]

teses = [
    {'familia': 'TARIFAS', 'rubrica': 'TARIFA BANCÁRIA - CESTA B EXPRESSO', 'lancamentos': CESTA},
    {'familia': 'TARIFAS', 'rubrica': 'PACOTE DE SERVIÇOS PADRONIZADO PRIORITÁRIOS', 'lancamentos': PACOTE},
]

autora = {
    'nome': 'CLIENTE EXEMPLO CAVALCANTE SANTANA', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.015-25', 'rg': '1000013-3',
    'orgao_expedidor_prefixo': 'CIN (CPF como Registro Geral)',
    'logradouro': 'Rua Manoel (RM) Jardim Floresta', 'numero': 's/nº',
    'bairro': 'Centro', 'cidade': 'Presidente Figueiredo', 'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '510965-5'}
renda = {'valor_float': 1621.00, 'descricao': 'aposentadoria pelo INSS'}

res = gerar_combinada(
    pasta=PASTA, nome_arquivo_base='INICIAL_Combinada_CLIENTE EXEMPLO',
    autora=autora, conta=conta, renda=renda, teses=teses,
    comarca='Presidente Figueiredo', uf='AM', eh_idoso=True,
)
print(f'CLIENTE EXEMPLO — combinada {res["totais"]["n_teses"]} teses')
print(f'  Total: R$ {res["totais"]["total_geral"]:.2f} / dobro R$ {res["totais"]["dobro_geral"]:.2f}')
print(f'  Dano moral: R$ {res["totais"]["dano_moral"]:.2f}')
print(f'  VC: R$ {res["totais"]["valor_causa"]:.2f}')
print(f'  Parágrafos removidos: {res["paras_removidos"]}, pós-fix: {res["pos_fix"]}')

doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Combinada_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']),
             ('Tese', f'COMBINADA TARIFAS ({len(teses)} rubricas)'),
             ('Comarca', 'Presidente Figueiredo/AM'),
             ('Arquivo', 'INICIAL_Combinada_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)
doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
linhas = [
    ('Comarca', 'Presidente Figueiredo/AM'),
    ('Prioridade', 'IDOSA 60 anos (13/03/1966)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CIN nova)', f'{autora["rg"]} — Registro Geral = CPF'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', 'R$ 1.621,00 (1 SM 2026)'),
]
for i, t in enumerate(teses, 1):
    linhas.append((f'Tese {i} — {t["rubrica"]}',
                   f'{t["n_lanc"]} lanç. {t["inicio"]}–{t["fim"]} = R$ {t["total"]:.2f} / R$ {t["dobro"]:.2f} dobro'))
linhas.extend([
    ('TOTAL combinado', f'R$ {res["totais"]["total_geral"]:.2f} / R$ {res["totais"]["dobro_geral"]:.2f} dobro'),
    ('Dano moral (2 teses x R$ 5.000)', f'R$ {res["totais"]["dano_moral"]:.2f}'),
    ('Valor da causa', f'R$ {res["totais"]["valor_causa"]:.2f}'),
])
for k, v in linhas:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('Inicial COMBINADA com 2 núcleos fáticos', '2 procurações distintas → 2 núcleos individualizados, 1 bloco doutrinário TARIFAS comum, 2 pedidos.'),
    ('IDOSA confirmada 60 anos', 'CIN nova, nascimento 13/03/1966, naturalidade Manaus/AM.'),
    ('VALOR BAIXO (dobro R$ 620,06)', 'Combinação com TÍTULO CAPITALIZAÇÃO (pasta separada) ainda recomendada para fortalecer o caso.'),
    ('CLIENTE EXEMPLO também tem TÍTULO em pasta separada', 'TÍTULO continua em ação SEPARADA por enquanto. Avaliar se procurador quer mega-combinada com TÍTULO.'),
    ('TETO JEC — coberto', f'VC R$ {res["totais"]["valor_causa"]:.2f}.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'CONFERIR visualmente combinada (2 núcleos TARIFAS, sem MORA/TÍTULO/APLIC).',
    'Conferir nome / CPF / nascimento.',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    f'Conferir VC R$ {res["totais"]["valor_causa"]:.2f} e dano moral R$ {res["totais"]["dano_moral"]:.2f}.',
    'Confirmar com cliente: nunca contratou pacote de serviços nem cesta de tarifas.',
    'Anexar 2 procurações + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
    'AVALIAR mega-combinada com TÍTULO (3ª tese em pasta separada).',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após conferência visual, ')
r = p.add_run('PROTOCOLAR.'); r.bold = True
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_COMBINADA_v1.docx')
doc_r.save(RELAT_OUT)
print(f'  -> {RELAT_OUT}')
