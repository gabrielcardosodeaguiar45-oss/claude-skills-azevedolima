"""Inicial COMBINADA TARIFAS — CLIENTE EXEMPLO RODRIGUES PINTO.

3 procurações distintas na subpasta TARIFAS (Bradesco):
  1. CARTÃO CRÉDITO ANUIDADE
  2. PACOTE DE SERVIÇOS - PADRONIZADO PRIORITÁRIOS
  3. TARIFA BANCÁRIA - CESTA B EXPRESSO

Comarca Maués/AM, IDOSA 60 anos, Ag 3706 / Conta 21299-7,
INSS R$ 947,57.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _combinada_helper import gerar_combinada
from docx import Document

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO RODRIGUES PINTO - Maurivã - TARIFAS\TARIFAS'

# Lançamentos da tabela única, separados por rubrica
CARTAO = [
    ('10/08/2022', 19.25),
]
PACOTE = [
    ('15/08/2024', 15.95), ('13/09/2024', 15.95), ('15/10/2024', 15.95),
    ('14/11/2024', 15.95), ('13/12/2024', 15.95),
    ('30/05/2025', 16.35), ('30/05/2025', 15.95), ('30/05/2025', 16.35),
    ('30/05/2025', 15.95), ('30/05/2025', 16.35),
    ('13/06/2025', 9.26), ('30/06/2025', 7.09), ('15/07/2025', 6.11),
    ('31/07/2025', 10.24), ('15/08/2025', 2.96), ('29/08/2025', 13.39),
    ('15/09/2025', 0.81), ('19/09/2025', 15.84), ('15/10/2025', 16.65),
    ('14/11/2025', 0.22), ('28/11/2025', 16.43), ('15/12/2025', 3.53),
    ('30/12/2025', 13.12), ('30/01/2026', 16.65), ('13/02/2026', 0.48),
    ('27/02/2026', 16.17),
]
CESTA = [
    ('13/05/2022', 44.50), ('15/06/2022', 1.20), ('30/06/2022', 43.30),
    ('18/07/2022', 44.50), ('17/08/2022', 44.50), ('15/09/2022', 44.50),
    ('18/10/2022', 49.90), ('18/11/2022', 49.90), ('15/12/2022', 49.90),
    ('13/01/2023', 49.90), ('15/02/2023', 49.90), ('15/03/2023', 49.90),
    ('14/04/2023', 49.90), ('15/05/2023', 49.90), ('15/06/2023', 49.90),
    ('14/07/2023', 49.90), ('15/08/2023', 51.60), ('15/09/2023', 51.60),
    ('13/10/2023', 51.60), ('14/11/2023', 51.60), ('15/12/2023', 51.60),
    ('15/01/2024', 51.60), ('15/02/2024', 56.75), ('15/03/2024', 56.75),
    ('15/04/2024', 56.75), ('15/05/2024', 56.75), ('14/06/2024', 56.75),
    ('15/07/2024', 56.75),
]

teses = [
    {'familia': 'TARIFAS', 'rubrica': 'CARTÃO CRÉDITO ANUIDADE', 'lancamentos': CARTAO},
    {'familia': 'TARIFAS', 'rubrica': 'PACOTE DE SERVIÇOS PADRONIZADO PRIORITÁRIOS', 'lancamentos': PACOTE},
    {'familia': 'TARIFAS', 'rubrica': 'TARIFA BANCÁRIA - CESTA B EXPRESSO', 'lancamentos': CESTA},
]

autora = {
    'nome': 'CLIENTE EXEMPLO RODRIGUES PINTO', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.009-19', 'rg': '1000007-7',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Ponta Alegre Apocuitaua', 'numero': '2601',
    'bairro': 'Rio Apocuitaua', 'cidade': 'Maués', 'cep': '69.190-000',
}
conta = {'agencia': '3706', 'numero': '21299-7'}
renda = {'valor_float': 947.57, 'descricao': 'aposentadoria pelo INSS'}

res = gerar_combinada(
    pasta=PASTA, nome_arquivo_base='INICIAL_Combinada_CLIENTE EXEMPLO',
    autora=autora, conta=conta, renda=renda, teses=teses,
    comarca='Maués', uf='AM', eh_idoso=True,
)
print(f'CLIENTE EXEMPLO — combinada {res["totais"]["n_teses"]} teses')
print(f'  Total: R$ {res["totais"]["total_geral"]:.2f} / dobro R$ {res["totais"]["dobro_geral"]:.2f}')
print(f'  Dano moral ({res["totais"]["n_teses"]} x R$ 5.000): R$ {res["totais"]["dano_moral"]:.2f}')
print(f'  VC: R$ {res["totais"]["valor_causa"]:.2f}')
print(f'  Parágrafos removidos: {res["paras_removidos"]}, pós-fix: {res["pos_fix"]}')
print(f'  -> {res["docx"]}')

# Relatório paralelo
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Combinada_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']),
             ('Tese', f'COMBINADA TARIFAS ({len(teses)} rubricas)'),
             ('Comarca', 'Maués/AM'),
             ('Arquivo', 'INICIAL_Combinada_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
linhas = [
    ('Comarca', 'Maués/AM'),
    ('Prioridade', 'IDOSA 60 anos (07/08/1965, RG OCR confirmado)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', 'R$ 947,57 (INSS último crédito 27/02/2026)'),
]
for i, t in enumerate(teses, 1):
    linhas.append((f'Tese {i} — {t["rubrica"]}',
                   f'{t["n_lanc"]} lanç. {t["inicio"]}–{t["fim"]} = R$ {t["total"]:.2f} simples / R$ {t["dobro"]:.2f} dobro'))
linhas.extend([
    ('TOTAL combinado', f'R$ {res["totais"]["total_geral"]:.2f} simples / R$ {res["totais"]["dobro_geral"]:.2f} dobro'),
    ('Dano moral (3 teses x R$ 5.000)', f'R$ {res["totais"]["dano_moral"]:.2f}'),
    ('Valor da causa', f'R$ {res["totais"]["valor_causa"]:.2f}'),
])
for k, v in linhas:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('Inicial COMBINADA com 3 núcleos fáticos individualizados',
     'Subpasta TARIFAS contém 3 procurações distintas (Cartão Anuidade + Pacote de '
     'Serviços + Cesta B Expresso). Inicial estruturada com 3 núcleos fáticos '
     'autônomos, 1 bloco doutrinário TARIFAS comum (todas seguem IRDR Tema 5 TJAM) '
     'e 3 pedidos de restituição em dobro.'),
    ('IDOSA confirmada 60 anos', 'RG 0946698-3, nascimento 07/08/1965 (60 anos completos em 07/08/2025).'),
    ('Renda < 1 SM — possível consignação', 'INSS R$ 947,57 < salário mínimo. Conferir HISCON.'),
    ('CARTÃO ANUIDADE com 1 só lançamento', 'Tese tarifária com APENAS R$ 19,25 em 10/08/2022. Pode reduzir o impacto, mas é parte da combinada.'),
    ('5 lançamentos PACOTE em 30/05/2025 — duplicidade', 'Confirmar no extrato.'),
    ('Estado civil — não informado', 'Notificação não traz. Placeholder OMITIDO.'),
    ('TETO JEC — coberto', f'VC R$ {res["totais"]["valor_causa"]:.2f} cabe folgadamente no JEC.'),
    ('CONFERIR pós-processamento', 'Verificar visualmente que blocos MORA, TÍTULO e APLIC NÃO aparecem na peça.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'CONFERIR visualmente combinada (3 núcleos TARIFAS, sem MORA/TÍTULO/APLIC).',
    'Conferir nome / CPF / RG / nascimento.',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    f'Conferir VC R$ {res["totais"]["valor_causa"]:.2f} e dano moral R$ {res["totais"]["dano_moral"]:.2f}.',
    'Confirmar com cliente: nunca contratou cartão Bradesco, pacote de serviços, cesta de tarifas.',
    'Anexar 3 procurações + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
    'MORA da CLIENTE EXEMPLO continua em ação SEPARADA (subpasta MORA E ENCARGOS, em 2. MORA).',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após conferência visual, ')
r = p.add_run('PROTOCOLAR.'); r.bold = True
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_COMBINADA_v1.docx')
doc_r.save(RELAT_OUT)
print(f'  -> {RELAT_OUT}')
