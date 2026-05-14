# -*- coding: utf-8 -*-
"""Lista todas as ocorrências de rubricas Bradesco hardcoded nos templates."""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
from docx import Document

VAULT_TEMPLATES = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates")

# Padrões a procurar (case-insensitive); ordem importa pra evitar overlap
PADROES = [
    # MORA — variantes literais (forma curta como aparece no extrato)
    (r"Mora\s+Cred\s+Pess(?!oal)", "MORA_LITERAL"),
    (r"MORA\s+CRED\s+PESS(?!OAL)", "MORA_LITERAL_UPPER"),
    # MORA — forma canônica/expandida
    (r"Mora\s+Cr[ée]dito\s+Pessoal", "MORA_CANONICA"),
    (r"Cr[ée]dito\s+Mora\s+Pessoal", "MORA_CANONICA_INV"),
    (r"MORA\s+CREDITO\s+PESSOAL", "MORA_CANONICA_UPPER"),
    # ENCARGO — variantes literais
    (r"Enc\.?\s+Lim\.?\s+Cr[ée]dito", "ENCARGO_LITERAL"),
    (r"ENC\.?\s+LIM\.?\s+CR[ÉE]DITO", "ENCARGO_LITERAL_UPPER"),
    # ENCARGO — forma canônica
    (r"Encargos?\s+Limite\s+Cr[ée]dito", "ENCARGO_CANONICA"),
    (r"ENCARGOS?\s+LIMITE\s+CR[ÉE]DITO", "ENCARGO_CANONICA_UPPER"),
]

def varrer(caminho_docx):
    doc = Document(caminho_docx)
    achados = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        for pat, tag in PADROES:
            for m in re.finditer(pat, t):
                # Pega contexto curto
                ini = max(0, m.start() - 30)
                fim = min(len(t), m.end() + 30)
                achados.append({
                    "par": i, "tag": tag, "match": m.group(0),
                    "contexto": t[ini:fim].replace("\n", " ")
                })
    # Tabelas
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    t = p.text
                    for pat, tag in PADROES:
                        for m in re.finditer(pat, t):
                            achados.append({
                                "par": f"T{ti}R{ri}C{ci}P{pi}", "tag": tag,
                                "match": m.group(0), "contexto": t[:80]
                            })
    return achados

for nome in ["inicial-mora.docx", "inicial-mora-encargo.docx", "inicial-combinada.docx"]:
    caminho = VAULT_TEMPLATES / nome
    if not caminho.exists():
        print(f"\n=== {nome}: NÃO EXISTE ===")
        continue
    print(f"\n=== {nome} ===")
    achados = varrer(caminho)
    if not achados:
        print("  (nenhuma ocorrência)")
        continue
    # Agrupar por tag
    from collections import defaultdict
    por_tag = defaultdict(list)
    for a in achados:
        por_tag[a["tag"]].append(a)
    for tag, lst in por_tag.items():
        print(f"  {tag}: {len(lst)} ocorrência(s)")
        for a in lst[:3]:
            print(f"    par {a['par']}: '{a['match']}' — ...{a['contexto'][:80]}...")
        if len(lst) > 3:
            print(f"    [+{len(lst)-3} adicional(is)]")
