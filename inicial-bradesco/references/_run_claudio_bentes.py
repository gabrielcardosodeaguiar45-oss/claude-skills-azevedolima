"""Inicial TARIFAS — CLIENTE EXEMPLO CARNEIRO.

Comarca Barreirinha/AM. Conta Bradesco Ag 3725/512140-0. Aposentado INSS
R$ 1.621,00 (último crédito 04/02/2026). Idade não confirmada (RG não
lido) — assumir NÃO IDOSO.

Tese: TARIFA BANCARIA CESTA B.EXPRESSO1 — 96 lançamentos entre
04/08/2021 e 03/03/2026 (total R$ 2.942,35; dobro R$ 5.884,70). VC =
dobro + R$ 15.000 dano moral = R$ 20.884,70. Cabe folgadamente no JEC.
"""
import io, sys, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\CLIENTE EXEMPLO CARNEIRO - Wilson - TARIFAS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_CLAUDIO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLAUDIO_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_claudio_lancs.json')) as f:
    LANCS_RAW = json.load(f)
LANCAMENTOS = sorted([(d, v) for d, t, v in LANCS_RAW],
                     key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO CARNEIRO', 'nacionalidade': 'brasileiro',
    'estado_civil': '', 'profissao': 'aposentado',
    'cpf': '000.000.004-14', 'rg': '1000003-3',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Olaria', 'numero': '32',
    'bairro': 'Santa Luzia', 'cidade': 'Barreirinha', 'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '512140-0'}
renda = {'valor_float': 1621.00}

tese = {'rubrica': 'TARIFA BANCÁRIA CESTA B.EXPRESSO1', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Barreirinha', uf='AM')
from extenso import extenso_cardinal
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA CESTA B.EXPRESSO1'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | dobro: R$ {calc["dobro"]:,.2f} | VC: R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
from extenso import extenso_cardinal
substituicoes_raw = {
    '{{remuneração}}': dados['remuneração'],
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
    '{{titulo}}': 'TARIFA BANCÁRIA CESTA B.EXPRESSO1',
    '{{numero_desconto}}': str(len(LANCAMENTOS)),
    '{{desconto_extenso}}': extenso_cardinal(len(LANCAMENTOS)),
    '{{total_descontos}}': dados['total_descontos'],
    '{{total_descontos_extenso}}': dados['total_descontos_extenso'],
}
fix = 0
for k, v in substituicoes_raw.items():
    if k in xml:
        cnt = xml.count(k)
        xml = xml.replace(k, v)
        fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes:
        z.writestr(n, buf[n])
print(f'Pós-fix: {fix} substituições')
print(f'OK -> {DOCX_OUT}')

# Relatório
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_CLAUDIO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA CESTA B.EXPRESSO1'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_CLAUDIO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Estado civil', '(omitido — confirmar)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 04/02/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos TARIFA', '96 (deduplicados — extrato PDF traz a rubrica em 2 linhas)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro (pleiteado)', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('IDADE — não confirmada',
     'RG não lido nesta sessão. Assume NÃO IDOSO. Conferir antes do protocolo.'),
    ('TEMPLATE TARIFAS aplicado normalmente',
     'Template inicial-tarifas.docx aplicado sem pós-processamento estrutural — pleito '
     'inclui repetição em dobro padrão (CDC art. 42 p.ún.).'),
    ('TETO JEC — coberto',
     f'VC R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.') + ' ≈ 13,75 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (512140-0 / 3725).',
           'Conferir comarca: Barreirinha/AM.', 'Confirmar idade no RG (>60? aplicar prioridade).',
           'Anexar 8-NOTIFICACAO + 8.1-COMPROVANTE NOTIFICAÇÃO (já estão na pasta).']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — '); r2 = p.add_run('PROTOCOLAR após confirmar idade.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
