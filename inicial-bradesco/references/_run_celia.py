"""Inicial TARIFAS — CLIENTE EXEMPLO RODRIGUES DA SILVA.

Comarca Maués/AM. Dados extraídos via OCR v2 (easyocr alta resolução +
parser robusto) sobre o extrato escaneado de 38 páginas.

Conta Bradesco Ag 3706 / 16649-9. INSS R$ 980,09 (último crédito
03/02/2026). 100 lançamentos TARIFA BANCÁRIA CESTA B.EXPRESSO (cheia
+ VR.PARCIAL) entre 15/01/2020 e 02/10/2025 — total R$ 2.519,99 /
dobro R$ 5.039,98. VC R$ 20.039,98 — cabe folgadamente no JEC.

REPROCESSAMENTO 09/05/2026: a versão v1 anterior (32 lançamentos,
R$ 558,49) tinha 3 falhas críticas:
  1. Parser exigia data + valor na mesma linha → perdeu lançamentos
     onde a data vinha em linha separada acima do bloco
  2. Filtro de valores zerados destruiu 15 lançamentos com OCR ruim
     que deveriam ter virado pendência manual
  3. Variantes "VR.PARCIAL CESTA B.EXPRESSO" não foram capturadas
     (~45 ocorrências perdidas)
Adicional: a v1 incluía erroneamente "TARIFA EMISSAO EXTRATO" (3
lançamentos: 04/10/2021, 02/02/2023 dup, 13/03/2025) que NÃO é
pleiteável. Versão v2 filtra corretamente apenas TARIFA BANCARIA.

Backup da lista antiga em _cliente exemplo_lancs.json.bak_pre_v2.
"""
import io, sys, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\CLIENTE EXEMPLO RODRIGUES DA SILVA - Maurivã - TARIFAS\TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_lancs.json')) as f:
    LANCAMENTOS = sorted(json.load(f),
                         key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO RODRIGUES DA SILVA', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.003-13', 'rg': '1000002-2',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua 10', 'numero': '818',
    'bairro': 'Nova Esperança', 'cidade': 'Maués', 'cep': '69.190-000',
}
conta = {'agencia': '3706', 'numero': '16649-9'}
renda = {'valor_float': 980.09}

tese = {'rubrica': 'TARIFA BANCÁRIA CESTA B.EXPRESSO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Maués', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA CESTA B.EXPRESSO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS (OCR) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | dobro: R$ {calc["dobro"]:,.2f} | VC: R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

# Relatório
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA CESTA B.EXPRESSO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]} (extraído via OCR — confirmar)'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 03/02/2026 via OCR)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos TARIFA', f'{len(LANCAMENTOS)} (cheios + VR.PARCIAL via OCR v2)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('OCR v2 (alta resolução + parser robusto) — REVISADO 09/05/2026',
     'O extrato 6-EXTRATO BANCÁRIO.pdf é PDF escaneado sem text-layer. Reprocessado '
     'via easyocr em DPI alto (max_dim 2200px) + parser que: (a) agrupa tokens em '
     'linhas visuais por coordenada Y, (b) usa data inferida do número de docto '
     'quando ausente na linha, (c) captura variantes "VR.PARCIAL CESTA B.EXPRESSO" '
     '(saldo parcial cobrado quando a conta não tem saldo cheio), (d) filtra '
     'corretamente "TARIFA EMISSAO EXTRATO" que NÃO é pleiteável. Versão v1 '
     'anterior (32 lançamentos / R$ 558,49) tinha falhas de parser e foi substituída.'),
    ('Tipos de lançamento capturados',
     f'Total {len(LANCAMENTOS)}. Inclui tarifas CHEIAS (CESTA B.EXPRESSO completa) e '
     'PARCIAIS (VR.PARCIAL — Bradesco cobra em parcelas quando a conta não tem '
     'saldo suficiente). Ambas configuram a mesma cobrança indevida (mesma '
     'tese IRDR Tema 5 TJAM).'),
    ('Conta/Agência extraídos via OCR',
     'Ag 3706 / Conta 16649-9 — confirmar antes do protocolo (cabeçalho do extrato OCR).'),
    ('Renda confirmada via OCR',
     'INSS R$ 980,09 (último crédito 03/02/2026). Histórico mostra valores entre R$ 797 '
     'e R$ 1.278 em 2024-2026.'),
    ('IDADE — não confirmada',
     'RG não lido. Assume NÃO IDOSO. Conferir antes do protocolo.'),
    ('TETO JEC — coberto',
     f'VC R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.') + '. Cabe no JEC.'),
    ('OUTRAS TESES detectadas no OCR',
     'O extrato também mostra MORA CREDITO PESSOAL e PARCELA CREDITO PESSOAL recorrentes. '
     'Cliente é candidata para inicial COMBINADA (TARIFA + MORA + ENCARGO). Decidir '
     'em batch separado se vai gerar inicial-combinada.docx.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir Conta/Agência (16649-9 / 3706).',
           'Conferir comarca: Maués/AM.', 'CONFERIR lançamentos vs extrato original (OCR pode ter falhas).',
           'Confirmar idade no RG.',
           'Anexar 8-NOTIFICAÇÃO + 8.1-COMPROVANTE NOTIFICAÇÃO (já estão na pasta).',
           'Avaliar inicial COMBINADA com MORA + ENCARGO.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas — CONFERIR OCR contra extrato original e '); r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
