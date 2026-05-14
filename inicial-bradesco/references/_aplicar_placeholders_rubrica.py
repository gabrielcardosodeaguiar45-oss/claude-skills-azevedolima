# -*- coding: utf-8 -*-
"""Aplica os 8 placeholders de rubrica nos 3 templates Bradesco que tratam
mora/encargo. Faz backup antes; preserva formatação via substituir_in_run."""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
from docx import Document
from helpers_docx import substituir_in_run

VAULT = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates")

# Mapa GLOBAL: literal → placeholder
# Ordem importa: variantes mais específicas/longas primeiro pra não dar conflito
MAPA_RUBRICAS = {
    # ==== MORA — formas canônicas/expandidas (mais específico primeiro) ====
    "MORA CREDITO PESSOAL": "{{rubrica_mora_canonica_caps}}",
    "MORA CRÉDITO PESSOAL": "{{rubrica_mora_canonica_caps}}",
    "Mora Crédito Pessoal": "{{rubrica_mora_canonica}}",
    "Mora Credito Pessoal":  "{{rubrica_mora_canonica}}",
    "Crédito Mora Pessoal":  "{{rubrica_mora_canonica}}",
    "Credito Mora Pessoal":  "{{rubrica_mora_canonica}}",

    # ==== MORA — formas literais/curtas ====
    "MORA CRED PESS": "{{rubrica_mora_caps}}",
    "Mora Cred Pess": "{{rubrica_mora}}",

    # ==== ENCARGO — formas canônicas/expandidas ====
    "ENCARGOS LIMITE DE CRED":  "{{rubrica_encargo_canonica_caps}}",
    "ENCARGOS LIMITE CRÉDITO":  "{{rubrica_encargo_canonica_caps}}",
    "ENCARGOS LIMITE CREDITO":  "{{rubrica_encargo_canonica_caps}}",
    "Encargos Limite Crédito":  "{{rubrica_encargo_canonica}}",
    "Encargos Limite Credito":  "{{rubrica_encargo_canonica}}",

    # ==== ENCARGO — formas literais/curtas ====
    "ENC. LIM. CRÉDITO": "{{rubrica_encargo_caps}}",
    "ENC. LIM. CREDITO": "{{rubrica_encargo_caps}}",
    "ENC LIM CRÉDITO":   "{{rubrica_encargo_caps}}",
    "ENC LIM CREDITO":   "{{rubrica_encargo_caps}}",
    "Enc. Lim. Crédito": "{{rubrica_encargo}}",
    "Enc. Lim. Credito": "{{rubrica_encargo}}",
    "Enc Lim Crédito":   "{{rubrica_encargo}}",
    "Enc Lim Credito":   "{{rubrica_encargo}}",

    # ==== Renomeação dos placeholders ANTIGOS de inicial-mora.docx ====
    "{{rubrica_curta_caps}}": "{{rubrica_mora_caps}}",
    "{{rubrica_curta}}":      "{{rubrica_mora}}",
}


def processar_template(caminho):
    print(f"\n=== {caminho.name} ===")
    backup = caminho.with_suffix(caminho.suffix + ".bak_pre_rubrica_var")
    if not backup.exists():
        shutil.copy2(caminho, backup)
        print(f"  Backup: {backup.name}")
    else:
        print(f"  (backup já existe — não sobrescrito)")

    doc = Document(caminho)
    n_par_modificados = 0

    # Body
    for p in doc.paragraphs:
        if substituir_in_run(p._p, MAPA_RUBRICAS):
            n_par_modificados += 1

    # Tabelas
    n_cell_mod = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if substituir_in_run(p._p, MAPA_RUBRICAS):
                        n_cell_mod += 1

    doc.save(caminho)
    print(f"  Parágrafos modificados (body): {n_par_modificados}")
    print(f"  Parágrafos modificados (tabelas): {n_cell_mod}")
    return n_par_modificados + n_cell_mod


total = 0
for nome in ["inicial-mora.docx", "inicial-mora-encargo.docx", "inicial-combinada.docx"]:
    caminho = VAULT / nome
    if not caminho.exists():
        print(f"\n=== {nome}: NÃO EXISTE ===")
        continue
    total += processar_template(caminho)

print(f"\nTotal de modificações: {total}")
