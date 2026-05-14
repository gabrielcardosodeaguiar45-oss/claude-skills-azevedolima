"""Inicial COMBINADA TARIFAS + MORA — EXEMPLO MANUEL CORDOVIL.

REAGRUPAMENTO conforme procurador (08/05/2026):
  Tese 1 (TARIFAS): CARTÃO CRÉDITO ANUIDADE + SERVIÇO CARTÃO PROTEGIDO (família TARIFAS)
  Tese 2 (MORA):    MORA CRÉDITO PESSOAL

Comarca Barreirinha/AM, IDOSO 66 anos (04/10/1959), Ag 3725 / Conta 2782-0,
INSS R$ 846,22.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _combinada_helper import gerar_combinada
from docx import Document

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\EXEMPLO MANUEL CORDOVIL - Wilson - TARIFA\TARIFA E MORA'

# Tese 1 (TARIFAS): CARTÃO ANUIDADE + SERVIÇO CARTÃO PROTEGIDO
CARTAO_E_SERVICO = sorted([
    # CARTÃO CRÉDITO ANUIDADE (15)
    ('05/05/2022', 2.54), ('02/06/2022', 16.71), ('06/06/2022', 19.25),
    ('05/07/2022', 4.32), ('02/08/2022', 14.93), ('05/08/2022', 0.48),
    ('02/09/2022', 18.77), ('06/09/2022', 7.35), ('04/10/2022', 11.90),
    ('05/10/2022', 5.65), ('03/11/2022', 13.60), ('07/11/2022', 7.12),
    ('02/12/2022', 12.13), ('05/12/2022', 4.35), ('03/01/2023', 14.90),
    # SERVIÇO CARTÃO PROTEGIDO (4)
    ('05/02/2021', 0.56), ('02/03/2021', 9.43), ('05/03/2021', 8.48),
    ('05/04/2021', 2.67),
], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

# Tese 2 (MORA): MORA CRÉDITO PESSOAL (14)
MORA_CRED = [
    ('04/07/2022', 21.09), ('04/10/2022', 297.86),
    ('03/01/2023', 296.47), ('04/04/2023', 297.88), ('03/05/2023', 296.48),
    ('02/08/2023', 108.41), ('04/09/2023', 118.97), ('03/01/2025', 152.04),
    ('04/02/2025', 145.65), ('07/03/2025', 153.34), ('02/04/2025', 151.88),
    ('05/05/2025', 153.31), ('03/06/2025', 151.58), ('02/07/2025', 150.12),
]

teses = [
    {'familia': 'TARIFAS',
     'rubrica': 'CARTÃO CRÉDITO ANUIDADE / SERVIÇO CARTÃO PROTEGIDO',
     'lancamentos': CARTAO_E_SERVICO},
    {'familia': 'MORA',
     'rubrica': 'MORA CRÉDITO PESSOAL',
     'lancamentos': MORA_CRED},
]

autora = {
    'nome': 'EXEMPLO MANUEL CORDOVIL', 'nacionalidade': 'brasileiro',
    'estado_civil': '', 'profissao': 'aposentado',
    'cpf': '000.000.018-28', 'rg': '1000016-6', 'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Pimentel Tavares', 'numero': '341',
    'bairro': 'CM Terra P do Limão', 'cidade': 'Barreirinha', 'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '2782-0'}
renda = {'valor_float': 846.22, 'descricao': 'aposentadoria pelo INSS'}

res = gerar_combinada(
    pasta=PASTA, nome_arquivo_base='INICIAL_Combinada_MANUEL_LAZARO',
    autora=autora, conta=conta, renda=renda, teses=teses,
    comarca='Barreirinha', uf='AM', eh_idoso=True,
)
t = res['totais']
print(f'EXEMPLO MANUEL — combinada {t["n_teses"]} teses (TARIFAS+MORA)')
print(f'  Total: R$ {t["total_geral"]:.2f} / dobro R$ {t["dobro_geral"]:.2f}')
print(f'  Dano moral: R$ {t["dano_moral_total"]:.2f} ({t["n_teses"]} x R$ {t["dano_moral_por_tese"]:.2f})')
print(f'  VC: R$ {t["valor_causa"]:.2f}')
print(f'  Parágrafos removidos: {res["paras_removidos"]}, pós-fix: {res["pos_fix"]}')
print(f'  -> {res["docx"]}')

# ===== Relatório paralelo =====
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Combinada_MANUEL_LAZARO', level=1)
for k, v in [('Cliente', autora['nome']),
             ('Tese', 'COMBINADA TARIFAS (CARTÃO+SERVIÇO) + MORA'),
             ('Comarca', 'Barreirinha/AM'),
             ('Arquivo', 'INICIAL_Combinada_MANUEL_LAZARO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
linhas = [
    ('Comarca', 'Barreirinha/AM'),
    ('Prioridade', 'IDOSO 66 anos (04/10/1959, RG OCR confirmado)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]} — naturalidade Barreirinha/AM'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', 'R$ 846,22 (INSS último crédito 02/07/2025)'),
]
for i, te in enumerate(teses, 1):
    linhas.append((f'Tese {i} ({te["familia"]}) — {te["rubrica"]}',
                   f'{te["n_lanc"]} lanç. {te["inicio"]}–{te["fim"]} = R$ {te["total"]:.2f} / R$ {te["dobro"]:.2f} dobro'))
linhas.extend([
    ('TOTAL combinado', f'R$ {t["total_geral"]:.2f} / R$ {t["dobro_geral"]:.2f} dobro'),
    ('Dano moral por tese', f'R$ {t["dano_moral_por_tese"]:.2f}'),
    ('Dano moral total', f'R$ {t["dano_moral_total"]:.2f} ({t["n_teses"]} teses)'),
    ('Valor da causa', f'R$ {t["valor_causa"]:.2f}'),
])
for k, v in linhas:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('Inicial COMBINADA — TARIFAS + MORA',
     '2 procurações distintas. CARTÃO ANUIDADE + SERVIÇO CARTÃO PROTEGIDO formam o '
     'núcleo TARIFAS (produtos de cartão); MORA CRED PESS forma o núcleo MORA. '
     '2 núcleos fáticos individualizados, 2 blocos doutrinários, 2 cabeçalhos de '
     'pedido com 4 sub-itens cada (declaratório / condenatório / subsidiário / '
     'dano moral).'),
    ('Dano moral pleiteado individualmente por tese',
     'Conforme modelo paradigma "PI Tarifas - Mora - Titulo.docx" da pasta, cada '
     'pedido pleiteia R$ 15.000 de dano moral por tese. Total cumulativo R$ 30.000 '
     'embora possa o juízo arbitrar valor único.'),
    ('PRESCRIÇÃO — corte 30/03/2021',
     '4 lançamentos SERVIÇO CARTÃO PROTEGIDO em fev-abr/2021 podem estar '
     'parcialmente prescritos (R$ 21,14). Restante (CARTÃO ANUIDADE 2022-2023 e '
     'MORA CRED 2022-2025) está válido. REVISAR.'),
    ('IDOSO confirmado 66 anos (RG OCR)',
     'RG 0590363-7 SSP/AM, naturalidade Barreirinha/AM.'),
    ('Estado civil — não informado', 'Notificação não traz. Placeholder OMITIDO.'),
    ('Renda < 1 SM — possível consignação',
     'INSS R$ 846,22 (02/07/2025) < salário mínimo. Conferir HISCON.'),
    ('VALORES MORA ELEVADOS — cheque especial',
     'Lançamentos chegam a R$ 296-298 em 2022/2023.'),
    ('CARTÃO ANUIDADE com lançamentos repetidos no mesmo mês',
     '2 lançamentos em jun/2022, ago/2022, set/2022, out/2022, nov/2022, dez/2022. '
     'Pode ser 2 cartões ou erro. Conferir extrato.'),
    ('CLIENTE TEM 1 PG ELETRON SEPARADA',
     'PG ELETRON SUDACRED em pasta separada. Mantida SEPARADA por solidariedade do '
     'terceiro.'),
    ('TETO JEC',
     f'VC R$ {t["valor_causa"]:.2f} ≈ {t["valor_causa"]/1518:.1f} SM. {"Cabe" if t["valor_causa"] < 60720 else "EXCEDE"} no JEC (40 SM = R$ 60.720).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'CONFERIR visualmente combinada (2 núcleos: TARIFAS+MORA; sem TÍTULO/APLIC).',
    'Conferir 2 núcleos fáticos com rótulo em negrito.',
    'Conferir 2 cabeçalhos de pedido + 4 sub-itens cada (a, b, c, d numeração automática).',
    'Conferir nome / CPF / RG / nascimento.',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Barreirinha/AM.',
    f'Conferir VC R$ {t["valor_causa"]:.2f} e dano moral total R$ {t["dano_moral_total"]:.2f}.',
    'Confirmar com cliente: nunca contratou cartão, seguro de cartão nem aceitou cheque especial.',
    'Anexar 2 procurações + 3-RG + 4-Hipossuficiência + 5-Comprovante + 5.1-Declaração + 5.2-RG proprietária + 6-Extrato + 7-Tabelas + 8-Notificação + 8.1-Comprovante.',
    'PG ELETRON SUDACRED segue em ação SEPARADA.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após conferência visual, ')
r = p.add_run('PROTOCOLAR.'); r.bold = True
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MANUEL_LAZARO_COMBINADA_v1.docx')
doc_r.save(RELAT_OUT)
print(f'  -> {RELAT_OUT}')
