"""Inicial MORA — CLIENTE EXEMPLO RODRIGUES DA SILVA.

Comarca Maués/AM. Conta Bradesco Ag 3706/16649-9. INSS R$ 980,09 (último
crédito 03/02/2026). Dados de identificação reaproveitados do batch TARIFAS
(mesmo extrato escaneado de 38 páginas, processado via easyocr).

Tese: MORA CREDITO PESSOAL — 37 lançamentos OCR-extraídos entre 03/01/2020
e 31/12/2025 (total R$ 5.330,75; dobro R$ 10.661,50). Dano moral R$ 15.000.
VC R$ 25.661,50 — cabe folgadamente no JEC.

PENDÊNCIAS:
1. PRESCRIÇÃO PARCIAL CDC — 18 lançamentos pré-06/05/2021 (R$ 2.888,10)
   estão prescritos pelo art. 27 CDC; 19 lançamentos pós-06/05/2021
   (R$ 2.442,65) estão vivos. Pelo art. 205 CC (10 anos) TODOS estão vivos.
   Inicial pleiteia O TOTAL apostando na corrente DECENAL acolhida em
   algumas câmaras do TJ-AM. Procurador deve estar preparado para defender.
2. OCR DE EXTRATO ESCANEADO — números podem ter pequenas variações de
   parsing. CONFERIR contra extrato original.
3. Cliente também tem TARIFA BANCÁRIA CESTA B.EXPRESSO no extrato (32
   lançamentos = R$ 558,49 já processados em batch separado).
"""
import io, sys, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\1. TESTE 2\CLIENTE EXEMPLO RODRIGUES DA SILVA - Maurivã - TARIFAS\MORA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Mora_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_MORA_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_mora_lancs.json')) as f:
    data = json.load(f)
LANCAMENTOS = sorted(data['todos'], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO RODRIGUES DA SILVA', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.003-13', 'rg': '1000002-2',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua 10', 'numero': '818',
    'bairro': 'Nova Esperança', 'cidade': 'Maués', 'cep': '69.190-000',
}
conta = {'agencia': '3706', 'numero': '16649-9'}
renda = {'valor_float': 980.09}

tese = {'rubrica': 'MORA CREDITO PESSOAL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Maués', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['rubrica_curta'] = 'Mora Cred Pess'
dados['rubrica_curta_caps'] = 'MORA CRED PESS'
dados['rubrica_completa'] = 'Mora Crédito Pessoal'
dados['rubrica_completa_caps'] = 'MORA CRÉDITO PESSOAL'
dados['titulo'] = 'MORA CRÉDITO PESSOAL'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

vivos = data['vivos']
prescritos = data['prescritos']

print(f'=== CLIENTE EXEMPLO — MORA (OCR) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | dobro: R$ {calc["dobro"]:,.2f} | VC: R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'  vivos CDC: {len(vivos)} = R$ {sum(v for _,v in vivos):.2f} | prescritos CDC: {len(prescritos)} = R$ {sum(v for _,v in prescritos):.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
       '{{cidade_filial}}': 'Maués',
       '{{uf_filial}}': 'AM',
       '{{uf_extenso}}': 'Amazonas'}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

# Relatório
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Mora_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CREDITO PESSOAL'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Mora_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]} (extraído via OCR)'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 03/02/2026 via OCR)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos MORA', f'{len(LANCAMENTOS)} (OCR — CONFERIR)'),
    ('  vivos CDC (≥ 06/05/2021)', f'{len(vivos)} = R$ {sum(v for _,v in vivos):.2f}'.replace('.', ',')),
    ('  prescritos CDC (< 06/05/2021)', f'{len(prescritos)} = R$ {sum(v for _,v in prescritos):.2f}'.replace('.', ',')),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO PARCIAL CDC — APOSTA NA TESE DECENAL',
     'Dos 37 lançamentos MORA, 18 (R$ 2.888,10, datas entre 03/01/2020 e 06/04/2021) '
     'estão prescritos pelo art. 27 CDC (5 anos). 19 (R$ 2.442,65, datas entre '
     '06/05/2021 e 31/12/2025) estão vivos. Inicial pleiteia O TOTAL R$ 5.330,75 '
     'apostando na corrente DECENAL (art. 205 CC) acolhida em algumas câmaras do '
     'TJ-AM. Procurador deve estar preparado para defender em recurso. Padrão '
     'validado no caso EXEMPLA RAIMUNDA (06/05/2026).'),
    ('OCR (easyocr) usado em extrato escaneado de 38 páginas',
     'O extrato Bradesco é PDF escaneado (CamScanner). Processado via easyocr (PT). '
     'Pequenas variações de parsing são possíveis. CONFERIR cruzando lista de 37 '
     'lançamentos com extrato original. Easyocr pode confundir caracteres ("0" vs '
     '"O", "3" vs "8") e pegar saldo em vez de débito em raros casos.'),
    ('IDADE — não confirmada',
     'RG não lido nesta sessão. Assume NÃO IDOSO. Conferir antes do protocolo.'),
    ('Cliente também tem TARIFA BANCÁRIA',
     'O extrato registra tese paralela TARIFA BANCÁRIA CESTA B.EXPRESSO (32 lanç. '
     '= R$ 558,49) já processada em batch separado. Avaliar se vale ação COMBINADA '
     '(MORA + TARIFA, dano moral R$ 5.000 × 2 = R$ 10.000) em vez de duas separadas.'),
    ('TETO JEC — coberto',
     f'VC R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.') + ' ≈ 16,89 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (16649-9 / 3706).',
           'Conferir comarca: Maués/AM.', 'CONFERIR OCR contra extrato original (37 MORA).',
           'Confirmar idade no RG.',
           'Decidir estratégia: 2 iniciais separadas (MORA + TARIFAS) ou 1 COMBINADA.',
           'Anexar 8-NOTIFICAÇÃO MORA + 8.1-COMPROVANTE NOTIFICAÇÃO.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA COM RESSALVA OCR + DECENAL — '); r2 = p.add_run('PROTOCOLAR após conferir lançamentos contra extrato original.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
