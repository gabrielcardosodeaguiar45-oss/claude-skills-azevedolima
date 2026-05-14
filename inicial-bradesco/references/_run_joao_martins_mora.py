"""Inicial MORA + ENCARGOS — EXEMPLO MARTINS DA SILVA.

Mesma autora da inicial TARIFAS (Manacapuru/AM, IDOSO, INSS R$ 1.621,00,
Ag 3707 Conta 8698-3, procuração a rogo). Pasta separada.

Tabela: 120 lançamentos! ENCARGOS LIMITE DE CRED + MORA CREDITO PESSOAL
entre 26/01/2017 e 03/11/2025. Total R$ 6.397,91 / dobro R$ 12.795,82.
VC R$ 27.795,82 — cabe no JEC (40 SM = R$ 60.720).

PRESCRIÇÃO: muitos lançamentos antes de 30/03/2021 (EAREsp 1.280.825).
Procurador decide se aplica corte ou pleiteia tudo.

Template `inicial-mora-encargo.docx`.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal, extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\2. MORA\EXEMPLO MARTINS DA SILVA - Ney Pedroza - TARIFAS\MORA E ENCARGOS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora-encargo.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_MoraEncargo_JOAO_MARTINS_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_JOAO_MARTINS_MORA_v1.docx')

ENCARGO = [
    ('26/01/2017', 8.27), ('31/01/2017', 35.80), ('02/02/2017', 40.10),
    ('31/03/2017', 24.41), ('03/04/2017', 40.48), ('02/05/2017', 6.63),
    ('01/06/2017', 1.55), ('03/07/2017', 13.95), ('31/08/2017', 0.79),
    ('01/09/2017', 39.99), ('02/10/2017', 11.16), ('01/11/2017', 10.24),
    ('01/02/2018', 16.43), ('29/03/2018', 7.96), ('02/04/2018', 37.16),
    ('02/05/2018', 27.78), ('01/06/2018', 4.83), ('02/07/2018', 1.22),
    ('01/08/2018', 33.51), ('03/09/2018', 2.19), ('01/10/2018', 24.30),
    ('01/11/2018', 39.23), ('03/12/2018', 1.81), ('04/01/2019', 12.38),
    ('01/02/2019', 26.87), ('05/02/2019', 10.94), ('01/03/2019', 6.88),
    ('23/04/2019', 9.56), ('07/05/2019', 40.26), ('28/06/2019', 43.84),
    ('05/07/2019', 34.56), ('01/08/2019', 39.72), ('02/09/2019', 17.27),
    ('01/10/2019', 1.59), ('01/11/2019', 2.91), ('02/12/2019', 4.10),
    ('31/01/2020', 4.68), ('28/02/2020', 27.79), ('31/03/2020', 23.86),
    ('23/04/2020', 27.89), ('04/05/2020', 22.23), ('03/08/2020', 2.35),
    ('01/09/2020', 3.55), ('01/10/2020', 1.76), ('09/11/2020', 19.67),
    ('01/12/2020', 18.15), ('29/01/2021', 13.02), ('01/02/2021', 19.22),
    ('01/03/2021', 21.34), ('01/04/2021', 16.05), ('31/05/2021', 26.28),
    ('01/06/2021', 23.42), ('01/07/2021', 3.09), ('02/08/2021', 1.63),
    ('01/09/2021', 3.09), ('01/10/2021', 1.50), ('01/11/2021', 2.38),
    ('01/12/2021', 2.35), ('03/01/2022', 1.57), ('01/02/2022', 1.42),
    ('02/03/2022', 1.07), ('01/04/2022', 1.75), ('16/05/2022', 1.51),
    ('01/06/2022', 11.13), ('01/07/2022', 1.80), ('01/08/2022', 1.95),
    ('01/09/2022', 0.70), ('31/10/2022', 1.94), ('01/11/2022', 21.61),
    ('01/12/2022', 23.16), ('01/02/2023', 22.41), ('01/03/2023', 18.96),
    ('28/04/2023', 25.51), ('02/05/2023', 19.86), ('01/06/2023', 0.02),
    ('01/08/2023', 9.84), ('29/09/2023', 25.72), ('02/10/2023', 22.33),
    ('01/11/2023', 22.03), ('01/12/2023', 0.14), ('02/01/2024', 10.45),
    ('01/02/2024', 19.86), ('03/02/2025', 0.63), ('01/04/2025', 14.91),
    ('02/05/2025', 12.04), ('02/06/2025', 0.07), ('03/11/2025', 0.04),
]
MORA = [
    ('26/01/2017', 191.73), ('31/01/2017', 194.30), ('31/03/2017', 194.30),
    ('31/03/2017', 213.80), ('31/10/2018', 35.35), ('31/10/2018', 107.11),
    ('31/10/2018', 177.83), ('04/01/2019', 112.84), ('31/01/2019', 189.12),
    ('31/01/2019', 244.84), ('31/05/2019', 244.84), ('31/07/2019', 244.84),
    ('31/10/2019', 54.82), ('31/10/2019', 244.84), ('31/01/2020', 65.12),
    ('31/01/2020', 69.12), ('31/01/2020', 74.50), ('31/01/2020', 121.07),
    ('31/01/2020', 123.74), ('31/01/2020', 172.62), ('31/01/2020', 244.83),
    ('28/02/2020', 123.15), ('31/03/2020', 12.74), ('31/03/2020', 68.40),
    ('31/03/2020', 78.24), ('31/03/2020', 123.75), ('31/03/2020', 127.16),
    ('31/03/2020', 172.62), ('31/03/2020', 181.29), ('31/03/2020', 255.73),
    ('23/04/2020', 122.15), ('23/04/2020', 239.99), ('31/07/2020', 310.73),
]

LANCAMENTOS = sorted(ENCARGO + MORA, key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'EXEMPLO MARTINS DA SILVA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'aposentado',
    'cpf': '000.000.013-23',
    'rg': '1000011-1',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua do Campinho',
    'numero': '1333',
    'bairro': 'Novo Manaca',
    'cidade': 'Manacapuru',
    'cep': '69.409-899',
}
conta = {'agencia': '3707', 'numero': '8698-3'}
renda = {'valor_float': 1621.00}

tese = {'rubrica': 'MORA CRED PESS / ENC LIM CRÉDITO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Manacapuru', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['rubrica_curta'] = 'Mora Cred Pess / Enc Lim Crédito'
dados['rubrica_curta_caps'] = 'MORA CRED PESS / ENC LIM CRÉDITO'
dados['rubrica_completa'] = 'Crédito Mora Pessoal / Encargos Limite de Crédito'
dados['rubrica_completa_caps'] = 'MORA CREDITO PESSOAL / ENCARGOS LIMITE DE CRÉDITO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== EXEMPLO MARTINS — MORA + ENCARGOS ===')
print(f'Lançamentos: {len(LANCAMENTOS)} ({len(MORA)} mora + {len(ENCARGO)} encargo), total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
def fmt_moeda(v):
    s = f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')
    return f'R$ {s}'
te=sum(v for _,v in ENCARGO); tm=sum(v for _,v in MORA)

with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {
    '{{remuneração}}': dados['remuneração'],
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
    '{{cidade_filial}}': 'Maués',
    '{{uf_filial}}': 'AM',
    '{{uf_extenso}}': 'Amazonas',
    '{{numero_desconto_encargo}}': str(len(ENCARGO)),
    '{{numero_desconto_mora}}': str(len(MORA)),
    '{{total_encargo}}': fmt_moeda(te),
    '{{total_mora}}': fmt_moeda(tm),
    '{{total_encargo_extenso}}': extenso_moeda(te),
    '{{total_mora_extenso}}': extenso_moeda(tm),
    '{{desconto_extenso_encargo}}': extenso_cardinal(len(ENCARGO)),
    '{{desconto_extenso_mora}}': extenso_cardinal(len(MORA)),
}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_MoraEncargo_JOAO_MARTINS', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CRED PESS + ENC LIM CRÉDITO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_MoraEncargo_JOAO_MARTINS_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 30/01/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} ({len(ENCARGO)} encargo + {len(MORA)} mora)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO CRÍTICA — corte 30/03/2021 (EAREsp 1.280.825 STJ)',
     'TODOS os 33 lançamentos MORA são de 2017-2020 (estão prescritos pelo art. 27 '
     'CDC). Dos ENCARGOS, ~50 são pré-30/03/2021 (também prescritos). Pós-30/03/2021: '
     '~37 lançamentos válidos (~R$ 380 simples / R$ 760 dobro). DECISÃO DO '
     'PROCURADOR: (a) pleitear TUDO testando art. 205 CC (10 anos — minoritária '
     'TJAM); (b) corte 5 anos reduzindo VC para ~R$ 15.760. Inicial gerada com '
     'TODOS — REVISAR.'),
    ('IDOSO — prioridade aplicada',
     'Notificação afirma "pessoa idosa" e "aposentado".'),
    ('PROCURAÇÃO ASSINADA A ROGO',
     'Pasta tem RG da rogada (Marilene Pereira da Silva) + 2 testemunhas (Rosiane e '
     'Deilson). Conferir validade.'),
    ('CLIENTE TEM 3 TESES SEPARADAS',
     'EXEMPLO MARTINS aparece em TARIFAS (já gerada), MORA (esta) e TÍTULO CAPITALIZAÇÃO. '
     'AVALIAR consolidação em inicial-combinada (Manacapuru adota combinação por '
     'padrão).'),
    ('VALORES MORA ELEVADOS',
     'Lançamentos MORA chegam a R$ 244-310. Indica cheque especial alto. Conferir '
     'extrato.'),
    ('TETO JEC — coberto',
     'VC R$ 27.795,82 ≈ 18,3 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'DECIDIR estratégia de prescrição (corte 30/03/2021 ou pleito decenal).',
    'AVALIAR consolidação com TARIFAS + TÍTULO em 1 só inicial-combinada.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade da procuração assinada a rogo.',
    'Conferir comarca: Manacapuru/AM.',
    'Confirmar com cliente: nunca contratou cheque especial.',
    'Anexar 2-Procurações (ENC LIM + MORA CRED PESS) + 3-RG + 3.1 a 3.3 (rogada + testemunhas) + 4-Hipossuficiência + 5-Comprovante + 5.1-Declaração + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir prescrição e consolidação. Após resolução, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
