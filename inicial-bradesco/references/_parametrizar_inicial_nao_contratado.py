"""Parametriza o template do Roque (1. PETIÇÃO INICIAL Gabriel.docx) substituindo
dados específicos por placeholders {{...}}, gerando inicial-jfba-base.docx no vault.

Uso: rodar 1 vez. Depois NÃO rodar mais — o template vai ser editado direto no vault.
"""
import io, sys, os, shutil, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from helpers_docx import substituir_in_run

def nfd(s):
    """Decompõe caracteres acentuados (ã → a + tilde) para casar com texto Word."""
    return unicodedata.normalize('NFD', s)

TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jfba-base.docx'

# === MAPA DE SUBSTITUIÇÕES ===
# (índice_paragrafo, dict {antigo: novo})
# Estratégia: trocar trechos curtos e literais por placeholders preservando formatação dos runs
SUBSTITUICOES = {
    # p000 — cabeçalho juízo
    0: {'Salvador/BA': '{{cidade_subsecao}}/{{uf_subsecao}}'},

    # p007 — qualificação completa do autor
    # Endereço é particionado para evitar problemas com caracteres especiais (município com ı + ́)
    7: {
        'ROQUE CUSTODIO DA SILVA': '{{nome_autor}}',
        'brasileiro, solteiro, aposentado': '{{nacionalidade}}, {{estado_civil}}, {{profissao}}',
        '000.000.001-11': '{{cpf_autor}}',
        '01.095.230-67': '{{rg_autor}}',
        'Rua da Buganvilla, s/n, bairro Jardim Imbassai':
            '{{logradouro_autor}}, {{numero_autor}}, bairro {{bairro_autor}}',
        'Mata de São João/BA': '{{cidade_autor}}/{{uf_autor}}',
        'CEP 48289-000': 'CEP {{cep_autor}}',
    },

    # p011 — polo passivo (banco + INSS)
    11: {
        'BANCO DAYCOVAL S/A, inscrito no CNPJ/MF sob o nº 62.232.889/0001 90, na Avenida Paulista 1793, Bela Vista, São Paulo /SC, Brasil, CEP: 01311-200':
            '{{banco_reu_nome}}, {{banco_reu_descricao_pj}}, inscrito no CNPJ/MF sob o nº {{banco_reu_cnpj}}, com sede na {{banco_reu_endereco}}',
        'Av. Sete de Setembro, 1078 - Mercês, Salvador/BA':
            '{{inss_endereco_subsecao}}',
    },

    # p013 — síntese fática do benefício
    13: {
        'aposentadoria por invalidez previdenciária': '{{tipo_beneficio}}',
        '024.044.285-7': '{{nb_beneficio}}',
        'agência 3706, conta corrente nº 0000211974':
            'agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}',
        'Banco Cooperativo Sicoob SA': '{{banco_pagador}}',
    },

    # p017 — descrição do contrato (BLOCO REPETÍVEL — para 1 contrato funciona direto;
    # para N contratos a skill duplica este parágrafo antes de substituir)
    17: {
        'na competência 05/2017': 'na competência {{contrato_competencia_inicio}}',
        'de um total de 72 parcelas': 'de um total de {{contrato_qtd_parcelas}} parcelas',
        'no valor de R$ 58,98': 'no valor de R$ {{contrato_valor_parcela}}',
        'no valor de R$ 4.246,56': 'no valor de R$ {{contrato_valor_emprestado}}',
        'tendo as parcelas sido encerradas em 02/2023': 'tendo as parcelas sido encerradas em {{contrato_competencia_fim}}',
        'contrato n° 237344 469': 'contrato n° {{contrato_numero}}',
        'banco Daycoval S.A': 'banco {{contrato_banco}}',
    },

    # p044 — justiça gratuita (renda)
    44: {'R$ 974,69': 'R$ {{valor_renda_liquida}}'},

    # p191 — quantum dano moral (parágrafo expositivo)
    191: {'R$ 15.000,00 (quinze mil reais)': 'R$ {{dano_moral_total}} ({{dano_moral_total_extenso}})'},

    # p229 — pedido DECLARAR (BLOCO REPETÍVEL)
    229: {
        'no valor de R$ 4.246,56': 'no valor de R$ {{contrato_valor_emprestado}}',
        'contrato nº- 237344 469': 'contrato nº {{contrato_numero}}',
        'descontos de R$ 58,98 mensais': 'descontos de R$ {{contrato_valor_parcela}} mensais',
        'com inclusão em 03/05/2017': 'com inclusão em {{contrato_data_inclusao}}',
        'início de desconto em 05/2017 a 02/2023': 'início de desconto em {{contrato_competencia_inicio}} a {{contrato_competencia_fim}}',
        'benefício previdenciário 024.044.285-7': 'benefício previdenciário {{nb_beneficio}}',
    },

    # p230 — pedido CONDENAR dano moral
    230: {'R$ 15.000,00 (quinze mil reais)': 'R$ {{dano_moral_total}} ({{dano_moral_total_extenso}})'},

    # p246 — valor da causa
    246: {
        'R$ 40.212,65 (quarenta mil, duzentos e doze reais e sessenta e cinco centavos)':
            'R$ {{valor_causa}} ({{valor_causa_extenso}})',
    },

    # p248 — cidade + data
    248: {'Salvador/BA, 7 de maio de 2026': '{{cidade_protocolo}}/{{uf_protocolo}}, {{data_protocolo}}'},
}

# === EXECUÇÃO ===
print(f'Carregando: {TEMPLATE}')
doc = Document(TEMPLATE)
print(f'Total parágrafos: {len(doc.paragraphs)}')

falhas = []
sucessos = 0
for idx, mapa in SUBSTITUICOES.items():
    if idx >= len(doc.paragraphs):
        falhas.append(f'p{idx} fora do range')
        continue
    p = doc.paragraphs[idx]
    texto_antes = p.text
    for antigo, novo in mapa.items():
        # tenta primeiro a forma direta, depois NFD (decomposta)
        chave = antigo if antigo in texto_antes else nfd(antigo)
        if chave not in texto_antes:
            falhas.append(f'p{idx}: NÃO ENCONTRADO "{antigo[:80]}..."')
            continue
        ok = substituir_in_run(p._element, {chave: novo})
        if ok:
            sucessos += 1
        else:
            falhas.append(f'p{idx}: substituir_in_run falhou para "{antigo[:60]}"')

print(f'\\n--- RESUMO ---')
print(f'Substituições bem-sucedidas: {sucessos}')
print(f'Falhas: {len(falhas)}')
for f in falhas:
    print(f'  ✗ {f}')

# Salvar (sobrescreve o copy original)
doc.save(TEMPLATE)
print(f'\\nSalvo em: {TEMPLATE}')

# Verificar placeholders presentes
print(f'\\n--- PLACEHOLDERS DETECTADOS ---')
import re
doc2 = Document(TEMPLATE)
texto_full = '\\n'.join(p.text for p in doc2.paragraphs)
placeholders = sorted(set(re.findall(r'\{\{[^}]+\}\}', texto_full)))
print(f'Total únicos: {len(placeholders)}')
for ph in placeholders:
    cnt = texto_full.count(ph)
    print(f'  {ph} (×{cnt})')
