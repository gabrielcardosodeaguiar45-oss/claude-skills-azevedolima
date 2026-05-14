# -*- coding: utf-8 -*-
"""Substitui:
1. 'aposentada' (descrição da parte autora) por {{condicao_socioeconomica}}
   em 5 ocorrências (par 125 mora; 131 mora-encargo; 151 tarifas; 155 pg-eletron;
   219 combinada).
2. 'Autora' (par 155 pg-eletron) por {{parte_autora}}.

NÃO altera 'aposentado/pensionista' em texto argumentativo genérico
(par 145 tarifas, 213 combinada) nem em jurisprudência."""
import sys, io, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from pathlib import Path
sys.path.insert(0, str(Path(__file__).parent))
from docx import Document
from helpers_docx import substituir_in_run

VAULT = Path(r"C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates")

# ---- Substituições GLOBAIS por arquivo (ancoradas em contexto pra não pegar
# jurisprudência) ----

# Padrão 1 — "Imperioso considerar..." aparece em 5 templates, sempre com
# "pessoa idosa, aposentada," — substituímos no escopo desse parágrafo.
PADRAO1_SRC = "pessoa idosa, aposentada, de poucos recursos"
PADRAO1_DST = "pessoa idosa, {{condicao_socioeconomica}}, de poucos recursos"

# Padrão 2 — "A condição pessoal da Autora, pessoa idosa, aposentada, de parcos recursos"
PADRAO2_SRC = "A condição pessoal da Autora, pessoa idosa, aposentada, de parcos recursos"
PADRAO2_DST = "A condição pessoal {{parte_autora_artigo}} {{parte_autora}}, pessoa idosa, {{condicao_socioeconomica}}, de parcos recursos"

# Apenas onde Padrão 2 estiver presente (pg-eletron par 155).

ALVOS = [
    "inicial-mora.docx",
    "inicial-mora-encargo.docx",
    "inicial-tarifas.docx",
    "inicial-aplic-invest.docx",
    "inicial-pg-eletron.docx",
    "inicial-combinada.docx",
]

for nome in ALVOS:
    caminho = VAULT / nome
    if not caminho.exists():
        continue

    # backup só se ainda não houver
    backup = caminho.with_suffix(caminho.suffix + ".bak_pre_condicao")
    if not backup.exists():
        shutil.copy2(caminho, backup)

    doc = Document(caminho)
    n = 0
    for p in doc.paragraphs:
        # Aplicar primeiro o padrão 2 (mais específico) e depois o padrão 1
        if substituir_in_run(p._p, {PADRAO2_SRC: PADRAO2_DST}):
            n += 1
        if substituir_in_run(p._p, {PADRAO1_SRC: PADRAO1_DST}):
            n += 1
    doc.save(caminho)
    print(f"{nome}: {n} substituição(ões)")

# Verificação final
print("\n=== Verificação ===")
for nome in ALVOS:
    d = Document(VAULT / nome)
    for i, p in enumerate(d.paragraphs):
        t = p.text
        if ("pessoa idosa, aposentada" in t or
            "condição pessoal da Autora" in t):
            print(f"  ⚠️ AINDA NÃO SUBSTITUÍDO: [{nome}:{i}] {t[:200]}")
        if "{{condicao_socioeconomica}}" in t or "{{parte_autora}}" in t:
            print(f"  ✓ [{nome}:{i}] {t[:160]}...")
