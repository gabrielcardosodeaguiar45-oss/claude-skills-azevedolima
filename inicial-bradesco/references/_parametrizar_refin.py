"""Parametriza inicial-jfba-refin-ativo.docx (template para REFINANCIAMENTO ATIVO).

Particularidades vs base:
- Vocabulário: "refinanciamento" no lugar de "empréstimo" em p15, p17, p244
- p18-p23: parágrafos extras FIXOS sobre HISCON não registrar valor liberado
  em refins → impossibilidade de saber o "troco"
- p98 (CDC): tem trecho específico "embora o valor tenha sido creditado à parte
  autora" — mantém literal (reconhece que houve crédito de troco)
- p245: pleito inclui DANO TEMPORAL adicional (R$ 5.000) além do dano moral
"""
import io, sys, os, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from helpers_docx import substituir_in_run

def nfd(s):
    return unicodedata.normalize('NFD', s)

TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates\inicial-jfba-refin-ativo.docx'

SUBSTITUICOES = {
    # p0 — juízo (atenção: "Subseção de  /BA" — tem 2 espaços antes da /)
    0: {'Subseção de  /BA': 'Subseção de {{cidade_subsecao}}/{{uf_subsecao}}'},

    # p7 — qualificação autor (com placeholders manuais "FULANO DE TAL", "xxxxxx" etc.)
    7: {
        'FULANO DE TAL': '{{nome_autor}}',
        'brasileiro, casado, aposentado': '{{nacionalidade}}, {{estado_civil}}, {{profissao}}',
        'xxx.xxx.xxx-xx': '{{cpf_autor}}',
        'Cédula de Identidade sob nº xxxxxxxx, órgão expedidor SSP/SC':
            'Cédula de Identidade sob nº {{rg_autor}}, órgão expedidor {{orgao_expedidor_autor}}',
        'rua tal, n° xx, bairro tal':
            '{{logradouro_autor}}, n° {{numero_autor}}, bairro {{bairro_autor}}',
        'em cidade/BA': 'em {{cidade_autor}}/{{uf_autor}}',
        'CEP xxxxx-xxx': 'CEP {{cep_autor}}',
    },

    # p11 — polo passivo (BRADESCO específico do exemplo → genérico)
    11: {
        'BANCO BRADESCO S/A, pessoa jurídica de direito privado, inscrita no CNPJ sob o nº 60.746.948/0320-73, com endereço na Avenida Sete de Setembro, nº 895, Centro, Salvador/BA, CEP 69.005-140':
            '{{banco_reu_nome}}, {{banco_reu_descricao_pj}}, inscrita no CNPJ sob o nº {{banco_reu_cnpj}}, com endereço na {{banco_reu_endereco}}',
        'Av. Sete de Setembro, 1078 - Mercês, Salvador/BA':
            '{{inss_endereco_subsecao}}',
    },

    # p13 — síntese benefício
    13: {
        'aposentadoria por idade': '{{tipo_beneficio}}',
        '149.139.433-9': '{{nb_beneficio}}',
        'agência 3706, conta corrente nº 0000211974':
            'agência {{agencia_pagador}}, conta corrente nº {{conta_pagador}}',
        'BANCO BRADESCO SA': '{{banco_pagador}}',
    },

    # p15 — frase intro (banco + número do refinanciamento)
    15: {
        'BANCO PAN S/A, CONTRATO Nº 3880089838':
            '{{banco_reu_nome}}, CONTRATO Nº {{contrato_numero}}',
    },

    # p17 — descrição do refinanciamento (parametriza os xxxxxxxx)
    17: {
        'na competência xxxxxxxx': 'na competência {{contrato_competencia_inicio}}',
        'de um total de xx parcelas': 'de um total de {{contrato_qtd_parcelas}} parcelas',
        'no valor de R$ xxx,xx (valor por extenso), relativas':
            'no valor de R$ {{contrato_valor_parcela}} ({{contrato_valor_parcela_extenso}}), relativas',
        'no valor de R$ xxx,xx (valor por extenso), contrato':
            'no valor de R$ {{contrato_valor_emprestado}} ({{contrato_valor_emprestado_extenso}}), contrato',
        'contrato n° xxxxxxx': 'contrato n° {{contrato_numero}}',
        'pelo banco xxxxx': 'pelo banco {{contrato_banco}}',
    },

    # p47 — renda líquida (vazio)
    47: {'no valor líquido de R$ (), conforme': 'no valor líquido de R$ {{valor_renda_liquida}}, conforme'},

    # p244 — pedido DECLARAR (refin ativo: tem valor por extenso direto + texto "empréstimo/refinanciamento")
    244: {
        'no valor de R$ 1.394,78 (um mil, trezentos e noventa e quatro reais e setenta e oito centavos)':
            'no valor de R$ {{contrato_valor_emprestado}} ({{contrato_valor_emprestado_extenso}})',
        'contrato nº 0123506012709': 'contrato nº {{contrato_numero}}',
        'descontos de R$ 31,21 (trinta e um reais e vinte e um centavos) mensais':
            'descontos de R$ {{contrato_valor_parcela}} ({{contrato_valor_parcela_extenso}}) mensais',
        'com inclusão em 24/07/2024': 'com inclusão em {{contrato_data_inclusao}}',
        'início de desconto em 01/08/2024': 'início de desconto em {{contrato_competencia_inicio}}',
        'benefício previdenciário 149.139.433-9': 'benefício previdenciário {{nb_beneficio}}',
    },

    # p245 — pedido CONDENAR (dano moral + dano temporal — particularidade do REF ATV)
    245: {
        'R$ 15.000,00 (quinze mil reais) e de R$ 5.000,00 (cinco mil reais) à título de indenização por danos temporais':
            'R$ {{dano_moral_total}} ({{dano_moral_total_extenso}}) e de R$ {{dano_temporal_total}} ({{dano_temporal_total_extenso}}) à título de indenização por danos temporais',
    },

    # p261 — valor da causa (vazio: "R$ (centavos).")
    261: {'Dá-se a causa o valor de R$ (centavos).': 'Dá-se a causa o valor de R$ {{valor_causa}} ({{valor_causa_extenso}}).'},

    # p263 — cidade + data
    263: {'Salvador/BA, 7 de maio de 2026': '{{cidade_protocolo}}/{{uf_protocolo}}, {{data_protocolo}}'},
}

print(f'Carregando: {TEMPLATE}')
doc = Document(TEMPLATE)
print(f'Total parágrafos: {len(doc.paragraphs)}')

falhas = []
sucessos = 0
for idx, mapa in SUBSTITUICOES.items():
    if idx >= len(doc.paragraphs):
        falhas.append(f'p{idx} fora do range')
        continue
    p = doc.paragraphs[idx]
    texto_antes = p.text
    for antigo, novo in mapa.items():
        chave = antigo if antigo in texto_antes else nfd(antigo)
        if chave not in texto_antes:
            falhas.append(f'p{idx}: NÃO ENCONTRADO "{antigo[:60]}..."')
            continue
        ok = substituir_in_run(p._element, {chave: novo})
        if ok:
            sucessos += 1
        else:
            falhas.append(f'p{idx}: substituir_in_run falhou para "{antigo[:60]}"')

print(f'\n--- RESUMO ---')
print(f'Substituições bem-sucedidas: {sucessos}')
print(f'Falhas: {len(falhas)}')
for f in falhas:
    print(f'  ✗ {f}')

doc.save(TEMPLATE)
print(f'\nSalvo em: {TEMPLATE}')

import re
doc2 = Document(TEMPLATE)
texto_full = '\n'.join(p.text for p in doc2.paragraphs)
placeholders = sorted(set(re.findall(r'\{\{[^}]+\}\}', texto_full)))
print(f'\n--- PLACEHOLDERS DETECTADOS ({len(placeholders)} únicos) ---')
for ph in placeholders:
    cnt = texto_full.count(ph)
    print(f'  {ph} (×{cnt})')
