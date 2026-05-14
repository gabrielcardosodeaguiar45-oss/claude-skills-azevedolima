"""Inicial APLIC.INVEST FACIL — CLIENTE EXEMPLO SOUZA DE FREITAS VIANA
(NOVA tese, DIFERENTE da ODONTOPREV/PG ELETRON já processada em 06/05/2026).

Comarca Caapiranga/AM. NÃO idosa (48 anos, nascida 08/02/1978 — confirmado
checkpoint anterior). Solteira. Servidora pública municipal Caapiranga
(TRANSF SALDO C/SAL P/CC R$ 3.575,38 mensal — confirmado checkpoint).
Conta Bradesco Ag 3707/501049-7.

Tese: 159 lançamentos APLIC.INVEST FACIL entre 28/01/2021 e 31/07/2024,
total R$ 213.040,06 (parseados do extrato). Audit: 491 RESGATEs
totalizando R$ 213.512,88. Saldo NEGATIVO R$ 472,82 (recebeu de volta
tudo + rentabilidade). Estratégia (b) PADRÃO confirmada.

VC R$ 15.000. Cabe folgadamente no JEC.
"""
import io, sys, os, copy, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\CLIENTE EXEMPLO SOUZA DE FREITAS VIANA - Ney Pedroza'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_APLIC_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_lancs.json')) as f:
    LANCAMENTOS = sorted(json.load(f), key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO SOUZA DE FREITAS VIANA', 'nacionalidade': 'brasileira',
    'estado_civil': 'solteira', 'profissao': 'servidora pública municipal',
    'cpf': '000.000.023-33', 'rg': '1000021-1',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'CM Parana do Mari', 'numero': '7955',
    'bairro': 'Zona Rural', 'cidade': 'Caapiranga', 'cep': '69.425-000',
}
conta = {'agencia': '3707', 'numero': '501049-7'}
renda = {'valor_float': 3575.38}

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Caapiranga', uf='AM')
dados['valor_causa'] = fmt_moeda_rs(15000.00)
dados['valor_causa_extenso'] = extenso_moeda(15000.00)
dados['remuneração'] = 'salário do serviço público municipal'

print(f'=== CLIENTE EXEMPLO — APLIC.INVEST (b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | VC: R$ 15.000,00'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
for p in list(doc.paragraphs):
    for m in MARKERS:
        if m in (p.text or ''): p._element.getparent().remove(p._element); break

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'reiterado durante todo o período em que a renda alimentar da parte autora ficou '
        'indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = ('No caso concreto, o extrato bancário registra 159 (cento e cinquenta e nove) '
      'ocorrências de aplicação automática entre 28/01/2021 e 31/07/2024 (mais de 3 anos). '
      'Em todos os meses, parcela substancial do salário da parte autora — recebida via '
      'transferência de saldo da conta-salário — foi automaticamente subtraída pelo banco '
      'réu sob a rubrica APLIC.INVEST FACIL, restando indisponível ao consumidor por '
      'período variável até o resgate manual. Embora os valores tenham sido restituídos '
      'via RESGATE INVEST FACIL ao longo do período (491 ocorrências de resgate), o cerne '
      'do dano moral não reside na perda patrimonial líquida — inexistente, pois o saldo '
      'agregado em 3 anos foi negativo (cliente recebeu de volta tudo e mais R$ 472,82 '
      'de rentabilidade) —, mas na privação reiterada da autodeterminação do consumidor '
      'sobre sua própria renda alimentar, mês após mês, durante 3 (três) anos consecutivos. '
      'Cada retenção mensal configura, autonomamente, prática abusiva vedada pelo art. 39, '
      'inciso VI, do Código de Defesa do Consumidor, sendo a recorrência sistêmica o fato '
      'gerador do abalo extrapatrimonial.')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'APLIC.INVEST FACIL — estratégia (b) PADRÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'),
             ('Arquivo', 'INICIAL_AplicInvest_CLIENTE EXEMPLO_v1.docx'),
             ('Observação', 'Cliente já tem inicial ODONTOPREV (PG ELETRON) processada em 06/05/2026; ESTA é OUTRA tese (APLIC.INVEST FACIL).')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '08/02/1978 — 48 anos (NÃO idosa) — confirmado em sessão anterior'),
    ('Estado civil', autora['estado_civil']),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (TRANSF SALDO C/SAL — servidora municipal)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC', '159 (mais de 3 anos)'),
    ('Lançamentos RESGATE', '491'),
    ('Total bruto APLIC', dados['total_descontos']),
    ('Saldo líquido', '-R$ 472,82 (cliente recebeu R$ 472,82 a mais via rentabilidade)'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('AUDITORIA APLIC vs RESGATE — confirmada (estratégia b)',
     '159 APLICs vs 491 RESGATEs. Saldo NEGATIVO R$ 472,82.'),
    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Bloco "Repetição do indébito" REMOVIDO. Parágrafo doutrinário REESCRITO. '
     'Parágrafo do caso concreto INSERIDO (159 retenções, 3 anos).'),
    ('SEGUNDA TESE da mesma cliente',
     'Cliente já tem inicial PG ELETRON ODONTOPREV processada em 06/05/2026 (R$ 879,90, '
     '1 lançamento). Esta APLIC.INVEST é segunda ação. Decidir se vai protocolar '
     'simultaneamente ou em momentos diferentes.'),
    ('TETO JEC — folgadamente coberto', 'VC R$ 15.000 ≈ 9,87 SM.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (501049-7 / 3707).',
           'Conferir comarca: Caapiranga/AM.', 'Conferir VC = R$ 15.000,00.',
           'Decidir momento do protocolo vs ODONTOPREV já protocolada.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — '); r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
