"""Inicial MORA + ENCARGOS — CLIENTE EXEMPLO RODRIGUES PINTO.

Mesma autora da inicial TARIFAS (Maués/AM, IDOSA, INSS R$ 947,57,
Ag 3706 Conta 21299-7). Pasta separada do MORA E ENCARGOS.

Tabela: 48 lançamentos. ENCARGOS LIMITE DE CRED (a maioria, 36 lanç.)
+ MORA CREDITO PESSOAL (12 lanç. a partir de 30/04/2025). Total
R$ 2.000,80 / dobro R$ 4.001,60. VC R$ 19.001,60 — cabe no JEC.

Template `inicial-mora-encargo.docx` (cobre as 2 rubricas).
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\2. MORA\CLIENTE EXEMPLO RODRIGUES PINTO - Maurivã - TARIFAS\MORA E ENCARGOS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora-encargo.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_MoraEncargo_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_MORA_v1.docx')

LANCAMENTOS = [
    ('08/06/2021', 0.06), ('07/07/2021', 0.48), ('06/08/2021', 0.92),
    ('08/09/2021', 1.15), ('07/10/2021', 0.69), ('08/11/2021', 0.66),
    ('06/05/2022', 14.37), ('07/06/2022', 0.20), ('07/07/2022', 19.32),
    ('05/08/2022', 1.91), ('08/09/2022', 2.74), ('07/10/2022', 1.65),
    ('08/11/2022', 1.56), ('07/12/2022', 1.59), ('06/01/2023', 1.68),
    ('07/02/2023', 1.97), ('07/03/2023', 2.10), ('10/04/2023', 2.66),
    ('08/05/2023', 2.33), ('07/06/2023', 2.51), ('07/07/2023', 0.78),
    ('07/08/2023', 1.53), ('08/09/2023', 2.14), ('06/10/2023', 1.83),
    ('08/11/2023', 4.09), ('07/12/2023', 10.61), ('08/01/2024', 1.76),
    ('07/02/2024', 2.06), ('07/03/2024', 1.87), ('05/04/2024', 1.82),
    ('08/05/2024', 0.30), ('07/06/2024', 1.12), ('05/07/2024', 1.83),
    ('09/09/2024', 0.30), ('07/11/2024', 0.64), ('06/12/2024', 0.60),
    ('30/04/2025', 4.27), ('30/04/2025', 7.61), ('30/04/2025', 30.52),
    ('30/04/2025', 386.63), ('30/04/2025', 68.69), ('30/04/2025', 114.70),
    ('30/04/2025', 255.56), ('30/04/2025', 211.76), ('30/05/2025', 134.61),
    ('30/05/2025', 226.12), ('30/05/2025', 240.67), ('30/01/2026', 225.83),
]

autora = {
    'nome': 'CLIENTE EXEMPLO RODRIGUES PINTO',
    'nacionalidade': 'brasileira',
    'estado_civil': '',
    'profissao': 'aposentada',
    'cpf': '000.000.009-19',
    'rg': '1000007-7',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Ponta Alegre Apocuitaua',
    'numero': '2601',
    'bairro': 'Rio Apocuitaua',
    'cidade': 'Maués',
    'cep': '69.190-000',
}
conta = {'agencia': '3706', 'numero': '21299-7'}
renda = {'valor_float': 947.57}

tese = {'rubrica': 'MORA CRED PESS / ENC LIM CRÉDITO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Maués', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['rubrica_curta'] = 'Mora Cred Pess / Enc Lim Crédito'
dados['rubrica_curta_caps'] = 'MORA CRED PESS / ENC LIM CRÉDITO'
dados['rubrica_completa'] = 'Crédito Mora Pessoal / Encargos Limite de Crédito'
dados['rubrica_completa_caps'] = 'MORA CREDITO PESSOAL / ENCARGOS LIMITE DE CRÉDITO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — MORA + ENCARGOS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
       '{{cidade_filial}}': 'Maués',
       '{{uf_filial}}': 'AM',
       '{{uf_extenso}}': 'Amazonas'}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_MoraEncargo_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CRED PESS + ENC LIM CRÉDITO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_MoraEncargo_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 27/02/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (36 ENC LIM + 12 MORA)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO — corte 30/03/2021 (EAREsp 1.280.825 STJ)',
     '6 lançamentos antes de 07/05/2021 (08/06/2021 a 08/11/2021) podem estar '
     'prescritos. Resto válido. CONFERIR antes do protocolo.'),
    ('IDOSA — prioridade aplicada',
     'Notificação afirma "pessoa idosa" e "aposentada".'),
    ('LANÇAMENTOS GIGANTES EM 30/04/2025',
     'Em 30/04/2025 a tabela registra 8 lançamentos altos (R$ 4,27 + 7,61 + 30,52 + '
     'R$ 386,63 + R$ 68,69 + R$ 114,70 + R$ 255,56 + R$ 211,76). Em 30/05/2025: '
     'mais 3 lançamentos altos (R$ 134,61 + R$ 226,12 + R$ 240,67). Cliente '
     'provavelmente teve cheque especial muito alto nesse período. CONFERIR no '
     'extrato — pode ter sido renegociação que gerou essa cascata de encargos.'),
    ('NOTIFICAÇÃO previa dano moral R$ 15.000',
     'Notificação pleiteia R$ 15.000 dano moral (1 tese isolada). Skill aplica idem. '
     'OK.'),
    ('CLIENTE TEM 2 TESES SEPARADAS',
     'Cliente também aparece na pasta TARIFAS (já gerada). Avaliar se quer '
     'consolidar em 1 só inicial-combinada. Inicial gerada como ISOLADA MORA + '
     'ENCARGOS (mesma família — IRDR 0004464 TJAM).'),
    ('TETO JEC — coberto',
     'VC R$ 19.001,60 ≈ 12,5 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'CONFERIR cascata de lançamentos em 30/04/2025 e 30/05/2025 no extrato.',
    'AVALIAR consolidação com TARIFAS em 1 só inicial-combinada.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Maués/AM.',
    'Conferir VC = R$ 19.001,60 e dano moral R$ 15.000,00.',
    'Confirmar com cliente: nunca contratou cheque especial.',
    'Anexar 2-Procurações (ENC LIM + MORA CRED PESS) + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após confirmar cascata e decidir consolidação, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
