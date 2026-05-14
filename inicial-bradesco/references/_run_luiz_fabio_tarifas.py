"""Inicial TARIFAS — CLIENTE EXEMPLO PEREIRA DA SILVA GUERRA.

Comarca Presidente Figueiredo/AM (Ag 3732 / Conta 1525-3). Pessoa
IDOSA (afirmação na notificação). Aposentado pelo INSS recente
(R$ 1.091,86 em 30/01/2026 — abaixo de 1 SM, sugere consignações).

Tabela: 19 lançamentos TARIFA BANCARIA - CESTA FACIL ECONOMICA entre
15/07/2020 e 15/04/2021. Total R$ 257,08 / dobro R$ 514,16. VC R$
15.514,16 — cabe folgadamente no JEC.

PRESCRIÇÃO CRÍTICA: TODOS os lançamentos são de 2020-2021. Considerando
art. 27 CDC (5 anos) e marco EAREsp 1.280.825 STJ (30/03/2021):
  - Antes de 30/03/2021: 15 lançamentos PRESCRITOS (~R$ 167)
  - A partir de 30/03/2021: 4 lançamentos válidos (~R$ 90)

Caso fraco. CONFIRMAR estratégia com procurador antes do protocolo.

PENDÊNCIA: documento de identificação é CNH (não RG). Nº 16138856
adotado da notificação.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO PEREIRA DA SILVA GUERRA - Ruth - TARIFAS\TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_LUIZ_FABIO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_LUIZ_FABIO_v1.docx')

LANCAMENTOS = [
    ('15/07/2020', 0.53), ('23/07/2020', 25.37), ('14/08/2020', 0.64),
    ('21/08/2020', 25.26), ('15/09/2020', 1.75), ('28/09/2020', 24.15),
    ('15/10/2020', 2.48), ('27/10/2020', 25.33), ('13/11/2020', 1.68),
    ('26/11/2020', 26.13), ('15/12/2020', 4.75), ('16/12/2020', 23.06),
    ('15/01/2021', 3.95), ('08/02/2021', 23.86), ('12/02/2021', 6.14),
    ('30/03/2021', 21.67), ('30/03/2021', 28.33), ('01/04/2021', 11.07),
    ('15/04/2021', 0.93),
]

autora = {
    'nome': 'LUIZ FÁBIO PEREIRA DA SILVA GUERRA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': '',
    'cpf': '000.000.016-26',
    'rg': '1000014-4',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Manoel (RM) do Urubuí',
    'numero': '23 A - KM 01',
    'bairro': 'Centro',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '1525-3'}
renda = {'valor_float': 1091.86}

tese = {'rubrica': 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'benefício do INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_LUIZ_FABIO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_LUIZ_FABIO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'NÃO IDOSO (44 anos — CNH mostra 23/10/1981). Notificação errou.'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CNH)', f'18138858 SSP/AM (notificação trouxe 16138856 — ERRO; CNH 286.0165139 confirma 18.138.858)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 30/01/2026 — abaixo 1 SM)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO CRÍTICA — caso fraco',
     'TODOS os 19 lançamentos são de 15/07/2020 a 15/04/2021. Hoje é 07/05/2026 '
     '(5 anos atrás = 07/05/2021). Considerando art. 27 CDC (5 anos) e marco EAREsp '
     '1.280.825 STJ (30/03/2021): ANTES de 30/03/2021 = 15 lançamentos prescritos '
     '(~R$ 167); A PARTIR de 30/03/2021 = 4 lançamentos válidos (~R$ 90). DECISÃO DO '
     'PROCURADOR: (a) pleitear TUDO testando art. 205 CC (10 anos — minoritária TJAM); '
     '(b) reduzir VC ao mínimo do art. 27 CDC (~R$ 180 dobro + R$ 15.000 = R$ 15.180); '
     '(c) DESISTIR do caso (valor não compensa custo). Inicial gerada com TODOS '
     '19 lançamentos — REVISAR antes do protocolo.'),
    ('NÃO IDOSO — prioridade REMOVIDA (44 anos)',
     'CNH mostra nascimento 23/10/1981 (44 anos completos em 07/05/2026). Notificação '
     'errou ao afirmar "pessoa idosa". Prioridade do art. 1.048 CPC FOI REMOVIDA. '
     'Naturalidade Manaus/AM. Filiação: LUIZ GUERRA DA SILVA / GENY PEREIRA DA SILVA. '
     '1ª habilitação 05/02/2001.'),
    ('RG CORRIGIDO — 18138858 SSP/AM (CNH)',
     'A notificação extrajudicial trouxe RG 16138856, mas a CNH mostra 18138858 SSP/AM '
     '(Nº Registro CNH 286.0165139). Inicial corrigida. Anexar a CNH como documento '
     'de identificação.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente.'),
    ('Renda baixa (< 1 SM) — possível consignação',
     'INSS R$ 1.091,86 (30/01/2026) é abaixo do salário mínimo 2026 (R$ 1.518). '
     'Indica descontos consignados pelo próprio INSS. Renda BRUTA pode ser maior '
     '(R$ 1.518 ou mais). Conferir HISCON. Reforça hipossuficiência.'),
    ('Sem comprovante de notificação (8.1)',
     'Pasta não tem 8.1-COMPROVANTE NOTIFICAÇÃO. Conferir se foi efetivamente '
     'enviada/recebida pela ouvidoria do Bradesco.'),
    ('NOTIFICAÇÃO previa dano moral R$ 5.000',
     'A notificação pleiteou R$ 5.000 (provavelmente projetando 2 teses combinadas: '
     'TARIFAS + MORA). Como esta inicial é ISOLADA, pleiteia R$ 15.000 conforme regra '
     '§ 9 da skill. CONFIRMAR com procurador.'),
    ('CLIENTE TEM 2 TESES SEPARADAS',
     'Cliente também aparece na pasta MORA. AVALIAR consolidação em 1 só inicial-'
     'combinada (Presidente Figueiredo adota combinação por padrão).'),
    ('TETO JEC — coberto',
     'VC R$ 15.514,16 ≈ 10,2 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AVALIAR estratégia frente à prescrição (caso fraco — pode não compensar).',
    'AVALIAR consolidação com MORA em 1 só inicial-combinada.',
    'Conferir nome / CPF / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Verificar comprovante de envio da notificação extrajudicial.',
    'Anexar CNH como documento de identificação.',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Confirmar com cliente: nunca contratou cesta fácil econômica.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir prescrição. Caso fraco — avaliar com procurador. ')
r2 = p.add_run('Ressalvas críticas.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
