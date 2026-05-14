"""Inicial TÍTULO DE CAPITALIZAÇÃO — MARIA DO LIVRAMENTO LIMA DOS SANTOS."""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\3. Título de capitalização\MARIA DO LIVRAMENTO LIMA DOS SANTOS - Ruth\TÍTULO DE CAPITALIZAÇÃO'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Titulo_MARIA_LIVRAMENTO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MARIA_LIVRAMENTO_TITULO_v1.docx')

LANCAMENTOS = [
    ('04/11/2020', 20.00), ('05/01/2021', 20.00), ('02/02/2021', 20.00),
    ('02/03/2021', 20.00), ('30/03/2021', 20.00), ('30/04/2021', 20.00),
    ('31/05/2021', 20.00), ('30/06/2021', 20.00), ('30/07/2021', 20.00),
    ('02/09/2021', 20.00), ('04/10/2021', 20.00), ('03/11/2021', 20.00),
    ('03/11/2021', 20.00), ('02/12/2021', 20.00), ('04/01/2022', 22.00),
    ('02/02/2022', 22.00), ('04/03/2022', 22.00), ('04/04/2022', 22.00),
    ('03/05/2022', 22.00), ('02/06/2022', 22.00), ('30/06/2022', 22.00),
    ('02/08/2022', 22.00), ('06/09/2022', 22.00), ('04/10/2022', 22.00),
    ('03/11/2022', 22.00), ('02/12/2022', 23.92), ('03/01/2023', 23.92),
    ('02/02/2023', 23.92), ('02/03/2023', 23.92), ('04/04/2023', 23.92),
    ('03/05/2023', 23.92), ('02/06/2023', 23.92), ('03/07/2023', 23.92),
    ('02/08/2023', 23.92), ('30/08/2023', 23.92),
]

autora = {
    'nome': 'MARIA DO LIVRAMENTO LIMA DOS SANTOS', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': 'aposentada',
    'cpf': '000.000.021-31', 'rg': '1000019-9', 'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Ramal do Rumo Certo, BR 174 - KM 165', 'numero': 's/nº',
    'bairro': 'Comunidade Boa União', 'cidade': 'Presidente Figueiredo', 'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '21528-7'}
renda = {'valor_float': 986.58}
tese = {'rubrica': 'TÍTULO DE CAPITALIZAÇÃO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TÍTULO DE CAPITALIZAÇÃO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== MARIA DO LIVRAMENTO — TÍTULO ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Titulo_MARIA_LIVRAMENTO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TÍTULO DE CAPITALIZAÇÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Titulo_MARIA_LIVRAMENTO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]}'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)}'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('TEMPLATE GENÉRICO ADAPTADO', 'Ajustar fundamentação para TÍTULO.'),
    ('PRESCRIÇÃO — corte 30/03/2021', '~5 lançamentos pré-30/03/2021. Pós: 30 válidos.'),
    ('IDOSA — prioridade aplicada', 'Notificação afirma "pessoa idosa" e "aposentada".'),
    ('CLIENTE TEM 3 TESES', 'MARIA também aparece em TARIFAS (já gerada) e PG ELETRON. AVALIAR consolidação.'),
    ('CASAL CO-DEMANDANTE', 'Endereço idêntico ao EXEMPLO MANUEL DOS SANTOS — provavelmente cônjuges. Avaliar litisconsórcio.'),
    ('TETO JEC — coberto', 'VC R$ 16.522,40 ≈ 10,9 SM.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'AJUSTAR fundamentação para TÍTULO.',
    'AVALIAR consolidação com TARIFAS + PG ELETRON.',
    'DECIDIR estratégia de prescrição.',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Confirmar com cliente: nunca contratou título de capitalização.',
    'Anexar 2-Procuração + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem ajustes. ')
r2 = p.add_run('Atenção.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
