"""Inicial TÍTULO DE CAPITALIZAÇÃO — EXEMPLO MARTINS DA SILVA.

Mesma autora (Manacapuru/AM, IDOSO, INSS R$ 1.621,00, Ag 3707 Conta
8698-3, procuração a rogo). Template adaptado: inicial-tarifas.docx
com {{titulo}} = "TÍTULO DE CAPITALIZAÇÃO".

Tabela: 104 lançamentos TITULO DE CAPITALIZACAO entre 31/01/2017 e
30/01/2026 (mensal por 9 anos!). Total R$ 3.080,37 / dobro R$ 6.160,74.
VC R$ 21.160,74 — cabe no JEC.

PRESCRIÇÃO: muitos lançamentos antes 30/03/2021. Procurador decide.

ATENÇÃO: template inicial-tarifas é GENÉRICO. A fundamentação jurídica
para TÍTULO DE CAPITALIZAÇÃO requer ajustes específicos (TC nº 415 SUSEP,
arts. 6º III, 39 VI, 52 CDC + jurisprudência sobre venda casada de
título). REVISAR antes do protocolo.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\3. Título de capitalização\EXEMPLO MARTINS DA SILVA - Ney Pedroza - TARIFAS\TÍTULO DE CAPITALIZAÇÃO'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Titulo_JOAO_MARTINS_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_JOAO_MARTINS_TITULO_v1.docx')

LANCAMENTOS = [
    ('31/01/2017', 55.48), ('31/03/2017', 57.33), ('02/05/2017', 29.59),
    ('30/05/2017', 29.59), ('30/06/2017', 29.59), ('31/07/2017', 29.59),
    ('31/08/2017', 29.59), ('02/10/2017', 29.59), ('30/10/2017', 29.59),
    ('30/11/2017', 29.59), ('02/01/2018', 29.59), ('30/01/2018', 29.59),
    ('28/02/2018', 29.59), ('02/04/2018', 29.47), ('30/04/2018', 29.47),
    ('30/05/2018', 29.47), ('02/07/2018', 29.47), ('30/07/2018', 29.47),
    ('30/08/2018', 29.47), ('31/10/2018', 58.94), ('19/11/2018', 500.00),
    ('30/11/2018', 29.47), ('31/01/2019', 58.94), ('23/04/2019', 20.00),
    ('31/05/2019', 20.00), ('31/07/2019', 40.00), ('30/08/2019', 20.00),
    ('30/09/2019', 20.00), ('31/10/2019', 20.00), ('02/12/2019', 20.00),
    ('30/12/2019', 20.00), ('31/01/2020', 20.00), ('23/04/2020', 40.00),
    ('30/04/2020', 20.86), ('01/06/2020', 20.86), ('30/06/2020', 20.86),
    ('31/07/2020', 20.86), ('31/08/2020', 20.86), ('30/09/2020', 20.86),
    ('30/10/2020', 20.86), ('30/11/2020', 20.86), ('30/12/2020', 20.86),
    ('26/02/2021', 20.86), ('01/03/2021', 20.86), ('30/03/2021', 20.86),
    ('30/04/2021', 21.80), ('31/05/2021', 21.80), ('30/06/2021', 21.80),
    ('30/07/2021', 21.80), ('30/08/2021', 21.80), ('30/09/2021', 21.80),
    ('01/11/2021', 21.80), ('01/11/2021', 21.80), ('30/11/2021', 21.80),
    ('30/12/2021', 21.80), ('31/01/2022', 21.80), ('02/03/2022', 21.80),
    ('30/03/2022', 21.80), ('02/05/2022', 24.10), ('30/05/2022', 24.10),
    ('30/06/2022', 24.10), ('01/08/2022', 24.10), ('30/08/2022', 24.10),
    ('30/09/2022', 24.10), ('31/10/2022', 24.10), ('30/11/2022', 24.10),
    ('02/01/2023', 24.10), ('31/01/2023', 24.10), ('28/02/2023', 24.10),
    ('31/03/2023', 24.10), ('02/05/2023', 25.49), ('30/05/2023', 25.49),
    ('30/06/2023', 25.49), ('31/07/2023', 25.49), ('31/08/2023', 25.49),
    ('02/10/2023', 25.49), ('31/10/2023', 25.49), ('30/11/2023', 25.49),
    ('02/01/2024', 25.49), ('31/01/2024', 25.49), ('29/02/2024', 25.49),
    ('01/04/2024', 25.49), ('29/04/2024', 20.00), ('31/05/2024', 20.00),
    ('01/07/2024', 20.00), ('30/07/2024', 20.00), ('30/08/2024', 20.00),
    ('30/09/2024', 20.00), ('30/10/2024', 20.00), ('02/12/2024', 20.00),
    ('30/12/2024', 20.00), ('30/01/2025', 20.00), ('28/02/2025', 20.00),
    ('31/03/2025', 20.00), ('30/04/2025', 20.91), ('30/05/2025', 20.91),
    ('30/06/2025', 20.91), ('30/07/2025', 20.91), ('01/09/2025', 20.91),
    ('30/09/2025', 20.91), ('30/10/2025', 20.91), ('01/12/2025', 20.91),
    ('30/12/2025', 20.91), ('30/01/2026', 20.91),
]

autora = {
    'nome': 'EXEMPLO MARTINS DA SILVA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'aposentado',
    'cpf': '000.000.013-23',
    'rg': '1000011-1',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua do Campinho',
    'numero': '1333',
    'bairro': 'Novo Manaca',
    'cidade': 'Manacapuru',
    'cep': '69.409-899',
}
conta = {'agencia': '3707', 'numero': '8698-3'}
renda = {'valor_float': 1621.00}

tese = {'rubrica': 'TÍTULO DE CAPITALIZAÇÃO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Manacapuru', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TÍTULO DE CAPITALIZAÇÃO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== EXEMPLO MARTINS — TÍTULO DE CAPITALIZAÇÃO ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Titulo_JOAO_MARTINS', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TÍTULO DE CAPITALIZAÇÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Titulo_JOAO_MARTINS_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 30/01/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (mensal por 9 anos)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('TEMPLATE GENÉRICO ADAPTADO — REVISAR FUNDAMENTAÇÃO',
     'Skill ainda não tem template próprio para TÍTULO DE CAPITALIZAÇÃO. Foi usado '
     'template inicial-tarifas.docx com {{titulo}} = "TÍTULO DE CAPITALIZAÇÃO". A '
     'fundamentação doutrinária menciona Resoluções BACEN 3.919/2010 e 4.196/2013 '
     '(adequadas para TARIFAS, não para TÍTULO). Para TÍTULO DE CAPITALIZAÇÃO é '
     'mais adequado citar: TC nº 415 SUSEP, Lei 11.795/2008 (consórcios), arts. 6º '
     'III, 39 VI, 52 CDC + jurisprudência sobre venda casada de título de '
     'capitalização. AJUSTAR antes do protocolo. IRDR 0005053 TJAM citado é genérico '
     'sobre tarifas — pode ou não cobrir TÍTULO. Conferir.'),
    ('LANÇAMENTO DE 19/11/2018 — R$ 500,00',
     'Em 19/11/2018 a tabela registra R$ 500,00 (valor anômalo no contexto de '
     'lançamentos mensais de R$ 20-30). Pode ser depósito inicial do título ou '
     'compra extraordinária de cota. Conferir extrato e descrição da operação.'),
    ('PRESCRIÇÃO CRÍTICA — corte 30/03/2021',
     '~46 lançamentos antes de 30/03/2021 podem estar prescritos (~R$ 1.450 simples). '
     'Pós-30/03/2021: 58 lançamentos válidos (~R$ 1.630 simples / R$ 3.260 dobro). '
     'Inicial gerada com TODOS — REVISAR.'),
    ('IDOSO — prioridade aplicada',
     'Notificação afirma "pessoa idosa".'),
    ('PROCURAÇÃO ASSINADA A ROGO',
     'Pasta tem RG da rogada (Marilene Pereira da Silva) + 2 testemunhas (Rosiane e '
     'Deilson).'),
    ('CLIENTE TEM 3 TESES SEPARADAS',
     'EXEMPLO MARTINS aparece em TARIFAS (já gerada), MORA (já gerada), TÍTULO (esta). '
     'AVALIAR consolidação em inicial-combinada.'),
    ('TETO JEC — coberto',
     'VC R$ 21.160,74 ≈ 13,9 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AJUSTAR fundamentação doutrinária para TÍTULO DE CAPITALIZAÇÃO (TC SUSEP).',
    'CONFERIR lançamento R$ 500 em 19/11/2018 no extrato.',
    'DECIDIR estratégia de prescrição.',
    'AVALIAR consolidação com TARIFAS + MORA.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade da procuração assinada a rogo.',
    'Conferir comarca: Manacapuru/AM.',
    'Confirmar com cliente: nunca contratou título de capitalização.',
    'Anexar 2-Procuração + 3-RG + 3.1 a 3.3 (rogada + testemunhas) + 4-Hipossuficiência + 5-Comprovante + 5.1-Declaração + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem AJUSTAR fundamentação para TÍTULO e decidir prescrição/consolidação. ')
r2 = p.add_run('Atenção crítica.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
