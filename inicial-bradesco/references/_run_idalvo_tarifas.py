"""Inicial TARIFAS — CLIENTE EXEMPLO SOUZA CUNHA.

Comarca Caapiranga/AM (Ag 3707, endereço Caapiranga). Pessoa IDOSA
(afirmação expressa na notificação extrajudicial). Servidor(a) recebia
TRANSF SALDO C/SAL P/CC da Prefeitura de Caapiranga (R$ 1.306,10 em
10/02/2025 — último crédito de salário no extrato; em 2026 não há
movimentação salarial registrada).

Tabela: 22 lançamentos TARIFA BANCARIA - CESTA B EXPRESSO entre
20/10/2021 e 30/07/2025. Total R$ 582,55 / dobro R$ 1.165,10. VC
R$ 16.165,10 — cabe folgadamente no JEC.

PENDÊNCIAS: (1) RG não consta na notificação — número precisa ser
extraído manualmente do 3-RG.pdf. (2) Estado civil omitido. (3) Idade
confirmar no RG.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\CLIENTE EXEMPLO SOUZA CUNHA - Ney Pedroza - TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

LANCAMENTOS = [
    ('20/10/2021', 38.60), ('12/11/2021', 38.60), ('13/05/2022', 0.14),
    ('10/06/2022', 44.36), ('26/04/2023', 49.90), ('26/04/2023', 49.90),
    ('15/02/2024', 2.84), ('12/04/2024', 53.91), ('12/04/2024', 56.75),
    ('15/04/2024', 1.78), ('24/05/2024', 54.97), ('15/05/2025', 0.15),
    ('16/05/2025', 16.00), ('22/05/2025', 4.00), ('26/05/2025', 43.45),
    ('13/06/2025', 22.45), ('16/06/2025', 41.15), ('15/07/2025', 30.85),
    ('17/07/2025', 5.03), ('21/07/2025', 1.00), ('28/07/2025', 12.00),
    ('30/07/2025', 14.72),
]

autora = {
    'nome': 'CLIENTE EXEMPLO SOUZA CUNHA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'servidor público municipal',
    'cpf': '000.000.012-22',
    'rg': '1000010-0',
    'orgao_expedidor_prefixo': 'CIN (CPF como Registro Geral)',
    'logradouro': 'CM Nossa Sra. Da Conceição',
    'numero': 's/nº',
    'bairro': 'Distrito de Campinas',
    'cidade': 'Caapiranga',
    'cep': '69.414-900',
}
conta = {'agencia': '3707', 'numero': '45809-0'}
renda = {'valor_float': 1306.10}

tese = {'rubrica': 'TARIFA BANCÁRIA - CESTA B EXPRESSO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Caapiranga', uf='AM')
dados['remuneração'] = 'salário do serviço público municipal'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA B EXPRESSO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA - CESTA B EXPRESSO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'NÃO IDOSO (49 anos — RG mostra 29/09/1976). Notificação errou ao afirmar "pessoa idosa".'),
    ('Sexo / nacionalidade', 'MASCULINO — "brasileiro" (notificação dizia "brasileira")'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG (CIN nova)', f'Registro Geral = CPF ({autora["rg"]}) — Carteira de Identidade Nacional unifica RG/CPF'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (TRANSF SALDO C/SAL — último crédito 10/02/2025)'),
    ('Período tabela', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('RG (CIN nova) — Registro Geral = CPF',
     'O RG do cliente é a nova Carteira de Identidade Nacional (CIN), em que o '
     'Registro Geral é o próprio número do CPF (000.000.012-22). A notificação '
     'extrajudicial não informou o número porque coincide com o CPF. Inicial corrigida: '
     'campo RG preenchido com 000.000.012-22 e órgão expedidor "CIN (CPF como Registro Geral)".'),
    ('NÃO IDOSO — 49 anos (29/09/1976) — prioridade REMOVIDA',
     'O RG mostra data de nascimento 29/09/1976, ou seja, 49 anos completos em 07/05/2026. '
     'A notificação extrajudicial errou ao afirmar "pessoa idosa". A prioridade do art. '
     '1.048, I, do CPC FOI REMOVIDA da inicial. O autor ainda pode pleitear gratuidade '
     '(art. 98 CPC) e a inversão do ônus probatório (CDC art. 6º VIII).'),
    ('SEXO MASCULINO — "brasileiro" (notificação dizia "brasileira")',
     'O RG mostra sexo M (masculino). A notificação extrajudicial errou ao usar femininos '
     '("brasileira", "domiciliada", "Notificante" etc.). A inicial foi corrigida para '
     '"brasileiro" e profissão masculina.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente.'),
    ('Profissão SERVIDOR PÚBLICO MUNICIPAL — Prefeitura de Caapiranga',
     'Recebia salário via Prefeitura Municipal de Caapiranga (TRANSF SALDO C/SAL '
     'P/CC). Último crédito 10/02/2025 R$ 1.306,10. Em 2026 não há movimentação '
     'salarial registrada no Bradesco — pode ter mudado de banco ou ter sido '
     'desligado. CONFIRMAR status atual com o cliente (cargo, vínculo, '
     'banco onde recebe).'),
    ('Renda CONFIRMADA R$ 1.306,10',
     'Último crédito de salário no extrato. Cliente tem 49 anos, servidor da '
     'Prefeitura de Caapiranga.'),
    ('TETO JEC — coberto',
     'VC R$ 16.165,10 ≈ 10,6 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
    ('Comarca Caapiranga — Justiça Estadual Delegada',
     'Endereço do autor em Caapiranga + Ag 3707 (Caapiranga) + salário via Prefeitura '
     'de Caapiranga. Comarca de Caapiranga sem JEF — Justiça Estadual Delegada.'),
    ('LANÇAMENTOS DE 26/04/2023 — 2 idênticos',
     'Em 26/04/2023 a tabela registra 2 lançamentos de R$ 49,90 cada. Pode ser cobrança '
     'em duplicidade no mesmo dia ou erro de digitação. Conferir no extrato original.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'Conferir nome / CPF (000.000.012-22 = RG via CIN nova).',
    'Conferir status profissional atual (servidor ativo? desligado?).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Caapiranga/AM (Justiça Estadual Delegada).',
    'Conferir VC = R$ 16.165,10 e dano moral R$ 15.000,00.',
    'Confirmar com cliente: nunca contratou cesta de tarifas Bradesco.',
    'Conferir status profissional atual (em 2026 sem salário no extrato).',
    'Anexar 2-Procuração + 3-RG + 4-Hipossuficiência + 5-Comprovante residência + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — RG (CIN nova) confirmado, sexo masculino corrigido, prioridade idoso REMOVIDA. Após confirmar status profissional, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
