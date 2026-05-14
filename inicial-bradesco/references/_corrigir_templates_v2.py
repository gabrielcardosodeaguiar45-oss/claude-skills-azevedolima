"""Round 2 — versão FINAL sem usar substituir_in_run em substituições onde
NOVO contém ANTIGO (causa loop infinito).

1. Adicionar marcador {{SE_IDOSO}} no início dos parágrafos condicionais
   de prioridade idoso.
2. Aplicar BOLD + SUBLINHADO aos placeholders {{contrato_numero}} e
   {{contratos_lista_breve}}.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stdout.reconfigure(line_buffering=True)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from lxml import etree
from copy import deepcopy

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

BASE_DIR = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates'
T = {
    'base':  os.path.join(BASE_DIR, 'inicial-jfba-base.docx'),
    'mult':  os.path.join(BASE_DIR, 'inicial-jfba-multiplos-avn-inativo.docx'),
    'refin': os.path.join(BASE_DIR, 'inicial-jfba-refin-ativo.docx'),
}


def fix6_marcador_se_idoso_direto(doc):
    """Insere {{SE_IDOSO}} no INÍCIO do primeiro <w:t> do parágrafo, sem
    usar substituir_in_run (que entraria em loop)."""
    n = 0
    padroes = ('Prioridade de tramitação: art. 1.048', 'A prioridade na tramitação,')
    for p in doc.paragraphs:
        if '{{SE_IDOSO}}' in p.text:
            continue  # já marcado
        if not any(pad in p.text for pad in padroes):
            continue
        # Localizar o primeiro <w:t> com texto
        primeiro_t = None
        for t in p._element.findall('.//' + W + 't'):
            if t.text:
                primeiro_t = t
                break
        if primeiro_t is None:
            continue
        # Prefixar com {{SE_IDOSO}}
        primeiro_t.text = '{{SE_IDOSO}}' + primeiro_t.text
        primeiro_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        n += 1
    return n


def fix7_bold_underline(doc):
    """Aplica bold + underline aos placeholders {{contrato_numero}} e
    {{contratos_lista_breve}}.

    Para cada parágrafo, percorre seus runs. Se encontrar um run com o
    placeholder dentro do texto, divide em até 3 runs (prefixo, placeholder
    formatado, sufixo).
    """
    placeholders = ('{{contrato_numero}}', '{{contratos_lista_breve}}')
    n_total = 0

    for p in doc.paragraphs:
        # Captura snapshot dos runs ANTES de modificar
        runs_lista = list(p._element.findall('.//' + W + 'r'))
        for r in runs_lista:
            t_el = r.find(W + 't')
            if t_el is None or not t_el.text:
                continue
            txt = t_el.text
            for ph in placeholders:
                if ph not in txt:
                    continue
                # Verifica se já está formatado
                rpr = r.find(W + 'rPr')
                tem_b = rpr is not None and rpr.find(W + 'b') is not None
                tem_u = rpr is not None and rpr.find(W + 'u') is not None
                if tem_b and tem_u:
                    continue

                # Divide em prefixo, placeholder, sufixo
                idx_ini = txt.find(ph)
                idx_fim = idx_ini + len(ph)
                prefixo = txt[:idx_ini]
                sufixo = txt[idx_fim:]

                # Se prefixo+sufixo ambos vazios → o run é SÓ o placeholder
                if not prefixo and not sufixo:
                    if rpr is None:
                        rpr = etree.Element(W + 'rPr')
                        r.insert(0, rpr)
                    if rpr.find(W + 'b') is None:
                        b = etree.Element(W + 'b')
                        rpr.append(b)
                    if rpr.find(W + 'u') is None:
                        u = etree.Element(W + 'u')
                        u.set(W + 'val', 'single')
                        rpr.append(u)
                    n_total += 1
                    continue

                # Caso geral: dividir em runs
                # 1) Atualizar t do run original com prefixo
                t_el.text = prefixo
                if prefixo and (prefixo[0].isspace() or prefixo[-1].isspace()):
                    t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                # 2) Criar novo run para placeholder com bold+underline
                parent = r.getparent()
                idx_pos = list(parent).index(r) + 1
                r_ph = etree.Element(W + 'r')
                if rpr is not None:
                    r_ph.append(deepcopy(rpr))
                rpr_ph = r_ph.find(W + 'rPr')
                if rpr_ph is None:
                    rpr_ph = etree.SubElement(r_ph, W + 'rPr')
                if rpr_ph.find(W + 'b') is None:
                    rpr_ph.append(etree.Element(W + 'b'))
                if rpr_ph.find(W + 'u') is None:
                    u = etree.Element(W + 'u')
                    u.set(W + 'val', 'single')
                    rpr_ph.append(u)
                t_ph = etree.SubElement(r_ph, W + 't')
                t_ph.text = ph
                t_ph.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                parent.insert(idx_pos, r_ph)
                idx_pos += 1
                # 3) Criar novo run para sufixo (se houver)
                if sufixo:
                    r_suf = etree.Element(W + 'r')
                    if rpr is not None:
                        r_suf.append(deepcopy(rpr))
                    t_suf = etree.SubElement(r_suf, W + 't')
                    t_suf.text = sufixo
                    if sufixo[0].isspace() or sufixo[-1].isspace():
                        t_suf.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    parent.insert(idx_pos, r_suf)

                n_total += 1
                # NÃO BREAK - mas como modificamos o XML, parar de processar este run
                break
    return n_total


def processar(label, path):
    print(f'\n████ {label} — {os.path.basename(path)} ████', flush=True)
    doc = Document(path)
    n6 = fix6_marcador_se_idoso_direto(doc)
    print(f'  fix6: {n6}', flush=True)
    n7 = fix7_bold_underline(doc)
    print(f'  fix7: {n7}', flush=True)
    doc.save(path)
    print(f'  SALVO', flush=True)


if __name__ == '__main__':
    for label, path in T.items():
        processar(label.upper(), path)
    print('\nFIM', flush=True)
