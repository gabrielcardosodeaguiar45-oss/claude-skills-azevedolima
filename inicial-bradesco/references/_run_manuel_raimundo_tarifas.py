"""Inicial TARIFAS — EXEMPLO MANUEL DOS SANTOS.

Comarca Presidente Figueiredo/AM (Ag 3732 / Conta 8799-8). Pessoa
IDOSA (afirmação na notificação). Procuração assinada A ROGO.

CASO MUITO FRACO: tabela tem APENAS 2 lançamentos (R$ 0,51 + R$ 5,00 =
R$ 5,51 simples / R$ 11,02 dobro), com gap de quase 5 anos entre eles
(14/05/2021 e 15/01/2026). Risco alto de improcedência por bagatela.

PENDÊNCIA CRÍTICA RENDA: extrato não tem nenhum crédito INSS/SALARIO/
BENEFÍCIO no período. Único movimento é TED do BANCO PAN em 22/04/2021
(provavelmente empréstimo). Cliente provavelmente recebe via outra
conta. RENDA pleiteada [A CONFIRMAR] — placeholder.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\EXEMPLO MANUEL DOS SANTOS - Ruth - TARIFAS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_MANUEL_RAIMUNDO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MANUEL_RAIMUNDO_v1.docx')

LANCAMENTOS = [
    ('14/05/2021', 0.51), ('15/01/2026', 5.00),
]

autora = {
    'nome': 'EXEMPLO MANUEL DOS SANTOS',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'aposentado',
    'cpf': '000.000.019-29',
    'rg': '1000017-7',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Ramal do Rumo Certo, BR 174 - KM 165',
    'numero': 's/nº',
    'bairro': 'Comunidade Boa União',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-500',
}
conta = {'agencia': '3732', 'numero': '8799-8'}
renda = {'valor_float': 1518.00}

tese = {'rubrica': 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== EXEMPLO MANUEL — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_MANUEL_RAIMUNDO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA - CESTA FÁCIL ECONÔMICA'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_MANUEL_RAIMUNDO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (PADRÃO 1 SM 2026 — extrato Bradesco SEM créditos de renda)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('CASO MUITO FRACO — apenas 2 lançamentos com gap de 5 anos',
     'A tabela registra APENAS 2 lançamentos: R$ 0,51 em 14/05/2021 e R$ 5,00 em '
     '15/01/2026. Total simples R$ 5,51, dobro R$ 11,02. Pleitear R$ 15.000 de dano '
     'moral por R$ 5,51 de prejuízo material é DESPROPORCIONAL e pode levar a '
     'improcedência por princípio da bagatela. Risco também de o juiz identificar '
     'litigância predatória. CONFIRMAR estratégia com procurador antes do protocolo.'),
    ('PRESCRIÇÃO — 1º lançamento de 14/05/2021',
     'O 1º lançamento (14/05/2021) está dentro dos 5 anos (corte 07/05/2021). '
     'Marginalmente válido. O 2º lançamento (15/01/2026) é recente e claramente '
     'válido. Sem prescrição relevante.'),
    ('RENDA — extrato SEM créditos INSS/SALARIO',
     'O extrato Bradesco NÃO mostra nenhum crédito de INSS, salário ou benefício no '
     'período. Único movimento é TED do BANCO PAN em 22/04/2021 (R$ 3.242,51 — '
     'provavelmente empréstimo). Cliente provavelmente recebe via OUTRA conta '
     '(Caixa, BB, Banco da Amazônia etc.). RENDA pleiteada usa PADRÃO 1 SM '
     '(R$ 1.518,00) como placeholder. CONFIRMAR com cliente e ajustar antes do '
     'protocolo.'),
    ('IDOSO CONFIRMADO — 77 anos (15/08/1948)',
     'RG confirmado via OCR: data nascimento 15/08/1948, 77 anos completos. '
     'Naturalidade Almeirim/PA. Filiação JUSTINO DIAS DOS SANTOS / HILDA BENAION DOS SANTOS. '
     'RG nº 0820921-9 SSP/AM (notificação trouxe 0820923-9 — diferença num dígito; '
     'inicial corrigida).'),
    ('PROCURAÇÃO ASSINADA A ROGO',
     'Pasta tem RG da rogada (Maria) + 2 testemunhas (José Felício e Nuberlandia). '
     'Conferir validade.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente.'),
    ('TETO JEC — coberto, mas pode ser ínfimo',
     'VC R$ 15.011,02 ≈ 9,9 SM. Cabe no JEC. Mas o pleito do dano moral pode ser '
     'reduzido pelo juiz dada a magnitude irrisória do prejuízo material.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AVALIAR fortemente se vale a pena protocolar (caso muito fraco).',
    'CONFIRMAR renda real do cliente (extrato Bradesco não tem créditos).',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade da procuração assinada a rogo.',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Confirmar com cliente: nunca contratou cesta fácil econômica.',
    'Anexar 2-Procuração + 3-RG + 3.1-RG rogada + 3.2/3.3-RGs testemunhas + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem confirmar renda e avaliar viabilidade. Caso muito fraco — pode levar a improcedência por bagatela. ')
r2 = p.add_run('Atenção crítica.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
