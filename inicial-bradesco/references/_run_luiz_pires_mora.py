"""Inicial MORA + ENCARGOS — CLIENTE EXEMPLO.

Mesma autora da TARIFAS (Presidente Figueiredo/AM, IDOSO, casado,
INSS R$ 988,00, Ag 3732 Conta 20304-1, RG SSP/AC, procuração a rogo).

Tabela: 78 lançamentos. ENCARGOS LIMITE DE CRED (~50) + SERVIÇO CARTÃO
PROTEGIDO (~28) entre 02/01/2020 e 02/12/2024. Total R$ 421,79 / dobro
R$ 843,58. VC R$ 15.843,58. NÃO tem MORA CRED PESSOAL — substitui
"mora" por "serviço cartão protegido" no template.

PRESCRIÇÃO: muitos lançamentos antes de 30/03/2021. Alertar.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal, extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\2. MORA\CLIENTE EXEMPLO - Ruth - TARIFA\ENCARGOS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora-encargo.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_EncargoServico_LUIZ_PIRES_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_LUIZ_PIRES_MORA_v1.docx')

ENCARGO = [
    ('02/01/2020', 1.46), ('03/02/2020', 0.66), ('02/03/2020', 0.21),
    ('01/06/2020', 0.08), ('03/08/2020', 0.08), ('01/09/2020', 0.12),
    ('01/10/2020', 1.20), ('03/11/2020', 0.87), ('01/12/2020', 0.28),
    ('04/01/2021', 1.20), ('01/02/2021', 0.57), ('01/03/2021', 2.04),
    ('01/04/2021', 0.88), ('03/05/2021', 0.08), ('01/06/2021', 0.55),
    ('01/07/2021', 0.86), ('02/08/2021', 1.80), ('01/09/2021', 2.44),
    ('01/10/2021', 2.72), ('01/11/2021', 5.62), ('01/12/2021', 3.71),
    ('03/01/2022', 1.13), ('01/02/2022', 1.41), ('02/03/2022', 1.21),
    ('01/04/2022', 1.93), ('02/05/2022', 0.58), ('01/06/2022', 0.06),
    ('01/07/2022', 0.16), ('03/10/2022', 0.30), ('01/11/2022', 2.15),
    ('01/12/2022', 3.42), ('02/01/2023', 2.45), ('01/02/2023', 5.76),
    ('01/03/2023', 3.73), ('03/04/2023', 2.66), ('02/05/2023', 2.25),
    ('01/06/2023', 2.84), ('03/07/2023', 2.52), ('01/08/2023', 21.25),
    ('04/09/2023', 8.47), ('02/10/2023', 2.74), ('01/11/2023', 3.03),
    ('01/12/2023', 3.18), ('02/01/2024', 1.77), ('01/02/2024', 2.78),
    ('01/03/2024', 0.90), ('01/04/2024', 11.82), ('02/05/2024', 5.52),
    ('01/07/2024', 0.70), ('01/08/2024', 4.36), ('09/09/2024', 23.74),
    ('01/10/2024', 15.20), ('01/11/2024', 4.18), ('02/12/2024', 10.40),
]
SERVICO = [
    ('03/01/2020', 9.99), ('03/02/2020', 9.99), ('03/03/2020', 9.99),
    ('03/04/2020', 9.99), ('04/05/2020', 9.99), ('03/06/2020', 9.99),
    ('03/07/2020', 9.99), ('03/08/2020', 9.99), ('03/09/2020', 9.99),
    ('05/10/2020', 9.99), ('03/11/2020', 9.99), ('03/12/2020', 9.99),
    ('04/01/2021', 9.99), ('03/02/2021', 9.99), ('03/03/2021', 9.99),
    ('05/04/2021', 9.99), ('03/05/2021', 9.99), ('04/06/2021', 9.99),
    ('05/07/2021', 9.99), ('03/08/2021', 9.99), ('06/09/2021', 9.99),
    ('04/10/2021', 9.99), ('03/11/2021', 9.99), ('03/12/2021', 9.99),
]

LANCAMENTOS = sorted(ENCARGO + SERVICO, key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO',
    'nacionalidade': 'brasileiro',
    'estado_civil': 'casado',
    'profissao': 'aposentado',
    'cpf': '000.000.017-27',
    'rg': '1000015-5',
    'orgao_expedidor_prefixo': 'SSP/AC',
    'logradouro': 'Av. Joaquim Cardoso',
    'numero': '646',
    'bairro': 'Aida Mendonça',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '20304-1'}
renda = {'valor_float': 988.00}

tese = {'rubrica': 'ENC LIM CRÉDITO / SERVIÇO CARTÃO PROTEGIDO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
# Mapeia: encargo = ENCARGOS LIMITE; mora = SERVIÇO CARTÃO PROTEGIDO
dados['rubrica_curta'] = 'Enc Lim Crédito / Serviço Cartão Protegido'
dados['rubrica_curta_caps'] = 'ENC LIM CRÉDITO / SERVIÇO CARTÃO PROTEGIDO'
dados['rubrica_completa'] = 'Encargos Limite de Crédito / Serviço Cartão Protegido'
dados['rubrica_completa_caps'] = 'ENCARGOS LIMITE DE CRÉDITO / SERVIÇO CARTÃO PROTEGIDO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — ENCARGO + SERVIÇO CARTÃO ===')
print(f'Lançamentos: {len(LANCAMENTOS)} ({len(ENCARGO)} encargo + {len(SERVICO)} serviço), total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
def fmt_moeda(v):
    s = f'{v:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.')
    return f'R$ {s}'
te=sum(v for _,v in ENCARGO); ts=sum(v for _,v in SERVICO)

with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {
    '{{remuneração}}': dados['remuneração'],
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
    '{{cidade_filial}}': 'Maués',
    '{{uf_filial}}': 'AM',
    '{{uf_extenso}}': 'Amazonas',
    '{{numero_desconto_encargo}}': str(len(ENCARGO)),
    '{{numero_desconto_mora}}': str(len(SERVICO)),
    '{{total_encargo}}': fmt_moeda(te),
    '{{total_mora}}': fmt_moeda(ts),
    '{{total_encargo_extenso}}': extenso_moeda(te),
    '{{total_mora_extenso}}': extenso_moeda(ts),
    '{{desconto_extenso_encargo}}': extenso_cardinal(len(ENCARGO)),
    '{{desconto_extenso_mora}}': extenso_cardinal(len(SERVICO)),
}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_EncargoServico_LUIZ_PIRES', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'ENCARGOS LIMITE + SERVIÇO CARTÃO PROTEGIDO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_EncargoServico_LUIZ_PIRES_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 06/12/2024)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} ({len(ENCARGO)} encargos + {len(SERVICO)} serviço cartão)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO CRÍTICA — corte 30/03/2021',
     'Aproximadamente 30 lançamentos pré-30/03/2021 podem estar prescritos. '
     'Pós-30/03/2021: 48 lançamentos válidos. DECISÃO DO PROCURADOR.'),
    ('TABELA NÃO TEM MORA CRED PESSOAL',
     'Apesar de a pasta se chamar "ENCARGOS" e a notificação mencionar genericamente '
     'descontos indevidos, a tabela só traz ENCARGOS LIMITE DE CRED + SERVIÇO CARTÃO '
     'PROTEGIDO. Template inicial-mora-encargo.docx adaptado: rubrica "mora" '
     'substituída por "serviço cartão protegido". REVISAR texto antes do protocolo.'),
    ('IDOSO — prioridade aplicada', 'Notificação afirma "pessoa idosa".'),
    ('PROCURAÇÃO ASSINADA A ROGO',
     'Pasta tem RG da rogada + 2 testemunhas (Evaristo e Nuberlândia). Conferir '
     'validade.'),
    ('CLIENTE TEM 4 TESES SEPARADAS',
     'CLIENTE EXEMPLO aparece em TARIFAS, MORA (esta), TÍTULO CAPITALIZAÇÃO e PG ELETRON. '
     'AVALIAR consolidação.'),
    ('TETO JEC — coberto',
     'VC R$ 15.843,58 ≈ 10,4 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'REVISAR adaptação "mora → serviço cartão protegido" no DOCX gerado.',
    'DECIDIR estratégia de prescrição.',
    'AVALIAR consolidação com TARIFAS + TÍTULO + PG ELETRON.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade da procuração assinada a rogo.',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Confirmar com cliente: nunca contratou cheque especial nem seguro de cartão.',
    'Anexar 2-Procurações (ENC LIM + SERVIÇO CARTÃO) + 3-RG + 3.1 a 3.3 (rogada + testemunhas) + 5-Comprovante + 5.1-Declaração + 5.2-RG proprietária + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem REVISAR adaptação do template e decidir prescrição. ')
r2 = p.add_run('Atenção crítica.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
