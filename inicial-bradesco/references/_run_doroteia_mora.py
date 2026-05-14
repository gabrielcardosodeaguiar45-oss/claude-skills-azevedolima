"""Inicial MORA — CLIENTE EXEMPLO VIANA DE VASCONCELOS.

Comarca Caapiranga/AM (Ag 3707 / Conta 59047-9). Pessoa IDOSA.
Aposentada pelo INSS R$ 1.089,91 (último crédito 27/01/2026).

Tabela: 2 lançamentos MORA CREDITO PESSOAL em 19/05/2025 e 20/05/2025.
Total simples R$ 545,15 / dobro R$ 1.090,30. VC R$ 16.090,30 — cabe
folgadamente no JEC.

Caso recente sem questão de prescrição.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\2. MORA\CLIENTE EXEMPLO VIANA DE VASCONCELOS - Ney Pedroza - TARIFAS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Mora_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

LANCAMENTOS = [
    ('19/05/2025', 250.00),
    ('20/05/2025', 295.15),
]

autora = {
    'nome': 'CLIENTE EXEMPLO VIANA DE VASCONCELOS',
    'nacionalidade': 'brasileira',
    'estado_civil': '',
    'profissao': 'aposentada',
    'cpf': '000.000.008-18',
    'rg': '1000006-6',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'CM São Jorge',
    'numero': 's/nº',
    'bairro': 'Centro',
    'cidade': 'Caapiranga',
    'cep': '69.425-000',
}
conta = {'agencia': '3707', 'numero': '59047-9'}
renda = {'valor_float': 1089.91}

tese = {'rubrica': 'MORA CRÉDITO PESSOAL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Caapiranga', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['rubrica_curta'] = 'Mora Cred Pess'
dados['rubrica_curta_caps'] = 'MORA CRED PESS'
dados['rubrica_completa'] = 'Crédito Mora Pessoal'
dados['rubrica_completa_caps'] = 'MORA CREDITO PESSOAL'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — MORA ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
       '{{cidade_filial}}': 'Maués',
       '{{uf_filial}}': 'AM',
       '{{uf_extenso}}': 'Amazonas'}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Mora_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CRÉDITO PESSOAL'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Mora_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 27/01/2026)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)}'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('IDOSA — prioridade aplicada',
     'Notificação afirma "pessoa idosa". Confirmar RG (≥ 60 anos).'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO.'),
    ('Renda < 1 SM — possível consignação',
     'INSS R$ 1.089,91 (27/01/2026) < salário mínimo. Conferir HISCON. Reforça '
     'hipossuficiência.'),
    ('Caso recente — sem prescrição',
     'Lançamentos de 19-20/05/2025 — nenhuma questão prescricional.'),
    ('Apenas 2 lançamentos',
     'A tabela registra apenas 2 lançamentos consecutivos (19 e 20/05/2025), '
     'totalizando R$ 545,15. Pode ser cobrança em duplicidade do mesmo encargo. '
     'Conferir extrato — pode haver mais lançamentos não compilados na tabela.'),
    ('TETO JEC — coberto',
     'VC R$ 16.090,30 ≈ 10,6 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Caapiranga/AM.',
    'Conferir VC = R$ 16.090,30 e dano moral R$ 15.000,00.',
    'Confirmar com cliente: nunca contratou cheque especial nem aceitou encargos de mora.',
    'Verificar se há outros lançamentos MORA no extrato não compilados na tabela.',
    'Anexar 2-Procuração + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após confirmar com cliente, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
