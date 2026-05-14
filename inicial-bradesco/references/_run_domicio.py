"""Inicial TARIFAS — CLIENTE EXEMPLO RENAN MATOS DOS SANTOS.

Comarca Barreirinha/AM. Conta Bradesco Ag 3725/11708-0. Idade não
confirmada. Renda não identificada no extrato — adotar R$ 1.518 (SM
2025) como ESTIMATIVA, alertar pendência.

Tese: TARIFA BANCÁRIA CESTA B.EXPRESSO4 — 24 lançamentos parseados do
extrato Internet Banking entre 26/01/2023 e 15/08/2025 (total R$ 410,57;
dobro R$ 821,14). VC R$ 15.821,14 — cabe no JEC. Valor pequeno mas
suficiente para R$ 15.000 dano moral.
"""
import io, sys, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\CLIENTE EXEMPLO RENAN MATOS DOS SANTOS - Wilson - TARIFAS\TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_lancs.json')) as f:
    LANCS_RAW = json.load(f)
LANCAMENTOS = sorted([(d, v) for d, t, v in LANCS_RAW],
                     key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO RENAN MATOS DOS SANTOS', 'nacionalidade': 'brasileiro',
    'estado_civil': '', 'profissao': 'aposentado',
    'cpf': '000.000.007-17', 'rg': '1000005-5',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua São Benedito', 'numero': 's/nº',
    'bairro': 'Centro, Distrito Pedras', 'cidade': 'Barreirinha', 'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '11708-0'}
renda = {'valor_float': 1518.00}  # ESTIMADO

tese = {'rubrica': 'TARIFA BANCÁRIA CESTA B.EXPRESSO4', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Barreirinha', uf='AM')
dados['remuneração'] = 'aposentadoria/benefício previdenciário'
dados['titulo'] = 'TARIFA BANCÁRIA CESTA B.EXPRESSO4'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO RENAN — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | VC: R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {
    '{{remuneração}}': dados['remuneração'],
    '{{valor_remuneração}}': dados['valor_remuneração'],
    '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

# Relatório
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA CESTA B.EXPRESSO4'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)
doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda usada (ESTIMADA)', f'{dados["valor_remuneração"]} (CONFIRMAR — não identificada no extrato)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos TARIFA', '24 (extrato Internet Banking)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v
doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('RENDA ESTIMADA — confirmar com cliente',
     'Extrato é Bradesco Internet Banking (formato condensado, sem text-layer claro de '
     'créditos). Adotada renda R$ 1.518 (SM 2025) como ESTIMATIVA. Confirmar renda real '
     'antes do protocolo.'),
    ('IDADE — não confirmada',
     'RG não lido. Assume NÃO IDOSO. Conferir.'),
    ('VALOR PEQUENO',
     'Total bruto R$ 410,57 (dobro R$ 821,14) é valor pequeno. O dano moral R$ 15.000 '
     'sustenta o pleito. Verificar se cliente tem outras teses no KIT que possam combinar.'),
    ('TETO JEC — coberto',
     f'VC R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.') + '. Cabe no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)
doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (11708-0 / 3725).',
           'Conferir comarca: Barreirinha/AM.', 'CONFIRMAR RENDA real.',
           'Confirmar idade no RG.',
           'Anexar 8-NOTIFICAÇÃO + 8.1-COMPROVANTE NOTIFICAÇÃO (já estão na pasta).']:
    doc_r.add_paragraph(it, style='List Number')
p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas — '); r2 = p.add_run('PROTOCOLAR após confirmar renda e idade.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
