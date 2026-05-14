"""Inicial TÍTULO DE CAPITALIZAÇÃO — CLIENTE EXEMPLO."""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\3. Título de capitalização\CLIENTE EXEMPLO - Ruth - TARIFA\TÍTULO DE CAPITALIZAÇÃO'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Titulo_LUIZ_PIRES_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_LUIZ_PIRES_TITULO_v1.docx')

LANCAMENTOS = [
    ('10/08/2020', 20.00), ('08/09/2020', 20.00), ('07/10/2020', 20.00),
    ('09/11/2020', 20.00), ('07/12/2020', 20.00), ('07/01/2021', 20.00),
    ('08/02/2021', 20.00), ('08/03/2021', 20.00), ('07/04/2021', 20.00),
    ('07/05/2021', 20.00), ('07/06/2021', 20.00), ('07/07/2021', 20.00),
    ('09/08/2021', 21.35), ('08/09/2021', 21.35), ('07/10/2021', 21.35),
    ('08/11/2021', 21.35), ('07/12/2021', 21.35), ('07/01/2022', 21.35),
    ('07/02/2022', 21.35), ('07/03/2022', 21.35), ('07/04/2022', 21.35),
    ('09/05/2022', 21.35), ('07/06/2022', 21.35), ('07/07/2022', 21.35),
    ('08/08/2022', 23.99), ('08/09/2022', 23.99), ('12/09/2022', 20.00),
    ('07/10/2022', 23.99), ('10/10/2022', 20.00), ('07/11/2022', 23.99),
    ('09/11/2022', 20.00), ('07/12/2022', 23.99), ('09/12/2022', 20.00),
    ('09/01/2023', 20.00), ('09/01/2023', 23.99), ('07/02/2023', 23.99),
    ('09/02/2023', 20.00), ('07/03/2023', 23.99), ('09/03/2023', 20.00),
    ('10/04/2023', 20.00), ('10/04/2023', 23.99), ('08/05/2023', 23.99),
    ('09/05/2023', 20.00), ('07/06/2023', 23.99), ('09/06/2023', 20.00),
    ('07/07/2023', 23.99), ('10/07/2023', 20.00), ('07/08/2023', 23.99),
    ('09/08/2023', 20.00), ('08/09/2023', 23.99), ('11/09/2023', 20.70),
    ('09/10/2023', 20.70), ('09/10/2023', 23.99), ('07/11/2023', 23.99),
    ('09/11/2023', 20.70), ('07/12/2023', 23.99), ('11/12/2023', 20.70),
    ('08/01/2024', 23.99), ('09/01/2024', 20.70), ('07/02/2024', 23.99),
    ('09/02/2024', 20.70), ('07/03/2024', 23.99), ('11/03/2024', 20.70),
    ('08/04/2024', 23.99), ('09/04/2024', 20.70), ('07/05/2024', 23.99),
    ('09/05/2024', 20.70), ('07/06/2024', 23.99), ('10/06/2024', 20.70),
    ('08/07/2024', 23.99), ('09/07/2024', 20.70), ('07/08/2024', 24.91),
    ('09/09/2024', 20.70), ('09/09/2024', 21.63), ('09/09/2024', 24.91),
    ('07/10/2024', 24.91), ('09/10/2024', 21.63), ('07/11/2024', 24.91),
    ('11/11/2024', 21.63), ('09/12/2024', 21.63), ('09/12/2024', 24.91),
]

autora = {
    'nome': 'CLIENTE EXEMPLO', 'nacionalidade': 'brasileiro',
    'estado_civil': 'casado', 'profissao': 'aposentado',
    'cpf': '000.000.017-27', 'rg': '1000015-5', 'orgao_expedidor_prefixo': 'SSP/AC',
    'logradouro': 'Av. Joaquim Cardoso', 'numero': '646',
    'bairro': 'Aida Mendonça', 'cidade': 'Presidente Figueiredo', 'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '20304-1'}
renda = {'valor_float': 988.00}
tese = {'rubrica': 'TÍTULO DE CAPITALIZAÇÃO', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TÍTULO DE CAPITALIZAÇÃO'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — TÍTULO ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Titulo_LUIZ_PIRES', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TÍTULO DE CAPITALIZAÇÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Titulo_LUIZ_PIRES_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]}'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (vários títulos paralelos)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('TEMPLATE GENÉRICO ADAPTADO', 'Sem template próprio para TÍTULO. Usado inicial-tarifas. Ajustar fundamentação.'),
    ('VÁRIOS TÍTULOS PARALELOS', 'A partir de 09/01/2023 a tabela mostra 2-3 títulos descontados no mesmo dia (R$ 20 + R$ 23,99 + R$ 20,70 + R$ 21,63 + R$ 24,91). Indica vários títulos vendidos sem autorização.'),
    ('PRESCRIÇÃO — corte 30/03/2021', '~10 lançamentos pré-30/03/2021 podem estar prescritos. Pós: 71 válidos.'),
    ('IDOSO — prioridade aplicada', 'Notificação afirma "pessoa idosa".'),
    ('PROCURAÇÃO ASSINADA A ROGO', 'RG da rogada + 2 testemunhas (Evaristo + Nuberlândia).'),
    ('CLIENTE TEM 4 TESES', 'CLIENTE EXEMPLO aparece em TARIFAS, MORA, TÍTULO (esta), PG ELETRON. AVALIAR consolidação.'),
    ('TETO JEC — coberto', 'VC R$ 18.542,86 ≈ 12,2 SM.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'AJUSTAR fundamentação para TÍTULO.',
    'AVALIAR consolidação com outras 3 teses.',
    'DECIDIR estratégia de prescrição.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade da procuração assinada a rogo.',
    'Confirmar com cliente: nunca contratou múltiplos títulos de capitalização.',
    'Anexar 2-Procuração + 3-RG + 3.1 a 3.3 + 5-Comprovante + 5.1-Declaração + 5.2-RG proprietária + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem ajustes. ')
r2 = p.add_run('Atenção.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
