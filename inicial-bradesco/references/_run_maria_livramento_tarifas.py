"""Inicial TARIFAS — MARIA DO LIVRAMENTO LIMA DOS SANTOS.

Comarca Presidente Figueiredo/AM (Ag 3732 / Conta 21528-7). Pessoa
IDOSA (notificação afirma "pessoa idosa" mais "aposentada"). Recebe
INSS R$ 986,58 (último crédito 03/02/2026 — líquido, abaixo de 1 SM).

Endereço idêntico ao EXEMPLO MANUEL DOS SANTOS (Ramal do Rumo Certo,
BR 174 - KM 165, Comunidade Boa União) — provavelmente cônjuges.

Tabela: 44 lançamentos TARIFA BANCARIA - CESTA BENEFIC entre 12/11/2021
e 03/11/2023. Total R$ 489,30 / dobro R$ 978,60. VC R$ 15.978,60 —
cabe folgadamente no JEC.

PENDÊNCIAS: cliente também tem TÍTULO CAPITALIZAÇÃO e PG ELETRON em
pastas separadas (3 teses no total).
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\MARIA DO LIVRAMENTO LIMA DOS SANTOS - Ruth\TARIFA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_MARIA_LIVRAMENTO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MARIA_LIVRAMENTO_v1.docx')

LANCAMENTOS = [
    ('12/11/2021', 4.42), ('02/12/2021', 15.18), ('15/12/2021', 8.87),
    ('04/01/2022', 10.73), ('14/01/2022', 1.33), ('02/02/2022', 18.27),
    ('04/03/2022', 19.60), ('15/03/2022', 4.45), ('04/04/2022', 15.15),
    ('14/04/2022', 8.90), ('03/05/2022', 11.20), ('13/05/2022', 19.81),
    ('02/06/2022', 0.29), ('15/06/2022', 20.10), ('15/07/2022', 4.48),
    ('02/08/2022', 15.62), ('15/08/2022', 8.79), ('06/09/2022', 11.31),
    ('15/09/2022', 3.09), ('04/10/2022', 17.01), ('14/10/2022', 7.41),
    ('03/11/2022', 13.04), ('14/11/2022', 1.37), ('02/12/2022', 19.08),
    ('15/12/2022', 3.43), ('03/01/2023', 17.02), ('13/01/2023', 0.46),
    ('02/02/2023', 19.99), ('15/02/2023', 5.50), ('02/03/2023', 14.95),
    ('15/03/2023', 11.92), ('04/04/2023', 8.53), ('14/04/2023', 3.96),
    ('03/05/2023', 17.14), ('15/05/2023', 2.56), ('02/06/2023', 18.54),
    ('15/06/2023', 2.35), ('30/06/2023', 18.75), ('14/07/2023', 21.10),
    ('15/08/2023', 0.01), ('17/08/2023', 21.09), ('15/09/2023', 21.10),
    ('13/10/2023', 10.85), ('03/11/2023', 10.55),
]

autora = {
    'nome': 'MARIA DO LIVRAMENTO LIMA DOS SANTOS',
    'nacionalidade': 'brasileira',
    'estado_civil': '',
    'profissao': 'aposentada',
    'cpf': '000.000.021-31',
    'rg': '1000019-9',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Ramal do Rumo Certo, BR 174 - KM 165',
    'numero': 's/nº',
    'bairro': 'Comunidade Boa União',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '21528-7'}
renda = {'valor_float': 986.58}

tese = {'rubrica': 'TARIFA BANCÁRIA - CESTA BENEFIC', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Presidente Figueiredo', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'TARIFA BANCÁRIA - CESTA BENEFIC'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== MARIA DO LIVRAMENTO — TARIFAS ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_MARIA_LIVRAMENTO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFA BANCÁRIA - CESTA BENEFIC'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_MARIA_LIVRAMENTO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSA — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 03/02/2026 — líquido)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO — corte 30/03/2021 (EAREsp 1.280.825 STJ)',
     'Tabela inicia 12/11/2021 — todos os 44 lançamentos pós 30/03/2021. SEM '
     'prescrição relevante.'),
    ('IDOSA — prioridade aplicada',
     'Notificação afirma "pessoa idosa" e "aposentada". Confirmar RG (≥ 60 anos).'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO. Endereço idêntico ao '
     'EXEMPLO MANUEL — provavelmente cônjuges.'),
    ('Renda LÍQUIDA < 1 SM — possível consignação',
     'INSS R$ 986,58 (03/02/2026) < salário mínimo. Indica consignados pelo próprio '
     'INSS. Renda BRUTA pode ser maior. Conferir HISCON.'),
    ('CLIENTE TEM 3 TESES SEPARADAS',
     'Cliente também aparece nas pastas: TÍTULO CAPITALIZAÇÃO (3) e PG ELETRON (4). '
     'AVALIAR consolidação em inicial-combinada (Presidente Figueiredo adota '
     'combinação por padrão).'),
    ('VALOR BAIXO',
     'Dobro R$ 978,60 está bem abaixo do limite R$ 400 da skill — combinação '
     'recomendada com TÍTULO + PG ELETRON.'),
    ('CASAL CO-DEMANDANTE',
     'Endereço da MARIA é IDÊNTICO ao do EXEMPLO MANUEL DOS SANTOS (Ramal do Rumo '
     'Certo, BR 174 - KM 165, Comunidade Boa União). Provavelmente cônjuges. '
     'Avaliar se é caso de litisconsórcio ativo facultativo (art. 113 CPC) com '
     'mesma autora coletiva.'),
    ('TETO JEC — coberto',
     'VC R$ 15.978,60 ≈ 10,5 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AVALIAR consolidação com TÍTULO CAPITALIZAÇÃO + PG ELETRON.',
    'AVALIAR litisconsórcio ativo facultativo com EXEMPLO MANUEL (mesmo endereço).',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Confirmar com cliente: nunca contratou cesta benefic.',
    'Anexar 2-Procuração + 3-RG + 4-Hipossuficiência + 5-Comprovante + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA — após decidir consolidação, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
