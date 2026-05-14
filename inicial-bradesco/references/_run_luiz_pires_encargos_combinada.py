"""Inicial COMBINADA TARIFAS+MORA — CLIENTE EXEMPLO (subpasta ENCARGOS).

2 procurações distintas:
  1. ENCARGOS LIMITE DE CRED (família MORA)
  2. SERVIÇO CARTÃO PROTEGIDO (família TARIFAS — seguro de cartão)

Comarca Pres. Figueiredo/AM, IDOSO 72 anos (03/10/1953), casado,
RG SSP/AC, procuração a rogo. Ag 3732 / Conta 20304-1, INSS R$ 988.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _combinada_helper import gerar_combinada
from docx import Document

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\2. MORA\CLIENTE EXEMPLO - Ruth - TARIFA\ENCARGOS'

ENCARGO = [
    ('02/01/2020', 1.46), ('03/02/2020', 0.66), ('02/03/2020', 0.21),
    ('01/06/2020', 0.08), ('03/08/2020', 0.08), ('01/09/2020', 0.12),
    ('01/10/2020', 1.20), ('03/11/2020', 0.87), ('01/12/2020', 0.28),
    ('04/01/2021', 1.20), ('01/02/2021', 0.57), ('01/03/2021', 2.04),
    ('01/04/2021', 0.88), ('03/05/2021', 0.08), ('01/06/2021', 0.55),
    ('01/07/2021', 0.86), ('02/08/2021', 1.80), ('01/09/2021', 2.44),
    ('01/10/2021', 2.72), ('01/11/2021', 5.62), ('01/12/2021', 3.71),
    ('03/01/2022', 1.13), ('01/02/2022', 1.41), ('02/03/2022', 1.21),
    ('01/04/2022', 1.93), ('02/05/2022', 0.58), ('01/06/2022', 0.06),
    ('01/07/2022', 0.16), ('03/10/2022', 0.30), ('01/11/2022', 2.15),
    ('01/12/2022', 3.42), ('02/01/2023', 2.45), ('01/02/2023', 5.76),
    ('01/03/2023', 3.73), ('03/04/2023', 2.66), ('02/05/2023', 2.25),
    ('01/06/2023', 2.84), ('03/07/2023', 2.52), ('01/08/2023', 21.25),
    ('04/09/2023', 8.47), ('02/10/2023', 2.74), ('01/11/2023', 3.03),
    ('01/12/2023', 3.18), ('02/01/2024', 1.77), ('01/02/2024', 2.78),
    ('01/03/2024', 0.90), ('01/04/2024', 11.82), ('02/05/2024', 5.52),
    ('01/07/2024', 0.70), ('01/08/2024', 4.36), ('09/09/2024', 23.74),
    ('01/10/2024', 15.20), ('01/11/2024', 4.18), ('02/12/2024', 10.40),
]
SERVICO = [
    ('03/01/2020', 9.99), ('03/02/2020', 9.99), ('03/03/2020', 9.99),
    ('03/04/2020', 9.99), ('04/05/2020', 9.99), ('03/06/2020', 9.99),
    ('03/07/2020', 9.99), ('03/08/2020', 9.99), ('03/09/2020', 9.99),
    ('05/10/2020', 9.99), ('03/11/2020', 9.99), ('03/12/2020', 9.99),
    ('04/01/2021', 9.99), ('03/02/2021', 9.99), ('03/03/2021', 9.99),
    ('05/04/2021', 9.99), ('03/05/2021', 9.99), ('04/06/2021', 9.99),
    ('05/07/2021', 9.99), ('03/08/2021', 9.99), ('06/09/2021', 9.99),
    ('04/10/2021', 9.99), ('03/11/2021', 9.99), ('03/12/2021', 9.99),
]

teses = [
    {'familia': 'TARIFAS', 'rubrica': 'SERVIÇO CARTÃO PROTEGIDO', 'lancamentos': SERVICO},
    {'familia': 'MORA',    'rubrica': 'ENCARGOS LIMITE DE CRÉDITO', 'lancamentos': ENCARGO},
]

autora = {
    'nome': 'CLIENTE EXEMPLO', 'nacionalidade': 'brasileiro',
    'estado_civil': 'casado', 'profissao': 'aposentado',
    'cpf': '000.000.017-27', 'rg': '1000015-5', 'orgao_expedidor_prefixo': 'SSP/AC',
    'logradouro': 'Av. Joaquim Cardoso', 'numero': '646',
    'bairro': 'Aida Mendonça', 'cidade': 'Presidente Figueiredo', 'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '20304-1'}
renda = {'valor_float': 988.00, 'descricao': 'aposentadoria pelo INSS'}

res = gerar_combinada(
    pasta=PASTA, nome_arquivo_base='INICIAL_Combinada_LUIZ_PIRES_Encargos',
    autora=autora, conta=conta, renda=renda, teses=teses,
    comarca='Presidente Figueiredo', uf='AM', eh_idoso=True,
)
print(f'CLIENTE EXEMPLO ENCARGOS — combinada {res["totais"]["n_teses"]} teses (TARIFAS+MORA)')
print(f'  Total: R$ {res["totais"]["total_geral"]:.2f} / dobro R$ {res["totais"]["dobro_geral"]:.2f}')
print(f'  Dano moral: R$ {res["totais"]["dano_moral_total"]:.2f}')
print(f'  VC: R$ {res["totais"]["valor_causa"]:.2f}')
print(f'  Parágrafos removidos: {res["paras_removidos"]}, pós-fix: {res["pos_fix"]}')

doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Combinada_LUIZ_PIRES_Encargos', level=1)
for k, v in [('Cliente', autora['nome']),
             ('Tese', 'COMBINADA TARIFAS (Serviço Cartão) + MORA (Encargos)'),
             ('Comarca', 'Presidente Figueiredo/AM'),
             ('Arquivo', 'INICIAL_Combinada_LUIZ_PIRES_Encargos_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
linhas = [
    ('Comarca', 'Presidente Figueiredo/AM'),
    ('Prioridade', 'IDOSO 72 anos (03/10/1953)'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]} — naturalidade Cruzeiro do Sul/AC'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', 'R$ 988,00 (INSS último 06/12/2024 — sem créditos em 2025/2026)'),
]
for i, t in enumerate(teses, 1):
    linhas.append((f'Tese {i} ({t["familia"]}) — {t["rubrica"]}',
                   f'{t["n_lanc"]} lanç. {t["inicio"]}–{t["fim"]} = R$ {t["total"]:.2f} / R$ {t["dobro"]:.2f} dobro'))
linhas.extend([
    ('TOTAL combinado', f'R$ {res["totais"]["total_geral"]:.2f} / R$ {res["totais"]["dobro_geral"]:.2f} dobro'),
    ('Dano moral (2 teses x R$ 5.000)', f'R$ {res["totais"]["dano_moral_total"]:.2f}'),
    ('Valor da causa', f'R$ {res["totais"]["valor_causa"]:.2f}'),
])
for k, v in linhas:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS', level=2)
for titulo, txt in [
    ('Inicial COMBINADA — 2 famílias (TARIFAS + MORA) ativas',
     '2 procurações distintas. SERVIÇO CARTÃO PROTEGIDO ativa o bloco TARIFAS '
     '(seguro de cartão é tarifa); ENCARGOS LIMITE DE CRED ativa o bloco MORA. '
     'Inicial tem 2 núcleos fáticos individualizados, 2 blocos doutrinários, '
     '2 pedidos de restituição.'),
    ('PRESCRIÇÃO CRÍTICA',
     'Tabela inicia 02/01/2020. ~30 lançamentos pré-30/03/2021 (EAREsp 1.280.825 STJ) '
     'podem estar prescritos pelo CDC. Estratégia decenal (art. 205 CC) acolhida em '
     'algumas câmaras do TJ-AM permite pleitear tudo. REVISAR.'),
    ('IDOSO 72 anos confirmado',
     'RG 0984937-8 SSP/AC, nascimento 03/10/1953, naturalidade Cruzeiro do Sul/AC.'),
    ('PROCURAÇÃO A ROGO',
     'RG da rogada + 2 testemunhas (Evaristo + Nuberlândia). Conferir validade.'),
    ('CLIENTE TEM 4 TESES TOTAIS',
     'CLIENTE EXEMPLO também tem TARIFAS (CESTA BENEFIC), TÍTULO e PG ELETRON em pastas '
     'separadas. Esta combinada cobre apenas as 2 da subpasta ENCARGOS. AVALIAR '
     'mega-combinada com TARIFAS+TÍTULO; PG ELETRON segue separada (terceiro solidário).'),
    ('TETO JEC — coberto', f'VC R$ {res["totais"]["valor_causa"]:.2f}.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'CONFERIR visualmente combinada (2 núcleos: TARIFAS+MORA; sem TÍTULO/APLIC).',
    'DECIDIR estratégia de prescrição.',
    'Conferir nome / CPF / RG / nascimento.',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir validade procuração a rogo.',
    f'Conferir VC R$ {res["totais"]["valor_causa"]:.2f}.',
    'Confirmar com cliente: nunca contratou seguro de cartão nem cheque especial.',
    'Anexar 2 procurações + 3-RG + 3.1-RG rogada + 3.2/3.3-RGs testemunhas + 5-Comprovante + 5.1-Declaração + 5.2-RG proprietária + 6-Extrato + 7-Tabela + 8-Notificação + 8.1-Comprovante.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir prescrição. Após decisão, ')
r = p.add_run('PROTOCOLAR.'); r.bold = True
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_LUIZ_PIRES_ENCARGOS_COMBINADA_v1.docx')
doc_r.save(RELAT_OUT)
print(f'  -> {RELAT_OUT}')
