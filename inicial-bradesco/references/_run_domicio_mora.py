"""Inicial MORA — CLIENTE EXEMPLO RENAN MATOS DOS SANTOS.

Comarca Barreirinha/AM. Conta Bradesco Ag 3725/11708-0. Renda R$ 1.518
(SM 2025) ESTIMADA — não identificada no extrato Internet Banking. Idade
não confirmada.

Tese: MORA CRED PESS — APENAS 3 lançamentos detectados em 2023 (R$ 84,43
total; dobro R$ 168,86). Dano moral R$ 15.000. VC R$ 15.168,86 — cabe no
JEC com valor próximo do piso.

PENDÊNCIAS CRÍTICAS:
1. VALOR MÍNIMO — apenas 3 lançamentos identificados. Extrato Internet
   Banking tem cobertura limitada. Cliente pode ter mais MORAs em períodos
   não baixados. Solicitar extrato completo de 2020-2025.
2. RUBRICA "Mora Operacao" detectada em 2024 (sem detalhamento). Pode ser
   tese adicional — investigar.
3. RENDA ESTIMADA R$ 1.518 — confirmar antes do protocolo.
4. IDADE não confirmada — assume NÃO IDOSO.
"""
import io, sys, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\1. TESTE 2\CLIENTE EXEMPLO RENAN MATOS DOS SANTOS - Wilson - TARIFAS\MORA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Mora_CLIENTE EXEMPLO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_CLIENTE EXEMPLO_MORA_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_cliente exemplo_mora_lancs.json')) as f:
    data = json.load(f)
LANCAMENTOS = sorted(data['todos'], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO RENAN MATOS DOS SANTOS', 'nacionalidade': 'brasileiro',
    'estado_civil': '', 'profissao': 'aposentado',
    'cpf': '000.000.007-17', 'rg': '1000005-5',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua São Benedito', 'numero': 's/nº',
    'bairro': 'Centro, Distrito Pedras', 'cidade': 'Barreirinha', 'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '11708-0'}
renda = {'valor_float': 1518.00}  # ESTIMADO

tese = {'rubrica': 'MORA CRED PESS', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Barreirinha', uf='AM')
dados['remuneração'] = 'aposentadoria/benefício previdenciário'
dados['rubrica_curta'] = 'Mora Cred Pess'
dados['rubrica_curta_caps'] = 'MORA CRED PESS'
dados['rubrica_completa'] = 'Mora Crédito Pessoal'
dados['rubrica_completa_caps'] = 'MORA CRÉDITO PESSOAL'
dados['titulo'] = 'MORA CRÉDITO PESSOAL'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — MORA ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | dobro: R$ {calc["dobro"]:,.2f} | VC: R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
       '{{cidade_filial}}': 'Barreirinha',
       '{{uf_filial}}': 'AM',
       '{{uf_extenso}}': 'Amazonas'}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

# Relatório
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Mora_CLIENTE EXEMPLO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CRED PESS'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Mora_CLIENTE EXEMPLO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda usada (ESTIMADA)', f'{dados["valor_remuneração"]} (CONFIRMAR)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos MORA', '3 (extrato Internet Banking limitado)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('VALOR MÍNIMO — solicitar extrato completo',
     'Apenas 3 lançamentos MORA CRED PESS detectados (R$ 84,43 total) entre '
     '25/01/2023 e 26/01/2023. O extrato Internet Banking traz apenas anos isolados '
     '(2021, 2023, 2024) com cobertura limitada. CONFIRMAR com cliente se houve mais '
     'MORAs em outros períodos. Solicitar extrato completo Bradesco Celular '
     '(modelo similar ao da CLIENTE EXEMPLO) para garantir detecção de TODOS os MORAs.'),
    ('RUBRICA "Mora Operacao" em 2024',
     'Página 3 do extrato (2024) lista no topo "Mora Operacao" como rubrica '
     'presente, mas sem detalhamento de lançamentos. Pode ser tese adicional. '
     'Investigar antes do protocolo.'),
    ('RENDA ESTIMADA — confirmar com cliente',
     'Adotada renda R$ 1.518 (SM 2025) como ESTIMATIVA. Extrato Internet Banking '
     'não traz CRÉDITO DE SALARIO/INSS identificável. Confirmar renda real.'),
    ('IDADE — não confirmada',
     'RG não lido. Assume NÃO IDOSO. Conferir.'),
    ('TETO JEC — coberto',
     f'VC R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.') + ' ≈ 9,99 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (11708-0 / 3725).',
           'Conferir comarca: Barreirinha/AM.',
           'SOLICITAR extrato completo Bradesco Celular para cobertura plena.',
           'Investigar tese "Mora Operacao" 2024.',
           'CONFIRMAR renda real.',
           'Confirmar idade no RG.',
           'Anexar 8-NOTIFICAÇÃO + 8.1-COMPROVANTE NOTIFICAÇÃO.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA COM RESSALVA VALOR MÍNIMO — '); r2 = p.add_run('PROTOCOLAR após solicitar extrato completo, confirmar renda e idade.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
