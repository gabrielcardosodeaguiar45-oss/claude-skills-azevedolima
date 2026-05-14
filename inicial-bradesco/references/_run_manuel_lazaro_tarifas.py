"""Inicial TARIFAS — EXEMPLO MANUEL CORDOVIL.

Comarca Barreirinha/AM (endereço autora) — agência 3725 (Barreirinha).
Pessoa IDOSA. Aposentado pelo INSS (R$ 846,22 último crédito 02/07/2025
— valor reduzido sugere consignações).

A subpasta da pasta TARIFAS chama-se "TARIFA E MORA" e contém DUAS
tabelas separadas. Esta inicial cobre APENAS o bloco TARIFAS (CARTÃO
CRÉDITO ANUIDADE — 15 lançamentos). A inicial MORA (SERVIÇO CARTÃO
PROTEGIDO + MORA CRED PESSOAL) é gerada por outro script
(_run_manuel_lazaro_mora.py).

Tabela TARIFAS: 15 lançamentos CARTAO CREDITO ANUIDADE entre 05/05/2022
e 03/01/2023. Total R$ 154,00 / dobro R$ 308,00. Valor isolado é baixo
— procurador pode optar por consolidar com MORA em inicial-combinada.

VC R$ 15.308,00 — cabe folgadamente no JEC.
"""
import io, sys, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao
from helpers_docx import aplicar_template
from extenso import extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0.0 - Procução de iniciais\1. TARIFAS\EXEMPLO MANUEL CORDOVIL - Wilson - TARIFA\TARIFA E MORA'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-tarifas.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Tarifas_MANUEL_LAZARO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_MANUEL_LAZARO_TARIFAS_v1.docx')

LANCAMENTOS = [
    ('05/05/2022', 2.54), ('02/06/2022', 16.71), ('06/06/2022', 19.25),
    ('05/07/2022', 4.32), ('02/08/2022', 14.93), ('05/08/2022', 0.48),
    ('02/09/2022', 18.77), ('06/09/2022', 7.35), ('04/10/2022', 11.90),
    ('05/10/2022', 5.65), ('03/11/2022', 13.60), ('07/11/2022', 7.12),
    ('02/12/2022', 12.13), ('05/12/2022', 4.35), ('03/01/2023', 14.90),
]

autora = {
    'nome': 'EXEMPLO MANUEL CORDOVIL',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'aposentado',
    'cpf': '000.000.018-28',
    'rg': '1000016-6',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Pimentel Tavares',
    'numero': '341',
    'bairro': 'CM Terra P do Limão',
    'cidade': 'Barreirinha',
    'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '2782-0'}
renda = {'valor_float': 846.22}

tese = {'rubrica': 'CARTÃO CRÉDITO ANUIDADE', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=True, competência='Barreirinha', uf='AM')
dados['remuneração'] = 'aposentadoria pelo INSS'
dados['titulo'] = 'CARTÃO CRÉDITO ANUIDADE'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== EXEMPLO MANUEL — TARIFAS (CARTAO ANUIDADE) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:.2f} / dobro: R$ {calc["dobro"]:.2f} / VC: R$ {calc["valor_causa"]:.2f}')

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Tarifas_MANUEL_LAZARO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'TARIFAS — CARTÃO CRÉDITO ANUIDADE'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Tarifas_MANUEL_LAZARO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Prioridade', 'IDOSO — art. 1.048, I, CPC'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda', f'{dados["valor_remuneração"]} (INSS último crédito 02/07/2025 — líquido)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos', f'{len(LANCAMENTOS)} (1 rubrica)'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS / CAMPOS A CONFERIR', level=2)
for titulo, txt in [
    ('CLIENTE TEM 3 TESES NA SUBPASTA "TARIFA E MORA"',
     'A subpasta "TARIFA E MORA" contém 2 tabelas: (i) CARTÃO ANUIDADE (esta inicial); '
     '(ii) MORA CRED + SERVIÇO CARTÃO PROTEGIDO (gerada em INICIAL_Mora_MANUEL_LAZARO). '
     'Cliente também tem PG ELETRON na pasta separada "4. Pagamento Eletrônico de '
     'Cobrança". TOTAL: 4 rubricas em 3 famílias de tese (TARIFAS + MORA + PG ELETRON). '
     'AVALIAR consolidação em inicial-combinada.docx (Barreirinha não está na lista '
     'de comarcas que combinam por padrão, mas valor baixo da TARIFAS justifica). '
     'Inicial gerada como ISOLADA TARIFAS — PROCURADOR DECIDE.'),
    ('IDOSO — prioridade aplicada',
     'Notificação afirma "pessoa idosa". Confirmar RG (≥ 60 anos).'),
    ('VALOR BAIXO — combinação aconselhada',
     'Dobro CARTAO ANUIDADE = R$ 308,00 (abaixo do limite R$ 400 da skill). Skill '
     'recomenda combinar com outras teses. Consolidação com MORA traria R$ 308 + '
     'R$ 5.032,44 = R$ 5.340 dobro + R$ 10k dano moral = R$ 15.340 VC.'),
    ('Estado civil — não informado',
     'Notificação não traz estado civil. Placeholder OMITIDO limpamente.'),
    ('Renda < 1 SM — possível consignação',
     'INSS R$ 846,22 (02/07/2025) é abaixo do salário mínimo. Indica descontos '
     'consignados pelo próprio INSS. Conferir HISCON.'),
    ('NOTIFICAÇÃO previa dano moral R$ 15.000',
     'A notificação extrajudicial pleiteia R$ 15.000 dano moral total (cobrindo as 3 '
     'rubricas: CARTAO ANUIDADE + SERVIÇO CARTÃO PROTEGIDO + MORA CRED PESS). Como '
     'esta inicial é ISOLADA TARIFAS, pleiteia R$ 15.000 conforme regra § 9 da skill. '
     'Se for consolidada com MORA, a regra § 9 diz R$ 5.000 × 2 = R$ 10.000.'),
    ('NOTIFICAÇÃO MENCIONA "Maués/AM"',
     'A notificação foi escrita no escritório de Maués/AM mas o ENDEREÇO da autora é '
     'Barreirinha/AM. A inicial usa Barreirinha (foro do consumidor — art. 101, I, '
     'CDC). Verificar.'),
    ('TETO JEC — coberto',
     'VC R$ 15.308,00 ≈ 10,1 SM. Cabe folgadamente no JEC.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST PRÉ-PROTOCOLO', level=2)
for it in [
    'AVALIAR consolidação com MORA + PG ELETRON em 1 só inicial-combinada.',
    'Conferir nome / CPF / RG / nascimento (≥ 60 anos).',
    f'Conferir Conta/Agência ({conta["numero"]} / {conta["agencia"]}).',
    'Conferir comarca: Barreirinha/AM.',
    'Confirmar com cliente: nunca contratou cartão de crédito Bradesco.',
    'Anexar 2-Procuração específica (CARTAO ANUIDADE) + 3-RG + 4-Hipossuficiência + 5-Comprovante + 5.1-Declaração domicílio + 6-Extrato + 7-Tabela CARTAO ANUIDADE + 8-Notificação + 8.1-Comprovante notificação.',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('NÃO PROTOCOLAR sem decidir consolidação. Após resolução, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
