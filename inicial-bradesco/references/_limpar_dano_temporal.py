"""Remove menções a 'dano temporal' / 'danos temporais' / 'morais e temporais'
dos 3 templates da pasta IniciaisNaoContratado/_templates/.

Aplica:
- TODOS os 3 templates: substituições simples no texto comum (ementa, síntese,
  manifestação volitiva, juros, pedidos)
- APENAS no REFIN: remoção de parágrafos inteiros (subtítulo p191, fundamentação
  p193-p206, trecho específico do p245)
"""
import io, sys, os, unicodedata
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from helpers_docx import substituir_in_run

def nfd(s):
    return unicodedata.normalize('NFD', s)


# === SUBSTITUIÇÕES COMUNS (aplica nos 3 templates) ===
SUBS_COMUNS = [
    # Ementa
    ('DANOS MORAIS E TEMPORAIS CONFIGURADOS', 'DANOS MORAIS CONFIGURADOS'),
    # Síntese fática + Manifestação volitiva + Pedidos juros
    ('reparação por danos morais e temporais', 'reparação por danos morais'),
    ('compensação por danos morais e temporais', 'compensação por danos morais'),
    ('sobre os danos morais e temporais', 'sobre os danos morais'),
    # Juros e correção
    ('desde o dano moral e temporal', 'desde o dano moral'),
    # Variantes da ementa segundo bloco (precisa olhar o p5 inteiro)
    ('danos morais e temporais', 'danos morais'),
]


# === REFIN — REMOÇÃO DE PARÁGRAFOS COMPLETOS ===
# Índices 0-based, calculados depois das substituições comuns
REMOVER_PARAGRAFOS_REFIN = [
    191,  # subtítulo "Dos danos temporais: a Teoria do Desvio Produtivo..."
    192,  # subtítulo continuação? (verificar)
    193, 194, 195, 196, 197, 198, 199, 200, 201, 202, 203, 204, 205, 206,
    # bloco fundamentação completo (Teoria do Desvio Produtivo) + pedido R$5k
]

# === REFIN — substituição no p245 (CONDENAR) — remover trecho dano temporal ===
SUBS_REFIN_245 = {
    ' e de R$ {{dano_temporal_total}} ({{dano_temporal_total_extenso}}) à título de indenização por danos temporais':
        '',
}


def aplicar_subs_paragrafo(p, subs_list):
    """Aplica substituições no parágrafo (lista de tuplas (antigo, novo))."""
    n = 0
    for antigo, novo in subs_list:
        if antigo in p.text or nfd(antigo) in p.text:
            chave = antigo if antigo in p.text else nfd(antigo)
            if substituir_in_run(p._element, {chave: novo}):
                n += 1
    return n


def remover_paragrafo(doc, idx):
    """Remove o parágrafo de índice idx do documento."""
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)


def processar(template_path, eh_refin=False):
    print(f'\n████ {os.path.basename(template_path)} ████')
    doc = Document(template_path)

    # 1. Aplicar substituições comuns
    total_subs = 0
    for i, p in enumerate(doc.paragraphs):
        n = aplicar_subs_paragrafo(p, SUBS_COMUNS)
        if n > 0:
            total_subs += n
    print(f'  Substituições comuns aplicadas: {total_subs}')

    if eh_refin:
        # 2. Substituir trecho do p245 (antes de remover blocos para não perder índice)
        # Precisamos fazer ANTES da remoção dos parágrafos (para manter índice)
        for antigo, novo in SUBS_REFIN_245.items():
            ok = substituir_in_run(doc.paragraphs[245]._element, {antigo: novo})
            if ok:
                print(f'  Trecho dano temporal removido do p245')

        # 3. Remover parágrafos do bloco "Dos danos temporais"
        # IMPORTANTE: remover de trás para frente para não bagunçar índices
        removidos = 0
        for idx in sorted(REMOVER_PARAGRAFOS_REFIN, reverse=True):
            if idx < len(doc.paragraphs):
                remover_paragrafo(doc, idx)
                removidos += 1
        print(f'  Parágrafos removidos: {removidos}')

    doc.save(template_path)
    print(f'  Salvo: {template_path}')


if __name__ == '__main__':
    base_dir = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisNaoContratado\_templates'

    processar(os.path.join(base_dir, 'inicial-jfba-base.docx'))
    processar(os.path.join(base_dir, 'inicial-jfba-multiplos-avn-inativo.docx'))
    processar(os.path.join(base_dir, 'inicial-jfba-refin-ativo.docx'), eh_refin=True)

    print('\n=== VERIFICAÇÃO PÓS-LIMPEZA ===')
    for nome in ['inicial-jfba-base.docx', 'inicial-jfba-multiplos-avn-inativo.docx', 'inicial-jfba-refin-ativo.docx']:
        path = os.path.join(base_dir, nome)
        doc = Document(path)
        print(f'\n--- {nome} ---')
        for i, p in enumerate(doc.paragraphs):
            for kw in ['temporal', 'temporais']:
                if kw in p.text.lower():
                    print(f'  RESIDUAL p{i}: {p.text[:200]}')
                    break
