"""Helper reutilizável para gerar inicial-combinada.docx com N teses.

Estrutura fiel ao modelo `PI Tarifas - Mora - Titulo.docx` da pasta do Lázaro:

  1. Núcleos fáticos: 1 parágrafo POR TESE com rótulo "PRIMEIRO NÚCLEO FÁTICO:"
     em NEGRITO + texto regular. Datas/rubrica/valores ganham bold+highlight.

  2. Pedidos: para cada tese ATIVA, substituir o {{BLOCO_PEDIDO_X}} pelo cabeçalho
     "Ao que se refere os descontos a título de "RUBRICA" Sejam julgados..."
     e os 4 {{__DELETE__}} subsequentes pelos 4 sub-itens reais (declaratório,
     condenatório, subsidiário, dano moral). Para teses INATIVAS, remover
     BLOCO_PEDIDO_X + os 4 __DELETE__ associados.

  3. Dano moral PEDIDO INDIVIDUAL = R$ 15.000 por tese (igual ao modelo
     paradigma). Dano moral TOTAL no VC = R$ 15.000 × N teses.

Uso:
    from _combinada_helper import gerar_combinada
    teses = [
        {'familia': 'TARIFAS', 'rubrica': 'CARTÃO CRÉDITO ANUIDADE',
         'lancamentos': [...]},
        {'familia': 'MORA', 'rubrica': 'MORA CRÉDITO PESSOAL',
         'lancamentos': [...]},
    ]
    gerar_combinada(pasta, nome_arquivo_base, autora, conta, renda, teses,
                    comarca, uf, eh_idoso)
"""
import os, sys, re, zipfile
from docx import Document
from lxml import etree

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W = '{' + W_NS + '}'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-combinada.docx'

ORDINAIS = {1: 'PRIMEIRO', 2: 'SEGUNDO', 3: 'TERCEIRO', 4: 'QUARTO', 5: 'QUINTO'}
# Regra § 9 da skill: 1 tese isolada = R$ 15.000; 2+ teses combinadas = R$ 5.000/tese
def _dano_moral_por_tese(n_teses):
    return 15000.00 if n_teses == 1 else 5000.00


def _get_text(p_element):
    return ''.join(t.text or '' for t in p_element.findall(f'.//{W}t'))


def _make_run(parent, text, *, bold=False, highlight=False, fonte='Cambria'):
    """Adiciona um w:r em parent com formatação configurável."""
    r = etree.SubElement(parent, W + 'r')
    rPr = etree.SubElement(r, W + 'rPr')
    rF = etree.SubElement(rPr, W + 'rFonts')
    rF.set(W + 'ascii', fonte); rF.set(W + 'hAnsi', fonte); rF.set(W + 'cs', fonte)
    if bold:
        etree.SubElement(rPr, W + 'b')
        etree.SubElement(rPr, W + 'bCs')
    if highlight:
        h = etree.SubElement(rPr, W + 'highlight')
        h.set(W + 'val', 'yellow')
    t = etree.SubElement(r, W + 't')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return r


def _substituir_runs(p_element, runs_data):
    """Remove os w:r existentes e adiciona novos. Preserva pPr/pStyle/numPr."""
    for r in list(p_element.findall(W + 'r')):
        p_element.remove(r)
    for rd in runs_data:
        _make_run(p_element, rd['text'], bold=rd.get('bold', False),
                  highlight=rd.get('highlight', False))


def _runs_nucleo_fatico(ordinal, tese):
    """Constrói os runs do parágrafo de núcleo fático."""
    n_ext = extenso_cardinal(tese['n_lanc'])
    val_ext = extenso_moeda(tese['total'])
    return [
        {'text': f'{ordinal} NÚCLEO FÁTICO: ',                    'bold': True},
        {'text': f'cobrança indevida sob a rubrica "{tese["rubrica"]}" em ', 'bold': False},
        {'text': str(tese['n_lanc']),                              'bold': True, 'highlight': True},
        {'text': f' ({n_ext}) lançamentos efetuados entre ',        'bold': False},
        {'text': tese['inicio'],                                    'bold': True, 'highlight': True},
        {'text': ' e ',                                             'bold': False},
        {'text': tese['fim'],                                       'bold': True, 'highlight': True},
        {'text': ', totalizando ',                                  'bold': False},
        {'text': fmt_moeda_rs(tese['total']),                       'bold': True, 'highlight': True},
        {'text': f' ({val_ext}).',                                  'bold': False},
    ]


def _runs_cabecalho_pedido(tese):
    """Cabeçalho 'Ao que se refere os descontos a título de "RUBRICA"...'"""
    return [
        {'text': 'Ao que se refere os descontos a título de ', 'bold': True},
        {'text': f'"{tese["rubrica"]}"',                       'bold': True},
        {'text': ' ',                                            'bold': False},
        {'text': 'Sejam julgados totalmente procedentes os pedidos iniciais, para o fim de:',
         'bold': False},
    ]


def _runs_subitens(tese, letra_subsidiaria, dano_moral_valor):
    """4 sub-itens: declaratório, condenatório, subsidiário, dano moral."""
    valor_simples = tese['total']
    valor_dobro   = tese['dobro']
    return [
        # 1. Declaratório
        [{'text': 'Declarar a inexistência da relação jurídica, bem como a nulidade de forma '
                  'definitiva, por ausência de contratação válida, bem como o cancelamento '
                  'definitivo de novos débitos da mesma espécie;', 'bold': False}],
        # 2. Condenatório (repetição em dobro)
        [{'text': 'Condenar o banco requerido à repetição do indébito, com a devolução em '
                  'dobro dos valores, com a incidência de correção monetária pelo INPC, a '
                  'partir de cada desconto indevido, e de juros de mora na monta de 1% ao mês '
                  'a contar da citação, cujo valor, até a presente data, é de ', 'bold': False},
         {'text': fmt_moeda_rs(valor_dobro), 'bold': True, 'highlight': True},
         {'text': f' ({extenso_moeda(valor_dobro)}).', 'bold': False}],
        # 3. Subsidiário
        [{'text': f'{letra_subsidiaria}.1) ', 'bold': False},
         {'text': 'Subsidiariamente, não sendo esse o entendimento de Vossa Excelência, '
                  'requer a repetição do indébito de forma simples, com a incidência de '
                  'correção monetária pelo INPC, a partir de cada desconto indevido, e de '
                  'juros de mora na monta de 1% ao mês a contar da citação, cujo valor, até '
                  'a presente data, é de ', 'bold': False},
         {'text': fmt_moeda_rs(valor_simples), 'bold': True, 'highlight': True},
         {'text': f' ({extenso_moeda(valor_simples)}).', 'bold': False}],
        # 4. Dano moral
        [{'text': 'Seja o banco requerido condenado ao pagamento de compensação por danos '
                  'morais causados à parte Autora, em razão da deslealdade, falta de '
                  'transparência, má-fé e conduta abusiva de seus prepostos, no valor de ',
          'bold': False},
         {'text': fmt_moeda_rs(dano_moral_valor), 'bold': True, 'highlight': True},
         {'text': f' ({extenso_moeda(dano_moral_valor)}), acrescidos de juros de mora a '
                  'partir da data do evento danoso, ou seja, da data do primeiro desconto '
                  'indevido (Súmula 54, STJ), e correção monetária pelo INPC, a partir do '
                  'arbitramento (Súmula 362, STJ).', 'bold': False}],
    ]


def gerar_combinada(pasta, nome_arquivo_base, autora, conta, renda, teses,
                    comarca='Maués', uf='AM', eh_idoso=False):
    docx_out = os.path.join(pasta, nome_arquivo_base + '_v1.docx')

    # ===== Pré-cálculo de cada tese =====
    for t in teses:
        t['total']    = sum(v for _, v in t['lancamentos'])
        t['dobro']    = t['total'] * 2
        t['n_lanc']   = len(t['lancamentos'])
        t['inicio']   = t['lancamentos'][0][0]
        t['fim']      = t['lancamentos'][-1][0]

    total_geral = sum(t['total'] for t in teses)
    dobro_geral = total_geral * 2
    n_teses     = len(teses)
    dano_moral_por_tese = _dano_moral_por_tese(n_teses)
    dano_moral_total = dano_moral_por_tese * n_teses
    valor_causa = dobro_geral + dano_moral_total

    familias_ativas   = set(t['familia'] for t in teses)
    familias_inativas = {'TARIFAS', 'MORA', 'TITULO', 'APLIC'} - familias_ativas

    # ===== Dados base =====
    todos_lanc = sum((t['lancamentos'] for t in teses), [])
    tese_dummy = {'rubrica': ' / '.join(t['rubrica'] for t in teses), 'lancamentos': todos_lanc}
    terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'',
                'cidade':'','uf':'','cep':''}
    dados, _ = montar_dados_padrao(autora=autora, conta=conta, renda=renda,
                                    tese=tese_dummy, terceiro=terceiro,
                                    eh_idoso=eh_idoso, competência=comarca, uf=uf)
    dados['valor_causa']               = fmt_moeda_rs(valor_causa)
    dados['valor_causa_extenso']       = extenso_moeda(valor_causa)
    dados['dano_moral_total']          = fmt_moeda_rs(dano_moral_total)
    dados['dano_moral_total_extenso']  = extenso_moeda(dano_moral_total)
    dados['total_descontos']           = fmt_moeda_rs(total_geral)
    dados['total_descontos_extenso']   = extenso_moeda(total_geral)
    dados['dobro_descontos']           = fmt_moeda_rs(dobro_geral)
    dados['dobro_descontos_extenso']   = extenso_moeda(dobro_geral)
    dados['remuneração']               = renda.get('descricao', 'aposentadoria pelo INSS')
    dados['nucleo_referencia_encargo'] = 'encargo'

    # Placeholders rubrica_mora* / rubrica_encargo* usados na seção IRDR
    # (citação do precedente TJAM 0004464-79.2023.8.04.0000) — só fazem
    # sentido quando a família MORA está ativa. Quando há 1 só rubrica MORA,
    # ambas as citações ('os encargos X e Y') recebem o mesmo valor.
    teses_mora = [t for t in teses if t['familia'] == 'MORA']
    if teses_mora:
        rub_mora_principal = teses_mora[0]['rubrica']
        rub_mora_secundaria = teses_mora[1]['rubrica'] if len(teses_mora) > 1 else rub_mora_principal
        dados['rubrica_mora']                   = rub_mora_principal
        dados['rubrica_mora_caps']              = rub_mora_principal.upper()
        dados['rubrica_mora_canonica']          = rub_mora_principal
        dados['rubrica_encargo']                = rub_mora_secundaria
        dados['rubrica_encargo_caps']           = rub_mora_secundaria.upper()
        dados['rubrica_encargo_canonica_caps']  = rub_mora_secundaria.upper()
    else:
        # Família MORA inativa — preenche com strings vazias para o
        # processar_paragrafo OMITIR limpamente (são parte do bloco MORA
        # que será removido inteiro pelo pós-processamento, mas o
        # aplicar_template precisa do dict pra não deixar {{...}} cru).
        dados['rubrica_mora']                   = ''
        dados['rubrica_mora_caps']              = ''
        dados['rubrica_mora_canonica']          = ''
        dados['rubrica_encargo']                = ''
        dados['rubrica_encargo_caps']           = ''
        dados['rubrica_encargo_canonica_caps']  = ''

    rubricas_lista = ', '.join(f'"{t["rubrica"]}"' for t in teses)
    dados['questao_em_discussao'] = (
        f'Saber se são devidos os descontos efetuados pelo banco réu, na conta '
        f'corrente da parte autora, sob as rubricas {rubricas_lista}, sem prova '
        f'documental de contratação válida ou anuência expressa do consumidor.'
    )
    dados['ratio_decidendi'] = (
        'A cobrança de tarifas bancárias e encargos acessórios em conta corrente '
        'do consumidor pressupõe a comprovação documental da prévia e expressa '
        'anuência, nos termos do art. 39, VI, do CDC, das Resoluções BACEN nº '
        '3.919/2010 e nº 4.196/2013, e dos IRDRs nº 0005053-71.2023.8.04.0000 '
        '(TJAM, tarifas) e nº 0004464-79.2023.8.04.0000 (TJAM, encargos de mora). '
        'A ausência dessa prova torna a cobrança indevida e enseja repetição em '
        'dobro dos valores debitados.'
    )
    dados['solucao_juridica'] = (
        'Declarar a inexistência de relação jurídica que legitime os débitos sob '
        'as rubricas impugnadas; condenar o réu à cessação imediata das cobranças, '
        'à restituição em dobro dos valores indevidamente debitados (art. 42, '
        'parágrafo único, do CDC) e ao pagamento de indenização por danos morais.'
    )
    sintese_total_txt = (
        f'No total, somando todos os {extenso_cardinal(n_teses)} núcleos fáticos, '
        f'foram debitados {fmt_moeda_rs(total_geral)} ({extenso_moeda(total_geral)}) '
        f'da conta da parte autora, valor que, restituído em dobro nos termos do '
        f'art. 42, parágrafo único, do CDC, totaliza {fmt_moeda_rs(dobro_geral)} '
        f'({extenso_moeda(dobro_geral)}).'
    )
    dados['sintese_total'] = sintese_total_txt
    # NÃO incluir 'nucleos_faticos' nem 'BLOCO_PEDIDO_*' no dict — assim o
    # aplicar_template deixa esses placeholders INTACTOS para o pós-processamento
    # via python-docx.

    # ===== Aplicar template (placeholders simples) =====
    # IMPORTANTE: strict=False aqui porque o template inicial-combinada usa
    # placeholders de seção ({{INICIO_BLOCO_X}}, {{FIM_BLOCO_X}},
    # {{BLOCO_PEDIDO_X}}, {{__DELETE__}}, {{nucleos_faticos}}) que são
    # processados depois via python-docx no pós-processamento abaixo. A
    # verificação final de residuais acontece no fim deste helper.
    res = aplicar_template(TEMPLATE, dados, docx_out, strict=False)

    # ===== Pós-processamento via python-docx =====
    doc = Document(docx_out)

    # Marcamos parágrafos para remoção via REFERÊNCIA (não por índice), porque vamos
    # inserir parágrafos novos no meio.
    paras = list(doc.paragraphs)
    para_remove = set()  # set de id(elemento) → guardar referência

    def marcar(p_obj):
        para_remove.add(id(p_obj._element))

    def texto(p_obj):
        return _get_text(p_obj._element)

    # 1) Remover blocos doutrinários inteiros das famílias INATIVAS
    for fam in familias_inativas:
        in_block = False
        for p in paras:
            t = texto(p)
            if f'{{{{INICIO_BLOCO_{fam}}}}}' in t:
                in_block = True; marcar(p); continue
            if f'{{{{FIM_BLOCO_{fam}}}}}' in t:
                marcar(p); in_block = False; continue
            if in_block:
                marcar(p)

    # 2) Remover marcadores INICIO/FIM_BLOCO_X de famílias ATIVAS (manter conteúdo)
    for fam in familias_ativas:
        for p in paras:
            t = texto(p)
            if f'{{{{INICIO_BLOCO_{fam}}}}}' in t or f'{{{{FIM_BLOCO_{fam}}}}}' in t:
                marcar(p)

    # 3) BLOCO_DIREITO_APLIC: remover se APLIC inativo
    if 'APLIC' in familias_inativas:
        for p in paras:
            if '{{BLOCO_DIREITO_APLIC}}' in texto(p):
                marcar(p)

    # 4) Mapear BLOCO_PEDIDO_X e seus 4 __DELETE__ associados (em ordem após o pedido)
    SUBS_POR_FAMILIA = {'TARIFAS': 4, 'MORA': 4, 'TITULO': 4, 'APLIC': 0}
    pedido_subs = {}  # família -> {'p_cab': obj, 'p_subs': [obj,...]}
    for i, p in enumerate(paras):
        t = texto(p)
        for fam in ('TARIFAS', 'MORA', 'TITULO', 'APLIC'):
            if f'{{{{BLOCO_PEDIDO_{fam}}}}}' in t:
                subs_objs = []
                target = SUBS_POR_FAMILIA[fam]
                j = i + 1
                while j < len(paras) and len(subs_objs) < target:
                    if '{{__DELETE__}}' in texto(paras[j]):
                        subs_objs.append(paras[j])
                    j += 1
                pedido_subs[fam] = {'p_cab': p, 'p_subs': subs_objs}
                if fam in familias_inativas:
                    marcar(p)
                    for s in subs_objs:
                        marcar(s)

    # 5) __DELETE__ que NÃO estão associados a um BLOCO_PEDIDO ativo: remover
    ids_subs_ativos = set()
    for fam in familias_ativas:
        if fam in pedido_subs:
            for s in pedido_subs[fam]['p_subs']:
                ids_subs_ativos.add(id(s._element))
    for p in paras:
        if '{{__DELETE__}}' in texto(p) and id(p._element) not in ids_subs_ativos:
            marcar(p)

    # 6) Substituir BLOCO_PEDIDO_X (cabeçalho) e __DELETE__ associados (sub-itens) das ATIVAS
    sub_itens_por_familia = {}  # para reinício de numeração depois
    for fam in familias_ativas:
        if fam not in pedido_subs:
            continue
        teses_fam = [t for t in teses if t['familia'] == fam]
        if not teses_fam:
            continue
        rubrica_combinada = ' / '.join(t['rubrica'] for t in teses_fam)
        total_fam   = sum(t['total'] for t in teses_fam)
        dobro_fam   = total_fam * 2
        n_lanc_fam  = sum(t['n_lanc'] for t in teses_fam)
        ini_fam     = teses_fam[0]['inicio']
        fim_fam     = teses_fam[-1]['fim']
        tese_fam = {'rubrica': rubrica_combinada, 'total': total_fam,
                    'dobro': dobro_fam, 'n_lanc': n_lanc_fam,
                    'inicio': ini_fam, 'fim': fim_fam}

        # Substituir cabeçalho preservando pPr/pStyle/numPr
        p_cab = pedido_subs[fam]['p_cab']
        _substituir_runs(p_cab._element, _runs_cabecalho_pedido(tese_fam))

        # Substituir os 4 sub-itens (declaratório, condenatório, subsidiário, dano moral)
        sub_runs_list = _runs_subitens(tese_fam, '', dano_moral_por_tese)
        sub_itens_por_familia[fam] = []
        for k, p_sub in enumerate(pedido_subs[fam]['p_subs']):
            if k >= len(sub_runs_list):
                break
            sub_itens_por_familia[fam].append(p_sub)
            runs_data = sub_runs_list[k]
            # Sub-item 3 (subsidiário): remover prefixo "X.1) " e ajustar capitalização
            if k == 2:
                runs_data = runs_data[1:]
                if runs_data:
                    runs_data[0]['text'] = 'Subsidiariamente, não sendo esse o entendimento ' \
                        'de Vossa Excelência, requer a repetição do indébito de forma simples, ' \
                        'com a incidência de correção monetária pelo INPC, a partir de cada ' \
                        'desconto indevido, e de juros de mora na monta de 1% ao mês a contar ' \
                        'da citação, cujo valor, até a presente data, é de '
            _substituir_runs(p_sub._element, runs_data)

    # 7) Substituir o parágrafo {{nucleos_faticos}} por N parágrafos individuais
    p_nuc_template = None
    for p in paras:
        if '{{nucleos_faticos}}' in texto(p):
            p_nuc_template = p; break
    if p_nuc_template is not None:
        import copy as _copy
        for k, t in enumerate(teses, 1):
            ord_str = ORDINAIS[k]
            novo_p = _copy.deepcopy(p_nuc_template._element)
            for r in list(novo_p.findall(W + 'r')):
                novo_p.remove(r)
            for rd in _runs_nucleo_fatico(ord_str, t):
                _make_run(novo_p, rd['text'], bold=rd.get('bold', False),
                          highlight=rd.get('highlight', False))
            p_nuc_template._element.addprevious(novo_p)
        marcar(p_nuc_template)

    # 8) Aplicar remoção via referência (não índice — já que paras inseridos mudam ordem)
    for p in paras:
        if id(p._element) in para_remove and p._element.getparent() is not None:
            p._element.getparent().remove(p._element)

    # 10) Limpeza de citações órfãs (parágrafos do "Do dano moral" citando teses inativas)
    ANCHORS = {
        'TITULO': ['título de capitalização', 'desconto de título de capitalização',
                    'TÍTULO DE CAPITALIZAÇÃO'],
        'MORA':   ['ENCARGOS LIMITE DE CRED', 'encargos limite de cred',
                    'Crédito Mora Pessoal', 'CRÉDITO MORA PESSOAL', 'Mora Cred Pess'],
        'APLIC':  ['APLIC.INVEST', 'aplic.invest', 'INVEST FACIL', 'INVEST FÁCIL'],
    }
    paras_pos = list(doc.paragraphs)
    idx_orfaos = set()
    for fam in familias_inativas:
        for ancora in ANCHORS.get(fam, []):
            for i, p in enumerate(paras_pos):
                t = _get_text(p._element)
                if ancora not in t:
                    continue
                idx_orfaos.add(i)
                for j in range(i + 1, min(i + 7, len(paras_pos))):
                    nxt = _get_text(paras_pos[j]._element)
                    if not nxt.strip():
                        idx_orfaos.add(j); continue
                    if (nxt.startswith('EMENTA') or
                        nxt.startswith('RECURSO INOMINADO') or
                        nxt.startswith('RECURSOS INOMINADOS') or
                        nxt.startswith('(Recurso') or
                        '[...]' in nxt or
                        nxt.startswith('No mesmo sentido') or
                        nxt.startswith('Esse é, inclusive') or
                        nxt.startswith('E também reconhece') or
                        len(nxt) > 250):
                        idx_orfaos.add(j)
                    else:
                        break
    for i in sorted(idx_orfaos, reverse=True):
        if i < len(paras_pos):
            p = paras_pos[i]
            if p._element.getparent() is not None:
                p._element.getparent().remove(p._element)

    # ===== Reiniciar numeração a/b/c em cada bloco de sub-itens =====
    # Objetivo: cada bloco de pedido deve começar pela letra "a)" (não continuar
    # do bloco anterior). Usamos numIds novos com <w:lvlOverride startOverride=1>.
    if sub_itens_por_familia:
        # 1) Coletar TODOS os numIds usados pelos sub-itens (em geral todos = mesmo numId,
        # ex.: 20). Vamos usar o primeiro encontrado como base.
        numId_base = None
        for fam, p_list in sub_itens_por_familia.items():
            for p in p_list:
                m = p._element.find(f'.//{W}numPr/{W}numId')
                if m is not None:
                    numId_base = m.get(W + 'val')
                    break
            if numId_base:
                break

        if numId_base:
            # 2) Inserir nova definição de numId no numbering.xml
            from docx.oxml.ns import qn
            num_part = doc.part.numbering_part
            if num_part is not None:
                num_xml = num_part.element
                # Achar abstractNumId usado por numId_base
                base_num = None
                for n in num_xml.findall(qn('w:num')):
                    if n.get(qn('w:numId')) == numId_base:
                        base_num = n; break
                abstract_id = '0'
                if base_num is not None:
                    abs_el = base_num.find(qn('w:abstractNumId'))
                    if abs_el is not None:
                        abstract_id = abs_el.get(qn('w:val'))

                # Maior numId atual
                existentes = [int(n.get(qn('w:numId'))) for n in num_xml.findall(qn('w:num'))]
                proximo = max(existentes) + 1 if existentes else 100

                # Para cada família ativa, criar um novo numId com startOverride=1
                fam_to_new_id = {}
                for fam in sub_itens_por_familia:
                    novo_num = etree.SubElement(num_xml, qn('w:num'))
                    novo_num.set(qn('w:numId'), str(proximo))
                    abs_e = etree.SubElement(novo_num, qn('w:abstractNumId'))
                    abs_e.set(qn('w:val'), abstract_id)
                    lvl_over = etree.SubElement(novo_num, qn('w:lvlOverride'))
                    lvl_over.set(qn('w:ilvl'), '0')
                    start_over = etree.SubElement(lvl_over, qn('w:startOverride'))
                    start_over.set(qn('w:val'), '1')
                    fam_to_new_id[fam] = str(proximo)
                    proximo += 1

                # 3) Aplicar o novo numId aos sub-itens de cada família
                for fam, p_list in sub_itens_por_familia.items():
                    novo_id = fam_to_new_id[fam]
                    for p_obj in p_list:
                        numId_el = p_obj._element.find(f'.//{W}numPr/{W}numId')
                        if numId_el is not None:
                            numId_el.set(W + 'val', novo_id)

    doc.save(docx_out)

    # ===== Pós-fix raw para placeholders quebrados entre runs =====
    # Limpar tags w:r/w:t intermediárias que quebram placeholders {{...}} e depois substituir
    with zipfile.ZipFile(docx_out, 'r') as z:
        nomes = z.namelist()
        buf = {n: z.read(n) for n in nomes}
    xml = buf['word/document.xml'].decode('utf-8')

    sub_extra = {
        '{{remuneração}}':              dados['remuneração'],
        '{{valor_remuneração}}':        dados['valor_remuneração'],
        '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso'],
        '{{questao_em_discussao}}':     dados['questao_em_discussao'],
        '{{ratio_decidendi}}':          dados['ratio_decidendi'],
        '{{solucao_juridica}}':         dados['solucao_juridica'],
        '{{sintese_total}}':            dados['sintese_total'],
        '{{nucleo_referencia_encargo}}': dados['nucleo_referencia_encargo'],
        '{{cidade_filial}}': 'Maués', '{{uf_filial}}': 'AM', '{{uf_extenso}}': 'Amazonas',
    }

    fix = 0
    # 1ª passada: substituições simples (placeholder íntegro)
    for k, v in sub_extra.items():
        if k in xml:
            cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt

    # 2ª passada: placeholders QUEBRADOS entre runs (ex.: "{{cidade_filial</w:t>...</w:r>...<w:t>}}")
    # estratégia: para cada chave residual, tenta uma regex tolerante a tags Word entre os caracteres
    placeholders_residuais = sorted(set(re.findall(r'\{\{[A-Za-zçãõéúíáÇÃÕÉÚÍÁ_]+', xml)))
    for ph_inicio in placeholders_residuais:
        nome = ph_inicio[2:]  # remove "{{"
        if nome not in [k.strip('{}') for k in sub_extra]:
            # tentar achar correspondência tolerante
            chave_dict = '{{' + nome + '}}'
            valor = sub_extra.get(chave_dict)
            if not valor:
                continue
        else:
            chave_dict = '{{' + nome + '}}'
            valor = sub_extra[chave_dict]
        # Construir regex que permite tags Word entre os caracteres do nome
        chars_pattern = ''.join(re.escape(c) + r'(?:<[^>]*>)*' for c in nome)
        pat = r'\{\{(?:<[^>]*>)*' + chars_pattern + r'\}\}'
        new_xml, n = re.subn(pat, valor, xml)
        if n:
            xml = new_xml; fix += n

    buf['word/document.xml'] = xml.encode('utf-8')
    os.remove(docx_out)
    with zipfile.ZipFile(docx_out, 'w', zipfile.ZIP_DEFLATED) as z:
        for n in nomes: z.writestr(n, buf[n])

    # ===== Verificação FINAL de residuais (após TODO o pós-processamento) =====
    # Equivale ao strict=True do aplicar_template, mas só roda agora que o
    # helper terminou de processar os placeholders de seção.
    from helpers_docx import PlaceholdersResiduaisError
    with zipfile.ZipFile(docx_out, 'r') as z:
        xml_final = z.read('word/document.xml').decode('utf-8')
    residuais_final = sorted(set(re.findall(r'\{\{([^{}]+)\}\}', xml_final)))
    if residuais_final:
        base, ext = os.path.splitext(docx_out)
        falha_path = base + '_FALHOU_PLACEHOLDERS' + ext
        if os.path.exists(falha_path):
            os.remove(falha_path)
        os.rename(docx_out, falha_path)
        raise PlaceholdersResiduaisError(residuais_final, falha_path)

    return {
        'docx': docx_out,
        'res': res,
        'paras_removidos': len(para_remove) + len(idx_orfaos),
        'pos_fix': fix,
        'totais': {
            'n_teses': n_teses,
            'total_geral': total_geral,
            'dobro_geral': dobro_geral,
            'dano_moral_total': dano_moral_total,
            'dano_moral_por_tese': dano_moral_por_tese,
            'valor_causa': valor_causa,
            'familias_ativas': familias_ativas,
            'teses': teses,
        },
    }
