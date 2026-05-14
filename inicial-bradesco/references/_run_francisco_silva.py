"""Gera inicial APLIC.INVEST FACIL do FRANCISCO DA SILVA FERREIRA.

Comarca: Presidente Figueiredo/AM (foro do domicílio art. 101 I CDC —
Comunidade Cacaia, BR 174 KM 156, Zona Rural). Idoso 68 anos
(nascido 17/03/1958). Estado civil OMITIDO (CNH não traz; procuração
não menciona). Conta Bradesco Ag 3732, conta 20531-1.

Renda: 2 benefícios INSS distintos creditados mensalmente em datas
diferentes (R$ 1.103,07 dia 06 + R$ 987,81 dia 07) = soma R$ 2.090,88
mensais. Padrão de aposentadoria + pensão por morte (presumido —
confirmar com cliente).

Tese APLIC.INVEST FACIL — 128 lançamentos entre 06/04/2020 e 09/08/2024,
total bruto aplicado R$ 147.816,26 (NÃO PLEITEAR — ciclo aplica-resgate
confirmado: 135 APLICs no extrato vs 324 RESGATEs, saldo líquido
NEGATIVO R$ 8,68 — cliente recebeu de volta tudo + R$ 8,68 de
rentabilidade).

ESTRATÉGIA (b) PADRÃO. VC R$ 15.000 (~10 SM) → cabe folgadamente no JEC
(40 SM = R$ 60.720). Sem (b), VC seria 15.000 + 295.632,52 = R$ 310 mil
— estouraria 5x o teto JEC.

Comprovante de residência está no nome de TERCEIRO (5.2 - RG
PROPRIETÁRIO indica que o imóvel é de outra pessoa). Há AUTODECLARAÇÃO
DE RESIDÊNCIA (5.1) e o nome do proprietário em separado.

Pasta KIT contém material para outras 5 teses (MORA CRED, PGTO ELETRÔN
ASPECIR, RMC, TARIFA, TÍTULO CAPITALIZAÇÃO). Notificação extrajudicial
APLIC tem (.docx) — se foi enviada de fato, juntar comprovante.
"""
import io, sys, os, json, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\FRANCISCO DA SILVA FERREIRA - Ruth (bradesco)\APLICAÇÃO INVEST FÁCIL'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_FRANCISCO_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_FRANCISCO_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_francisco_lancs.json'), encoding='utf-8') as f:
    LANCAMENTOS = sorted(json.load(f), key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'FRANCISCO DA SILVA FERREIRA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',  # CNH não traz, procuração não menciona — omitir
    'profissao': 'aposentado',
    'cpf': '000.000.011-21',
    'rg': '1000009-9',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Comunidade Cacaia – LD 11237, BR 174 – LM 156',
    'numero': 's/nº',
    'bairro': 'Zona Rural',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '20531-1'}
renda = {'valor_float': 2090.88}  # 2 benefícios INSS somados

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome': '', 'cnpj': '', 'logradouro': '', 'numero': '',
            'bairro': '', 'cidade': '', 'uf': '', 'cep': ''}

dados, calc = montar_dados_padrao(
    autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
    eh_idoso=True, competência='Presidente Figueiredo', uf='AM',
)

DANO_MORAL = 15000.00
VALOR_CAUSA_B = 15000.00
dados['valor_causa'] = fmt_moeda_rs(VALOR_CAUSA_B)
dados['valor_causa_extenso'] = extenso_moeda(VALOR_CAUSA_B)
dados['remuneração'] = 'aposentadoria pelo INSS (com pensão por morte acumulada)'

print('=== FRANCISCO DA SILVA FERREIRA — APLIC.INVEST FACIL (estratégia b) ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total bruto: R$ {calc["total"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
print(f'VC (b): R$ {VALOR_CAUSA_B:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-processamento
from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
removidos = 0
for p in list(doc.paragraphs):
    txt = p.text or ''
    for m in MARKERS:
        if m in txt:
            p._element.getparent().remove(p._element); removidos += 1; break
print(f'Removidos: {removidos}')

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'reiterado durante todo o período em que a renda alimentar da parte autora ficou '
        'indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = ('No caso concreto, o extrato bancário registra 128 (cento e vinte e oito) ocorrências '
      'de aplicação automática entre 06/04/2020 e 09/08/2024 (mais de 4 anos). Em todos os meses, '
      'parcela substancial da renda da parte autora — recebida em conta-salário/benefício — foi '
      'automaticamente subtraída pelo banco réu sob a rubrica APLIC.INVEST FACIL, restando '
      'indisponível ao consumidor pelo prazo de 1 (um) a 5 (cinco) dias até o resgate manual. '
      'Embora os valores tenham sido restituídos via RESGATE INVEST FACIL ao longo do período '
      '(324 ocorrências de resgate), o cerne do dano moral não reside na perda patrimonial '
      'líquida — inexistente, pois o saldo agregado em 4 anos foi de aproximadamente zero —, '
      'mas na privação reiterada da autodeterminação do consumidor sobre sua própria renda '
      'alimentar, mês após mês, durante 4 (quatro) anos consecutivos. Cada retenção mensal '
      'configura, autonomamente, prática abusiva vedada pelo art. 39, inciso VI, do Código de '
      'Defesa do Consumidor, sendo a recorrência sistêmica o fato gerador do abalo extrapatrimonial.')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_FRANCISCO', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'APLIC.INVEST FACIL — estratégia (b) PADRÃO'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_AplicInvest_FRANCISCO_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '17/03/1958 — 68 anos (IDOSO)'),
    ('Estado civil', '(omitido)'),
    ('Endereço', f'{autora["logradouro"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda usada (extrato)', f'{dados["valor_remuneração"]} (soma de 2 benefícios INSS mensais)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC (tabela)', '128'),
    ('Lançamentos APLIC (extrato)', '135 — diferença 7 lançamentos a confirmar'),
    ('Lançamentos RESGATE (extrato)', '324'),
    ('Total bruto aplicado (NÃO pleiteado)', dados['total_descontos']),
    ('Saldo líquido (APLIC - RESGATE)', '-R$ 8,68 (cliente recebeu R$ 8,68 a mais via rentabilidade)'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('AUDITORIA APLIC vs RESGATE — confirmada',
     'Extrato confirma ciclo aplica-resgate: 135 APLIC.INVEST FACIL e 324 RESGATE INVEST FACIL '
     'totalizando R$ 150.893,07 vs R$ 150.901,75. Saldo NEGATIVO R$ 8,68 (rentabilidade). '
     'Estratégia (b) PADRÃO aplicada.'),
    ('TABELA vs EXTRATO — diferença de 7 lançamentos',
     'A tabela 7-TABELA.pdf traz 128 lançamentos APLIC totalizando R$ 147.816,26. O extrato '
     'detecta 135 APLICs totalizando R$ 150.893,07 (~R$ 3.000 a mais). Possíveis hipóteses: '
     '(i) tabela truncou alguns lançamentos; (ii) extrato detectou alguns lançamentos com valores '
     'diferentes da tabela. CONFERIR antes do protocolo se a inicial precisa usar tabela ou extrato.'),
    ('RENDA — 2 benefícios INSS distintos',
     'Cliente recebe MENSALMENTE 2 créditos INSS distintos: R$ 1.103,07 (em torno do dia 06) + '
     'R$ 987,81 (em torno do dia 07). Total mensal R$ 2.090,88. Padrão típico de '
     'aposentadoria + pensão por morte. CONFIRMAR COM CLIENTE quais são os 2 benefícios. '
     'Para a Justiça Gratuita está sendo usada a SOMA (R$ 2.090,88) — se houver razão para '
     'mostrar separado, ajustar. Pode trazer questionamento do juízo sobre hipossuficiência '
     '(2 benefícios ainda enquadram, mas recomenda-se justificar despesas).'),
    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Template inicial-aplic-invest.docx pós-processado: 7 parágrafos do bloco "Repetição do '
     'indébito" REMOVIDOS, parágrafo doutrinário REESCRITO, parágrafo do caso concreto '
     '(128 retenções, 4+ anos) INSERIDO.'),
    ('COMPROVANTE DE RESIDÊNCIA — imóvel de TERCEIRO',
     'Pasta tem 5.1 - AUTODECLARAÇÃO DE RESIDÊNCIA + 5.2 - RG PROPRIETÁRIO indicando que o '
     'imóvel é de terceiro (não do autor). Confirmar vínculo (familiar? aluguel?) e juntar '
     'declaração formalizada se houver dúvida.'),
    ('NOTIFICAÇÃO EXTRAJUDICIAL — formato .docx',
     'Pasta tem "Notificação Extrajudicial - APLICAÇÃO INVEST FÁCIL.docx" mas SEM comprovante '
     'de envio juntado. Avaliar se já foi enviada (e juntar comprovante depois) ou se precisa '
     'enviar agora antes do protocolo. Bloco "Do prévio requerimento administrativo" '
     '(parágrafos 25-29 do template) MENCIONA notificação — se NÃO foi enviada, AJUSTAR no DOCX.'),
    ('PASTA KIT — outras 5 teses do mesmo cliente',
     'KIT contém procurações + tabelas + notificações para: MORA CRED + ENCARGO, PGTO ELETRÔN '
     'COBRANÇA ASPECIR, RMC (com HISCON e HISCRE), TARIFA BANCÁRIA, TÍTULO DE CAPITALIZAÇÃO. '
     '5 teses adicionais — decidir em batch separado se vai gerar iniciais combinadas '
     '(Presidente Figueiredo adota combinação) ou separadas.'),
    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 ≈ 9,87 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720). Sem estratégia (b), '
     'VC seria R$ 310k — estouraria 5x o teto.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'Conferir nome / CPF / RG / nascimento.',
    'Conferir conta/agência (20531-1 / 3732).',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'Conferir prioridade idoso aplicada.',
    'Conferir VC = R$ 15.000,00 (apenas dano moral).',
    'Conferir parágrafos doutrinários "Repetição do indébito" REMOVIDOS.',
    'Conferir parágrafo caso concreto inserido (128 retenções, 4+ anos).',
    'CONFIRMAR 2 benefícios INSS com cliente.',
    'Resolver TABELA vs EXTRATO (128 vs 135 lançamentos).',
    'Avaliar envio/comprovante notificação extrajudicial.',
    'Decidir sobre as outras 5 teses do KIT (combinada ou separadas).',
    'Anexar AUTODECLARAÇÃO + RG PROPRIETÁRIO (comprovante residência).',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com ressalvas — após pendências, ')
r2 = p.add_run('PROTOCOLAR.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
