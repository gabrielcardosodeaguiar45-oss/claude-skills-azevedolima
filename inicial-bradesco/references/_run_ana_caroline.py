"""Inicial MORA — CLIENTE EXEMPLO SEIXAS DE SOUZA.

Comarca Barreirinha/AM. Conta Bradesco Ag 3725/7359-8. Notificação chama de
"pessoa idosa", mas extrato traz CRÉDITO DE SALÁRIO (servidora/empregada),
provavelmente NÃO idosa — confirmar RG.

Tese: MORA CREDITO PESSOAL — 24 lançamentos entre 12/02/2020 e 04/01/2021
(total R$ 3.554,71; dobro R$ 7.109,42). Dano moral R$ 15.000.
VC R$ 22.109,42 — cabe no JEC.

PENDÊNCIAS CRÍTICAS:
1. PRESCRIÇÃO CDC — TODOS os 24 MORA estão fora do prazo de 5 anos do art. 27
   CDC (último 04/01/2021, prazo até 04/01/2026; hoje 06/05/2026 → 4 meses
   prescritos). Inicial é gerada apostando na corrente DECENAL (art. 205 CC,
   acolhida em algumas câmaras do TJ-AM, ver caso EXEMPLA RAIMUNDA).
2. RENDA ATUAL DESCONHECIDA — último crédito de salário foi 28/01/2021
   R$ 966,63. Conta zerada de 02/2021 até 09/02/2026, com OPERAÇÕES VENCIDAS
   pendentes (R$ 4.863,98). Adotada renda R$ 966,63 do último crédito como
   ESTIMATIVA — confirmar renda atual com cliente antes do protocolo.
"""
import io, sys, os, json
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda, extenso_cardinal

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\1. TESTE 2\CLIENTE EXEMPLO SEIXAS DE SOUZA - Wilson - TARIFAS\MORA CRED PESS'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-mora.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_Mora_ANA_CAROLINE_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_ANA_CAROLINE_v1.docx')

with open(os.path.join(os.path.dirname(__file__), '_ana_caroline_lancs.json')) as f:
    data = json.load(f)
LANCAMENTOS = sorted(data['mora'], key=lambda x: tuple(reversed([int(p) for p in x[0].split('/')])))

autora = {
    'nome': 'CLIENTE EXEMPLO SEIXAS DE SOUZA', 'nacionalidade': 'brasileira',
    'estado_civil': '', 'profissao': '',
    'cpf': '000.000.002-12', 'rg': '1000001-1',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Nove de Junho', 'numero': '498',
    'bairro': 'Centro', 'cidade': 'Barreirinha', 'cep': '69.160-000',
}
conta = {'agencia': '3725', 'numero': '7359-8'}
renda = {'valor_float': 966.63}  # ESTIMATIVA - último crédito 28/01/2021

tese = {'rubrica': 'MORA CREDITO PESSOAL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome':'','cnpj':'','logradouro':'','numero':'','bairro':'','cidade':'','uf':'','cep':''}

dados, calc = montar_dados_padrao(autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
                                  eh_idoso=False, competência='Barreirinha', uf='AM')
dados['remuneração'] = 'salário'
# Placeholders MORA específicos
dados['rubrica_curta'] = 'Mora Cred Pess'
dados['rubrica_curta_caps'] = 'MORA CRED PESS'
dados['rubrica_completa'] = 'Mora Crédito Pessoal'
dados['rubrica_completa_caps'] = 'MORA CRÉDITO PESSOAL'
dados['titulo'] = 'MORA CRÉDITO PESSOAL'
dados['numero_desconto'] = str(len(LANCAMENTOS))
dados['desconto_extenso'] = extenso_cardinal(len(LANCAMENTOS))

print(f'=== CLIENTE EXEMPLO — MORA ===')
print(f'Lançamentos: {len(LANCAMENTOS)}, total: R$ {calc["total"]:,.2f} | dobro: R$ {calc["dobro"]:,.2f} | VC: R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))

res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

# Pós-fix raw
import zipfile
with zipfile.ZipFile(DOCX_OUT, 'r') as z:
    nomes = z.namelist()
    buf = {n: z.read(n) for n in nomes}
xml = buf['word/document.xml'].decode('utf-8')
sub = {'{{remuneração}}': dados['remuneração'],
       '{{valor_remuneração}}': dados['valor_remuneração'],
       '{{valor_remuneração_extenso}}': dados['valor_remuneração_extenso']}
fix = 0
for k, v in sub.items():
    if k in xml:
        cnt = xml.count(k); xml = xml.replace(k, v); fix += cnt
buf['word/document.xml'] = xml.encode('utf-8')
os.remove(DOCX_OUT)
with zipfile.ZipFile(DOCX_OUT, 'w', zipfile.ZIP_DEFLATED) as z:
    for n in nomes: z.writestr(n, buf[n])
print(f'Pós-fix: {fix}')
print(f'OK -> {DOCX_OUT}')

# Relatório paralelo
from docx import Document
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_Mora_ANA_CAROLINE', level=1)
for k, v in [('Cliente', autora['nome']), ('Tese', 'MORA CREDITO PESSOAL'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'), ('Arquivo', 'INICIAL_Mora_ANA_CAROLINE_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda usada (ESTIMATIVA)', f'{dados["valor_remuneração"]} (último crédito 28/01/2021 — CONFIRMAR renda atual)'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos MORA', '24'),
    ('Total bruto', dados['total_descontos']),
    ('Total em dobro', dados['dobro_descontos']),
    ('Dano moral', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('PRESCRIÇÃO CDC — APOSTA NA TESE DECENAL (art. 205 CC)',
     'TODOS os 24 lançamentos MORA estão entre 12/02/2020 e 04/01/2021. Pelo art. 27 '
     'CDC (5 anos) o último estaria prescrito desde 04/01/2026 (4 meses antes do '
     'protocolo em 06/05/2026). Pelo art. 205 CC (10 anos, repetição de indébito por '
     'enriquecimento sem causa) TODOS estariam vivos. Esta inicial APOSTA na corrente '
     'DECENAL acolhida em algumas câmaras do TJ-AM. Procurador deve estar preparado '
     'para defender a tese decenal em recurso. Padrão validado no caso RAIMUNDA '
     'RODRIGUES (06/05/2026, sessão anterior).'),
    ('RENDA ATUAL DESCONHECIDA',
     'O extrato mostra CRÉDITO DE SALARIO recorrente até 28/01/2021 R$ 966,63. A '
     'partir de 02/2021 a conta ficou ZERADA até 09/02/2026, quando aparece o '
     'lançamento OPERACOES VENCIDAS R$ 4.863,98. Adotada renda R$ 966,63 do último '
     'crédito como ESTIMATIVA. CONFIRMAR com cliente: (a) onde recebe salário hoje? '
     '(b) trocou de conta? (c) está ciente da dívida em aberto na conta?'),
    ('IDADE — não confirmada',
     'RG não lido. Notificação chama de "pessoa idosa", mas extrato traz CRÉDITO DE '
     'SALARIO (típico de empregada/servidora). Texto da notificação é padrão e não '
     'confiável. Assume NÃO IDOSO. Conferir RG.'),
    ('TETO JEC — coberto',
     f'VC R$ {calc["valor_causa"]:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.') + ' ≈ 14,56 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),
    ('Procurador Patrick Willian (OAB/AM A2638)',
     'Notificação extrajudicial é assinada pelo Dr. Patrick Willian da Silva, do '
     'escritório em Joaçaba-SC. Contrato de serviços com a equipe Wilson da pasta '
     'principal — confirmar quem assina a inicial.'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in ['Conferir nome / CPF / RG.', 'Conferir conta/agência (7359-8 / 3725).',
           'Conferir comarca: Barreirinha/AM.', 'CONFIRMAR renda atual.',
           'Confirmar idade no RG.',
           'CONFIRMAR estratégia decenal — preparar argumentação anti-prescrição CDC.',
           'Anexar 8-NOTIFICAÇÃO + 8.1-COMPROVANTE NOTIFICAÇÃO.']:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA COM RESSALVA DECENAL — '); r2 = p.add_run('PROTOCOLAR após confirmar renda, idade e estratégia anti-prescrição.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
