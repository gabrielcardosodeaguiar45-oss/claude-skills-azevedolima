"""Gera inicial APLIC.INVEST FACIL do EXEMPLO DA SILVA DE SOUSA.

CASO ATÍPICO — APENAS 2 lançamentos APLIC.INVEST FACIL isolados em 2020
e 2022 (R$ 487,41 + R$ 7.572,26 = R$ 8.059,67), saldo líquido zero. NÃO
configura "recorrência sistêmica" típica da estratégia (b) padrão.

ALERTA CRÍTICO #1: a planilha 7 - TABELA.xlsx do cliente tem 3 abas
(MORA CRED PESSOAL R$ 2.175,18 + TARIFA BANCÁRIA CESTA R$ 333,86 +
APLIC.INVEST FÁCIL R$ 8.059,67). Recomendação fortemente sugerida:
processar como INICIAL COMBINADA (MORA + TARIFA + APLIC) via template
inicial-combinada.docx, em vez de APLIC isolado. Isso fortalece a peça
e, dado que Presidente Figueiredo adota combinação, mantém o pleito
viável. Esta inicial APLIC isolada é gerada apenas porque o batch atual
foi orientado a processar APLIC.INVEST especificamente — DECISÃO FINAL
do procurador sobre se vai protocolar APLIC sozinha ou combinada.

Comarca: Presidente Figueiredo/AM. NÃO IDOSO (nascido 12/11/1966 →
59 anos). Motorista profissional (CNH AD com EAR). Estado civil
omitido. Conta Bradesco Ag 3732 / 512518-9.

Renda: NÃO IDENTIFICADA no extrato APLIC (só 1 ano de 2020). Adotando
R$ 1.500 como valor estimado para Justiça Gratuita — CONFIRMAR COM
CLIENTE renda atual (2026).

Pasta APLIC.INVEST FÁCIL contém TUDO (procuração, RG, comprovante,
autodeclaração, RG proprietário, extrato, notificação .docx).
Notificação extrajudicial está em formato .docx — CONFERIR se já foi
enviada (e juntar comprovante) ou enviar agora.

ESTRATÉGIA (b) ADAPTADA: dano moral fundamentado em 2 ocorrências
autônomas (não em recorrência sistêmica). VC R$ 15.000.
"""
import io, sys, os, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _pipeline_caso import montar_dados_padrao, fmt_moeda_rs
from helpers_docx import aplicar_template
from extenso import extenso_moeda

PASTA = r'C:\Users\gabri\OneDrive\Área de Trabalho\APP - BRADESCO\0. TESTE 1\EXEMPLO DA SILVA DE SOUSA - Ruth (bradesco)\APLICAÇÃO INVEST FÁCIL'
TEMPLATE = r'C:\Users\gabri\OneDrive\Documentos\Obsidian Vault\Modelos\IniciaisBradesco\_templates\inicial-aplic-invest.docx'
DOCX_OUT = os.path.join(PASTA, 'INICIAL_AplicInvest_JOSE_ELISSON_v1.docx')
RELAT_OUT = os.path.join(PASTA, '_RELATORIO_PENDENCIAS_JOSE_ELISSON_v1.docx')

autora = {
    'nome': 'EXEMPLO DA SILVA DE SOUSA',
    'nacionalidade': 'brasileiro',
    'estado_civil': '',
    'profissao': 'motorista profissional',
    'cpf': '000.000.014-24',
    'rg': '1000012-2',
    'orgao_expedidor_prefixo': 'SSP/AM',
    'logradouro': 'Rua Itapiranga',
    'numero': '03',
    'bairro': 'José Dutra',
    'cidade': 'Presidente Figueiredo',
    'cep': '69.735-000',
}
conta = {'agencia': '3732', 'numero': '512518-9'}
renda = {'valor_float': 1500.00}  # ESTIMADO — confirmar com cliente

LANCAMENTOS = [
    ('10/03/2020', 487.41),
    ('30/08/2022', 7572.26),
]

tese = {'rubrica': 'APLIC.INVEST FACIL', 'lancamentos': LANCAMENTOS}
terceiro = {'nome': '', 'cnpj': '', 'logradouro': '', 'numero': '',
            'bairro': '', 'cidade': '', 'uf': '', 'cep': ''}

dados, calc = montar_dados_padrao(
    autora=autora, conta=conta, renda=renda, tese=tese, terceiro=terceiro,
    eh_idoso=False, competência='Presidente Figueiredo', uf='AM',
)
DANO_MORAL = 15000.00
VALOR_CAUSA_B = 15000.00
dados['valor_causa'] = fmt_moeda_rs(VALOR_CAUSA_B)
dados['valor_causa_extenso'] = extenso_moeda(VALOR_CAUSA_B)
dados['remuneração'] = 'sua atividade laboral'

print('=== EXEMPLO DA SILVA DE SOUSA — APLIC.INVEST (2 lançamentos isolados) ===')
print(f'Total bruto: R$ {calc["total"]:,.2f} | VC (b): R$ {VALOR_CAUSA_B:,.2f}'.replace(',', '#').replace('.', ',').replace('#', '.'))
res = aplicar_template(TEMPLATE, dados, DOCX_OUT)
print(f'Template: {res["modificados"]} modif, residuais: {res["residuais"] or "nenhum"}')

from docx import Document
from lxml import etree
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
doc = Document(DOCX_OUT)

MARKERS = ['Repetição do indébito', 'Caso se verifique a existência de valores indevidamente',
           'Na cobrança de débitos, o consumidor inadimplente', 'O consumidor cobrado em quantia indevida tem direito',
           'a restituição em dobro se faz necessária como penalidade', 'deve a requerida ser condenada a restituir em dobro',
           'Havendo a retenção dos valores a título de investimento']
removidos = 0
for p in list(doc.paragraphs):
    for m in MARKERS:
        if m in (p.text or ''):
            p._element.getparent().remove(p._element); removidos += 1; break
print(f'Removidos: {removidos}')

T102 = ('A cobrança indevida não decorre de engano justificável, mas de modelo operacional '
        'estruturado para funcionar sem contratação inequívoca, configurando ato ilícito '
        'durante o período em que a renda da parte autora ficou indisponível para uso imediato.')
for p in doc.paragraphs:
    if 'Além disso, havendo lançamentos, cobranças ou perdas vinculadas' in p.text:
        for re_ in list(p._element.findall(W + 'r')): p._element.remove(re_)
        r = etree.SubElement(p._element, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = T102
        break

TC = ('No caso concreto, o extrato bancário registra 2 (duas) ocorrências de aplicação '
      'automática: uma em 10/03/2020, no valor de R$ 487,41, e outra em 30/08/2022, no '
      'valor expressivo de R$ 7.572,26. Em ambos os casos, parcela substancial dos '
      'recursos disponíveis em conta-corrente foi automaticamente subtraída pelo banco '
      'réu sob a rubrica APLIC.INVEST FACIL, sem qualquer autorização expressa, '
      'pré-elaboração de orçamento, ou anuência informada da parte autora. Embora os '
      'valores tenham sido posteriormente restituídos via RESGATE INVEST FACIL, o cerne '
      'do dano moral não reside na perda patrimonial líquida — inexistente, pois o saldo '
      'agregado foi restituído —, mas na violação autônoma da autodeterminação do '
      'consumidor sobre seu próprio patrimônio em cada uma das ocasiões. Cada retenção '
      'configura, autonomamente, prática abusiva vedada pelo art. 39, inciso VI, do '
      'Código de Defesa do Consumidor, sendo o dano moral configurado pela prática '
      'reiterada (ainda que pontual) e pela apropriação não autorizada de valor expressivo '
      '(R$ 7.572,26 em uma única ocasião).')
for p in doc.paragraphs:
    if 'Alegar que os valores permaneciam disponíveis e não geraram saldo negativo' in p.text:
        np = copy.deepcopy(p._element)
        for r_ in list(np.findall(W + 'r')): np.remove(r_)
        r = etree.SubElement(np, W + 'r')
        rPr = etree.SubElement(r, W + 'rPr'); rF = etree.SubElement(rPr, W + 'rFonts')
        rF.set(W + 'ascii', 'Cambria'); rF.set(W + 'hAnsi', 'Cambria'); rF.set(W + 'cs', 'Cambria')
        t = etree.SubElement(r, W + 't'); t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = TC
        p._element.addnext(np)
        break

doc.save(DOCX_OUT)
print(f'OK -> {DOCX_OUT}')

# Relatório
doc_r = Document()
doc_r.add_heading('RELATÓRIO DE PENDÊNCIAS — INICIAL_AplicInvest_JOSE_ELISSON', level=1)
for k, v in [('Cliente', autora['nome']),
             ('Tese', 'APLIC.INVEST FACIL — estratégia (b) ADAPTADA p/ 2 lançamentos isolados'),
             ('Comarca', f'{dados["competência"]}/{dados["uf"]}'),
             ('Arquivo', 'INICIAL_AplicInvest_JOSE_ELISSON_v1.docx')]:
    p = doc_r.add_paragraph(); p.add_run(k + ': ').bold = True; p.add_run(v)

doc_r.add_heading('1. RESUMO', level=2)
tbl = doc_r.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 1'
tbl.rows[0].cells[0].text = 'Campo'; tbl.rows[0].cells[1].text = 'Valor'
for k, v in [
    ('Comarca', f'{dados["competência"]}/AM'),
    ('Nome / CPF', f'{autora["nome"]} / {autora["cpf"]}'),
    ('RG', f'{autora["rg"]} {autora["orgao_expedidor_prefixo"]}'),
    ('Nascimento / Idade', '12/11/1966 — 59 anos (NÃO idoso)'),
    ('Estado civil', '(omitido — confirmar)'),
    ('Profissão', autora['profissao']),
    ('Endereço', f'{autora["logradouro"]}, {autora["numero"]}, {autora["bairro"]}, {autora["cidade"]}/AM, CEP {autora["cep"]}'),
    ('Conta / Agência', f'{conta["numero"]} / {conta["agencia"]}'),
    ('Renda usada (ESTIMADA)', dados['valor_remuneração'] + ' — confirmar'),
    ('Período', f'{dados["inicio_desconto"]} a {dados["fim_desconto"]}'),
    ('Lançamentos APLIC', '2 (10/03/2020 R$ 487,41 + 30/08/2022 R$ 7.572,26)'),
    ('Total bruto aplicado (NÃO pleiteado)', dados['total_descontos']),
    ('Saldo líquido (extrato)', '-R$ 0,11 (cliente recebeu de volta tudo)'),
    ('Dano moral pleiteado', dados['dano_moral_total']),
    ('Valor da causa', dados['valor_causa']),
]:
    row = tbl.add_row().cells; row[0].text = k; row[1].text = v

doc_r.add_heading('2. PENDÊNCIAS CRÍTICAS', level=2)
for titulo, txt in [
    ('CASO ATÍPICO — APENAS 2 LANÇAMENTOS APLIC.INVEST',
     'Diferentemente dos demais casos APLIC do batch (que têm dezenas de retenções mensais '
     'configurando "recorrência sistêmica"), este cliente tem APENAS 2 ocorrências: '
     'R$ 487,41 (10/03/2020) e R$ 7.572,26 (30/08/2022). A estratégia (b) PADRÃO da '
     'SKILL.md assume múltiplas retenções para fundamentar dano moral. Aqui o argumento '
     'foi adaptado para "2 ocorrências autônomas + valor expressivo (R$ 7.572,26) em uma '
     'única ocasião". RISCO: juízo pode entender que 2 lançamentos não configuram '
     'situação grave o suficiente para R$ 15.000 dano moral, podendo reduzir o quantum '
     'ou julgar improcedente.'),

    ('RECOMENDAÇÃO FORTE — INICIAL COMBINADA com MORA + TARIFA',
     'A planilha 7-TABELA.xlsx do cliente tem 3 abas: MORA CRED PESSOAL (R$ 2.175,18), '
     'TARIFA BANCÁRIA CESTA (R$ 333,86) e APLIC.INVEST FÁCIL (R$ 8.059,67). '
     'RECOMENDA-SE FORTEMENTE processar como INICIAL COMBINADA (template '
     'inicial-combinada.docx) ao invés de APLIC isolada — isso fortalece a peça, '
     'agrega mais teses, e Presidente Figueiredo adota combinação. Esta inicial APLIC '
     'isolada foi gerada apenas porque o batch atual focou em APLIC.INVEST. DECISÃO '
     'FINAL do procurador sobre qual modalidade protocolar. Se optar pela combinada, '
     'descartar este DOCX e processar via batch combinada.'),

    ('AUDITORIA APLIC vs RESGATE — confirmada (sem prejuízo material)',
     'Extrato confirma: 2 APLICs e 4 RESGATEs totalizando R$ 8.059,67 vs R$ 8.059,78. '
     'Saldo líquido NEGATIVO R$ 0,11 (cliente recebeu R$ 0,11 a mais). Estratégia (b).'),

    ('RENDA ATUAL — NÃO IDENTIFICADA no extrato',
     'O extrato APLIC.INVEST cobre apenas 2020 e não traz INSS/SALARIO/PREFEITURA. '
     'Adotando R$ 1.500 como valor ESTIMADO para a Justiça Gratuita. CONFIRMAR COM '
     'CLIENTE a renda mensal atual (2026). Cliente é motorista profissional (CNH AD '
     'com observação EAR — exerce atividade remunerada).'),

    ('TEMPLATE PADRÃO ADAPTADO — pós-processamento aplicado',
     'Template inicial-aplic-invest.docx pós-processado: 7 parágrafos do bloco '
     '"Repetição do indébito" REMOVIDOS, parágrafo doutrinário REESCRITO, parágrafo do '
     'caso concreto INSERIDO com narrativa adaptada para 2 ocorrências autônomas (não '
     'recorrência sistêmica).'),

    ('NOTIFICAÇÃO EXTRAJUDICIAL — formato .docx',
     'Pasta tem "Notificação Extrajudicial - EXEMPLO DA SILVA DE SOUSA - AM - APLICAÇÃO '
     'INVEST FÁCIL.docx" mas SEM comprovante de envio. Avaliar envio antes do protocolo '
     'e juntar comprovante.'),

    ('TETO JEC — folgadamente coberto',
     'VC R$ 15.000 ≈ 9,87 SM. Cabe folgadamente no JEC (40 SM = R$ 60.720).'),

    ('COMPROVANTE DE RESIDÊNCIA — imóvel de TERCEIRO',
     'Pasta tem 5 - COMPROVANTE + 5.1 - AUTODECLARAÇÃO + 5.2 - RG PROPRIETÁRIO. '
     'Imóvel de terceiro. Confirmar vínculo (familiar? aluguel?).'),
]:
    p = doc_r.add_paragraph(style='List Bullet')
    p.add_run(titulo + ': ').bold = True; p.add_run(txt)

doc_r.add_heading('3. CHECKLIST', level=2)
for it in [
    'Conferir nome / CPF / RG / nascimento.',
    'Conferir conta/agência (512518-9 / 3732).',
    'Conferir comarca: Presidente Figueiredo/AM.',
    'CONFIRMAR RENDA com cliente.',
    'DECIDIR: APLIC isolada (este DOCX) ou COMBINADA com MORA + TARIFA?',
    'Conferir VC = R$ 15.000,00.',
    'Conferir parágrafos doutrinários "Repetição do indébito" REMOVIDOS.',
    'Conferir parágrafo caso concreto adaptado (2 ocorrências, não recorrência).',
    'Avaliar envio de notificação extrajudicial.',
    'Confirmar comprovante de residência (imóvel de terceiro).',
]:
    doc_r.add_paragraph(it, style='List Number')

p = doc_r.add_paragraph(); p.add_run('Conclusão: ').bold = True
p.add_run('APTA com RESSALVA CRÍTICA — recomenda-se INICIAL COMBINADA (MORA + TARIFA + APLIC) ao invés de APLIC isolada. '
          'Se procurador insistir em APLIC isolada, ')
r2 = p.add_run('PROTOCOLAR com risco aumentado.'); r2.bold = True
doc_r.save(RELAT_OUT)
print(f'OK -> {RELAT_OUT}')
