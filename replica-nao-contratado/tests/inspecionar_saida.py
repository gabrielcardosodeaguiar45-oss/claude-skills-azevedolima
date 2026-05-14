"""
Inspeciona o DOCX gerado pelo test_injecao_sintetica.py.
Lista paragrafos numerados e marca grifo amarelo / negrito para conferencia visual textual.
"""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

THIS = Path(__file__).resolve()
PATH_DOCX = THIS.parent / "saida_teste_injecao.docx"

doc = Document(str(PATH_DOCX))

print(f"\n{'='*70}")
print(f"INSPECAO TEXTUAL: {PATH_DOCX.name}")
print(f"{'='*70}\n")

# Processa paragrafos no nivel raiz do documento, em ordem
# (tabelas e paragrafos seguem ordem do DOCX se iterarmos pelo body)

idx_par = 0
idx_tab = 0
for elem in doc.element.body.iterchildren():
    tag = elem.tag.rsplit("}", 1)[-1]
    if tag == "p":
        # paragrafo
        p_text = ""
        marks = []
        for child in elem.iter():
            ctag = child.tag.rsplit("}", 1)[-1]
            if ctag == "t" and child.text:
                p_text += child.text
        # checa runs do paragrafo correspondente
        if idx_par < len(doc.paragraphs):
            par_obj = doc.paragraphs[idx_par]
            for run in par_obj.runs:
                if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    marks.append("AMARELO")
                if run.font.bold:
                    marks.append("BOLD")
                if run.font.italic:
                    marks.append("ITAL")
            idx_par += 1
        marks_str = ",".join(set(marks)) if marks else ""
        prefix = f"[P{idx_par:02d}]"
        if marks_str:
            prefix += f"({marks_str})"
        # Trunca para visualizacao
        snippet = p_text[:200].replace("\n", " ")
        if len(p_text) > 200:
            snippet += "..."
        print(f"{prefix:30s} {snippet}")
    elif tag == "tbl":
        idx_tab += 1
        # conta linhas
        if idx_tab - 1 < len(doc.tables):
            tab = doc.tables[idx_tab - 1]
            n_linhas = len(tab.rows)
            n_cols = len(tab.rows[0].cells) if n_linhas > 0 else 0
            primeira_celula = tab.rows[0].cells[0].text if n_linhas > 0 else ""
            print(f"[TABELA {idx_tab}: {n_linhas}x{n_cols}, header[0]='{primeira_celula[:50]}']")

print(f"\n{'='*70}")
print(f"Total: {idx_par} paragrafos, {idx_tab} tabelas")
print(f"{'='*70}\n")
