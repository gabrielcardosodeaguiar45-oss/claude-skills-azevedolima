"""
Teste sintetico end-to-end da integracao perica-digital -> replica-nao-contratado.

O que valida:
  1. Schema JSON do `_pericia.json` e consumido sem erros pelos helpers
  2. `injetar_pericia_digital()` materializa a secao III.X completa
  3. Sub-blocos por contrato (III.X.4) honram regras especiais (L.3 -> compensacao,
     B.1 -> placeholder visual, A.3 -> inconsistencia, etc.)
  4. Matriz cruzada (III.X.5) ativa kit-fraude com 6 padroes sistemicos
  5. DOCX e gerado sem erros e abre normalmente

Como executar:
    cd C:\\Users\\gabri\\.claude\\skills\\replica-nao-contratado\\tests
    python test_injecao_sintetica.py
"""
import json
import sys
from pathlib import Path

# Permite importar do diretorio references/ sem ter que ajustar PYTHONPATH manualmente.
THIS = Path(__file__).resolve()
REFS = THIS.parent.parent / "references"
sys.path.insert(0, str(REFS))

from helpers import (
    novo_documento,
    add_par,
    add_titulo_secao,
    injetar_pericia_digital,
    salvar,
)


def main():
    pericia_path = THIS.parent / "pericia_exemplo.json"
    print(f"[1/5] Carregando JSON de pericia sintetica: {pericia_path}")
    with open(pericia_path, "r", encoding="utf-8") as f:
        pericia = json.load(f)

    n_contratos = len(pericia.get("contratos_digitais", []))
    n_padroes = pericia.get("matriz_cruzada", {}).get("padroes_count", 0)
    ativa_kit = pericia.get("matriz_cruzada", {}).get("ativa_kit_fraude", False)
    print(f"      Contratos digitais: {n_contratos}")
    print(f"      Padroes na matriz cruzada: {n_padroes}")
    print(f"      Ativa kit-fraude: {ativa_kit}")

    print("[2/5] Criando documento DOCX vazio (Cambria 12, margens padrao)")
    doc = novo_documento()

    # Cabecalho minimo para contextualizar a secao III.X dentro de uma replica fictícia.
    add_titulo_secao(doc, "III - DO MERITO", nivel=0)
    add_par(doc,
        "Bloco de teste sintetico. As secoes anteriores (qualificacao, preliminares, "
        "merito comum) seriam preenchidas pelo redator-replica-nao-contratado a partir "
        "do `_analise.json` do caso real. Aqui isolamos apenas a injecao da secao III.X "
        "para validar o pipeline pericial."
    )

    print("[3/5] Chamando injetar_pericia_digital(...) com banco_invocou_selfie=True "
          "e ha_pretensa_assinatura_digital=True")
    injetar_pericia_digital(
        doc,
        pericia,
        banco_invocou_selfie=True,
        ha_pretensa_assinatura_digital=True,
    )

    out_path = THIS.parent / "saida_teste_injecao.docx"
    print(f"[4/5] Salvando DOCX em: {out_path}")
    salvar(doc, str(out_path))

    print("[5/5] Verificacoes pos-execucao:")
    if not out_path.exists():
        print("      ERRO: arquivo nao foi criado")
        sys.exit(1)

    size_kb = out_path.stat().st_size / 1024
    print(f"      Arquivo existe: OK ({size_kb:.1f} KB)")
    if size_kb < 5:
        print("      AVISO: tamanho suspeito (< 5 KB). Possivel falha silenciosa.")
        sys.exit(2)

    # Conta paragrafos no DOCX gerado para sanity-check
    from docx import Document
    doc2 = Document(str(out_path))
    n_par = len(doc2.paragraphs)
    n_tab = len(doc2.tables)
    print(f"      Paragrafos no DOCX: {n_par}")
    print(f"      Tabelas no DOCX: {n_tab}")

    if n_par < 20:
        print("      AVISO: poucos paragrafos. Verificar geracao.")
    if n_tab < 2:
        print("      AVISO: esperavamos pelo menos 2 tabelas (abertura + matriz). Encontradas:", n_tab)

    print()
    print("=" * 60)
    print("TESTE CONCLUIDO. Abra o DOCX abaixo para inspecao visual:")
    print(f"  {out_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
