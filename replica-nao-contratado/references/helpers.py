"""
Helpers Python para gerar réplica DOCX no padrão do escritório
Azevedo Lima & Rebonatto — réplicas em NÃO CONTRATADO.

Uso:
    from helpers import (
        novo_documento, add_par, add_titulo_secao,
        add_citacao, add_aviso_tese_pendente, add_placeholder_manual,
        add_placeholder_grafotecnica, add_mapa_cadeia, add_tabela_hashes,
        add_tabela_lucro, add_assinatura_advogado, salvar,
    )

    doc = novo_documento()
    add_par(doc, "Texto do parágrafo em prosa contínua.")
    add_titulo_secao(doc, "PRELIMINARES", nivel=0)
    add_citacao(doc, "Trecho jurisprudencial...")
    salvar(doc, "REPLICA_BANCO_AUTORA.docx")

NOTA: a numeração contínua de parágrafos foi REMOVIDA do padrão visual.
A peça é redigida em prosa contínua, com parágrafos separados por linha em
branco. As funções `add_par_numerado` e `ContadorParagrafos` permanecem por
compatibilidade com scripts antigos, mas `add_par_numerado` agora apenas
delega a `add_par` (não imprime mais o número).

Convenções de marcação no texto (processadas automaticamente):
    <H>texto</H>   → grifa em AMARELO (use para slots preenchidos e adaptações)
    <B>texto</B>   → NEGRITO inline
    **texto**      → convertido para <B>texto</B> antes do processamento
    *texto*        → asteriscos simples são REMOVIDOS (sem itálico inline);
                     evita que markdown de pilotos vaze como caracteres literais
    <!-- cmt -->   → comentários HTML são removidos
    <H></H>/<B></B> vazios são silenciosamente eliminados (ocorrem quando um
                     slot é substituído por string em branco). Marcações
                     ANINHADAS como <B><H>x</H></B> são suportadas via parser
                     de estado.

Slots de adaptação:
    {{CHAVE}}      → substituído via `aplicar_slots(texto, slots)`. Aceita
                     qualquer caractere (acentos, cedilha, dígitos, _) dentro
                     das chaves: {{NEGAÇÃO_TESES}}, {{VALOR_LIBERADO}}, etc.
                     Slots vazios desaparecem; slots não declarados ficam
                     como {{CHAVE}} no texto e são detectáveis via
                     `slots_residuais(texto)`.

Tabelas markdown nos pilotos:
    | col1 | col2 | → bloco detectado por `parse_tabela_markdown(linhas)` e
    |---|---|        renderizado como tabela DOCX por `add_tabela_markdown`.
    | a | b |        As células passam pelo processador de marcações, então
                     `**bold**` e `<H>amarelo</H>` funcionam dentro de células.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_COLOR_INDEX
import re


# ============================================================
# UTILITÁRIOS GERAIS
# ============================================================

class ContadorParagrafos:
    """Contador para numeração contínua de parágrafos (1., 2., 3., ...).
    Use uma única instância por documento."""

    def __init__(self):
        self.n = 0

    def proximo(self):
        self.n += 1
        return self.n

    def reset(self):
        self.n = 0


# ============================================================
# CRIAÇÃO E CONFIGURAÇÃO DO DOCUMENTO
# ============================================================

def novo_documento():
    """Cria documento DOCX novo com configuração padrão do escritório."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
    style = doc.styles['Normal']
    style.font.name = 'Cambria'
    style.font.size = Pt(12)
    return doc


# ============================================================
# PROCESSAMENTO DE MARCAÇÕES INLINE
# ============================================================

def md_para_marcacao(texto):
    """Converte sintaxe Markdown para tags inline antes do processamento.

    `**bold**`     -> `<B>...</B>`
    `*italic*`     -> remove os asteriscos (mantem o texto sem itálico, pois
                       muitos pilotos usam asteriscos apenas para destaque
                       editorial e o texto fica mais limpo sem itálico).
    `<!-- ... -->` -> removido (comentários HTML usados nos pilotos).

    Também elimina marcações vazias `<H></H>` e `<B></B>` (que ocorrem
    quando um slot foi substituído por string vazia).
    """
    if not texto:
        return texto
    # 1) Remove comentários HTML
    texto = re.sub(r'<!--.*?-->', '', texto, flags=re.DOTALL)
    # 2) **bold** -> <B>bold</B>  (lazy para suportar múltiplos no parágrafo)
    texto = re.sub(r'\*\*(.+?)\*\*', r'<B>\1</B>', texto)
    # 3) *italic* -> remove os asteriscos preservando o texto.
    #    Não pode capturar `**` (já tratado) nem cruzar quebra de linha.
    texto = re.sub(r'(?<!\*)\*(?!\*)([^*\n]+?)(?<!\*)\*(?!\*)', r'\1', texto)
    # 4) Remove marcações vazias provenientes de slots em branco
    texto = re.sub(r'<H>\s*</H>', '', texto)
    texto = re.sub(r'<B>\s*</B>', '', texto)
    # 5) Limpeza final: remove asteriscos órfãos remanescentes de markdown
    #    quebrado (ex.: `**texto**{{slot_vazio}}**` deixa um `**` pendurado).
    #    Como pilotos jurídicos não usam `*` como caractere literal, é seguro
    #    eliminar qualquer asterisco residual.
    texto = re.sub(r'\*+', '', texto)
    return texto


def _emit_run(p, texto, bold, italic, highlight, font_size):
    """Cria um run no parágrafo com as flags de estado correntes."""
    if not texto:
        return
    run = p.add_run(texto)
    run.font.name = 'Cambria'
    run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _processa_marcacoes(p, texto, bold=False, italic=False, font_size=12):
    """Processa <H>...</H> (grifo amarelo) e <B>...</B> (negrito) inline.

    Suporta ANINHAMENTO arbitrário: `<B><H>x</H></B>`, `<H><B>x</B></H>`,
    sequências combinadas etc. — o parser opera por máquina de estado e abre
    um novo run a cada transição de marcação, ao invés de tratar cada par de
    tags como um bloco indivisível.

    Também tolera tags soltas (uma tag de fechamento sem abertura simplesmente
    desativa o estado correspondente).
    """
    if not texto:
        return
    # Pré-limpeza: aplica conversão Markdown e remove marcações vazias
    texto = md_para_marcacao(texto)

    pattern = re.compile(r'<(/?)(B|H)>')

    state_bold = bold
    state_highlight = False
    state_italic = italic

    pos = 0
    for m in pattern.finditer(texto):
        # Emite o texto entre a posição atual e o início do match
        chunk = texto[pos:m.start()]
        if chunk:
            _emit_run(p, chunk, state_bold, state_italic, state_highlight, font_size)
        # Atualiza estado conforme a tag encontrada
        is_close = (m.group(1) == '/')
        tag = m.group(2)
        if tag == 'B':
            state_bold = not is_close
        elif tag == 'H':
            state_highlight = not is_close
        pos = m.end()
    # Resto após a última tag
    chunk = texto[pos:]
    if chunk:
        _emit_run(p, chunk, state_bold, state_italic, state_highlight, font_size)


# ============================================================
# PARÁGRAFOS — VERSÃO SIMPLES (sem numeração)
# ============================================================

def add_par(doc, texto, bold=False, align='justify', indent_left=None):
    """Adiciona parágrafo SEM numeração contínua. Use <H>texto</H> para grifo
    amarelo, <B>texto</B> para negrito inline.

    align: 'justify' (padrão), 'center', 'right', 'left'
    indent_left: cm de recuo à esquerda (opcional)
    """
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    _processa_marcacoes(p, texto, bold=bold)
    return p


# ============================================================
# PARÁGRAFOS — VERSÃO NUMERADA (1., 2., 3., ...)
# ============================================================

def add_par_numerado(doc, contador, texto, bold=False, align='justify',
                     indent_left=None):
    """DEPRECIADO. Mantido para compatibilidade com scripts antigos.

    A numeração contínua de parágrafos foi REMOVIDA do padrão visual do
    escritório. Esta função agora apenas delega a `add_par`, ignorando o
    contador. O contador é incrementado mesmo assim para preservar a
    semântica de scripts que reportam o total ao final.
    """
    if contador is not None and hasattr(contador, 'proximo'):
        contador.proximo()
    return add_par(doc, texto, bold=bold, align=align, indent_left=indent_left)


# ============================================================
# TÍTULOS DE SEÇÃO
# ============================================================

def add_titulo_secao(doc, texto, nivel=1):
    """Adiciona título de seção em negrito.

    nivel=0: centralizado (para títulos maiores como PRELIMINARES, MÉRITO)
    nivel=1: alinhado à esquerda (para subtítulos)
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if nivel == 0 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    run.font.bold = True
    return p


# ============================================================
# CITAÇÕES JURISPRUDENCIAIS / NORMATIVAS
# ============================================================

def add_citacao(doc, texto):
    """Adiciona citação jurisprudencial em recuo de 4cm, Cambria 11pt itálico.
    Use <B>texto</B> para destacar partes em negrito (ex.: ementa).
    Use <H>texto</H> para grifar em amarelo (slots preenchidos)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r'(<B>.*?</B>|<H>.*?</H>)', texto)
    for part in parts:
        if not part:
            continue
        if part.startswith('<B>') and part.endswith('</B>'):
            inner = part[3:-4]
            run = p.add_run(inner)
            run.font.name = 'Cambria'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.italic = True
        elif part.startswith('<H>') and part.endswith('</H>'):
            inner = part[3:-4]
            run = p.add_run(inner)
            run.font.name = 'Cambria'
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run = p.add_run(part)
            run.font.name = 'Cambria'
            run.font.size = Pt(11)
            run.font.italic = True
    return p


# ============================================================
# PLACEHOLDERS — para teses sem piloto, blocos manuais e grafotécnica
# ============================================================

def add_aviso_tese_pendente(doc, titulo_tese, motivo):
    """Para teses arguidas pelo banco que NÃO têm piloto pronto no vault.
    Cria título + parágrafo amarelo de aviso + 3 linhas em branco para
    preenchimento manual.
    """
    add_titulo_secao(doc, titulo_tese)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"[TESE A SER DESENVOLVIDA — {motivo}]")
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for _ in range(3):
        doc.add_paragraph()


def add_placeholder_manual(doc, descricao):
    """Marca posição de bloco padrão que o usuário insere manualmente.

    Exemplos de uso:
        add_placeholder_manual(doc, "bloco padrão de fraude sistêmica do INSS")
        add_placeholder_manual(doc, "Lei 14.063/2020 — 3 tipos de assinatura")
        add_placeholder_manual(doc, "Selfie liveness — 11 requisitos NT INSS")
        add_placeholder_manual(doc, "Validador ITI — assinatura inválida")
        add_placeholder_manual(doc, "Modus operandi kit fraude")
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f">>> [INSERIR MANUALMENTE — {descricao}] <<<")
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for _ in range(2):
        doc.add_paragraph()


def add_placeholder_grafotecnica(doc, contrato_numero, contrato_localizacao,
                                 procuracao_localizacao, rg_localizacao):
    """Insere placeholder para análise grafotécnica manual de um contrato físico.

    A análise grafotécnica conclusiva é decisão humana — a skill apenas marca
    a posição e enumera os documentos a comparar.
    """
    titulo_msg = f"Da análise grafotécnica do contrato nº {contrato_numero}"
    add_titulo_secao(doc, titulo_msg)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    msg = (
        f"[ANÁLISE GRAFOTÉCNICA — PREENCHER MANUALMENTE NO CHAT\n"
        f"Comparar visualmente:\n"
        f"  - Assinatura na CCB do contrato {contrato_numero} ({contrato_localizacao})\n"
        f"  - Assinatura padrão na Procuração ({procuracao_localizacao})\n"
        f"  - Assinatura padrão no RG/CNH ({rg_localizacao})\n"
        f"Apontar: traçado, fluidez, pressão, fragmentação, similitudes/divergências.\n"
        f"Indicar se há indício de decalque (preto-e-branco, ausência de variação de pressão).]"
    )
    run = p.add_run(msg)
    run.font.name = 'Cambria'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    for _ in range(3):
        doc.add_paragraph()


# ============================================================
# TABELAS GENÉRICAS
# ============================================================

def add_tabela(doc, headers, linhas, estilo='Light Grid Accent 1', processa_marcacao=True):
    """Adiciona tabela com cabeçalho + linhas.

    headers: lista de strings (cabeçalho)
    linhas: lista de listas de strings
    processa_marcacao: se True (padrão), as células passam por `_processa_marcacoes`
                        e `md_para_marcacao` — assim `<H>`, `<B>`, `**bold**`,
                        `*italic*`, `<!-- cmt -->` e slots vazios são tratados
                        igual aos parágrafos. Se False, escreve texto cru.
    """
    table = doc.add_table(rows=len(linhas) + 1, cols=len(headers))
    table.style = estilo
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        if processa_marcacao:
            hdr[i].text = ''
            _processa_marcacoes(hdr[i].paragraphs[0], str(h), bold=True)
        else:
            hdr[i].text = str(h)
    for i, linha in enumerate(linhas, 1):
        cells = table.rows[i].cells
        for j, val in enumerate(linha):
            if processa_marcacao:
                cells[j].text = ''
                _processa_marcacoes(cells[j].paragraphs[0], str(val))
            else:
                cells[j].text = str(val)
    return table


# ============================================================
# TABELA A PARTIR DE MARKDOWN (pilotos do vault contêm tabelas | col | col |)
# ============================================================

def _eh_separador_markdown(linha):
    """Detecta linha separadora de tabela markdown: `|---|---|` ou `|:--:|---:|`."""
    return bool(re.match(r'^\s*\|[\s:|\-]+\|\s*$', linha))


def parse_tabela_markdown(linhas):
    """Detecta tabela markdown em uma lista de linhas.

    Retorna (headers, rows) com strings limpas (incluindo marcações inline
    `<H>`, `<B>`, `**bold**`), ou None se as linhas não formam tabela.
    """
    # Normaliza
    linhas = [l for l in (l.rstrip() for l in linhas) if l.strip()]
    if len(linhas) < 2:
        return None
    primeiro = linhas[0].strip()
    segundo = linhas[1].strip()
    if not (primeiro.startswith('|') and primeiro.endswith('|')):
        return None
    if not _eh_separador_markdown(segundo):
        return None
    # Headers
    headers = [c.strip() for c in primeiro.strip('|').split('|')]
    # Body
    rows = []
    for ln in linhas[2:]:
        ln = ln.strip()
        if ln.startswith('|') and ln.endswith('|') and not _eh_separador_markdown(ln):
            cells = [c.strip() for c in ln.strip('|').split('|')]
            rows.append(cells)
        else:
            break
    if not rows:
        return None
    return headers, rows


def add_tabela_markdown(doc, texto_md, estilo='Light Grid Accent 1'):
    """Converte uma tabela em formato markdown (texto OU lista de linhas) em
    tabela DOCX. As células passam pelo processador de marcações inline.

    Retorna True se converteu, False se o input não é tabela markdown válida.
    """
    if isinstance(texto_md, str):
        linhas = texto_md.splitlines()
    else:
        linhas = list(texto_md)
    parsed = parse_tabela_markdown(linhas)
    if not parsed:
        return False
    headers, rows = parsed
    add_tabela(doc, headers=headers, linhas=rows, estilo=estilo, processa_marcacao=True)
    return True


# ============================================================
# SUBSTITUIÇÃO DE SLOTS {{CHAVE}} (aceita acentos e cedilha)
# ============================================================

# Regex permissivo: aceita qualquer caractere exceto chaves dentro de {{ }}.
# Isso cobre slots como {{NEGAÇÃO_TESES}}, {{COMARCA_AÇÃO}}, etc.
RE_SLOT = re.compile(r'\{\{([^{}\s][^{}]*?)\}\}')


def aplicar_slots(texto, slots, marcar_amarelo=True):
    """Substitui `{{CHAVE}}` pelo valor de `slots[CHAVE]`.

    - Slots cujo valor é string vazia desaparecem (não geram `<H></H>`).
    - Quando `marcar_amarelo=True` (padrão), envolve a substituição em
      `<H>...</H>` para grifo amarelo no DOCX.
    - Slots não declarados em `slots` permanecem como `{{CHAVE}}` no texto;
      use `slots_residuais()` para detectá-los após a substituição.
    """
    def rep(m):
        chave = m.group(1).strip()
        if chave in slots:
            valor = slots[chave]
            if valor == "" or valor is None:
                return ""
            valor_str = str(valor)
            return f"<H>{valor_str}</H>" if marcar_amarelo else valor_str
        return m.group(0)
    return RE_SLOT.sub(rep, texto)


def slots_residuais(texto):
    """Retorna o set de chaves `{{X}}` que ainda permanecem no texto após
    `aplicar_slots`. Útil para emitir warnings antes de salvar o DOCX.
    """
    return set(RE_SLOT.findall(texto))


# ============================================================
# COMPONENTES VISUAIS ESPECÍFICOS DA RÉPLICA
# ============================================================

def add_mapa_cadeia(doc, etapas):
    """Insere mapa visual da cadeia contratual em formato centralizado.

    etapas: lista de strings, cada uma representando uma etapa da cadeia.
    Setas e indicações de paralelismo são adicionadas pelo chamador
    como elementos da própria lista.

    Exemplo:
        add_mapa_cadeia(doc, [
            "11915198 (físico)",
            "10/10/2019 — Campos Novos/SC",
            "R$ 3.551,64 / 72 parcelas / R$ 96,45",
            "↓ refinanciado em 19/07/2021 ↓",
            "20031787 (digital, ATIVO)",
            "16/07/2021 — Solida Serviço Administrativ",
            "R$ 4.345,16 / 84 parcelas / R$ 96,45 — troco R$ 1.229,79",
            "+ contratado em paralelo no mesmo dia",
            "20032423 (digital, ATIVO)",
            "16/07/2021 — Solida Serviço Administrativ",
            "R$ 750,70 / 84 parcelas / R$ 19,20",
        ])
    """
    add_titulo_secao(doc, "Mapa visual da cadeia contratual")
    for etapa in etapas:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(etapa)
        run.font.name = 'Cambria'
        run.font.size = Pt(11)
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.add_paragraph()


def add_tabela_hashes(doc, contratos):
    """Tabela inventário dos 6 hashes SHA-256 entre contratos digitais.

    contratos: lista de dicts, cada um com:
        {
            'numero': '20031787',
            'codigo_verificacao': 'AC55-61D7-95BD-3579',
            'hash_envelope': 'cecdc5a7...',
            'hash_ccb': '72a4f485...',
            'hash_cet': '15a69c2d...',
            'hash_termo_inss': '27d53572...',
            'hash_cadastro': '131910d0...',
            'hash_evidencias': '0D013F65...',
        }

    Renderiza em formato Componente x Contrato 1 x Contrato 2 x ...
    """
    if not contratos:
        return
    componentes = [
        ("Código de verificação", "codigo_verificacao"),
        ("Hash do Documento (envelope)", "hash_envelope"),
        ("Hash da CCB", "hash_ccb"),
        ("Hash da CET", "hash_cet"),
        ("Hash do Termo de Autorização Consulta INSS", "hash_termo_inss"),
        ("Hash da Abertura/Renovação de Cadastro", "hash_cadastro"),
        ("Hash das Evidências", "hash_evidencias"),
    ]
    headers = ["Componente"] + [f"Contrato {c['numero']}" for c in contratos]
    linhas = []
    for nome, chave in componentes:
        linha = [nome]
        for c in contratos:
            linha.append(c.get(chave, ""))
        linhas.append(linha)
    return add_tabela(doc, headers, linhas)


def add_tabela_lucro(doc, contratos_com_ted):
    """Tabela do lucro estimado do banco — APENAS dos contratos com TED comprovado.

    contratos_com_ted: lista de dicts com:
        {
            'numero': '000020 031787',
            'depositado': 1229.79,
            'total_pago': 8103.80,
        }
    Calcula automaticamente o lucro e o percentual.
    """
    if not contratos_com_ted:
        return
    headers = ["CONTRATO", "Valor depositado sem autorização",
               "Valor total a ser pago", "Lucro estimado do banco",
               "Percentual de lucro"]
    linhas = []
    total_dep = 0.0
    total_pago = 0.0
    total_lucro = 0.0

    def fmt_brl(v):
        return f"R$ {v:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

    for c in contratos_com_ted:
        dep = float(c['depositado'])
        pago = float(c['total_pago'])
        lucro = pago - dep
        perc = (lucro / dep * 100) if dep > 0 else 0
        linhas.append([
            c['numero'],
            fmt_brl(dep),
            fmt_brl(pago),
            fmt_brl(lucro),
            f"~ {perc:.2f}%",
        ])
        total_dep += dep
        total_pago += pago
        total_lucro += lucro
    if len(contratos_com_ted) > 1:
        linhas.append([
            "TOTAL",
            fmt_brl(total_dep),
            fmt_brl(total_pago),
            fmt_brl(total_lucro),
            "—",
        ])
    return add_tabela(doc, headers, linhas)


# ============================================================
# CABEÇALHO E ASSINATURA
# ============================================================

def add_cabecalho_simples(doc, vara_completa, processo):
    """Cabeçalho enxuto: 2 linhas centralizadas.

    Exemplo:
        add_cabecalho_simples(doc,
            "Vara do Único Ofício da Comarca de Feira Grande/AL",
            "0700273-13.2026.8.02.0060")
    """
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(0)
    run1 = p1.add_run(f"Ao Juízo da {vara_completa}")
    run1.font.name = 'Cambria'
    run1.font.size = Pt(12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    run2 = p2.add_run(f"Processo nº {processo}")
    run2.font.name = 'Cambria'
    run2.font.size = Pt(12)


def add_apresentacao_enxuta(doc, nome_autor, banco_referencia,
                            ja_qualificado_genero='qualificado'):
    """Apresentação em 1 frase. Não usar parágrafo formal com data de protocolo
    e fls. — só nome do autor e referência da contestação.

    ja_qualificado_genero: 'qualificado' (masc) ou 'qualificada' (fem)

    Exemplo:
        add_apresentacao_enxuta(doc,
            "EDINETE GENEZIO DA SILVA SANTOS",
            "nas fls. 214-224",
            ja_qualificado_genero='qualificada')
    """
    add_par(doc,
        f"<H>{nome_autor}</H>, já {ja_qualificado_genero} nos autos do processo, "
        f"vem à presença de Vossa Excelência, por intermédio de seus procuradores "
        f"constituídos, apresentar réplica da contestação apresentada pelo "
        f"banco réu <H>{banco_referencia}</H>, conforme razões de fato e de "
        f"direito que seguem."
    )


def add_assinatura_advogado(doc, cidade_filial, data_extenso,
                            nome_advogado, oab):
    """Assinatura final no padrão da equipe.

    cidade_filial: cidade da FILIAL do escritório que atende a parte autora,
                   não a cidade do juízo.
    data_extenso: ex.: "4 de maio de 2026"
    nome_advogado: nome do advogado que assinou a INICIAL
    oab: ex.: "OAB/AL 20.906A"
    """
    p_data = doc.add_paragraph()
    p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_data.paragraph_format.space_before = Pt(24)
    p_data.paragraph_format.space_after = Pt(36)
    run_d = p_data.add_run(f"{cidade_filial}, {data_extenso}")
    run_d.font.name = 'Cambria'
    run_d.font.size = Pt(12)
    run_d.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Linha de assinatura
    p_lin = doc.add_paragraph()
    p_lin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lin.paragraph_format.space_after = Pt(0)
    run_lin = p_lin.add_run("________________________________________")
    run_lin.font.name = 'Cambria'
    run_lin.font.size = Pt(12)

    # Nome do advogado
    p_nome = doc.add_paragraph()
    p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nome.paragraph_format.space_after = Pt(0)
    run_n = p_nome.add_run(nome_advogado)
    run_n.font.name = 'Cambria'
    run_n.font.size = Pt(12)
    run_n.font.bold = True
    run_n.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # OAB
    p_oab = doc.add_paragraph()
    p_oab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_oab.paragraph_format.space_after = Pt(0)
    run_o = p_oab.add_run(oab)
    run_o.font.name = 'Cambria'
    run_o.font.size = Pt(12)
    run_o.font.highlight_color = WD_COLOR_INDEX.YELLOW


# ============================================================
# SÍNTESE PROCESSUAL EM BULLETS
# ============================================================

def add_sintese_processual(doc, teses_do_banco):
    """Adiciona seção 'Síntese processual' com bullets curtos das teses do banco.

    teses_do_banco: lista de strings curtas, sem explicações longas entre
                    parênteses. Padrão visual da equipe.

    Exemplo:
        add_sintese_processual(doc, [
            "Decadência",
            "Prescrição quinquenal — supressio",
            "Falta de interesse de agir",
            "Ausência de juntada de extrato bancário",
            "Procuração genérica",
            "Impugnação ao valor da causa",
            "Validade do negócio jurídico",
            "Compensação/devolução do valor recebido em conta",
        ])
    """
    add_titulo_secao(doc, "SÍNTESE PROCESSUAL", nivel=0)
    add_par(doc, "Resumo da contestação")
    for tese in teses_do_banco:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(0)
        # Sem grifo — esses bullets representam o que o banco arguiu
        run = p.add_run(f"{tese};")
        run.font.name = 'Cambria'
        run.font.size = Pt(12)


# ============================================================
# INJEÇÃO DE PERÍCIA DIGITAL — seção III.X (contratos digitais)
# Consome o `_pericia.json` gerado por `pipeline_pericia_digital.py`.
# Estrutura segue `tabela-mestre-achado-piloto.md`:
#   III.X    — Abertura + tabela de identificação
#   III.X.1  — Moldura técnica geral [placeholders manuais]
#   III.X.2  — Régua biométrica [se banco invocou selfie]
#   III.X.3  — Régua de assinatura digital [se há pretensa assinatura digital]
#   III.X.4  — Inconsistências individuais [sub-bloco por contrato]
#   III.X.5  — Matriz cruzada [se ≥2 contratos digitais]
#   III.X.6  — Insuficiência probatória [piloto: insuficiencia-probatoria-prova-unilateral]
#   III.X.7  — Robôs + pedido de perícia [piloto: robos-falsificar-assinatura]
# ============================================================

# Mapa de variantes para o título da seção do achado dentro do sub-bloco
_TITULOS_ACHADO = {
    "A.1": "Do campo de e-mail vazio na CCB",
    "A.2": "Do e-mail cadastrado com placeholder",
    "A.3": "Da identificação eletrônica por e-mail do próprio banco",
    "B.1": "Da assinatura inválida no validador ITI",
    "B.2": "Da impossibilidade de validação no ITI",
    "C.1": "Do hash divergente — efeito avalanche",
    "C.2": "Da ausência de código hash",
    "C.3": "Do hash idêntico entre contratos",
    "D.1": "Dos metadados — criação posterior à data alegada",
    "D.2": "Dos metadados — software automatizado",
    "D.3": "Dos metadados — modificação pós-assinatura",
    "E.1": "Do IP em rede privada (RFC 1918)",
    "E.2": "Do IP geolocalizado fora da residência",
    "E.3": "Do IP coincidente com a sede do correspondente",
    "G.1": "Da sessão compartilhada entre contratos",
    "G.2": "Do horário de aceite humanamente impossível",
    "G.3": "Da trilha de auditoria com status incompleto",
    "H.1": "Da ausência de selfie",
    "H.2": "Da selfie reutilizada entre contratos",
    "H.3": "Da selfie sem liveness adequado",
    "H.4": "Do RG com impossibilidade de assinar versus selfie apresentada",
    "I.1": "Do correspondente bancário em cidade distante",
    "I.2": "Do correspondente bancário em Manaus versus residência em Maués",
    "I.3": "Dos múltiplos contratos no mesmo correspondente no mesmo dia",
    "I.4": "Do correspondente comum entre bancos distintos",
    "J.1": "Do telefone com DDD divergente da residência",
    "J.2": "Dos telefones distintos entre contratos da mesma autora",
    "J.3": "Do SMS de formalização com DDD divergente",
    "K.1": "Do contrato impugnado não juntado pelo banco",
    "K.2": "Do contrato citado mas não anexado",
    "L.1": "Do valor do TED divergente da CCB",
    "L.2": "Do horário do TED incompatível com aceite",
    "L.3": "Do depósito em conta diversa daquela em que recebia o INSS",
    "L.4": "Da ausência de comprovante de TED",
    "L.5": "Do comprovante sem NSU/EndToEndID",
}


def add_secao_digital_abertura(doc, contratos_digitais):
    """III.X — Abertura: impugnação geral + tabela de identificação dos contratos digitais.

    contratos_digitais: lista de dicts conforme schema `_pericia.json`:
        [{"numero", "ade", "data_alegada", "valor_liberado", "evento_pdf", "status"}, ...]
    """
    add_titulo_secao(doc, "Dos contratos digitais", nivel=0)

    add_par(doc,
        "A parte autora desconhece totalmente os supostos contratos digitais apresentados pelo "
        "banco reu, jamais tendo manifestado vontade de contratar emprestimo consignado, sequer "
        "acessado plataforma digital do banco para tal fim."
    )

    linhas = []
    for c in contratos_digitais:
        linhas.append([
            f"{c.get('numero','?')}{(' (ADE ' + c['ade'] + ')') if c.get('ade') else ''}",
            c.get("data_alegada", "?"),
            f"R$ {c.get('valor_liberado', 0.0):,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
            c.get("evento_pdf", "?"),
            c.get("status", "?"),
        ])
    add_tabela(doc,
        headers=["Contrato", "Data alegada", "Valor liberado", "Localização nos autos", "Status"],
        linhas=linhas
    )


def add_secao_moldura_geral(doc):
    """III.X.1 — Moldura técnica geral [PILOTO FIXO MANUAL]."""
    add_titulo_secao(doc, "Da moldura técnica geral da contratação digital")
    add_placeholder_manual(doc,
        "PILOTO FIXO — Lei 14.063/2020 (3 tipos de assinatura eletrônica) + "
        "MP 2.200-2/2001 (assinatura digital ICP-Brasil e equivalentes) + "
        "Resolução CMN 5.057/2022 (correspondentes bancários e contratação eletrônica). "
        "Bloco doutrinário padrão do escritório — copiar do acervo."
    )


def add_secao_regua_biometrica(doc, aplicavel: bool):
    """III.X.2 — Régua biométrica [PILOTO FIXO MANUAL, apenas se banco invocou selfie/biometria]."""
    if not aplicavel:
        return
    add_titulo_secao(doc, "Da régua biométrica — 11 requisitos de selfie e liveness")
    add_placeholder_manual(doc,
        "PILOTO FIXO — 11 requisitos da Nota Técnica do INSS (IN 138/2022 e IN 162/2024) "
        "+ ISO/IEC 30107-3 (liveness) + IEEE Std 2790-2020. "
        "Bloco doutrinário padrão — copiar do acervo."
    )


def add_secao_regua_assinatura(doc, aplicavel: bool):
    """III.X.3 — Régua de assinatura digital [PILOTO FIXO MANUAL, apenas se há pretensa assinatura digital]."""
    if not aplicavel:
        return
    add_titulo_secao(doc, "Da régua da assinatura digital — Validador ITI")
    add_placeholder_manual(doc,
        "PILOTO FIXO — Validador oficial do ITI (validar.iti.gov.br), MP 2.200-2/2001 art. 10 §1º, "
        "Lei 14.063/2020. Bloco doutrinário padrão — copiar do acervo. "
        "ANEXAR: print do validador para cada contrato impugnado, demonstrando assinatura INVÁLIDA."
    )


def _formatar_valor(v):
    """Helper para formatar valores na tabela do sub-bloco."""
    if v is None:
        return "—"
    if isinstance(v, list):
        return ", ".join(str(x) for x in v) if v else "—"
    if isinstance(v, dict):
        return str(v)
    return str(v)


def add_subbloco_contrato_digital(doc, contrato_pericia: dict):
    """III.X.4 — Sub-bloco de inconsistências individuais por contrato (opção A, cirúrgica).

    Recebe um item da lista `contratos_digitais` do `_pericia.json`. Insere:
    1. Título do sub-bloco com número do contrato
    2. Tabela compacta de achados aplicáveis (só os com piloto_acionado != null e risco != BAIXO)
    3. Um parágrafo por achado, com o texto do achado em amarelo + referência ao piloto
    """
    num = contrato_pericia.get("numero", "?")
    ade = contrato_pericia.get("ade")
    titulo = f"Do contrato nº {num}"
    if ade:
        titulo += f" (ADE {ade})"
    add_titulo_secao(doc, titulo)

    achados = contrato_pericia.get("achados", {}) or {}
    aplicaveis = []
    for codigo, ach in achados.items():
        if not isinstance(ach, dict):
            continue
        piloto = ach.get("piloto_acionado")
        risco = ach.get("risco")
        if piloto and risco in ("ALTO", "MEDIO", "manual"):
            aplicaveis.append((codigo, ach))

    if not aplicaveis:
        add_par(doc,
            "<H>Nenhum achado pericial alto/médio detectado neste contrato a partir da "
            "documentação juntada pelo banco. A impugnação é sustentada pelos blocos da "
            "moldura técnica geral e pela inversão do ônus probatório (Tema 1.061 STJ).</H>"
        )
        return

    # Tabela compacta de achados
    linhas = []
    for codigo, ach in aplicaveis:
        variante = ach.get("variante") or codigo
        titulo_v = _TITULOS_ACHADO.get(variante, codigo)
        linhas.append([
            variante,
            titulo_v,
            ach.get("risco", "?"),
        ])
    add_tabela(doc,
        headers=["Cód.", "Achado", "Risco"],
        linhas=linhas
    )

    # Um parágrafo por achado com texto + piloto correspondente
    for codigo, ach in aplicaveis:
        variante = ach.get("variante") or codigo
        texto = ach.get("texto_achado") or ""
        piloto = ach.get("piloto_acionado") or ""
        # Sub-título do achado
        add_par(doc,
            f"<B>{variante} — {_TITULOS_ACHADO.get(variante, codigo)}.</B> <H>{texto}</H>"
        )

        # Para variantes especiais, observação adicional
        if variante == "L.3":
            add_par(doc,
                "<H>Aciona-se a tese da compensação: o depósito não é prova de contratação, "
                "mas elemento da própria fraude. Em eventual procedência, autoriza-se a "
                "compensação dos valores efetivamente creditados para evitar enriquecimento "
                "ilícito.</H>"
            )
        if variante == "B.1" and ach.get("placeholder_visual"):
            add_par(doc, ach["placeholder_visual"])

        # Wikilink para o piloto (apenas comentário)
        if piloto and not piloto.startswith("FLEX:"):
            run_p = doc.add_paragraph()
            run_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = run_p.add_run(f"[piloto: {piloto}]")
            r.font.name = 'Cambria'
            r.font.size = Pt(10)
            r.font.italic = True


def add_secao_matriz_cruzada(doc, matriz: dict):
    """III.X.5 — Matriz cruzada (apenas se ≥2 contratos digitais)."""
    if not matriz or matriz.get("padroes_count", 0) == 0:
        return

    add_titulo_secao(doc, "Da matriz cruzada entre os contratos digitais")

    add_par(doc,
        "Confrontando os contratos digitais entre si — IP, sessão, correspondente, horário "
        "do aceite, hashes SHA-256 e selfie — emergem padrões sistêmicos que reforçam a "
        "tese de fraude estruturada e descaracterizam a pretensa individualização técnica "
        "das contratações alegada pelo banco."
    )

    # Tabela visual da matriz
    linhas = []
    for col in matriz.get("tabela_comparativa", []):
        if not col.get("padrao_detectado"):
            continue
        valores_str = ""
        for ln in col.get("linhas", []):
            val = ln.get("valor", "")
            valores_str += f"\n{ln.get('contrato','?')}: {val}"
        linhas.append([
            col.get("campo", "?").upper(),
            valores_str.strip(),
            "✓ padrão detectado",
        ])

    if linhas:
        add_tabela(doc,
            headers=["Campo", "Valor por contrato", "Padrão"],
            linhas=linhas
        )

    # Padrões em texto
    add_par(doc, f"<H>{matriz.get('observacao_padrao','')}</H>")

    # Aciona kit-fraude se padroes_count >= 3
    if matriz.get("ativa_kit_fraude"):
        add_par(doc,
            "<B>O conjunto de padrões sistêmicos acima caracteriza esquema estruturado de "
            "fraude.</B> A reutilização do mesmo IP, da mesma sessão, da mesma selfie e/ou "
            "do mesmo correspondente, em contratos celebrados em sequência ou no mesmo dia, "
            "é incompatível com fluxo humano legítimo de contratação."
        )
        run_p = doc.add_paragraph()
        r = run_p.add_run("[piloto: merito-probatorio-digital/kit-fraude]")
        r.font.name = 'Cambria'; r.font.size = Pt(10); r.font.italic = True
    elif matriz.get("ativa_cadeia_custodia"):
        add_par(doc,
            "<B>A presença de hash idêntico ou de outros campos compartilhados entre os "
            "contratos rompe a cadeia de custódia digital</B> que deveria assegurar a "
            "individualização técnica de cada operação."
        )
        run_p = doc.add_paragraph()
        r = run_p.add_run("[piloto: merito-probatorio-digital/cadeia-custodia-digital-inexistente]")
        r.font.name = 'Cambria'; r.font.size = Pt(10); r.font.italic = True


def injetar_pericia_digital(doc, pericia: dict, banco_invocou_selfie: bool = True,
                              ha_pretensa_assinatura_digital: bool = True):
    """Função orquestradora que materializa a seção III.X completa no DOCX a partir do `_pericia.json`.

    Sequência fixa:
      III.X    Abertura + tabela
      III.X.1  Moldura geral
      III.X.2  Régua biométrica (se aplicável)
      III.X.3  Régua de assinatura (se aplicável)
      III.X.4  Sub-blocos por contrato
      III.X.5  Matriz cruzada (se ≥2 contratos digitais)

    III.X.6 (insuficiência probatória) e III.X.7 (robôs) ficam por conta do script-redator,
    que os adiciona após esta chamada usando os pilotos correspondentes do vault.
    """
    contratos = pericia.get("contratos_digitais", [])
    if not contratos:
        return  # nada a fazer — caso só físico

    # III.X — Abertura
    add_secao_digital_abertura(doc, contratos)

    # III.X.1 — Moldura geral (sempre)
    add_secao_moldura_geral(doc)

    # III.X.2 — Régua biométrica (condicional)
    add_secao_regua_biometrica(doc, banco_invocou_selfie)

    # III.X.3 — Régua de assinatura (condicional)
    add_secao_regua_assinatura(doc, ha_pretensa_assinatura_digital)

    # III.X.4 — Sub-blocos por contrato
    add_titulo_secao(doc, "Das inconsistências individuais por contrato")
    for c in contratos:
        add_subbloco_contrato_digital(doc, c)

    # III.X.5 — Matriz cruzada (apenas se ≥2 contratos)
    if len(contratos) >= 2:
        add_secao_matriz_cruzada(doc, pericia.get("matriz_cruzada", {}))


# ============================================================
# SALVAR
# ============================================================

def salvar(doc, caminho):
    """Salva o documento no caminho especificado."""
    doc.save(caminho)
    print(f"OK - Replica salva em: {caminho}")
